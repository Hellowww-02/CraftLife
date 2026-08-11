"""
CraftLife Desktop Version  —  MainPyQt6.py  v1.0
PyQt6 Windows / Linux / macOS
Install : pip install PyQt6
Run     : python MainPyQt6.py
"""

# ══════════════════════════════════════════════════════════════════════════════
# STRUKTUR FILE  (Table of Contents) — MainPyQt6.py
# ══════════════════════════════════════════════════════════════════════════════
# 1.  IMPORTS & KONFIGURASI
# 2.  ENGINE & CORE        : SoundEngine, TimeSync, LoadingDialog, AppState
# 3.  THEME ENGINE         : _T, apply_theme, build_ss (+PyQtDarkTheme), _accent_overlay
# 4.  UI SHELL             : TopBar, NavBar, MainWindow
# 5.  TASK SYSTEM          : AddTaskDialog, EditTaskDialog, FolderWidget, FolderDialog,
#                            DraggableCard, TaskPage
# 6.  SPORT                : AddSportActivityDialog, SportTrackPage
# 7.  ECONOMY              : AddEconomyDialog, EconomyPage, AddEditDebtDialog,
#                            AddEditSavingDialog, AddToSavingDialog, AddInvestmentDialog,
#                            AddSubscriptionDialog
# 8.  SHOP / PETS / GUILD  : ShopPage, PetsPage, GuildPage
# 9.  DASHBOARD & RANK     : DashboardPage, ProgressRing, RankListDialog
# 10. ACHIEVEMENT/PROFILE  : AchievementPage, ProfilePage, SettingsPage, NotifPopup
# 11. SOCIAL               : LeaderboardPage, FriendsPage, ChatDialog, GuildChatDialog,
#                            FriendProfileDialog
# 12. AUTH                 : LoginWindow, ResetPasswordBySecurityDialog,
#                            ChooseResetMethodDialog, ResetPasswordByBackupCodeDialog
# 13. HEALTH & FOOD        : AddFoodDialog, AddRecipeDialog, SetGoalsDialog,
#                            RecipeManagerDialog, HealthGoalsDialog, HealthFoodPage
# 14. CALENDAR/NOTES/RMND  : CalendarPage, NotesTextEdit, NotesPage, RemindersPage,
#                            ReminderDialog
# 15. MUSIC                : MusicPage
# 16. ENTRY POINT          : main()
# ══════════════════════════════════════════════════════════════════════════════


# =============================================================================
# STANDARD LIBRARY IMPORTS
# =============================================================================
import calendar as calmod
import json
import os
import sys
import tempfile
import time as time_module
import traceback
from datetime import date, datetime, timedelta
from functools import partial
from io import BytesIO
import re
import random

import mutagen
from mutagen.mp3 import MP3
from mutagen.flac import FLAC
from mutagen.mp4 import MP4
from mutagen.oggvorbis import OggVorbis
import requests

# =============================================================================
# THIRD-PARTY IMPORTS (PyQt6)
# =============================================================================
from PyQt6.QtCore import (
    QDate, QDateTime, QEasingCurve, QMimeData, QPropertyAnimation, QRect,
    Qt, QTimer, QUrl, pyqtSignal
)
from PyQt6.QtGui import QColor, QCursor, QDrag, QFont, QFontMetrics, QIcon, QKeySequence, QPen, QShortcut, QTextCharFormat, QTextCursor, QPainter, QPixmap
from PyQt6.QtMultimedia import QAudioOutput, QMediaPlayer, QMediaDevices
from PyQt6.QtWidgets import (
    QAbstractItemView, QApplication, QButtonGroup, QCheckBox, QColorDialog,
    QComboBox, QDateEdit, QDateTimeEdit, QDialog, QDialogButtonBox, QDoubleSpinBox,
    QFileDialog, QFormLayout, QFrame, QGraphicsOpacityEffect, QGridLayout,
    QGroupBox, QHBoxLayout, QHeaderView, QInputDialog, QLabel, QLayout, QLineEdit,
    QListWidget, QListWidgetItem, QMainWindow, QMenu, QMessageBox,
    QProgressBar, QPushButton, QRadioButton, QScrollArea, QSizePolicy,
    QSlider, QSpacerItem, QSpinBox, QSplitter, QStackedWidget, QSystemTrayIcon,
    QTabBar, QTabWidget, QTableWidget, QTableWidgetItem, QTextEdit,
    QTimeEdit, QToolButton, QTreeWidget, QTreeWidgetItem,
    QTreeWidgetItemIterator, QVBoxLayout, QWidget, QGraphicsView, QGraphicsScene, 
    QGraphicsProxyWidget
)

# =============================================================================
# INTERNAL MODULES
# =============================================================================
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from applog import get_logger
import mathtools
import database as db
from food_data import get_food_name
from holidays import get_holiday_name, get_holidays_for_year
from translations import get_text
try:
    import learning_helper as lh
    LEARNING_AVAILABLE = True
except ImportError as e:
    LEARNING_AVAILABLE = False
    lh = None
    print(f"[Learning] helper not available: {e}")

log = get_logger("ui")

# ═══════════════════════════════════════════════════════════════════
#  OPTIONAL IMPORTS FOR EXPORT (Excel, Word, PDF, Charts)
# ═══════════════════════════════════════════════════════════════════
try:
    import matplotlib.pyplot as plt
    import openpyxl
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.utils import get_column_letter
    from openpyxl.chart import BarChart, Reference
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    EXPORT_IMPORTS_OK = True
except ImportError as e:
    EXPORT_IMPORTS_OK = False
    print(f"Peringatan: Beberapa library ekspor tidak terinstall: {e}")

# Matplotlib tertanam di Qt (grafik korelasi kesehatan↔produktivitas)
try:
    import matplotlib
    matplotlib.use("QtAgg")
    from matplotlib.figure import Figure
    from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
    MPL_QT_OK = True
except Exception as e:
    MPL_QT_OK = False
    print(f"Peringatan: matplotlib Qt backend tidak tersedia: {e}")

def resource_path(relative_path):
    """Mendapatkan path absolut untuk file, baik saat running dari source maupun setelah di-exe."""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)


def get_appdata_session_path():
    """Mendapatkan path untuk session.json di AppData."""
    if getattr(sys, 'frozen', False):
        # Jika dijalankan sebagai .exe, simpan session di %APPDATA%\CraftLife
        appdata = os.getenv('APPDATA')
        base_dir = os.path.join(appdata, 'CraftLife')
        os.makedirs(base_dir, exist_ok=True)
        return os.path.join(base_dir, 'session.json')
    else:
        # Jika dijalankan sebagai skrip Python biasa, simpan di folder yang sama dengan skrip
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), 'session.json')

SESSION_FILE = get_appdata_session_path()

def get_icon_path(filename):
    """Mengembalikan path absolut ke file ikon, bekerja baik di development maupun PyInstaller .exe"""
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, 'icons', filename)

def setup_error_handling():
    log_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Error.txt")
    def handler(err_type, err_value, tb):
        err_text = "".join(traceback.format_exception(err_type, err_value, tb))
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"\n--- {datetime.now()} ---\n")
            f.write(err_text)
        log.critical(f"CRASH tidak tertangani:\n{err_text}")
        # ── BACKUP DATABASE SAAT CRASH ──
        try:
            import database as db
            db.backup_database()
        except Exception:
            log.exception("Backup otomatis saat crash gagal")
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Icon.Critical)
        msg.setWindowTitle(tr("app_error_title"))
        msg.setText(f"Terjadi kesalahan fatal!\n\n{err_value}")
        msg.setInformativeText("Detail error disimpan di Error.txt.\nDatabase telah dibackup otomatis.")
        msg.setStandardButtons(QMessageBox.StandardButton.Ok)
        msg.exec()
        sys.exit(1)
    sys.excepthook = handler

# ══════════════════════════════════════════════════════════════════════════════
# ANIMATION POOL — prevents GC from killing animations mid-flight
# ══════════════════════════════════════════════════════════════════════════════
_anim_pool: list = []

def fade_in(widget, ms: int = 220):
    """Fade a widget from opacity 0 → 1. NEVER call on QMainWindow or QDialog directly."""
    try:
        if hasattr(widget, '_fading') and widget._fading:
            return
        widget._fading = True
        eff = QGraphicsOpacityEffect()
        widget.setGraphicsEffect(eff)
        a = QPropertyAnimation(eff, b"opacity")
        a.setDuration(ms)
        a.setStartValue(0.0)
        a.setEndValue(1.0)
        a.setEasingCurve(QEasingCurve.Type.OutCubic)
        _anim_pool.append(a)

        def _done():
            if a in _anim_pool:
                _anim_pool.remove(a)
            widget._fading = False
            # Optionally remove effect after animation to avoid paint warnings
            widget.setGraphicsEffect(None)

        a.finished.connect(_done)
        a.start()
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════════════════════
# SOUND ENGINE
# ══════════════════════════════════════════════════════════════════════════════
_snd_queue = None
_snd_worker = None

def _ensure_sound_worker():
    global _snd_queue, _snd_worker
    if _snd_worker is not None:
        return
    import threading, queue
    _snd_queue = queue.Queue()
    def _run():
        try:
            import winsound
        except Exception:
            winsound = None
        while True:
            item = _snd_queue.get()
            if item is None:
                break
            freq, dur = item
            try:
                if winsound is not None and sys.platform == "win32":
                    winsound.Beep(freq, dur)
            except Exception:
                pass
    _snd_worker = threading.Thread(target=_run, daemon=True)
    _snd_worker.start()

def _enqueue_beep(freq, dur):
    if sys.platform != "win32":
        return
    try:
        _ensure_sound_worker()
        _snd_queue.put((freq, dur))
    except Exception:
        pass

class SoundEngine:
    enabled: bool = True

    @staticmethod
    def _beep(freq: int, dur: int):
        if not SoundEngine.enabled:
            return
        # Non-blocking: antri beep ke worker thread agar UI tidak freeze
        # (winsound.Beep bersifat sinkron & bisa nge-hang di beberapa sistem).
        _enqueue_beep(max(37, min(32767, freq)), max(30, dur))

    @staticmethod
    def complete(): SoundEngine._beep(880, 60); SoundEngine._beep(1100, 80)
    @staticmethod
    def level_up():
        for f in [523, 659, 784, 1047]: SoundEngine._beep(f, 80)
    @staticmethod
    def buy(): SoundEngine._beep(600, 80); SoundEngine._beep(800, 80)
    @staticmethod
    def error(): SoundEngine._beep(200, 180)
    @staticmethod
    def boss_hit(): SoundEngine._beep(150, 200)
    @staticmethod
    def boss_dead():
        for f in [300, 500, 800, 1200]: SoundEngine._beep(f, 85)
    @staticmethod
    def click(): SoundEngine._beep(700, 30)
    @staticmethod
    def notify(): SoundEngine._beep(520, 50); SoundEngine._beep(720, 50)

SND = SoundEngine()

# ══════════════════════════════════════════════════════════════════════════════
# Time Synchronization with online server to prevent local time exploits (e.g. changing system clock to bypass timers)
# ══════════════════════════════════════════════════════════════════════════════
class TimeSync:
    _offset = 0
    _last_sync = 0
    _zone = "Asia/Jakarta"
    _last_attempt = 0
    _backoff = 60
    DEBUG = True

    @classmethod
    def sync(cls):
        """Sinkronkan waktu dengan server (NON-BLOCKING, jalan di background)."""
        now = time_module.time()
        if cls.DEBUG:
            print(f"[TimeSync] Sync dipanggil. Backoff tersisa: {cls._backoff:.0f}s")
        if now - cls._last_attempt < cls._backoff:
            return False
        cls._last_attempt = now
        try:
            import threading
            threading.Thread(target=cls._sync_blocking, daemon=True).start()
        except Exception:
            cls._sync_blocking()
        return True

    @classmethod
    def _sync_blocking(cls):
        """Network call (blocking) — HANYA dipanggil dari thread terpisah."""
        apis = [
            f"https://time.now/developer/api/timezone/{cls._zone}",
            f"https://timeapi.world/api/timezone/{cls._zone}",
            "https://1.1.1.1/cdn-cgi/trace",
        ]
        for url in apis:
            try:
                resp = requests.get(url, timeout=5)
                if resp.status_code != 200:
                    continue
                server_unixtime = None
                if "time.now" in url or "timeapi.world" in url:
                    server_unixtime = resp.json().get("unixtime")
                else:
                    for line in resp.text.splitlines():
                        if line.startswith("ts="):
                            server_unixtime = int(float(line.split("=")[1])); break
                    else:
                        continue
                if server_unixtime is None:
                    continue
                cls._offset = server_unixtime - int(time_module.time())
                cls._last_sync = time_module.time()
                cls._backoff = 3600
                if cls.DEBUG:
                    print(f"[TimeSync] Sync sukses (offset={cls._offset}s)")
                return
            except Exception as e:
                if cls.DEBUG:
                    print(f"[TimeSync] gagal {type(e).__name__}: {e}")
                continue
        cls._backoff = min(cls._backoff * 2, 3600)
        if cls.DEBUG:
            print(f"[TimeSync] SEMUA API GAGAL, backoff={cls._backoff:.0f}s")

    @classmethod
    def get_current_time(cls):
        """Mendapatkan datetime objek dari waktu server yang sudah disinkronkan"""
        if cls._last_sync == 0:
            if cls.DEBUG:
                print("[TimeSync] ⚠️ Belum pernah sync, mencoba sync...")
            cls.sync()
        local_now = time_module.time()
        server_now = local_now + cls._offset
        dt = datetime.fromtimestamp(server_now)
        
        if cls.DEBUG:
            print(f"[TimeSync] 🕐 get_current_time() -> {dt.strftime('%Y-%m-%d %H:%M:%S')} (offset={cls._offset})")
        
        return dt

    @classmethod
    def get_formatted_time(cls):
        """Format waktu HH:MM:SS"""
        dt = cls.get_current_time()
        return dt.strftime("%H:%M:%S")

# ══════════════════════════════════════════════════════════════════════════════
# THEME ENGINE
# ══════════════════════════════════════════════════════════════════════════════
_theme: dict = {
    "primary": "#5a8a2e", "light": "#7bbf3e",
    "bg": "#1a1a1a", "panel": "#2d2d2d", "border": "#444",
    "accent": "#80c000", "text": "#e8e8e8", "muted": "#888",
}

def _T(key: str) -> str:
    return _theme.get(key, "#888")

def apply_theme(theme_dict: dict):
    global _theme
    _theme = theme_dict


# ══════════════════════════════════════════════════════════════════════════════
#  ♿ AKSESIBILITAS — skala font global & mode kontras tinggi
# ══════════════════════════════════════════════════════════════════════════════
_FONT_SCALE = 1.0          # 0.8 – 1.4
_HIGH_CONTRAST = False


def set_font_scale_factor(pct) -> None:
    """Atur faktor skala font global dari persen (80–140)."""
    global _FONT_SCALE
    try:
        _FONT_SCALE = max(0.8, min(1.4, int(pct) / 100.0))
    except (TypeError, ValueError):
        _FONT_SCALE = 1.0


def _font_scale_overlay() -> str:
    """Aturan QSS tambahan di urutan terakhir supaya menang dari basis.
    Juga menyamakan ukuran font (otomatis ukuran emoji inline) pada tombol,
    toolbar, dan semua tab/subtab di seluruh aplikasi."""
    px = max(9, round(13 * _FONT_SCALE))
    return (
        f"\n/* === font scale {_FONT_SCALE:.2f} + emoji/ikon normalization === */\n"
        f"QWidget {{ font-size: {px}px; }}\n"
        f"QPushButton, QToolButton, QRadioButton, QCheckBox {{ font-size: {px}px; }}\n"
        f"QTabBar::tab {{ font-size: {px}px; }}\n"
        f"QToolBox::tab {{ font-size: {px}px; }}\n"
    )


def _high_contrast_overlay() -> str:
    if not _HIGH_CONTRAST:
        return ""
    return (
        "\n/* === high contrast === */\n"
        "QWidget { background: #000000; color: #ffff00; }\n"
        "QFrame#card, QGroupBox { background: #000000; border: 2px solid #ffff00; }\n"
        "QPushButton { color: #ffff00; background: #000000; border: 2px solid #ffff00; }\n"
        "QPushButton:hover { background: #333300; }\n"
        "QPushButton:disabled { color: #777700; border-color: #777700; }\n"
        "QLineEdit, QTextEdit, QPlainTextEdit, QComboBox, QSpinBox, QDoubleSpinBox,"
        " QListWidget, QTableWidget, QTreeWidget, QTabWidget::pane {"
        " color: #ffff00; background: #000000; border: 2px solid #ffff00;"
        " selection-background-color: #333300; }\n"
        "QLabel#sub, QLabel#muted { color: #ffe97a; }\n"
        "QProgressBar { background: #000000; border: 2px solid #ffff00; color: #ffff00; }\n"
        "QProgressBar::chunk { background: #ffff00; }\n"
        "QScrollBar { background: #000000; }\n"
        "QScrollBar::handle { background: #ffff00; }\n"
        "QMenu { background: #000000; color: #ffff00; border: 2px solid #ffff00; }\n"
        "QMenu::item:selected { background: #333300; }\n"
    )


def apply_accessibility(user: dict) -> None:
    """Terapkan preferensi aksesibilitas user (font scale + kontras tinggi).
    Panggil SETELAH apply_theme(), sebelum build_ss()."""
    global _theme, _HIGH_CONTRAST
    user = user or {}
    set_font_scale_factor(user.get("font_scale", 100))
    _HIGH_CONTRAST = bool(user.get("high_contrast", 0))
    if _HIGH_CONTRAST:
        t = dict(_theme)
        t.update({
            "bg": "#000000", "bg2": "#0a0a0a", "bg3": "#141400",
            "panel": "#000000", "border": "#ffff00",
            "text": "#ffff00", "muted": "#ffe97a", "label": "#ffff00",
            "primary": "#ffff00", "light": "#ffff66",
            "accent": "#ffff00", "glow": "#ffff00",
        })
        _theme = t

def _is_light_theme() -> bool:
    """Deteksi tema terang/gelap dari warna text (text gelap => light mode)."""
    try:
        h = _theme.get("text", "#ffffff").lstrip("#")
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        lum = (0.299 * r + 0.587 * g + 0.114 * b) / 255
        return lum < 0.5
    except Exception:
        return False

def _accent_overlay() -> str:
    """Aturan accent ber-theme yang ditumpuk di atas basis profesional PyQtDarkTheme."""
    p = _T("primary"); l = _T("light"); acc = _T("accent")
    a2 = _theme.get("accent2") or p; a3 = _theme.get("accent3") or l
    brd = _T("border"); pan = _T("panel"); txt = _T("text"); mut = _T("muted")
    return f"""
QFrame#card {{ background: {pan}; border: 1px solid {brd}; border-radius: 16px; }}
QFrame#card:hover {{ border: 1px solid {a3}; }}
QPushButton#solid {{ color: #ffffff; border: none; border-radius: 10px;
    background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 {p}, stop:1 {a2}); }}
QPushButton#solid:hover {{ background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 {l}, stop:1 {a3}); }}
QPushButton#danger {{ color: #ff6b6b; border: 1px solid rgba(255,80,80,90); }}
QPushButton#danger:hover {{ background: #ff5a5a; color: #ffffff; border: none; }}
QPushButton#gold {{ color: #ffd166; border: 1px solid rgba(255,200,80,90); }}
QPushButton#gold:hover {{ background: #f5b133; color: #1a1500; border: none; }}
QPushButton#diamond {{ color: #4dd9e0; border: 1px solid rgba(60,220,230,90); }}
QPushButton#diamond:hover {{ background: #4dd9e0; color: #00252a; border: none; }}
QLabel#section {{ color: {txt}; font-size: 17px; font-weight: 700; }}
QLabel#sub {{ color: {mut}; font-size: 12px; }}
QLabel#chip_hp {{ color: #ff6b6b; font-weight: 600; font-size: 13px; }}
QLabel#chip_mp {{ color: #4da6ff; font-weight: 600; font-size: 13px; }}
QLabel#chip_gold {{ color: #ffd166; font-weight: 600; font-size: 13px; }}
#navindicator {{ border: 1px solid {acc}; border-radius: 12px;
    background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 rgba(124,92,255,95), stop:1 rgba(34,211,238,95)); }}
/* === density & spacing (ringkas & padat) === */
QWidget {{ font-size: 13px; }}
QPushButton {{ padding: 6px 16px; min-height: 32px; border-radius: 8px; }}
QLineEdit, QTextEdit, QPlainTextEdit, QComboBox, QSpinBox {{ padding: 6px 10px; border-radius: 8px; }}
QListWidget::item {{ padding: 6px 8px; }}
QGroupBox {{ padding: 12px 12px 10px 12px; margin-top: 16px; border-radius: 11px; }}
QTabBar::tab {{ padding: 8px 18px; }}
QProgressBar {{ height: 12px; border-radius: 6px; }}
QFrame#card {{ border-radius: 14px; }}
QLabel#section {{ font-size: 16px; }}
"""

def build_ss() -> str:
    """Stylesheet aplikasi. Pakai PyQtDarkTheme (basis profesional) bila tersedia,
    ditumpuk accent sesuai tema. Fallback ke _legacy_full_ss() bila lib belum terpasang."""
    overlay = _accent_overlay()
    a11y = _font_scale_overlay() + _high_contrast_overlay()
    try:
        import qdarktheme
        mode = "light" if _is_light_theme() else "dark"
        base = qdarktheme.load_stylesheet(mode)
        return base + "\n/* === CraftLife accent overlay === */\n" + overlay + a11y
    except Exception as e:
        try:
            print(f"[Theme] PyQtDarkTheme tidak tersedia ({e}); pakai stylesheet bawaan.")
        except Exception:
            pass
        return _legacy_full_ss() + a11y

def _legacy_full_ss() -> str:
    p = _T("primary"); l = _T("light"); bg = _T("bg"); pan = _T("panel")
    brd = _T("border"); acc = _T("accent"); txt = _T("text"); mut = _T("muted")
    bg2 = _theme.get("bg2") or bg
    bg3 = _theme.get("bg3") or bg
    a2 = _theme.get("accent2") or p
    a3 = _theme.get("accent3") or l
    return f"""
/* ── Global / Aurora ── */
QMainWindow, QDialog {{
    background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
        stop:0 {bg}, stop:0.5 {bg2}, stop:1 {bg3});
    color: {txt};
    font-family: 'Inter','Segoe UI','SF Pro Text',Arial,sans-serif;
    font-size: 13px;
}}
QWidget {{ background: transparent; color: {txt}; }}
QToolTip {{ background: {pan}; color: {txt}; border: 1px solid {brd}; border-radius: 8px; padding: 6px 10px; }}
/* ── Tabs ── */
QTabWidget::pane {{ border: none; background: transparent; top:-1px; }}
QTabBar {{ background: transparent; }}
QTabBar::tab {{ background: transparent; color: {mut}; padding: 9px 16px;
    border: none; border-bottom: 2px solid transparent; font-weight: 600; font-size: 12px; min-height: 30px; }}
QTabBar::tab:selected {{ color: {l}; border-bottom: 2px solid {l}; }}
QTabBar::tab:hover:!selected {{ color: {txt}; border-bottom: 2px solid {brd}; }}
/* ── Buttons ── */
QPushButton {{ background: {pan}; color: {l}; border: 1px solid {brd}; border-radius: 10px;
    padding: 8px 16px; font-weight: 600; font-size: 12px; min-height: 34px; }}
QPushButton:hover {{ background: {pan}; color: #ffffff; border: 1px solid {a3}; }}
QPushButton:pressed {{ background: {bg}; }}
QPushButton:disabled {{ color: {mut}; border-color: {brd}; }}
QPushButton#danger {{ background: {pan}; color: #ff6b6b; border: 1px solid rgba(255,80,80,90); }}
QPushButton#danger:hover {{ color: #fff; border: none;
    background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #ff5a5a, stop:1 #e23636); }}
QPushButton#gold {{ background: {pan}; color: #ffd166; border: 1px solid rgba(255,200,80,90); }}
QPushButton#gold:hover {{ color: #1a1500; border: none;
    background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #ffd24d, stop:1 #f5b133); }}
QPushButton#diamond {{ background: {pan}; color: #4dd9e0; border: 1px solid rgba(60,220,230,90); }}
QPushButton#diamond:hover {{ color: #00252a; border: none;
    background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #4dd9e0, stop:1 #2bb6bd); }}
QPushButton#solid {{ color: #ffffff; border: none; border-radius: 10px;
    background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 {p}, stop:1 {a2}); }}
QPushButton#solid:hover {{ background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 {l}, stop:1 {a3}); }}
QPushButton#flat {{ background: transparent; color: {mut}; border: none; padding: 6px 10px; }}
QPushButton#flat:hover {{ color: {txt}; background: {pan}; border-radius: 10px; }}
QToolButton {{ background: transparent; color: {mut}; border: none; border-radius: 9px; padding: 6px; }}
QToolButton:hover {{ background: {pan}; color: {txt}; }}
/* ── Inputs ── */
QLineEdit {{ background: {pan}; color: {txt}; border: 1px solid {brd}; border-radius: 10px;
    padding: 8px 10px; font-size: 13px; selection-background-color: {p}; selection-color: #fff; min-height: 30px; }}
QTextEdit {{ background: {pan}; color: {txt}; border: 1px solid {brd}; border-radius: 10px;
    padding: 8px 10px; font-size: 13px; selection-background-color: {p}; selection-color: #fff; }}
QComboBox {{ background: {pan}; color: {txt}; border: 1px solid {brd}; border-radius: 10px;
    padding: 8px 10px; font-size: 13px; min-height: 30px; }}
QSpinBox {{ background: {pan}; color: {txt}; border: 1px solid {brd}; border-radius: 10px;
    padding: 6px 10px; font-size: 13px; min-height: 30px; }}
QLineEdit:focus, QTextEdit:focus, QPlainTextEdit:focus {{ border: 1px solid {a3}; }}
QComboBox:focus, QSpinBox:focus {{ border: 1px solid {a3}; }}
QComboBox::drop-down {{ border: none; width: 26px; }}
QComboBox QAbstractItemView {{ background: {pan}; color: {txt}; border: 1px solid {brd}; border-radius: 10px;
    selection-background-color: {p}; selection-color: #fff; padding: 4px; outline: none; }}
/* ── Scroll ── */
QScrollArea {{ border: none; background: transparent; }}
QScrollBar:vertical {{ background: transparent; width: 10px; margin: 2px; }}
QScrollBar::handle:vertical {{ background: {a3}; border-radius: 4px; min-height: 28px; }}
QScrollBar::handle:vertical:hover {{ background: {l}; }}
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{ height: 0; }}
QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {{ background: transparent; }}
QScrollBar:horizontal {{ background: transparent; height: 10px; margin: 2px; }}
QScrollBar::handle:horizontal {{ background: {a3}; border-radius: 4px; min-width: 28px; }}
QScrollBar::handle:horizontal:hover {{ background: {l}; }}
QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{ width: 0; }}
/* ── Progress ── */
QProgressBar {{ background: {bg}; border: 1px solid {brd}; border-radius: 7px; height: 14px;
    text-align: center; font-size: 10px; color: {mut}; }}
QProgressBar::chunk {{ border-radius: 6px;
    background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 {a2}, stop:1 {a3}); }}
/* ── Cards ── */
QFrame#card {{ background: {pan}; border: 1px solid {brd}; border-radius: 16px; }}
QFrame#card:hover {{ border: 1px solid {a3}; }}
/* ── Lists ── */
QListWidget {{ background: {pan}; border: 1px solid {brd}; border-radius: 12px; color: {txt}; outline: none; }}
QListWidget::item {{ padding: 8px; border-bottom: 1px solid {brd}; }}
QListWidget::item:selected {{ color: #ffffff; background: {p}; }}
QListWidget::item:hover {{ background: {bg}; }}
/* ── Groups ── */
QGroupBox {{ color: {txt}; font-weight: 600; border: 1px solid {brd}; border-radius: 14px;
    margin-top: 14px; padding-top: 12px; background: {pan}; }}
QGroupBox::title {{ subcontrol-origin: margin; left: 12px; padding: 0 6px; color: {l}; }}
/* ── Checkboxes / Radio ── */
QCheckBox {{ color: {txt}; font-size: 13px; spacing: 8px; }}
QCheckBox::indicator, QRadioButton::indicator {{ width: 18px; height: 18px;
    border: 2px solid {brd}; border-radius: 5px; background: {bg}; }}
QCheckBox::indicator:checked, QRadioButton::indicator:checked {{ background: {l}; border-color: {l}; }}
QRadioButton::indicator {{ border-radius: 9px; }}
/* ── Labels ── */
QLabel#section {{ color: {txt}; font-size: 17px; font-weight: 700; }}
QLabel#sub {{ color: {mut}; font-size: 12px; }}
QLabel#chip_hp {{ color: #ff6b6b; font-weight: 600; font-size: 13px; }}
QLabel#chip_mp {{ color: #4da6ff; font-weight: 600; font-size: 13px; }}
QLabel#chip_gold {{ color: #ffd166; font-weight: 600; font-size: 13px; }}
/* ── Menus ── */
QMenu {{ background: {pan}; color: {txt}; border: 1px solid {brd}; border-radius: 12px; padding: 4px; }}
QMenu::item {{ padding: 6px 18px; border-radius: 6px; }}
QMenu::item:selected {{ color: #ffffff; background: {p}; }}
"""


def build_overworld_ss() -> str:
    """Stylesheet statis tema Overworld (hardcoded, tidak berubah)"""
    return """
    QMainWindow, QDialog, QWidget {
        background: #1a1a1a;
        color: #e8e8e8;
        font-family: 'Segoe UI', Arial, sans-serif;
        font-size: 13px;
    }
    QTabWidget::pane { border: 1px solid #444; background: #1a1a1a; }
    QTabBar::tab {
        background: #2d2d2d; color: #888;
        padding: 10px 16px;
        border: 1px solid #444; border-bottom: none;
        font-weight: bold; font-size: 12px;
        min-height: 30px;
    }
    QTabBar::tab:selected { background: #1a1a1a; color: #7bbf3e; border-bottom: 2px solid #7bbf3e; }
    QPushButton {
        background: #2d2d2d; color: #7bbf3e;
        border: 1px solid #5a8a2e; border-radius: 6px;
        padding: 8px 16px;
        font-weight: bold; font-size: 12px;
        min-height: 34px;
    }
    QPushButton:hover { background: #5a8a2e; color: #fff; border-color: #7bbf3e; }
    QPushButton#solid { background: #5a8a2e; color: #fff; border-color: #7bbf3e; }
    QPushButton#danger { background: #2a0808; color: #e05050; border-color: #8a2e2e; }
    QPushButton#gold { background: #2e2500; color: #f0a800; border-color: #8a7000; }
    QPushButton#diamond { background: #00252a; color: #4dd9e0; border-color: #006a6a; }
    QLineEdit, QTextEdit, QComboBox, QSpinBox {
        background: #111; color: #e8e8e8;
        border: 1px solid #444; border-radius: 5px;
        padding: 8px 10px; font-size: 13px;
        min-height: 30px;
    }
    QProgressBar {
        background: #1a1a1a; border: 1px solid #444;
        border-radius: 4px; height: 14px;
        text-align: center; font-size: 10px; color: #888;
    }
    QProgressBar::chunk { background: #80c000; border-radius: 3px; }
    QFrame#card { background: #2d2d2d; border: 1px solid #444; border-radius: 8px; }
    QGroupBox { color: #7bbf3e; font-weight: bold; border: 1px solid #444; border-radius: 6px; margin-top: 12px; padding-top: 10px; }
    QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 6px; background: #1a1a1a; }
    QCheckBox { color: #e8e8e8; font-size: 13px; spacing: 6px; }
    QCheckBox::indicator { width: 18px; height: 18px; border: 2px solid #5a8a2e; border-radius: 4px; background: #111; }
    QCheckBox::indicator:checked { background: #7bbf3e; border-color: #7bbf3e; }
    QLabel#section { color: #7bbf3e; font-size: 14px; font-weight: bold; }
    QLabel#sub { color: #888; font-size: 12px; }
    QLabel#chip_hp { color: #e05050; font-weight: bold; font-size: 13px; }
    QLabel#chip_mp { color: #4da6ff; font-weight: bold; font-size: 13px; }
    QLabel#chip_gold { color: #f0a800; font-weight: bold; font-size: 13px; }
    QListWidget { background: #1a1a1a; border: 1px solid #444; border-radius: 6px; color: #e8e8e8; }
    QListWidget::item { padding: 8px; border-bottom: 1px solid #444; }
    QListWidget::item:selected { background: #5a8a2e; color: #fff; }
    QScrollArea { border: none; background: transparent; }
    QScrollBar:vertical { background: #1a1a1a; width: 7px; border-radius: 3px; }
    QScrollBar::handle:vertical { background: #444; border-radius: 3px; }
    QScrollBar::handle:vertical:hover { background: #5a8a2e; }
    """
# ══════════════════════════════════════════════════════════════════════════════
#  Loading Animation
# ══════════════════════════════════════════════════════════════════════════════
class LoadingDialog(QDialog):
    """Dialog loading sederhana dengan animasi titik-titik"""
    def __init__(self, message="Loading...", parent=None):
        super().__init__(parent)
        
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setModal(True)
        self.setMinimumSize(300, 100)
        self.setStyleSheet("background: #2d2d2d; border: 1px solid #5a8a2e; border-radius: 10px;")
        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.label = QLabel(message)
        self.label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.label.setStyleSheet("color: #e8e8e8; font-size: 14px;")
        layout.addWidget(self.label)
        
        self.dot_count = 0
        self.timer = QTimer()
        self.timer.timeout.connect(self._animate)
        self.timer.start(500)

    def _animate(self):
        self.dot_count = (self.dot_count + 1) % 4
        dots = "." * self.dot_count
        base_text = self.label.text().split(".")[0]
        self.label.setText(f"{base_text}{dots}")
    
    def closeEvent(self, e):
        self.timer.stop()
        super().closeEvent(e)


def save_session(user_id, username):
    """Simpan session 'Remember Me' memakai TOKEN ACAK (bukan password_hash).

    File di disk hanya berisi token sekali pakai yang:
      - diverifikasi terhadap hash-nya di database,
      - punya masa berlaku (db.SESSION_TOKEN_DAYS hari),
      - hangus saat logout / ganti password / akun di-lock.
    """
    try:
        token = db.create_session_token(user_id)
        data = {
            "user_id": user_id,
            "username": username,
            "token": token,
            "timestamp": datetime.now().isoformat()
        }
        # Tulis ke file sementara lalu rename untuk atomic operation
        temp_file = SESSION_FILE + ".tmp"
        with open(temp_file, "w") as f:
            json.dump(data, f)
        os.replace(temp_file, SESSION_FILE)  # atomic di Windows
        try:
            os.chmod(SESSION_FILE, 0o600)   # hanya pemilik yang bisa baca (POSIX)
        except OSError:
            pass
    except Exception as e:
        log.error(f"Save session error: {e}")

def load_session():
    """Load session jika token valid & belum kedaluwarsa."""
    if not os.path.exists(SESSION_FILE):
        return None
    try:
        with open(SESSION_FILE, "r") as f:
            data = json.load(f)
        # Validasi struktur
        if not all(k in data for k in ("user_id", "username", "token")):
            clear_session()
            return None
        # Verifikasi token ke database (hash + expiry)
        if db.validate_session_token(data["user_id"], data["token"]):
            user = db.get_user(data["user_id"])
            if user:
                return {"user_id": data["user_id"], "username": data["username"]}
        # Session tidak valid, hapus file
        clear_session()
        return None
    except Exception as e:
        log.warning(f"Load session gagal: {e}")
        clear_session()
        return None

def clear_session():
    """Hapus file session DAN hanguskan token-nya di database."""
    try:
        if os.path.exists(SESSION_FILE):
            try:
                with open(SESSION_FILE, "r") as f:
                    data = json.load(f)
                if "user_id" in data and "token" in data:
                    db.delete_session_token(data["user_id"], data["token"])
            except Exception:
                pass
            os.remove(SESSION_FILE)
    except Exception as e:
        log.warning(f"Clear session error: {e}")
# ══════════════════════════════════════════════════════════════════════════════
#  APP STATE  ─ single source of truth, avoids restart-to-sync bugs
# ══════════════════════════════════════════════════════════════════════════════
class AppState:
    user_id: int = 0
    _lang: str = "en"
    _cbs: list = []
    _lang_cbs: list = []

    @classmethod
    def set_user(cls, uid: int):
        cls.user_id = uid
        t = db.get_user_theme(uid)
        apply_theme(t)
        SoundEngine.enabled = bool(db.get_user(uid).get("sound_enabled", 1))
        # Load language from DB
        cls._lang = db.get_user_language(uid)

    @classmethod
    def get_language(cls):
        return cls._lang

    @classmethod
    def set_language(cls, lang_code):
        if lang_code not in ("id", "en"):
            return
        cls._lang = lang_code
        if cls.user_id:
            db.set_user_language(cls.user_id, lang_code)
        # Trigger all language callbacks
        for cb in cls._lang_cbs:
            try:
                cb()
            except Exception:
                pass

    @classmethod
    def register_lang_cb(cls, cb):
        if cb not in cls._lang_cbs:
            cls._lang_cbs.append(cb)

    @classmethod
    def unregister_lang_cb(cls, cb):
        if cb in cls._lang_cbs:
            cls._lang_cbs.remove(cb)

    @classmethod
    def user(cls) -> dict:
        return db.get_user(cls.user_id) if cls.user_id else {}

    @classmethod
    def refresh(cls):
        for cb in list(cls._cbs):
            try: cb()
            except Exception: pass

    @classmethod
    def register(cls, cb):
        if cb not in cls._cbs: cls._cbs.append(cb)

    @classmethod
    def unregister(cls, cb):
        if cb in cls._cbs: cls._cbs.remove(cb)

# ══════════════════════════════════════════════════════════════════════════════
#  SMALL HELPERS
# ══════════════════════════════════════════════════════════════════════════════
# Fungsi global untuk memudahkan akses terjemahan
def tr(key, **kwargs):
    """Get translated text for current user language."""
    lang = AppState.get_language()
    text = get_text(key, lang)
    if kwargs:
        return text.format(**kwargs)
    return text

def _lbl(text, obj="", size=13, bold=False):
    if size <= 0: size = 10
    size = max(6, round(size * _FONT_SCALE))
    w = QLabel(text)
    w.setFont(QFont("Segoe UI", size, QFont.Weight.Bold if bold else QFont.Weight.Normal))
    if obj: w.setObjectName(obj)
    return w

def _btn(text, obj="", slot=None, h=38):
    b = QPushButton(text)
    if obj:
        b.setObjectName(obj)
    if slot:
        b.clicked.connect(slot)
    b.setMinimumHeight(h)
    return b

from PyQt6.QtWidgets import QGraphicsDropShadowEffect

def _card() -> QFrame:
    f = QFrame()
    f.setObjectName("card")
    # Catatan: QGraphicsDropShadowEffect sengaja TIDAK dipakai per-kartu karena
    # sangat berat saat direbuild tiap load() (sumbu utama delay pindah tab).
    # Tampilan modern sudah diatur via QSS (#card: border-radius + border + bg glass).
    return f

# ── Skala ikon/emoji konsisten seluruh aplikasi (medium) ──
ICON_INLINE = 14   # ikon sejajar teks: navbar, chips, folder, tab, tombol, list
ICON_CARD   = 24   # ikon kartu: stat shop/pet/achievement/task/economy cards
ICON_HERO   = 30   # avatar, rank, boss & hero dialog


def _fmt_qty(v) -> str:
    """Format kuantitas: bilangan bulat tanpa desimal ('x1'), pecahan 1 desimal
    ('x0.5') — supaya tidak tampil 'x1.0'."""
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v)
    return str(int(f)) if f.is_integer() else f"{f:.1f}"

def _emoji_label(text: str, px: int = ICON_CARD) -> QLabel:
    """Render emoji ke pixmap di ukuran PASTI (px) memakai font 'Segoe UI Emoji'.
    Bebas dari quirk scaling emoji -> selalu tampil sebesar px di sistem apa pun.
    Ukuran dinormalisasi ke tier medium (ICON_INLINE / ICON_CARD / ICON_HERO)."""
    lbl = QLabel()
    size = px + 6
    pm = QPixmap(size, size)
    pm.fill(Qt.GlobalColor.transparent)
    p = QPainter(pm)
    p.setRenderHint(QPainter.RenderHint.Antialiasing)
    p.setRenderHint(QPainter.RenderHint.TextAntialiasing)
    f = QFont("Segoe UI Emoji"); f.setPixelSize(px)
    p.setFont(f)
    p.drawText(pm.rect(), Qt.AlignmentFlag.AlignCenter, text)
    p.end()
    lbl.setPixmap(pm)
    lbl.setFixedSize(size, size)
    return lbl

def _scrolled(inner):
    sa = QScrollArea()
    sa.setWidgetResizable(True)
    sa.setWidget(inner)
    sa.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
    return sa

def _sep():
    f = QFrame()
    f.setFrameShape(QFrame.Shape.HLine)
    f.setStyleSheet(f"color: {_T('border')};")
    return f

def _combo(items):
    cb = QComboBox()
    cb.setMinimumHeight(42)
    for text, data in items:
        cb.addItem(text, data)
    return cb

def _parse_positive_amount(text: str):
    """Parse jumlah uang (float > 0) dari input teks.

    Return float jika valid, None jika kosong/bukan angka/<=0.
    Dipakai semua dialog ekonomi agar validasi tidak duplikatif.
    """
    try:
        val = float((text or "").strip().replace(",", "."))
    except (ValueError, AttributeError):
        return None
    return val if val > 0 else None


def _input(placeholder="", password=False):
    f = QLineEdit()
    f.setPlaceholderText(placeholder)
    f.setMinimumHeight(42)
    if password: f.setEchoMode(QLineEdit.EchoMode.Password)
    return f

def _show(parent, title, msg, kind="info"):
    dlg = QDialog(parent)
    dlg.setWindowTitle(title)
    dlg.setMinimumWidth(380)
    dlg.setMinimumHeight(140)
    dlg.setStyleSheet(build_ss())
    lay = QVBoxLayout(dlg)
    lay.setContentsMargins(24,20,24,20)
    lay.setSpacing(14)
    icon = {"info":"💬","success":"✅","error":"❌","warning":"⚠️","levelup":"🎉"}.get(kind,"💬")
    lbl = QLabel(f"{icon}  {msg}")
    lbl.setWordWrap(True)
    lbl.setStyleSheet(f"color: {_T('text')}; font-size:13px;")
    lay.addWidget(lbl)
    ok = _btn(tr("msg_ok"), "solid", dlg.accept, 40)
    lay.addWidget(ok)
    dlg.exec()



# ══════════════════════════════════════════════════════════════════════════════
#  TOP BAR
# ══════════════════════════════════════════════════════════════════════════════
class TopBar(QWidget):
    def __init__(self, on_notif, on_profile, on_toggle_theme=None):
        super().__init__()
        self._on_toggle_theme = on_toggle_theme

        self.setMinimumHeight(72)
        self._update_bg()

        lay = QHBoxLayout(self)
        lay.setContentsMargins(16, 8, 16, 8)
        lay.setSpacing(12)

        self._logo = QLabel(tr("app_logo"))
        self._logo.setFont(QFont("Segoe UI", 18, QFont.Weight.Bold))
        self._logo.setStyleSheet(f"color: {_T('light')};")
        lay.addWidget(self._logo)
        lay.addSpacing(8)

        # XP column
        xp_col = QVBoxLayout()
        xp_col.setSpacing(3)
        self._xp_lbl = QLabel(tr("level_label", lvl=1))
        self._xp_lbl.setStyleSheet(
            f"color: {_T('accent')}; font-size: 11px; font-weight: bold;")
        self._xp_bar = QProgressBar()
        self._xp_bar.setMinimumHeight(10)
        self._xp_bar.setStyleSheet(
            f"QProgressBar::chunk {{ background: {_T('accent')}; border-radius: 4px; }}")
        xp_col.addWidget(self._xp_lbl)
        xp_col.addWidget(self._xp_bar)
        xp_w = QWidget()
        xp_w.setLayout(xp_col)
        xp_w.setSizePolicy(QSizePolicy.Policy.Expanding,
                            QSizePolicy.Policy.Preferred)
        lay.addWidget(xp_w)
        lay.addSpacing(8)

        # Stat chips
        self._hp_lbl   = self._chip("❤️ --",   "chip_hp")
        self._mp_lbl   = self._chip("💙 -- MP", "chip_mp")
        self._gold_lbl = self._chip("💰 --",    "chip_gold")
        self._time_lbl = QLabel("--:--:--")
        self._time_lbl.setObjectName("chip_time")
        self._time_lbl.setStyleSheet(
            f"background: {_T('panel')}; border: 1px solid {_T('border')};"
            f" border-radius: 6px; padding: 3px 10px;"
            f" font-family: monospace; font-size: 13px;"
        )
        lay.addWidget(self._time_lbl)
        for w in [self._hp_lbl, self._mp_lbl, self._gold_lbl]:
            lay.addWidget(w)

        self._theme_toggle = _btn("🌙", slot=self._do_toggle_theme)
        self._theme_toggle.setMinimumWidth(44)
        self._theme_toggle.setMinimumHeight(34)
        self._theme_toggle.setToolTip(tr("theme_toggle_tip"))
        lay.addWidget(self._theme_toggle)
        self._notif_btn = _btn("🔔", slot=on_notif)
        self._notif_btn.setMinimumWidth(44)
        self._notif_btn.setMinimumHeight(34)
        lay.addWidget(self._notif_btn)

        prof_btn = _btn("👤", slot=on_profile)
        prof_btn.setMinimumWidth(44)
        prof_btn.setMinimumHeight(34)
        lay.addWidget(prof_btn)

        exit_btn = _btn("🚪", slot=QApplication.instance().quit)
        exit_btn.setMinimumWidth(44)
        exit_btn.setMinimumHeight(34)
        lay.addWidget(exit_btn)

        logout_btn = _btn(tr("logout"), slot=self._logout)
        logout_btn.setMinimumWidth(70)
        logout_btn.setMinimumHeight(34)
        lay.addWidget(logout_btn)

        self.refresh()

        # Timer untuk jam digital (update setiap detik)
        self._time_timer = QTimer()
        self._time_timer.timeout.connect(self._update_time)
        self._time_timer.start(1000)
        self._update_time()  # langsung tampilkan

    # ── Timer ────────────────────────────────────────────────────────────────
    def _update_time(self):
        try:
            self._time_lbl.setText(TimeSync.get_formatted_time())
            if TimeSync._last_sync > 0:
                self._time_lbl.setStyleSheet(
                    f"background: {_T('panel')}; border: 1px solid {_T('border')};"
                    f" border-radius: 6px; padding: 3px 10px;"
                    f" font-family: monospace; font-size: 13px;"
                )
        except Exception:
            self._time_lbl.setText(datetime.now().strftime("%H:%M:%S"))
            self._time_lbl.setStyleSheet("color: #e05050;")

    # ── LogOut ────────────────────────────────────────────────────────────────
    def _logout(self):
        # ── FORCE CHECKPOINT SEBELUM LOGOUT ──
        db.force_checkpoint()
        
        clear_session()
        main_win = self.window()
        for widget in QApplication.topLevelWidgets():
            if isinstance(widget, LoginWindow):
                widget.show()
                break
        else:
            login = LoginWindow()
            login.show()
        main_win.close()

    # ── internal ──────────────────────────────────────────────────────────────

    def _chip(self, text: str, obj: str) -> QLabel:
        w = QLabel(text)
        w.setObjectName(obj)
        w.setStyleSheet(
            f"background: {_T('panel')}; border: 1px solid {_T('border')};"
            f" border-radius: 6px; padding: 3px 10px;"
        )
        return w

    def _update_bg(self):
        self.setStyleSheet(
            f"background: qlineargradient(x1:0,y1:0,x2:1,y2:0,"
            f"stop:0 {_T('bg')},stop:1 {_T('panel')});"
            f"border-bottom: 2px solid {_T('primary')};"
        )

    # ── public ────────────────────────────────────────────────────────────────
    def refresh(self):
        u = AppState.user()
        if not u:
            return
        lvl  = u["level"]
        xp   = u["xp"]
        need = lvl * 150
        dn   = u.get("display_name", "") or u.get("username", "")
        self._xp_lbl.setText(tr("level_xp_format", lvl=lvl, name=dn, xp=xp, need=need))
        self._xp_bar.setMaximum(need)
        self._xp_bar.setValue(int(xp))

        if db.is_account_locked(AppState.user_id):
            self._lock_indicator = QLabel("🔒 LOCKED")
            self._lock_indicator.setStyleSheet("color:#e05050; font-weight:bold; margin-left:10px;")
        else:
            pass

        self._hp_lbl.setText(tr("hp_format", hp=u['hp'], max_hp=u['max_hp']))

        # MP chip shows skill name
        cls   = u.get("avatar_class", "warrior")
        skill = db.CLASS_SKILLS.get(cls, {})
        self._mp_lbl.setText(tr("mp_format", mp=u['mp'], max_mp=u['max_mp'], skill=skill.get('name','?')))

        self._gold_lbl.setText(tr("gold_format", gold=u['gold']))
        notifs = db.get_notifications(AppState.user_id)
        self._notif_btn.setText(
            tr("notif_button", count=len(notifs)) if notifs else tr("notif_button_empty"))

    def load(self):
        """Update UI text when language changes"""
        self._logo.setText(tr("app_logo"))
        self.refresh() 

    def retheme(self):
        self._update_bg()
        # Logo
        self._logo.setStyleSheet(f"color: {_T('light')};")
        # XP row
        self._xp_lbl.setStyleSheet(
            f"color: {_T('accent')}; font-size: 11px; font-weight: bold;")
        self._xp_bar.setStyleSheet(
            f"QProgressBar::chunk {{ background: {_T('accent')}; border-radius:4px; }}")
        # Stat chips — wajib di-restyle saat theme berubah
        _chip_ss = (
            f"background: {_T('panel')}; border: 1px solid {_T('border')};"
            f" border-radius: 6px; padding: 3px 10px;"
        )
        self._hp_lbl.setStyleSheet(_chip_ss)
        self._mp_lbl.setStyleSheet(_chip_ss)
        self._gold_lbl.setStyleSheet(_chip_ss)
        # Jam digital
        self._time_lbl.setStyleSheet(
            f"background: {_T('panel')}; border: 1px solid {_T('border')};"
            f" border-radius: 6px; padding: 3px 10px;"
            f" font-family: monospace; font-size: 13px;"
        )
        self._refresh_theme_toggle()
        self.refresh()

    def _do_toggle_theme(self):
        if self._on_toggle_theme:
            self._on_toggle_theme()

    def _refresh_theme_toggle(self):
        if not hasattr(self, "_theme_toggle"):
            return
        try:
            cur = (db.get_user(AppState.user_id) or {}).get("theme", "modern_dark")
        except Exception:
            cur = "modern_dark"
        self._theme_toggle.setText("☀️" if cur == "modern_dark" else "🌙")


# ══════════════════════════════════════════════════════════════════════════════
#  NAV BAR  (left sidebar)
# ══════════════════════════════════════════════════════════════════════════════
class NavBar(QWidget):
    tab_changed = pyqtSignal(str)

    ICON_MAP = {
        "dashboard": "🏠",
        "habits": "⛏",
        "dailies": "📅",
        "todos": "📜",
        "sport": "🏅",
        "economy": "💰",
        "health_food": "💚",
        "calendar": "📅",
        "notes": "📝",
        "learning": "📚",
        "reminders": "⏰",
        "music": "🎵",
        "pomodoro": "🍅",
        "crafting": "🔨",
        "shop": "🏪",
        "pets": "🐾",
        "friends": "👥",
        "guild": "⚔️",
        "achievements": "🏆",
        "profile": "🎭",
        "settings": "⚙️",
        "leaderboard": "🎖️",
    }

    _TABS = [
        ("nav_dashboard",   "dashboard"),
        ("nav_habits",      "habits"),
        ("nav_dailies",     "dailies"),
        ("nav_quests",      "todos"),
        ("nav_pomodoro",    "pomodoro"),
        ("nav_sport",       "sport"),
        ("nav_crafting",    "crafting"),
        ("nav_economy",     "economy"),
        ("nav_health_food", "health_food"),
        ("nav_calendar", "calendar"),
        ("nav_notes", "notes"),
        ("nav_learning", "learning"),
        ("nav_reminders", "reminders"),
        ("nav_music", "music"),
        ("nav_shop",        "shop"),
        ("nav_pets",        "pets"),
        ("nav_friends",     "friends"),
        ("nav_guild",       "guild"),
        ("nav_achievements","achievements"),
        ("nav_profile",     "profile"),
        ("nav_settings",    "settings"),
        ("nav_leaderboard", "leaderboard")
    ]

    def __init__(self):
        super().__init__()
        self.setMinimumWidth(60)
        self.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Expanding)
        self._btns: dict = {}
        self._active = ""
        self._indicator = None
        self._indicator_anim = None
        self._build()
        AppState.register_lang_cb(self.reload_texts)
        self.reload_texts()
        self._select("dashboard")

    def load(self):
        """Update UI text when language changes"""
        self.reload_texts()

    def _build(self):
        lay = QVBoxLayout(self)
        lay.setAlignment(Qt.AlignmentFlag.AlignTop)
        lay.setContentsMargins(0, 6, 0, 6)
        lay.setSpacing(2)
        for key_label, key in self._TABS:
            b = QPushButton()
            b.setCheckable(True)
            b.setMinimumHeight(40)
            b.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
            b.setStyleSheet("text-align: center; padding: 4px 2px; font-size: 13px;")
            b.clicked.connect(lambda _, k=key: self._select(k))
            lay.addWidget(b)
            self._btns[key] = b
        lay.addStretch()
        self._indicator = QFrame(self)
        self._indicator.setObjectName("navindicator")
        self._indicator.setVisible(False)
        self._indicator.lower()
        self.retheme()

    def reload_texts(self):
        for key_label, key in self._TABS:
            b = self._btns.get(key)
            if b:
                icon = self.ICON_MAP.get(key, "\u2753")
                b.setText(f"{icon}\n{tr(key_label)}")

    def _style(self, active: bool) -> str:
        if active:
            return (f"QPushButton {{ background: transparent; color: #ffffff;"
                    f" border: none; border-radius: 0;"
                    f" font-size: 13px; font-weight: bold;"
                    f" border-left: 3px solid {_T('light')}; }}"
                    f"QPushButton:hover {{ background: rgba(255,255,255,18); }}")
        return (f"QPushButton {{ background: transparent; color: {_T('muted')};"
                f" border: none; border-radius: 0;"
                f" font-size: 13px; font-weight: bold; }}"
                f"QPushButton:hover {{ background: rgba(255,255,255,12);"
                f" color: {_T('text')}; }}")

    def _select(self, key: str):
        self._active = key
        for k, b in self._btns.items():
            b.setChecked(k == key)
            b.setStyleSheet(self._style(k == key))
        SND.click()
        self._move_indicator(key, animate=True)
        self.tab_changed.emit(key)

    def _move_indicator(self, key: str, animate: bool = True):
        if self._indicator is None:
            return
        b = self._btns.get(key)
        if b is None:
            return
        try:
            g = b.geometry()
            if g.height() < 5:
                return
            target = QRect(g.x() + 5, g.y() + 4, g.width() - 10, g.height() - 8)
            self._indicator.setVisible(True)
            if animate and self._indicator.isVisible():
                if self._indicator_anim is not None:
                    self._indicator_anim.stop()
                a = QPropertyAnimation(self._indicator, b"geometry", self)
                a.setDuration(220)
                a.setStartValue(self._indicator.geometry())
                a.setEndValue(target)
                a.setEasingCurve(QEasingCurve.Type.OutCubic)
                a.start()
                self._indicator_anim = a
            else:
                self._indicator.setGeometry(target)
        except Exception:
            pass

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self._move_indicator(self._active, animate=False)

    def retheme(self):
        self.setStyleSheet(
            f"background: qlineargradient(x1:0,y1:0,x2:0,y2:1,"
            f" stop:0 {_T('panel')}, stop:1 {_T('bg')});"
            f"border-right: 1px solid {_T('border')};")
        if self._indicator is not None:
            self._indicator.setStyleSheet(
                f"#navindicator {{ border-radius: 12px; border: 1px solid {_T('accent')};"
                f" background: qlineargradient(x1:0,y1:0,x2:1,y2:1,"
                f" stop:0 rgba(124,92,255,95), stop:1 rgba(34,211,238,95)); }}")
        for k, b in self._btns.items():
            b.setStyleSheet(self._style(k == self._active))
        self._move_indicator(self._active, animate=False)


# ══════════════════════════════════════════════════════════════════════════════
#  ADD TASK DIALOG  (fixed: proper heights, scroll if needed)
# ══════════════════════════════════════════════════════════════════════════════
class AddTaskDialog(QDialog):
    def __init__(self, mode: str, user_id: int, parent=None):
        super().__init__(parent)
        self.mode    = mode
        self.user_id = user_id
        titles = {"habit": tr("dialog_add_habit"),
                  "daily": tr("dialog_add_daily"),
                  "todo":  tr("dialog_add_todo")}
        self.setWindowTitle(titles.get(mode, tr("dialog_add")))
        self.setMinimumWidth(460)
        self.setMinimumHeight(460)
        self.setMaximumHeight(700)
        self.setStyleSheet(build_ss())
        self._build()

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(14)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))
        lay.addWidget(_sep())

        # Name
        lay.addWidget(_lbl(tr("dialog_name"), size=12))
        self._name = _input(tr("dialog_name_placeholder"))
        lay.addWidget(self._name)

        # Icon
        lay.addWidget(_lbl(tr("dialog_icon"), size=12))
        icons = [
            (tr("task_icon_combat"), "⚔️"),
            (tr("task_icon_study"), "📚"),
            (tr("task_icon_exercise"), "🏃"),
            (tr("task_icon_health"), "🍎"),	
            (tr("task_icon_sleep"), "💤"),	
            (tr("task_icon_mindfulness"), "🧘"),	
            (tr("task_icon_hydration"), "💧"),	
            (tr("task_icon_growth"), "🌱"),	
            (tr("task_icon_focus"), "🎯"),	
            (tr("task_icon_ideas"), "💡"),
            (tr("task_icon_quest"), "📜"),	
            (tr("task_icon_build"), "🏗️")
        ]
        self._icon = _combo(icons)
        lay.addWidget(self._icon)

        # Difficulty / priority
        if self.mode == "todo":
            lay.addWidget(_lbl(tr("dialog_priority"), size=12))
            opts = [
                (tr("task_priority_trivial_text"), "trivial"),
                (tr("task_priority_easy_text"), "easy"),
                (tr("task_priority_medium_text"), "medium"),
                (tr("task_priority_hard_text"), "hard"),
            ]
        else:
            lay.addWidget(_lbl(tr("dialog_difficulty"), size=12))
            opts = [
                (tr("task_diff_easy_text"), "easy"),
                (tr("task_diff_medium_text"), "medium"),
                (tr("task_diff_hard_text"), "hard"),
                (tr("task_diff_epic_text"), "epic"),
            ]
        self._diff = _combo(opts)
        self._diff.setCurrentIndex(1)
        lay.addWidget(self._diff)

        # Notes
        lay.addWidget(_lbl(tr("dialog_notes"), size=12))
        self._notes = _input(tr("dialog_notes_placeholder"))
        lay.addWidget(self._notes)

        # Folder
        lay.addWidget(_lbl(tr("dialog_folder"), size=12))
        _folders = db.get_task_folders(self.user_id, self.mode)
        _fopts = [(tr("dialog_no_folder"), None)] + [(f"{fd['icon']}  {fd['name']}", fd["id"]) for fd in _folders]
        self._folder = _combo(_fopts)
        lay.addWidget(self._folder)

        # Recurrence: pilih hari pengulangan (khusus habit & daily)
        if self.mode in ("habit", "daily"):
            lay.addWidget(_lbl(tr("recur_days_label"), size=12))
            self._recur = WeekdaySelector("")
            lay.addWidget(self._recur)
            lay.addWidget(_lbl(tr("recur_every_day_hint"), "sub", 11))

        lay.addSpacing(8)
        ok = _btn(tr("dialog_add"), "solid", self._save, 46)
        lay.addWidget(ok)
        self._name.returnPressed.connect(self._save)

        sa = _scrolled(content)
        root.addWidget(sa)

    def _save(self):
        name = self._name.text().strip()
        if not name:
            _show(self, tr("msg_error"), tr("msg_name_empty"), "error")
            return
        icon  = self._icon.currentData()
        diff  = self._diff.currentData()
        notes = self._notes.text()
        folder_id = self._folder.currentData()
        repeat = self._recur.get_days_str() if self.mode in ("habit", "daily") else ""
        if self.mode == "habit":
            db.add_habit(self.user_id, name, icon, diff, 1, 1, notes, repeat)
            if folder_id:
                new_id = max(i["id"] for i in db.get_habits(self.user_id))
                db.set_item_folder(self.user_id, "habit", new_id, folder_id)
        elif self.mode == "daily":
            db.add_daily(self.user_id, name, icon, diff, notes, repeat)
            if folder_id:
                new_id = max(i["id"] for i in db.get_dailies(self.user_id))
                db.set_item_folder(self.user_id, "daily", new_id, folder_id)
        else:
            db.add_todo(self.user_id, name, diff, icon, None, notes)
            if folder_id:
                new_id = max(i["id"] for i in db.get_todos(self.user_id))
                db.set_item_folder(self.user_id, "todo", new_id, folder_id)
        SND.complete()
        self.accept()


# ══════════════════════════════════════════════════════════════════════════════
#  Edit Task Dialog  (Habits / Dailies / Todos)
# ══════════════════════════════════════════════════════════════════════════════
class EditTaskDialog(QDialog):
    def __init__(self, mode, user_id, item, parent=None):
        super().__init__(parent)
        self.mode = mode
        self.user_id = user_id
        self.item = item
        self.setWindowTitle(tr("dialog_edit_mode", mode=mode.capitalize()))
        self.setMinimumWidth(460)
        self.setMinimumHeight(460)
        self.setMaximumHeight(700)
        self.setStyleSheet(build_ss())
        self._build()

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0,0,0,0)
        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24,24,24,24)
        lay.setSpacing(14)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))
        lay.addWidget(_sep())

        # Name
        lay.addWidget(_lbl(tr("dialog_name"), size=12))
        self._name = _input(tr("dialog_name_placeholder"))
        if self.item:
            self._name.setText(self.item["name"])
        lay.addWidget(self._name)

        # Icon (sama seperti AddTaskDialog)
        lay.addWidget(_lbl(tr("dialog_icon"), size=12))
        icons = [
            ("⚔️  Combat",     "⚔️"),
            ("📚  Study",      "📚"),
            ("🏃  Exercise",   "🏃"),
            ("🍎  Health",     "🍎"),
            ("💤  Sleep",      "💤"),
            ("🧘  Mindfulness","🧘"),
            ("💧  Hydration",  "💧"),
            ("🌱  Growth",     "🌱"),
            ("🎯  Focus",      "🎯"),
            ("💡  Ideas",      "💡"),
            ("📜  Quest",      "📜"),
            ("🏗️  Build",      "🏗️"),
        ]
        self._icon = _combo(icons)
        # set current icon
        index = self._icon.findData(self.item["icon"])
        if index >= 0: self._icon.setCurrentIndex(index)
        lay.addWidget(self._icon)

        # Difficulty / priority
        if self.mode == "todo":
            lay.addWidget(_lbl(tr("dialog_priority"), size=12))
            opts = [
                (tr("task_priority_trivial_text"), "trivial"), 
                (tr("task_priority_easy_text"), "easy"), 
                (tr("task_priority_medium_text"), "medium"), 
                (tr("task_priority_hard_text"), "hard")
            ]
        else:
            lay.addWidget(_lbl(tr("dialog_difficulty"), size=12))
            opts = [
                (tr("task_diff_easy_text"), "easy"), 
                (tr("task_diff_medium_text"), "medium"), 
                (tr("task_diff_hard_text"), "hard"), 
                (tr("task_diff_epic_text"), "epic")
            ]
        self._diff = _combo(opts)
        idx = self._diff.findData(self.item.get("difficulty") or self.item.get("priority","medium"))
        if idx >= 0: self._diff.setCurrentIndex(idx)
        lay.addWidget(self._diff)

        # Notes
        lay.addWidget(_lbl(tr("dialog_notes"), size=12))
        self._notes = _input(tr("dialog_notes_placeholder2"))
        if self.item:
            self._notes.setText(self.item["notes"])
        lay.addWidget(self._notes)

        # Folder
        lay.addWidget(_lbl(tr("dialog_folder"), size=12))
        _folders = db.get_task_folders(self.user_id, self.mode)
        _fopts = [(tr("dialog_no_folder"), None)] + [
            (f"{fd['icon']}  {fd['name']}", fd["id"]) for fd in _folders]
        self._folder = _combo(_fopts)
        cur_fid = self.item.get("folder_id")
        if cur_fid:
            idx = next((i for i,(_, d) in enumerate(_fopts) if d == cur_fid), 0)
            self._folder.setCurrentIndex(idx)
        lay.addWidget(self._folder)

        # Recurrence: pilih hari pengulangan (khusus habit & daily)
        if self.mode in ("habit", "daily"):
            lay.addWidget(_lbl(tr("recur_days_label"), size=12))
            self._recur = WeekdaySelector(self.item.get("repeat_days", "") if self.item else "")
            lay.addWidget(self._recur)
            lay.addWidget(_lbl(tr("recur_every_day_hint"), "sub", 11))

        lay.addSpacing(8)
        ok = _btn(tr("dialog_save"), "solid", self._save, 46)
        lay.addWidget(ok)
        sa = _scrolled(content)
        root.addWidget(sa)

    def _save(self):
        name = self._name.text().strip()
        if not name:
            _show(self, tr("msg_error"), tr("msg_name_empty"), "error")
            return
        icon = self._icon.currentData()
        diff = self._diff.currentData()
        notes = self._notes.text()
        folder_id = self._folder.currentData()
        recur_kw = {}
        if self.mode in ("habit", "daily"):
            recur_kw["repeat_days"] = self._recur.get_days_str()
        if self.mode == "habit":
            db.update_habit(self.item["id"], self.user_id, name=name, icon=icon, difficulty=diff, notes=notes, **recur_kw)
        elif self.mode == "daily":
            db.update_daily(self.item["id"], self.user_id, name=name, icon=icon, difficulty=diff, notes=notes, **recur_kw)
        else:
            db.update_todo(self.item["id"], self.user_id, name=name, icon=icon, priority=diff, notes=notes)
        db.set_item_folder(self.user_id, self.mode, self.item["id"], folder_id)
        SND.complete()
        self.accept()

# ══════════════════════════════════════════════════════════════════════════════
#  TASK ICON CATEGORIES  (sub-tab categories untuk habits/dailies/todos)
# ══════════════════════════════════════════════════════════════════════════════
TASK_ICON_CATEGORIES = {
    "⚔️":  {"name": "Combat",      "icon": "⚔️"},
    "📚":  {"name": "Study",       "icon": "📚"},
    "🏃":  {"name": "Exercise",    "icon": "🏃"},
    "🍎":  {"name": "Health",      "icon": "🍎"},
    "💤":  {"name": "Sleep",       "icon": "💤"},
    "🧘":  {"name": "Mindfulness", "icon": "🧘"},
    "💧":  {"name": "Hydration",   "icon": "💧"},
    "🌱":  {"name": "Growth",      "icon": "🌱"},
    "🎯":  {"name": "Focus",       "icon": "🎯"},
    "💡":  {"name": "Ideas",       "icon": "💡"},
    "📜":  {"name": "Quest",       "icon": "📜"},
    "🏗️": {"name": "Build",       "icon": "🏗️"},
}


# ══════════════════════════════════════════════════════════════════════════════
#  FOLDER HEADER WIDGET  (collapsible section for organizing cards)
# ══════════════════════════════════════════════════════════════════════════════
class FolderWidget(QWidget):
    """Collapsible folder widget - menggunakan panah bawaan Qt untuk toggle."""

    def __init__(self, folder: dict, mode: str, user_id: int,
                 on_reload, parent):
        super().__init__(parent)
        self.folder = folder
        self.mode = mode
        self.user_id = user_id
        self.on_reload = on_reload
        self.folder_id = folder["id"] 
        self._collapsed = bool(folder.get("collapsed", 0))
        self._item_count = 0
        self._build()
        self.setAcceptDrops(True)

    def dragEnterEvent(self, event):
        if event.mimeData().hasFormat("application/x-craftlife-card"):
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event):
        raw = event.mimeData().data("application/x-craftlife-card")
        info = json.loads(raw.data().decode())
        if info["mode"] == self.mode and info["user_id"] == self.user_id:
            source_folder = info.get("current_folder_id")
            target_folder = self.folder_id
            if source_folder == target_folder:
                # Drop di folder yang sama → tidak ada reorder, biarkan
                event.acceptProposedAction()
            else:
                # Pindahkan item ke folder ini (posisi terakhir)
                db.move_item_to_folder(self.user_id, self.mode, info["item_id"], target_folder)
                self.on_reload()
                event.acceptProposedAction()
        else:
            event.ignore()

    def _build(self):
        self._root = QVBoxLayout(self)
        self._root.setContentsMargins(0, 2, 0, 6)
        self._root.setSpacing(0)

        # Header card — pakai _card() agar konsisten
        hdr = _card()
        hdr.setStyleSheet(
            f"QFrame#card {{"
            f" background:{_T('panel')};"
            f" border:1px solid {_T('primary')};"
            f" border-left:4px solid {_T('light')};"
            f" border-radius:8px; }}"
            f"QFrame#card:hover {{"
            f" border-color:{_T('light')};"
            f" border-left-color:{_T('accent')}; }}"
        )

        row = QHBoxLayout(hdr)
        row.setContentsMargins(10, 10, 10, 10)
        row.setSpacing(8)

        # Tombol toggle dengan panah Qt


        self._toggle_btn = QToolButton()
        self._toggle_btn.setArrowType(Qt.ArrowType.RightArrow if self._collapsed else Qt.ArrowType.DownArrow)
        self._toggle_btn.setMinimumSize(28, 28)
        self._toggle_btn.setStyleSheet("border: none; background: transparent;")
        self._toggle_btn.clicked.connect(self._toggle)
        row.addWidget(self._toggle_btn)

        # Icon folder
        ico_lbl = _emoji_label(self.folder.get("icon", "\U0001f4c1"), ICON_CARD)
        ico_lbl.setMinimumWidth(34)
        row.addWidget(ico_lbl)

        # Nama folder + jumlah item
        info_col = QVBoxLayout()
        info_col.setSpacing(1)
        self._name_lbl = QLabel(self.folder["name"])
        self._name_lbl.setStyleSheet(
            f"color:{_T('text')}; font-weight:bold; font-size:14px;"
            f" background:transparent; border:none;")
        info_col.addWidget(self._name_lbl)
        self._count_lbl = QLabel(tr("folder_item_count", count=0))
        self._count_lbl.setStyleSheet(
            f"color:{_T('muted')}; font-size:11px;"
            f" background:transparent; border:none;")
        info_col.addWidget(self._count_lbl)
        row.addLayout(info_col, 1)

        # Tombol aksi (edit, duplikat, hapus)
        edit_btn = _btn("\u270f\ufe0f", h=36)
        edit_btn.setMinimumWidth(36)
        edit_btn.setToolTip(tr("folder_tooltip_edit"))
        edit_btn.clicked.connect(self._edit_folder)
        row.addWidget(edit_btn)

        dup_btn = _btn("\U0001f4cb", h=36)
        dup_btn.setMinimumWidth(36)
        dup_btn.setToolTip(tr("folder_tooltip_dup"))
        dup_btn.clicked.connect(self._dup_folder)
        row.addWidget(dup_btn)

        del_btn = _btn("\U0001f5d1", "danger", h=36)
        del_btn.setMinimumWidth(36)
        del_btn.setToolTip(tr("folder_tooltip_del"))
        del_btn.clicked.connect(self._del_folder)
        row.addWidget(del_btn)

        self._root.addWidget(hdr)

        # Area konten — dengan indentasi
        self._content = QWidget(self)
        self._content_lay = QVBoxLayout(self._content)
        self._content_lay.setContentsMargins(24, 4, 0, 0)
        self._content_lay.setSpacing(6)
        self._content.setVisible(not self._collapsed)
        self._root.addWidget(self._content)

    def add_card(self, card: QWidget):
        self._content_lay.addWidget(card)
        self._item_count += 1
        self._count_lbl.setText(
            f"{self._item_count} item{'s' if self._item_count > 1 else ''}")

    def _toggle(self):
        self._collapsed = not self._collapsed
        # Ubah panah sesuai status
        self._toggle_btn.setArrowType(Qt.ArrowType.RightArrow if self._collapsed else Qt.ArrowType.DownArrow)
        self._content.setVisible(not self._collapsed)
        # Simpan status ke database
        db.update_task_folder(self.folder["id"], self.user_id,
                              collapsed=int(self._collapsed))

    def _edit_folder(self):
        dlg = FolderDialog(self.mode, self.user_id,
                           existing=self.folder, parent=self)
        if dlg.exec():
            self.on_reload()

    def _dup_folder(self):
        r = db.duplicate_task_folder(self.user_id, self.folder["id"], self.mode)
        SND.notify() if r["ok"] else SND.error()
        self.on_reload()

    def _del_folder(self):
        db.delete_task_folder(self.user_id, self.folder["id"], self.mode)
        SND.click()
        self.on_reload()


# ══════════════════════════════════════════════════════════════════════════════
#  FOLDER DIALOG  (add / edit folder)
# ══════════════════════════════════════════════════════════════════════════════
class FolderDialog(QDialog):
    """Dialog untuk membuat atau mengedit folder."""

    FOLDER_ICONS = [
        (tr("folder_icon_default"), "📁"),
        (tr("folder_icon_favorite"), "⭐"),
        (tr("folder_icon_priority"), "🔥"),
        (tr("folder_icon_exercise"), "💪"),
        (tr("folder_icon_study"), "📖"),
        (tr("folder_icon_goal"), "🎯"),
        (tr("folder_icon_night"), "🌙"),
        (tr("folder_icon_morning"), "☀️"),
        (tr("folder_icon_hobby"), "🎮"),
        (tr("folder_icon_work"), "💼"),
        (tr("folder_icon_home"), "🏠"),
        (tr("folder_icon_health"), "🌿"),
    ]

    def __init__(self, mode: str, user_id: int,
                 existing=None, parent=None):
        super().__init__(parent)
        self.mode     = mode
        self.user_id  = user_id
        self.existing = existing
        title = "✏️ Edit Folder" if existing else tr("folder_create_new")
        self.setWindowTitle(title)
        self.setMinimumWidth(380)
        self.setStyleSheet(build_ss())
        self._build()

    def _build(self):
        lay = QVBoxLayout(self)
        lay.setContentsMargins(24, 20, 24, 20)
        lay.setSpacing(14)

        lay.addWidget(_lbl(self.windowTitle(), "section", 13, True))
        lay.addWidget(_sep())

        lay.addWidget(_lbl(tr("folder_name"), size=12))
        self._name = _input(tr("folder_name_placeholder"))
        if self.existing:
            self._name.setText(self.existing["name"])
        lay.addWidget(self._name)

        lay.addWidget(_lbl(tr("dialog_icon"), size=12))
        self._icon = _combo(self.FOLDER_ICONS)
        if self.existing:
            idx = next((i for i, (_, d) in enumerate(self.FOLDER_ICONS)
                        if d == self.existing.get("icon", "📁")), 0)
            self._icon.setCurrentIndex(idx)
        lay.addWidget(self._icon)

        ok_text = tr("dialog_save") if self.existing else tr("folder_create_btn")
        ok = _btn(ok_text, "solid", self._save, 44)
        lay.addWidget(ok)

    def _save(self):
        name = self._name.text().strip()
        if not name:
            _show(self, tr("msg_error"), tr("folder_name_empty"), "error")
            return
        icon = self._icon.currentData()
        if self.existing:
            db.update_task_folder(self.existing["id"], self.user_id,
                                  name=name, icon=icon)
        else:
            db.add_task_folder(self.user_id, self.mode, name, icon)
        SND.complete()
        self.accept()

# ══════════════════════════════════════════════════════════════════════════════
#  Draggable Card Widget  (untuk drag & drop ke folder lain)
# ══════════════════════════════════════════════════════════════════════════════
class DraggableCard(QFrame):
    """Kartu yang bisa di-drag untuk dipindahkan ke folder lain."""
    def __init__(self, item_id: int, mode: str, user_id: int, current_folder_id: int, content_widget: QWidget, parent=None):
        super().__init__(parent)
        self.item_id = item_id
        self.mode = mode
        self.user_id = user_id
        self.current_folder_id = current_folder_id
        self.setAcceptDrops(True)
        self.setObjectName("card")
        self.setStyleSheet(f"QFrame#card {{ background:{_T('panel')}; border:1px solid {_T('border')}; border-radius:8px; }}")
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(content_widget)
        self.drag_start_pos = None

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.drag_start_pos = event.pos()

    def mouseMoveEvent(self, event):
        if not (event.buttons() & Qt.MouseButton.LeftButton):
            return
        if self.drag_start_pos is None:
            return
        if (event.pos() - self.drag_start_pos).manhattanLength() < QApplication.startDragDistance():
            return
        drag = QDrag(self)
        mime = QMimeData()
        data = {
            "item_id": self.item_id,
            "mode": self.mode,
            "user_id": self.user_id,
            "current_folder_id": self.current_folder_id
        }
        mime.setData("application/x-craftlife-card", json.dumps(data).encode())
        drag.setMimeData(mime)
        drag.exec(Qt.DropAction.MoveAction)

    def dragEnterEvent(self, event):
        if event.mimeData().hasFormat("application/x-craftlife-card"):
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event):
        if not event.mimeData().hasFormat("application/x-craftlife-card"):
            event.ignore()
            return
        raw = event.mimeData().data("application/x-craftlife-card")
        info = json.loads(raw.data().decode())
        if info["user_id"] != self.user_id or info["mode"] != self.mode:
            event.ignore()
            return

        if info["item_id"] == self.item_id:
            event.acceptProposedAction()
            return

        source_folder = info.get("current_folder_id")
        target_folder = self.current_folder_id

        if source_folder == target_folder:
            # Reorder: pindahkan source relatif terhadap target
            r = db.reorder_item_relative(self.user_id, self.mode, info["item_id"], self.item_id)
            if not r.get("ok"):
                print(f"Reorder error: {r.get('msg')}")
        else:
            # Pindahkan ke folder target
            db.move_item_to_folder(self.user_id, self.mode, info["item_id"], target_folder)

        # Cari parent terdekat yang punya method load()
        parent = self.parent()
        while parent and not hasattr(parent, "load"):
            parent = parent.parent()
        if parent and hasattr(parent, "load"):
            parent.load()
        event.acceptProposedAction()

# ══════════════════════════════════════════════════════════════════════════════
#  TASK PAGE  (habits / dailies / todos)  — with sub-tabs + folders
# ══════════════════════════════════════════════════════════════════════════════
class TaskPage(QWidget):
    def __init__(self, user_id: int, mode: str):
        super().__init__()
        self.user_id = user_id
        self.mode    = mode
        self.card_widgets = {}
        self._build()
        AppState.register(self.load)
        AppState.register_lang_cb(self.load)

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(20, 16, 20, 16)
        root.setSpacing(10)

        titles = {"habit": tr("task_title_habit"),
                  "daily": tr("task_title_daily"),
                  "todo":  tr("task_title_todo")}
        hdr = QHBoxLayout()
        hdr.addWidget(_lbl(titles[self.mode], "section", 14, True))
        hdr.addStretch()
        folder_btn = _btn(tr("task_folder_button"), h=36)
        folder_btn.setMinimumWidth(110)
        folder_btn.clicked.connect(self._open_folder_add)
        hdr.addWidget(folder_btn)
        if self.mode == "habit":
            tpl_btn = _btn(tr("template_btn"), h=36)
            tpl_btn.setMinimumWidth(110)
            tpl_btn.clicked.connect(self._open_templates)
            hdr.addWidget(tpl_btn)
        add = _btn(tr("task_add_button"), "solid", self._open_add)
        add.setMinimumWidth(130)
        hdr.addWidget(add)
        root.addLayout(hdr)
        root.addWidget(_sep())

        # Filter bar
        filter_widget = QWidget()
        filter_layout = QHBoxLayout(filter_widget)
        filter_layout.setContentsMargins(0, 0, 0, 0)
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText(tr("task_search"))
        self.search_input.textChanged.connect(self.load)
        self.filter_combo = QComboBox()
        if self.mode == "todo":
            self.filter_combo.addItem(tr("task_filter_all"), "all")
            self.filter_combo.addItem(tr("task_priority_trivial"), "trivial")
            self.filter_combo.addItem(tr("task_priority_easy"), "easy")
            self.filter_combo.addItem(tr("task_priority_medium"), "medium")
            self.filter_combo.addItem(tr("task_priority_hard"), "hard")
        else:
            self.filter_combo.addItem(tr("task_filter_all"), "all")
            self.filter_combo.addItem(tr("task_difficulty_easy"), "easy")
            self.filter_combo.addItem(tr("task_difficulty_medium"), "medium")
            self.filter_combo.addItem(tr("task_difficulty_hard"), "hard")
            self.filter_combo.addItem(tr("task_difficulty_epic"), "epic")
        self.filter_combo.currentIndexChanged.connect(self.load)
        filter_layout.addWidget(self.search_input)
        filter_layout.addWidget(self.filter_combo)
        root.addWidget(filter_widget)

        # Tab widget: Semua + per kategori icon
        self._tabs = QTabWidget()
        self._inner_all = QWidget()
        self._lay_all   = QVBoxLayout(self._inner_all)
        self._lay_all.setSpacing(8)
        self._lay_all.addStretch()
        self._tabs.addTab(_scrolled(self._inner_all), "🗂 " + tr("task_tab_all"))

        self._cat_lays: dict   = {}
        self._cat_inners: dict  = {}
        for key, cat in TASK_ICON_CATEGORIES.items():
            inner = QWidget()
            lay   = QVBoxLayout(inner)
            lay.setSpacing(8)
            lay.addStretch()
            self._cat_lays[key]   = lay
            self._cat_inners[key] = inner
            self._tabs.addTab(_scrolled(inner), f"{cat['icon']} {cat['name']}")

        root.addWidget(self._tabs, 1)
        self.load()

    def _clear_lay(self, lay: QVBoxLayout):
        while lay.count() > 1:
            it = lay.takeAt(0)
            if it.widget():
                it.widget().deleteLater()

    def _get_items(self):
        if self.mode == "habit":
            return db.get_habits(self.user_id)
        elif self.mode == "daily":
            return db.get_dailies(self.user_id)
        else:
            return db.get_todos(self.user_id)

    def load(self):
        _sa = self._tabs.currentWidget()
        try:
            self._saved_scroll = _sa.verticalScrollBar().value() if isinstance(_sa, QScrollArea) else 0
        except Exception:
            self._saved_scroll = 0
        self.card_widgets.clear()
        if not AppState.user_id:
            return
        db.reset_daily_tasks(self.user_id)

        self._clear_lay(self._lay_all)
        for lay in self._cat_lays.values():
            self._clear_lay(lay)

        all_items = self._get_items()
        if self.mode == "todo":
            all_items.sort(key=lambda x: (x.get('done', 0), -x.get('sort_order', 0)))
        else:
            all_items.sort(key=lambda x: (x.get('done_today', 0), -x.get('sort_order', 0)))
        search     = self.search_input.text().lower()
        filter_diff = self.filter_combo.currentData()
        if search:
            all_items = [i for i in all_items if search in i["name"].lower()]
        if filter_diff != "all":
            all_items = [i for i in all_items
                         if i.get("difficulty", i.get("priority", "")).lower() == filter_diff]

        folders = db.get_task_folders(self.user_id, self.mode)

        self._render_to_layout(self._lay_all, self._inner_all, all_items, folders)
        by_cat: dict = {}
        for item in all_items:
            icon_key = item.get("icon", "").replace("\ufe0f", "").strip()
            matched = None
            for k in self._cat_lays:
                if k.replace("\ufe0f", "").strip() == icon_key:
                    matched = k
                    break
            if matched:
                by_cat.setdefault(matched, []).append(item)

        for key, lay in self._cat_lays.items():
            cat_items = by_cat.get(key, [])
            self._render_to_layout(lay, self._cat_inners[key], cat_items, folders)
        try:
            _sa2 = self._tabs.currentWidget()
            if isinstance(_sa2, QScrollArea):
                _val = getattr(self, "_saved_scroll", 0)
                QTimer.singleShot(0, lambda v=_val: _sa2.verticalScrollBar().setValue(v))
        except Exception:
            pass

    # ========== RENDER WITH DROP AREA ==========


    def _render_to_layout(self, lay: QVBoxLayout, container: QWidget, items: list, folders: list):
        if not items and not folders:
            e_icon = {"habit": "⛏️", "daily": "📅", "todo": "📜"}[self.mode]
            e_msg  = {"habit": tr("task_empty_habit"),
                      "daily": tr("task_empty_daily"),
                      "todo": tr("task_empty_todo")}[self.mode]
            el = _lbl(f"{e_icon}  {e_msg}", "sub", 13)
            el.setAlignment(Qt.AlignmentFlag.AlignCenter)
            el.setStyleSheet(f"color: {_T('muted')}; padding: 40px;")
            lay.insertWidget(0, el)
            return

        insert_pos = 0
        folder_items: dict = {f["id"]: [] for f in folders}
        ungrouped: list = []
        for item in items:
            fid = item.get("folder_id")
            if fid and fid in folder_items:
                folder_items[fid].append(item)
            else:
                ungrouped.append(item)

        for folder in folders:
            fw = FolderWidget(folder, self.mode, self.user_id, self.load, parent=container)
            cards_in_folder = folder_items.get(folder["id"], [])
            if not cards_in_folder:
                empty_lbl = QLabel("   📭  " + tr("folder_empty"))
                empty_lbl.setStyleSheet(f"color:{_T('muted')}; font-size:12px; padding:6px 0;")
                fw.add_card(empty_lbl)
            else:
                for item in cards_in_folder:
                    fw.add_card(self._create_card_widget(item))
            lay.insertWidget(insert_pos, fw)
            insert_pos += 1

        for item in ungrouped:
            lay.insertWidget(insert_pos, self._create_card_widget(item))
            insert_pos += 1

        # ── Drop area untuk "Tanpa Folder" ──
        drop_area = QFrame()
        drop_area.setAcceptDrops(True)
        drop_area.setMinimumHeight(50)
        drop_area.setStyleSheet(f"""
            QFrame {{
                border: 2px dashed {_T('border')};
                border-radius: 8px;
                background: {_T('bg')};
                margin-top: 8px;
            }}
            QFrame:hover {{
                border-color: {_T('accent')};
                background: {_T('panel')};
            }}
        """)
        drop_label = QLabel(tr("drop_here_to_remove_folder") if hasattr(tr, "__call__") else "📂 Taruh di sini untuk keluarkan dari folder")
        drop_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        drop_label.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        drop_area_layout = QVBoxLayout(drop_area)
        drop_area_layout.addWidget(drop_label)

        def _drag_enter(e):
            if e.mimeData().hasFormat("application/x-craftlife-card"):
                e.acceptProposedAction()
            else:
                e.ignore()
        drop_area.dragEnterEvent = _drag_enter

        def _drop(e):
            raw = e.mimeData().data("application/x-craftlife-card")
            info = json.loads(raw.data().decode())
            if info["mode"] == self.mode and info["user_id"] == self.user_id:
                db.set_item_folder(self.user_id, self.mode, info["item_id"], None)
                self.load()
                e.acceptProposedAction()
            else:
                e.ignore()
        drop_area.dropEvent = _drop

        lay.insertWidget(insert_pos, drop_area)

    def _refresh_topbar(self):
        """Refresh HANYA top bar (XP/gold/HP) tanpa reload page -> scroll tetap di tempat."""
        try:
            self.window()._topbar.refresh()
        except Exception:
            try: AppState.refresh()
            except Exception: pass

    def _update_item(self, item_id: int):
        """Update hanya satu kartu tanpa reload seluruh halaman"""
        items = self._get_items()
        item = next((i for i in items if i['id'] == item_id), None)
        if not item:
            self.load()
            return

        card = self.card_widgets.get(item_id)
        if not card:
            self.load()
            return

        # Hapus konten lama
        while card.layout().count():
            child = card.layout().takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Buat konten baru
        new_content = self._build_card_content(item)
        card.layout().addWidget(new_content)
        card.current_folder_id = item.get('folder_id')   # update untuk drag

        # Update top bar (hanya topbar, bukan reload page)
        self._refresh_topbar()

    # ========== MAKE CARD (dengan DraggableCard) ==========
    def _build_card_content(self, item: dict) -> QWidget:
        """Buat widget konten kartu (tanpa wrapper DraggableCard)"""
        done = bool(item.get("done_today") or item.get("done", False))

        content = QWidget()
        row = QHBoxLayout(content)
        row.setContentsMargins(14, 10, 14, 10)
        row.setSpacing(10)

        # Icon
        ico = _emoji_label(item["icon"], ICON_CARD)
        row.addWidget(ico)

        # Info block
        info = QVBoxLayout()
        info.setSpacing(2)

        name_style = (f"text-decoration:line-through; color:{_T('muted')};"
                    if done else f"color:{_T('text')}; font-weight:bold;")
        nm = QLabel(item["name"])
        nm.setStyleSheet(f"font-size:14px; {name_style}")
        info.addWidget(nm)

        if self.mode == "todo":
            val = item.get("priority", "medium")
            color_map = {"trivial": _T("muted"), "easy": _T("light"),
                        "medium": "#f0a800", "hard": "#e05050"}
            label_text = tr(f"task_priority_{val}")
        else:
            val = item.get("difficulty", "medium")
            color_map = {"easy": "#80c000", "medium": "#f0a800",
                        "hard": "#e05050", "epic": "#a97fff"}
            label_text = tr(f"task_difficulty_{val}")
        diff_lbl = QLabel(label_text)
        diff_lbl.setStyleSheet(f"color: {color_map.get(val, '#888')}; font-size: 10px; font-weight: bold;")
        info.addWidget(diff_lbl)

        streak_txt = (tr("task_streak_days", streak=item['streak'])
                    if item.get("streak", 0) > 0 else "")
        sub = QLabel(tr("task_reward_format", xp=item['xp_reward'], gold=item['gold_reward'], streak=streak_txt))
        sub.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        info.addWidget(sub)

        fail_streak = item.get("fail_streak", 0)
        if fail_streak > 0:
            fail_lbl = QLabel(tr("task_fail_streak_days", streak=fail_streak))
            fail_lbl.setStyleSheet(f"color:#e05050; font-size:10px; font-weight:bold;")
            info.addWidget(fail_lbl)

        if item.get("notes"):
            notes_lbl = QLabel(f"📝 {item['notes']}")
            notes_lbl.setWordWrap(True)
            notes_lbl.setStyleSheet(f"color:{_T('muted')}; font-size:11px; font-style:italic;")
            info.addWidget(notes_lbl)

        # History
        history = db.get_task_last_history(self.user_id, self.mode, item["id"])
        if history:
            action = history["action"]
            date_str = history["action_date"]
            emoji = "✅" if action == "success" else "❌" if action == "fail" else "⏭️"
            color = "#80c000" if action == "success" else "#e05050" if action == "fail" else "#f0a800"
            status_label = QLabel(tr("task_status_history", emoji=emoji, date=date_str))
            status_label.setStyleSheet(f"color:{color}; font-size:10px;")
            info.addWidget(status_label)
        else:
            status_label = QLabel(tr("task_status_never"))
            status_label.setStyleSheet(f"color:{_T('muted')}; font-size:10px;")
            info.addWidget(status_label)

        row.addLayout(info, 1)

        # Tombol aksi (sama seperti sebelumnya, tapi sekarang tidak dibungkus DraggableCard)
        if self.mode == "habit":
            ck = _btn(tr("task_btn_done") if done else tr("task_btn_check"), h=36)
            ck.setEnabled(not done)
            if done:
                ck.setStyleSheet(f"background:{_T('border')}; color:{_T('muted')}; border-color:{_T('border')};")
            ck.setMinimumWidth(92)
            ck.clicked.connect(lambda _, i=item["id"]: self._do("up", i))
            row.addWidget(ck)

            edit_btn = _btn("✏️", h=36)
            edit_btn.setMinimumWidth(36)
            edit_btn.clicked.connect(lambda _, i=item["id"]: self._edit(i))
            row.addWidget(edit_btn)

        elif self.mode == "daily":
            ck = _btn(tr("task_btn_done") if done else tr("task_btn_check"), h=36)
            ck.setEnabled(not done)
            if done:
                ck.setStyleSheet(f"background:{_T('border')}; color:{_T('muted')}; border-color:{_T('border')};")
            ck.setMinimumWidth(92)
            ck.clicked.connect(lambda _, i=item["id"]: self._do_daily(i))
            row.addWidget(ck)

            # ── Tombol Freeze ──
            freeze_btn = _btn("🧊 Freeze", h=36)
            freeze_btn.setMinimumWidth(70)
            freeze_btn.clicked.connect(lambda _, i=item["id"]: self._add_freeze(i))
            row.addWidget(freeze_btn)

            nb = _btn(tr("task_btn_fail"), "danger", h=36)
            nb.setEnabled(not done)
            if done:
                nb.setStyleSheet(f"background:{_T('border')}; color:{_T('muted')}; border-color:{_T('border')};")
            nb.setMinimumWidth(82)
            nb.clicked.connect(lambda _, i=item["id"]: self._fail_daily(i))
            row.addWidget(nb)

            # Tampilkan freeze slots
            freeze_slots = item.get("freeze_slots", 0)
            freeze_label = QLabel(f"🧊 Freeze: {freeze_slots}/3")
            freeze_label.setStyleSheet(f"color:{_T('accent')}; font-size:10px; font-weight:bold;")
            info.addWidget(freeze_label)

            edit_btn = _btn("✏️", h=36)
            edit_btn.setMinimumWidth(36)
            edit_btn.clicked.connect(lambda _, i=item["id"]: self._edit(i))
            row.addWidget(edit_btn)

        else:  # todo
            cb = QCheckBox()
            cb.setChecked(done)
            cb.setEnabled(not done)
            cb.stateChanged.connect(lambda _, i=item["id"]: self._do_todo(i))
            row.addWidget(cb)

            edit_btn = _btn("✏️", h=36)
            edit_btn.setMinimumWidth(36)
            edit_btn.clicked.connect(lambda _, i=item["id"]: self._edit(i))
            row.addWidget(edit_btn)

        # Tombol hapus, duplikat, folder
        dl = _btn(tr("task_btn_delete"), "danger", h=36)
        dl.setMinimumWidth(36)
        dl.clicked.connect(lambda _, i=item["id"]: self._delete(i))
        row.addWidget(dl)

        dup_btn = _btn(tr("task_btn_duplicate"), h=36)
        dup_btn.setMinimumWidth(36)
        dup_btn.clicked.connect(lambda _, i=item["id"]: self._duplicate(i))
        row.addWidget(dup_btn)

        folder_btn = _btn(tr("task_btn_folder"), h=36)
        folder_btn.setMinimumWidth(36)
        folder_btn.clicked.connect(lambda _, iid=item["id"], fid=item.get("folder_id"):
                                self._move_to_folder(iid, fid))
        row.addWidget(folder_btn)

        up_btn = _btn("⬆", h=36)
        up_btn.setMinimumWidth(36)
        up_btn.clicked.connect(lambda _, i=item["id"]: self._reorder(i, "up"))
        row.addWidget(up_btn)

        down_btn = _btn("⬇", h=36)
        down_btn.setMinimumWidth(36)
        down_btn.clicked.connect(lambda _, i=item["id"]: self._reorder(i, "down"))
        row.addWidget(down_btn)

        return content

    def _create_card_widget(self, item: dict) -> DraggableCard:
        """Bungkus konten dengan DraggableCard dan simpan referensi"""
        content = self._build_card_content(item)
        card = DraggableCard(
            item_id=item["id"],
            mode=self.mode,
            user_id=self.user_id,
            current_folder_id=item.get("folder_id"),
            content_widget=content,
            parent=self
        )
        self.card_widgets[item["id"]] = card
        return card

    # ========== AKSI LAINNYA (sama seperti sebelumnya) ==========
    def _move_to_folder(self, item_id: int, current_folder_id: int = None):
        folders = db.get_task_folders(self.user_id, self.mode)
        folder_names = [tr("dialog_no_folder")] + [f["name"] for f in folders]
        folder_ids = [None] + [f["id"] for f in folders]

        dlg = QDialog(self)
        dlg.setWindowTitle(tr("dialog_move_folder"))
        dlg.setMinimumSize(300, 200)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        combo = QComboBox()
        for name in folder_names:
            combo.addItem(name)
        if current_folder_id is not None:
            try:
                idx = folder_ids.index(current_folder_id)
                combo.setCurrentIndex(idx)
            except ValueError:
                pass
        btn_ok = _btn(tr("dialog_move"), "solid", dlg.accept)
        layout.addWidget(QLabel(tr("dialog_select_folder")))
        layout.addWidget(combo)
        layout.addWidget(btn_ok)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            selected_id = folder_ids[combo.currentIndex()]
            db.set_item_folder(self.user_id, self.mode, item_id, selected_id)
            QTimer.singleShot(0, self.load)
            SND.notify()

    def _do(self, direction: str, iid: int):
        r = db.complete_habit(self.user_id, iid, direction)
        if not r.get("ok", True) and r.get("msg"):
            _show(self, tr("msg_info"), r["msg"])
            return
        if direction == "up":
            SND.complete()
            msg = tr("habit_complete_msg", xp=r.get('xp_gained',0), gold=r.get('gold_gained',0))
            if r.get("leveled_up"):
                SND.level_up()
                celebrate_levelup(self)
                msg += f"\n🎉 {tr('level_up_msg', lvl=r['new_level'])}"
            _show(self, tr("task_habit_success"), msg, "success")
        else:
            SND.error()
            lost_hp = r.get("lost_hp", 5)
            _show(self, tr("task_hp_loss"), f"💔 -{lost_hp:.0f} HP " + tr("task_bad_habit"), "warning")
        self._refresh_topbar()
        self._update_item(iid)

    def _do_daily(self, iid: int):
        # Recurrence: daily di luar jadwal hari ini tidak bisa diselesaikan
        item = next((d for d in db.get_dailies(self.user_id) if d["id"] == iid), None)
        if item and not db.is_due_today(item.get("repeat_days", "")):
            _show(self, tr("info_title"), tr("recur_not_today"), "info")
            return
        r = db.complete_daily(self.user_id, iid)
        if not r.get("ok", True) and r.get("msg"):
            _show(self, tr("msg_info"), r["msg"])
            return
        SND.complete()
        msg = tr("daily_complete_msg", xp=r.get('xp_gained',0))
        if r.get("leveled_up"):
            SND.level_up()
            celebrate_levelup(self)
            msg += f"\n🎉 {tr('level_up_msg', lvl=r['new_level'])}"
        _show(self, tr("task_daily_success"), msg, "success")
        self._refresh_topbar()
        self._update_item(iid)

    def _fail_daily(self, iid: int):
        r = db.fail_daily(self.user_id, iid)
        if not r.get("ok", True) and r.get("msg"):
            _show(self, tr("msg_info"), r["msg"])
            return

        if r.get("freeze_used"):
            SND.notify()
            _show(self, tr("task_daily_freeze_title"), r["msg"], "success")
            self._refresh_topbar()
            self._update_item(iid)
            return

        penalty_type = r.get("penalty_type", "hp")
        penalty_amount = r.get("penalty_amount", 0)
        lost_hp = r.get("lost_hp", 0)

        if penalty_type == "hp":
            msg = tr("task_daily_fail", hp=lost_hp)
        elif penalty_type == "gold":
            msg = tr("penalty_gold_popup", gold=penalty_amount)
        elif penalty_type == "xp":
            msg = tr("penalty_xp_popup_percent", xp=penalty_amount, new_level=r.get("new_level", 1))
        else:
            msg = tr("penalty_none_popup")

        SND.error()
        _show(self, tr("task_daily_fail_title"), msg, "warning")
        self._refresh_topbar()
        self._update_item(iid)

    def _add_freeze(self, daily_id):
        loading = LoadingDialog(tr("proces_freeze"), self)
        loading.show()
        QApplication.processEvents()
        r = db.add_freeze_to_daily(self.user_id, daily_id)
        loading.accept()
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        self._refresh_topbar()
        self._update_item(daily_id)

    def _do_todo(self, iid: int):
        r = db.complete_todo(self.user_id, iid)
        if not r.get("ok"):
            return
        SND.complete()
        msg = tr("quest_complete_msg", text=tr('task_quest_complete'), xp=r.get('xp_gained',0))
        if r.get("leveled_up"):
            SND.level_up()
            celebrate_levelup(self)
            msg += f"\n🎉 {tr('level_up_msg', lvl=r['new_level'])}"
        _show(self, tr("task_quest_success"), msg, "success")
        self._refresh_topbar()
        self._update_item(iid)

    def _open_templates(self):
        dlg = HabitTemplateDialog(self.user_id, self)
        if dlg.exec():
            self.load()

    def _delete(self, iid: int):
        fns = {"habit": db.delete_habit,
               "daily": db.delete_daily,
               "todo":  db.delete_todo}
        getter = {"habit": db.get_habits,
                  "daily": db.get_dailies,
                  "todo":  db.get_todos}[self.mode]
        name = next((it["name"] for it in getter(self.user_id)
                     if it["id"] == iid), "?")
        r = fns[self.mode](self.user_id, iid)
        SND.click()
        self.load()
        # Tawarkan undo via toast (item sudah masuk trash di DB)
        tid = r.get("trash_id") if isinstance(r, dict) else None
        mw = self.window()
        if tid and hasattr(mw, "show_undo_toast"):
            mw.show_undo_toast(name, tid)

    def _duplicate(self, iid):
        if self.mode == "habit":
            db.duplicate_habit(self.user_id, iid)
        elif self.mode == "daily":
            db.duplicate_daily(self.user_id, iid)
        else:
            db.duplicate_todo(self.user_id, iid)
        SND.click()
        self.load()

    def _edit(self, iid):
        items_map = {
            "habit": db.get_habits,
            "daily": db.get_dailies,
            "todo":  db.get_todos,
        }
        all_items = items_map[self.mode](self.user_id)
        item = next((i for i in all_items if i["id"] == iid), None)
        if not item:
            return
        dlg = EditTaskDialog(self.mode, self.user_id, item, self)
        if dlg.exec():
            self.load()

    def _open_add(self):
        dlg = AddTaskDialog(self.mode, self.user_id, self)
        if dlg.exec():
            self.load()

    def _open_folder_add(self):
        dlg = FolderDialog(self.mode, self.user_id, parent=self)
        if dlg.exec():
            self.load()

    def _reorder(self, item_id, direction):
        # Tentukan status berdasarkan mode
        status = None
        items = self._get_items()
        item = next((i for i in items if i["id"] == item_id), None)
        if item:
            if self.mode == "todo":
                status = item.get("done", 0)
            elif self.mode in ("habit", "daily", "sport"):
                status = item.get("done_today", 0)

        # Balik arah untuk mode yang menggunakan sort_order DESC di tampilan
        actual_direction = direction
        if self.mode in ("todo", "habit", "daily", "sport"):
            actual_direction = "down" if direction == "up" else "up"

        r = db.reorder_item(self.user_id, self.mode, item_id, actual_direction, status)
        if r.get("ok"):
            if r.get("moved", False):
                SND.click()
                QTimer.singleShot(0, self.load)
            # else tidak ada perubahan, tidak perlu reload
        else:
            SND.error()
            _show(self, tr("msg_error"), r.get("msg", "Gagal mengubah urutan"), "error")

    def _toggle_left_panel(self, collapsed):
        # Cari left widget (index 0 di splitter)
        try:
            left = self.findChild(QWidget, "learning_left")
            if not left:
                # Fallback: ambil widget pertama di splitter
                left = self.sender().parent().findChild(QSplitter).widget(0) if hasattr(self, 'sender') else None
            # Simpan sizes
            if not hasattr(self, '_splitter'):
                # Cari splitter
                for child in self.findChildren(QSplitter):
                    self._splitter = child
                    break
            if hasattr(self, '_splitter'):
                sizes = self._splitter.sizes()
                if collapsed:
                    self._splitter.setSizes([0, 600, 320])
                    self.btn_collapse_left.setText("Sources ▶")
                else:
                    self._splitter.setSizes([280, 400, 320])
                    self.btn_collapse_left.setText("◀ Sources")
        except Exception as e:
            print(f"Toggle left failed: {e}")

    def _toggle_right_panel(self, collapsed):
        try:
            if not hasattr(self, '_splitter'):
                for child in self.findChildren(QSplitter):
                    self._splitter = child
                    break
            if hasattr(self, '_splitter'):
                if collapsed:
                    self._splitter.setSizes([280, 400, 0])
                    self.btn_collapse_right.setText("◀ Studio")
                else:
                    self._splitter.setSizes([280, 400, 320])
                    self.btn_collapse_right.setText("Studio ▶")
        except Exception as e:
            print(f"Toggle right failed: {e}")

    def closeEvent(self, e):
        AppState.unregister(self.load)
        AppState.unregister_lang_cb(self.load)
        super().closeEvent(e)


# ══════════════════════════════════════════════════════════════════════════════
#  SPORT TRACK PAGE  — sub-halaman per jenis olahraga
# ══════════════════════════════════════════════════════════════════════════════
class AddSportActivityDialog(QDialog):
    """Dialog untuk tambah/edit aktivitas sport."""
    def __init__(self, user_id: int, item=None, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.item    = item          # None = add, dict = edit
        self.setWindowTitle(
            tr("sport_edit_title") if item else tr("sport_add_title"))
        self.setMinimumWidth(480)
        self.setMinimumHeight(480)
        self.setMaximumHeight(700)
        self.setStyleSheet(build_ss())
        self._build()

    # MET values for professional calorie calculation
    MET_VALUES = {
        "running": 9.8, "gym": 6.0, "cycling": 7.5, "swimming": 8.0,
        "yoga": 3.0, "football": 7.0, "calisthenics": 5.0,
        "martial_arts": 8.0, "badminton": 5.5, "other": 4.0
    }
    INTENSITY_FACTOR = {"easy": 0.85, "medium": 1.0, "hard": 1.25, "epic": 1.5}

    def _calc_calories(self):
        try:
            sport = self._sport_type.currentData()
            met = self.MET_VALUES.get(sport, 4.0)
            weight = self.weight_input.value()
            duration = self.duration.value()
            intensity = self._diff.currentData()
            factor = self.INTENSITY_FACTOR.get(intensity, 1.0)
            # Formula: MET * weight(kg) * hours * factor
            hours = duration / 60.0
            cals = met * weight * hours * factor
            return int(round(cals))
        except:
            return 0

    def _update_calc_preview(self, *_):
        if not hasattr(self, 'calc_preview'):
            return
        if self.auto_calc.isChecked():
            cals = self._calc_calories()
            self.calories_burned.setValue(cals)
            sport = self._sport_type.currentData()
            met = self.MET_VALUES.get(sport, 4.0)
            intensity = self._diff.currentData()
            factor = self.INTENSITY_FACTOR.get(intensity, 1.0)
            per_min = cals / max(1, self.duration.value())
            self.calc_preview.setText(f"🔥 {cals} kcal  •  {per_min:.1f} kcal/min  •  MET {met} × {factor:.2f}")
            # Update info colors
            self.calories_burned.setEnabled(False)
            self.calories_burned.setStyleSheet(f"background: {_T('panel')}; color: {_T('accent')}; font-weight:bold;")
        else:
            self.calc_preview.setText(f"✏️ Manual: {self.calories_burned.value()} kcal")
            self.calories_burned.setEnabled(True)
            self.calories_burned.setStyleSheet("")

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)

        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(12)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))
        lay.addWidget(_sep())

        # Nama
        lay.addWidget(_lbl(tr("sport_activity_name"), size=12))
        self._name = _input(tr("sport_activity_ph"))
        if self.item:
            self._name.setText(self.item["name"])
        lay.addWidget(self._name)

        # Jenis olahraga
        lay.addWidget(_lbl(tr("sport_type_label"), size=12))
        sport_opts = [
            (f"{v['icon']}  {tr('sport_type_' + k)}", k)
            for k, v in db.SPORT_TYPES.items()
        ]   
        self._sport_type = _combo(sport_opts)
        if self.item:
            idx = self._sport_type.findData(self.item.get("sport_type","running"))
            if idx >= 0: self._sport_type.setCurrentIndex(idx)
        lay.addWidget(self._sport_type)

        # Intensitas - with factor hint
        lay.addWidget(_lbl(tr("sport_intensity") + "  ( memengaruhi kalori )", size=12))
        diff_opts = [
            (tr("sport_intensity_easy") + " ×0.85", "easy"),
            (tr("sport_intensity_medium") + " ×1.00", "medium"),
            (tr("sport_intensity_hard") + " ×1.25", "hard"),
            (tr("sport_intensity_epic") + " ×1.50", "epic"),
        ]
        self._diff = _combo(diff_opts)
        if self.item:
            idx = self._diff.findData(self.item.get("difficulty","medium"))
            if idx >= 0: self._diff.setCurrentIndex(idx)
        else:
            self._diff.setCurrentIndex(1)   # medium default
        lay.addWidget(self._diff)

        # === Professional Calorie Section ===
        cal_group = QGroupBox("🔥 Kalkulator Kalori Profesional")
        cal_group.setStyleSheet(f"QGroupBox {{ color: {_T('accent')}; font-weight:bold; }}")
        cal_lay = QVBoxLayout(cal_group)
        cal_lay.setSpacing(8)

        # Weight + Duration row
        row1 = QHBoxLayout()
        w_col = QVBoxLayout()
        w_col.addWidget(_lbl("Berat Badan (kg)", size=11))
        self.weight_input = QDoubleSpinBox()
        self.weight_input.setRange(30.0, 200.0)
        self.weight_input.setValue(65.0)
        self.weight_input.setSuffix(" kg")
        self.weight_input.setDecimals(1)
        # Try to fetch last weight from health logs
        try:
            import database as _db
            u = _db.get_user(self.user_id)
            # try health_logs weight
            conn = _db.get_conn()
            rw = conn.execute("SELECT weight_kg FROM health_logs WHERE user_id=? AND weight_kg IS NOT NULL ORDER BY log_date DESC LIMIT 1", (self.user_id,)).fetchone()
            conn.close()
            if rw and rw["weight_kg"]:
                self.weight_input.setValue(float(rw["weight_kg"]))
        except:
            pass
        if self.item and self.item.get("weight_kg"):
            try:
                self.weight_input.setValue(float(self.item.get("weight_kg")))
            except:
                pass
        w_col.addWidget(self.weight_input)
        row1.addLayout(w_col)

        d_col = QVBoxLayout()
        d_col.addWidget(_lbl(tr("sport_duration"), size=11))
        dur_row = QHBoxLayout()
        self.duration = QSpinBox()
        self.duration.setRange(1, 480)
        self.duration.setSuffix(tr("unit_minutes"))
        self.duration.setValue(self.item.get("duration_minutes", 30) if self.item else 30)
        dur_row.addWidget(self.duration)
        # Quick duration buttons
        for mins in (15, 30, 60):
            b = _btn(f"{mins}m", h=32)
            b.setMinimumWidth(45)
            b.clicked.connect(lambda _, v=mins: self.duration.setValue(v))
            dur_row.addWidget(b)
        d_col.addLayout(dur_row)
        row1.addLayout(d_col)
        cal_lay.addLayout(row1)

        # Duration slider
        self.duration_slider = QSlider(Qt.Orientation.Horizontal)
        self.duration_slider.setRange(5, 180)
        self.duration_slider.setValue(self.duration.value())
        self.duration_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.duration_slider.setTickInterval(15)
        cal_lay.addWidget(self.duration_slider)
        # Sync slider <-> spinbox
        self.duration.valueChanged.connect(lambda v: self.duration_slider.setValue(v))
        self.duration_slider.valueChanged.connect(lambda v: self.duration.setValue(v))

        # Auto calc checkbox + preview
        self.auto_calc = QCheckBox("✨ Hitung otomatis (MET × berat × durasi × intensitas)")
        self.auto_calc.setChecked(True)
        self.auto_calc.setStyleSheet(f"color: {_T('text')}; font-size:11px;")
        cal_lay.addWidget(self.auto_calc)

        self.calc_preview = QLabel("🔥 -- kcal")
        self.calc_preview.setStyleSheet(f"color: {_T('accent')}; font-size:12px; font-weight:bold; background: {_T('panel')}; padding:8px; border-radius:8px;")
        self.calc_preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        cal_lay.addWidget(self.calc_preview)

        # Kalori burned - now auto-linked
        lay_cal = QHBoxLayout()
        lay_cal.addWidget(_lbl(tr("sport_calories_label"), size=12))
        lay_cal.addStretch()
        self.calories_burned = QSpinBox()
        self.calories_burned.setRange(0, 5000)
        self.calories_burned.setSuffix(tr("unit_kcal"))
        self.calories_burned.setValue(self.item.get("calories_burned", 100) if self.item else 100)
        self.calories_burned.setMinimumHeight(36)
        self.calories_burned.setStyleSheet(f"font-size:14px; font-weight:bold;")
        lay_cal.addWidget(self.calories_burned)
        cal_lay.addLayout(lay_cal)

        # Connect auto updates
        self._sport_type.currentIndexChanged.connect(self._update_calc_preview)
        self._diff.currentIndexChanged.connect(self._update_calc_preview)
        self.weight_input.valueChanged.connect(self._update_calc_preview)
        self.duration.valueChanged.connect(self._update_calc_preview)
        self.auto_calc.toggled.connect(self._update_calc_preview)
        self.calories_burned.valueChanged.connect(lambda v: self.calc_preview.setText(f"✏️ Manual: {v} kcal") if not self.auto_calc.isChecked() else None)

        lay.addWidget(cal_group)
        # Initial calc
        QTimer.singleShot(0, self._update_calc_preview)

        # Catatan
        lay.addWidget(_lbl(tr("dialog_notes"), size=12))
        self._notes = _input(tr("dialog_notes_placeholder"))
        if self.item:
            self._notes.setText(self.item.get("notes", ""))
        lay.addWidget(self._notes)

        # Folder
        lay.addWidget(_lbl(tr("dialog_folder"), size=12))
        _folders = db.get_task_folders(self.user_id, "sport")
        _fopts = [(tr("dialog_no_folder"), None)] + [
            (f"{fd['icon']}  {fd['name']}", fd["id"]) for fd in _folders]
        self._folder = _combo(_fopts)
        if self.item:
            cur_fid = self.item.get("folder_id")
            if cur_fid:
                idx = next((i for i,(_, d) in enumerate(_fopts) if d == cur_fid), 0)
                self._folder.setCurrentIndex(idx)
        lay.addWidget(self._folder)

        lay.addSpacing(6)
        # Info footer
        info_lbl = QLabel("💡 Rumus: Kalori = MET × Berat(kg) × Durasi(jam) × Faktor Intensitas")
        info_lbl.setStyleSheet(f"color: {_T('muted')}; font-size:10px; font-style:italic;")
        info_lbl.setWordWrap(True)
        lay.addWidget(info_lbl)

        ok_lbl = tr("dialog_save") if self.item else tr("dialog_add")
        ok = _btn(ok_lbl, "solid", self._save, 46)
        lay.addWidget(ok)
        self._name.returnPressed.connect(self._save)

        root.addWidget(_scrolled(content))

    def _save(self):
        name = self._name.text().strip()
        if not name:
            _show(self, tr("msg_error"), tr("msg_name_empty"), "error")
            return
        sport_type = self._sport_type.currentData()
        diff = self._diff.currentData()
        notes = self._notes.text()
        icon = db.SPORT_TYPES.get(sport_type, {}).get("icon", "🏅")
        calories_burned = self.calories_burned.value()
        duration = self.duration.value()
        folder_id = self._folder.currentData()
        
        if self.item:
            db.update_sport_activity(
                self.item["id"], self.user_id,
                name=name, sport_type=sport_type, icon=icon,
                difficulty=diff, notes=notes,
                calories_burned=calories_burned, duration_minutes=duration
            )
            db.set_item_folder(self.user_id, "sport", self.item["id"], folder_id)
        else:
            db.add_sport_activity(
                self.user_id, name, sport_type, icon, diff, notes,
                calories_burned=calories_burned, duration_minutes=duration
            )
            if folder_id:
                new_id = max(a["id"] for a in db.get_sport_activities(self.user_id))
                db.set_item_folder(self.user_id, "sport", new_id, folder_id)
        SND.complete()
        self.accept()


# Warna tier rank reps (kunci = key di db.SPORT_REP_RANKS)
SPORT_RANK_COLORS = {
    "rookie": "#9aa0a6", "bronze": "#cd7f32", "silver": "#c0c0c0",
    "gold": "#f0c040", "platinum": "#4dd9e0", "diamond": "#6fb7ff",
    "master": "#a97fff", "mythic": "#ff6fd8",
}


class SportRepsChartWidget(QWidget):
    """Bar chart reps harian (7 hari, zero-fill) untuk header Sport page."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self._lay = QVBoxLayout(self)
        self._lay.setContentsMargins(0, 0, 0, 0)
        self._lay.setSpacing(4)
        self.setMinimumHeight(170)

    def _clear(self):
        while self._lay.count():
            it = self._lay.takeAt(0)
            if it.widget():
                it.widget().deleteLater()

    def refresh(self, series: list):
        """series: [{date, reps}] dari db.get_sport_rep_series."""
        self._clear()
        has_data = any(s["reps"] for s in series)
        if not MPL_QT_OK or not has_data:
            lbl = _lbl(tr("sport_reps_chart_title") + " — " +
                       tr("economy_daily_chart_empty"), "sub", 10)
            lbl.setWordWrap(True)
            lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lbl.setStyleSheet(f"color: {_T('muted')}; padding: 18px;")
            self._lay.addWidget(lbl)
            return
        xs = list(range(len(series)))
        reps = [s["reps"] for s in series]
        dates = [s["date"] for s in series]
        muted = _T("muted")

        fig = Figure(figsize=(6.4, 1.7), tight_layout=True)
        fig.patch.set_alpha(0.0)
        ax = fig.add_subplot(111)
        ax.set_facecolor("none")
        ax.bar(xs, reps, width=0.55, color="#4da6ff", alpha=0.85)
        ax.set_xticks(xs)
        ax.set_xticklabels([d[5:] for d in dates], fontsize=7, color=muted)
        ax.tick_params(colors=muted, labelsize=7)
        for spine in ax.spines.values():
            spine.set_visible(False)
        canvas = FigureCanvas(fig)
        canvas.setStyleSheet("background: transparent;")
        self._lay.addWidget(canvas)


class LogSportRepsDialog(QDialog):
    """Catat sesi reps (set × reps) untuk satu aktivitas — push-up, pull-up, dll.
    Menampilkan rank saat ini, riwayat sesi terakhir, dan merayakan kenaikan rank."""
    def __init__(self, user_id: int, activity: dict, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        self.activity = activity
        self.setWindowTitle(tr("sport_log_reps_title", name=activity["name"]))
        self.setMinimumWidth(460)
        self.setStyleSheet(build_ss())
        self._build()

    def _total(self):
        return self._sets.value() * self._reps.value()

    def _update_total(self, *_):
        self._total_lbl.setText(tr("sport_log_reps_total", total=self._total()))

    def _build(self):
        lay = QVBoxLayout(self)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(12)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))

        # Rank saat ini
        total_now = db.get_sport_rep_total(self.user_id, self.activity["id"])
        rank = db.get_rep_rank(total_now)
        rank_color = SPORT_RANK_COLORS.get(rank["key"], "#9aa0a6")
        rank_lbl = QLabel(f"{rank['icon']} {tr('sport_rank_' + rank['key'])}   ·   "
                          + tr("sport_reps_total_label", n=total_now))
        rank_lbl.setStyleSheet(f"color:{rank_color}; font-weight:bold; font-size:12px;")
        lay.addWidget(rank_lbl)

        # Quick add buttons
        quick = QHBoxLayout()
        for n in (5, 10, 25, 50):
            b = _btn(f"+{n}", h=34)
            b.setMinimumWidth(60)
            b.clicked.connect(lambda _, v=n: self._reps.setValue(self._reps.value() + v))
            quick.addWidget(b)
        quick.addStretch()
        lay.addLayout(quick)

        # Set × reps
        form = QHBoxLayout()
        form.addWidget(QLabel(tr("sport_log_reps_sets")))
        self._sets = QSpinBox()
        self._sets.setRange(1, 50)
        self._sets.setValue(1)
        self._sets.setMinimumHeight(38)
        form.addWidget(self._sets)
        form.addSpacing(10)
        form.addWidget(QLabel("× " + tr("sport_log_reps_reps")))
        self._reps = QSpinBox()
        self._reps.setRange(1, 1000)
        self._reps.setValue(10)
        self._reps.setMinimumHeight(38)
        form.addWidget(self._reps)
        self._total_lbl = QLabel()
        self._total_lbl.setStyleSheet(f"color:{_T('accent')}; font-weight:bold; font-size:14px;")
        form.addWidget(self._total_lbl)
        form.addStretch()
        lay.addLayout(form)
        self._sets.valueChanged.connect(self._update_total)
        self._reps.valueChanged.connect(self._update_total)
        self._update_total()

        # Catatan
        self._note = _input(tr("dialog_notes_placeholder"))
        lay.addWidget(self._note)

        hint = QLabel(tr("sport_log_reps_hint"))
        hint.setWordWrap(True)
        hint.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        lay.addWidget(hint)

        # Riwayat sesi terakhir
        hist = db.get_sport_rep_history(self.user_id, self.activity["id"], 5)
        if hist:
            hl = QLabel("\n".join(
                f"💪 {h['reps']} reps ×{h['sets']} set · {h['log_date']}"
                + (f" — {h['note']}" if h.get("note") else "")
                for h in hist))
            hl.setStyleSheet(f"color:{_T('muted')}; font-size:10px;")
            lay.addWidget(hl)

        ok = _btn(tr("dialog_save"), "solid", self._save, 46)
        lay.addWidget(ok)

    def _save(self):
        r = db.add_sport_rep_log(self.user_id, self.activity["id"],
                                 self._total(), self._sets.value(),
                                 note=self._note.text())
        if not r.get("ok"):
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", ""), "error")
            return
        SND.complete()
        _show(self, tr("berhasil_title"),
              tr("sport_reps_logged", reps=self._total(), total=r["total_reps"]), "success")
        if r.get("rank_up"):
            rank_name = tr("sport_rank_" + r["rank_after"]["key"])
            SND.levelup() if hasattr(SND, "levelup") else SND.complete()
            _show(self, tr("berhasil_title"), tr("sport_rank_up", rank=rank_name), "success")
        self.accept()


class SportTrackPage(QWidget):
    """Halaman utama SportTrack dengan sub-tab per jenis olahraga."""
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self.card_widgets = {}
        self.mode = "sport"
        self._rep_stats = {}     # {activity_id: {total, today, week, rank}}
        self._build()
        AppState.register(self.load)

    def _mini_stat(self, title, color):
        """Kartu stat mini untuk strip reps di header Sport."""
        card = _card()
        card.setMinimumHeight(74)
        v = QVBoxLayout(card)
        v.setContentsMargins(12, 8, 12, 8)
        v.setSpacing(2)
        t = QLabel(title)
        t.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        t.setWordWrap(True)
        val = QLabel("0")
        val.setStyleSheet(f"color:{color}; font-size:18px; font-weight:bold;")
        v.addWidget(t)
        v.addWidget(val)
        card.value_label = val
        return card

    # ── build UI ──────────────────────────────────────────────────────────────
    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(20, 16, 20, 16)
        root.setSpacing(10)

        # Header (tetap di atas, tidak di-scroll)
        hdr = QHBoxLayout()
        hdr.addWidget(_lbl(tr("sport_title"), "section", 14, True))
        hdr.addStretch()
        folder_btn = _btn(tr("task_folder_button"), h=36)
        folder_btn.setMinimumWidth(110)
        folder_btn.clicked.connect(self._open_folder_add)
        hdr.addWidget(folder_btn)
        add_btn = _btn(tr("task_add_button"), "solid", self._open_add)
        add_btn.setMinimumWidth(130)
        hdr.addWidget(add_btn)
        root.addLayout(hdr)
        root.addWidget(_sep())

        # ========== BODY SCROLL - SEMUA MASUK SCROLL (FIX: tidak menutupi setengah layar) ==========
        body_scroll = QScrollArea()
        body_scroll.setWidgetResizable(True)
        body_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        body_scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")
        body_container = QWidget()
        body_lay = QVBoxLayout(body_container)
        body_lay.setContentsMargins(0, 8, 0, 0)
        body_lay.setSpacing(10)

        # Sport Level bar (sekarang di dalam scroll)
        lvl_row = QHBoxLayout()
        self._sport_lvl_lbl = QLabel(tr("sport_level_default"))
        self._sport_lvl_lbl.setStyleSheet(
            "color:#f0a800; font-weight:bold; font-size:13px;")
        self._sport_xp_bar = QProgressBar()
        self._sport_xp_bar.setMinimumHeight(12)
        self._sport_xp_bar.setTextVisible(False)
        self._sport_xp_bar.setStyleSheet(
            "QProgressBar { background:#222; border:1px solid #555;"
            " border-radius:5px; }"
            "QProgressBar::chunk { background:#f0a800; border-radius:4px; }")
        lvl_row.addWidget(self._sport_lvl_lbl)
        lvl_row.addWidget(self._sport_xp_bar, 1)
        body_lay.addLayout(lvl_row)

        # ── Stats strip reps: hari ini, 7 hari, rank terbaik ── (compact, inside scroll)
        stats_w = QWidget()
        stats_h = QHBoxLayout(stats_w)
        stats_h.setContentsMargins(0, 0, 0, 0)
        stats_h.setSpacing(8)
        self._reps_today_card = self._mini_stat(tr("sport_reps_today"), "#80c000")
        self._reps_week_card = self._mini_stat(tr("sport_reps_week"), "#4da6ff")
        self._best_rank_card = self._mini_stat(tr("sport_best_rank"), "#f0c040")
        for c in (self._reps_today_card, self._reps_week_card, self._best_rank_card):
            c.setMinimumHeight(60)
            stats_h.addWidget(c)
        body_lay.addWidget(stats_w)

        # ── Calorie stats row (compact, inside scroll) ──
        cal_stats_w = QWidget()
        cal_h = QHBoxLayout(cal_stats_w)
        cal_h.setContentsMargins(0, 0, 0, 0)
        cal_h.setSpacing(8)
        self._cal_today_card = self._mini_stat(tr("sport_cal_today"), "#ff6b35")
        self._cal_week_card = self._mini_stat(tr("sport_cal_week"), "#ff8c42")
        self._cal_avg_card = self._mini_stat(tr("sport_cal_avg"), "#a78bfa")
        for c in (self._cal_today_card, self._cal_week_card, self._cal_avg_card):
            c.setMinimumHeight(60)
            cal_h.addWidget(c)
        body_lay.addWidget(cal_stats_w)

        self._reps_chart = SportRepsChartWidget()
        self._reps_chart.setMinimumHeight(140)
        body_lay.addWidget(self._reps_chart)

        # ── Baris filter (search + difficulty) ──
        filter_widget = QWidget()
        filter_layout = QHBoxLayout(filter_widget)
        filter_layout.setContentsMargins(0, 0, 0, 0)
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText(tr("task_search"))
        self.search_input.textChanged.connect(self.load)
        self.filter_combo = QComboBox()
        self.filter_combo.addItem(tr("sport_filter_all"), "all")
        self.filter_combo.addItem(tr("sport_difficulty_easy"), "easy")
        self.filter_combo.addItem(tr("sport_difficulty_medium"), "medium")
        self.filter_combo.addItem(tr("sport_difficulty_hard"), "hard")
        self.filter_combo.addItem(tr("sport_difficulty_epic"), "epic")
        self.filter_combo.currentTextChanged.connect(self.load)
        filter_layout.addWidget(self.search_input)
        filter_layout.addWidget(self.filter_combo)
        body_lay.addWidget(filter_widget)

        # Tab widget — "Semua" + satu tab per jenis olahraga
        self._tabs = QTabWidget()
        self._tabs.setMinimumHeight(320)

        # Tab "Semua"
        self._inner_all = QWidget()
        self._lay_all   = QVBoxLayout(self._inner_all)
        self._lay_all.setSpacing(8)
        self._lay_all.addStretch()
        self._tabs.addTab(_scrolled(self._inner_all), tr("sport_tab_all"))

        # Satu tab per jenis olahraga
        self._sport_lays: dict   = {}
        self._sport_inners: dict  = {}
        for key, sport in db.SPORT_TYPES.items():
            inner = QWidget()
            lay   = QVBoxLayout(inner)
            lay.setSpacing(8)
            lay.addStretch()
            self._sport_lays[key]   = lay
            self._sport_inners[key] = inner
            self._tabs.addTab(_scrolled(inner),
                            f"{sport['icon']} {tr('sport_type_' + key)}")

        body_lay.addWidget(self._tabs, 1)
        body_scroll.setWidget(body_container)
        root.addWidget(body_scroll, 1)
        self.load()

    # ── helpers ───────────────────────────────────────────────────────────────
    def _clear_lay(self, lay: QVBoxLayout):
        while lay.count() > 1:
            item = lay.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

    def _render_sport_to_layout(self, lay: QVBoxLayout, container: QWidget,
                                items: list, folders: list, empty_msg: str):
        """Render sport items ke layout dengan dukungan folder collapsible."""
        if not items and not folders:
            el = _lbl(empty_msg, "sub", 13)
            el.setAlignment(Qt.AlignmentFlag.AlignCenter)
            el.setStyleSheet(f"color:{_T('muted')}; padding:30px;")
            lay.insertWidget(0, el)
            return

        insert_pos = 0
        folder_items: dict = {f["id"]: [] for f in folders}
        ungrouped: list    = []
        for item in items:
            fid = item.get("folder_id")
            if fid and fid in folder_items:
                folder_items[fid].append(item)
            else:
                ungrouped.append(item)

        for folder in folders:
            fw = FolderWidget(folder, "sport", self.user_id, self.load, parent=container)
            cards_in_folder = folder_items.get(folder["id"], [])
            if not cards_in_folder:
                empty_lbl = QLabel("   📭  " + tr("folder_empty"))
                empty_lbl.setStyleSheet(
                    f"color:{_T('muted')}; font-size:12px; padding:6px 0;")
                fw.add_card(empty_lbl)
            else:
                for item in cards_in_folder:
                    fw.add_card(self._create_card_widget(item))
            lay.insertWidget(insert_pos, fw)
            insert_pos += 1

        for item in ungrouped:
            lay.insertWidget(insert_pos, self._create_card_widget(item))
            insert_pos += 1

        # ── Drop area untuk "Tanpa Folder" ──
        drop_area = QFrame()
        drop_area.setAcceptDrops(True)
        drop_area.setMinimumHeight(50)
        drop_area.setStyleSheet(f"""
            QFrame {{
                border: 2px dashed {_T('border')};
                border-radius: 8px;
                background: {_T('bg')};
                margin-top: 8px;
            }}
            QFrame:hover {{
                border-color: {_T('accent')};
                background: {_T('panel')};
            }}
        """)
        drop_label = QLabel(tr("drop_here_to_remove_folder"))
        drop_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        drop_label.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        drop_area_layout = QVBoxLayout(drop_area)
        drop_area_layout.addWidget(drop_label)

        def _drag_enter(e):
            if e.mimeData().hasFormat("application/x-craftlife-card"):
                e.acceptProposedAction()
            else:
                e.ignore()
        drop_area.dragEnterEvent = _drag_enter

        def _drop(e):
            raw = e.mimeData().data("application/x-craftlife-card")
            info = json.loads(raw.data().decode())
            if info["mode"] == self.mode and info["user_id"] == self.user_id:
                db.set_item_folder(self.user_id, self.mode, info["item_id"], None)
                self.load()
                e.acceptProposedAction()
            else:
                e.ignore()
        drop_area.dropEvent = _drop

        lay.insertWidget(insert_pos, drop_area)

    def _refresh_topbar(self):
        """Refresh HANYA top bar (XP/gold/HP) tanpa reload page -> scroll tetap di tempat."""
        try:
            self.window()._topbar.refresh()
        except Exception:
            try: AppState.refresh()
            except Exception: pass

    def _update_item(self, item_id: int):
        activities = db.get_sport_activities(self.user_id)
        item = next((a for a in activities if a['id'] == item_id), None)
        if not item:
            self.load()
            return

        card = self.card_widgets.get(item_id)
        if not card:
            self.load()
            return

        while card.layout().count():
            child = card.layout().takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        new_content = self._build_card_content(item)
        card.layout().addWidget(new_content)
        card.current_folder_id = item.get('folder_id')
        # Update top bar (hanya topbar, bukan reload page)
        self._refresh_topbar()

    # ── load / refresh ────────────────────────────────────────────────────────
    def load(self):
        _sa = self._tabs.currentWidget()
        try:
            self._saved_scroll = _sa.verticalScrollBar().value() if isinstance(_sa, QScrollArea) else 0
        except Exception:
            self._saved_scroll = 0
        self.card_widgets.clear()
        if not AppState.user_id:
            return
        db.reset_daily_tasks(self.user_id)

        # Update sport level bar
        u = AppState.user()
        sport_lvl = u.get("sport_level", 1) or 1
        sport_xp  = u.get("sport_xp", 0)  or 0
        needed    = sport_lvl * 100
        self._sport_lvl_lbl.setText(
            tr("sport_level", lvl=sport_lvl, xp=sport_xp, need=needed))
        self._sport_xp_bar.setMaximum(needed)
        self._sport_xp_bar.setValue(sport_xp)

        # Bersihkan semua layout
        self._clear_lay(self._lay_all)
        for lay in self._sport_lays.values():
            self._clear_lay(lay)

        activities = db.get_sport_activities(self.user_id)
        folders    = db.get_task_folders(self.user_id, "sport")

        # ── Refresh stats reps (strip header + chart) ──
        self._rep_stats = db.get_user_sport_rep_stats(self.user_id)
        today_total = sum(s["today"] for s in self._rep_stats.values())
        week_total = sum(s["week"] for s in self._rep_stats.values())
        self._reps_today_card.value_label.setText(f"{today_total} reps")
        self._reps_week_card.value_label.setText(f"{week_total} reps")
        # ── Calorie stats (FIX 7) ──
        try:
            tot_cals_today = 0
            tot_cals_week = 0
            total_per_min = []
            for a in activities:
                c = a.get("calories_burned",0) or 0
                d = a.get("duration_minutes",0) or 0
                # Estimate today/week based on streak & done_today
                if a.get("done_today"):
                    tot_cals_today += c
                # Assume 3x per week if streak >0 else 0 for week estimate
                if a.get("streak",0) > 0:
                    # Simple: calories * min(streak,7) for week
                    tot_cals_week += c * min(a.get("streak",0), 7) // 3 + c
                if d and c:
                    total_per_min.append(c/d)
            avg_per_min = sum(total_per_min)/len(total_per_min) if total_per_min else 0
            self._cal_today_card.value_label.setText(f"{tot_cals_today} kcal")
            self._cal_week_card.value_label.setText(f"{tot_cals_week} kcal")
            self._cal_avg_card.value_label.setText(f"{avg_per_min:.1f} kcal/min")
        except Exception as e:
            pass
        # Rank terbaik di antara semua aktivitas
        best_key, best_idx, best_name = "rookie", -1, ""
        for a in activities:
            st = self._rep_stats.get(a["id"])
            if st and st["rank"]["index"] > best_idx:
                best_idx = st["rank"]["index"]
                best_key = st["rank"]["key"]
                best_name = a["name"]
        best_txt = tr("sport_rank_" + best_key)
        if best_name:
            best_txt += f" ({best_name})"
        self._best_rank_card.value_label.setText(best_txt)
        self._reps_chart.refresh(db.get_sport_rep_series(self.user_id))

        # Filter berdasarkan search & difficulty
        search = self.search_input.text().lower()
        filter_diff = self.filter_combo.currentData()
        if search:
            activities = [a for a in activities if search in a["name"].lower()]
        if filter_diff != "all":
            activities = [a for a in activities if a.get("difficulty", "").lower() == filter_diff]

        # ── Tab "Semua" ───────────────────────────────────────────────────────
        self._render_sport_to_layout(
            self._lay_all, self._inner_all, activities, folders,
            tr("sport_empty"))

        # ── Tab per jenis ─────────────────────────────────────────────────────
        by_sport: dict = {}
        for a in activities:
            by_sport.setdefault(a["sport_type"], []).append(a)

        for key, lay in self._sport_lays.items():
            sport_nm = tr("sport_type_" + key)
            self._render_sport_to_layout(
                lay, self._sport_inners[key], by_sport.get(key, []), folders,
                tr("sport_no_activity", sport_nm=sport_nm))
        try:
            _sa2 = self._tabs.currentWidget()
            if isinstance(_sa2, QScrollArea):
                _val = getattr(self, "_saved_scroll", 0)
                QTimer.singleShot(0, lambda v=_val: _sa2.verticalScrollBar().setValue(v))
        except Exception:
            pass

    # ── card builder ──────────────────────────────────────────────────────────


    def _build_card_content(self, item: dict) -> QWidget:
        done = bool(item.get("done_today", False))

        content = QWidget()
        row = QHBoxLayout(content)
        row.setContentsMargins(14, 10, 14, 10)
        row.setSpacing(10)

        sport_data = db.SPORT_TYPES.get(item["sport_type"], {"icon": "🏅"})
        ico = _emoji_label(item.get("icon") or sport_data["icon"], ICON_CARD)
        row.addWidget(ico)

        info = QVBoxLayout()
        info.setSpacing(2)
        name_style = (f"text-decoration:line-through; color:{_T('muted')};"
                    if done else f"color:{_T('text')}; font-weight:bold;")
        nm = QLabel(item["name"])
        nm.setStyleSheet(f"font-size:14px; {name_style}")
        info.addWidget(nm)

        sport_nm = tr("sport_type_" + item["sport_type"])
        streak_txt = (tr("task_streak_days", streak=item['streak'])
                    if item.get("streak", 0) > 0 else "")
        sub = QLabel(tr("sport_reward_format", sport=sport_nm, xp=item['xp_reward'],
                        gold=item['gold_reward'], sp=item['sport_points_reward'], streak=streak_txt))
        sub.setTextFormat(Qt.TextFormat.RichText)
        sub.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        info.addWidget(sub)

        # ── Professional Calorie Row (FIX 7) ──
        cals = item.get("calories_burned", 0) or 0
        dur = item.get("duration_minutes", 0) or 0
        if cals:
            cals_per_min = cals / max(1, dur) if dur else 0
            cal_row = QHBoxLayout()
            cal_row.setSpacing(8)
            cal_chip = QLabel(f"🔥 {cals} kcal")
            cal_chip.setStyleSheet("color:#ff6b35; font-weight:bold; font-size:11px; background: rgba(255,107,53,18); border-radius:6px; padding:2px 8px;")
            cal_row.addWidget(cal_chip)
            if dur:
                dur_lbl = QLabel(f"⏱ {dur} min • {cals_per_min:.1f} kcal/min")
                dur_lbl.setStyleSheet(f"color:{_T('muted')}; font-size:10px;")
                cal_row.addWidget(dur_lbl)
            MET_MAP = {"running":9.8,"gym":6.0,"cycling":7.5,"swimming":8.0,"yoga":3.0,"football":7.0,"calisthenics":5.0,"martial_arts":8.0,"badminton":5.5,"other":4.0}
            met = MET_MAP.get(item.get("sport_type"), 4.0)
            met_lbl = QLabel(f"MET {met}")
            met_lbl.setStyleSheet(f"color:{_T('accent')}; font-size:9px; border:1px solid {_T('accent')}; border-radius:6px; padding:1px 6px;")
            cal_row.addWidget(met_lbl)
            # Weekly estimate
            try:
                streak = item.get("streak",0) or 0
                weekly_est = cals * (7 if streak>3 else 3)
                week_lbl = QLabel(f"📅 ~{weekly_est} kcal/minggu")
                week_lbl.setStyleSheet(f"color:{_T('muted')}; font-size:9px;")
                cal_row.addWidget(week_lbl)
            except:
                pass
            cal_row.addStretch()
            info.addLayout(cal_row)

        # ── Baris RANK reps per variasi (chip + progres ke rank berikutnya) ──
        st = getattr(self, "_rep_stats", {}).get(item["id"])
        rank = st["rank"] if st else db.get_rep_rank(0)
        total_reps = st["total"] if st else 0
        rank_color = SPORT_RANK_COLORS.get(rank["key"], "#9aa0a6")

        rank_row = QHBoxLayout()
        rank_row.setSpacing(8)
        rank_chip = QLabel(f"{rank['icon']} {tr('sport_rank_' + rank['key'])}")
        rank_chip.setStyleSheet(
            f"color:{rank_color}; font-weight:bold; font-size:11px; "
            f"background:{_T('panel')}; border:1px solid {rank_color}; "
            f"border-radius:9px; padding:1px 8px;")
        rank_row.addWidget(rank_chip)

        reps_info = QLabel(tr("sport_reps_total_label", n=total_reps))
        reps_info.setStyleSheet(f"color:{_T('muted')}; font-size:10px;")
        rank_row.addWidget(reps_info)
        if st and st["today"]:
            today_lbl = QLabel(tr("sport_reps_chip", n=st["today"]))
            today_lbl.setStyleSheet(f"color:#80c000; font-size:10px; font-weight:bold;")
            rank_row.addWidget(today_lbl)
        rank_row.addStretch()
        info.addLayout(rank_row)

        # Progres ke rank berikutnya
        if rank["next_reps"]:
            pbar = QProgressBar()
            pbar.setRange(rank["min_reps"], rank["next_reps"])
            pbar.setValue(min(total_reps, rank["next_reps"]))
            pbar.setFixedHeight(10)
            pbar.setTextVisible(False)
            pbar.setStyleSheet(
                "QProgressBar { background:#222; border:1px solid #555; border-radius:5px; }"
                f"QProgressBar::chunk {{ background:{rank_color}; border-radius:4px; }}")
            info.addWidget(pbar)
            next_key = db.SPORT_REP_RANKS[rank["index"] + 1][1]
            next_lbl = QLabel(tr("sport_rank_progress", cur=total_reps,
                                 next=rank["next_reps"],
                                 next_rank=tr("sport_rank_" + next_key)))
            next_lbl.setStyleSheet(f"color:{_T('muted')}; font-size:9px;")
            info.addWidget(next_lbl)
        elif total_reps > 0:
            max_lbl = QLabel(f"👑 {tr('sport_rank_max')}")
            max_lbl.setStyleSheet(f"color:{rank_color}; font-size:10px; font-weight:bold;")
            info.addWidget(max_lbl)

        if item.get("notes"):
            nl = QLabel(f"📝 {item['notes']}")
            nl.setWordWrap(True)
            nl.setStyleSheet(f"color:{_T('muted')}; font-size:11px; font-style:italic;")
            info.addWidget(nl)

        history = db.get_task_last_history(self.user_id, "sport", item["id"])
        if history:
            date_str = history["action_date"]
            emoji = "✅" if history["action"] == "success" else "❌" if history["action"] == "fail" else "⏭️"
            color = "#80c000" if history["action"] == "success" else "#e05050" if history["action"] == "fail" else "#f0a800"
            status_label = QLabel(tr("task_status_history", emoji=emoji, date=date_str))
            status_label.setStyleSheet(f"color:{color}; font-size:10px;")
            info.addWidget(status_label)
        else:
            status_label = QLabel(tr("task_status_never"))
            status_label.setStyleSheet(f"color:{_T('muted')}; font-size:10px;")
            info.addWidget(status_label)

        row.addLayout(info, 1)

        reps_btn = _btn(tr("sport_log_reps_btn"), "solid", h=36)
        reps_btn.setMinimumWidth(104)
        reps_btn.clicked.connect(lambda _, a=item: self._open_log_reps(a))
        row.addWidget(reps_btn)

        ck = _btn(tr("sport_btn_done") if done else tr("sport_btn_complete"), h=36)
        ck.setEnabled(not done)
        if done:
            ck.setStyleSheet(f"background:{_T('border')}; color:{_T('muted')}; border-color:{_T('border')};")
        ck.setMinimumWidth(100)
        ck.clicked.connect(lambda _, i=item["id"]: self._complete(i))
        row.addWidget(ck)

        edit_btn = _btn("✏️", h=36)
        edit_btn.setMinimumWidth(36)
        edit_btn.clicked.connect(lambda _, i=item["id"]: self._edit(i))
        row.addWidget(edit_btn)

        dl = _btn("🗑", "danger", h=36)
        dl.setMinimumWidth(36)
        dl.clicked.connect(lambda _, i=item["id"]: self._delete(i))
        row.addWidget(dl)

        dup_btn = _btn("📋", h=36)
        dup_btn.setMinimumWidth(36)
        dup_btn.clicked.connect(lambda _, i=item["id"]: self._duplicate(i))
        row.addWidget(dup_btn)

        folder_btn = _btn("📁", h=36)
        folder_btn.setMinimumWidth(36)
        folder_btn.clicked.connect(lambda _, iid=item["id"], fid=item.get("folder_id"): self._move_to_folder(iid, fid))
        row.addWidget(folder_btn)

        up_btn = _btn("⬆", h=36)
        up_btn.setMinimumWidth(36)
        up_btn.clicked.connect(lambda _, i=item["id"]: self._reorder(i, "up"))
        row.addWidget(up_btn)

        down_btn = _btn("⬇", h=36)
        down_btn.setMinimumWidth(36)
        down_btn.clicked.connect(lambda _, i=item["id"]: self._reorder(i, "down"))
        row.addWidget(down_btn)

        return content

    def _create_card_widget(self, item: dict) -> DraggableCard:
        content = self._build_card_content(item)
        card = DraggableCard(
            item_id=item["id"],
            mode="sport",
            user_id=self.user_id,
            current_folder_id=item.get("folder_id"),
            content_widget=content,
            parent=self
        )
        self.card_widgets[item["id"]] = card
        return card

    # ── actions ───────────────────────────────────────────────────────────────

    def _open_log_reps(self, activity: dict):
        """Buka dialog pencatatan reps (set × reps) untuk aktivitas ini."""
        dlg = LogSportRepsDialog(self.user_id, activity, self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)
            self._refresh_topbar()

    def _move_to_folder(self, item_id: int, current_folder_id: int = None):
        folders = db.get_task_folders(self.user_id, "sport")   # ✅
        folder_names = [tr("folder_no_folder")] + [f["name"] for f in folders]
        folder_ids = [None] + [f["id"] for f in folders]

        dlg = QDialog(self)
        dlg.setWindowTitle(tr("dialog_move_folder"))
        dlg.setMinimumSize(300, 200)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        combo = QComboBox()
        for name in folder_names:
            combo.addItem(name)
        if current_folder_id is not None:
            try:
                idx = folder_ids.index(current_folder_id)
                combo.setCurrentIndex(idx)
            except ValueError:
                pass
        btn_ok = _btn(tr("dialog_move"), "solid", dlg.accept)
        layout.addWidget(QLabel(tr("dialog_select_folder")))
        layout.addWidget(combo)
        layout.addWidget(btn_ok)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            selected_id = folder_ids[combo.currentIndex()]
            db.set_item_folder(self.user_id, "sport", item_id, selected_id)
            QTimer.singleShot(0, self.load)
            SND.notify()

    def _complete(self, iid: int):
        r = db.complete_sport_activity(self.user_id, iid)
        if not r.get("ok", True) and r.get("msg"):
            _show(self, tr("info_title"), r["msg"])
            return
        SND.complete()
        msg = tr("sport_complete_message",
        xp=r.get('xp_gained', 0),
        gold=r.get('gold_gained', 0),
        sp=r.get('sport_points_gained', 0))
        if r.get("leveled_up"):
            msg += tr("level_up_sport", lvl=r['new_level'])
        if r.get("sport_leveled_up"):
            msg += tr("sport_level_up", lvl=r['new_sport_level'])
        _show(self, tr("sport_complete_title"), msg, "success")
        self._refresh_topbar()
        self._update_item(iid)

    def _edit(self, iid: int):
        activities = db.get_sport_activities(self.user_id)
        item = next((a for a in activities if a["id"] == iid), None)
        if not item:
            return
        dlg = AddSportActivityDialog(self.user_id, item, self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _delete(self, iid: int):
        db.delete_sport_activity(self.user_id, iid)
        SND.click()
        QTimer.singleShot(0, self.load)

    def _duplicate(self, activity_id):
        r = db.duplicate_sport_activity(self.user_id, activity_id)
        if r["ok"]:
             SND.click()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        QTimer.singleShot(0, self.load)

    def _open_add(self):
        dlg = AddSportActivityDialog(self.user_id, parent=self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _open_folder_add(self):
        dlg = FolderDialog("sport", self.user_id, parent=self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _reorder(self, item_id, direction):
        r = db.reorder_item(self.user_id, self.mode, item_id, direction)
        if r.get("ok"):
            if r.get("moved", False):
                SND.click()
                QTimer.singleShot(0, self.load)
            # else tidak ada perubahan, tidak perlu reload
        else:
            SND.error()
            _show(self, tr("msg_error"), r.get("msg", "Gagal mengubah urutan"), "error")

    def closeEvent(self, e):
        AppState.unregister(self.load)
        super().closeEvent(e)


# ══════════════════════════════════════════════════════════════════════════════
#  ECONOMY PAGE  — track income & expense with folders, categories, sub-tabs
# ══════════════════════════════════════════════════════════════════════════════

class AddEconomyDialog(QDialog):
    """Dialog untuk tambah/edit item ekonomi."""
    def __init__(self, user_id: int, item=None, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.item = item
        self.setWindowTitle(tr("economy_transaction_title_edit") if item else tr("economy_transaction_title_add"))
        self.setMinimumWidth(480)
        self.setMinimumHeight(520)
        self.setMaximumHeight(700)
        self.setStyleSheet(build_ss())
        self._build()
        # Jika edit, konversi amount dari IDR ke mata uang user (tampil bulat)
        if self.item:
            user_curr = db.get_user_currency(self.user_id)
            amount_idr = self.item['amount']
            amount_usr = db.convert_from_idr(amount_idr, user_curr)
            self._amount.setText(f"{int(round(amount_usr))}")

    def currency_symbol(self):
        curr = db.get_user_currency(self.user_id)
        symbols = {"IDR": "Rp", "USD": "$", "EUR": "€"}
        return symbols.get(curr, "Rp")

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)

        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(14)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))
        lay.addWidget(_sep())

        # Nama
        lay.addWidget(_lbl(tr("economy_transaction_name"), size=12))
        self._name = _input(tr("economy_transaction_name_ph"))
        if self.item:
            self._name.setText(self.item["name"])
        lay.addWidget(self._name)

        # Icon (pilihan singkat)
        lay.addWidget(_lbl(tr("dialog_icon"), size=12))
        icon_choices = [
            (tr("economy_icon_income"), "💰"),
            (tr("economy_icon_expense"), "💸"),
            (tr("economy_icon_salary"), "🏦"),
            (tr("economy_icon_food"), "🍔"),
            (tr("economy_icon_transport"), "🚗"),
            (tr("economy_icon_rent"), "🏠"),
            (tr("economy_icon_education"), "📚"),
            (tr("economy_icon_health"), "💊"),
            (tr("economy_icon_entertainment"), "🎮"),
            (tr("economy_icon_investment"), "📈"),
            (tr("economy_icon_gift"), "🎁"),
            (tr("economy_icon_shopping"), "🛒")
        ]
        self._icon = _combo(icon_choices)
        if self.item:
            idx = self._icon.findData(self.item.get("icon", "💰"))
            if idx >= 0:
                self._icon.setCurrentIndex(idx)
        lay.addWidget(self._icon)

        # Tipe (income/expense)
        lay.addWidget(_lbl(tr("economy_transaction_type"), size=12))
        self._type = _combo([(tr("economy_type_income"), "income"), (tr("economy_type_expense"), "expense")])
        if self.item:
            idx = self._type.findData(self.item.get("type", "expense"))
            if idx >= 0:
                self._type.setCurrentIndex(idx)
        lay.addWidget(self._type)

        # Jumlah
        lay.addWidget(_lbl(tr("economy_amount_label", symbol=self.currency_symbol()), size=12))
        self._amount = QLineEdit()
        self._amount.setPlaceholderText("0")
        self._amount.setMinimumHeight(42)
        if self.item:
            self._amount.setText(str(self.item["amount"]))
        lay.addWidget(self._amount)

        # Kategori (bisa custom atau pilih dari yang sudah ada)
        lay.addWidget(_lbl(tr("economy_category"), size=12))
        self._category = QLineEdit()
        self._category.setPlaceholderText(tr("economy_category_ph"))
        if self.item:
            self._category.setText(self.item["category"])
        # Tambahkan label penjelasan
        cat_hint = QLabel(tr("economy_category_hint"))
        cat_hint.setStyleSheet(f"color:{_T('muted')}; font-size:10px; font-style:italic;")
        lay.addWidget(self._category)
        lay.addWidget(cat_hint)
        # Saran kategori dari DB
        self._suggest_cat = QComboBox()
        self._suggest_cat.setMinimumHeight(36)
        self._suggest_cat.addItem(tr("economy_category_suggest"), None)
        try:
            cats = db.get_economy_categories(self.user_id)
            for cat in cats:
                self._suggest_cat.addItem(cat, cat)
            self._suggest_cat.currentIndexChanged.connect(self._on_suggest)
        except Exception:
            pass
        lay.addWidget(self._suggest_cat)

        # Tanggal
        lay.addWidget(_lbl(tr("economy_date_label"), size=12))
        self._date = QLineEdit()
        self._date.setPlaceholderText("YYYY-MM-DD")
        if self.item:
            self._date.setText(self.item["date"])
        else:
            from datetime import date
            self._date.setText(date.today().isoformat())
        self._date.setMinimumHeight(42)
        lay.addWidget(self._date)

        # Catatan
        lay.addWidget(_lbl(tr("dialog_notes"), size=12))
        self._notes = _input(tr("dialog_notes_placeholder"))
        if self.item:
            self._notes.setText(self.item.get("notes", ""))
        lay.addWidget(self._notes)

        # Folder
        lay.addWidget(_lbl(tr("economy_folder_label"), size=12))
        _folders = db.get_task_folders(self.user_id, "economy")
        _fopts = [(tr("dialog_no_folder"), None)] + [(f"{fd['icon']}  {fd['name']}", fd["id"]) for fd in _folders]
        self._folder = _combo(_fopts)
        if self.item and self.item.get("folder_id"):
            cur_fid = self.item["folder_id"]
            idx = next((i for i, (_, d) in enumerate(_fopts) if d == cur_fid), 0)
            self._folder.setCurrentIndex(idx)
        lay.addWidget(self._folder)

        lay.addSpacing(8)
        ok_lbl = tr("dialog_save") if self.item else tr("dialog_add")
        ok = _btn(ok_lbl, "solid", self._save, 46)
        lay.addWidget(ok)
        self._name.returnPressed.connect(self._save)

        root.addWidget(_scrolled(content))

    def _on_suggest(self, idx):
        if idx >= 0:
            cat = self._suggest_cat.currentData()
            if cat:
                self._category.setText(cat)

    def _save(self):
        name = self._name.text().strip()
        if not name:
            _show(self, tr("msg_error"), tr("msg_name_empty"), "error")
            return
        
        # Ambil jumlah dalam mata uang user
        amount_user = _parse_positive_amount(self._amount.text())
        if amount_user is None:
            _show(self, tr("msg_error"), tr("msg_invalid_amount"), "error")
            return

        # Konversi ke IDR untuk disimpan
        user_currency = db.get_user_currency(self.user_id)
        amount_idr = db.convert_to_idr(amount_user, user_currency)
        
        icon = self._icon.currentData()
        type_ = self._type.currentData()
        cat = self._category.text().strip()
        if not cat:
            cat = "other"
        date_str = self._date.text().strip()
        if not date_str:
            from datetime import date
            date_str = date.today().isoformat()
        notes = self._notes.text()
        folder_id = self._folder.currentData()
        
        if self.item:
            db.update_economy_item(self.item["id"], self.user_id,
                                name=name, icon=icon, type=type_, amount=amount_idr,
                                category=cat, date=date_str, notes=notes, folder_id=folder_id)
        else:
            db.add_economy_item(self.user_id, name, icon, type_, amount_idr, cat, date_str, notes, folder_id)
        
        SND.complete()
        self.accept()


class EconomyTrendWidget(QWidget):
    """Mini bar chart income vs expense per hari (7 hari berurutan).
    Fallback ke teks jika matplotlib tidak tersedia atau belum ada data."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self._lay = QVBoxLayout(self)
        self._lay.setContentsMargins(0, 0, 0, 0)
        self._lay.setSpacing(4)
        self.setMinimumHeight(180)

    def _clear(self):
        while self._lay.count():
            it = self._lay.takeAt(0)
            if it.widget():
                it.widget().deleteLater()

    def refresh(self, series: list):
        """series: [{date, income, expense}] dari db.get_economy_daily_series."""
        self._clear()
        has_data = any(s["income"] or s["expense"] for s in series)
        if not MPL_QT_OK or not has_data:
            lbl = _lbl(tr("economy_daily_chart_empty"), "sub", 11)
            lbl.setWordWrap(True)
            lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lbl.setStyleSheet(f"color: {_T('muted')}; padding: 24px;")
            self._lay.addWidget(lbl)
            return

        xs = list(range(len(series)))
        incomes = [s["income"] for s in series]
        expenses = [s["expense"] for s in series]
        dates = [s["date"] for s in series]
        muted = _T("muted")
        w = 0.38

        fig = Figure(figsize=(6.4, 1.9), tight_layout=True)
        fig.patch.set_alpha(0.0)
        ax = fig.add_subplot(111)
        ax.set_facecolor("none")
        ax.bar([x - w / 2 for x in xs], incomes, width=w, color="#80c000",
               alpha=0.85, label="💰")
        ax.bar([x + w / 2 for x in xs], expenses, width=w, color="#e05050",
               alpha=0.85, label="💸")
        ax.set_xticks(xs)
        ax.set_xticklabels([d[5:] for d in dates], fontsize=7, color=muted)
        ax.tick_params(colors=muted, labelsize=7)
        for spine in ax.spines.values():
            spine.set_visible(False)
        leg = ax.legend(fontsize=7, frameon=False, loc="upper left",
                        labelcolor=muted)

        canvas = FigureCanvas(fig)
        canvas.setStyleSheet("background: transparent;")
        self._lay.addWidget(canvas)


class EconomyPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        
        self.user_id = user_id
        self.mode = "economy"
        self.currency = "IDR"
        self.currency_symbol = "Rp"
        self.card_widgets = {}
        self._daily_date = date.today()   # tanggal terpilih di tab Harian
        self._load_currency_settings()
        self._build()
        AppState.register(self.load)

    def currency_symbol(self):
        curr = db.get_user_currency(self.user_id)
        symbols = {"IDR": "Rp", "USD": "$", "EUR": "€"}
        return symbols.get(curr, "Rp")

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(20, 16, 20, 16)
        root.setSpacing(10)

        # Header
        hdr = QHBoxLayout()
        hdr.addWidget(_lbl(tr("economy_title"), "section", 14, True))
        hdr.addStretch()
        folder_btn = _btn(tr("economy_folder"), h=36)
        folder_btn.setMinimumWidth(110)
        folder_btn.clicked.connect(self._open_folder_add)
        hdr.addWidget(folder_btn)
        add_btn = _btn(tr("economy_add"), "solid", self._open_add)
        add_btn.setMinimumWidth(130)
        hdr.addWidget(add_btn)
        debt_add_btn = _btn(tr("economy_add_debt"), "gold", self._add_debt)
        debt_add_btn.setMinimumWidth(140)
        hdr.addWidget(debt_add_btn)
        saving_add_btn = _btn(tr("economy_add_saving"), "gold", self._add_saving)
        saving_add_btn.setMinimumWidth(150)
        hdr.addWidget(saving_add_btn)
        root.addLayout(hdr)
        root.addWidget(_sep())

        # Summary cards (total income, expense, balance)
        summary_widget = QWidget()
        summary_layout = QHBoxLayout(summary_widget)
        summary_layout.setSpacing(12)
        self.income_card = self._stat_card(tr("economy_total_income"), "0", "#80c000")
        self.income_card.setToolTip(tr("economy_tooltip_income"))
        self.expense_card = self._stat_card(tr("economy_total_expense"), "0", "#e05050")
        self.expense_card.setToolTip(tr("economy_tooltip_expense"))
        self.balance_card = self._stat_card(tr("economy_balance"), "0", "#4da6ff")
        self.balance_card.setToolTip(tr("economy_tooltip_balance"))
        self.debt_label = QLabel()
        self.net_worth_label = QLabel()
        summary_layout.addWidget(self.income_card)
        summary_layout.addWidget(self.expense_card)
        summary_layout.addWidget(self.balance_card)
        # Baris info hutang & tabungan
        info_row = QWidget()
        info_layout = QHBoxLayout(info_row)
        info_layout.setContentsMargins(0, 8, 0, 0)
        info_layout.setSpacing(20)

        self.debt_label = QLabel(tr("economy_total_debt", amount="0"))
        self.debt_label.setStyleSheet(f"color: {_T('muted')}; font-size: 12px;")
        self.saving_label = QLabel(tr("economy_total_saving", amount="0"))
        self.saving_label.setStyleSheet(f"color: {_T('muted')}; font-size: 12px;")

        info_layout.addWidget(self.debt_label)
        info_layout.addStretch()
        info_layout.addWidget(self.saving_label)

        root.addWidget(info_row)
        root.addWidget(summary_widget)

        # Filter bar
        filter_widget = QWidget()
        filter_layout = QHBoxLayout(filter_widget)
        filter_layout.setContentsMargins(0, 0, 0, 0)
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText(tr("economy_search"))
        self.search_input.textChanged.connect(self.load)
        self.type_combo = QComboBox()
        self.type_combo.addItem(tr("economy_filter_all"), None)
        self.type_combo.addItem(tr("economy_filter_income"), "income")
        self.type_combo.addItem(tr("economy_filter_expense"), "expense")
        self.type_combo.currentTextChanged.connect(self.load)
        self.category_combo = QComboBox()
        self.category_combo.addItem(tr("economy_all_categories"), "all")
        self.category_combo.setToolTip(tr("economy_tooltip_category"))
        self.category_combo.currentIndexChanged.connect(self.load)
        filter_layout.addWidget(self.search_input)
        filter_layout.addWidget(self.type_combo)
        filter_layout.addWidget(self.category_combo)
        root.addWidget(filter_widget)

        # Tab widget: "Semua" + per kategori dinamis
        self._tabs = QTabWidget()
        self._inner_all = QWidget()
        self._lay_all = QVBoxLayout(self._inner_all)
        self._lay_all.setSpacing(8)
        self._lay_all.addStretch()
        self._tabs.addTab(_scrolled(self._inner_all), tr("economy_tab_all"))

        # Tab Harian: nav tanggal + kartu harian + chart tren + list transaksi
        self._inner_daily = QWidget()
        self._lay_daily_root = QVBoxLayout(self._inner_daily)
        self._lay_daily_root.setSpacing(8)
        self._lay_daily_root.setAlignment(Qt.AlignmentFlag.AlignTop)
        self._tabs.addTab(_scrolled(self._inner_daily), tr("economy_tab_daily"))

        self._inner_debt = QWidget()
        self._lay_debt = QVBoxLayout(self._inner_debt)
        self._lay_debt.setSpacing(8)
        self._lay_debt.setAlignment(Qt.AlignmentFlag.AlignTop)
        self._tabs.addTab(_scrolled(self._inner_debt), tr("economy_tab_debt"))

        self._inner_debnotes = QWidget()
        self._lay_debnotes = QVBoxLayout(self._inner_debnotes)
        self._lay_debnotes.setSpacing(8)
        self._lay_debnotes.setAlignment(Qt.AlignmentFlag.AlignTop)
        self._tabs.addTab(_scrolled(self._inner_debnotes), tr("economy_tab_debtnotes"))

        self._inner_saving = QWidget()
        self._lay_saving = QVBoxLayout(self._inner_saving)
        self._lay_saving.setSpacing(8)
        self._lay_saving.setAlignment(Qt.AlignmentFlag.AlignTop)
        self._tabs.addTab(_scrolled(self._inner_saving), tr("economy_tab_saving"))

        self._inner_invest = QWidget()
        self._lay_invest = QVBoxLayout(self._inner_invest)
        self._lay_invest.setSpacing(8)
        self._lay_invest.setAlignment(Qt.AlignmentFlag.AlignTop)
        self._tabs.addTab(_scrolled(self._inner_invest), tr("economy_tab_invest"))

        self._inner_subscription = QWidget()
        self._lay_subscription = QVBoxLayout(self._inner_subscription)
        self._lay_subscription.setSpacing(8)
        self._lay_subscription.setAlignment(Qt.AlignmentFlag.AlignTop)
        self._tabs.addTab(_scrolled(self._inner_subscription), tr("economy_tab_subscription"))

        self._cat_lays = {}
        self._cat_inners = {}
        root.addWidget(self._tabs, 1)
        self.load()

    def _create_card_widget(self, item: dict) -> DraggableCard:
        card = self._make_card(item)
        self.card_widgets[item["id"]] = card
        return card

    def _add_investment(self):
        dlg = AddInvestmentDialog(self.user_id, self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _collect_return(self, invest_id):
        from PyQt6.QtWidgets import QInputDialog
        dlg = QInputDialog(self)
        dlg.setWindowTitle(tr("economy_invest_return_title"))
        dlg.setLabelText(tr("economy_invest_return_label", symbol=self.currency_symbol))
        dlg.setDoubleDecimals(0)
        dlg.setDoubleRange(0, 1e9)
        dlg.setDoubleValue(0)
        if dlg.exec():
            amount_user = dlg.doubleValue()
            if amount_user > 0:
                user_currency = db.get_user_currency(self.user_id)
                amount_idr = db.convert_to_idr(amount_user, user_currency)
                r = db.add_investment_return(invest_id, self.user_id, amount_idr)
                if r["ok"]:
                    SND.complete()
                    symbol = self.currency_symbol
                    _show(self, tr("return_title"),
                        tr("investment_return_detail", amount=f"{symbol}{amount_user:.0f}", total=self.format_currency(r['new_amount'])),
                        "success")
                else:
                    SND.error()
                    _show(self, tr("gagal_title"), r["msg"], "error")
                QTimer.singleShot(0, self.load)

    def _withdraw_investment(self, invest_id):
        reply = QMessageBox.question(self, tr("confirm_title"), tr("economy_invest_withdraw_confirm"),
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            r = db.withdraw_investment(invest_id, self.user_id)
            if r["ok"]:
                SND.complete()
                amount_usr = self.convert_from_idr(r['amount'])
                _show(self, tr("berhasil_title"), tr("berhasil_invest_withdraw", symbol=self.currency_symbol, amount=amount_usr), "success")
                AppState.refresh()
            else:
                SND.error()
                _show(self, tr("gagal_title"), r["msg"], "error")
            QTimer.singleShot(0, self.load)


    def _delete_investment(self, invest_id):
        reply = QMessageBox.question(self, tr("confirm_title"), tr("economy_invest_delete_confirm"), 
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_investment(invest_id, self.user_id)
            SND.click()
            QTimer.singleShot(0, self.load)

    def _add_subscription(self):
        dlg = AddSubscriptionDialog(self.user_id, parent=self)   
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _renew_subscription(self, sub_id):
        # Renew manual (pengguna klik)
        r = db.renew_subscription(sub_id, self.user_id, auto_pay=False)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        QTimer.singleShot(0, self.load)
        AppState.refresh()

    def _edit_subscription(self, sub):
        dlg = AddSubscriptionDialog(self.user_id, sub, self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _delete_subscription(self, sub_id):
        reply = QMessageBox.question(self, tr("confirm_title"), tr("economy_sub_delete_confirm"),
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_subscription(sub_id, self.user_id)
            SND.click()
            QTimer.singleShot(0, self.load)

    def _make_investment_card(self, inv: dict) -> QFrame:
        f = _card()
        row = QHBoxLayout(f)
        row.setContentsMargins(14, 10, 14, 10)
        row.setSpacing(10)

        icon = QLabel(inv["icon"])
        icon.setFont(QFont("Segoe UI", 14))
        icon.setMinimumWidth(38)
        row.addWidget(icon)

        info = QVBoxLayout()
        info.setSpacing(2)
        name = QLabel(inv["name"])
        name.setStyleSheet(f"font-size:14px; font-weight:bold; color:{_T('text')};")
        info.addWidget(name)

        amount = QLabel(self.format_currency(inv["amount"]))
        amount.setStyleSheet(f"color:{_T('accent')}; font-size:13px;")
        info.addWidget(amount)

        date_str = tr("economy_invest_date_label", date=inv['invested_date'][:10])
        date_lbl = QLabel(date_str)
        date_lbl.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        info.addWidget(date_lbl)

        if inv.get("notes"):
            note = QLabel(f"📝 {inv['notes']}")
            note.setWordWrap(True)
            note.setStyleSheet(f"color:{_T('muted')}; font-size:11px; font-style:italic;")
            info.addWidget(note)

        row.addLayout(info, 1)

        collect_btn = _btn(tr("economy_invest_collect"), "solid", h=36)
        collect_btn.clicked.connect(lambda _, iid=inv["id"]: self._collect_return(iid))
        row.addWidget(collect_btn)

        withdraw_btn = _btn(tr("economy_invest_withdraw_btn"), "diamond", h=36)
        withdraw_btn.clicked.connect(lambda _, iid=inv["id"]: self._withdraw_investment(iid))
        row.addWidget(withdraw_btn)

        del_btn = _btn("🗑", "danger", h=36)
        del_btn.clicked.connect(lambda _, iid=inv["id"]: self._delete_investment(iid))
        row.addWidget(del_btn)

        return f

    def _make_subscription_card(self, sub: dict) -> QFrame:
        f = _card()
        today = date.today().isoformat()
        overdue = sub["due_date"] < today

        if overdue:
            f.setStyleSheet(
                f"QFrame#card {{ background: {_T('panel')}; "
                f"border: 2px solid #e05050; border-left: 6px solid #e05050; "
                f"border-radius: 8px; }}"
            )
        else:
            f.setStyleSheet(f"QFrame#card {{ background: {_T('panel')}; "
                            f"border: 1px solid {_T('border')}; border-radius: 8px; }}")

        row = QHBoxLayout(f)
        row.setContentsMargins(14, 10, 14, 10)
        row.setSpacing(10)

        icon = QLabel(sub["icon"])
        icon.setFont(QFont("Segoe UI", 14))
        icon.setMinimumWidth(38)
        row.addWidget(icon)

        info = QVBoxLayout()
        info.setSpacing(2)
        name = QLabel(sub["name"])
        name.setStyleSheet(f"font-size:14px; font-weight:bold; color:{_T('text')};")
        info.addWidget(name)

        amount = QLabel(tr("subscription_amount_period", amount=self.format_currency(sub['amount']), period=sub['period']))
        amount.setStyleSheet(f"color:{_T('accent')}; font-size:12px;")
        info.addWidget(amount)

        due = tr("economy_sub_due_format", date=sub['due_date'])
        due_lbl = QLabel(due)
        due_lbl.setStyleSheet(f"color:{'#e05050' if overdue else _T('muted')}; font-size:11px;")
        info.addWidget(due_lbl)

        if sub.get("notes"):
            note = QLabel(f"📝 {sub['notes']}")
            note.setWordWrap(True)
            note.setStyleSheet(f"color:{_T('muted')}; font-size:11px; font-style:italic;")
            info.addWidget(note)

        row.addLayout(info, 1)

        if overdue:
            renew_btn = _btn(tr("economy_sub_renew"), "gold", h=36)
            renew_btn.clicked.connect(lambda _, sid=sub["id"]: self._renew_subscription(sid))
            row.addWidget(renew_btn)

        edit_btn = _btn("✏️", h=36)
        edit_btn.clicked.connect(lambda _, s=sub: self._edit_subscription(s))
        row.addWidget(edit_btn)

        del_btn = _btn("🗑", "danger", h=36)
        del_btn.clicked.connect(lambda _, sid=sub["id"]: self._delete_subscription(sid))
        row.addWidget(del_btn)

        return f

    def _load_investments(self):
        # Bersihkan layout
        while self._lay_invest.count():
            item = self._lay_invest.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        
        investments = db.get_investments(self.user_id)
        if not investments:
            empty = QLabel(tr("economy_invest_empty"))
            empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            empty.setStyleSheet(f"color:{_T('muted')}; padding:30px;")
            self._lay_invest.addWidget(empty)
        else:
            for inv in investments:
                card = self._make_investment_card(inv)
                self._lay_invest.addWidget(card)
        
        # Tombol tambah investasi
        add_btn = _btn(tr("economy_invest_add"), "gold", self._add_investment)
        add_btn.setMinimumHeight(40)
        self._lay_invest.addWidget(add_btn)

    def _load_subscriptions(self):
        while self._lay_subscription.count():
            item = self._lay_subscription.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        
        # Cek overdue & auto-renewal
        db.check_all_subscriptions(self.user_id)
        subs = db.get_subscriptions(self.user_id, active_only=True)
        if not subs:
            empty = QLabel(tr("economy_sub_empty"))
            empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            empty.setStyleSheet(f"color:{_T('muted')}; padding:30px;")
            self._lay_subscription.addWidget(empty)
        else:
            for sub in subs:
                card = self._make_subscription_card(sub)
                self._lay_subscription.addWidget(card)
        
        add_btn = _btn(tr("economy_sub_add"), "gold", self._add_subscription)
        add_btn.setMinimumHeight(40)
        self._lay_subscription.addWidget(add_btn)

    def _add_saving(self):
        dlg = AddEditSavingDialog(self.user_id, parent=self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _load_savings(self):
        # Bersihkan layout tabungan
        while self._lay_saving.count():
            item = self._lay_saving.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        savings = db.get_savings(self.user_id)
        if not savings:
            empty = QLabel(tr("economy_saving_empty"))
            empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            empty.setStyleSheet(f"color:{_T('muted')}; padding:30px;")
            self._lay_saving.addWidget(empty)
        else:
            for saving in savings:
                card = self._make_saving_card(saving)
                self._lay_saving.addWidget(card)

    def _make_saving_card(self, saving: dict) -> QFrame:
        f = _card()
        row = QHBoxLayout(f)
        row.setContentsMargins(14, 10, 14, 10)
        row.setSpacing(10)

        icon = QLabel(saving["icon"])
        icon.setFont(QFont("Segoe UI", 14))
        icon.setMinimumWidth(38)
        row.addWidget(icon)

        info = QVBoxLayout()
        info.setSpacing(2)
        nm = QLabel(saving["name"])
        nm.setStyleSheet(f"font-size:14px; font-weight:bold; color:{_T('text')};")
        info.addWidget(nm)

        target_text = ""
        if saving["target_amount"] > 0:
            percent = (saving["current_amount"] / saving["target_amount"]) * 100
            target_text = "  " + tr("saving_target_format", amount=self.format_currency(saving['target_amount']), percent=percent)
        if saving.get("target_date"):
            target_text += "  " + tr("saving_deadline_format", date=saving['target_date'])
        sub = QLabel(tr("saving_current_amount", amount=self.format_currency(saving['current_amount']), target_text=target_text))
        sub.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        info.addWidget(sub)

        if saving.get("notes"):
            nl = QLabel(f"📝 {saving['notes']}")
            nl.setWordWrap(True)
            nl.setStyleSheet(f"color:{_T('muted')}; font-size:11px; font-style:italic;")
            info.addWidget(nl)

        if saving["target_amount"] > 0:
            pb = QProgressBar()
            target = int(saving["target_amount"])
            if target > 2147483647:
                target = 2147483647
            pb.setMaximum(target)
            pb.setValue(int(saving["current_amount"]))
            pb.setMinimumHeight(10)
            pb.setTextVisible(False)
            info.addWidget(pb)

        row.addLayout(info, 1)

        add_btn = _btn(tr("dialog_add"), "gold", h=36)
        add_btn.setMinimumWidth(80)
        add_btn.clicked.connect(lambda _, sid=saving["id"], cur=saving["current_amount"]: self._add_to_saving(sid, cur))
        row.addWidget(add_btn)

        withdraw_btn = _btn(tr("economy_saving_withdraw_btn"), "diamond", h=36)
        withdraw_btn.setMinimumWidth(80)
        withdraw_btn.clicked.connect(lambda _, sid=saving["id"], name=saving["name"], cur=saving["current_amount"]: self._withdraw_from_saving(sid, name, cur))
        row.addWidget(withdraw_btn)

        edit_btn = _btn("✏️", h=36)
        edit_btn.setMinimumWidth(36)
        edit_btn.clicked.connect(lambda _, s=saving: self._edit_saving(s))
        row.addWidget(edit_btn)

        del_btn = _btn("🗑", "danger", h=36)
        del_btn.setMinimumWidth(36)
        del_btn.clicked.connect(lambda _, sid=saving["id"]: self._delete_saving(sid))
        row.addWidget(del_btn)

        return f

    def _add_to_saving(self, saving_id, current_amount_idr):
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("economy_saving_add_funds"))
        dlg.setMinimumSize(350, 200)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)

        layout.addWidget(QLabel(tr("saving_balance_label", balance=self.format_currency(current_amount_idr))))
        layout.addWidget(QLabel(tr("saving_add_amount_label", symbol=self.currency_symbol)))
        
        amount_spin = QDoubleSpinBox()
        amount_spin.setRange(0, 1_000_000_000)
        amount_spin.setValue(10000)
        amount_spin.setPrefix(f"{self.currency_symbol}: ")
        amount_spin.setMinimumHeight(42)
        layout.addWidget(amount_spin)

        btn = _btn(tr("economy_saving_add_confirm_btn"), "solid")
        btn.clicked.connect(lambda: self._process_add_to_saving(saving_id, amount_spin.value(), dlg))
        layout.addWidget(btn)
        
        dlg.exec()

    def _process_add_to_saving(self, saving_id, amount_user, dialog):
        if amount_user <= 0:
            _show(self, tr("msg_error"), tr("economy_amount_gt_zero"), "error")
            return
        user_currency = db.get_user_currency(self.user_id)
        amount_idr = db.convert_to_idr(amount_user, user_currency)
        r = db.add_to_saving(saving_id, self.user_id, amount_idr)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), tr("berhasil_saving_add", symbol=self.currency_symbol, amount=amount_user), "success")
            dialog.accept()
            QTimer.singleShot(0, self.load)
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _edit_saving(self, saving):
        dlg = AddEditSavingDialog(self.user_id, saving, self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _delete_saving(self, saving_id):
        reply = QMessageBox.question(self, tr("confirm_title"), tr("economy_saving_delete_confirm"), 
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_saving(saving_id, self.user_id)
            SND.click()
            QTimer.singleShot(0, self.load)

    def _withdraw_from_saving(self, saving_id, saving_name, current_amount_idr):
        """Dialog untuk menarik uang dari tabungan dengan konversi mata uang"""
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("economy_saving_withdraw_title", name=saving_name))
        dlg.setMinimumSize(400, 260)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)

        current_user = self.convert_from_idr(current_amount_idr)
        layout.addWidget(QLabel(tr("saving_withdraw_balance_label", symbol=self.currency_symbol, balance=current_user)))
        layout.addWidget(QLabel(tr("economy_saving_withdraw_label")))
        amount_spin = QDoubleSpinBox()
        amount_spin.setRange(0, current_user)
        amount_spin.setPrefix(f"{self.currency_symbol}: ")
        amount_spin.setValue(0)
        layout.addWidget(amount_spin)

        cb_income = QCheckBox(tr("economy_saving_withdraw_record"))
        cb_income.setChecked(True)
        layout.addWidget(cb_income)

        btn_ok = _btn(tr("economy_saving_withdraw_btn"), "solid")
        btn_ok.clicked.connect(lambda: self._do_withdraw(saving_id, amount_spin.value(), cb_income.isChecked(), dlg))
        layout.addWidget(btn_ok)
        dlg.exec()

    def _do_withdraw(self, saving_id, amount_user, add_to_income, dialog):
        if amount_user <= 0:
            _show(self, tr("msg_error"), tr("economy_amount_gt_zero"), "error")
            return
        user_currency = db.get_user_currency(self.user_id)
        amount_idr = db.convert_to_idr(amount_user, user_currency)
        r = db.withdraw_from_saving(saving_id, self.user_id, amount_idr)
        if r["ok"]:
            SND.complete()
            if add_to_income:
                from datetime import date
                today = date.today().isoformat()
                db.add_economy_item(
                    self.user_id,
                    name="Penarikan Tabungan",
                    icon="🏦",
                    type_="income",
                    amount=amount_idr,
                    category="Tabungan",
                    date_str=today,
                    notes=f"Penarikan dari tabungan ID {saving_id}"
                )
                _show(self, tr("berhasil_title"), tr("berhasil_saving_withdraw", symbol=self.currency_symbol, amount=amount_user), "success")
            else:
                _show(self, tr("berhasil_title"), tr("berhasil_saving_withdraw2", symbol=self.currency_symbol, amount=amount_user), "success")
            dialog.accept()
            QTimer.singleShot(0, self.load)
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _stat_card(self, title, value, color):
        card = _card()
        card.setMinimumHeight(90)
        layout = QVBoxLayout(card)
        layout.setContentsMargins(12, 12, 12, 12)
        lbl_title = QLabel(title)
        lbl_title.setStyleSheet(f"color: {_T('muted')}; font-size: 12px;")
        lbl_value = QLabel(value)
        lbl_value.setStyleSheet(f"color: {color}; font-size: 22px; font-weight: bold;")
        layout.addWidget(lbl_title)
        layout.addWidget(lbl_value)
        # simpan referensi untuk update
        card.value_label = lbl_value
        return card

    def _clear_lay(self, lay: QVBoxLayout):
        while lay.count() > 1:
            item = lay.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

    def _daily_prev(self):
        from datetime import timedelta
        self._daily_date -= timedelta(days=1)
        self.load()

    def _daily_next(self):
        from datetime import timedelta
        self._daily_date += timedelta(days=1)
        self.load()

    def _daily_today(self):
        self._daily_date = date.today()
        self.load()

    def _render_daily(self):
        """Render ulang TAB HARIAN: navigasi tanggal ◀ ▶, 4 kartu stat harian
        (pemasukan/pengeluaran/selisih/saldo s.d. tanggal itu), chart tren
        7 hari, dan daftar transaksi tanggal terpilih. Ringkasan all-time di
        atas halaman TETAP utuh — ini lapisan harian yang lebih terarah."""
        while self._lay_daily_root.count():
            it = self._lay_daily_root.takeAt(0)
            if it.widget():
                it.widget().deleteLater()

        date_str = self._daily_date.isoformat()

        # ── Navigasi tanggal ──
        nav = QWidget()
        nlay = QHBoxLayout(nav)
        nlay.setContentsMargins(0, 0, 0, 0)
        prev_btn = _btn("◀", h=36)
        prev_btn.setMinimumWidth(46)
        prev_btn.clicked.connect(self._daily_prev)
        nlay.addWidget(prev_btn)

        date_txt = f"📅 {date_str}"
        if self._daily_date == date.today():
            date_txt += f"  {tr('economy_daily_today_note')}"
        dlbl = QLabel(date_txt)
        dlbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        dlbl.setStyleSheet(f"font-size:14px; font-weight:bold; color:{_T('text')};")
        nlay.addWidget(dlbl, 1)

        next_btn = _btn("▶", h=36)
        next_btn.setMinimumWidth(46)
        next_btn.clicked.connect(self._daily_next)
        nlay.addWidget(next_btn)
        today_btn = _btn(tr("economy_daily_today_btn"), h=36)
        today_btn.setMinimumWidth(90)
        today_btn.clicked.connect(self._daily_today)
        nlay.addWidget(today_btn)
        self._lay_daily_root.addWidget(nav)

        # ── 4 kartu stat harian ──
        tot = db.get_economy_daily_totals(self.user_id, date_str)
        running = db.get_economy_balance_until(self.user_id, date_str)
        cards = QWidget()
        clay = QHBoxLayout(cards)
        clay.setContentsMargins(0, 0, 0, 0)
        clay.setSpacing(10)
        c_inc = self._stat_card(tr("economy_daily_income"),
                                self.format_currency(tot["income"]), "#80c000")
        c_exp = self._stat_card(tr("economy_daily_expense"),
                                self.format_currency(tot["expense"]), "#e05050")
        net_color = "#80c000" if tot["net"] >= 0 else "#e05050"
        c_net = self._stat_card(tr("economy_daily_net"),
                                self.format_currency(tot["net"]), net_color)
        c_run = self._stat_card(tr("economy_daily_running"),
                                self.format_currency(running), "#f0c040")
        # Referensi untuk pembaruan/pemantauan (dipakai juga smoke test)
        self._daily_inc_lbl = c_inc.value_label
        self._daily_exp_lbl = c_exp.value_label
        self._daily_net_lbl = c_net.value_label
        self._daily_run_lbl = c_run.value_label
        for c in (c_inc, c_exp, c_net, c_run):
            clay.addWidget(c)
        self._lay_daily_root.addWidget(cards)

        # ── Chart tren 7 hari (berakhir di tanggal terpilih) ──
        self._lay_daily_root.addWidget(_lbl(tr("economy_daily_chart_title"), "sub", 11))
        self._daily_chart = EconomyTrendWidget()
        self._daily_chart.refresh(db.get_economy_daily_series(self.user_id, date_str, 7))
        self._lay_daily_root.addWidget(self._daily_chart)

        # ── Daftar transaksi tanggal terpilih ──
        self._lay_daily_root.addWidget(_lbl(tr("economy_daily_list_title"), "sub", 11))
        items = db.get_economy_items_by_date(self.user_id, date_str)
        folders = db.get_task_folders(self.user_id, "economy")
        host = QWidget()
        hlay = QVBoxLayout(host)
        hlay.setContentsMargins(0, 0, 0, 0)
        hlay.setSpacing(8)
        hlay.addStretch()
        self._render_to_layout(hlay, host, items, folders, tr("economy_daily_empty"))
        self._lay_daily_root.addWidget(host)

    def _render_to_layout(self, lay: QVBoxLayout, container: QWidget, items: list, folders: list, empty_msg: str):
        if not items and not folders:
            el = _lbl(empty_msg, "sub", 13)
            el.setAlignment(Qt.AlignmentFlag.AlignCenter)
            el.setStyleSheet(f"color:{_T('muted')}; padding:30px;")
            lay.insertWidget(0, el)
            return

        insert_pos = 0
        folder_items = {f["id"]: [] for f in folders}
        ungrouped = []
        for item in items:
            fid = item.get("folder_id")
            if fid and fid in folder_items:
                folder_items[fid].append(item)
            else:
                ungrouped.append(item)

        for folder in folders:
            fw = FolderWidget(folder, "economy", self.user_id, self.load, parent=container)
            cards = folder_items.get(folder["id"], [])
            if not cards:
                empty_lbl = QLabel("   📭  " + tr("folder_empty"))
                empty_lbl.setStyleSheet(f"color:{_T('muted')}; font-size:12px; padding:6px 0;")
                fw.add_card(empty_lbl)
            else:
                for item in cards:
                    fw.add_card(self._create_card_widget(item))
            lay.insertWidget(insert_pos, fw)
            insert_pos += 1

        for item in ungrouped:
            lay.insertWidget(insert_pos, self._create_card_widget(item))
            insert_pos += 1

        # ── Drop area untuk "Tanpa Folder" ──
        drop_area = QFrame()
        drop_area.setAcceptDrops(True)
        drop_area.setMinimumHeight(50)
        drop_area.setStyleSheet(f"""
            QFrame {{
                border: 2px dashed {_T('border')};
                border-radius: 8px;
                background: {_T('bg')};
                margin-top: 8px;
            }}
            QFrame:hover {{
                border-color: {_T('accent')};
                background: {_T('panel')};
            }}
        """)
        drop_label = QLabel(tr("drop_here_to_remove_folder"))
        drop_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        drop_label.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        drop_area_layout = QVBoxLayout(drop_area)
        drop_area_layout.addWidget(drop_label)

        def _drag_enter(e):
            if e.mimeData().hasFormat("application/x-craftlife-card"):
                e.acceptProposedAction()
            else:
                e.ignore()
        drop_area.dragEnterEvent = _drag_enter

        def _drop(e):
            raw = e.mimeData().data("application/x-craftlife-card")
            info = json.loads(raw.data().decode())
            if info["mode"] == self.mode and info["user_id"] == self.user_id:
                db.set_item_folder(self.user_id, self.mode, info["item_id"], None)
                self.load()
                e.acceptProposedAction()
            else:
                e.ignore()
        drop_area.dropEvent = _drop

        lay.insertWidget(insert_pos, drop_area)

    def _make_card(self, item: dict) -> QFrame:
        content = QWidget()
        row = QHBoxLayout(content)
        row.setContentsMargins(14, 10, 14, 10)
        row.setSpacing(10)

        ico = _emoji_label(item["icon"], 28)
        row.addWidget(ico)

        info = QVBoxLayout()
        info.setSpacing(2)
        nm = QLabel(item["name"])
        nm.setStyleSheet(f"font-size:14px; font-weight:bold; color:{_T('text')};")
        info.addWidget(nm)

        tipe = tr("economy_type_income") if item["type"] == "income" else tr("economy_type_expense")
        color = "#80c000" if item["type"] == "income" else "#e05050"
        amount_str = f"+{self.format_currency(item['amount'])}" if item["type"] == "income" else f"-{self.format_currency(item['amount'])}"
        sub = QLabel(tr("economy_item_format", color=color, type=tipe, amount=amount_str, category=item['category'], date=item['date']))
        sub.setTextFormat(Qt.TextFormat.RichText)
        sub.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        info.addWidget(sub)

        if item.get("notes"):
            nl = QLabel(f"📝 {item['notes']}")
            nl.setWordWrap(True)
            nl.setStyleSheet(f"color:{_T('muted')}; font-size:11px; font-style:italic;")
            info.addWidget(nl)

        row.addLayout(info, 1)

        edit_btn = _btn("✏️", h=36)
        edit_btn.setMinimumWidth(36)
        edit_btn.clicked.connect(lambda _, i=item["id"]: self._edit(i))
        row.addWidget(edit_btn)

        dl = _btn("🗑", "danger", h=36)
        dl.setMinimumWidth(36)
        dl.clicked.connect(lambda _, i=item["id"]: self._delete(i))
        row.addWidget(dl)

        dup_btn = _btn("📋", h=36)
        dup_btn.setMinimumWidth(36)
        dup_btn.clicked.connect(lambda _, i=item["id"]: self._duplicate(i))
        row.addWidget(dup_btn)

        folder_btn = _btn("📁", h=36)
        folder_btn.setMinimumWidth(36)
        folder_btn.clicked.connect(lambda _, iid=item["id"], fid=item.get("folder_id"): self._move_to_folder(iid, fid))
        row.addWidget(folder_btn)

        # Tombol panah untuk reorder (naik/turun)
        up_btn = _btn("⬆", h=36)
        up_btn.setMinimumWidth(36)
        up_btn.clicked.connect(lambda _, i=item["id"]: self._reorder(i, "up"))
        row.addWidget(up_btn)

        down_btn = _btn("⬇", h=36)
        down_btn.setMinimumWidth(36)
        down_btn.clicked.connect(lambda _, i=item["id"]: self._reorder(i, "down"))
        row.addWidget(down_btn)

        card = DraggableCard(
            item_id=item["id"],
            mode="economy",
            user_id=self.user_id,
            current_folder_id=item.get("folder_id"),
            content_widget=content,
            parent=self
        )
        return card

    def _create_card_widget(self, item: dict) -> DraggableCard:
        """Bungkus kartu dan simpan referensi."""
        card = self._make_card(item)
        self.card_widgets[item["id"]] = card
        return card

    def _reorder(self, item_id, direction):
        r = db.reorder_item(self.user_id, self.mode, item_id, direction)
        if r.get("ok"):
            if r.get("moved", False):
                SND.click()
                QTimer.singleShot(0, self.load)
            # else tidak ada perubahan, tidak perlu reload
        else:
            SND.error()
            _show(self, tr("msg_error"), r.get("msg", "Gagal mengubah urutan"), "error")

    def _move_to_folder(self, item_id: int, current_folder_id: int = None):
        folders = db.get_task_folders(self.user_id, "economy")
        folder_names = [tr("folder_no_folder")] + [f["name"] for f in folders]
        folder_ids = [None] + [f["id"] for f in folders]
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("dialog_move_folder"))
        dlg.setMinimumSize(300, 200)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        combo = QComboBox()
        for name in folder_names:
            combo.addItem(name)
        if current_folder_id is not None:
            try:
                idx = folder_ids.index(current_folder_id)
                combo.setCurrentIndex(idx)
            except ValueError:
                pass
        btn_ok = _btn(tr("dialog_move"), "solid", dlg.accept)
        layout.addWidget(QLabel(tr("dialog_select_folder")))
        layout.addWidget(combo)
        layout.addWidget(btn_ok)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            selected_id = folder_ids[combo.currentIndex()]
            db.update_economy_item(item_id, self.user_id, folder_id=selected_id)
            QTimer.singleShot(0, self.load)
            SND.notify()

    def load(self):
        if not AppState.user_id:
            return

        self.card_widgets.clear()
        self._load_currency_settings()
        self._check_and_apply_debt_penalties()

        # Update summary cards
        summary = db.get_economy_summary(self.user_id)
        self.income_card.value_label.setText(self.format_currency(summary['total_income']))
        self.expense_card.value_label.setText(self.format_currency(summary['total_expense']))
        self.balance_card.value_label.setText(self.format_currency(summary['balance']))
        unpaid_total = db.get_total_unpaid_debt(self.user_id)
        self.debt_label.setText(tr("economy_total_debt", amount=self.format_currency(unpaid_total)))
        savings_list = db.get_savings(self.user_id)
        total_savings = sum(s['current_amount'] for s in savings_list) if savings_list else 0
        self.saving_label.setText(tr("economy_total_saving", amount=self.format_currency(total_savings)))
        self.saving_label.setStyleSheet(f"color: {_T('muted')}; font-size: 12px;")

        # Refresh category combo
        current_cat = self.category_combo.currentData()
        self.category_combo.blockSignals(True)
        self.category_combo.clear()
        self.category_combo.addItem(tr("economy_all_categories"), "all")
        try:
            cats = db.get_economy_categories(self.user_id)
            for cat in cats:
                self.category_combo.addItem(cat, cat)
        except Exception:
            pass
        if current_cat and current_cat != "all":
            idx = self.category_combo.findData(current_cat)
            if idx >= 0:
                self.category_combo.setCurrentIndex(idx)
        self.category_combo.blockSignals(False)

        while self._lay_debt.count():
            item = self._lay_debt.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        debts = db.get_debts(self.user_id, include_paid=False)  # hanya yang belum lunas
        unpaid_total = db.get_total_unpaid_debt(self.user_id)

        # Filter items
        type_filter = None
        type_text = self.type_combo.currentData()
        if type_text == "income":
            type_filter = "income"
        elif type_text == "expense":
            type_filter = "expense"

        cat_filter = self.category_combo.currentData()
        if cat_filter == "all":
            cat_filter = None
        search = self.search_input.text().strip() or None

        items = db.get_economy_items(self.user_id, type_filter, cat_filter, search)
        folders = db.get_task_folders(self.user_id, "economy")

        # Bersihkan semua tab
        self._clear_lay(self._lay_all)
        # Hapus tab kategori dinamis (7 tab statis dipertahankan:
        # Semua, Harian, Hutang, Catatan Hutang, Tabungan, Investasi, Langganan)
        while self._tabs.count() > 7:
            self._tabs.removeTab(7)
        self._cat_lays.clear()
        self._cat_inners.clear()

        # Tab "Semua"
        self._render_to_layout(self._lay_all, self._inner_all, items, folders,
                               tr("economy_empty"))

        # Refresh tab hutang

        if not debts:
            empty = QLabel(tr("economy_debt_empty"))
            empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            empty.setStyleSheet(f"color:{_T('muted')}; padding:30px;")
            self._lay_debt.addWidget(empty)
        else:
            for debt in debts:
                card = self._make_debt_card(debt)
                self._lay_debt.addWidget(card)

        # ── Tab Harian: pemasukan/pengeluaran/saldo per tanggal terpilih ──
        self._render_daily()

        # ── Tab Catatan Hutang (piutang: orang lain berhutang KE user) ──
        while self._lay_debnotes.count():
            dn_item = self._lay_debnotes.takeAt(0)
            if dn_item.widget():
                dn_item.widget().deleteLater()
        debnotes = db.get_debt_notes(self.user_id)

        dn_head = QWidget()
        dn_hlay = QHBoxLayout(dn_head)
        dn_hlay.setContentsMargins(0, 0, 0, 0)
        dn_unpaid_total = db.get_total_unpaid_debt_notes(self.user_id)
        dn_total_lbl = QLabel(tr("debnote_total_unpaid", amount=self.format_currency(dn_unpaid_total)))
        dn_total_lbl.setStyleSheet(f"color:{_T('muted')}; font-size:12px; font-weight:bold;")
        dn_hlay.addWidget(dn_total_lbl, 1)
        dn_add_btn = _btn(tr("debnote_add_btn"), "solid", h=36)
        dn_add_btn.clicked.connect(self._add_debnote)
        dn_hlay.addWidget(dn_add_btn)
        self._lay_debnotes.addWidget(dn_head)

        if not debnotes:
            dn_empty = QLabel(tr("debnote_empty"))
            dn_empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            dn_empty.setStyleSheet(f"color:{_T('muted')}; padding:30px;")
            self._lay_debnotes.addWidget(dn_empty)
        else:
            for note in debnotes:
                self._lay_debnotes.addWidget(self._make_debnote_card(note))

        # Tab per kategori (hanya dari items yang ada)
        by_cat = {}
        for it in items:
            cat = it["category"]
            by_cat.setdefault(cat, []).append(it)
        for cat, cat_items in by_cat.items():
            inner = QWidget()
            lay = QVBoxLayout(inner)
            lay.setSpacing(8)
            lay.addStretch()
            self._cat_lays[cat] = lay
            self._cat_inners[cat] = inner
            self._tabs.addTab(_scrolled(inner), f"📁 {cat}")
            self._render_to_layout(lay, inner, cat_items, folders, f"Tidak ada transaksi di kategori {cat}")

        # Tampilkan jumlah hutang overdue sebagai label (bukan popup agar tidak mengganggu)
        overdue_count = db.get_overdue_debts_count(self.user_id)
        if overdue_count > 0:
            self.debt_label.setText(
                tr("economy_overdue_warning", count=overdue_count)
            )
            self.debt_label.setStyleSheet("color: #e05050; font-size: 12px; font-weight: bold;")
        else:
            self.debt_label.setText(tr("economy_total_debt", amount=self.format_currency(unpaid_total)))
            self.debt_label.setStyleSheet(f"color: {_T('muted')}; font-size: 12px;")
      
        self._load_savings()
        self._load_investments()
        self._load_subscriptions()

    def _check_and_apply_debt_penalties(self):
        conn = db.get_conn()
        today = date.today().isoformat()
        try:
            debts = conn.execute("""
                SELECT id, due_date, name FROM debts
                WHERE user_id=? AND is_paid=0 AND penalty_applied=0 AND due_date IS NOT NULL
            """, (self.user_id,)).fetchall()
            debts = [dict(d) for d in debts]
        finally:
            conn.close()

        penalties = []  # untuk menyimpan pesan
        for debt in debts:
            due = debt["due_date"]
            # Guard: due_date adalah input teks bebas — lewati format invalid
            # (mis. '2026-08-00' tersimpan di DB lama) daripada crash satu halaman
            try:
                due_dt = datetime.strptime(due, "%Y-%m-%d").date()
            except (ValueError, TypeError):
                log.warning(f"due_date hutang tidak valid (id={debt.get('id')}): {due!r} — dilewati")
                continue
            if due_dt < date.today():
                days_late = (date.today() - due_dt).days
                if days_late >= 3:
                    r = db.apply_late_debt_penalty(self.user_id, debt["id"], days_late)
                    if r.get("ok"):
                        penalties.append(
                            tr("debt_penalty_detail", due=due, days=days_late, gold=r['gold_lost'])
                        )
        if penalties:
            SND.error()
            msg = "\n".join(penalties)
            _show(self, tr("penalty_title"), msg, "warning")

    def _make_debt_card(self, debt: dict) -> QFrame:
        today = date.today().isoformat()
        is_overdue = debt.get("due_date") and debt["due_date"] < today

        f = _card()
        if is_overdue:
            f.setStyleSheet(
                f"QFrame#card {{ background: {_T('panel')}; "
                f"border: 2px solid #e05050; border-left: 6px solid #e05050; "
                f"border-radius: 8px; }}"
            )
        else:
            f.setStyleSheet(f"QFrame#card {{ background: {_T('panel')}; "
                            f"border: 1px solid {_T('border')}; border-radius: 8px; }}")

        row = QHBoxLayout(f)
        row.setContentsMargins(14, 10, 14, 10)
        row.setSpacing(10)

        icon = QLabel("💸")
        icon.setFont(QFont("Segoe UI", 14))
        icon.setMinimumWidth(38)
        row.addWidget(icon)

        info = QVBoxLayout()
        info.setSpacing(2)

        nm = QLabel(debt["name"])
        nm.setStyleSheet(f"font-size:14px; font-weight:bold; color:{_T('text')};")
        info.addWidget(nm)

        amount_str = self.format_currency(debt['amount'])
        due_str = f"  ⏰ Tenggat: {debt['due_date']}" if debt.get("due_date") else ""
        sub = QLabel(amount_str + due_str)
        sub.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        info.addWidget(sub)

        penalty = debt.get("penalty_applied", 0)
        if penalty:
            penalty_lbl = QLabel(tr("debt_penalty_label", amount=self.format_currency(debt.get('penalty_amount',0))))
            penalty_lbl.setStyleSheet("color:#e05050; font-size:10px;")
            info.addWidget(penalty_lbl)

        if debt.get("notes"):
            nl = QLabel(f"📝 {debt['notes']}")
            nl.setWordWrap(True)
            nl.setStyleSheet(f"color:{_T('muted')}; font-size:11px; font-style:italic;")
            info.addWidget(nl)

        row.addLayout(info, 1)

        pay_btn = _btn(tr("economy_btn_pay"), "gold", h=36)
        pay_btn.setMinimumWidth(80)
        pay_btn.clicked.connect(lambda _, d=debt["id"]: self._pay_debt(d))
        row.addWidget(pay_btn)

        installment_btn = _btn(tr("economy_btn_installment"), "diamond", h=36)
        installment_btn.setMinimumWidth(80)
        installment_btn.clicked.connect(lambda _, d=debt: self._installment_debt(d))
        row.addWidget(installment_btn)

        edit_btn = _btn("✏️", h=36)
        edit_btn.setMinimumWidth(36)
        edit_btn.clicked.connect(lambda _, d=debt: self._edit_debt(d))
        row.addWidget(edit_btn)

        del_btn = _btn("🗑", "danger", h=36)
        del_btn.setMinimumWidth(36)
        del_btn.clicked.connect(lambda _, d=debt["id"]: self._delete_debt(d))
        row.addWidget(del_btn)

        return f

    def _make_debnote_card(self, note: dict) -> QFrame:
        """Kartu catatan hutang orang lain KE user (piutang)."""
        paid = note.get("status") == "paid"
        f = _card()
        f.setStyleSheet(
            f"QFrame#card {{ background: {_T('panel')}; "
            f"border: 1px solid {'#3fae5a' if paid else _T('border')}; "
            f"border-left: 6px solid {'#3fae5a' if paid else '#f0a800'}; "
            f"border-radius: 8px; }}"
        )

        row = QHBoxLayout(f)
        row.setContentsMargins(14, 10, 14, 10)
        row.setSpacing(10)
        row.addWidget(_emoji_label("✅" if paid else "📒", ICON_CARD))

        info = QVBoxLayout()
        info.setSpacing(2)

        nm = QLabel(note["person_name"])
        nm.setStyleSheet(f"font-size:14px; font-weight:bold; color:{_T('text')};")
        info.addWidget(nm)

        sub_txt = self.format_currency(note["amount"])
        if note.get("date"):
            sub_txt += f"  📅 {note['date']}"
        sub = QLabel(sub_txt)
        sub.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        info.addWidget(sub)

        if note.get("notes"):
            nl = QLabel(f"📝 {note['notes']}")
            nl.setWordWrap(True)
            nl.setStyleSheet(f"color:{_T('muted')}; font-size:11px; font-style:italic;")
            info.addWidget(nl)

        if paid:
            status_txt = tr("debnote_status_paid")
            if note.get("paid_at"):
                status_txt += f" · {str(note['paid_at'])[:10]}"
        else:
            status_txt = tr("debnote_status_unpaid")
        status = QLabel(status_txt)
        status.setStyleSheet(
            f"color:{'#3fae5a' if paid else '#f0a800'}; font-size:11px; font-weight:bold;")
        info.addWidget(status)

        row.addLayout(info, 1)

        if not paid:
            settle_btn = _btn(tr("debnote_settle_btn"), "solid", h=36)
            settle_btn.setMinimumWidth(120)
            settle_btn.clicked.connect(lambda _, n=note: self._settle_debnote(n))
            row.addWidget(settle_btn)

        del_btn = _btn("🗑", "danger", h=36)
        del_btn.setMinimumWidth(36)
        del_btn.clicked.connect(lambda _, nid=note["id"]: self._delete_debnote(nid))
        row.addWidget(del_btn)

        return f

    def _add_debnote(self):
        dlg = AddDebtNoteDialog(self.user_id, self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)
            AppState.refresh()

    def _settle_debnote(self, note):
        reply = QMessageBox.question(
            self, tr("confirm_title"),
            tr("debnote_settle_confirm", name=note["person_name"]),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply != QMessageBox.StandardButton.Yes:
            return
        r = db.settle_debt_note(self.user_id, note["id"])
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"),
                  tr("debnote_settled", name=note["person_name"]), "success")
            QTimer.singleShot(0, self.load)
            AppState.refresh()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", ""), "error")

    def _delete_debnote(self, note_id):
        reply = QMessageBox.question(
            self, tr("confirm_title"), tr("debnote_delete_confirm"),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_debt_note(self.user_id, note_id)
            SND.click()
            QTimer.singleShot(0, self.load)
            AppState.refresh()

    def _installment_debt(self, debt):
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("economy_debt_installment_title", name=debt['name']))
        dlg.setMinimumSize(400, 250)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        layout.setSpacing(12)

        remaining_text = tr("debt_remaining", amount=self.format_currency(debt['amount']))
        layout.addWidget(QLabel(remaining_text))
        layout.addWidget(QLabel(tr("economy_debt_installment_label")))
        amount_spin = QDoubleSpinBox()
        amount_spin.setRange(1, debt['amount'])
        amount_spin.setPrefix(f"{self.currency_symbol}: ")
        amount_spin.setValue(min(10000, debt['amount']))
        amount_spin.setMinimumHeight(42)
        layout.addWidget(amount_spin)

        info = QLabel(tr("economy_debt_installment_hint"))
        info.setWordWrap(True)
        info.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        layout.addWidget(info)

        btn = _btn(tr("economy_debt_installment_btn"), "solid")
        btn.clicked.connect(lambda: self._process_installment(debt['id'], amount_spin.value(), dlg))
        layout.addWidget(btn)
        dlg.exec()

    def _process_installment(self, debt_id, amount_user, dialog):
        if amount_user <= 0:
            _show(self, tr("msg_error"), tr("installment_amount_zero"), "error")
            return
        user_currency = db.get_user_currency(self.user_id)
        amount_idr = db.convert_to_idr(amount_user, user_currency)
        loading = LoadingDialog("Memproses cicilan...", self)
        loading.show()
        QApplication.processEvents()
        r = db.pay_debt_installment(debt_id, self.user_id, amount_idr)
        loading.accept()
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            dialog.accept()
            QTimer.singleShot(0, self.load)
            AppState.refresh()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _edit(self, item_id):
        items = db.get_economy_items(self.user_id)
        item = next((i for i in items if i["id"] == item_id), None)
        if not item:
            return
        dlg = AddEconomyDialog(self.user_id, item, self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _delete(self, item_id):
        db.delete_economy_item(self.user_id, item_id)
        SND.click()
        QTimer.singleShot(0, self.load)

    def _duplicate(self, item_id):
        r = db.duplicate_economy_item(self.user_id, item_id)
        if r["ok"]:
            SND.click()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", tr("duplicate_failed")), "error")
        QTimer.singleShot(0, self.load)

    def _open_add(self):
        dlg = AddEconomyDialog(self.user_id, parent=self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _open_folder_add(self):
        dlg = FolderDialog("economy", self.user_id, parent=self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _add_debt(self):
        dlg = AddEditDebtDialog(self.user_id, parent=self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _edit_debt(self, debt):
        dlg = AddEditDebtDialog(self.user_id, debt, self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _delete_debt(self, debt_id):
        reply = QMessageBox.question(self, tr("confirm_title"), tr("economy_debt_delete_confirm"), 
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_debt(debt_id, self.user_id)
            SND.click()
            QTimer.singleShot(0, self.load)

    def _pay_debt(self, debt_id):
        loading = LoadingDialog("Memproses pelunasan...", self)
        loading.show()
        QApplication.processEvents()
        r = db.mark_debt_paid(debt_id, self.user_id)
        loading.accept()
        if r["ok"]:
            SND.complete()
            conn = db.get_conn()
            debt = conn.execute("SELECT amount FROM debts WHERE id=?", (debt_id,)).fetchone()
            conn.close()
            if debt:
                amount_usr = self.convert_from_idr(debt["amount"])
                _show(self, tr("berhasil_title"), tr("berhasil_debt_paid", symbol=self.currency_symbol, amount=amount_usr), "success")
            else:
                _show(self, tr("berhasil_title"), r["msg"], "success")
            AppState.refresh()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        QTimer.singleShot(0, self.load)

    def _load_currency_settings(self):
        curr = db.get_user_currency(self.user_id)
        self.currency = curr
        self.CURRENCY_RATES = {"IDR": 1, "USD": 17800, "EUR": 20700}
        self.currency_symbol = {"IDR": "Rp", "USD": "$", "EUR": "€"}[curr]

    def convert_from_idr(self, amount_idr):
        """Konversi dari IDR ke mata uang user (untuk tampilan)."""
        rate = self.CURRENCY_RATES[self.currency]
        return amount_idr / rate

    def convert_to_idr(self, amount):
        """Konversi dari mata uang user ke IDR (untuk penyimpanan)."""
        rate = self.CURRENCY_RATES[self.currency]
        return amount * rate

    def format_currency(self, amount_idr):
        """Kembalikan string mata uang yang sudah dikonversi + simbol."""
        converted = self.convert_from_idr(amount_idr)
        return f"{self.currency_symbol} {converted:,.0f}"

    def closeEvent(self, e):
        AppState.unregister(self.load)
        super().closeEvent(e)


# ========== DEBT DIALOG ==========
class AddEditDebtDialog(QDialog):
    def __init__(self, user_id: int, debt=None, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.debt = debt
        self.setWindowTitle(tr("debt_edit_title") if debt else tr("debt_add_title"))
        self.setMinimumWidth(480)
        self.setMinimumHeight(420)
        self.setStyleSheet(build_ss())
        self._build()
        if self.debt:
            user_curr = db.get_user_currency(self.user_id)
            amount_idr = self.debt['amount']
            amount_usr = db.convert_from_idr(amount_idr, user_curr)
            self._amount.setText(f"{int(round(amount_usr))}")

    def currency_symbol(self):
        curr = db.get_user_currency(self.user_id)
        symbols = {"IDR": "Rp", "USD": "$", "EUR": "€"}
        return symbols.get(curr, "Rp")

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(14)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))
        lay.addWidget(_sep())

        # Nama hutang
        lay.addWidget(_lbl(tr("economy_debt_name_label"), size=12) )
        self._name = _input(tr("economy_debt_name_ph"))
        if self.debt:
            self._name.setText(self.debt["name"])
        lay.addWidget(self._name)

        # Jumlah
        lay.addWidget(_lbl(tr("economy_amount_label", symbol=self.currency_symbol()), size=12))
        self._amount = QLineEdit()
        self._amount.setPlaceholderText("0")
        self._amount.setMinimumHeight(42)
        if self.debt:
            self._amount.setText(str(self.debt["amount"]))
        lay.addWidget(self._amount)

        # Tenggat waktu (opsional)
        lay.addWidget(_lbl(tr("economy_debt_deadline_label"), size=12))
        self._due_date = QLineEdit()
        self._due_date.setPlaceholderText("YYYY-MM-DD")
        from datetime import date, timedelta
        if self.debt and self.debt.get("due_date"):
            self._due_date.setText(self.debt["due_date"])
        else:
            self._due_date.setText((date.today() + timedelta(days=14)).isoformat())
        lay.addWidget(self._due_date)

        # Catatan
        lay.addWidget(_lbl(tr("dialog_notes"), size=12))
        self._notes = _input(tr("dialog_notes_placeholder"))
        if self.debt:
            self._notes.setText(self.debt.get("notes", ""))
        lay.addWidget(self._notes)

        lay.addSpacing(8)
        ok_lbl = tr("dialog_save") if self.debt else tr("economy_debt_add_title")
        ok = _btn(ok_lbl, "solid", self._save, 46)
        lay.addWidget(ok)
        self._name.returnPressed.connect(self._save)

        root.addWidget(_scrolled(content))

    def _save(self):
        name = self._name.text().strip()
        if not name:
            _show(self, tr("msg_error"), tr("economy_debt_nama_not_empty"), "error")
            return
        amount_user = _parse_positive_amount(self._amount.text())
        if amount_user is None:
            _show(self, tr("msg_error"), tr("msg_invalid_amount"), "error")
            return
        
        user_currency = db.get_user_currency(self.user_id)
        amount_idr = db.convert_to_idr(amount_user, user_currency)
        
        due_date = self._due_date.text().strip()
        if due_date == "":
            due_date = None
        # Validasi format tanggal (mencegah nilai invalid seperti '2026-08-00'
        # yang bikin crash pengecekan penalty di halaman Economy)
        if due_date is not None:
            try:
                datetime.strptime(due_date, "%Y-%m-%d")
            except ValueError:
                SND.error()
                _show(self, tr("gagal_title"), tr("economy_debt_date_invalid"), "error")
                return
        notes = self._notes.text()

        if self.debt:
            db.update_debt(self.debt["id"], self.user_id,
                        name=name, amount=amount_idr, due_date=due_date, notes=notes)
            SND.complete()
            _show(self, tr("berhasil_title"), tr("economy_debt_done_update"), "success")
        else:
            r = db.add_debt(self.user_id, name, amount_idr, due_date, notes)
            if r["ok"]:
                SND.complete()
                _show(self, tr("berhasil_title"), tr("economy_succes_add_debt", name=name), "success")
            else:
                SND.error()
                _show(self, tr("gagal_title"), r["msg"], "error")
                return
        self.accept()

# ========== DEBT NOTE DIALOG (Catatan Hutang / piutang) ==========
class AddDebtNoteDialog(QDialog):
    """Catat hutang ORANG LAIN ke user. Saat dibuat, nominal tercatat sebagai
    expense; saat ditandai lunas (tab Catatan Hutang), otomatis jadi income."""
    def __init__(self, user_id: int, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        self.setWindowTitle(tr("debnote_add_title"))
        self.setMinimumWidth(480)
        self.setMinimumHeight(420)
        self.setStyleSheet(build_ss())
        self._build()

    def currency_symbol(self):
        curr = db.get_user_currency(self.user_id)
        symbols = {"IDR": "Rp", "USD": "$", "EUR": "€"}
        return symbols.get(curr, "Rp")

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(14)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))
        lay.addWidget(_sep())

        # Nama peminjam
        lay.addWidget(_lbl(tr("debnote_person_label"), size=12))
        self._name = _input(tr("debnote_person_ph"))
        lay.addWidget(self._name)

        # Nominal
        lay.addWidget(_lbl(tr("economy_amount_label", symbol=self.currency_symbol()), size=12))
        self._amount = QLineEdit()
        self._amount.setPlaceholderText("0")
        self._amount.setMinimumHeight(42)
        lay.addWidget(self._amount)

        # Tanggal pinjam
        lay.addWidget(_lbl(tr("debnote_date_label"), size=12))
        self._date = QLineEdit()
        self._date.setPlaceholderText("YYYY-MM-DD")
        self._date.setText(date.today().isoformat())
        self._date.setMinimumHeight(42)
        lay.addWidget(self._date)

        # Catatan
        lay.addWidget(_lbl(tr("dialog_notes"), size=12))
        self._notes = _input(tr("dialog_notes_placeholder"))
        lay.addWidget(self._notes)

        hint = QLabel(tr("debnote_hint"))
        hint.setWordWrap(True)
        hint.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        lay.addWidget(hint)

        lay.addSpacing(8)
        ok = _btn(tr("dialog_save"), "solid", self._save, 46)
        lay.addWidget(ok)
        self._name.returnPressed.connect(self._save)
        self._amount.returnPressed.connect(self._save)

        root.addWidget(_scrolled(content))

    def _save(self):
        name = self._name.text().strip()
        if not name:
            SND.error()
            _show(self, tr("msg_error"), tr("debnote_person_empty"), "error")
            return
        amount_user = _parse_positive_amount(self._amount.text())
        if amount_user is None:
            SND.error()
            _show(self, tr("msg_error"), tr("msg_invalid_amount"), "error")
            return
        date_str = self._date.text().strip() or date.today().isoformat()
        try:
            datetime.strptime(date_str, "%Y-%m-%d")
        except ValueError:
            SND.error()
            _show(self, tr("gagal_title"), tr("economy_debt_date_invalid"), "error")
            return
        amount_idr = db.convert_to_idr(amount_user, db.get_user_currency(self.user_id))
        r = db.add_debt_note(self.user_id, name, amount_idr, date_str, self._notes.text())
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), tr("debnote_added", name=name), "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", ""), "error")
            return
        self.accept()

# ========== SAVINGS DIALOG ==========
class AddEditSavingDialog(QDialog):
    def __init__(self, user_id: int, saving=None, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.saving = saving
        self.setWindowTitle(tr("saving_edit_title") if saving else tr("saving_add_title"))
        self.setMinimumWidth(480)
        self.setMinimumHeight(520)
        self.setStyleSheet(build_ss())
        self._build()
        if self.saving:
            user_curr = db.get_user_currency(self.user_id)
            current_idr = self.saving.get('current_amount', 0)
            target_idr = self.saving.get('target_amount', 0)
            current_usr = db.convert_from_idr(current_idr, user_curr)
            target_usr = db.convert_from_idr(target_idr, user_curr)
            self._current.setValue(current_usr)
            self._target.setValue(target_usr)

    def currency_symbol(self):
        curr = db.get_user_currency(self.user_id)
        symbols = {"IDR": "Rp", "USD": "$", "EUR": "€"}
        return symbols.get(curr, "Rp")

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(14)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))
        lay.addWidget(_sep())

        # Nama tabungan
        lay.addWidget(_lbl(tr("economy_saving_name_label"), size=12))
        self._name = _input(tr("economy_saving_name_ph"))
        if self.saving:
            self._name.setText(self.saving["name"])
        lay.addWidget(self._name)

        # Icon
        lay.addWidget(_lbl(tr("dialog_icon"), size=12))
        icon_choices = [
            (tr("saving_icon_bank"), "🏦"), 
            (tr("saving_icon_target"), "🎯"),
            (tr("saving_icon_vacation"), "🏖️"),
            (tr("saving_icon_vehicle"), "🚗"),
            (tr("saving_icon_house"), "🏠"),
            (tr("saving_icon_education"), "📚"),
            (tr("saving_icon_health"), "💊"),
            (tr("saving_icon_shopping"), "🛒"),
            (tr("saving_icon_investment"), "💰"),
            (tr("saving_icon_gift"), "🎁"),
            (tr("saving_icon_pet"), "🐾"),
            (tr("saving_icon_food"), "🍽️")
        ]
        self._icon = _combo(icon_choices)
        if self.saving:
            idx = self._icon.findData(self.saving.get("icon", "🏦"))
            if idx >= 0:
                self._icon.setCurrentIndex(idx)
        lay.addWidget(self._icon)

        # Jumlah saat ini
        lay.addWidget(_lbl(tr("economy_saving_current_amount", symbol=self.currency_symbol()), size=12))
        self._current = QDoubleSpinBox()
        self._current.setRange(0, 1_000_000_000)
        self._current.setValue(self.saving["current_amount"] if self.saving else 0)
        self._current.setPrefix(f"{self.currency_symbol()}: ")
        self._current.setMinimumHeight(42)
        lay.addWidget(self._current)

        # Target jumlah
        lay.addWidget(_lbl(tr("economy_saving_target_amount_label"), size=12))
        self._target = QDoubleSpinBox()
        self._target.setRange(0, 1_000_000_000)
        self._target.setValue(self.saving["target_amount"] if self.saving and self.saving.get("target_amount") else 0)
        self._target.setPrefix(f"{self.currency_symbol()}: ")
        self._target.setMinimumHeight(42)
        lay.addWidget(self._target)

        # Target tanggal
        lay.addWidget(_lbl(tr("economy_saving_target_date_label"), size=12))
        self._target_date = QLineEdit()
        from datetime import date, timedelta
        if self.saving and self.saving.get("target_date"):
            self._target_date.setText(self.saving["target_date"])
        else:
            self._target_date.setText((date.today() + timedelta(days=90)).isoformat())
        self._target_date.setMinimumHeight(42)
        lay.addWidget(self._target_date)

        # Catatan
        lay.addWidget(_lbl(tr("dialog_notes"), size=12))
        self._notes = _input(tr("dialog_notes_placeholder"))
        if self.saving:
            self._notes.setText(self.saving.get("notes", ""))
        lay.addWidget(self._notes)

        lay.addSpacing(8)
        ok_lbl = tr("dialog_save") if self.saving else tr("economy_saving_add_btn")
        ok = _btn(ok_lbl, "solid", self._save, 46)
        lay.addWidget(ok)
        self._name.returnPressed.connect(self._save)

        root.addWidget(_scrolled(content))

    def _save(self):
        name = self._name.text().strip()
        if not name:
            _show(self, tr("msg_error"), tr("saving_name_empty"), "error")
            return
        icon = self._icon.currentData()
        
        current_user = self._current.value()
        target_user = self._target.value()
        
        user_currency = db.get_user_currency(self.user_id)
        current_idr = db.convert_to_idr(current_user, user_currency)
        target_idr = db.convert_to_idr(target_user, user_currency)
        
        target_date = self._target_date.text().strip()
        if target_date == "":
            target_date = None
        notes = self._notes.text()
        
        if self.saving:
            db.update_saving(self.saving["id"], self.user_id,
                            name=name, icon=icon, current_amount=current_idr,
                            target_amount=target_idr, target_date=target_date, notes=notes)
            SND.complete()
            _show(self, tr("berhasil_title"), "Tabungan berhasil diupdate!", "success")
        else:
            r = db.add_saving(self.user_id, name, icon, target_idr, current_idr, target_date, notes)
            if r["ok"]:
                SND.complete()
                _show(self, tr("berhasil_title"), f"Tabungan '{name}' ditambahkan!", "success")
            else:
                SND.error()
                _show(self, tr("gagal_title"), r["msg"], "error")
                return
        self.accept()

class AddToSavingDialog(QDialog):
    def __init__(self, saving_id, user_id, current_amount, parent=None):
        super().__init__(parent)
        
        self.saving_id = saving_id
        self.user_id = user_id
        self.current_amount = current_amount
        self.setWindowTitle(tr("saving_add_funds_title"))
        self.setMinimumSize(350, 200)
        self.setStyleSheet(build_ss())
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)
        layout.addWidget(_lbl(f"Jumlah yang ingin ditambahkan: ({self.currency_symbol()})", size=12))
        self.amount = QDoubleSpinBox()
        self.amount.setRange(0, 1_000_000_000)
        self.amount.setValue(10000)
        self.amount.setPrefix(f"{self.currency_symbol()}: ")
        self.amount.setMinimumHeight(42)
        layout.addWidget(self.amount)
        btn = _btn(tr("economy_saving_add_confirm_btn"), "solid", self._add, 40)
        layout.addWidget(btn)

    def currency_symbol(self):
        curr = db.get_user_currency(self.user_id)
        symbols = {"IDR": "Rp", "USD": "$", "EUR": "€"}
        return symbols.get(curr, "Rp")

    def _add(self):
        amt_user = self.amount.value()
        if amt_user <= 0:
            _show(self, tr("msg_error"), "Jumlah harus lebih dari 0", "error")
            return
        user_currency = db.get_user_currency(self.user_id)
        amt_idr = db.convert_to_idr(amt_user, user_currency)
        r = db.add_to_saving(self.saving_id, self.user_id, amt_idr)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), tr("berhasil_saving_add", symbol=self.currency_symbol(), amount=f"{amt_user:.0f}"), "success")
            self.accept()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

# =========== INVESTMENT DIALOG =========== #
class AddInvestmentDialog(QDialog):
    def __init__(self, user_id, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.setWindowTitle(tr("economy_invest_add"))
        self.setMinimumWidth(450)
        self.setMinimumHeight(420)
        self.setStyleSheet(build_ss())
        self._build()

    def currency_symbol(self):
        curr = db.get_user_currency(self.user_id)
        symbols = {"IDR": "Rp", "USD": "$", "EUR": "€"}
        return symbols.get(curr, "Rp")

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(14)

        lay.addWidget(_lbl(tr("investment_add_title"), "section", 14, True))
        lay.addWidget(_sep())

        lay.addWidget(_lbl(tr("economy_invest_name"), size=12))
        self.name = _input(tr("economy_invest_name_ph"))
        lay.addWidget(self.name)

        lay.addWidget(_lbl(tr("dialog_icon"), size=12))
        icons = [
            (tr("invest_icon_stock"), "📈"), 
            (tr("invest_icon_deposit"), "🏦"), 
            (tr("invest_icon_gold"), "💰"), 
            (tr("invest_icon_property"), "🏠"), 
            (tr("invest_icon_crypto"), "🚀")
        ]
        self.icon = _combo(icons)
        lay.addWidget(self.icon)

        lay.addWidget(_lbl(tr("economy_invest_amount_label", symbol=self.currency_symbol()), size=12))
        self.amount = QDoubleSpinBox()
        self.amount.setRange(1000, 1e9)
        self.amount.setValue(50000)
        self.amount.setPrefix(f"{self.currency_symbol()}: ")
        lay.addWidget(self.amount)

        lay.addWidget(_lbl(tr("dialog_notes"), size=12))
        self.notes = _input()
        lay.addWidget(self.notes)

        btn = _btn(tr("economy_invest_btn"), "solid", self._save)
        lay.addWidget(btn)
        root.addWidget(_scrolled(content))

    def _save(self):
        name = self.name.text().strip()
        if not name:
            _show(self, tr("msg_error"), tr("invest_name_empty"), "error")
            return
        amount_user = self.amount.value()
        if amount_user <= 0:
            _show(self, tr("msg_error"), tr("invest_amount_positive"), "error")
            return
        user_currency = db.get_user_currency(self.user_id)
        amount_idr = db.convert_to_idr(amount_user, user_currency)
        icon = self.icon.currentData()
        notes = self.notes.text()
        r = db.add_investment(self.user_id, name, icon, amount_idr, notes)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            self.accept()
        else:
            SND.error()
            msg = r["msg"]
            import re
            match = re.search(r"Saldo: (\d+)", msg)
            if match:
                saldo_idr = int(match.group(1))
                saldo_usr = db.convert_from_idr(saldo_idr, user_currency)
                saldo_str = f"{self.currency_symbol()} {saldo_usr:,.0f}"
                msg = msg.replace(f"Saldo: {saldo_idr}", f"Saldo: {saldo_str}")
            _show(self, tr("gagal_title"), msg, "error")

# ========== AddSubscriptionDialog ========== #
class AddSubscriptionDialog(QDialog):
    def __init__(self, user_id, sub=None, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.sub = sub
        self.setWindowTitle(tr("subscription_edit_title") if sub else tr("subscription_add_title"))
        self.setMinimumWidth(480)
        self.setMinimumHeight(520)
        self.setStyleSheet(build_ss())
        self._build()
        if self.sub:
            user_curr = db.get_user_currency(self.user_id)
            amount_idr = self.sub['amount']
            amount_usr = db.convert_from_idr(amount_idr, user_curr)
            self.amount.setValue(amount_usr)

    def currency_symbol(self):
        curr = db.get_user_currency(self.user_id)
        symbols = {"IDR": "Rp", "USD": "$", "EUR": "€"}
        return symbols.get(curr, "Rp")

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(14)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))
        lay.addWidget(_sep())

        # Nama
        lay.addWidget(_lbl(tr("economy_sub_name_label"), size=12))
        self.name = QLineEdit()
        self.name.setPlaceholderText(tr("sub_name_placeholder"))
        if self.sub:
            self.name.setText(self.sub["name"])
        lay.addWidget(self.name)

        # Icon
        lay.addWidget(_lbl(tr("dialog_icon"), size=12))
        icons = [
            (tr("sub_icon_tv"), "📺"), 
            (tr("sub_icon_music"), "🎵"), 
            (tr("sub_icon_ai"), "🤖"), 
            (tr("sub_icon_book"), "📚"), 
            (tr("sub_icon_gym"), "🏋️"), 
            (tr("sub_icon_general"), "📅")
        ]
        self.icon = _combo(icons)
        if self.sub:
            idx = self.icon.findData(self.sub.get("icon", "📅"))
            if idx >= 0:
                self.icon.setCurrentIndex(idx)
        lay.addWidget(self.icon)

        # Amount
        lay.addWidget(_lbl(tr("economy_sub_cost_label", symbol=self.currency_symbol()), size=12))
        self.amount = QDoubleSpinBox()
        self.amount.setRange(1000, 1e7)
        self.amount.setValue(self.sub["amount"] if self.sub else 10000)
        self.amount.setPrefix(f"{self.currency_symbol()}: ")
        lay.addWidget(self.amount)

        # Due date
        lay.addWidget(_lbl(tr("economy_sub_due_label"), size=12))
        self.due_date = QLineEdit()
        from datetime import date, timedelta
        default_date = (date.today() + timedelta(days=30)).isoformat() if not self.sub else self.sub["due_date"]
        self.due_date.setText(default_date)
        lay.addWidget(self.due_date)

        # Period
        lay.addWidget(_lbl(tr("economy_sub_period_label"), size=12))
        period_opts = [(tr("sub_period_monthly"), "monthly"), (tr("sub_period_yearly"), "yearly"), (tr("sub_period_onetime"), "one-time")]
        self.period = _combo(period_opts)
        if self.sub:
            idx = self.period.findData(self.sub["period"])
            if idx >= 0:
                self.period.setCurrentIndex(idx)
        lay.addWidget(self.period)
        self.period.currentIndexChanged.connect(self._on_period_changed)

        # Recurring checkbox
        self.recurring_cb = QCheckBox(tr("economy_sub_autorenew"))
        self.recurring_cb.setChecked(self.sub["is_recurring"] if self.sub else True)
        lay.addWidget(self.recurring_cb)

        # Notes
        lay.addWidget(_lbl(tr("dialog_notes"), size=12))
        self.notes = QLineEdit()
        self.notes.setPlaceholderText(tr("dialog_notes_placeholder"))
        if self.sub:
            self.notes.setText(self.sub.get("notes", ""))
        lay.addWidget(self.notes)

        btn = _btn(tr("dialog_save"), "solid", self._save)
        lay.addWidget(btn)
        root.addWidget(_scrolled(content))

    def _save(self):
        name = self.name.text().strip()
        if not name:
            _show(self, tr("msg_error"), "Nama layanan harus diisi", "error")
            return
        icon = self.icon.currentData()
        amount_user = self.amount.value()
        if amount_user <= 0:
            _show(self, tr("msg_error"), "Biaya harus lebih dari 0", "error")
            return
        user_currency = db.get_user_currency(self.user_id)
        amount_idr = db.convert_to_idr(amount_user, user_currency)
        
        due_date = self.due_date.text().strip()
        if not due_date:
            from datetime import date, timedelta
            due_date = (date.today() + timedelta(days=30)).isoformat()
        period = self.period.currentData()
        is_recurring = self.recurring_cb.isChecked()
        if period == 'one-time':
            is_recurring = False
        notes = self.notes.text()
        
        if self.sub:
            db.update_subscription(self.sub["id"], self.user_id,
                                name=name, icon=icon, amount=amount_idr, due_date=due_date,
                                period=period, is_recurring=is_recurring, notes=notes)
            SND.complete()
            _show(self, tr("berhasil_title"), "Subscription berhasil diupdate", "success")
        else:
            r = db.add_subscription(self.user_id, name, icon, amount_idr, due_date, period, is_recurring, notes)
            if r["ok"]:
                SND.complete()
                _show(self, tr("berhasil_title"), f"Subscription '{name}' ditambahkan", "success")
            else:
                SND.error()
                _show(self, tr("gagal_title"), r.get("msg", "Gagal menambah subscription"), "error")
                return
        self.accept()

    def _on_period_changed(self):
        period = self.period.currentData()
        if period == 'one-time':
            self.recurring_cb.setChecked(False)
            self.recurring_cb.setEnabled(False)
        else:
            self.recurring_cb.setEnabled(True)

# ══════════════════════════════════════════════════════════════════════════════
#  SHOP PAGE  (real buffs, use consumables)
# ══════════════════════════════════════════════════════════════════════════════
class ShopPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        
        self.user_id = user_id
        self._build()
        AppState.register(self.load)

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(20, 16, 20, 16)
        root.setSpacing(10)
        root.addWidget(_lbl(tr("shop_title"), "section", 14, True))
        root.addWidget(_sep())

        self._buff_bar = QLabel("")
        self._buff_bar.setWordWrap(True)
        self._buff_bar.setStyleSheet(
            f"background: {_T('panel')}; color: {_T('accent')};"
            f" border: 1px solid {_T('border')};"
            f" border-radius: 6px; padding: 8px 12px; font-size: 12px;")
        root.addWidget(self._buff_bar)

        self._tabs = QTabWidget()
        self._items_inner = QWidget()
        self._items_grid  = QGridLayout(self._items_inner)
        self._items_grid.setSpacing(8)
        self._pets_inner  = QWidget()
        self._pets_grid   = QGridLayout(self._pets_inner)
        self._pets_grid.setSpacing(8)
        self._tabs.addTab(_scrolled(self._items_inner), tr("shop_tab_items"))
        self._tabs.addTab(_scrolled(self._pets_inner),  tr("shop_tab_pets"))
        root.addWidget(self._tabs)
        self.load()

    def _clear_grid(self, grid: QGridLayout):
        while grid.count():
            item = grid.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

    def load(self):
        if not AppState.user_id:
            return
        self._clear_grid(self._items_grid)
        self._clear_grid(self._pets_grid)
        u   = AppState.user()
        inv = db.get_inventory(self.user_id)
        owned = {i["item_id"]: i for i in inv}

        # ── Buff summary ──────────────────────────────────────────────────────
        buffs = db.get_all_active_buffs(self.user_id)
        if buffs:
            self._buff_bar.setText("⚡ Buff Aktif :  " + "  ·  ".join(buffs))
        else:
            self._buff_bar.setText("⚡ Buff Aktif :  Tidak ada buff aktif.")

        # ── Items ─────────────────────────────────────────────────────────────
        # Filter: item craft_only disembunyikan; item seasonal hanya saat event
        visible_items = [(iid, item) for iid, item in db.SHOP_ITEMS.items()
                         if db.is_shop_item_visible(iid)]
        COLS = 4
        for idx, (iid, item) in enumerate(visible_items):
            f   = _card()
            cl  = QVBoxLayout(f)
            cl.setContentsMargins(10, 10, 10, 10)
            cl.setSpacing(4)

            cl.addWidget(QLabel(item["icon"],
                                alignment=Qt.AlignmentFlag.AlignCenter))
            nm = QLabel(item["name"])
            nm.setAlignment(Qt.AlignmentFlag.AlignCenter)
            nm.setStyleSheet(
                f"font-size:12px; font-weight:bold; color:{_T('text')};")
            cl.addWidget(nm)
            # Badge untuk item seasonal (event terbatas)
            if item.get("seasonal"):
                sb = QLabel(tr("shop_seasonal_badge"))
                sb.setAlignment(Qt.AlignmentFlag.AlignCenter)
                sb.setStyleSheet("font-size:9px; color:#2dd4bf; font-weight:bold;")
                cl.addWidget(sb)
            bd = QLabel(item["buff_desc"])
            bd.setAlignment(Qt.AlignmentFlag.AlignCenter)
            bd.setWordWrap(True)
            bd.setStyleSheet(f"font-size:10px; color:{_T('accent')};")
            cl.addWidget(bd)
            tp = QLabel(tr(f"shop_type_{item['type']}"))
            tp.setAlignment(Qt.AlignmentFlag.AlignCenter)
            tp.setStyleSheet(f"font-size:10px; color:{_T('muted')};")
            cl.addWidget(tp)

            if iid in owned:
                qty = owned[iid]["quantity"]
                inv_id = owned[iid]["id"]
                
                # Label "Dimiliki"
                ol = QLabel(tr("shop_owned"))
                ol.setAlignment(Qt.AlignmentFlag.AlignCenter)
                ol.setStyleSheet(f"color:{_T('light')}; font-size:11px; font-weight:bold;")
                cl.addWidget(ol)
                
                # Tampilkan tombol berdasarkan tipe
                if item["type"] == "consumable":
                    ub = _btn(tr("shop_use", qty=qty), "diamond", h=30)
                    ub.clicked.connect(lambda _, i=iid: self._use(i))
                    cl.addWidget(ub)

                    bb = _btn(tr("shop_buy_again"), "gold", h=30)
                    bb.clicked.connect(lambda _, i=iid: self._buy(i))
                    cl.addWidget(bb)

                    sell_price = int(item["cost"] * 0.1)
                    sell_price = max(1, sell_price)
                    sell_btn = _btn(tr("shop_sell"), h=30)
                    sell_btn.setMinimumWidth(60)
                    sell_btn.clicked.connect(lambda _, inv=owned[iid]["id"], name=item['name'], cost=item['cost'], qty=qty: 
                                            self._sell_item(inv, name, cost, qty))
                    cl.addWidget(sell_btn)
                    # Tampilkan harga jual
                    price_lbl = QLabel(tr("shop_sell_price", gold=sell_price))
                    price_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    price_lbl.setStyleSheet(f"color:{_T('muted')}; font-size:10px;")
                    cl.addWidget(price_lbl)
                else:
                    # Tombol Jual
                    sell_price = int(item["cost"] * 0.1)
                    sell_price = max(1, sell_price)
                    sell_btn = _btn(tr("shop_sell"), h=30)
                    sell_btn.setMinimumWidth(60)
                    sell_btn.clicked.connect(lambda _, inv=inv_id, name=item['name'], cost=item['cost'], qty=qty: 
                                            self._sell_item(inv, name, cost, qty))
                    cl.addWidget(sell_btn)
                    
                    # Tampilkan harga jual
                    price_lbl = QLabel(tr("shop_sell_price", gold=sell_price))
                    price_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    price_lbl.setStyleSheet(f"color:{_T('muted')}; font-size:10px;")
                    cl.addWidget(price_lbl)

                    # ✨ Enchanting (equipment saja)
                    elvl = owned[iid].get("enchant_level", 0) or 0
                    if elvl > 0:
                        el = QLabel(tr("enchant_level_tag", lvl=elvl))
                        el.setAlignment(Qt.AlignmentFlag.AlignCenter)
                        el.setStyleSheet("color:#a97fff; font-size:11px; font-weight:bold;")
                        cl.addWidget(el)
                    if elvl >= db.ENCHANT_MAX_LEVEL:
                        mx = QLabel(tr("enchant_max_tag"))
                        mx.setAlignment(Qt.AlignmentFlag.AlignCenter)
                        mx.setStyleSheet("color:#a97fff; font-size:10px; font-weight:bold;")
                        cl.addWidget(mx)
                    else:
                        cost_next = db.enchant_cost(elvl)
                        key = "enchant_btn" if elvl > 0 else "enchant_first_btn"
                        eb = _btn(tr(key, lvl=elvl + 1, cost=cost_next), h=30)
                        eb.clicked.connect(lambda _, i=iid: self._enchant(i))
                        cl.addWidget(eb)
            else:
                # Belum dimiliki: tampilkan harga dan tombol Beli
                cl.addWidget(QLabel(
                    f"💰 {item['cost']} G",
                    alignment=Qt.AlignmentFlag.AlignCenter))
                bb = _btn(tr("shop_buy"), "gold", h=30)
                bb.clicked.connect(lambda _, i=iid: self._buy(i))
                cl.addWidget(bb)
                
            self._items_grid.addWidget(f, idx // COLS, idx % COLS)

        # ── Pets ──────────────────────────────────────────────────────────────
        user_pets  = db.get_user_pets(self.user_id)
        owned_pets = {p["pet_id"] for p in user_pets}
        active_pets = {p["pet_id"] for p in user_pets if p["is_active"]}
        for idx, (pid, pet) in enumerate(db.PETS_DATA.items()):
            f   = _card()
            cl  = QVBoxLayout(f)
            cl.setContentsMargins(10, 10, 10, 10)
            cl.setSpacing(4)
            ico = _emoji_label(pet["icon"], ICON_CARD)
            ico.setAlignment(Qt.AlignmentFlag.AlignCenter)
            cl.addWidget(ico)
            nm = QLabel(pet["name"])
            nm.setAlignment(Qt.AlignmentFlag.AlignCenter)
            nm.setStyleSheet(
                f"font-size:12px; font-weight:bold; color:{_T('text')};")
            cl.addWidget(nm)
            bns = QLabel(pet["bonus"])
            bns.setAlignment(Qt.AlignmentFlag.AlignCenter)
            bns.setStyleSheet("font-size:10px; color:#4dd9e0;")
            bns.setWordWrap(True)
            cl.addWidget(bns)
            if pid in owned_pets:
                if pid in active_pets:
                    # Label AKTIF
                    al = QLabel(tr("shop_active"))
                    al.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    al.setStyleSheet("color:#4dd9e0; font-size:11px; font-weight:bold;")
                    cl.addWidget(al)
                    # Tombol Unequip
                    uq = _btn(tr("shop_unequip"), "danger", h=30)
                    uq.clicked.connect(lambda _, p=pid: self._unequip(p))
                    cl.addWidget(uq)
                else:
                    eq = _btn(tr("shop_equip"), "diamond", h=30)
                    eq.clicked.connect(lambda _, p=pid: self._equip(p))
                    cl.addWidget(eq)
            else:
                cl.addWidget(QLabel(
                    f"💰 {pet['cost']} G",
                    alignment=Qt.AlignmentFlag.AlignCenter))
                ab = _btn(tr("shop_adopt"), "gold", h=30)
                ab.clicked.connect(lambda _, p=pid: self._adopt(p))
                cl.addWidget(ab)
            self._pets_grid.addWidget(f, idx // 3, idx % 3)

    def _enchant(self, item_id):
        """Enchant equipment dengan XP (max +5, tiap level +12% kekuatan buff)."""
        elvl = db.get_enchant_level(self.user_id, item_id)
        cost = db.enchant_cost(elvl)
        u = AppState.user()
        if u.get("xp", 0) < cost:
            SND.error()
            _show(self, tr("gagal_title"),
                  tr("db_enchant_no_xp", cost=cost), "error")
            return
        r = db.enchant_item(self.user_id, item_id)
        if r["ok"]:
            SND.level_up()
            _show(self, tr("enchant_success_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()
        self.load()

    def _buy(self, iid):
        r = db.buy_item(self.user_id, iid)
        if r["ok"]:
            SND.buy()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()

    def _sell_item(self, inv_id, item_name, item_cost, max_qty):
        """Dialog untuk menjual item dengan pilihan quantity."""
        sell_price_per_item = int(item_cost * 0.1)
        sell_price_per_item = max(1, sell_price_per_item)
        
        # Dialog pilih quantity
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("shop_sell_title"))
        dlg.setMinimumSize(350, 200)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)
        
        layout.addWidget(QLabel(tr("shop_sell_confirm", name=item_name, qty=1, gold=sell_price_per_item)))
        
        # Spin untuk quantity
        qty_spin = QSpinBox()
        qty_spin.setRange(1, max_qty)
        qty_spin.setValue(1)
        qty_spin.setSuffix(f" (max {max_qty})")
        qty_spin.setMinimumHeight(42)
        layout.addWidget(qty_spin)
        
        # Label total harga
        total_label = QLabel(tr("shop_sell_price", gold=sell_price_per_item))
        total_label.setStyleSheet(f"color:{_T('accent')}; font-weight:bold;")
        layout.addWidget(total_label)
        
        def update_total(value):
            total = value * sell_price_per_item
            total_label.setText(tr("shop_sell_price", gold=total))

        qty_spin.valueChanged.connect(update_total)
        
        btn_layout = QHBoxLayout()
        cancel_btn = _btn(tr("btn_cancel"), "flat", dlg.reject)
        sell_btn = _btn(tr("shop_sell"), "gold", dlg.accept)
        btn_layout.addWidget(cancel_btn)
        btn_layout.addWidget(sell_btn)
        layout.addLayout(btn_layout)
        
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        
        qty = qty_spin.value()
        
        # Konfirmasi akhir
        total_gold = qty * sell_price_per_item
        reply = QMessageBox.question(
            self,
            tr("shop_sell_title"),
            tr("shop_sell_confirm", name=item_name, qty=qty, gold=total_gold),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        # Proses penjualan
        r = db.sell_item(self.user_id, inv_id, qty)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            AppState.refresh()
            self.load()       
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _use(self, iid):
        r = db.use_item(self.user_id, iid)
        if r.get("ok"):
            SND.complete()
            _show(self, tr("item_used_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"),
                  r.get("msg", "Item tidak bisa digunakan."), "error")
        AppState.refresh()

    def _adopt(self, pid):
        r = db.adopt_pet(self.user_id, pid)
        if r["ok"]:
            SND.buy()
            _show(self, tr("adopt_success"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()

    def _equip(self, pid):
        r = db.equip_pet(self.user_id, pid)
        SND.notify()
        _show(self, "Pet Aktif", r["msg"], "success")
        AppState.refresh()

    def _unequip(self, pet_id):
        r = db.unequip_pet(self.user_id, pet_id)
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()
        QTimer.singleShot(0, self.load)

    def closeEvent(self, e):
        AppState.unregister(self.load)
        super().closeEvent(e)


# ================================================================== #
class PetsPage(QWidget):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self._build()
        AppState.register(self.load)
        

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 16, 20, 16)
        layout.setSpacing(10)
        self.title = _lbl(tr("pets_title"), "section", 14, True)
        layout.addWidget(self.title)
        layout.addWidget(_sep())
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.container = QWidget()
        self.grid = QGridLayout(self.container)
        self.grid.setSpacing(12)
        # Atur alignment agar card tidak meregang
        self.grid.setAlignment(Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft)
        # Buat kolom dengan stretch yang sama (agar card terdistribusi merata)
        for col in range(3):
            self.grid.setColumnStretch(col, 1)
        self.grid.setColumnStretch(0, 1)
        self.grid.setColumnStretch(1, 1)
        self.grid.setColumnStretch(2, 1)
        self.scroll.setWidget(self.container)
        layout.addWidget(self.scroll)
        self.load()

    def load(self):
        # Bersihkan grid terlebih dahulu
        while self.grid.count():
            item = self.grid.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Ambil data pets dari database
        pets = db.get_user_pets(self.user_id)
        
        if not pets:
            empty = _lbl(tr("pets_empty"), "sub", 13)
            empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.grid.addWidget(empty, 0, 0)
            return

        # Hitung jumlah pet aktif & info batas
        active_count = len([p for p in pets if p["is_active"]])
        user_level = AppState.user().get("level", 1)
        max_pets = 2 if user_level >= 25 else 1
        status_text = tr("pets_max_2") if user_level >= 25 else tr("pets_max_1")
        
        # Buat widget info (ditampilkan sebagai card di atas daftar pet)
        info_widget = _card()
        info_layout = QVBoxLayout(info_widget)
        info_label = QLabel(tr("pets_active_info", active=active_count, max=max_pets, level=user_level, status=status_text))
        info_label.setStyleSheet(f"color:{_T('light')}; font-size:13px; font-weight:bold;")
        info_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        info_layout.addWidget(info_label)
        # Tambahkan info widget ke grid (baris 0, kolom 0, lebar 3 kolom)
        self.grid.addWidget(info_widget, 0, 0, 1, 3)

        # Mulai baris berikutnya untuk kartu pet (mulai dari baris 1)
        row = 1
        col = 0
        for i, p in enumerate(pets):
            pet_data = db.PETS_DATA.get(p["pet_id"], {})
            card = _card()
            card.setMinimumWidth(260)
            vlay = QVBoxLayout(card)
            vlay.setSpacing(4)
            vlay.setContentsMargins(8, 8, 8, 8)

            # Header baris
            top = QHBoxLayout()
            ico = _emoji_label(pet_data.get("icon", "🐾"), ICON_CARD)
            top.addWidget(ico)
            name = QLabel(pet_data.get("name", p["pet_id"]))
            name.setStyleSheet("font-size:13px; font-weight:bold;")
            top.addWidget(name, 1)
            if p["is_active"]:
                active_lbl = QLabel(tr("shop_active"))
                active_lbl.setStyleSheet("color:#80c000; font-size:10px;")
                top.addWidget(active_lbl)
            vlay.addLayout(top)

            # Level & EXP
            lvl_lbl = QLabel(tr("pets_level_label", level=p['level']))
            lvl_lbl.setStyleSheet("color:#f0a800; font-size:11px;")
            vlay.addWidget(lvl_lbl)
            exp_bar = QProgressBar()
            exp_needed = p["level"] * 100
            exp_bar.setMaximum(exp_needed)
            exp_bar.setValue(int(p["exp"]))
            exp_bar.setFormat(tr("pets_exp_format", exp=p['exp'], need=exp_needed))
            exp_bar.setMinimumHeight(8)
            vlay.addWidget(exp_bar)

            # Hunger
            hunger_bar = QProgressBar()
            hunger_bar.setMaximum(100)
            hunger_bar.setValue(int(p["hunger"]))
            hunger_bar.setFormat(tr("pets_hunger", hunger=p["hunger"]))
            hunger_bar.setMinimumHeight(12)
            hunger_bar.setStyleSheet("QProgressBar::chunk { background: #f0a800; }")
            vlay.addWidget(hunger_bar)

            # Buff
            base_buff = pet_data.get("base_buff", {})
            level = p["level"]
            scale = 1 + (level - 1) * 0.1   # Harus sama dengan scaling di database
            buff_lines = []
            if "xp_pct" in base_buff:
                val = base_buff['xp_pct'] * scale
                buff_lines.append(tr("pets_buff_xp_format", val=val))
            if "gold_pct" in base_buff:
                val = base_buff['gold_pct'] * scale
                buff_lines.append(tr("pets_buff_gold_format", val=val))
            if "boss_dmg" in base_buff:
                val = base_buff['boss_dmg'] * scale
                buff_lines.append(tr("pets_buff_dmg_format", val=val))
            if "hp_reduc" in base_buff:
                val = base_buff['hp_reduc'] * scale
                buff_lines.append(tr("pets_buff_reduc_format", val=val))
            if buff_lines:
                buff_label = QLabel(" | ".join(buff_lines))
                buff_label.setStyleSheet("font-size:10px; color:#4dd9e0;")
                vlay.addWidget(buff_label)

            # Tombol
            btn_row = QHBoxLayout()
            feed_cost = 30
            feed_btn = _btn(tr("pets_feed", cost=feed_cost), h=28)
            feed_btn.setMinimumWidth(70)
            feed_btn.clicked.connect(lambda _, pid=p["pet_id"]: self._feed(pid))
            train_cost = 25 + (p["level"] - 1) * 5
            train_btn = _btn(tr("pets_train", cost=train_cost), h=28)
            train_btn.setMinimumWidth(70)
            train_btn.clicked.connect(lambda _, pid=p["pet_id"]: self._train(pid))
            if p["is_active"]:
                unequip_btn = _btn(tr("shop_unequip"), "danger", h=28)
                unequip_btn.setMinimumWidth(80)
                unequip_btn.clicked.connect(lambda _, pid=p["pet_id"]: self._unequip(pid))
                btn_row.addWidget(unequip_btn)
            else:
                equip_btn = _btn(tr("shop_equip"), "diamond", h=28)
                equip_btn.setMinimumWidth(80)
                equip_btn.clicked.connect(lambda _, pid=p["pet_id"]: self._equip(pid))
                btn_row.addWidget(equip_btn)
            btn_row.addWidget(feed_btn)
            btn_row.addWidget(train_btn)
            vlay.addLayout(btn_row)

            # Tambahkan kartu ke grid (mulai dari row 1)
            self.grid.addWidget(card, row, col)
            col += 1
            if col >= 3:
                col = 0
                row += 1

    def _feed(self, pet_id):
        r = db.feed_pet(self.user_id, pet_id)
        if r["ok"]:
            msg = tr("pet_fed", name=r['name'], cost=r['cost'])
            _show(self, tr("berhasil_title"), msg, "success")
        else:
            _show(self, tr("gagal_title"), r["msg"], "error")
        QTimer.singleShot(0, self.load)
        AppState.refresh()

    def _train(self, pet_id):
        r = db.train_pet(self.user_id, pet_id)
        if r["ok"]:
            msg = tr("pet_trained", name=r['name'], exp=r['exp_gained'], cost=r['cost'])
            _show(self, tr("berhasil_title"), msg, "success")
            if r.get("leveled_up"):
                SND.level_up()
        else:
            _show(self, tr("gagal_title"), r["msg"], "error")
        QTimer.singleShot(0, self.load)
        AppState.refresh()

    def _equip(self, pet_id):
        r = db.equip_pet(self.user_id, pet_id)
        self._show_result(r)
        QTimer.singleShot(0, self.load)
        AppState.refresh()

    def _unequip(self, pet_id):
        r = db.unequip_pet(self.user_id, pet_id)
        self._show_result(r)
        QTimer.singleShot(0, self.load)
        AppState.refresh()

    def _show_result(self, r):
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def closeEvent(self, e):
        AppState.unregister(self.load)
        super().closeEvent(e)

# ══════════════════════════════════════════════════════════════════════════════
# GUILD / BOSS PAGE — FIXED
# ══════════════════════════════════════════════════════════════════════════════
class GuildPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        _guild_scroll = QScrollArea()
        _guild_scroll.setWidgetResizable(True)
        _guild_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        _guild_container = QWidget()
        self._root = QVBoxLayout(_guild_container)
        self._root.setContentsMargins(20, 16, 20, 16)
        self._root.setSpacing(10)
        _guild_scroll.setWidget(_guild_container)
        _outer = QVBoxLayout(self)
        _outer.setContentsMargins(0, 0, 0, 0)
        _outer.addWidget(_guild_scroll)
        self._boss_info = None
        self._tier_cb = None
        self._boss_cb = None
        AppState.register(self.load)
        self.load()

    def _clear(self):
        while self._root.count():
            item = self._root.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

    # ---------- KOMPONEN DASHBOARD ----------
    def _make_header(self, guild: dict, user: dict) -> QWidget:
        w = QWidget()
        layout = QHBoxLayout(w)
        layout.setContentsMargins(0, 0, 0, 0)

        info = QVBoxLayout()
        info.setSpacing(2)
        name_label = QLabel(tr("guild_name_header", name=guild.get('name', 'Guild')))
        name_label.setStyleSheet(f"font-size:18px; font-weight:bold; color:{_T('light')};")
        info.addWidget(name_label)
        id_label = QLabel(tr("guild_id_level", id=guild.get('id', '?'), level=guild.get('level', 1)))
        id_label.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        info.addWidget(id_label)
        if guild.get('description'):
            desc = QLabel(tr("guild_description_label", desc=guild['description']))
            desc.setWordWrap(True)
            desc.setStyleSheet(f"color:{_T('text')}; font-size:12px;")
            info.addWidget(desc)

        layout.addLayout(info, 1)

        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(6)

        chat_btn = _btn(tr("guild_chat_title"), "diamond", self._open_guild_chat, 34)
        chat_btn.setMinimumWidth(80)
        btn_layout.addWidget(chat_btn)

        if user["id"] == guild.get("leader_id"):
            edit_btn = _btn(tr("guild_edit_desc_btn"), "flat", self._edit_guild_bio, 34)
            btn_layout.addWidget(edit_btn)

        leave_btn = _btn(tr("guild_leave_btn"), "danger", self._leave, 34)
        leave_btn.setMinimumWidth(80)
        btn_layout.addWidget(leave_btn)

        layout.addLayout(btn_layout)
        return w

    def _make_stats(self, guild: dict, members: list) -> QWidget:
        w = QWidget()
        grid = QGridLayout(w)
        grid.setSpacing(10)

        level = guild.get('level', 1)
        exp = guild.get('exp', 0)
        need = level * 500
        member_count = len(members)
        buff_xp = guild.get('buff_xp', 0)
        buff_gold = guild.get('buff_gold', 0)
        buff_damage = guild.get('buff_damage', 0)
        crit_chance = guild.get('crit_chance', 0)

        cards = [
            (tr("guild_stats_level"), tr("guild_stats_level_value", level=level), "#80c000"),
            (tr("guild_stats_members"), tr("guild_stats_members_value", count=member_count), "#4da6ff"),
            (tr("guild_stats_bonus_xp"), tr("guild_stats_bonus_xp_value", xp=buff_xp), "#f0a800"),
            (tr("guild_stats_bonus_gold"), tr("guild_stats_bonus_gold_value", gold=buff_gold), "#f0a800"),
            (tr("guild_stats_bonus_damage"), tr("guild_stats_bonus_damage_value", damage=buff_damage), "#e05050"),
            (tr("guild_stats_bonus_crit"), tr("guild_stats_bonus_crit_value", crit=crit_chance), "#a97fff"),
        ]

        for i, (title, value, color) in enumerate(cards):
            card = self._stat_card(title, value, color)
            grid.addWidget(card, i // 3, i % 3)

        exp_card = _card()
        exp_layout = QVBoxLayout(exp_card)
        exp_layout.setContentsMargins(12, 10, 12, 10)
        exp_layout.addWidget(_lbl(tr("guild_exp_progress", exp=exp, need=need), size=12))
        pb = QProgressBar()
        pb.setMaximum(need)
        pb.setValue(int(exp))
        pb.setMinimumHeight(16)
        pb.setStyleSheet("QProgressBar::chunk { background: #80c000; }")
        exp_layout.addWidget(pb)
        grid.addWidget(exp_card, len(cards)//3 + 1, 0, 1, 3)

        return w

    def _stat_card(self, title: str, value: str, color: str) -> QFrame:
        card = _card()
        card.setMinimumHeight(70)
        layout = QVBoxLayout(card)
        layout.setContentsMargins(10, 8, 10, 8)
        lbl_title = QLabel(title)
        lbl_title.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        lbl_value = QLabel(value)
        lbl_value.setStyleSheet(f"color:{color}; font-size:18px; font-weight:bold;")
        layout.addWidget(lbl_title)
        layout.addWidget(lbl_value)
        return card

    def _make_members_section(self, guild: dict, members: list, user: dict) -> QGroupBox:
        group = QGroupBox(tr("guild_members", count=len(members)))
        layout = QGridLayout(group)
        layout.setSpacing(10)

        leader = None
        regular = []
        for m in members:
            if m["id"] == guild.get("leader_id"):
                leader = m
            else:
                regular.append(m)

        row = 0
        if leader:
            leader_card = self._make_member_card(leader, user, guild, is_leader=True)
            layout.addWidget(leader_card, row, 0, 1, 2)
            row += 1

        for i, m in enumerate(regular):
            card = self._make_member_card(m, user, guild, is_leader=False)
            layout.addWidget(card, row + i // 4, i % 4)
        return group

    def _make_member_card(self, member: dict, current_user: dict, guild: dict, is_leader: bool) -> QFrame:
        f = _card()
        f.setMinimumWidth(120)
        f.setMaximumHeight(130)
        f.setStyleSheet(
            f"QFrame#card {{ background:{_T('panel')}; "
            f"border:1px solid {'#f0a800' if is_leader else _T('border')}; "
            f"border-radius:6px; }}"
        )
        layout = QVBoxLayout(f)
        layout.setContentsMargins(6, 5, 6, 5)
        layout.setSpacing(2)

        top = QHBoxLayout()
        avatar = _emoji_label(member.get('avatar_emoji', '⚔️'), ICON_CARD)
        top.addWidget(avatar)
        name = QLabel(member['display_name'])
        name.setStyleSheet(f"font-size:12px; font-weight:bold; color:{_T('text')};")
        top.addWidget(name, 1)
        if is_leader:
            crown = QLabel("👑")
            crown.setStyleSheet("font-size:14px;")
            top.addWidget(crown)
        layout.addLayout(top)

        level_label = QLabel(tr("level_abbr", level=member['level']))
        level_label.setStyleSheet(f"color:{_T('muted')}; font-size:10px;")
        layout.addWidget(level_label)

        hp_pb = QProgressBar()
        hp_pb.setMaximum(int(member['max_hp']))
        hp_pb.setValue(int(member['hp']))
        hp_pb.setMinimumHeight(6)
        hp_pb.setTextVisible(False)
        hp_pb.setStyleSheet(
            f"QProgressBar::chunk {{ background:{'#7bbf3e' if member['hp'] > 0 else '#e05050'}; }}"
        )
        layout.addWidget(hp_pb)

        if current_user["id"] == guild.get("leader_id") and member["id"] != current_user["id"]:
            btn_row = QHBoxLayout()
            btn_row.setSpacing(4)
            kick_btn = _btn(tr("guild_kick"), "danger", h=24)
            kick_btn.setMinimumWidth(45)
            kick_btn.clicked.connect(lambda _, uid=member["id"]: self._kick_member(uid))
            transfer_btn = _btn(tr("guild_transfer"), "gold", h=24)
            transfer_btn.setMinimumWidth(60)
            transfer_btn.clicked.connect(lambda _, uid=member["id"]: self._transfer_leadership(uid))
            btn_row.addWidget(kick_btn)
            btn_row.addWidget(transfer_btn)
            layout.addLayout(btn_row)

        return f

    def _make_boss_section(self, boss: dict, guild: dict, user: dict) -> QGroupBox:
        group = QGroupBox(tr("guild_boss_battle"))
        layout = QVBoxLayout(group)

        if boss:
            tier_color = db.BOSS_TIER_COLOR.get(boss.get('boss_tier', 'normal'), '#f0a800')
            title = QLabel(tr("boss_title_format", icon=boss['boss_icon'], name=boss['boss_name'], tier=boss.get('boss_tier', '?').upper()))
            title.setStyleSheet(f"font-size:15px; font-weight:bold; color:{tier_color};")
            layout.addWidget(title)

            hp_label = QLabel(tr("guild_boss_hp", hp=boss['boss_hp'], max_hp=boss['boss_max_hp']))
            hp_label.setStyleSheet(f"color:{_T('text')};")
            layout.addWidget(hp_label)

            pb = QProgressBar()
            pb.setMaximum(int(boss['boss_max_hp']))
            pb.setValue(int(boss['boss_hp']))
            pb.setMinimumHeight(20)
            pb.setStyleSheet(f"QProgressBar::chunk {{ background:{tier_color}; border-radius:6px; }}")
            layout.addWidget(pb)

            dmg_bonus = user.get('boss_damage_bonus', 0)
            atk_info = QLabel(tr("guild_boss_atk_info", atk=boss['boss_attack'], bonus=dmg_bonus, total=25+dmg_bonus))
            atk_info.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
            layout.addWidget(atk_info)

            # Info ultimate
            cls = user.get("avatar_class", "warrior")
            ult_name = tr(f"boss_ultimate_name_{cls}")
            ult_label = QLabel(tr("raid_ultimate_label", name=ult_name))
            ult_label.setStyleSheet(f"color:{_T('accent')}; font-weight:bold;")
            layout.addWidget(ult_label)

            if user["hp"] <= 0:
                warn = QLabel(tr("guild_hp_zero"))
                warn.setWordWrap(True)
                warn.setStyleSheet("color:#e05050; font-weight:bold;")
                layout.addWidget(warn)
                heal_btn = _btn(tr("guild_quick_heal"), "gold", self._quick_heal, 40)
                layout.addWidget(heal_btn)
            else:
                # ── Tombol aksi baru (Light, Heavy, Block, Ultimate) ──
                action_layout = QGridLayout()
                action_layout.setSpacing(10)
                action_layout.setContentsMargins(0, 10, 0, 0)

                # Light Attack
                light_btn = _btn(tr("boss_action_light_label"), slot=lambda: self._perform_action("light"), h=44)
                light_btn.setToolTip(tr("boss_action_light_tip"))
                action_layout.addWidget(light_btn, 0, 0)

                # Heavy Attack (15 MP)
                heavy_btn = _btn(tr("boss_action_heavy_label"), slot=lambda: self._perform_action("heavy"), h=44)
                heavy_btn.setToolTip(tr("boss_action_heavy_tip"))
                action_layout.addWidget(heavy_btn, 0, 1)

                # Block
                block_btn = _btn(tr("boss_action_block_label"), slot=lambda: self._perform_action("block"), h=44)
                block_btn.setToolTip(tr("boss_action_block_tip"))
                action_layout.addWidget(block_btn, 1, 0)

                # Ultimate (50 MP)
                ultimate_btn = _btn(tr("boss_action_ultimate_label"), slot=lambda: self._perform_action("ultimate"), h=44)
                ultimate_btn.setToolTip(tr("boss_action_ultimate_tip"))
                action_layout.addWidget(ultimate_btn, 1, 1)

                layout.addLayout(action_layout)
        else:
            layout.addWidget(_lbl(tr("guild_boss_none"), "sub", 12))
            self._boss_selector(layout, guild, user)

        return group

    def _boss_selector(self, lay: QVBoxLayout, guild: dict, user: dict):
        # Simpan konteks untuk refresh setelah custom boss dibuat
        self._cboss_guild_id = guild.get("id")
        self._cboss_user_level = user.get("level", 1)

        tier_row = QHBoxLayout()
        tier_row.addWidget(_lbl(tr("guild_boss_filter"), size=12))
        tier_opts = [(tr("guild_boss_all"), "all")] \
            + [(t.title(), t) for t in db.BOSS_TIER_ORDER] \
            + [(tr("cboss_custom_tag").title(), "custom")]
        self._tier_cb = _combo(tier_opts)
        self._tier_cb.currentIndexChanged.connect(lambda: self._fill_boss_cb(user["level"]) or self._update_boss_info())
        tier_row.addWidget(self._tier_cb)
        cboss_btn = _btn(tr("cboss_btn"), "gold", h=32,
                         slot=self._open_custom_boss)
        tier_row.addWidget(cboss_btn)
        tier_row.addStretch()
        lay.addLayout(tier_row)

        self._boss_cb = QComboBox()
        self._boss_cb.setMinimumHeight(42)
        self._boss_info = QLabel("")
        self._boss_info.setWordWrap(True)
        self._boss_info.setTextFormat(Qt.TextFormat.RichText)

        self._fill_boss_cb(user["level"])
        self._boss_cb.currentIndexChanged.connect(self._update_boss_info)
        lay.addWidget(self._boss_cb)
        lay.addWidget(self._boss_info)

        is_leader = (user["id"] == guild.get("leader_id"))
        start_btn = _btn(tr("guild_start_boss"), "danger" if is_leader else "flat", h=44)
        if is_leader:
            start_btn.clicked.connect(self._start_boss)
        else:
            start_btn.setEnabled(False)
            start_btn.setText(tr("guild_only_leader"))
        lay.addWidget(start_btn)

        # Tambahkan info spyglass jika dimiliki
        u = AppState.user()
        if u.get("has_spyglass", 0):
            info = QLabel(tr("guild_spyglass_active"))
            info.setStyleSheet(f"color:{_T('accent')}; font-size:11px; font-style:italic;")
            lay.addWidget(info)
            # Tampilkan boss stats lebih detail
            if self._boss_cb and self._boss_cb.count() > 0:
                self._boss_info.setText(
                    self._boss_info.text() + "\n🔭 " + tr("guild_spyglass_detail")
                )

    def _make_actions(self, user: dict, guild: dict) -> QWidget:
        w = QWidget()
        layout = QHBoxLayout(w)
        layout.setContentsMargins(0, 0, 0, 0)

        skill = db.CLASS_SKILLS.get(user.get("avatar_class", "warrior"), {})
        skill_text = tr("guild_skill_info",
                        mp=user['mp'], max_mp=user['max_mp'],
                        skill_icon=skill.get('icon', '❓'),
                        skill_name=skill.get('name', 'Unknown'),
                        skill_cost=skill.get('mp_cost', 0))
        skill_label = QLabel(skill_text)
        skill_label.setStyleSheet(f"color:{_T('text')}; font-size:12px;")
        layout.addWidget(skill_label, 1)

        skill_btn = _btn(tr("guild_use_skill_with_icon", icon=skill.get('icon', '⚡')), "diamond", self._skill, 36)
        skill_btn.setMinimumWidth(140)
        layout.addWidget(skill_btn)

        return w

    def _add_requests_and_invites(self, parent_layout: QVBoxLayout, guild: dict, user: dict):
        invites = db.get_guild_invites(self.user_id)
        if invites:
            invite_group = QGroupBox(tr("guild_invites"))
            vlay = QVBoxLayout(invite_group)
            for inv in invites:
                row = QHBoxLayout()
                row.addWidget(QLabel(tr("guild_invite_from", name=inv['guild_name'])))
                accept_btn = _btn(tr("guild_accept"), h=28)
                accept_btn.clicked.connect(lambda _, iid=inv["id"]: self._accept_invite(iid))
                reject_btn = _btn(tr("guild_reject"), "danger", h=28)
                reject_btn.clicked.connect(lambda _, iid=inv["id"]: self._reject_invite(iid))
                row.addWidget(accept_btn)
                row.addWidget(reject_btn)
                vlay.addLayout(row)
            parent_layout.addWidget(invite_group)

        if user["id"] == guild.get("leader_id"):
            requests = db.get_guild_requests(guild["id"])
            if requests:
                req_group = QGroupBox(tr("guild_join_requests"))
                vlay = QVBoxLayout(req_group)
                for req in requests:
                    row = QHBoxLayout()
                    row.addWidget(QLabel(tr("guild_join_request_format", name=req['display_name'], username=req['username'])))
                    accept_btn = _btn(tr("guild_accept"), h=28)
                    accept_btn.clicked.connect(lambda _, rid=req["id"]: self._accept_join(rid))
                    reject_btn = _btn(tr("guild_reject"), "danger", h=28)
                    reject_btn.clicked.connect(lambda _, rid=req["id"]: self._reject_join(rid))
                    row.addWidget(accept_btn)
                    row.addWidget(reject_btn)
                    vlay.addLayout(row)
                parent_layout.addWidget(req_group)

        conn = db.get_conn()
        transfers = conn.execute(
            "SELECT * FROM guild_leader_transfers WHERE guild_id=? AND status='pending'",
            (guild["id"],)
        ).fetchall()
        conn.close()
        if transfers:
            trans_group = QGroupBox(tr("guild_leader_inherit_box"))
            vlay = QVBoxLayout(trans_group)
            for t in transfers:
                row = QHBoxLayout()
                row.addWidget(QLabel(tr("guild_leader_inherit_msg")))
                accept_btn = _btn(tr("guild_accept"), h=28)
                accept_btn.clicked.connect(lambda _, tid=t["id"]: self._accept_leader(tid))
                row.addWidget(accept_btn)
                vlay.addLayout(row)
            parent_layout.addWidget(trans_group)

    # ---------- LOAD ----------
    def load(self):
        if not AppState.user_id:
            return
        if AppState.user().get("is_admin", 0):
            self._clear()
            msg = QLabel(tr("guild_admin_warning"))
            msg.setAlignment(Qt.AlignmentFlag.AlignCenter)
            msg.setStyleSheet(f"color:{_T('muted')}; font-size:14px; padding:40px;")
            self._root.addWidget(msg)
            return
        self._clear()
        u = AppState.user()
        gid = u.get("guild_id")
        if not gid:
            self._no_guild()
            return

        data = db.get_guild(gid)
        guild = data.get("guild", {})
        members = data.get("members", [])
        boss = data.get("boss")

        dashboard = QWidget()
        dashboard_layout = QVBoxLayout(dashboard)
        dashboard_layout.setSpacing(12)

        dashboard_layout.addWidget(self._make_header(guild, u))
        dashboard_layout.addWidget(self._make_stats(guild, members))
        dashboard_layout.addWidget(self._make_members_section(guild, members, u))
        dashboard_layout.addWidget(self._make_boss_section(boss, guild, u))
        dashboard_layout.addWidget(self._make_actions(u, guild))
        self._add_requests_and_invites(dashboard_layout, guild, u)

        # Tampilkan hadiah boss yang belum diklaim
        self._show_unclaimed_rewards(dashboard_layout)

        self._root.addWidget(dashboard)
        self._root.addStretch()

    # ---------- ACTION METHODS (tetap dari kode lama, hanya beberapa disesuaikan) ----------
    def _no_guild(self):
        self._root.addWidget(_lbl(tr("guild_no_guild"), "sub", 13))
        self._root.addSpacing(8)

        cg = QGroupBox(tr("guild_create"))
        cl = QVBoxLayout(cg)
        n_in = _input(tr("guild_name"))
        d_in = _input(tr("guild_desc"))
        cl.addWidget(_lbl(tr("dialog_name"), size=12)); cl.addWidget(n_in)
        cl.addWidget(_lbl(tr("guild_desc"), size=12)); cl.addWidget(d_in)
        def _create():
            n = n_in.text().strip()
            if n:
                r = db.create_guild(self.user_id, n, d_in.text())
                SND.notify()
                _show(self, tr("guild_created_title"), r["msg"], "success")
                AppState.refresh()
        cl.addWidget(_btn(tr("guild_create_btn"), "solid", _create, 40))
        self._root.addWidget(cg)

        rg = QGroupBox(tr("guild_request"))
        rl = QVBoxLayout(rg)
        sp = QSpinBox()
        sp.setRange(1, 99999)
        sp.setMinimumHeight(42)
        rl.addWidget(sp)
        def _request():
            r = db.send_guild_request(self.user_id, sp.value())
            if r["ok"]:
                SND.notify()
                _show(self, tr("berhasil_title"), r["msg"], "success")
            else:
                SND.error()
                _show(self, tr("gagal_title"), r["msg"], "error")
        rl.addWidget(_btn(tr("guild_request_btn"), "solid", _request, 40))
        self._root.addWidget(rg)

    def _show_unclaimed_rewards(self, parent_layout: QVBoxLayout):
        rewards = db.get_unclaimed_boss_rewards(self.user_id)
        if not rewards:
            return
        group = QGroupBox(tr("guild_unclaimed_rewards"))
        vlay = QVBoxLayout(group)
        for r in rewards:
            row = QHBoxLayout()
            row.addWidget(QLabel(tr("guild_reward_format", name=r['boss_name'], xp=r['xp_reward'], gold=r['gold_reward'])))
            claim_btn = _btn(tr("guild_claim"), "solid", h=30)
            claim_btn.clicked.connect(lambda _, rid=r["id"]: self._claim_reward(rid))
            row.addWidget(claim_btn)
            vlay.addLayout(row)
        parent_layout.addWidget(group)

    def _claim_reward(self, reward_id):
        r = db.claim_boss_reward(reward_id, self.user_id)
        if r["ok"]:
            SND.complete()
            if r.get("leveled_up"):
                SND.level_up()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()
        QTimer.singleShot(0, self.load)

    def _open_guild_chat(self):
        if AppState.user().get("is_admin", 0):
            _show(self, tr("info_title"), tr("guild_chat_admin_block"), "warning")
            return
        u = AppState.user()
        gid = u.get("guild_id")
        if not gid: return
        dlg = GuildChatDialog(gid, self.user_id, self)
        dlg.exec()

    def _edit_guild_bio(self):
        from PyQt6.QtWidgets import QInputDialog
        u = AppState.user()
        gid = u.get("guild_id")
        if not gid: return
        guild_data = db.get_guild(gid)
        old_desc = guild_data.get("guild", {}).get("description", "")
        new_desc, ok = QInputDialog.getMultiLineText(self, tr("guild_decs_edit_title"), tr("guild_new_desc"), old_desc)
        if ok and new_desc != old_desc:
            db.update_guild(gid, description=new_desc)
            SND.notify()
            _show(self, tr("berhasil_title"), tr("guild_desc_done"), "success")
            QTimer.singleShot(0, self.load)

    def _leave(self):
        r = db.leave_guild_with_transfer(self.user_id)
        SND.click()
        _show(self, tr("guild_title"), r["msg"])
        AppState.refresh()
        QTimer.singleShot(0, self.load)

    def _kick_member(self, target_id):
        u = AppState.user()
        r = db.kick_guild_member(u.get("guild_id"), u["id"], target_id)
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        QTimer.singleShot(0, self.load)

    def _transfer_leadership(self, new_leader_id):
        reply = QMessageBox.question(self, tr("guild_transfer_title"), tr("guild_transfer_confirm_msg"),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        u = AppState.user()
        guild_id = u.get("guild_id")
        if not guild_id:
            _show(self, tr("msg_error"), tr("db_guild_leave_not_in"), "error")
            return
        loading = LoadingDialog("Memproses transfer...", self)
        loading.show()
        QApplication.processEvents()
        r = db.transfer_guild_leadership(guild_id, u["id"], new_leader_id)
        loading.accept()
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            AppState.refresh()
            QTimer.singleShot(0, self.load)
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _fill_boss_cb(self, user_level: int):
        if not hasattr(self, "_boss_cb") or self._boss_cb is None:
            return
        self._boss_cb.blockSignals(True)
        self._boss_cb.clear()
        tier = self._tier_cb.currentData() if self._tier_cb else "all"
        guild_id = getattr(self, "_cboss_guild_id", None)
        for bid, bd in db.get_all_bosses_for_guild(guild_id).items():
            if tier != "all" and bd["tier"] != tier:
                continue
            # Boss seasonal hanya muncul saat event-nya aktif
            if not db.is_boss_available(bid):
                continue
            lock = "🔒 " if user_level < bd.get("min_level", 1) else ""
            self._boss_cb.addItem(
                tr("guild_boss_selector_item", lock=lock, icon=bd['icon'], name=bd['name'], tier=bd['tier'].upper(), hp=bd['hp'], min_level=bd['min_level']),
                bid)
        self._boss_cb.blockSignals(False)

    def _open_custom_boss(self):
        """Dialog buat custom boss untuk guild ini."""
        dlg = CustomBossDialog(self.user_id,
                               getattr(self, "_cboss_guild_id", None), self)
        if dlg.exec() == QDialog.DialogCode.Accepted and getattr(dlg, "created", False):
            lvl = getattr(self, "_cboss_user_level", 1)
            self._fill_boss_cb(lvl)
            self._update_boss_info()

    def _update_boss_info(self):
        """Update info boss di label. Jika tidak punya spyglass, hanya tampilkan nama dan tier."""
        if not isinstance(self._boss_info, QLabel):
            return
        if not self._boss_cb or self._boss_cb.count() == 0:
            return
        bid = self._boss_cb.currentData()
        bd = db.get_effective_boss(bid) or {}
        u = AppState.user()
        has_spyglass = u.get("has_spyglass", 0)
        ok = u.get("level", 1) >= bd.get("min_level", 1)
        
        if has_spyglass:
            # Tampilkan detail lengkap dengan HP, ATK, XP, Gold
            tc = db.BOSS_TIER_COLOR.get(bd.get("tier", "normal"), "#f0a800")
            text = tr("guild_boss_info_format",
                color=tc,
                icon=bd.get('icon', '🐉'),
                name=bd.get('name', 'Unknown'),
                tier=bd.get('tier', 'normal').upper(),
                hp=bd.get('hp', 0),
                atk=bd.get('atk', 0),
                xp=bd.get('xp', 0),
                gold=bd.get('gold', 0),
                min_level=bd.get('min_level', 1),
                ok="✅" if ok else "🔒")
            text += f"\n🔭 {tr('guild_spyglass_detail')}"
        else:
            # Tanpa spyglass: hanya nama, tier, dan status level
            tier_color = db.BOSS_TIER_COLOR.get(bd.get("tier", "normal"), "#f0a800")
            text = (f"<span style='color:{tier_color};font-weight:bold;'>"
                    f"{bd.get('icon', '🐉')} {bd.get('name', 'Unknown')} "
                    f"[{bd.get('tier', 'normal').upper()}]</span>\n")
            text += f"Min Level: {bd.get('min_level', 1)}  { '✅' if ok else '🔒' }"
            text += "\n\n🔒 Beli Spyglass di Shop untuk melihat statistik boss!"
        
        self._boss_info.setText(text)

    def _start_boss(self):
        u   = AppState.user()
        bid = self._boss_cb.currentData()
        r = db.start_boss(u.get("guild_id"), bid, u)
        if r["ok"]:
            SND.boss_hit()
            _show(self, tr("boss_appear_title"), r["msg"], "warning")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()

    def _attack(self):
        u = AppState.user()
        if u["hp"] <= 0:
            SND.error()
            _show(self, tr("hp_habis_title"), tr("guild_hp_zero_msg"), "error")
            return
        r = db.attack_boss(self.user_id, u.get("guild_id"), 25)
        if not r.get("ok"):
            SND.error()
            _show(self, tr("guild_cant_attack"), r["msg"], "error")
            return
        if r.get("defeated"):
            SND.boss_dead()
            _show(self, tr("victory_title"), r["msg"], "success")
            QTimer.singleShot(0, self.load)
        else:
            SND.boss_hit()
            u2 = AppState.user()
            has_spyglass = u2.get("has_spyglass", 0)
            
            user_dmg = r.get("user_damage", 0)
            user_crit = r.get("user_critical", False)
            user_dmg_text = f"⚔️ {user_dmg} DMG" + (" ⚡CRITICAL!" if user_crit else "")
            
            boss_dmg = r.get("boss_damage", 0)
            boss_crit = r.get("boss_critical", False)
            boss_dmg_text = f"💥 {boss_dmg} DMG" + (" ⚡CRITICAL!" if boss_crit else "")
            
            if has_spyglass:
                msg = f"{user_dmg_text}\nBoss HP: {r.get('boss_hp_left',0):.0f}/{r.get('boss_max_hp',0):.0f}\n"
                msg += f"Boss menyerang balik: {boss_dmg_text}\n❤️ HP kamu: {u2['hp']}"
            else:
                actual_damage = r.get('actual_damage', 0)
                msg = f"{user_dmg_text}\n"
                msg += f"Kamu kehilangan {actual_damage} HP!\n❤️ HP kamu: {u2['hp']}"
            
            reduc = u2.get("hp_damage_reduction", 0)
            if reduc > 0 and not has_spyglass:
                msg += f"\n🛡️ Damage reduction: -{reduc:.0f}"
            elif reduc > 0 and has_spyglass:
                msg += f" (reduksi: -{reduc:.0f})"
            
            if r.get("revived"):
                msg += "\n🗿 Totem menyelamatkanmu!"
            
            _show(self, tr("attack_title"), msg, "info" if not r.get("revived") else "success")
        AppState.refresh()

    def _quick_heal(self):
        r = db.use_item(self.user_id, "golden_apple")
        if r.get("ok"):
            SND.complete()
            _show(self, tr("hp_restored_title"), r["msg"], "success")
        else:
            _show(self, tr("no_item_title"), tr("no_golden_apple"), "warning")
        AppState.refresh()

    def _skill(self):
        r = db.use_class_skill(self.user_id)
        if r["ok"]:
            SND.notify()
            _show(self, tr("skill_used_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()

    def _accept_invite(self, invite_id):
        r = db.accept_invite(self.user_id, invite_id)
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()
        QTimer.singleShot(0, self.load)

    def _reject_invite(self, invite_id):
        r = db.reject_invite(self.user_id, invite_id)
        if r["ok"]:
            SND.notify()
            _show(self, tr("info_title"), r["msg"], "info")
        QTimer.singleShot(0, self.load)

    def _accept_join(self, request_id):
        u = AppState.user()
        r = db.accept_guild_request(u.get("guild_id"), u["id"], request_id)
        self._show_result(r)
        QTimer.singleShot(0, self.load)

    def _reject_join(self, request_id):
        u = AppState.user()
        r = db.reject_guild_request(u.get("guild_id"), u["id"], request_id)
        self._show_result(r)
        QTimer.singleShot(0, self.load)

    def _accept_leader(self, transfer_id):
        r = db.accept_leader_transfer(self.user_id, transfer_id)
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        QTimer.singleShot(0, self.load)

    def _show_result(self, r):
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _perform_action(self, action):
        """Panggil attack_boss dengan action tertentu."""
        u = AppState.user()
        if u["hp"] <= 0:
            SND.error()
            _show(self, tr("hp_habis_title"), tr("guild_hp_zero_msg"), "error")
            return

        # ── CEK MP UNTUK SEMUA AKSI YANG MEMBUTUHKAN MP ──
        if action in ("heavy", "ultimate", "block"):
            if action == "heavy":
                mp_cost = 5
            elif action == "ultimate":
                mp_cost = 50
            else:  # block
                mp_cost = 5
            if u.get("mp", 0) < mp_cost:
                SND.error()
                _show(self, tr("db_mp_insufficient_title"), tr("db_mp_insufficient_msg", cost=mp_cost, mp=u['mp']), "error")
                return

        loading = LoadingDialog(tr("boss_loading"), self)
        loading.show()
        QApplication.processEvents()

        r = db.attack_boss(self.user_id, u.get("guild_id"), action)
        loading.accept()

        if not r.get("ok"):
            SND.error()
            _show(self, tr("guild_cant_attack"), r.get("msg", "Error"), "error")
            return

        if r.get("defeated"):
            SND.boss_dead()
            msg = r.get("msg", "Boss defeated!")
            if r.get("extra_effect"):
                msg += "\n" + r["extra_effect"]
            _show(self, tr("victory_title"), msg, "success")
            QTimer.singleShot(0, self.load)
            AppState.refresh()
            return

        # Tampilkan hasil
        SND.boss_hit()
        u2 = AppState.user()
        has_spyglass = u2.get("has_spyglass", 0)

        user_dmg = r.get("user_damage", 0)
        user_crit = r.get("user_critical", False)
        user_dmg_text = f"⚔️ {user_dmg} DMG" + (" " + tr("boss_critical_mark") if user_crit else "")

        # BLOCK – ditangani terpisah
        if action == "block":
            reduction = r.get("block_reduction", 0)
            msg = tr("boss_block_result", reduction=reduction)
            _show(self, tr("boss_block_title"), msg, "info")
            QTimer.singleShot(0, self.load)
            AppState.refresh()
            return

        # ── BUAT PESAN SERANGAN ──
        if has_spyglass:
            boss_dmg = r.get("boss_damage", 0)
            boss_crit = r.get("boss_critical", False)
            boss_dmg_text = tr("boss_damage_text", dmg=boss_dmg) + (" " + tr("boss_critical_mark") if boss_crit else "")
            msg = tr("boss_attack_spyglass",
                    user_dmg=user_dmg,
                    boss_hp_left=r.get('boss_hp_left', 0),
                    boss_max_hp=r.get('boss_max_hp', 0),
                    boss_dmg_text=boss_dmg_text,
                    user_hp=u2['hp'])
        else:
            actual_damage = r.get('actual_damage', 0)
            msg = tr("boss_attack_no_spyglass",
                    user_dmg=user_dmg,
                    actual_damage=actual_damage,
                    user_hp=u2['hp'])

        # ── TAMBAHKAN INFORMASI TAMBAHAN ──
        if r.get("block_reduction", 0) > 0:
            msg += "\n" + tr("boss_block_active_info", reduction=r['block_reduction'])

        reduc = u2.get("hp_damage_reduction", 0)
        if reduc > 0:
            if has_spyglass:
                msg += " " + tr("boss_reduction_info", reduction=reduc)
            else:
                msg += "\n" + tr("boss_reduction_info", reduction=reduc)

        if r.get("shield_used"):
            msg += "\n" + tr("boss_shield_used")

        if r.get("revived"):
            msg += "\n" + tr("attack_totem_revive")

        # ── EFFECT ULTIMATE ──
        if r.get("action") == "ultimate" and r.get("extra_effect"):
            msg += "\n" + r["extra_effect"]

        _show(self, tr("attack_title"), msg, "info" if not r.get("revived") else "success")
        AppState.refresh()
        QTimer.singleShot(0, self.load)

    def _show_team_selection(self):
        """Dialog untuk memilih anggota tim raid (maks 5 termasuk leader)."""
        u = AppState.user()
        guild_id = u.get("guild_id")
        if not guild_id:
            return
        data = db.get_guild(guild_id)
        members = data.get("members", [])
        # Leader otomatis masuk
        leader = next((m for m in members if m["id"] == u["id"]), None)
        if not leader:
            return
        
        # Pilihan boss
        bid = self._boss_cb.currentData()
        if not bid:
            _show(self, tr("msg_error"), "Pilih boss terlebih dahulu!", "error")
            return
        boss = db.BOSSES.get(bid, {})
        min_lvl = boss.get("min_level", 1)
        max_lvl = boss.get("max_level", 999)
        
        # Tampilkan dialog dengan checkbox
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("raid_team_selection"))
        dlg.setMinimumSize(400, 300)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel(f"Pilih anggota tim (maks 4 anggota + leader).\nMin level: {min_lvl}, Max level: {max_lvl}"))
        
        checkboxes = []
        for m in members:
            if m["id"] == u["id"]:
                continue
            # Filter level
            if m["level"] < min_lvl or m["level"] > max_lvl:
                continue
            cb = QCheckBox(f"{m['display_name']} (Lv.{m['level']})")
            cb.setChecked(False)
            layout.addWidget(cb)
            checkboxes.append((cb, m["id"]))
        
        # Tombol ok
        btn_ok = _btn(tr("raid_start_btn"), "solid", dlg.accept)
        layout.addWidget(btn_ok)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        
        # Kumpulkan yang dipilih
        selected = [u["id"]]  # leader selalu masuk
        for cb, uid in checkboxes:
            if cb.isChecked():
                selected.append(uid)
                if len(selected) >= 5:
                    break
        self.selected_team = selected
        # Mulai boss dengan tim ini
        self._start_boss_with_team()

    def _start_boss_with_team(self):
        bid = self._boss_cb.currentData()
        if not bid:
            _show(self, tr("msg_error"), tr("raid_select_boss_first"), "error")
            return
        
        # Refresh user data
        u = AppState.user()
        if not u or not isinstance(u, dict) or not u.get("id"):
            # Coba refresh dari database langsung
            uid = AppState.user_id
            if uid:
                u = db.get_user(uid)
            if not u or not isinstance(u, dict):
                _show(self, tr("msg_error"), "Data user tidak valid.", "error")
                return
        
        uid = u.get("id")
        if not uid:
            _show(self, tr("msg_error"), "ID user tidak ditemukan.", "error")
            return
        guild_id = u.get("guild_id")
        if not guild_id:
            _show(self, tr("msg_error"), "Kamu tidak berada di guild.", "error")
            return
        
        if not self.selected_team:
            self.selected_team = [uid]
        elif uid not in self.selected_team:
            self.selected_team.insert(0, uid)
        self.selected_team = self.selected_team[:5]
        
        r = db.start_boss(guild_id, bid, u, self.selected_team)
        if r["ok"]:
            SND.boss_hit()
            _show(self, tr("boss_appear_title"), r["msg"], "warning")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()
        QTimer.singleShot(0, self.load)

    def closeEvent(self, e):
        AppState.unregister(self.load)
        super().closeEvent(e)
        
# ══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD PAGE
# ══════════════════════════════════════════════════════════════════════════════
class DashboardPage(QWidget):
    """Halaman Dashboard / Home yang modern, rapi, dan tidak overlap."""

    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self.cache = {"data": None, "timestamp": 0}
        self.cache_lifetime = 15
        self._build()
        AppState.register(self.load)

    def _build(self):
        self.main_layout = QVBoxLayout(self)
        self.main_layout.setContentsMargins(20, 16, 20, 16)
        self.main_layout.setSpacing(14)

        # ---- HEADER ----
        header = QWidget()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)

        title = _lbl(tr("dashboard_title"), size=18, bold=True)
        title.setStyleSheet(f"color: {_T('light')};")
        header_layout.addWidget(title)
        header_layout.addStretch()

        wrapped_btn = _btn(tr("wrapped_btn", year=date.today().year),
                           "gold", self._open_wrapped, h=32)
        header_layout.addWidget(wrapped_btn)

        widgets_btn = _btn(tr("dashwidgets_btn"), h=32,
                           slot=self._open_widget_dialog)
        header_layout.addWidget(widgets_btn)

        refresh_btn = _btn("🔄", h=32, slot=self.load)
        refresh_btn.setFixedWidth(60)
        header_layout.addWidget(refresh_btn)

        self.main_layout.addWidget(header)

        # ---- HEATMAP AKTIVITAS (persisten, tidak ikut rebuild konten) ----
        self._heatmap_card = QFrame()
        self._heatmap_card.setObjectName("card")
        hm_lay = QVBoxLayout(self._heatmap_card)
        hm_lay.setContentsMargins(16, 14, 16, 14)
        hm_lay.setSpacing(8)
        hm_lay.addWidget(_lbl(tr("heatmap_title"), "section", 13, True))
        self._heatmap = HeatmapWidget()
        hm_lay.addWidget(self._heatmap, 0, Qt.AlignmentFlag.AlignHCenter)
        # Legend: Sedikit □□□□□ Banyak
        self._heatmap_legend = QWidget()
        legend = QHBoxLayout(self._heatmap_legend)
        legend.setContentsMargins(0, 0, 0, 0)
        legend.setSpacing(4)
        legend.addStretch()
        legend.addWidget(_lbl(tr("heatmap_less"), "sub", 10))
        for i, c in enumerate(self._heatmap._colors()[1:]):
            box = QFrame()
            box.setFixedSize(11, 11)
            box.setStyleSheet(f"background: {c}; border-radius: 2px;")
            legend.addWidget(box)
        legend.addWidget(_lbl(tr("heatmap_more"), "sub", 10))
        hm_lay.addWidget(self._heatmap_legend)
        # CATATAN: kartu widget TIDAK diparkir di main_layout (itu menutupi
        # konten). Mereka disisipkan ke dalam flow content_layout oleh
        # _attach_widget_cards() setiap _render — tepat setelah kartu statistik.

        # ---- 💡 INSIGHT OTOMATIS (persisten) ----
        self._insights_card = QFrame()
        self._insights_card.setObjectName("card")
        ins_lay = QVBoxLayout(self._insights_card)
        ins_lay.setContentsMargins(16, 14, 16, 14)
        ins_lay.setSpacing(6)
        ins_lay.addWidget(_lbl(tr("insights_title"), "section", 13, True))
        self._insights_body = QLabel("")
        self._insights_body.setWordWrap(True)
        self._insights_body.setStyleSheet(
            f"color: {_T('text')}; font-size: 12px; line-height: 150%;")
        ins_lay.addWidget(self._insights_body)
        self._insights_compact = False

        # ---- 😴 GRAFIK TIDUR ↔ PRODUKTIVITAS (persisten) ----
        self._health_card = QFrame()
        self._health_card.setObjectName("card")
        hc_lay = QVBoxLayout(self._health_card)
        hc_lay.setContentsMargins(16, 14, 16, 14)
        hc_lay.setSpacing(8)
        hc_lay.addWidget(_lbl(tr("healthchart_title"), "section", 13, True))
        self._health_chart = HealthChartWidget()
        hc_lay.addWidget(self._health_chart)

        # ---- Registry widget yang bisa diatur user (urutan/visibilitas/compact) ----
        # Disisipkan ke content_layout di _render (bukan ditambahkan ke main_layout).
        self._widget_anchor = 3   # setelah greeting, rank card, stats grid
        self._widget_cards = {
            "heatmap":      (self._heatmap_card, self._set_heatmap_compact),
            "insights":     (self._insights_card, self._set_insights_compact),
            "health_chart": (self._health_card, self._set_health_compact),
        }

        # ---- SCROLL AREA ----
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")

        self.content = QWidget()
        self.content_layout = QVBoxLayout(self.content)
        self.content_layout.setSpacing(14)
        self.content_layout.setContentsMargins(0, 0, 0, 0)

        scroll.setWidget(self.content)
        self.main_layout.addWidget(scroll, 1)

        self.load()

    # ========== ⚙️ WIDGET CONFIG (urutan, visibilitas, compact) ==========
    def _detach_widget_cards(self):
        """Lepaskan kartu widget dari content_layout SEBELUM _clear_layout,
        supaya tidak ikut ke-deleteLater saat render ulang."""
        if not hasattr(self, "_widget_cards"):
            return
        for card, _setter in self._widget_cards.values():
            self.content_layout.removeWidget(card)
            card.setVisible(False)

    def _attach_widget_cards(self, start_index: int):
        """Sisipkan kartu widget ke DALAM flow content_layout (tengah-tengah
        kartu dashboard) sesuai konfigurasi user: urutan, visibilitas, compact."""
        if not hasattr(self, "_widget_cards"):
            return
        try:
            cfg_list = db.get_dashboard_widgets(self.user_id)
        except Exception:
            cfg_list = []
        cfg = {w.get("key"): w for w in cfg_list if isinstance(w, dict)}
        order = [w.get("key") for w in cfg_list
                 if isinstance(w, dict) and w.get("key") in self._widget_cards]
        # Key baru yang belum ada di konfigurasi tetap ditampilkan di akhir
        for k in self._widget_cards:
            if k not in order:
                order.append(k)
        idx = max(0, start_index)
        for k in order:
            card, setter = self._widget_cards[k]
            conf = cfg.get(k, {"visible": True, "compact": False})
            try:
                setter(bool(conf.get("compact", False)))
            except Exception:
                pass
            card.setVisible(bool(conf.get("visible", True)))
            self.content_layout.insertWidget(idx, card)
            idx += 1

    def _apply_widget_config(self):
        """Terapkan ulang konfigurasi widget (dipanggil dialog Atur Widget)."""
        self._detach_widget_cards()
        self._attach_widget_cards(getattr(self, "_widget_anchor", 3))

    def _set_heatmap_compact(self, compact: bool):
        self._heatmap.set_compact(compact)
        self._heatmap_legend.setVisible(not compact)

    def _set_insights_compact(self, compact: bool):
        self._insights_compact = compact
        self._refresh_insights()

    def _set_health_compact(self, compact: bool):
        self._health_chart.set_compact(compact)

    def _refresh_insights(self):
        if not hasattr(self, "_insights_body"):
            return
        try:
            ins = db.get_insights(self.user_id)
        except Exception as e:
            log.error(f"Insights gagal: {e}")
            return
        if not ins.get("has_data"):
            self._insights_body.setText(tr("insights_no_data"))
            return
        lines = []
        if ins.get("top_weekday"):
            lines.append(tr("insights_top_day", day=ins["top_weekday"],
                            n=ins["top_weekday_count"]))
        if ins.get("best_day"):
            lines.append(tr("insights_best_day", n=ins["best_day_count"],
                            date=ins["best_day"]))
        lines.append(tr("insights_longest", n=ins.get("longest_streak", 0)))
        lines.append(tr("insights_active", n=ins.get("active_streaks", 0)))
        lines.append(tr("insights_focus", n=ins.get("focus_minutes", 0)))
        if self._insights_compact:
            lines = lines[:2]
        self._insights_body.setText("\n".join(l.replace("**", "") for l in lines))

    def _refresh_health_chart(self):
        if hasattr(self, "_health_chart"):
            try:
                self._health_chart.refresh(
                    db.get_health_productivity_series(self.user_id))
            except Exception as e:
                log.error(f"Health chart gagal: {e}")

    def _open_wrapped(self):
        YearWrappedDialog(self.user_id, self).exec()

    def _open_widget_dialog(self):
        dlg = DashboardWidgetDialog(self.user_id, self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            self._apply_widget_config()

    # ========== HELPERS ==========
    def _currency_symbol(self):
        curr = db.get_user_currency(self.user_id)
        return {"IDR": "Rp", "USD": "$", "EUR": "€"}.get(curr, "Rp")

    def _convert_currency(self, amount_idr):
        curr = db.get_user_currency(self.user_id)
        return db.convert_from_idr(amount_idr, curr)

    def _format_currency(self, amount_idr):
        converted = self._convert_currency(amount_idr)
        return f"{self._currency_symbol()} {converted:,.0f}"

    # ========== CARD BUILDERS ==========
    def _stat_card(self, icon: str, title: str, value: str, subtitle: str = "", color: str = None):
        card = _card()
        card.setMinimumHeight(70)
        card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)

        if color:
            card.setStyleSheet(f"""
                QFrame#card {{
                    background: {_T('panel')};
                    border: 1px solid {_T('border')};
                    border-left: 4px solid {color};
                    border-radius: 6px;
                }}
            """)

        layout = QHBoxLayout(card)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(8)

        icon_lbl = _emoji_label(icon, ICON_CARD)
        layout.addWidget(icon_lbl)

        info = QVBoxLayout()
        info.setSpacing(1)
        title_lbl = QLabel(title)
        title_lbl.setStyleSheet(f"color: {_T('muted')}; font-size: 10px;")
        info.addWidget(title_lbl)

        value_lbl = QLabel(value)
        value_lbl.setStyleSheet(f"color: {_T('text')}; font-size: 16px; font-weight: bold;")
        value_lbl.setWordWrap(True)
        info.addWidget(value_lbl)

        if subtitle:
            sub_lbl = QLabel(subtitle)
            sub_lbl.setStyleSheet(f"color: {_T('muted')}; font-size: 10px;")
            info.addWidget(sub_lbl)

        layout.addLayout(info, 1)
        return card

    def _progress_ring(self, value: int, max_value: int, label: str, color: str = "#80c000", size: int = 80):
        """Buat progress ring menggunakan QPainter custom."""
        widget = QWidget()
        widget.setFixedSize(size + 20, size + 40)
        widget.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)

        layout = QVBoxLayout(widget)
        layout.setSpacing(4)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.setContentsMargins(0, 0, 0, 0)

        # Custom ring
        ring = ProgressRing(value, max_value, color, size)
        layout.addWidget(ring)

        # Label
        lbl = QLabel(label)
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl.setStyleSheet(f"color: {_T('text')}; font-size: 10px; font-weight: bold; max-width: {size}px;")
        layout.addWidget(lbl)

        # Simpan referensi untuk update (opsional)
        widget.ring = ring
        widget.label = lbl

        return widget

    # ========== LOAD & RENDER ==========
    def load(self):
        if not AppState.user_id:
            return

        # Refresh heatmap (murah, selalu dari DB)
        try:
            self._heatmap.set_data(db.get_activity_heatmap(self.user_id))
        except Exception as e:
            log.error(f"Heatmap refresh gagal: {e}")

        # Refresh insight & grafik kesehatan (widget persisten)
        self._refresh_insights()
        self._refresh_health_chart()

        import time
        now = time.time()

        if self.cache.get("data") and (now - self.cache["timestamp"]) < self.cache_lifetime:
            data = self.cache["data"]
        else:
            s = db.get_stats(self.user_id)
            u = s["user"]
            health = db.get_health_summary(self.user_id)
            ss = db.get_sport_stats(self.user_id)
            eco_summary = db.get_economy_summary(self.user_id)
            eco_weekly = db.get_economy_weekly(self.user_id)
            weekly = s["weekly"]
            food_stats = db.get_food_summary_stats(self.user_id)
            health_logs = db.get_health_logs(self.user_id, days=7)
            achievements = db.get_user_achievements(self.user_id)

            data = {
                "u": u, "s": s, "health": health, "ss": ss,
                "eco_summary": eco_summary, "eco_weekly": eco_weekly,
                "weekly": weekly, "food_stats": food_stats,
                "health_logs": health_logs, "achievements": achievements
            }
            self.cache["data"] = data
            self.cache["timestamp"] = now

        self._render(data)

    def _clear_layout(self, layout: QLayout):
        """Hapus semua widget dan sub-layout secara rekursif."""
        if layout is None:
            return
        while layout.count():
            item = layout.takeAt(0)
            if item.widget():
                widget = item.widget()
                widget.setParent(None)
                widget.deleteLater()
            elif item.layout():
                self._clear_layout(item.layout())
                item.layout().deleteLater()

    def _render(self, data):
        # Lepaskan kartu widget persisten (heatmap/insights/grafik) SEBELUM
        # content dibersihkan — jangan sampai ikut deleteLater
        self._detach_widget_cards()
        # Bersihkan layout secara rekursif
        self._clear_layout(self.content_layout)

        u = data["u"]
        s = data["s"]
        health = data["health"]
        ss = data["ss"]
        weekly = data["weekly"]
        achievements = data["achievements"]

        # ---- INISIALISASI VARIABEL UNTUK PROGRESS RING ----
        lvl = u.get("level", 1)
        xp = u.get("xp", 0)
        need = lvl * 150

        sport_lvl = ss.get("sport_level", 1)
        sport_xp = ss.get("sport_xp", 0)
        sport_need = sport_lvl * 100

        hp = u.get("hp", 0)
        max_hp = u.get("max_hp", 1)

        # Calorie progress (hari ini)
        today = date.today().isoformat()
        nutri = db.get_nutrition_summary(self.user_id, today)
        goals = db.get_nutrition_goals(self.user_id)
        cal_value = nutri.get('calories', 0)
        cal_max = goals.get('daily_calories', 2000)
        if cal_max <= 0:
            cal_max = 2000

        # ---- GREETING ----
        greeting = QLabel(tr("dashboard_greeting", name=u.get("display_name", u.get("username", "User"))))
        greeting.setStyleSheet(f"color: {_T('text')}; font-size: 15px; font-weight: bold;")
        self.content_layout.addWidget(greeting)

        # ---- RANK CARD ----
        rank_data = db.calculate_rank(self.user_id)
        rank_widget = self._rank_card(rank_data)
        self.content_layout.addWidget(rank_widget)

        # ---- STAT CARDS (4 kolom, 2 baris) ----
        stats_grid = QGridLayout()
        stats_grid.setSpacing(10)
        stats_grid.setContentsMargins(0, 0, 0, 0)

        stats_data = [
            ("⭐", tr("dashboard_level"), str(u["level"]), "", "#f0a800"),
            ("💰", tr("dashboard_gold"), f"{u['gold']:.0f}", "", "#f0a800"),
            ("❤️", tr("dashboard_hp"), f"{u['hp']}/{u['max_hp']}", "", "#e05050"),
            ("💙", tr("dashboard_mp"), f"{u['mp']}/{u['max_mp']}", "", "#4da6ff"),
            ("🔥", tr("dashboard_streak"), str(s["max_streak"]), tr("dashboard_streak_days"), "#ff6b00"),
            ("📜", tr("dashboard_tasks_done"), str(u.get("total_tasks_completed", 0)), "", "#80c000"),
            ("👹", tr("dashboard_boss_killed"), str(s["bosses_killed"]), "", "#a97fff"),
            ("🏅", tr("dashboard_sport_level"), str(ss["sport_level"]), f"{ss['sport_xp']}/{ss['sport_level']*100} SP", "#f0a800"),
        ]

        for i, (icon, title, value, subtitle, color) in enumerate(stats_data):
            card = self._stat_card(icon, title, value, subtitle, color)
            row = i // 4
            col = i % 4
            stats_grid.addWidget(card, row, col)
            stats_grid.setColumnStretch(col, 1)

        self.content_layout.addLayout(stats_grid)

        # ---- WIDGET DASHBOARD: heatmap/insights/grafik kesehatan ----
        # Disisipkan DI SINI, tengah-tengah kartu statistik (bukan menutupi
        # bagian atas halaman). Urutan/visibilitas mengikuti "⚙️ Atur Widget".
        self._widget_anchor = self.content_layout.count()
        self._attach_widget_cards(self._widget_anchor)

        # ---- PROGRESS RINGS (2x2 grid) ----
        rings_group = QGroupBox(tr("dashboard_progress"))
        rings_group.setStyleSheet(f"""
            QGroupBox {{
                color: {_T('light')};
                border: 1px solid {_T('border')};
                border-radius: 8px;
                margin-top: 8px;
                padding-top: 4px;
                background: {_T('panel')};
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px;
            }}
        """)

        rings_layout = QGridLayout(rings_group)
        rings_layout.setSpacing(16)
        rings_layout.setContentsMargins(16, 16, 16, 16)
        rings_layout.setHorizontalSpacing(30)
        rings_layout.setVerticalSpacing(16)

        # Ring 1: XP
        ring1 = self._progress_ring(xp, need, tr("dashboard_level_progress"), "#80c000", 75)
        rings_layout.addWidget(ring1, 0, 0)

        # Ring 2: HP
        ring2 = self._progress_ring(hp, max_hp, tr("dashboard_hp_progress"), "#e05050", 75)
        rings_layout.addWidget(ring2, 0, 1)

        # Ring 3: Sport
        ring3 = self._progress_ring(sport_xp, sport_need, tr("dashboard_sport_progress"), "#f0a800", 75)
        rings_layout.addWidget(ring3, 1, 0)

        # Ring 4: Calories
        ring4 = self._progress_ring(cal_value, cal_max, tr("dashboard_calories"), "#4da6ff", 75)
        rings_layout.addWidget(ring4, 1, 1)

        rings_layout.setColumnStretch(0, 1)
        rings_layout.setColumnStretch(1, 1)

        self.content_layout.addWidget(rings_group)

        # ---- QUICK ACTIONS (3 kolom, 2 baris) ----
        actions_group = QGroupBox(tr("dashboard_quick_actions"))
        actions_group.setStyleSheet(f"""
            QGroupBox {{
                color: {_T('light')};
                border: 1px solid {_T('border')};
                border-radius: 8px;
                margin-top: 8px;
                padding-top: 4px;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px;
            }}
        """)

        actions_layout = QGridLayout(actions_group)
        actions_layout.setSpacing(10)
        actions_layout.setContentsMargins(12, 12, 12, 12)

        actions = [
            ("⛏", tr("dashboard_habits"), "habits", "#5a8a2e"),
            ("📅", tr("dashboard_dailies"), "dailies", "#4da6ff"),
            ("📜", tr("dashboard_quests"), "todos", "#a97fff"),
            ("🏅", tr("dashboard_sport"), "sport", "#f0a800"),
            ("💰", tr("dashboard_economy"), "economy", "#f0a800"),
            ("🏆", tr("dashboard_achievements"), "achievements", "#ff6b00"),
        ]

        row, col = 0, 0
        for icon, label, page, color in actions:
            btn = QPushButton(f"{icon}\n{label}")
            btn.setMinimumHeight(54)
            btn.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {_T('panel')};
                    border: 1px solid {color};
                    border-radius: 8px;
                    color: {_T('text')};
                    font-weight: bold;
                    font-size: 10px;
                    padding: 4px;
                }}
                QPushButton:hover {{
                    background: {color};
                    color: #fff;
                }}
            """)
            btn.clicked.connect(lambda checked, p=page: self._navigate_to(p))
            actions_layout.addWidget(btn, row, col)
            col += 1
            if col >= 3:
                col = 0
                row += 1

        self.content_layout.addWidget(actions_group)

        # ---- WEEKLY CHART (jika ada matplotlib) ----
        if EXPORT_IMPORTS_OK and weekly:
            try:
                import matplotlib.pyplot as plt
                from io import BytesIO
                from PyQt6.QtGui import QPixmap

                chart_group = QGroupBox(tr("dashboard_weekly_chart"))
                chart_group.setStyleSheet(f"""
                    QGroupBox {{
                        color: {_T('light')};
                        border: 1px solid {_T('border')};
                        border-radius: 8px;
                        margin-top: 8px;
                        padding-top: 4px;
                    }}
                    QGroupBox::title {{
                        subcontrol-origin: margin;
                        left: 12px;
                        padding: 0 6px;
                    }}
                """)
                chart_layout = QVBoxLayout(chart_group)

                days = [row['day'][5:] for row in weekly]
                xp_vals = [row['xp'] or 0 for row in weekly]
                gold_vals = [row['gold'] or 0 for row in weekly]

                fig, ax = plt.subplots(figsize=(6, 2.2))
                ax.bar(days, xp_vals, label=tr("dashboard_xp"), color='#80c000', alpha=0.7)
                ax.bar(days, gold_vals, label=tr("dashboard_gold"), color='#f0a800', alpha=0.7)
                ax.set_facecolor('#2d2d2d')
                fig.patch.set_facecolor('#2d2d2d')
                ax.tick_params(colors='#e8e8e8')
                ax.xaxis.label.set_color('#e8e8e8')
                ax.yaxis.label.set_color('#e8e8e8')
                ax.title.set_color('#7bbf3e')
                plt.setp(ax.get_xticklabels(), rotation=20, ha='right')
                ax.legend(facecolor='#2d2d2d', edgecolor='#444', labelcolor='#e8e8e8', fontsize=8)

                buf = BytesIO()
                plt.savefig(buf, format='png', dpi=80, bbox_inches='tight')
                buf.seek(0)
                plt.close(fig)

                pixmap = QPixmap()
                pixmap.loadFromData(buf.read())
                chart_label = QLabel()
                chart_label.setPixmap(pixmap.scaledToWidth(550, Qt.TransformationMode.SmoothTransformation))
                chart_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                chart_layout.addWidget(chart_label)

                self.content_layout.addWidget(chart_group)
            except Exception as e:
                print(f"Chart error: {e}")

        # ---- RECENT ACTIVITY ----
        activity_group = QGroupBox(tr("dashboard_recent_activity"))
        activity_group.setStyleSheet(f"""
            QGroupBox {{
                color: {_T('light')};
                border: 1px solid {_T('border')};
                border-radius: 8px;
                margin-top: 8px;
                padding-top: 4px;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px;
            }}
        """)
        activity_layout = QVBoxLayout(activity_group)

        activity_list = QListWidget()
        activity_list.setMinimumHeight(90)
        activity_list.setStyleSheet(f"""
            QListWidget {{
                background: {_T('bg')};
                border: none;
                color: {_T('text')};
            }}
            QListWidget::item {{
                padding: 2px 6px;
                border-bottom: 1px solid {_T('border')};
                font-size: 11px;
            }}
        """)

        for entry in s["recent_log"][:8]:
            action_text = tr(f"action_{entry['action']}", default=entry['action'])
            activity_list.addItem(f"[{entry['created_at'][11:16]}]  {action_text}  —  {entry['detail']}")

        if not s["recent_log"]:
            activity_list.addItem(tr("dashboard_no_activity"))

        activity_layout.addWidget(activity_list)
        self.content_layout.addWidget(activity_group)

        # ---- HEALTH SUMMARY ----
        health_group = QGroupBox(tr("dashboard_health_summary"))
        health_group.setStyleSheet(f"""
            QGroupBox {{
                color: {_T('light')};
                border: 1px solid {_T('border')};
                border-radius: 8px;
                margin-top: 8px;
                padding-top: 4px;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px;
            }}
        """)
        health_layout = QGridLayout(health_group)
        health_layout.setSpacing(6)
        health_layout.setContentsMargins(12, 12, 12, 12)

        health_data = [
            (tr("dashboard_avg_steps"), f"{health['avg_steps']:,}", "🚶"),
            (tr("dashboard_avg_sleep"), f"{health['avg_sleep']:.1f} jam", "😴"),
            (tr("dashboard_avg_water"), f"{health['avg_water']:,} ml", "💧"),
            (tr("dashboard_avg_weight"), f"{health['avg_weight']:.1f} kg", "⚖️"),
            (tr("dashboard_avg_hr"), f"{health['avg_hr']} bpm", "❤️"),
            (tr("dashboard_days_recorded"), f"{health['days_recorded']}", "📆"),
        ]

        for i, (label, value, icon) in enumerate(health_data):
            lbl = QLabel(f"{icon} {label}: <b>{value}</b>")
            lbl.setStyleSheet(f"color: {_T('text')}; font-size: 11px; padding: 2px;")
            health_layout.addWidget(lbl, i // 3, i % 3)

        self.content_layout.addWidget(health_group)

        # ---- FOOTER ----
        footer = QLabel(tr("dashboard_footer"))
        footer.setAlignment(Qt.AlignmentFlag.AlignCenter)
        footer.setStyleSheet(f"color: {_T('muted')}; font-size: 10px; padding: 6px;")
        self.content_layout.addWidget(footer)

    def _show_rank_list(self):
        """Buka dialog daftar rank."""
        dlg = RankListDialog(self.user_id, self)
        dlg.exec()

    def _rank_card(self, rank_data: dict) -> QWidget:
        card = _card()
        card.setStyleSheet(f"""
            QFrame#card {{
                background: qlineargradient(x1:0,y1:0,x2:1,y2:0,
                    stop:0 #2d1b0e, stop:1 #1a1a1a);
                border: 2px solid #f0a800;
                border-radius: 12px;
                padding: 4px;
            }}
        """)
        
        layout = QHBoxLayout(card)
        layout.setContentsMargins(20, 14, 20, 14)
        layout.setSpacing(16)
        
        # Icon besar
        icon = _emoji_label(rank_data["rank_icon"], ICON_HERO)
        layout.addWidget(icon)
        
        # Info rank
        info = QVBoxLayout()
        info.setSpacing(2)
        
        name_text = tr(rank_data["rank_name_key"])
        name = QLabel(name_text)
        name.setStyleSheet(f"color: #f0a800; font-size: 20px; font-weight: bold;")
        info.addWidget(name)
        
        desc_text = tr(rank_data["rank_desc_key"])
        desc = QLabel(desc_text)
        desc.setStyleSheet(f"color: {_T('muted')}; font-size: 12px;")
        info.addWidget(desc)
        
        score_text = tr("rank_score_label", score=rank_data["score"], max=rank_data["max_score"])
        score = QLabel(score_text)
        score.setStyleSheet(f"color: {_T('text')}; font-size: 12px;")
        info.addWidget(score)
        
        # Progress ke rank berikutnya
        progress_widget = QWidget()
        progress_layout = QVBoxLayout(progress_widget)
        progress_layout.setContentsMargins(0, 0, 0, 0)
        progress_layout.setSpacing(2)
        
        if rank_data.get("has_next", True) and rank_data["next_rank_name_key"]:
            next_name = tr(rank_data["next_rank_name_key"])
            prog_text = tr("rank_next_progress", 
                        icon=rank_data["next_rank_icon"], 
                        name=next_name, 
                        progress=rank_data["progress"])
        else:
            prog_text = tr("rank_max_label")
        
        prog_label = QLabel(prog_text)
        prog_label.setStyleSheet(f"color: {_T('muted')}; font-size: 11px;")
        progress_layout.addWidget(prog_label)
        
        pb = QProgressBar()
        pb.setMaximum(100)
        pb.setValue(rank_data["progress"])
        pb.setMinimumHeight(8)
        pb.setStyleSheet(f"""
            QProgressBar {{
                background: {_T('bg')};
                border: 1px solid {_T('border')};
                border-radius: 4px;
                height: 8px;
            }}
            QProgressBar::chunk {{
                background: qlineargradient(x1:0,y1:0,x2:1,y2:0,
                    stop:0 #f0a800, stop:1 #ff6b00);
                border-radius: 4px;
            }}
        """)
        pb.setTextVisible(False)
        progress_layout.addWidget(pb)
        
        info.addLayout(progress_layout)
        layout.addLayout(info, 1)
        
        # Tombol lihat semua rank
        btn_rank_list = _btn("🏆 " + tr("rank_view_all"), h=36)
        btn_rank_list.setFixedWidth(140)
        btn_rank_list.clicked.connect(self._show_rank_list)
        layout.addWidget(btn_rank_list)
        
        return card
    def _navigate_to(self, page_key):
        """Navigasi ke halaman lain."""
        main_win = self.window()
        if hasattr(main_win, "_nav"):
            main_win._nav._select(page_key)
        if hasattr(main_win, "_switch"):
            main_win._switch(page_key)

    def closeEvent(self, e):
        AppState.unregister(self.load)
        super().closeEvent(e)

class ProgressRing(QWidget):
    def __init__(self, value=0, max_value=100, color="#80c000", size=80, parent=None):
        super().__init__(parent)
        self._value = value
        self._max_value = max_value
        self._color = QColor(color)
        self._size = size
        self._percent = 0
        self.setFixedSize(size + 20, size + 40)
        self._update_percent()

    def set_value(self, value):
        self._value = value
        self._update_percent()
        self.update()

    def set_max(self, max_value):
        self._max_value = max_value
        self._update_percent()
        self.update()

    def _update_percent(self):
        if self._max_value > 0:
            self._percent = int((min(self._value, self._max_value) / self._max_value) * 100)
        else:
            self._percent = 0

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        margin = 10
        size = self._size
        rect = QRect(margin, margin, size, size)
        center = rect.center()
        radius = size // 2
        pen_width = 8

        # Background
        painter.setPen(QPen(QColor(60, 60, 60), pen_width, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap))
        painter.drawEllipse(center, radius - pen_width//2, radius - pen_width//2)

        # Progress
        angle = int(360 * self._percent / 100)
        if angle > 0:
            painter.setPen(QPen(self._color, pen_width, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap))
            start_angle = -90 * 16
            span_angle = -angle * 16
            painter.drawArc(rect, start_angle, span_angle)

        # Text
        painter.setPen(QPen(self._color, 1))
        font = QFont("Segoe UI", max(11, size//6), QFont.Weight.Bold)
        painter.setFont(font)
        text = f"{self._percent}%"
        metrics = QFontMetrics(font)
        text_rect = metrics.boundingRect(text)
        text_x = center.x() - text_rect.width() // 2
        text_y = center.y() + text_rect.height() // 2 - 2
        painter.drawText(text_x, text_y, text)

class RankListDialog(QDialog):
    """Dialog daftar semua rank dan progres user."""
    
    def __init__(self, user_id: int, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        self.setWindowTitle(tr("rank_dialog_title"))
        self.setMinimumSize(550, 500)
        self.setStyleSheet(build_ss())
        self._build()
        self._load_data()
    
    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 16, 20, 16)
        layout.setSpacing(10)
        
        # Header
        hdr = QHBoxLayout()
        hdr.addWidget(_lbl(tr("rank_dialog_title"), size=16, bold=True))
        hdr.addStretch()
        close_btn = _btn(tr("btn_close"), "flat", self.accept)
        hdr.addWidget(close_btn)
        layout.addLayout(hdr)
        layout.addWidget(_sep())
        
        # Scroll area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")
        
        self.container = QWidget()
        self.container_layout = QVBoxLayout(self.container)
        self.container_layout.setSpacing(8)
        self.container_layout.setContentsMargins(0, 0, 0, 0)
        
        scroll.setWidget(self.container)
        layout.addWidget(scroll, 1)
    
    def _load_data(self):
        # Ambil data rank user
        rank_data = db.calculate_rank(self.user_id)
        user_score = rank_data["score"]
        current_rank = rank_data["rank"]
        
        # Daftar rank (hardcoded agar konsisten)
        RANKS = [
            {"min_score": 0,  "name_key": "rank_pemula",          "icon": "🥚", "desc_key": "rank_desc_pemula"},
            {"min_score": 10, "name_key": "rank_penambang",       "icon": "⛏️", "desc_key": "rank_desc_penambang"},
            {"min_score": 20, "name_key": "rank_penjelajah",      "icon": "🪓", "desc_key": "rank_desc_penjelajah"},
            {"min_score": 30, "name_key": "rank_petualang",       "icon": "⚔️", "desc_key": "rank_desc_petualang"},
            {"min_score": 40, "name_key": "rank_ksatria",         "icon": "🛡️", "desc_key": "rank_desc_ksatria"},
            {"min_score": 50, "name_key": "rank_veteran",         "icon": "⭐", "desc_key": "rank_desc_veteran"},
            {"min_score": 60, "name_key": "rank_legenda",         "icon": "🌟", "desc_key": "rank_desc_legenda"},
            {"min_score": 70, "name_key": "rank_raja",            "icon": "👑", "desc_key": "rank_desc_raja"},
            {"min_score": 80, "name_key": "rank_penguasa_naga",   "icon": "🐉", "desc_key": "rank_desc_penguasa_naga"},
            {"min_score": 90, "name_key": "rank_dewa",            "icon": "🌌", "desc_key": "rank_desc_dewa"},
        ]
        
        for i, rank in enumerate(RANKS):
            # Status rank
            is_unlocked = user_score >= rank["min_score"]
            is_current = (i == current_rank)
            
            # Buat card
            card = _card()
            if is_current:
                card.setStyleSheet(f"""
                    QFrame#card {{
                        background: {_T('panel')};
                        border: 2px solid #f0a800;
                        border-radius: 8px;
                    }}
                """)
            elif is_unlocked:
                card.setStyleSheet(f"""
                    QFrame#card {{
                        background: {_T('panel')};
                        border: 1px solid {_T('border')};
                        border-radius: 8px;
                    }}
                """)
            else:
                card.setStyleSheet(f"""
                    QFrame#card {{
                        background: {_T('bg')};
                        border: 1px dashed {_T('border')};
                        border-radius: 8px;
                        opacity: 0.6;
                    }}
                """)
            
            row = QHBoxLayout(card)
            row.setContentsMargins(12, 10, 12, 10)
            row.setSpacing(10)
            
            # Icon
            icon_lbl = QLabel(rank["icon"])
            icon_lbl.setFont(QFont("Segoe UI", 28))
            icon_lbl.setFixedWidth(44)
            row.addWidget(icon_lbl)
            
            # Info
            info = QVBoxLayout()
            info.setSpacing(1)
            
            name_text = tr(rank["name_key"])
            name_lbl = QLabel(name_text)
            name_lbl.setStyleSheet(f"color: {_T('text')}; font-size: 14px; font-weight: bold;")
            info.addWidget(name_lbl)
            
            desc_text = tr(rank["desc_key"])
            desc_lbl = QLabel(desc_text)
            desc_lbl.setStyleSheet(f"color: {_T('muted')}; font-size: 11px;")
            info.addWidget(desc_lbl)
            
            # Score required
            req_lbl = QLabel(tr("rank_required_score", score=rank["min_score"]))
            req_lbl.setStyleSheet(f"color: {_T('muted')}; font-size: 10px;")
            info.addWidget(req_lbl)
            
            row.addLayout(info, 1)
            
            # Status label
            if is_current:
                status = QLabel(tr("rank_current_badge"))
                status.setStyleSheet("color: #f0a800; font-weight: bold; font-size: 11px; background: #2e2500; padding: 2px 10px; border-radius: 10px;")
            elif is_unlocked:
                status = QLabel(tr("rank_unlocked_badge"))
                status.setStyleSheet("color: #80c000; font-weight: bold; font-size: 11px; background: #002a00; padding: 2px 10px; border-radius: 10px;")
            else:
                status = QLabel(tr("rank_locked_badge"))
                status.setStyleSheet("color: #e05050; font-weight: bold; font-size: 11px; background: #2a0000; padding: 2px 10px; border-radius: 10px;")
            
            row.addWidget(status)
            
            self.container_layout.addWidget(card)
        
        # Progress score di footer
        footer = QLabel(tr("rank_footer_progress", score=user_score, max=100))
        footer.setAlignment(Qt.AlignmentFlag.AlignCenter)
        footer.setStyleSheet(f"color: {_T('muted')}; font-size: 11px; padding: 8px;")
        self.container_layout.addWidget(footer)

# ══════════════════════════════════════════════════════════════════════════════
#  Achievement Page - Menampilkan daftar achievement
# ══════════════════════════════════════════════════════════════════════════════
class AchievementPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        
        self.user_id = user_id
        self._build()
        AppState.register(self.load)
    
    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 16, 20, 16)
        layout.setSpacing(10)
        layout.addWidget(_lbl(tr("achievement_title"), "section", 14, True))
        layout.addWidget(_sep())

        filter_widget = QWidget()
        filter_layout = QHBoxLayout(filter_widget)
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText(tr("achievement_search"))
        self.search_input.textChanged.connect(self.load)
        self.category_combo = QComboBox()
        # Gunakan data, bukan teks untuk filter
        self.category_combo.addItem(tr("achievement_all"), "all")
        self.category_combo.addItem(tr("achievement_category_level"), "level")
        self.category_combo.addItem(tr("achievement_category_focus"), "focus")
        self.category_combo.addItem(tr("achievement_category_crafting"), "crafting")
        self.category_combo.addItem(tr("achievement_category_habit"), "habit")
        self.category_combo.addItem(tr("achievement_category_daily"), "daily")
        self.category_combo.addItem(tr("achievement_category_todo"), "todo")
        self.category_combo.addItem(tr("achievement_category_sport"), "sport")
        self.category_combo.addItem(tr("achievement_category_economy"), "economy")
        self.category_combo.addItem(tr("achievement_category_pet"), "pet")
        self.category_combo.addItem(tr("achievement_category_guild"), "guild")
        self.category_combo.addItem(tr("achievement_category_boss"), "boss")
        self.category_combo.addItem(tr("achievement_category_social"), "social")
        self.category_combo.addItem(tr("achievement_category_health"), "health")
        self.category_combo.addItem(tr("achievement_category_nutrition"), "nutrition")
        self.category_combo.addItem(tr("achievement_category_special"), "special")
        self.category_combo.currentIndexChanged.connect(self.load)
        filter_layout.addWidget(self.search_input)
        filter_layout.addWidget(self.category_combo)
        layout.addWidget(filter_widget)

        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.container = QWidget()
        self.grid = QGridLayout(self.container)
        self.grid.setSpacing(12)
        self.grid.setAlignment(Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft)
        for col in range(3):
            self.grid.setColumnStretch(col, 1)
        self.scroll.setWidget(self.container)
        layout.addWidget(self.scroll)
        self.load()

    def load(self):
        while self.grid.count():
            item = self.grid.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        achievements = db.get_user_achievements(self.user_id)
        search_text = self.search_input.text().lower()
        category_data = self.category_combo.currentData()
        if category_data != "all":
            achievements = [a for a in achievements if a["category"] == category_data]
        if search_text:
            def _ach_matches(a):
                tn, td = db.tr_achievement(a, AppState.get_language())
                return (search_text in a["name"].lower()
                        or search_text in a["description"].lower()
                        or search_text in tn.lower()
                        or search_text in td.lower())
            achievements = [a for a in achievements if _ach_matches(a)]

        if not achievements:
            empty = _lbl(tr("achievement_empty"), "sub", 13)
            empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.grid.addWidget(empty, 0, 0)
            return

        col = 0
        row = 0
        for ach in achievements:
            card = self._make_achievement_card(ach)
            self.grid.addWidget(card, row, col)
            col += 1
            if col >= 3:
                col = 0
                row += 1
     
    def _make_achievement_card(self, ach: dict) -> QFrame:
        card = _card()
        card.setMinimumWidth(260)
        layout = QVBoxLayout(card)
        layout.setSpacing(1)
        layout.setContentsMargins(12, 12, 12, 12)
        
        # Nama & deskripsi SELALU dari translations.py via helper sentral
        name_text, desc_text = db.tr_achievement(ach, AppState.get_language())

        # Header: icon + nama
        header = QHBoxLayout()
        icon = _emoji_label(ach["icon"], ICON_CARD)
        header.addWidget(icon)
        name = QLabel(name_text)
        name.setStyleSheet(f"font-size:14px; font-weight:bold; color:{_T('text')};")
        header.addWidget(name, 1)
        layout.addLayout(header)
        
        # Deskripsi
        desc = QLabel(desc_text)
        desc.setWordWrap(True)
        desc.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        layout.addWidget(desc)
        
        # Progress bar
        progress = ach.get("progress", 0)
        req = ach["requirement_value"]
        percent = min(100, int(progress / req * 100)) if req > 0 else 0
        pb = QProgressBar()
        pb.setMaximum(req)
        pb.setValue(int(min(progress, req)))
        pb.setFormat(tr("achievement_progress_format", progress=int(progress), req=int(req), percent=percent))
        pb.setMinimumHeight(16)
        layout.addWidget(pb)
        
        # Reward info
        reward_text = tr("achievement_reward_format", xp=ach['xp_reward'], gold=ach['gold_reward'])
        reward_lbl = QLabel(reward_text)
        reward_lbl.setStyleSheet(f"color:{_T('accent')}; font-size:10px;")
        layout.addWidget(reward_lbl)
        
        # Status
        unlocked = ach.get("unlocked_at") is not None
        if unlocked:
            status = QLabel(tr("achievement_unlocked"))
            status.setStyleSheet("color:#80c000; font-weight:bold;")
            layout.addWidget(status)
            if not ach.get("claimed", 0):
                claim_btn = _btn(tr("achievement_claim"), "gold", h=28)
                claim_btn.clicked.connect(lambda _, aid=ach["id"]: self._claim_reward(aid))
                layout.addWidget(claim_btn)
        else:
            status = QLabel(tr("achievement_locked"))
            status.setStyleSheet("color:#e05050; font-size:10px;")
            layout.addWidget(status)
        
        return card
    
    def _claim_reward(self, ach_id):
        r = db.claim_achievement_reward(self.user_id, ach_id)
        if r["ok"]:
            SND.complete()
            _show(self, tr("achievement_claimed"), r["msg"], "success")
            AppState.refresh()
            QTimer.singleShot(0, self.load)
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
    
    def closeEvent(self, e):
        AppState.unregister(self.load)
        super().closeEvent(e)


# ══════════════════════════════════════════════════════════════════════════════
#  PROFILE PAGE  (instant sync — no restart needed)
# ══════════════════════════════════════════════════════════════════════════════
class ProfilePage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self._root = QVBoxLayout(self)
        self.class_buttons = {}
        self._root.setContentsMargins(20, 16, 20, 16)
        self._root.setSpacing(10)
        # Inisialisasi atribut untuk rebirth agar aman
        self.rebirth_info_label = None
        self.rebirth_conditions_label = None
        self._build()
        self.load()
        AppState.register(self.load)

    def _clear(self):
        while self._root.count():
            i = self._root.takeAt(0)
            if i.widget():
                i.widget().deleteLater()

    def _build(self):
        # ========== Buat semua widget di sini ==========
        self._root.addWidget(_lbl(tr("profile_title"), "section", 14, True))
        self._root.addWidget(_sep())

        inner = QWidget()
        il    = QVBoxLayout(inner)
        il.setSpacing(14)
        il.setContentsMargins(0, 0, 0, 0)

        # ── Avatar display ────────────────────────────────────────────────────
        av_row = QHBoxLayout()
        av_row.setSpacing(18)
        self.av_icon = QLabel("⚔️")
        self.av_icon.setFont(QFont("Segoe UI", 30))
        self.av_icon.setMinimumSize(86, 86)
        self.av_icon.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.av_icon.setStyleSheet(
            f"background: {_T('primary')};"
            f" border-radius: 12px;"
            f" border: 2px solid {_T('light')};")
        av_row.addWidget(self.av_icon)

        av_info = QVBoxLayout()
        av_info.setSpacing(4)
        self.av_display = _lbl("—", size=18, bold=True)
        self.av_username = _lbl("@username", "sub", 12)
        self.av_class_level = _lbl("Class · Level 0", "sub", 13)
        self.av_bio = _lbl("Bio", "sub", 12)
        self.av_joined = _lbl("Joined: ", "sub", 11)
        av_info.addWidget(self.av_display)
        av_info.addWidget(self.av_username)
        av_info.addWidget(self.av_class_level)
        av_info.addWidget(self.av_bio)
        av_info.addWidget(self.av_joined)
        av_row.addLayout(av_info, 1)
        il.addLayout(av_row)

        # ── Edit profil ───────────────────────────────────────────────────────
        eg = QGroupBox(tr("profile_edit"))
        el = QVBoxLayout(eg)
        el.setSpacing(8)
        self._dn  = _input(tr("profile_display_ph"))
        self._bio = _input(tr("profile_bio"))
        el.addWidget(_lbl(tr("profile_display_name_label"), size=12))
        el.addWidget(self._dn)
        el.addWidget(_lbl(tr("profile_bio_label_short"), size=12))
        el.addWidget(self._bio)
        # 🎖️ Title selector (terbuka dari statistik; tampil di leaderboard)
        el.addWidget(_lbl(tr("title_selector_label"), size=12))
        self._title_cb = QComboBox()
        el.addWidget(self._title_cb)
        self._title_hint = _lbl("", "sub", 10)
        self._title_hint.setWordWrap(True)
        el.addWidget(self._title_hint)
        self._fill_title_cb()
        el.addWidget(_btn(tr("talents_btn"), "gold", self._open_talents, 36))
        el.addWidget(_btn(tr("profile_save"), "solid", self._save_profile, 40))
        il.addWidget(eg)

        # ── Security Question ────────────────────────────────────────────────
        sqg = QGroupBox(tr("profile_security"))
        sql = QVBoxLayout(sqg)
        self.question_combo = QComboBox()
        for i in range(1, 8):
            self.question_combo.addItem(tr(f"security_q{i}"))
        self.answer_input = _input(tr("profile_security_answer"))
        save_sq = _btn(tr("profile_save_security_btn"), "solid", self._save_security)
        sql.addWidget(self.question_combo)
        sql.addWidget(self.answer_input)
        sql.addWidget(save_sq)
        il.addWidget(sqg)

        # ── Backup Codes ─────────────────────────────────────────────────────
        bcg = QGroupBox(tr("profile_backup_codes"))
        bcl = QVBoxLayout(bcg)
        self.backup_codes_label = QLabel(tr("profile_backup_desc"))
        self.backup_codes_label.setWordWrap(True)
        self.backup_codes_label.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        bcl.addWidget(self.backup_codes_label)
        self.generate_backup_btn = _btn(tr("profile_generate_backup"), "solid", self._generate_backup_codes)
        bcl.addWidget(self.generate_backup_btn)
        self.backup_codes_display = QTextEdit()
        self.backup_codes_display.setReadOnly(True)
        self.backup_codes_display.setMinimumHeight(100)
        self.backup_codes_display.setVisible(False)
        bcl.addWidget(self.backup_codes_display)
        il.addWidget(bcg)

        # ── Class picker ──────────────────────────────────────────────────────
        cg = QGroupBox(tr("profile_class"))
        cl_lay = QGridLayout(cg)
        cl_lay.setSpacing(8)
        for i, (cid, cdata) in enumerate(db.AVATAR_CLASSES.items()):
            f = _card()
            cv = QVBoxLayout(f)
            cv.setContentsMargins(10, 8, 10, 8)
            cv.setSpacing(4)
            cv.addWidget(QLabel(cdata["icon"], alignment=Qt.AlignmentFlag.AlignCenter))
            nm = QLabel(tr(f"class_{cid}_name"))
            nm.setAlignment(Qt.AlignmentFlag.AlignCenter)
            nm.setStyleSheet(f"font-size:12px; font-weight:bold; color:{_T('text')};")
            cv.addWidget(nm)
            bn = QLabel(tr(f"class_{cid}_bonus"))
            bn.setAlignment(Qt.AlignmentFlag.AlignCenter)
            bn.setWordWrap(True)
            bn.setStyleSheet(f"font-size:10px; color:{_T('muted')};")
            cv.addWidget(bn)
            # Tombol select class
            sb = _btn(tr("profile_select_class"), h=30)
            sb.clicked.connect(lambda _, c=cid: self._set_class(c))
            cv.addWidget(sb)
            self.class_buttons[cid] = sb   # ← SIMPAN REFERENSI
            cl_lay.addWidget(f, i // 3, i % 3)
        il.addWidget(cg)

        # ── Color picker ──────────────────────────────────────────────────────
        colg = QGroupBox(tr("profile_color"))
        col_lay = QHBoxLayout(colg)
        col_lay.setContentsMargins(12, 12, 12, 12)
        col_lay.setSpacing(8)
        colors = [
            ("#5a8a2e",tr("color_green")), ("#d04020",tr("color_red")), ("#4da6ff",tr("color_blue")),
            ("#f0a800", tr("color_gold")),  ("#9a50e0",tr("color_purple")),  ("#4dd9e0",tr("color_cyan")),
            ("#e8e8e8",tr("color_white")), ("#ff6a00",tr("color_orange")),
        ]
        for hex_c, name in colors:
            cb = QPushButton(name)
            cb.setMinimumHeight(36)
            cb.setStyleSheet(
                f"background:{hex_c}; color:#fff;"
                f" border:2px solid transparent; border-radius:6px;"
                f" font-size:11px; font-weight:bold;")
            cb.clicked.connect(lambda _, c=hex_c: self._set_color(c))
            col_lay.addWidget(cb)
        il.addWidget(colg)

        # ── Emoji picker ──────────────────────────────────────────────────────
        emg = QGroupBox(tr("profile_emoji"))
        em_lay = QHBoxLayout(emg)
        em_lay.setContentsMargins(12, 12, 12, 12)
        em_lay.setSpacing(6)
        emojis = ["⚔️","🧙","🏹","💊","🗡️","🛡️","🔮","🌟","👑","🐉","🦊","🐺"]
        for em in emojis:
            eb = QPushButton(em)
            eb.setMinimumSize(42, 42)
            eb.setStyleSheet(
                f"font-size:20px; background:{_T('panel')};"
                f" border:1px solid {_T('border')}; border-radius:6px;")
            eb.clicked.connect(lambda _, e=em: self._set_emoji(e))
            em_lay.addWidget(eb)
        il.addWidget(emg)

        # ── Redeem Code ──────────────────────────────────────────────────────
        rcg = QGroupBox(tr("profile_redeem"))
        rcl = QVBoxLayout(rcg)
        self.redeem_input = _input(tr("profile_redeem_placeholder"))
        redeem_btn = _btn(tr("profile_redeem_btn"), "gold", self._redeem_code)
        rcl.addWidget(self.redeem_input)
        rcl.addWidget(redeem_btn)
        rcl.addWidget(_lbl(tr("profile_redeem_desc"), "sub", 10))
        il.addWidget(rcg)

        # ── Admin badge (jika admin) ────────────────────────────────────────
        self.admin_badge = QLabel(tr("admin_mode_active"))
        self.admin_badge.setWordWrap(True)
        self.admin_badge.setStyleSheet("color:#e05050; font-weight:bold; background:#2a0808; padding:10px; border-radius:8px; margin:5px;")
        self.admin_badge.setVisible(False)
        il.addWidget(self.admin_badge)

        # ========== REBIRTH SECTION ==========
        rebirth_group = QGroupBox(tr("profile_rebirth_title"))
        rebirth_layout = QVBoxLayout(rebirth_group)

        self.rebirth_info_label = QLabel()
        self.rebirth_info_label.setWordWrap(True)
        self.rebirth_info_label.setStyleSheet(f"color: {_T('accent')}; font-weight: bold;")
        rebirth_layout.addWidget(self.rebirth_info_label)

        self.rebirth_conditions_label = QLabel()
        self.rebirth_conditions_label.setWordWrap(True)
        self.rebirth_conditions_label.setStyleSheet(f"color: {_T('muted')}; font-size: 11px;")
        rebirth_layout.addWidget(self.rebirth_conditions_label)

        rebirth_btn = _btn(tr("profile_rebirth_btn"), "diamond", self._rebirth)
        rebirth_btn.setMinimumHeight(40)
        rebirth_layout.addWidget(rebirth_btn)

        il.addWidget(rebirth_group)

        # ── Change password ───────────────────────────────────────────────────
        pg = QGroupBox(tr("profile_change_pw"))
        pl = QVBoxLayout(pg)
        pl.setSpacing(8)
        self._old_pw = _input(tr("profile_old_pw_placeholder"), True)
        self._new_pw = _input(tr("profile_new_pw_placeholder"), True)
        pl.addWidget(self._old_pw)
        pl.addWidget(self._new_pw)
        pl.addWidget(_btn(tr("profile_change_pw_btn"), "gold", self._change_pw, 40))
        il.addWidget(pg)

        # === LOCK ACCOUNT SECTION ===
        lock_group = QGroupBox(tr("profile_lock_account"))
        lock_layout = QVBoxLayout(lock_group)

        # Cek status lock
        is_locked = db.is_account_locked(self.user_id)

        if is_locked:
            # Tampilkan status locked + tombol unlock
            status_lbl = QLabel(tr("profile_account_locked"))
            status_lbl.setStyleSheet("color:#e05050; font-weight:bold;")
            lock_layout.addWidget(status_lbl)
            
            unlock_btn = _btn(tr("profile_unlock_account"), "gold", self._unlock_account, 40)
            lock_layout.addWidget(unlock_btn)
        else:
            # Tampilkan tombol lock
            lock_btn = _btn(tr("profile_lock_account"), "diamond", self._lock_account, 40)
            lock_layout.addWidget(lock_btn)

        il.addWidget(lock_group)  # masukkan ke dalam layout utama (il adalah QVBoxLayout)

        # ── Delete account ──────────────────────────────────────────────────
        delete_group = QGroupBox(tr("profile_delete_account"))
        delete_layout = QVBoxLayout(delete_group)
        delete_layout.addWidget(_lbl(tr("profile_delete_warning"), "sub", 12))
        delete_btn = _btn(tr("profile_delete_btn"), "danger", self._delete_account)
        delete_layout.addWidget(delete_btn)
        il.addWidget(delete_group)

        il.addStretch()
        self._root.addWidget(_scrolled(inner))
        fade_in(inner, 200)

    def load(self):
        if not AppState.user_id:
            return
        u = AppState.user()
        cls = db.AVATAR_CLASSES.get(u.get("avatar_class", "warrior"), {})

        # Update avatar
        self.av_icon.setText(u.get("avatar_emoji", "⚔️"))
        self.av_icon.setStyleSheet(
            f"background: {u.get('avatar_color', _T('primary'))};"
            f" border-radius: 12px;"
            f" border: 2px solid {_T('light')};")
        self.av_display.setText(u.get("display_name", "—"))
        self.av_username.setText(f"@{u.get('username','')}")
        self.av_class_level.setText(
            tr("profile_class_level", icon=cls.get('icon',''), name=cls.get('name',''), level=u['level'])
        )
        self.av_bio.setText(u.get("bio", tr("profile_no_bio")))
        self.av_joined.setText(
            (tr("profile_joined_prefix")) + f"{u.get('created_at','')[:10]}"
        )

        # Update display name & bio input fields
        self._dn.setText(u.get("display_name", ""))
        self._bio.setText(u.get("bio", ""))

        # Update admin badge
        is_admin = u.get("is_admin", 0)
        self.admin_badge.setVisible(is_admin)

        # Update class picker buttons
        # (cari child widgets di class picker dan set active label)
        class_group = self.findChild(QGroupBox, "profile_class")  # belum di-set objectName
        # Lebih mudah: kita update di load dengan mencari tombol yang aktif

        # Update rebirth info (pastikan label sudah dibuat)
        if self.rebirth_info_label is not None:
            rebirth_count = u.get("rebirth_count", 0)
            xp_bonus = rebirth_count * 10
            gold_bonus = rebirth_count * 5
            self.rebirth_info_label.setText(tr("profile_rebirth_info", count=rebirth_count, xp_bonus=xp_bonus, gold_bonus=gold_bonus))

        if self.rebirth_conditions_label is not None:
            conn = db.get_conn()
            ach_count = conn.execute(
                "SELECT COUNT(*) FROM user_achievements WHERE user_id=? AND unlocked_at IS NOT NULL",
                (self.user_id,)
            ).fetchone()[0]
            pet_count = conn.execute("SELECT COUNT(*) FROM user_pets WHERE user_id=?", (self.user_id,)).fetchone()[0]
            item_count = conn.execute("SELECT COUNT(*) FROM inventory WHERE user_id=?", (self.user_id,)).fetchone()[0]
            conn.close()
            conditions_text = (
                f"{'✅' if ach_count >= 10 else '❌'} {tr('profile_rebirth_cond_achievements', count=ach_count, need=10)}\n"
                f"{'✅' if u['level'] >= 25 else '❌'} {tr('profile_rebirth_cond_level', level=u['level'], need=25)}\n"
                f"{'✅' if pet_count >= 2 else '❌'} {tr('profile_rebirth_cond_pets', count=pet_count, need=2)}\n"
                f"{'✅' if item_count >= 6 else '❌'} {tr('profile_rebirth_cond_items', count=item_count, need=6)}"
            )
            self.rebirth_conditions_label.setText(conditions_text)

        class_group = None
        for child in self.findChildren(QGroupBox):
            if child.title() == tr("profile_class"):
                class_group = child
                break
        if class_group:
            grid = class_group.layout()
            if isinstance(grid, QGridLayout):
                for i in range(grid.count()):
                    item = grid.itemAt(i)
                    if item.widget():
                        card = item.widget()
                        if isinstance(card, QFrame):
                            for child_widget in card.findChildren(QLabel):
                                if child_widget.text() in [tr("profile_active_label"), "✔ Aktif"]:
                                    pass
                            pass

        # Update class picker buttons
        u = AppState.user()
        current_class = u.get("avatar_class", "warrior")
        for cid, btn in self.class_buttons.items():
            if cid == current_class:
                btn.setText(tr("profile_active_label"))   # "✔ Aktif"
                btn.setStyleSheet(
                    f"background: {_T('primary')}; color: #fff; "
                    f"border: 2px solid {_T('light')}; font-weight: bold;"
                )
            else:
                btn.setText(tr("profile_select_class"))
                btn.setStyleSheet("")  # reset ke default

    # ── actions ──────────────────────────────────────────────────────────────
    def _lock_account(self):
        """Lock akun: minta password, lock, lalu logout."""
        password, ok = QInputDialog.getText(
            self, tr("profile_lock_account"), tr("profile_lock_confirm"),
            QLineEdit.EchoMode.Password
        )
        if not ok or not password:
            return
        
        r = db.lock_account(self.user_id, password)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            # Logout
            clear_session()
            db.force_checkpoint()
            main_win = self.window()
            if isinstance(main_win, QMainWindow):
                main_win.close()
            # Tampilkan login
            for widget in QApplication.topLevelWidgets():
                if isinstance(widget, LoginWindow):
                    widget.show()
                    break
            else:
                login = LoginWindow()
                login.show()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _unlock_account(self):
        """Unlock akun: minta password, unlock, refresh UI."""
        password, ok = QInputDialog.getText(
            self, tr("profile_unlock_account"), tr("profile_unlock_confirm"),
            QLineEdit.EchoMode.Password
        )
        if not ok or not password:
            return
        
        r = db.unlock_account(self.user_id, password)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            # Refresh halaman profile agar tombol berubah
            self.load()
            # Refresh topbar agar status lock hilang (opsional)
            AppState.refresh()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _rebirth(self):
        # Cek syarat terlebih dahulu dengan preview
        u = AppState.user()
        conn = db.get_conn()
        ach_count = conn.execute(
            "SELECT COUNT(*) FROM user_achievements WHERE user_id=? AND unlocked_at IS NOT NULL",
            (self.user_id,)
        ).fetchone()[0]
        pet_count = conn.execute("SELECT COUNT(*) FROM user_pets WHERE user_id=?", (self.user_id,)).fetchone()[0]
        item_count = conn.execute("SELECT COUNT(*) FROM inventory WHERE user_id=?", (self.user_id,)).fetchone()[0]
        conn.close()

        # Cek apakah syarat terpenuhi
        conditions_met = (
            ach_count >= 10 and
            u["level"] >= 25 and
            pet_count >= 2 and
            item_count >= 6
        )

        if not conditions_met:
            # Tampilkan detail syarat yang belum terpenuhi
            missing = []
            if ach_count < 10:
                missing.append(tr("profile_rebirth_cond_achievements", count=ach_count, need=10))
            if u["level"] < 25:
                missing.append(tr("profile_rebirth_cond_level", level=u["level"], need=25))
            if pet_count < 2:
                missing.append(tr("profile_rebirth_cond_pets", count=pet_count, need=2))
            if item_count < 6:
                missing.append(tr("profile_rebirth_cond_items", count=item_count, need=6))
            _show(self, tr("msg_error"), "\n".join(missing), "error")
            return

        # Konfirmasi
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("profile_rebirth_confirm_title"))
        dlg.setMinimumSize(450, 350)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)

        layout.addWidget(_lbl(tr("profile_rebirth_confirm_warning"), size=14, bold=True))
        layout.addWidget(_lbl(tr("profile_rebirth_confirm_detail"), "sub", 12))
        layout.addWidget(_lbl(tr("profile_rebirth_confirm_benefit", xp=10, gold=5), "sub", 12))

        layout.addWidget(_lbl(tr("reset_confirm_type_label"), size=12))
        confirm_input = QLineEdit()
        confirm_input.setPlaceholderText(tr("profile_rebirth_confirm_placeholder"))
        confirm_input.setMinimumHeight(40)
        layout.addWidget(confirm_input)

        btn_layout = QHBoxLayout()
        cancel_btn = _btn(tr("btn_cancel"), "flat", dlg.reject)
        confirm_btn = _btn(tr("profile_rebirth_confirm_btn"), "diamond", dlg.accept)
        btn_layout.addWidget(cancel_btn)
        btn_layout.addWidget(confirm_btn)
        layout.addLayout(btn_layout)

        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        if confirm_input.text().strip().upper() != "REBIRTH":
            _show(self, tr("msg_error"), tr("reset_confirm_invalid"), "error")
            return

        # Proses Rebirth
        loading = LoadingDialog(tr("profile_rebirth_loading"), self)
        loading.show()
        QApplication.processEvents()

        try:
            r = db.perform_rebirth(self.user_id)
            loading.accept()
            if r["ok"]:
                SND.level_up()  # sukses besar
                _show(self, tr("profile_rebirth_success_title"), r["msg"], "success")
                # Refresh data user dan UI
                AppState.refresh()
                self.load()
                # Refresh semua halaman agar buff terupdate
                main_win = self.window()
                if hasattr(main_win, "_pages"):
                    for page in main_win._pages.values():
                        if hasattr(page, "load"):
                            try:
                                page.load()
                            except Exception:
                                pass
            else:
                SND.error()
                _show(self, tr("gagal_title"), r["msg"], "error")
        except Exception as e:
            loading.accept()
            SND.error()
            _show(self, tr("msg_error"), tr("profile_rebirth_error", error=str(e)), "error")

    def _redeem_code(self):
        code = self.redeem_input.text().strip().upper()
        if not code:
            _show(self, tr("msg_error"), tr("redeem_code_empty"), "error")
            return
        # FIX 4: Jika kode admin, minta password dulu
        try:
            conn = db.get_conn()
            row = conn.execute("SELECT reward_type FROM redeem_codes WHERE code=?", (code,)).fetchone()
            conn.close()
            is_admin_code = row and row["reward_type"] == "admin"
        except:
            is_admin_code = (code == "ADMINADMINADMIN")
        if is_admin_code:
            pwd, ok = QInputDialog.getText(self, tr("redeem_admin_password_title"), tr("redeem_admin_password_prompt"), QLineEdit.EchoMode.Password)
            if not ok or not pwd:
                return
            # Verifikasi password
            u = db.get_user(self.user_id)
            try:
                from database import _verify_password
                valid = _verify_password(pwd, u.get("password_hash",""))
            except:
                valid = db.login_user(u.get("username",""), pwd).get("ok", False)
            if not valid:
                SND.error()
                _show(self, tr("gagal_title"), tr("redeem_admin_password_wrong"), "error")
                return
        loading = LoadingDialog("Memverifikasi kode...", self)
        loading.show()
        QApplication.processEvents()
        r = db.redeem_code(self.user_id, code)
        loading.accept()
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            self.load()  # reload halaman
            AppState.refresh()
            # Jika menjadi admin, mungkin perlu reload seluruh window agar efek langsung terasa
            if "MODE ADMIN" in r["msg"]:
                # Refresh topbar dan nav mungkin cukup
                main_win = self.window()
                if hasattr(main_win, "_retheme"):
                    main_win._retheme()
                # Tampilkan notifikasi tambahan
                _show(self, tr("admin_mode_title"), tr("admin_mode_msg"), "warning")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        self.redeem_input.clear()

    def _delete_account(self):
        # Dialog konfirmasi
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("profile_delete_title"))
        dlg.setMinimumSize(450, 300)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)

        layout.addWidget(_lbl(tr("profile_delete_warning_label"), size=14, bold=True))
        layout.addWidget(_lbl(tr("profile_delete_warning2"), "sub", 12))
        warning_text = QLabel(
            tr("delete_account_warning_list")
        )
        warning_text.setStyleSheet(f"color: #e05050; font-size: 12px; margin: 10px;")
        layout.addWidget(warning_text)

        layout.addWidget(_lbl(tr("delete_account_confirm"), size=12))
        confirm_input = QLineEdit()
        confirm_input.setPlaceholderText(tr("profile_delete_ph"))
        confirm_input.setMinimumHeight(40)
        layout.addWidget(confirm_input)

        layout.addWidget(_lbl(tr("delete_account_password"), size=12))
        password_input = QLineEdit()
        password_input.setEchoMode(QLineEdit.EchoMode.Password)
        password_input.setMinimumHeight(40)
        layout.addWidget(password_input)

        btn_layout = QHBoxLayout()
        cancel_btn = _btn(tr("profile_delete_cancel"), "flat", dlg.reject)
        delete_btn = _btn(tr("profile_delete_permanent"), "danger", dlg.accept)
        btn_layout.addWidget(cancel_btn)
        btn_layout.addWidget(delete_btn)
        layout.addLayout(btn_layout)

        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        confirm_text = confirm_input.text().strip()
        password = password_input.text()

        if confirm_text != "DELETE ACCOUNT":
            _show(self, tr("profile_confirm_title"), tr("profile_confirm_delete_type"), "error")
            return

        # Proses delete account
        loading = LoadingDialog(tr("delete_account_loading"), self)
        loading.show()
        QApplication.processEvents()
        r = db.delete_account(self.user_id, password)
        loading.accept()

        if r["ok"]:
            SND.error()  # efek suara
            _show(self, tr("account_deleted_title"), r["msg"], "warning")
            # Hapus session
            clear_session()
            # Tutup main window dan kembali ke login
            main_win = self.window()
            if isinstance(main_win, QMainWindow):
                main_win.close()
            # Tampilkan login
            for widget in QApplication.topLevelWidgets():
                if isinstance(widget, LoginWindow):
                    widget.show()
                    break
            else:
                login = LoginWindow()
                login.show()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _open_talents(self):
        """Buka pohon talent class (dialog). Refresh profil setelah unlock."""
        TalentTreeDialog(self.user_id, self).exec()
        AppState.refresh()

    def _fill_title_cb(self):
        """Isi combo title: yang terbuka bisa dipilih, yang terkunci tampil sebagai hint."""
        if not hasattr(self, "_title_cb"):
            return
        u = db.get_user(self.user_id)
        cur = u.get("selected_title") or ""
        self._title_cb.blockSignals(True)
        self._title_cb.clear()
        self._title_cb.addItem(tr("title_none"), "")
        sel_idx = 0
        locked_hints = []
        for i, t in enumerate(db.get_unlocked_titles(self.user_id), start=1):
            if t["unlocked"]:
                self._title_cb.addItem(t["name"], t["key"])
                if t["key"] == cur:
                    sel_idx = self._title_cb.count() - 1
            else:
                locked_hints.append(
                    tr("title_locked_hint", name=t["name"],
                       target=t["target"], current=t["current"]))
        self._title_cb.setCurrentIndex(sel_idx)
        self._title_cb.blockSignals(False)
        self._title_hint.setText("\n".join(locked_hints[:4]) if locked_hints else "")

    def _save_profile(self):
        db.set_avatar(self.user_id,
                      bio=self._bio.text(),
                      display_name=self._dn.text().strip())
        if hasattr(self, "_title_cb"):
            db.set_title(self.user_id, self._title_cb.currentData())
        SND.notify()
        _show(self, tr("berhasil_title"), tr("profile_update_success"), "success")
        AppState.refresh()   # ← instant sync
        self._fill_title_cb()

    def _set_class(self, cls_id: str):
        r = db.change_class(self.user_id, cls_id)
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()

    def _set_color(self, color: str):
        db.set_avatar(self.user_id, color=color)
        SND.click()
        AppState.refresh()   # ← instant sync

    def _set_emoji(self, emoji: str):
        db.set_avatar(self.user_id, emoji=emoji)
        SND.click()
        AppState.refresh()   # ← instant sync

    def _change_pw(self):
        r = db.change_password(
            self.user_id, self._old_pw.text(), self._new_pw.text())
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            self._old_pw.clear()
            self._new_pw.clear()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _save_security(self):
        question = self.question_combo.currentText()
        answer = self.answer_input.text().strip()
        if not answer:
            _show(self, tr("msg_error"), tr("security_questions_not_empty"), "error")
            return
        db.set_security_question(self.user_id, question, answer)
        SND.notify()
        _show(self, tr("berhasil_title"), tr("security_questions_saved"), "success")
        self.answer_input.clear()

    def _generate_backup_codes(self):
        codes = db.generate_backup_codes(self.user_id, num_codes=5)
        if codes:
            msg = tr("backup_codes_intro") + "\n\n"
            for i, code in enumerate(codes, 1):
                msg += f"{i}. {code}\n"
            msg += "\n\n" + tr("backup_codes_warning")
            self.backup_codes_display.setText(msg)
            self.backup_codes_display.setVisible(True)
            SND.notify()
            _show(self, tr("backup_codes_title"), msg, "success")
        else:
            _show(self, tr("msg_error"), tr("backup_code_generate_fail"), "error")

    def closeEvent(self, e):
        AppState.unregister(self.load)
        super().closeEvent(e)


# ══════════════════════════════════════════════════════════════════════════════
#  SETTINGS PAGE  (5 themes + sound toggle)
# ══════════════════════════════════════════════════════════════════════════════
class SettingsPage(QWidget):
    theme_changed = pyqtSignal()
    language_changed = pyqtSignal()

    def __init__(self, user_id: int):
        super().__init__()
        
        self.user_id = user_id
        self._admin_was_shown = False
        self._build()
        AppState.register_lang_cb(self._retranslate)
        AppState.register(self._check_admin_panel)

    def _retranslate(self):
        self.section_title.setText(tr("settings_title"))
        self.theme_group.setTitle(tr("settings_theme"))
        self.sound_group.setTitle(tr("settings_sound"))
        self.currency_group.setTitle(tr("settings_currency"))
        self.lang_group.setTitle(tr("settings_language"))
        self.db_group.setTitle(tr("settings_database"))
        self._snd.setText(tr("settings_sound_enable"))
        self.exit_btn.setText(tr("settings_exit"))
        self.sound_hint.setText(tr("settings_sound_hint"))
        self.db_path_label.setText(tr("settings_db_path", path=db.DB_PATH))
        self.lang_combo.setItemText(0, tr("lang_id"))
        self.lang_combo.setItemText(1, tr("lang_en"))
        if hasattr(self, "a11y_group"):
            self.a11y_group.setTitle(tr("a11y_group"))
            self._a11y_font_lbl.setText(tr("a11y_font_scale"))
            self.a11y_font_hint.setText(tr("a11y_font_apply_hint"))
            self._hc.setText(tr("a11y_high_contrast"))

    def load(self):
        self._retranslate()
        self._refresh_theme_selection()

    def _check_admin_panel(self):
        """Tampilkan admin panel jika user baru menjadi admin tanpa restart."""
        is_admin = bool(AppState.user().get("is_admin", 0))
        if is_admin and not self._admin_was_shown:
            self._admin_was_shown = True
            # Cari layout container dalam scroll area
            scroll = self.findChild(QScrollArea)
            if scroll and scroll.widget():
                container_layout = scroll.widget().layout()
                if container_layout:
                    admin_group = QGroupBox(tr("admin_panel"))
                    admin_layout = QVBoxLayout(admin_group)
                    row1 = QHBoxLayout()
                    self.debug_xp = QSpinBox()
                    self.debug_xp.setRange(0, 100000)
                    self.debug_xp.setValue(1000)
                    self.debug_xp.setSuffix(tr("unit_xp"))
                    add_xp_btn = _btn(tr("admin_add_xp"), "solid", self._debug_add_xp)
                    row1.addWidget(self.debug_xp)
                    row1.addWidget(add_xp_btn)
                    self.debug_gold = QSpinBox()
                    self.debug_gold.setRange(0, 100000)
                    self.debug_gold.setValue(500)
                    self.debug_gold.setSuffix(tr("unit_gold"))
                    add_gold_btn = _btn(tr("admin_add_gold"), "solid", self._debug_add_gold)
                    row1.addWidget(self.debug_gold)
                    row1.addWidget(add_gold_btn)
                    admin_layout.addLayout(row1)
                    row2 = QHBoxLayout()
                    fill_hp_btn = _btn(tr("admin_fill_hp_mp"), "gold", self._debug_fill_hp_mp)
                    set_max_level_btn = _btn(tr("admin_max_level"), "diamond", self._debug_set_max_level)
                    row2.addWidget(fill_hp_btn)
                    row2.addWidget(set_max_level_btn)
                    admin_layout.addLayout(row2)
                    row3 = QHBoxLayout()
                    reset_tasks_btn = _btn(tr("admin_complete_tasks"), "danger", self._debug_complete_all_tasks)
                    row3.addWidget(reset_tasks_btn)
                    admin_layout.addLayout(row3)
                    pet_cheat_group = QGroupBox(tr("admin_pet_cheat"))
                    pet_cheat_layout = QVBoxLayout(pet_cheat_group)
                    btn_level_up_pets = _btn(tr("admin_pet_level_up"), "diamond", self._debug_level_up_pets)
                    pet_cheat_layout.addWidget(btn_level_up_pets)
                    exp_row = QHBoxLayout()
                    self.debug_pet_exp = QSpinBox()
                    self.debug_pet_exp.setRange(1, 10000)
                    self.debug_pet_exp.setValue(100)
                    self.debug_pet_exp.setSuffix(tr("unit_exp"))
                    add_exp_btn = _btn(tr("admin_pet_add_exp"), "gold", self._debug_add_exp_pets)
                    exp_row.addWidget(self.debug_pet_exp)
                    exp_row.addWidget(add_exp_btn)
                    pet_cheat_layout.addLayout(exp_row)
                    feed_btn = _btn(tr("admin_pet_feed"), "solid", self._debug_feed_pets)
                    pet_cheat_layout.addWidget(feed_btn)
                    admin_layout.addWidget(pet_cheat_group)
                    warn = QLabel(tr("admin_warning"))
                    warn.setWordWrap(True)
                    warn.setStyleSheet("color:#e05050; font-size:11px;")
                    admin_layout.addWidget(warn)
                    # Sisipkan sebelum stretch (index -2 biasanya stretch + exit_btn)
                    insert_idx = container_layout.count() - 2
                    if insert_idx < 0:
                        insert_idx = 0
                    container_layout.insertWidget(insert_idx, admin_group)

    def closeEvent(self, e):
        AppState.unregister(self._check_admin_panel)
        super().closeEvent(e)

    def _build(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 16, 20, 16)
        layout.setSpacing(14)

        self.section_title = _lbl("", "section", 14, True)
        layout.addWidget(self.section_title)
        layout.addWidget(_sep())

        u = db.get_user(self.user_id)
        cur = u.get("theme", "modern_dark")

        self.theme_group = QGroupBox("")
        tl = QVBoxLayout(self.theme_group)
        tl.setSpacing(10)
        self._theme_radios = {}
        for key, td in db.THEMES.items():
            row = QHBoxLayout()
            preview = QLabel("●")
            preview.setStyleSheet(f"color:{td['glow']}; font-size:16px;")
            preview.setMinimumWidth(30)
            rb = QRadioButton(f"{td['label']}   ")
            rb.setChecked(key == cur)
            rb.setStyleSheet(f"font-size:13px; color:{_T('text')};")
            rb.toggled.connect(lambda checked, k=key: self._apply(k) if checked else None)
            row.addWidget(preview)
            row.addWidget(rb, 1)
            self._theme_radios[key] = rb
            tl.addLayout(row)
        layout.addWidget(self.theme_group)

        self.sound_group = QGroupBox("")
        sl = QVBoxLayout(self.sound_group)
        self._snd = QCheckBox("")
        self._snd.setChecked(bool(u.get("sound_enabled", 1)))
        self._snd.stateChanged.connect(self._toggle_snd)
        self.sound_hint = _lbl("", "sub", 11)
        sl.addWidget(self._snd)
        sl.addWidget(self.sound_hint)
        layout.addWidget(self.sound_group)

        self.currency_group = QGroupBox("")
        cl = QVBoxLayout(self.currency_group)
        self.currency_combo = QComboBox()
        self.currency_combo.addItems([tr("currency_idr"), "USD ($)", "EUR (€)"])
        current_curr = db.get_user_currency(self.user_id)
        curr_map = {"IDR": 0, "USD": 1, "EUR": 2}
        self.currency_combo.setCurrentIndex(curr_map.get(current_curr, 0))
        self.currency_combo.currentIndexChanged.connect(self._change_currency)
        cl.addWidget(self.currency_combo)
        layout.addWidget(self.currency_group)

        self.lang_group = QGroupBox("")
        lang_layout = QVBoxLayout(self.lang_group)
        self.lang_combo = QComboBox()
        self.lang_combo.addItem("", "id")
        self.lang_combo.addItem("", "en")
        current_lang = AppState.get_language()
        idx = 0 if current_lang == "id" else 1
        self.lang_combo.setCurrentIndex(idx)
        self.lang_combo.currentIndexChanged.connect(self._change_language)
        lang_layout.addWidget(self.lang_combo)
        layout.addWidget(self.lang_group)

        # ── ♿ Aksesibilitas: skala font + kontras tinggi ──────────────────────
        self.a11y_group = QGroupBox("")
        al = QVBoxLayout(self.a11y_group)
        al.setSpacing(8)
        frow = QHBoxLayout()
        self._a11y_font_lbl = _lbl("", size=12)
        frow.addWidget(self._a11y_font_lbl)
        self.font_scale_combo = QComboBox()
        self._font_scale_opts = [80, 90, 100, 110, 120, 130, 140]
        for pct in self._font_scale_opts:
            self.font_scale_combo.addItem(f"{pct}%", pct)
        try:
            cur_scale = int(u.get("font_scale", 100) or 100)
        except (TypeError, ValueError):
            cur_scale = 100
        if cur_scale not in self._font_scale_opts:
            cur_scale = 100
        self.font_scale_combo.setCurrentIndex(self._font_scale_opts.index(cur_scale))
        self.font_scale_combo.currentIndexChanged.connect(self._change_font_scale)
        frow.addWidget(self.font_scale_combo, 1)
        al.addLayout(frow)
        self.a11y_font_hint = _lbl("", "sub", 10)
        self.a11y_font_hint.setWordWrap(True)
        al.addWidget(self.a11y_font_hint)
        self._hc = QCheckBox("")
        self._hc.setChecked(bool(u.get("high_contrast", 0)))
        self._hc.stateChanged.connect(self._toggle_high_contrast)
        al.addWidget(self._hc)
        layout.addWidget(self.a11y_group)

        reset_group = QGroupBox(tr("settings_reset_progress"))
        reset_layout = QVBoxLayout(reset_group)
        reset_layout.addWidget(_lbl(tr("settings_reset_warning"), "sub", 12))
        reset_btn = _btn(tr("settings_reset_btn"), "danger", self._reset_progress)
        reset_layout.addWidget(reset_btn)
        layout.addWidget(reset_group)

        if AppState.user().get("is_admin", 0):
            self._admin_was_shown = True
            admin_group = QGroupBox(tr("admin_panel"))
            admin_layout = QVBoxLayout(admin_group)
            row1 = QHBoxLayout()
            self.debug_xp = QSpinBox()
            self.debug_xp.setRange(0, 100000)
            self.debug_xp.setValue(1000)
            self.debug_xp.setSuffix(tr("unit_xp"))
            add_xp_btn = _btn(tr("admin_add_xp"), "solid", self._debug_add_xp)
            row1.addWidget(self.debug_xp)
            row1.addWidget(add_xp_btn)
            self.debug_gold = QSpinBox()
            self.debug_gold.setRange(0, 100000)
            self.debug_gold.setValue(500)
            self.debug_gold.setSuffix(tr("unit_gold"))
            add_gold_btn = _btn(tr("admin_add_gold"), "solid", self._debug_add_gold)
            row1.addWidget(self.debug_gold)
            row1.addWidget(add_gold_btn)
            admin_layout.addLayout(row1)
            row2 = QHBoxLayout()
            fill_hp_btn = _btn(tr("admin_fill_hp_mp"), "gold", self._debug_fill_hp_mp)
            set_max_level_btn = _btn(tr("admin_max_level"), "diamond", self._debug_set_max_level)
            row2.addWidget(fill_hp_btn)
            row2.addWidget(set_max_level_btn)
            admin_layout.addLayout(row2)
            row3 = QHBoxLayout()
            reset_tasks_btn = _btn(tr("admin_complete_tasks"), "danger", self._debug_complete_all_tasks)
            row3.addWidget(reset_tasks_btn)
            admin_layout.addLayout(row3)
            pet_cheat_group = QGroupBox(tr("admin_pet_cheat"))
            pet_cheat_layout = QVBoxLayout(pet_cheat_group)
            btn_level_up_pets = _btn(tr("admin_pet_level_up"), "diamond", self._debug_level_up_pets)
            pet_cheat_layout.addWidget(btn_level_up_pets)
            exp_row = QHBoxLayout()
            self.debug_pet_exp = QSpinBox()
            self.debug_pet_exp.setRange(1, 10000)
            self.debug_pet_exp.setValue(100)
            self.debug_pet_exp.setSuffix(tr("unit_exp"))
            add_exp_btn = _btn(tr("admin_pet_add_exp"), "gold", self._debug_add_exp_pets)
            exp_row.addWidget(self.debug_pet_exp)
            exp_row.addWidget(add_exp_btn)
            pet_cheat_layout.addLayout(exp_row)
            feed_btn = _btn(tr("admin_pet_feed"), "solid", self._debug_feed_pets)
            pet_cheat_layout.addWidget(feed_btn)
            admin_layout.addWidget(pet_cheat_group)
            warn = QLabel(tr("admin_warning"))
            warn.setWordWrap(True)
            warn.setStyleSheet("color:#e05050; font-size:11px;")
            admin_layout.addWidget(warn)
            layout.addWidget(admin_group)

        # ── Data Management ──
        data_group = QGroupBox(tr("settings_data_management"))
        data_layout = QVBoxLayout(data_group)

        export_btn = _btn(tr("settings_export_tracker"), "solid", self._export_tracker)
        import_btn = _btn(tr("settings_import_tracker"), "diamond", self._import_tracker)

        # Nonaktifkan untuk admin
        is_admin = bool(AppState.user().get("is_admin", 0))
        export_btn.setEnabled(not is_admin)
        import_btn.setEnabled(not is_admin)
        if is_admin:
            export_btn.setToolTip(tr("admin_export_blocked"))
            import_btn.setToolTip(tr("admin_import_blocked"))

        data_layout.addWidget(export_btn)
        data_layout.addWidget(import_btn)
        data_layout.addStretch()
        layout.addWidget(data_group)

        self.db_group = QGroupBox("")
        dl = QVBoxLayout(self.db_group)
        self.db_path_label = _lbl("", "sub", 11)
        dl.addWidget(self.db_path_label)
        layout.addWidget(self.db_group)
        backup_btn = _btn(tr("settings_backup_now"), "solid", self._manual_backup)
        dl.addWidget(backup_btn)

        layout.addStretch()
        self.exit_btn = _btn("", "danger", QApplication.instance().quit)
        layout.addWidget(self.exit_btn)

        scroll.setWidget(container)
        main_layout.addWidget(scroll)
        self._retranslate()

    def _export_tracker(self):
        """Ekspor data tracker ke file JSON."""
        try:
            data = db.export_tracker_data(self.user_id)
        except PermissionError as e:
            _show(self, tr("msg_error"), str(e), "error")
            return
        except Exception as e:
            _show(self, tr("msg_error"), tr("export_failed", error=str(e)), "error")
            return
        
        path, _ = QFileDialog.getSaveFileName(
            self,
            tr("save_json"),
            f"craftlife_tracker_{datetime.now().strftime('%Y%m%d')}.json",
            "JSON Files (*.json)"
        )
        if not path:
            return
        if not path.endswith(".json"):
            path += ".json"
        
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            _show(self, tr("berhasil_title"), tr("export_success"), "success")
        except Exception as e:
            _show(self, tr("msg_error"), tr("export_failed", error=str(e)), "error")

    def _import_tracker(self):
        """Impor data tracker dari file JSON."""
        path, _ = QFileDialog.getOpenFileName(
            self,
            tr("open_json"),
            "",
            "JSON Files (*.json)"
        )
        if not path:
            return
        
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            _show(self, tr("msg_error"), tr("invalid_json"), "error")
            return
        
        if "version" not in data or "tables" not in data:
            _show(self, tr("msg_error"), tr("invalid_format"), "error")
            return
        
        reply = QMessageBox.question(
            self,
            tr("confirm_title"),
            tr("import_confirm_warning"),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        try:
            db.import_tracker_data(self.user_id, data)
            _show(self, tr("berhasil_title"), tr("import_success"), "success")
            # Refresh semua halaman agar data baru tampil
            main = self.window()
            if hasattr(main, "_pages"):
                for page in main._pages.values():
                    if hasattr(page, "load"):
                        try:
                            page.load()
                        except Exception:
                            pass
            if hasattr(main, "_topbar"):
                main._topbar.refresh()
        except PermissionError as e:
            _show(self, tr("msg_error"), str(e), "error")
        except Exception as e:
            _show(self, tr("msg_error"), tr("import_failed", error=str(e)), "error")

    def _manual_backup(self):
        path = db.backup_database()
        if path:
            _show(self, tr("berhasil_title"), f"Backup berhasil disimpan di:\n{path}", "success")
        else:
            _show(self, tr("gagal_title"), "Backup gagal!", "error")

    def _reset_progress(self):
        # FIX 5: Minta password akun dulu sebelum reset (dengan translasi)
        pwd, ok = QInputDialog.getText(self, tr("reset_verify_password_title"), tr("reset_verify_password_prompt"), QLineEdit.EchoMode.Password)
        if not ok or not pwd:
            return
        u = db.get_user(self.user_id)
        try:
            from database import _verify_password
            valid = _verify_password(pwd, u.get("password_hash",""))
        except:
            valid = db.login_user(u.get("username",""), pwd).get("ok", False)
        if not valid:
            SND.error()
            _show(self, tr("gagal_title"), tr("redeem_admin_password_wrong"), "error")
            return
        # Dialog konfirmasi dengan input verifikasi
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("reset_confirm_title"))
        dlg.setMinimumSize(450, 300)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)

        layout.addWidget(_lbl(tr("reset_confirm_warning"), size=14, bold=True))
        layout.addWidget(_lbl(tr("reset_confirm_detail"), "sub", 12))
        layout.addWidget(_lbl(tr("reset_confirm_type_label"), size=12))

        confirm_input = QLineEdit()
        confirm_input.setPlaceholderText(tr("reset_confirm_placeholder"))
        confirm_input.setMinimumHeight(40)
        layout.addWidget(confirm_input)

        btn_layout = QHBoxLayout()
        cancel_btn = _btn(tr("btn_cancel"), "flat", dlg.reject)
        reset_btn = _btn(tr("reset_confirm_btn"), "danger", dlg.accept)
        btn_layout.addWidget(cancel_btn)
        btn_layout.addWidget(reset_btn)
        layout.addLayout(btn_layout)

        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        if confirm_input.text().strip().upper() != "RESET PROGRESS":
            _show(self, tr("msg_error"), tr("reset_confirm_invalid"), "error")
            return

        # Proses reset
        loading = LoadingDialog(tr("reset_loading"), self)
        loading.show()
        QApplication.processEvents()

        try:
            db.reset_user_progress(self.user_id)
            SND.error()
            _show(self, tr("reset_success_title"), tr("reset_success_msg"), "warning")
            clear_session()
            main_win = self.window()
            if isinstance(main_win, QMainWindow):
                main_win.close()
            for widget in QApplication.topLevelWidgets():
                if isinstance(widget, LoginWindow):
                    widget.show()
                    break
            else:
                login = LoginWindow()
                login.show()
        except Exception as e:
            SND.error()
            _show(self, tr("msg_error"), tr("reset_error", error=str(e)), "error")
        finally:
            loading.accept()

    def _change_language(self, idx):
        new_lang = self.lang_combo.itemData(idx)
        current_lang = AppState.get_language()
        if new_lang == current_lang:
            return

        # Ambil teks dalam bahasa baru (tanpa mengubah bahasa saat ini)
        title = get_text("settings_language_restart_title", new_lang)
        msg = get_text("settings_language_restart_msg", new_lang)
        yes_btn = get_text("settings_language_restart_yes", new_lang)
        no_btn = get_text("settings_language_restart_no", new_lang)

        reply = QMessageBox.question(
            self,
            title,
            msg,
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            # Simpan bahasa ke database dan AppState
            AppState.set_language(new_lang)
            db.set_user_language(AppState.user_id, new_lang)

            # Restart aplikasi
            self._restart_app()

    def _restart_app(self):
        """Restart aplikasi menggunakan QProcess."""
        from PyQt6.QtCore import QProcess
        import sys

        # Path ke executable Python dan skrip utama
        if getattr(sys, 'frozen', False):
            # Jika dijalankan sebagai .exe
            program = sys.executable
            arguments = []
        else:
            # Jika dijalankan sebagai skrip Python
            program = sys.executable
            arguments = [sys.argv[0]]

        QProcess.startDetached(program, arguments)
        QApplication.quit()

    def _prompt_restart(self):
        """Prompt restart generik untuk perubahan pengaturan (tema, font,
        kontras, suara, mata uang) — pola yang sama seperti pergantian bahasa."""
        reply = QMessageBox.question(
            self,
            tr("settings_language_restart_title"),
            tr("settings_change_restart_msg"),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            self._restart_app()

    def _apply(self, key: str):
        db.set_user_theme(self.user_id, key)
        t = db.THEMES[key]
        apply_theme(t)
        self.theme_changed.emit()
        self._prompt_restart()

    def _refresh_theme_selection(self):
        """Sinkronkan radio tema di Settings dengan tema aktif (mis. setelah toggle TopBar)."""
        cur = db.get_user(self.user_id).get("theme", "modern_dark")
        radios = getattr(self, "_theme_radios", {})
        for key, rb in radios.items():
            rb.blockSignals(True)
            rb.setChecked(key == cur)
            rb.blockSignals(False)

    def _toggle_snd(self, state: int):
        enabled = bool(state)
        SoundEngine.enabled = enabled
        db.set_user_settings(self.user_id, sound_enabled=enabled)
        if enabled:
            SND.notify()
        self._prompt_restart()

    def _change_currency(self, idx):
        curr = ["IDR", "USD", "EUR"][idx]
        db.set_user_currency(self.user_id, curr)
        SND.notify()
        self.theme_changed.emit()
        self._prompt_restart()

    def _change_font_scale(self, idx):
        """Simpan skala font + terapkan parsial (stylesheet & font app),
        lalu tawarkan restart agar perubahan diterapkan sepenuhnya."""
        pct = self.font_scale_combo.currentData()
        if pct is None:
            return
        db.set_font_scale(self.user_id, pct)
        win = self.window()
        if hasattr(win, "_retheme"):
            win._retheme(show_msg=False)
        SND.notify()
        self._prompt_restart()

    def _toggle_high_contrast(self, state):
        db.set_high_contrast(self.user_id, bool(state))
        win = self.window()
        if hasattr(win, "_retheme"):
            win._retheme(show_msg=False)
        SND.notify()
        self._prompt_restart()

    def _debug_add_xp(self):
        amount = self.debug_xp.value()
        r = db.gain_xp_gold(self.user_id, amount, 0)
        if r["ok"]:
            SND.complete()
            _show(self, tr("debug_title"), tr("debug_xp_added", amount=amount), "success")
            AppState.refresh()
        else:
            SND.error()
            _show(self, tr("msg_error"), r.get("msg", "Gagal"), "error")

    def _debug_add_gold(self):
        amount = self.debug_gold.value()
        r = db.gain_xp_gold(self.user_id, 0, amount)
        if r["ok"]:
            SND.complete()
            _show(self, tr("debug_title"), tr("debug_gold_added", amount=amount), "success")
            AppState.refresh()
        else:
            SND.error()
            _show(self, tr("msg_error"), r.get("msg", "Gagal"), "error")

    def _debug_fill_hp_mp(self):
        u = AppState.user()
        db.update_user(self.user_id, hp=u["max_hp"], mp=u["max_mp"])
        SND.complete()
        _show(self, tr("debug_title"), tr("debug_hp_mp_restored"), "success")
        AppState.refresh()

    def _debug_set_max_level(self):
        target_level = 50
        u = AppState.user()
        current_level = u["level"]
        if current_level >= target_level:
            _show(self, tr("info_title"), tr("debug_level_already", level=current_level), "info")
            return
        needed_xp = 0
        for lvl in range(current_level, target_level):
            needed_xp += lvl * 150
        db.gain_xp_gold(self.user_id, needed_xp, 0)
        _show(self, tr("debug_title"), tr("debug_level_set", target=target_level), "success")
        AppState.refresh()

    def _debug_complete_all_tasks(self):
        habits = db.get_habits(self.user_id)
        for h in habits:
            if not h["done_today"]:
                db.complete_habit(self.user_id, h["id"], "up")
        dailies = db.get_dailies(self.user_id)
        for d in dailies:
            if not d["done_today"]:
                db.complete_daily(self.user_id, d["id"])
        todos = db.get_todos(self.user_id)
        for t in todos:
            if not t["done"]:
                db.complete_todo(self.user_id, t["id"])
        SND.complete()
        _show(self, tr("debug_title"), tr("debug_tasks_done"), "success")
        AppState.refresh()

    def _debug_level_up_pets(self):
        loading = LoadingDialog("Menaikkan level pet...", self)
        loading.show()
        QApplication.processEvents()
        r = db.admin_level_up_all_pets(self.user_id)
        loading.accept()
        if r["ok"]:
            SND.complete()
            _show(self, tr("cheat_title"), r["msg"], "success")
            AppState.refresh()
            main_win = self.window()
            if hasattr(main_win, "_pages") and "pets" in main_win._pages:
                main_win._pages["pets"].load()
        else:
            SND.error()
            _show(self, tr("msg_error"), r["msg"], "error")

    def _debug_add_exp_pets(self):
        amount = self.debug_pet_exp.value()
        loading = LoadingDialog("Menambah EXP pet...", self)
        loading.show()
        QApplication.processEvents()
        r = db.admin_add_exp_all_pets(self.user_id, amount)
        loading.accept()
        if r["ok"]:
            SND.complete()
            _show(self, tr("cheat_title"), r["msg"], "success")
            AppState.refresh()
            main_win = self.window()
            if hasattr(main_win, "_pages") and "pets" in main_win._pages:
                main_win._pages["pets"].load()
        else:
            SND.error()
            _show(self, tr("msg_error"), r["msg"], "error")

    def _debug_feed_pets(self):
        loading = LoadingDialog("Mengisi hunger pet...", self)
        loading.show()
        QApplication.processEvents()
        r = db.admin_feed_all_pets(self.user_id)
        loading.accept()
        if r["ok"]:
            SND.complete()
            _show(self, tr("cheat_title"), r["msg"], "success")
            main_win = self.window()
            if hasattr(main_win, "_pages") and "pets" in main_win._pages:
                main_win._pages["pets"].load()
        else:
            SND.error()
            _show(self, tr("msg_error"), r["msg"], "error")

# ══════════════════════════════════════════════════════════════════════════════
#  NOTIFICATION POPUP
# ══════════════════════════════════════════════════════════════════════════════
class NotifPopup(QDialog):
    def __init__(self, user_id: int, parent=None):
        super().__init__(parent)
        
        self.setWindowTitle(tr("notif_window_title"))
        self.setMinimumSize(420, 400)
        self.setStyleSheet(build_ss())
        lay = QVBoxLayout(self)
        lay.setContentsMargins(20, 20, 20, 20)
        lay.setSpacing(10)
        lay.addWidget(_lbl(tr("notif_title"), "section", 14, True))
        lw = QListWidget()
        notifs = db.get_notifications(user_id, unread_only=False)
        db.mark_read(user_id)
        if not notifs:
            lw.addItem(tr("notif_empty"))
        else:
            icons = {"levelup": "🎉", "success": "✅",
                     "danger":  "💀", "info":    "💬",
                     "warning": "⚠️"}
            for n in notifs:
                ic = icons.get(n["type"], "💬")
                lw.addItem(
                    f"{ic}  {n['message']}\n"
                    f"      {n['created_at'][:16]}")
        lay.addWidget(lw)
        lay.addWidget(_btn(tr("notif_close"), "solid", self.accept, 40))
        # fade only the inner widget, not the dialog
        fade_in(lw, 200)

# ══════════════════════════════════════════════════════════════════════════════
#  LeaderBoard 
# ══════════════════════════════════════════════════════════════════════════════
class LeaderboardPage(QWidget):
    def __init__(self):
        super().__init__()
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 16, 20, 16)
        layout.setSpacing(10)
        layout.addWidget(_lbl(tr("leaderboard_title"), "section", 14, True))
        layout.addWidget(_sep())

        self.table = QTableWidget()
        self.table.setColumnCount(7)
        self._update_headers()

        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.table.setSelectionMode(QTableWidget.SelectionMode.NoSelection)
        self.table.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        self.table.horizontalHeader().setSectionsMovable(False)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table.verticalHeader().setVisible(False)
        self.table.setAlternatingRowColors(True)

        layout.addWidget(self.table)
        AppState.register(self.load)
        AppState.register_lang_cb(self._retranslate)
        self.load()

    def _update_headers(self):
        self.table.setHorizontalHeaderLabels([
            tr("leaderboard_col_user"),
            tr("leaderboard_col_level"),
            tr("leaderboard_col_xp"),
            tr("leaderboard_col_gold"),
            tr("leaderboard_col_sport"),
            tr("leaderboard_col_pet"),
            tr("leaderboard_col_rebirth")
        ])

    def _retranslate(self):
        self._update_headers()

    def load(self):
        # Terapkan style agar background sesuai tema
        self.table.setStyleSheet(f"""
            QTableWidget {{
                background-color: {_T("panel")};
                alternate-background-color: {_T("border")};
                gridline-color: {_T("border")};
                border: 1px solid {_T("border")};
                border-radius: 6px;
                color: {_T("text")};
            }}
            QHeaderView::section {{
                background-color: {_T("primary")};
                color: {_T("light")};
                padding: 8px;
                font-weight: bold;
                border: 1px solid {_T("border")};
            }}
        """)

        data = db.get_leaderboard_for_user(AppState.user_id)
        self.table.setRowCount(len(data))
        # Map title key → nama terlokalisasi
        lang = AppState.get_language()
        tmap = {t["key"]: (t["name"][0] if lang == "id" else t["name"][1])
                for t in db.TITLES}
        for i, r in enumerate(data):
            name_txt = r["display_name"] or r["username"]
            title_txt = tmap.get(r.get("selected_title") or "", "")
            if title_txt:
                name_txt = f"{name_txt}  {title_txt}"
            self.table.setItem(i, 0, QTableWidgetItem(name_txt))
            self.table.setItem(i, 1, QTableWidgetItem(str(r["level"])))
            self.table.setItem(i, 2, QTableWidgetItem(str(r["total_xp_earned"])))
            self.table.setItem(i, 3, QTableWidgetItem(f"{r['gold']:.0f}"))
            sport_lv = r.get("sport_level", 1) or 1
            sport_item = QTableWidgetItem(f"Lv.{sport_lv}")
            sport_item.setForeground(QColor("#f0a800"))
            self.table.setItem(i, 4, sport_item)
            self.table.setItem(i, 5, QTableWidgetItem(str(r["pet_count"])))
            rebirth_item = QTableWidgetItem(str(r.get("rebirth_count", 0)))
            rebirth_item.setForeground(QColor("#ff6b00"))
            self.table.setItem(i, 6, rebirth_item)

    def closeEvent(self, e):
        AppState.unregister(self.load)
        AppState.unregister_lang_cb(self._retranslate)
        super().closeEvent(e)

# ══════════════════════════════════════════════════════════════════════════════
#  Friends Page (Add Friend System)
# ══════════════════════════════════════════════════════════════════════════════
class FriendsPage(QWidget):
    def __init__(self, user_id):
        super().__init__()
        
        self.user_id = user_id
        self.main_layout = None  # referensi layout utama
        self._build()
        AppState.register(self.load)

    def _build(self):
        self.main_layout = QVBoxLayout(self)
        self.main_layout.setContentsMargins(20, 16, 20, 16)
        self.main_layout.setSpacing(10)
        self.main_layout.addWidget(_lbl(tr("friends_title"), "section", 14, True))
        self.main_layout.addWidget(_sep())

        # Form tambah teman
        add_form = QHBoxLayout()
        self.friend_username = _input(tr("friends_add_placeholder"))
        add_btn = _btn(tr("friends_add_btn"), "solid", self._send_request)
        add_form.addWidget(self.friend_username)
        add_form.addWidget(add_btn)
        self.main_layout.addLayout(add_form)

        # Daftar permintaan masuk
        self.pending_group = QGroupBox(tr("friends_pending"))
        self.pending_layout = QVBoxLayout(self.pending_group)
        self.main_layout.addWidget(self.pending_group)

        # Daftar teman
        self.friends_group = QGroupBox(tr("friends_list"))
        self.friends_layout = QVBoxLayout(self.friends_group)
        self.main_layout.addWidget(self.friends_group)

        # ⚔️ PvP Streak Battle
        self.pvp_group = QGroupBox(tr("pvp_section"))
        self.pvp_layout = QVBoxLayout(self.pvp_group)
        self.main_layout.addWidget(self.pvp_group)

        self.main_layout.addStretch()
        self.load()

    def _clear_layout(self):
        """Hapus semua widget dan sub-layout dari layout utama."""
        if self.main_layout:
            while self.main_layout.count():
                item = self.main_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
                elif item.layout():
                    # Bersihkan sub-layout juga
                    sub_layout = item.layout()
                    while sub_layout.count():
                        sub_item = sub_layout.takeAt(0)
                        if sub_item.widget():
                            sub_item.widget().deleteLater()
                    sub_layout.deleteLater()

    def load(self):
        if not AppState.user_id:
            return

        # ========== CEK ADMIN ==========
        if AppState.user().get("is_admin", 0):
            self._clear_layout()
            msg = QLabel(tr("friends_admin_block"))
            msg.setAlignment(Qt.AlignmentFlag.AlignCenter)
            msg.setStyleSheet(f"color:{_T('muted')}; font-size:14px; padding:40px;")
            self.main_layout.addWidget(msg)
            return

        # ========== NON-ADMIN: TAMPILKAN UI NORMAL ==========
        # Bersihkan isi pending_layout dan friends_layout (tapi jangan hapus group box-nya)
        while self.pending_layout.count():
            item = self.pending_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
            elif item.layout():
                sub = item.layout()
                while sub.count():
                    sub_item = sub.takeAt(0)
                    if sub_item.widget():
                        sub_item.widget().deleteLater()
                sub.deleteLater()

        while self.friends_layout.count():
            item = self.friends_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
            elif item.layout():
                sub = item.layout()
                while sub.count():
                    sub_item = sub.takeAt(0)
                    if sub_item.widget():
                        sub_item.widget().deleteLater()
                sub.deleteLater()

        # Pending requests
        pending = db.get_pending_friend_requests(self.user_id)
        if pending:
            self.pending_group.setVisible(True)
            for req in pending:
                row = QHBoxLayout()
                row.addWidget(QLabel(f"📨 {req['display_name']} (@{req['username']})"))
                accept = _btn(tr("guild_accept"), h=28)
                accept.clicked.connect(lambda _, rid=req["id"]: self._accept(rid))
                reject = _btn(tr("guild_reject"), "danger", h=28)
                reject.clicked.connect(lambda _, rid=req["id"]: self._reject(rid))
                row.addWidget(accept)
                row.addWidget(reject)
                self.pending_layout.addLayout(row)
        else:
            self.pending_group.setVisible(False)

        # Friends list
        friends = db.get_friends(self.user_id)
        if friends:
            self.friends_group.setVisible(True)
            for f in friends:
                row = QHBoxLayout()
                row.addWidget(QLabel(tr("friend_level_format", emoji=f['avatar_emoji'], name=f['display_name'], level=f['level'])))
                row.addStretch()
                pvp_btn = _btn("⚔️", "gold", h=28)
                pvp_btn.setToolTip(tr("pvp_btn"))
                pvp_btn.clicked.connect(lambda _, fid=f["id"]: self._challenge_pvp(fid))
                chat_btn = _btn(tr("friends_chat_btn"), h=28)
                chat_btn.clicked.connect(lambda _, fid=f["id"], name=f["display_name"]: self._open_chat(fid, name))
                profile_btn = _btn("👤", h=28)
                profile_btn.clicked.connect(lambda _, fid=f["id"]: self._view_profile(fid))
                kick_btn = _btn(tr("friends_remove_btn"), "danger", h=28)
                kick_btn.clicked.connect(lambda _, fid=f["id"]: self._remove_friend(fid))
                row.addWidget(chat_btn)
                row.addWidget(profile_btn)
                row.addWidget(pvp_btn)
                row.addWidget(kick_btn)
                self.friends_layout.addLayout(row)
        else:
            self.friends_group.setVisible(False)

        self._render_pvp()

    # ── ⚔️ PvP STREAK BATTLE ─────────────────────────────────────────────────
    def _clear_pvp_layout(self):
        while self.pvp_layout.count():
            item = self.pvp_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
            elif item.layout():
                sub = item.layout()
                while sub.count():
                    si = sub.takeAt(0)
                    if si.widget():
                        si.widget().deleteLater()

    def _render_pvp(self):
        self._clear_pvp_layout()
        try:
            items = db.get_pvp_challenges(self.user_id)
        except Exception:
            items = []
        u = AppState.user() or {}
        my_name = u.get("display_name") or u.get("username", "")
        shown = 0
        for it in items:
            if it["status"] == "declined":
                continue
            shown += 1
            if shown > 6:
                break
            card = _card()
            box = QVBoxLayout(card)
            box.setSpacing(6)
            status = it["status"]
            if status == "pending" and not it["is_challenger"]:
                box.addWidget(_lbl(f"⚔️ {it['opponent_name']}", size=12, bold=True))
                btns = QHBoxLayout()
                acc = _btn(tr("pvp_accept"), "solid", h=30)
                acc.clicked.connect(
                    lambda _, cid=it["id"]: self._respond_pvp(cid, True))
                dec = _btn(tr("pvp_decline"), "danger", h=30)
                dec.clicked.connect(
                    lambda _, cid=it["id"]: self._respond_pvp(cid, False))
                btns.addWidget(acc)
                btns.addWidget(dec)
                box.addLayout(btns)
            elif status == "pending":
                lbl = _lbl(tr("pvp_pending_out", name=it["opponent_name"]),
                           "sub", 11)
                lbl.setWordWrap(True)
                box.addWidget(lbl)
            elif status == "active":
                lbl = _lbl(tr("pvp_score_line", me=my_name,
                              opp=it["opponent_name"],
                              ms=it.get("my_score", 0),
                              os=it.get("opponent_score", 0),
                              days=it.get("days_left", 0)), size=11)
                lbl.setWordWrap(True)
                box.addWidget(lbl)
            elif status == "finished":
                if it.get("winner_id") is None:
                    key = "pvp_finished_tie"
                elif it.get("winner_id") == self.user_id:
                    key = "pvp_finished_win"
                else:
                    key = "pvp_finished_lose"
                lbl = _lbl(tr(key, name=it["opponent_name"],
                              ms=it.get("my_score", 0),
                              os=it.get("opponent_score", 0)), size=11)
                lbl.setWordWrap(True)
                box.addWidget(lbl)
            else:
                continue
            self.pvp_layout.addWidget(card)
        if shown == 0:
            none_lbl = _lbl(tr("pvp_none"), "sub", 11)
            none_lbl.setWordWrap(True)
            self.pvp_layout.addWidget(none_lbl)

    def _respond_pvp(self, challenge_id: int, accept: bool):
        r = db.respond_pvp_challenge(challenge_id, self.user_id, accept)
        if r.get("ok"):
            SND.notify()
        QTimer.singleShot(0, self.load)

    def _challenge_pvp(self, friend_id: int):
        r = db.send_pvp_challenge(self.user_id, friend_id)
        self._show_result(r)
        QTimer.singleShot(0, self.load)

    def _open_chat(self, friend_id, friend_name):
        if AppState.user().get("is_admin", 0):
            _show(self, tr("info_title"), tr("chat_admin_block"), "warning")
            return
        dlg = ChatDialog(self.user_id, friend_id, friend_name, self)
        dlg.exec()

    def _view_profile(self, friend_id):
        dlg = FriendProfileDialog(friend_id, self)
        dlg.exec()

    def _remove_friend(self, friend_id):
        r = db.remove_friend(self.user_id, friend_id)
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        QTimer.singleShot(0, self.load)

    def _send_request(self):
        username = self.friend_username.text().strip()
        if not username:
            _show(self, tr("msg_error"), tr("msg_enter_username"), "error")
            return
        r = db.send_friend_request(self.user_id, username)
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            self.friend_username.clear()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _accept(self, req_id):
        r = db.accept_friend_request(self.user_id, req_id)
        self._show_result(r)
        QTimer.singleShot(0, self.load)

    def _reject(self, req_id):
        r = db.reject_friend_request(self.user_id, req_id)
        self._show_result(r)
        QTimer.singleShot(0, self.load)

    def _show_result(self, r):
        if r["ok"]:
            SND.notify()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def closeEvent(self, e):
        AppState.unregister(self.load)
        super().closeEvent(e)

# ══════════════════════════════════════════════════════════════════════════════
#  ChatDialog (Direct Messaging System with real-time updates)
# ══════════════════════════════════════════════════════════════════════════════
class ChatDialog(QDialog):
    def __init__(self, user_id, friend_id, friend_name, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.friend_id = friend_id
        self.setWindowTitle(tr("chat_with", name=friend_name))
        self.setMinimumSize(400, 500)
        self.setStyleSheet(build_ss())
        self._build()
        self._load_messages()
        self.timer = QTimer()
        self.timer.timeout.connect(self._load_messages)
        self.timer.start(3000)

    def _build(self):
        layout = QVBoxLayout(self)
        
        # Tombol Clear Chat
        btn_layout = QHBoxLayout()
        clear_btn = _btn(tr("chat_clear_self"), "danger", self._clear_chat, 32)
        btn_layout.addStretch()
        btn_layout.addWidget(clear_btn)
        layout.addLayout(btn_layout)
        
        self.chat_area = QTextEdit()
        self.chat_area.setReadOnly(True)
        layout.addWidget(self.chat_area)
        
        input_layout = QHBoxLayout()
        self.message_input = QLineEdit()
        send_btn = _btn(tr("chat_send_btn"), "solid", self._send_message)
        input_layout.addWidget(self.message_input)
        input_layout.addWidget(send_btn)
        layout.addLayout(input_layout)

    def _clear_chat(self):
        reply = QMessageBox.question(self, tr("confirm_title"), tr("chat_clear_confirm_self"), 
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.clear_friend_chat(self.user_id, self.friend_id)
            self._load_messages()
            SND.notify()

    def _load_messages(self):
        messages = db.get_messages(self.user_id, self.friend_id)
        self.chat_area.clear()
        for m in messages:
            sender = tr("chat_you") if m["sender_id"] == self.user_id else tr("chat_friend")
            self.chat_area.append(tr("chat_message_format", time=m['created_at'][11:16], sender=sender, message=m['message']))
        db.mark_messages_read(self.user_id, self.friend_id)

    def _send_message(self):
        msg = self.message_input.text().strip()
        if msg:
            server_now = TimeSync.get_current_time().isoformat()
            db.send_message(self.user_id, self.friend_id, msg, server_now)
            self.message_input.clear()
            self._load_messages()

    def closeEvent(self, e):
        self.timer.stop()
        super().closeEvent(e)

# ══════════════════════════════════════════════════════════════════════════════
#  GuildChatDialog (Direct Messaging Guild System with real-time updates)
# ══════════════════════════════════════════════════════════════════════════════
class GuildChatDialog(QDialog):
    def __init__(self, guild_id, user_id, parent=None):
        super().__init__(parent)
        
        self.guild_id = guild_id
        self.user_id = user_id
        conn = db.get_conn()
        leader_row = conn.execute("SELECT leader_id FROM guilds WHERE id=?", (guild_id,)).fetchone()
        self.is_leader = leader_row and leader_row["leader_id"] == user_id
        conn.close()
        self.setWindowTitle(tr("guild_chat_title"))
        self.setMinimumSize(450, 550)
        self.setStyleSheet(build_ss())
        self._build()
        self._load_messages()
        self.timer = QTimer()
        self.timer.timeout.connect(self._load_messages)
        self.timer.start(3000)

    def _build(self):
        layout = QVBoxLayout(self)
        
        # Tombol Clear All jika leader
        if self.is_leader:
            btn_layout = QHBoxLayout()
            clear_all_btn = _btn(tr("chat_clear_all"), "danger", self._clear_all, 32)
            btn_layout.addStretch()
            btn_layout.addWidget(clear_all_btn)
            layout.addLayout(btn_layout)
        
        self.chat_area = QTextEdit()
        self.chat_area.setReadOnly(True)
        layout.addWidget(self.chat_area)
        
        input_layout = QHBoxLayout()
        self.message_input = QLineEdit()
        send_btn = _btn(tr("chat_send_btn"), "solid", self._send_message)
        input_layout.addWidget(self.message_input)
        input_layout.addWidget(send_btn)
        layout.addLayout(input_layout)

    def _clear_all(self):
        reply = QMessageBox.question(self, tr("confirm_title"), tr("chat_clear_all_confirm"), 
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.clear_guild_chat(self.guild_id)
            self._load_messages()
            SND.notify()

    def _load_messages(self):
        msgs = db.get_guild_messages(self.guild_id)
        self.chat_area.clear()
        for m in msgs:
            self.chat_area.append(tr("guild_chat_message_format", time=m['created_at'][11:16], name=m['display_name'], message=m['message']))

    def _send_message(self):
        msg = self.message_input.text().strip()
        if msg:
            server_now = TimeSync.get_current_time().isoformat()
            db.send_guild_message(self.guild_id, self.user_id, msg, server_now)
            self.message_input.clear()
            self._load_messages()

    def closeEvent(self, e):
        self.timer.stop()
        super().closeEvent(e)

# ══════════════════════════════════════════════════════════════════════════════
#  FriendProfileDialog
# ══════════════════════════════════════════════════════════════════════════════
class FriendProfileDialog(QDialog):
    def __init__(self, friend_id, parent=None):
        super().__init__(parent)
        
        self.friend_id = friend_id
        self.setWindowTitle(tr("friend_profile_title"))
        self.setMinimumSize(400, 500)
        self.setStyleSheet(build_ss())
        self._build()
        self._load()

    def _build(self):
        layout = QVBoxLayout(self)
        self.content = QWidget()
        self.content_layout = QVBoxLayout(self.content)
        scroll = _scrolled(self.content)
        layout.addWidget(scroll)

    def _load(self):
        while self.content_layout.count():
            item = self.content_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        d = db.get_friend_profile_details(self.friend_id)
        if not d:
            return
        u = d["user"]
        stats = d["stats"]

        # ── Header: avatar + identitas ──
        avatar = _emoji_label(u.get("avatar_emoji", "⚔️"), ICON_HERO)
        avatar.setAlignment(Qt.AlignmentFlag.AlignCenter)
        avatar.setStyleSheet(f"background: {u.get('avatar_color', '#5a8a2e')}; border-radius: 12px; padding: 10px;")
        self.content_layout.addWidget(avatar)
        name_lbl = _lbl(u.get("display_name", "No name"), size=16, bold=True)
        name_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.content_layout.addWidget(name_lbl)
        user_lbl = _lbl(f"@{u.get('username', '')}", "sub")
        user_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.content_layout.addWidget(user_lbl)

        # Chips: level + kelas + title
        chips = QHBoxLayout()
        chips.setSpacing(8)
        chips.addStretch()
        lvl_chip = QLabel(f"⚔️ Lv {d['level']}")
        lvl_chip.setStyleSheet(f"color:#f0c040; font-weight:bold; font-size:12px; "
                               f"background:{_T('panel')}; border:1px solid {_T('border')}; "
                               f"border-radius:10px; padding:3px 10px;")
        chips.addWidget(lvl_chip)
        if d["avatar_class"]:
            cls_txt = tr(f"class_{d['avatar_class']}")
            cls_chip = QLabel(f"🎭 {cls_txt}")
            cls_chip.setStyleSheet(f"color:{_T('text')}; font-size:12px; background:{_T('panel')}; "
                                   f"border:1px solid {_T('border')}; border-radius:10px; padding:3px 10px;")
            chips.addWidget(cls_chip)
        if d["selected_title"]:
            t_chip = QLabel(f"🏅 {d['selected_title']}")
            t_chip.setStyleSheet(f"color:#4dd9e0; font-size:12px; font-weight:bold; background:{_T('panel')}; "
                                 f"border:1px solid {_T('border')}; border-radius:10px; padding:3px 10px;")
            chips.addWidget(t_chip)
        chips.addStretch()
        self.content_layout.addLayout(chips)

        # Info ringkas: bio, guild, rebirth, sport level, join date
        info_card = _card()
        iv = QVBoxLayout(info_card)
        if u.get("bio"):
            bio = QLabel(f"💬 {u['bio']}")
            bio.setWordWrap(True)
            bio.setStyleSheet(f"color:{_T('muted')}; font-style:italic;")
            iv.addWidget(bio)
        if d["guild_name"]:
            iv.addWidget(QLabel(tr("friend_guild", name=d["guild_name"])))
        mini = []
        mini.append(tr("friend_sport_level", lvl=d["sport_level"]))
        mini.append(tr("friend_rebirth", count=d["rebirth_count"]))
        if d["join_date"]:
            mini.append(tr("friend_joined", date=d["join_date"]))
        ml = QLabel("   ·   ".join(mini))
        ml.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        ml.setWordWrap(True)
        iv.addWidget(ml)
        self.content_layout.addWidget(info_card)

        # ── Progres XP ──
        self.content_layout.addWidget(_lbl(tr("friend_xp_progress"), "sub", 11))
        xpbar = QProgressBar()
        xpbar.setRange(0, max(1, d["xp_needed"]))
        xpbar.setValue(min(d["xp"], d["xp_needed"]))
        xpbar.setFormat(f"{d['xp']}/{d['xp_needed']} XP  (" + tr("friend_total_xp") +
                        f": {int(d['total_xp_earned'])} )")
        xpbar.setMinimumHeight(18)
        self.content_layout.addWidget(xpbar)

        # ── Progres achievement ──
        self.content_layout.addWidget(_lbl(
            tr("friend_achievements_progress", done=d["achievements_done"],
               total=d["achievements_total"]), "sub", 11))
        abar = QProgressBar()
        abar.setRange(0, max(1, d["achievements_total"]))
        abar.setValue(d["achievements_done"])
        abar.setFormat(f"{d['achievements_done']}/{d['achievements_total']}")
        abar.setMinimumHeight(18)
        self.content_layout.addWidget(abar)

        # 6 achievement terbaru
        if d["latest_achievements"]:
            self.content_layout.addWidget(_lbl(tr("friend_latest_achievements"), "sub", 11))
            grid_w = QWidget()
            grid = QGridLayout(grid_w)
            grid.setContentsMargins(0, 0, 0, 0)
            grid.setSpacing(6)
            lang = AppState.get_language()
            for i, ach in enumerate(d["latest_achievements"]):
                chip = QFrame()
                chip.setStyleSheet(f"background:{_T('panel')}; border:1px solid {_T('border')}; border-radius:8px;")
                ch = QHBoxLayout(chip)
                ch.setContentsMargins(8, 4, 8, 4)
                ch.addWidget(_emoji_label(ach["icon"], ICON_INLINE))
                ach_name, _ = db.tr_achievement(ach, lang)
                nm = QLabel(ach_name)
                nm.setStyleSheet("font-size:11px;")
                nm.setWordWrap(True)
                ch.addWidget(nm, 1)
                grid.addWidget(chip, i // 2, i % 2)
            self.content_layout.addWidget(grid_w)

        # ── Statistik detail ──
        self.content_layout.addWidget(_lbl(tr("friend_stats_title"), "sub", 11))
        data = [
            (tr("stats_habit_today"), f"{stats['habits_done_today']}/{stats['habits_total']}"),
            (tr("stats_daily_today"), f"{stats['dailies_done_today']}/{stats['dailies_total']}"),
            (tr("stats_quest_done"), f"{stats['todos_done']}/{stats['todos_total']}"),
            (tr("friend_tasks_done"), str(d["tasks_done"])),
            (tr("stats_max_streak"), str(stats["max_streak"])),
            (tr("stats_boss_killed"), str(stats["bosses_killed"])),
            (tr("stats_pets"), str(stats["pet_count"])),
            (tr("friend_pomodoro"), f"{d['pomodoro_minutes']} min"),
        ]
        grid2_w = QWidget()
        grid2 = QGridLayout(grid2_w)
        grid2.setContentsMargins(0, 0, 0, 0)
        grid2.setSpacing(6)
        for i, (label, value) in enumerate(data):
            cell = QFrame()
            cell.setStyleSheet(f"background:{_T('panel')}; border:1px solid {_T('border')}; border-radius:8px;")
            cv = QVBoxLayout(cell)
            cv.setContentsMargins(10, 6, 10, 6)
            cv.setSpacing(2)
            vl = QLabel(value)
            vl.setStyleSheet(f"color:{_T('text')}; font-size:14px; font-weight:bold;")
            ll = QLabel(label)
            ll.setStyleSheet(f"color:{_T('muted')}; font-size:10px;")
            ll.setWordWrap(True)
            cv.addWidget(vl)
            cv.addWidget(ll)
            grid2.addWidget(cell, i // 2, i % 2)
        self.content_layout.addWidget(grid2_w)
        self.content_layout.addStretch()


# ══════════════════════════════════════════════════════════════════════════════
#  Reset Password via Pertanyaan Keamanan 
# ══════════════════════════════════════════════════════════════════════════════
class ResetPasswordBySecurityDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        
        self.setWindowTitle(tr("reset_password_security_title"))
        self.setMinimumSize(450, 400)
        self.setStyleSheet(build_overworld_ss())
        self.user_id = None
        self.security_question = None
        self._build()
    
    def _build(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        
        self.step = 1  # 1: username, 2: jawaban, 3: reset
        self._step1()
    
    def _step1(self):
        self._clear()
        layout = self.layout()
        layout.addWidget(_lbl(tr("reset_password_username"), size=12))
        self.username_input = _input(tr("login_username"))
        layout.addWidget(self.username_input)
        btn = _btn(tr("reset_password_check"), "solid", self._check_username)
        layout.addWidget(btn)
        self.msg_label = _lbl("", "sub", 12)
        layout.addWidget(self.msg_label)
    
    def _check_username(self):
        username = self.username_input.text().strip()
        if not username:
            self.msg_label.setText(tr("reset_username_empty"))
            return
        conn = db.get_conn()
        row = conn.execute(
            "SELECT id, security_question FROM users WHERE username=?",
            (username.lower(),)
        ).fetchone()
        conn.close()
        if not row:
            self.msg_label.setText(tr("reset_username_notfound"))
            return
        if not row["security_question"]:
            self.msg_label.setText(tr("reset_no_sq"))
            return
        self.user_id = row["id"]
        self.security_question = row["security_question"]
        self._step2() 

    def _step2(self):
        self._clear()
        layout = self.layout()
        layout.addWidget(_lbl(tr("reset_password_security_question"), size=12))
        layout.addWidget(_lbl(self.security_question, size=12, bold=True))
        self.answer_input = _input(tr("reset_password_answer"))
        layout.addWidget(self.answer_input)
        btn = _btn(tr("reset_password_verify"), "solid", self._verify)

        layout.addWidget(btn)
        self.msg_label = _lbl("", "sub", 12)
        layout.addWidget(self.msg_label)
    
    def _verify(self):
        answer = self.answer_input.text().strip()
        if not answer:
            self.msg_label.setText(tr("reset_answer_empty"))
            return
        if db.verify_security_answer(self.user_id, answer):
            self._step3()
        else:
            self.msg_label.setText(tr("reset_answer_wrong"))
    
    def _step3(self):
        self._clear()
        if not self.user_id:  
            self.msg_label.setText(tr("reset_error"))
            return
        layout = self.layout()
        layout.addWidget(_lbl(tr("reset_password_new"), size=12))   
        self.new_pass = _input(tr("reset_password_new"), password=True)
        self.confirm_pass = _input(tr("reset_password_confirm"), password=True)
        layout.addWidget(self.new_pass)
        layout.addWidget(self.confirm_pass)
        btn = _btn(tr("reset_password_btn"), "solid", self._reset)
        layout.addWidget(btn)
        self.msg_label = _lbl("", "sub", 12)
        layout.addWidget(self.msg_label)
    
    def _reset(self):
        p1 = self.new_pass.text()
        p2 = self.confirm_pass.text()
        if len(p1) < db.PASSWORD_MIN_LEN:
            self.msg_label.setText(tr("login_pw_min"))
            return
        if p1 != p2:
            self.msg_label.setText(tr("login_pw_mismatch"))
            return
        r = db.reset_password_by_security(self.user_id, p1)
        if r["ok"]:
            _show(self, tr("berhasil_title"), tr("reset_success_msg"), "success")
            self.accept()
        else:
            self.msg_label.setText(r["msg"])
    
    def _clear(self):
        # Hapus semua widget dari layout
        layout = self.layout()
        while layout.count():
            item = layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

# ══════════════════════════════════════════════════════════════════════════════
#  LOGIN WINDOW  (complex & user-friendly — no fade_in on self)
# ══════════════════════════════════════════════════════════════════════════════
class LoginWindow(QDialog):
    logged_in = pyqtSignal(dict)

    def __init__(self):
        super().__init__()
        self.setWindowTitle(tr("login_title"))
        self.setMinimumSize(500, 620)
        self.setWindowIcon(QIcon(get_icon_path('craftlife.ico')))
        self.setStyleSheet(build_ss())
        self._build()
        # Sembunyikan dulu, cek session
        self.hide()
        from PyQt6.QtCore import QTimer
        QTimer.singleShot(10, self._check_auto_login)

    def _check_auto_login(self):
        session = load_session()
        if session:
            user = db.get_user(session["user_id"])
            if user:
                self.logged_in.emit(user)
                self.accept()  
                return
        self.show()

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Banner (plain widget, safe to fade_in)
        banner = QWidget()
        banner.setMinimumHeight(110)
        banner.setStyleSheet(
            f"background: qlineargradient(x1:0,y1:0,x2:1,y2:1,"
            f"stop:0 {_T('bg')},stop:1 {_T('panel')});"
            f"border-bottom: 2px solid {_T('primary')};")
        bl = QVBoxLayout(banner)
        bl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        t1 = _lbl(tr("app_logo"), size=24, bold=True)
        t1.setAlignment(Qt.AlignmentFlag.AlignCenter)
        t1.setStyleSheet(f"color: {_T('light')};")
        t2 = _lbl(tr("app_tagline"), size=12)
        t2.setAlignment(Qt.AlignmentFlag.AlignCenter)
        bl.addWidget(t1)
        bl.addWidget(t2)
        root.addWidget(banner)
        fade_in(banner, 300)   # fade on banner widget — safe

        body = QWidget()
        body_lay = QVBoxLayout(body)
        body_lay.setContentsMargins(40, 20, 40, 20)
        body_lay.setSpacing(12)

        tabs = QTabWidget()
        tabs.addTab(self._login_tab(),    tr("login_tab_login_label"))
        tabs.addTab(self._register_tab(), tr("login_tab_register_label"))
        body_lay.addWidget(tabs)
        root.addWidget(body)

    def _login_tab(self) -> QWidget:
        w   = QWidget()
        lay = QVBoxLayout(w)
        lay.setContentsMargins(0, 16, 0, 0)
        lay.setSpacing(12)

        self._l_user = _input(tr("login_username"))
        self._l_pass = _input(tr("login_password"), True)
        self._l_msg  = _lbl("", "sub", 12)
        self._l_msg.setStyleSheet("color: #e05050;")
        self._l_msg.setWordWrap(True)

        self._remember_cb = QCheckBox(tr("login_remember_text"))
        self._remember_cb.setStyleSheet(f"color: {_T('text')};")
        lay.addWidget(self._remember_cb)

        ok = _btn(tr("login_button_text"), "solid", self._do_login, 48)
        for w_ in [self._l_user, self._l_pass, ok, self._l_msg]:
            lay.addWidget(w_)
        lay.addStretch()

        forgot_btn = _btn(tr("login_forgot_btn"), "flat", self._forgot_password)
        lay.addWidget(forgot_btn)
        self._l_pass.returnPressed.connect(self._do_login)
        self._l_user.returnPressed.connect(self._l_pass.setFocus)
        return w

    def _forgot_password(self):
        dlg = ChooseResetMethodDialog(self)
        dlg.exec()

    def _register_tab(self) -> QWidget:
        w   = QWidget()
        lay = QVBoxLayout(w)
        lay.setContentsMargins(0, 16, 0, 0)
        lay.setSpacing(10)

        self._r_user  = _input(tr("register_username"))
        self._r_disp  = _input(tr("register_display"))
        self._r_pass = _input(tr("register_password"), True)
        self._r_pass2 = _input(tr("register_confirm"), True)
        self._r_bio   = _input(tr("register_bio"))

        lay.addWidget(_lbl(tr("register_class_label"), size=12))
        self._r_class = QComboBox()
        self._r_class.setMinimumHeight(42)
        for cid, cdata in db.AVATAR_CLASSES.items():
            from translations import TRANSLATIONS
            # Ambil terjemahan nama class dan bonus
            name_key = f"class_{cid}_name"
            bonus_key = f"class_{cid}_bonus"
            name_text = tr(name_key) if name_key in TRANSLATIONS else cdata['name']
            bonus_text = tr(bonus_key) if bonus_key in TRANSLATIONS else cdata['bonus']
            self._r_class.addItem(
                f"{cdata['icon']}  {name_text}  —  {bonus_text}",
                cid)

        self._r_msg = _lbl("", "sub", 12)
        self._r_msg.setWordWrap(True)

        # Indikator kekuatan password (live mengikuti ketikan)
        self._r_meter = PasswordStrengthMeter(self._r_pass)

        ok = _btn(tr("register_button_text"), "solid", self._do_register, 48)
        for w_ in [self._r_user, self._r_disp, self._r_pass,
                   self._r_pass2, self._r_meter, self._r_bio, self._r_class,
                   ok, self._r_msg]:
            lay.addWidget(w_)
        lay.addStretch()
        return w

    def _do_login(self):
        loading = LoadingDialog(tr("login_processing"), self)
        loading.show()
        QApplication.processEvents()
        
        try:
            r = db.login_user(self._l_user.text(), self._l_pass.text())
            if r["ok"]:
                SND.complete()
                if self._remember_cb.isChecked():
                    user = r["user"]
                    save_session(user["id"], user["username"])
                loading.accept()
                self.logged_in.emit(r["user"])
                self.accept()
            else:
                loading.accept()
                SND.error()
                self._l_msg.setText(r["msg"])
        except Exception as e:
            loading.accept()
            self._l_msg.setText(tr("error_occurred", error=e))

    def _do_register(self):
        u = self._r_user.text().strip()
        if not u or " " in u:
            self._r_msg.setStyleSheet("color:#e05050;")
            self._r_msg.setText(tr("register_username_invalid"))
            return
        if len(self._r_pass.text()) < db.PASSWORD_MIN_LEN:
            self._r_msg.setStyleSheet("color:#e05050;")
            self._r_msg.setText(tr("login_pw_min"))
            return
        if self._r_pass.text() != self._r_pass2.text():
            self._r_msg.setStyleSheet("color:#e05050;")
            self._r_msg.setText(tr("login_pw_mismatch"))
            return
        r = db.register_user(
            u, self._r_pass.text(),
            self._r_disp.text().strip(),
            self._r_bio.text(),
            self._r_class.currentData(),
        )
        if r["ok"]:
            SND.level_up()
            self._r_msg.setStyleSheet(f"color:{_T('light')};")
            self._r_msg.setText("✅ " + r["msg"] + tr("register_success_login_now"))
        else:
            SND.error()
            self._r_msg.setStyleSheet("color:#e05050;")
            self._r_msg.setText(r["msg"])


# ══════════════════════════════════════════════════════════════════════════════
#  PASSWORD STRENGTH METER ─ indikator kekuatan password (live)
# ══════════════════════════════════════════════════════════════════════════════
def _password_score(pw: str) -> int:
    """Skor kekuatan password 0..4 (0 = sangat lemah)."""
    if not pw:
        return 0
    score = 0
    if len(pw) >= 8:
        score += 1
    if len(pw) >= 12:
        score += 1
    kinds = sum(bool(re.search(p, pw))
                for p in (r"[a-z]", r"[A-Z]", r"\d", r"[^A-Za-z0-9]"))
    if kinds >= 2:
        score += 1
    if kinds >= 3:
        score += 1
    if kinds == 4 and len(pw) >= 10:
        score += 1
    return min(score, 4)


class PasswordStrengthMeter(QWidget):
    """Bar 4 segmen + label — update otomatis mengikuti input password."""

    _LEVELS = [
        ("pw_strength_very_weak", "#e05050"),
        ("pw_strength_weak",      "#e08050"),
        ("pw_strength_fair",      "#e0b050"),
        ("pw_strength_strong",    "#7ac74c"),
        ("pw_strength_very_strong", "#4caf7d"),
    ]

    def __init__(self, password_input, parent=None):
        super().__init__(parent)
        lay = QHBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0)
        lay.setSpacing(8)

        bar_row = QHBoxLayout()
        bar_row.setSpacing(3)
        self._bars = []
        for _ in range(4):
            seg = QFrame()
            seg.setFixedHeight(5)
            seg.setMinimumWidth(26)
            seg.setStyleSheet(f"background:{_T('border')}; border-radius:2px;")
            self._bars.append(seg)
            bar_row.addWidget(seg)
        bar_row.addStretch()

        self._lbl = QLabel()
        self._lbl.setStyleSheet(f"color:{_T('sub')}; font-size:11px;")

        lay.addLayout(bar_row, 1)
        lay.addWidget(self._lbl)

        password_input.textChanged.connect(self.update_strength)
        self.update_strength(password_input.text())

    def update_strength(self, pw: str):
        score = _password_score(pw)
        key, color = self._LEVELS[score]
        for i, seg in enumerate(self._bars):
            c = color if i < score else _T("border")
            seg.setStyleSheet(f"background:{c}; border-radius:2px;")
        self._lbl.setText(tr(key) if pw else tr("pw_strength_hint"))
        self._lbl.setStyleSheet(
            f"color:{color if pw else _T('sub')}; font-size:11px;")


# ══════════════════════════════════════════════════════════════════════════════
#  Forgot Password
# ══════════════════════════════════════════════════════════════════════════════
class ChooseResetMethodDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        
        self.setWindowTitle(tr("login_forgot"))
        self.setMinimumSize(350, 200)
        self.setStyleSheet(build_overworld_ss())
        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        layout.addWidget(_lbl(tr("reset_method_title"), size=12))
        btn_security = _btn(tr("forgot_security_btn"), "solid", self._choose_security)
        btn_backup = _btn(tr("forgot_backup_btn"), "solid", self._choose_backup)
        layout.addWidget(btn_security)
        layout.addWidget(btn_backup)
        cancel = _btn(tr("forgot_cancel_btn"), "flat", self.reject)
        layout.addWidget(cancel)
    
    def _choose_security(self):
        self.accept()
        dlg = ResetPasswordBySecurityDialog(self.parent())
        dlg.exec()
    
    def _choose_backup(self):
        self.accept()
        dlg = ResetPasswordByBackupCodeDialog(self.parent())
        dlg.exec()


# ══════════════════════════════════════════════════════════════════════════════
#  Backup Code Reset Password
# ══════════════════════════════════════════════════════════════════════════════
class ResetPasswordByBackupCodeDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        
        self.setWindowTitle(tr("reset_password_backup_title"))
        self.setMinimumSize(450, 350)
        self.setStyleSheet(build_overworld_ss())
        self.user_id = None
        self._build()
    
    def _build(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        self._step1()
    
    def _step1(self):
        self._clear()
        layout = self.layout()
        layout.addWidget(_lbl(tr("reset_password_username"), size=12))
        self.username_input = _input(tr("login_username"))
        layout.addWidget(self.username_input)
        btn = _btn(tr("check_backup_code_btn"), "solid", self._check_username)
        layout.addWidget(btn)
        self.msg_label = _lbl("", "sub", 12)
        layout.addWidget(self.msg_label)
    
    def _check_username(self):
        username = self.username_input.text().strip()
        if not username:
            self.msg_label.setText(tr("reset_username_empty"))
            return
        conn = db.get_conn()
        row = conn.execute("SELECT id FROM users WHERE username=?", (username.lower(),)).fetchone()
        conn.close()
        if not row:
            self.msg_label.setText(tr("reset_username_notfound"))
            return
        self.user_id = row["id"]
        # Cek apakah user memiliki backup codes yang belum dipakai
        codes = db.get_user_backup_codes(self.user_id, only_unused=True)
        if not codes:
            self.msg_label.setText(tr("reset_no_bc_long"))
            return
        self._step2()
    
    def _step2(self):
        self._clear()
        layout = self.layout()
        layout.addWidget(_lbl(tr("reset_backup_code_label"), size=12))
        self.code_input = _input(tr("reset_password_backup_code"))
        layout.addWidget(self.code_input)
        btn = _btn(tr("reset_password_backup_btn"), "solid", self._verify)
        layout.addWidget(btn)
        self.msg_label = _lbl("", "sub", 12)
        layout.addWidget(self.msg_label)
    
    def _verify(self):
        code = self.code_input.text().strip().upper()
        if not code:
            self.msg_label.setText(tr("reset_bc_empty"))
            return
        if db.verify_backup_code(self.user_id, code):
            self._step3()
        else:
            self.msg_label.setText(tr("reset_bc_invalid"))
    
    def _step3(self):
        self._clear()
        layout = self.layout()
        layout.addWidget(_lbl(tr("reset_password_new"), size=12))
        self.new_pass = _input(tr("reset_password_new"), password=True)
        self.confirm_pass = _input(tr("reset_password_confirm"), password=True)
        layout.addWidget(self.new_pass)
        layout.addWidget(self.confirm_pass)
        btn = _btn(tr("reset_backup_reset_btn"), "solid", self._reset)
        layout.addWidget(btn)
        self.msg_label = _lbl("", "sub", 12)
        layout.addWidget(self.msg_label)
    
    def _reset(self):
        p1 = self.new_pass.text()
        p2 = self.confirm_pass.text()
        if len(p1) < db.PASSWORD_MIN_LEN:
            self.msg_label.setText(tr("login_pw_min"))
            return
        if p1 != p2:
            self.msg_label.setText(tr("login_pw_mismatch"))
            return
        r = db.reset_password_with_backup_code(self.user_id, p1)
        if r["ok"]:
            _show(self, tr("berhasil_title"), tr("reset_success_msg"), "success")
            self.accept()
        else:
            self.msg_label.setText(r["msg"])
    
    def _clear(self):
        layout = self.layout()
        while layout.count():
            item = layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

# ==================================================================
# FOOD PAGE (Calorie Tracker)
# ==================================================================
class AddFoodDialog(QDialog):
    """Dialog untuk mencatat makanan atau menambah makanan custom."""
    def __init__(self, user_id, mode="log", parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.mode = mode
        self.all_foods = []  # akan diisi list of dict
        self.setWindowTitle(tr("food_log") if mode == "log" else tr("food_add_custom"))
        self.setMinimumWidth(500)
        self.setMinimumHeight(550)
        self.setStyleSheet(build_ss())
        self._build()

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(14)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))
        lay.addWidget(_sep())

        if self.mode == "log":
            # ========== FILTER KATEGORI ==========
            lay.addWidget(_lbl(tr("food_filter_category"), size=12))
            self.category_filter = QComboBox()
            self.category_filter.setMinimumHeight(42)
            self.category_filter.addItem(tr("food_all_categories"), "all")
            self.category_filter.addItem(tr("food_category_main"), "main")
            self.category_filter.addItem(tr("food_category_protein"), "protein")
            self.category_filter.addItem(tr("food_category_veg"), "veg")
            self.category_filter.addItem(tr("food_category_drink"), "drink")
            self.category_filter.addItem(tr("food_category_snack"), "snack")
            self.category_filter.addItem(tr("food_category_international"), "international")
            self.category_filter.currentIndexChanged.connect(self._filter_food_combo)
            lay.addWidget(self.category_filter)

            # ========== SEARCH BAR ==========
            lay.addWidget(_lbl(tr("food_search_label"), size=12))
            self.search_input = QLineEdit()
            self.search_input.setPlaceholderText(tr("food_search_ph"))
            self.search_input.setMinimumHeight(42)
            self.search_input.textChanged.connect(self._filter_food_combo)
            lay.addWidget(self.search_input)

            # ========== COMBO MAKANAN ==========
            lay.addWidget(_lbl(tr("food_ingredient_select_label_short"), size=12))
            self.food_combo = QComboBox()
            self.food_combo.setMinimumHeight(42)
            lay.addWidget(self.food_combo)

            # Ambil data makanan
            foods = db.get_food_items(self.user_id)
            lang = AppState.get_language()   # ambil bahasa aktif
            self.all_foods = []
            for f in foods:
                display_name = get_food_name(f['name'], lang)  # terjemahkan
                display = f"{f['icon']} {display_name} - {f['calories']:.0f} kcal"
                self.all_foods.append({
                    "id": f["id"],
                    "name": f['name'],
                    "icon": f['icon'],
                    "display": display
                })

            # ========== PORS ==========
            lay.addWidget(_lbl(tr("food_serving_label"), size=12))
            self.serving = QDoubleSpinBox()
            self.serving.setMinimum(0.25)
            self.serving.setMaximum(10)
            self.serving.setSingleStep(0.25)
            self.serving.setValue(1)
            self.serving.setMinimumHeight(42)
            lay.addWidget(self.serving)

            # ========== JENIS MAKAN ==========
            lay.addWidget(_lbl(tr("food_meal_type"), size=12))
            self.meal_type = _combo([
                (tr("food_meal_breakfast"), "breakfast"),
                (tr("food_meal_lunch"), "lunch"),
                (tr("food_meal_dinner"), "dinner"),
                (tr("food_meal_snack"), "snack")
            ])
            lay.addWidget(self.meal_type)

            # ========== TANGGAL ==========
            lay.addWidget(_lbl(tr("food_date"), size=12))
            self.log_date = QLineEdit()
            self.log_date.setText(date.today().isoformat())
            self.log_date.setMinimumHeight(42)
            lay.addWidget(self.log_date)

            # ========== CATATAN ==========
            lay.addWidget(_lbl(tr("dialog_notes"), size=12))
            self.notes = _input()
            lay.addWidget(self.notes)

            btn = _btn(tr("food_log_btn"), "solid", self._log_food, 46)
            lay.addWidget(btn)

            # Isi combo awal
            self._filter_food_combo("")

        else:  # mode "add" untuk makanan custom
            lay.addWidget(_lbl(tr("food_name_label"), size=12))
            self.name = _input(tr("food_custom_ph"))
            lay.addWidget(self.name)

            lay.addWidget(_lbl(tr("food_icon_label"), size=12))
            self.icon = _input("🍎")
            lay.addWidget(self.icon)

            lay.addWidget(_lbl(tr("food_calories_label"), size=12))
            self.calories = QDoubleSpinBox()
            self.calories.setMaximum(2000)
            self.calories.setMinimum(0)
            self.calories.setValue(100)
            lay.addWidget(self.calories)

            lay.addWidget(_lbl(tr("food_protein_label"), size=12))
            self.protein = QDoubleSpinBox()
            self.protein.setMaximum(200)
            self.protein.setValue(10)
            lay.addWidget(self.protein)

            lay.addWidget(_lbl(tr("food_carbs_label"), size=12))
            self.carbs = QDoubleSpinBox()
            self.carbs.setMaximum(200)
            self.carbs.setValue(20)
            lay.addWidget(self.carbs)

            lay.addWidget(_lbl(tr("food_fat_label"), size=12))
            self.fat = QDoubleSpinBox()
            self.fat.setMaximum(100)
            self.fat.setValue(5)
            lay.addWidget(self.fat)

            btn = _btn(tr("food_save_btn"), "solid", self._add_food, 46)
            lay.addWidget(btn)

        root.addWidget(_scrolled(content))

    # ========== FUNGSI FILTER ==========
    def _filter_food_combo(self, text=""):
        """Filter daftar makanan berdasarkan teks pencarian dan kategori."""
        if not self.all_foods:
            return
        self.food_combo.blockSignals(True)
        self.food_combo.clear()

        search_text = self.search_input.text().strip().lower() if hasattr(self, 'search_input') else ""
        category = self.category_filter.currentText() if hasattr(self, 'category_filter') else tr("food_all_categories")

        for item in self.all_foods:
            # Filter berdasarkan teks
            if search_text and search_text not in item["display"].lower() and search_text not in item["name"].lower():
                continue
            # Filter berdasarkan kategori
            if not self._is_in_category(item["name"], item["icon"], category):
                continue
            self.food_combo.addItem(item["display"], item["id"])

        self.food_combo.blockSignals(False)

    def _is_in_category(self, name: str, icon: str, category: str) -> bool:
        """Cek apakah makanan termasuk dalam kategori berdasarkan nama dan icon."""
        name_lower = name.lower()
        icon_lower = icon.lower()

        if category == tr("food_all_categories"):
            return True

        # Makanan Utama
        if category == tr("food_category_main"):
            keywords = ["nasi", "mie", "mi ", "roti", "burger", "pizza", "sandwich",
                        "spaghetti", "pasta", "bubur", "ketupat", "lontong", "noodle",
                        "donut", "croissant", "bagel", "taco", "burrito", "quesadilla"]
            return any(k in name_lower for k in keywords)

        # Lauk
        if category == tr("food_category_protein"):
            keywords = ["ayam", "daging", "ikan", "telur", "tahu", "tempe", "sapi",
                        "kambing", "bebek", "udang", "cumi", "kepiting", "sate",
                        "rendang", "gulai", "tuna", "salmon", "tofu", "egg", "chicken",
                        "beef", "pork", "lamb", "shrimp", "crab", "meat"]
            return any(k in name_lower for k in keywords)

        # Sayur & Buah
        if category == tr("food_category_veg"):
            keywords = ["sayur", "sayuran", "buah", "salad", "tumis", "capcay", "gado",
                        "karedok", "pecel", "lalapan", "apel", "pisang", "jeruk",
                        "mangga", "semangka", "alpukat", "anggur", "stroberi", "nanas",
                        "melon", "pepaya", "jambu", "wortel", "kangkung", "bayam",
                        "brokoli", "kubis", "timun", "tomat", "lettuce", "vegetable",
                        "fruit", "apple", "banana", "orange", "mango", "watermelon"]
            return any(k in name_lower for k in keywords)

        # Minuman
        if category == tr("food_category_drink"):
            keywords = ["air", "kopi", "teh", "jus", "susu", "soda", "milk", "coffee",
                        "tea", "juice", "smoothie", "coklat panas", "matcha", "latte",
                        "cappuccino", "es ", "wedang", "bandrek", "bajigur", "soda",
                        "lemonade", "milkshake", "yogurt drink", "kefir", "kombucha",
                        "boba", "bubble tea", "thai tea", "chai", "horlicks", "ovaltine"]
            return any(k in name_lower for k in keywords)

        # Snack & Dessert
        if category == tr("food_category_snack"):
            keywords = ["keripik", "kue", "es krim", "coklat", "chocolate", "cake",
                        "cookie", "biscuit", "donat", "pastry", "puding", "pudding",
                        "cheesecake", "brownie", "muffin", "cupcake", "waffle", "pancake",
                        "ice cream", "gelato", "sorbet", "pie", "tart", "croissant",
                        "kripik", "popcorn", "permen", "candy", "wafer", "biskuit",
                        "martabak", "terang bulan", "klepon", "lupis", "getuk", "nagasari"]
            return any(k in name_lower for k in keywords)

        # Internasional
        if category == tr("food_category_international"):
            keywords = ["katsu", "sushi", "ramen", "curry", "tempura", "takoyaki",
                        "okonomiyaki", "teriyaki", "gyoza", "dim sum", "pho", "tom yum",
                        "pad thai", "bibimbap", "bulgogi", "kimchi", "naan", "biryani",
                        "falafel", "hummus", "paella", "croissant", "baguette", "lasagna",
                        "carbonara", "alfredo", "bolognese", "doner", "shawarma", "kebab"]
            return any(k in name_lower for k in keywords)

        return True

    # ========== FUNGSI LOG FOOD ==========
    def _log_food(self):
        food_id = self.food_combo.currentData()
        if food_id is None:
            _show(self, tr("msg_error"), tr("food_select_first"), "error")
            return
        serving = self.serving.value()
        meal_type = self.meal_type.currentData()
        log_date = self.log_date.text().strip()
        notes = self.notes.text()
        if not log_date:
            log_date = date.today().isoformat()
        r = db.log_food(self.user_id, food_id, serving, meal_type, log_date, notes)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), tr("calories_logged", cal=r['calories']), "success")
            self.accept()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", "Gagal mencatat makanan"), "error")

    def _add_food(self):
        name = self.name.text().strip()
        icon = self.icon.text().strip() or "🍎"
        cal = self.calories.value()
        pro = self.protein.value()
        carb = self.carbs.value()
        fat = self.fat.value()
        if not name:
            _show(self, tr("msg_error"), tr("food_name_empty"), "error")
            return
        r = db.add_custom_food(self.user_id, name, icon, cal, pro, carb, fat)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), tr("food_added_success", name=name), "success")
            self.accept()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", "Gagal menambah makanan"), "error")


class AddRecipeDialog(QDialog):
    def __init__(self, user_id, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.selected_items = []  # list of (food_id, quantity)
        self.setWindowTitle(tr("recipe_add_title"))
        self.setMinimumWidth(500)
        self.setMinimumHeight(600)
        self.setStyleSheet(build_ss())
        self._build()

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(12)

        lay.addWidget(_lbl(tr("recipe_add_title"), "section", 14, True))
        lay.addWidget(_sep())

        # Nama resep
        lay.addWidget(_lbl(tr("food_recipe_name_label"), size=12))
        self.name = _input(tr("food_recipe_name_ph"))
        lay.addWidget(self.name)

        # Icon
        lay.addWidget(_lbl(tr("dialog_icon"), size=12))
        self.icon = _input("🍲")
        lay.addWidget(self.icon)

        # Porsi standar
        lay.addWidget(_lbl(tr("food_recipe_serving_label"), size=12))
        self.serving = QDoubleSpinBox()
        self.serving.setRange(0.5, 10)
        self.serving.setValue(1)
        self.serving.setSuffix(tr("unit_serving"))
        lay.addWidget(self.serving)

        # Daftar bahan
        lay.addWidget(_lbl(tr("food_recipe_ingredients_label"), size=12))
        self.ingredients_list = QListWidget()
        self.ingredients_list.setMaximumHeight(120)
        lay.addWidget(self.ingredients_list)

        # Tombol tambah bahan
        add_ing_btn = _btn(tr("food_recipe_add_ingredient"), h=32)
        add_ing_btn.clicked.connect(self._add_ingredient)
        lay.addWidget(add_ing_btn)

        # Catatan
        lay.addWidget(_lbl(tr("dialog_notes"), size=12))
        self.notes = _input(tr("food_recipe_instructions_ph"))
        lay.addWidget(self.notes)

        lay.addSpacing(8)
        save_btn = _btn(tr("food_recipe_save"), "solid", self._save)
        lay.addWidget(save_btn)

        root.addWidget(_scrolled(content))

    def _add_ingredient(self):
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("food_ingredient_picker_title"))
        dlg.setMinimumSize(450, 400)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        layout.setSpacing(8)

        layout.addWidget(QLabel(tr("food_ingredient_search_label")))
        search_input = QLineEdit()
        search_input.setPlaceholderText(tr("food_ingredient_search_ph"))
        search_input.setMinimumHeight(36)
        layout.addWidget(search_input)

        layout.addWidget(QLabel(tr("food_ingredient_select_label")))
        combo = QComboBox()
        combo.setMinimumHeight(42)
        layout.addWidget(combo)

        # Ambil semua makanan
        all_foods = db.get_food_items(self.user_id)
        # Simpan dalam bentuk list of (id, display_text, icon, name)
        foods_data = [(f["id"], f"{f['icon']} {f['name']} - {f['calories']:.0f} kcal", f["icon"], f["name"]) for f in all_foods]

        def update_combo(text=""):
            combo.blockSignals(True)
            combo.clear()
            search_text = text.strip().lower()
            for fid, display, icon, name in foods_data:
                if search_text == "" or search_text in name.lower() or search_text in display.lower():
                    combo.addItem(display, fid)
            combo.blockSignals(False)

        # Isi combo awal (semua)
        update_combo("")

        # Hubungkan event textChanged
        search_input.textChanged.connect(update_combo)

        layout.addWidget(QLabel(tr("food_ingredient_qty_label")))
        qty = QDoubleSpinBox()
        qty.setRange(0.1, 10)
        qty.setValue(1)
        qty.setSuffix(tr("unit_quantity"))
        layout.addWidget(qty)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(dlg.accept)
        buttons.rejected.connect(dlg.reject)
        layout.addWidget(buttons)

        # Jalankan dialog
        if dlg.exec() == QDialog.DialogCode.Accepted:
            food_id = combo.currentData()
            if food_id is None:
                return
            # Cari nama makanan untuk ditampilkan di list
            food_name = next((disp for fid, disp, _, _ in foods_data if fid == food_id), "")
            quantity = qty.value()
            self.selected_items.append((food_id, quantity))
            self.ingredients_list.addItem(f"{food_name} x{_fmt_qty(quantity)}")

    def _save(self):
        name = self.name.text().strip()
        if not name:
            _show(self, tr("msg_error"), tr("food_recipe_name_empty"), "error")
            return
        if not self.selected_items:
            _show(self, tr("msg_error"), tr("food_recipe_no_ingredient"), "error")
            return
        icon = self.icon.text().strip() or "🍲"
        serving = self.serving.value()
        notes = self.notes.text()
        r = db.add_recipe(self.user_id, name, icon, serving, notes, self.selected_items)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), tr("food_recipe_saved", name=name), "success")
            self.accept()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", "Gagal menyimpan resep"), "error")

class SetGoalsDialog(QDialog):
    """Dialog untuk mengatur target nutrisi harian."""
    def __init__(self, user_id, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.setWindowTitle(tr("food_goals_title_window"))
        self.setMinimumSize(400, 350)
        self.setStyleSheet(build_ss())
        self._build()

    def _build(self):
        goals = db.get_nutrition_goals(self.user_id)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(12)

        layout.addWidget(_lbl(tr("health_daily_targets"), "section", 14, True))
        layout.addWidget(_sep())

        layout.addWidget(_lbl(tr("food_calories_label"), size=12))
        self.calories = QSpinBox()
        self.calories.setRange(500, 10000)
        self.calories.setValue(goals["daily_calories"])
        layout.addWidget(self.calories)

        layout.addWidget(_lbl(tr("food_protein_label"), size=12))
        self.protein = QSpinBox()
        self.protein.setRange(0, 500)
        self.protein.setValue(goals["daily_protein"])
        layout.addWidget(self.protein)

        layout.addWidget(_lbl(tr("food_carbs_label"), size=12))
        self.carbs = QSpinBox()
        self.carbs.setRange(0, 500)
        self.carbs.setValue(goals["daily_carbs"])
        layout.addWidget(self.carbs)

        layout.addWidget(_lbl(tr("food_fat_label"), size=12))
        self.fat = QSpinBox()
        self.fat.setRange(0, 200)
        self.fat.setValue(goals["daily_fat"])
        layout.addWidget(self.fat)

        btn = _btn(tr("food_save_goals"), "solid", self._save)
        layout.addWidget(btn)

    def _save(self):
        db.update_nutrition_goals(self.user_id, self.calories.value(), self.protein.value(),
                                  self.carbs.value(), self.fat.value())
        SND.complete()
        self.accept()


class RecipeManagerDialog(QDialog):
    def __init__(self, user_id, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.setWindowTitle(tr("recipe_add_title"))
        self.setMinimumSize(500, 400)
        self.setStyleSheet(build_ss())
        self._build()
        self.load()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)

        # Header
        hdr = QHBoxLayout()
        hdr.addWidget(_lbl(tr("food_recipe_manager_title"), "section", 14, True))
        add_btn = _btn(tr("food_recipe_add_btn"), "solid", self._add_recipe)
        hdr.addWidget(add_btn)
        layout.addLayout(hdr)
        layout.addWidget(_sep())

        self.recipe_list = QListWidget()
        self.recipe_list.itemDoubleClicked.connect(self._view_recipe)
        layout.addWidget(self.recipe_list)

        # Tombol aksi
        btn_layout = QHBoxLayout()
        use_btn = _btn(tr("food_recipe_log_today"), "gold", self._log_selected_recipe)
        delete_btn = _btn(tr("food_recipe_delete"), "danger", self._delete_selected)
        btn_layout.addWidget(use_btn)
        btn_layout.addWidget(delete_btn)
        layout.addLayout(btn_layout)

    def load(self):
        self.recipe_list.clear()
        recipes = db.get_recipes(self.user_id)
        for r in recipes:
            self.recipe_list.addItem(f"{r['icon']} {r['name']} ({r['serving_size']} porsi)")

    def _add_recipe(self):
        dlg = AddRecipeDialog(self.user_id, self)
        if dlg.exec():
            QTimer.singleShot(0, self.load)

    def _view_recipe(self, item):
        recipe_name = item.text().split(' ', 1)[1].split(' (')[0]
        recipes = db.get_recipes(self.user_id)
        recipe = next((r for r in recipes if r['name'] == recipe_name), None)
        if recipe:
            details = db.get_recipe_details(recipe['id'])
            if details:
                msg = f"🍲 {details['recipe']['name']}\nPorsi standar: {details['recipe']['serving_size']}\n\nBahan:\n"
                for i in details['items']:
                    msg += f"- {i['name']} x{_fmt_qty(i['quantity'])} porsi\n"
                if details['recipe']['notes']:
                    msg += f"\nCatatan: {details['recipe']['notes']}"
                QMessageBox.information(self, tr("food_recipe_detail_title"), msg)

    def _log_selected_recipe(self):
        selected = self.recipe_list.currentItem()
        if not selected:
            return
        recipe_name = selected.text().split(' ', 1)[1].split(' (')[0]
        recipes = db.get_recipes(self.user_id)
        recipe = next((r for r in recipes if r['name'] == recipe_name), None)
        if not recipe:
            return
        # Tanyakan berapa porsi yang dimakan
        multiplier, ok = QInputDialog.getDouble(self, tr("food_serving_dialog_title"), tr("food_serving_dialog_msg", name=recipe['name']), 1, 0.25, 10, 1)
        if not ok:
            return
        # Tanyakan jenis makan
        meal_types = {"breakfast": tr("food_meal_breakfast"), "lunch": tr("food_meal_lunch"), "dinner": tr("food_meal_dinner"), "snack": tr("food_meal_snack")}
        meal, ok = QInputDialog.getItem(self, "Jenis Makan", "Pilih waktu makan:", list(meal_types.values()), 0, False)
        if not ok:
            return
        meal_key = {v: k for k, v in meal_types.items()}[meal]
        log_date = date.today().isoformat()
        r = db.log_recipe(self.user_id, recipe['id'], multiplier, meal_key, log_date, "")
        if r["ok"]:
            SND.complete()
            _show(self, tr('berhasil_title'), tr('food_recipe_log_success', name=recipe['name'], calories=r['calories']), "success")
            self.accept()  # tutup dialog, refresh FoodPage
        else:
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", tr("food_recipe_log_fail")), "error")

    def _delete_selected(self):
        selected = self.recipe_list.currentItem()
        if not selected:
            return
        recipe_name = selected.text().split(' ', 1)[1].split(' (')[0]
        recipes = db.get_recipes(self.user_id)
        recipe = next((r for r in recipes if r['name'] == recipe_name), None)
        if recipe:
            reply = QMessageBox.question(self, tr("confirm_title"), tr("food_recipe_delete_confirm", name=recipe_name), QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                db.delete_recipe(self.user_id, recipe['id'])
                SND.click()
                QTimer.singleShot(0, self.load)

# ═════════════════════════════════════════════════════════════════════════════
#  HEALTH DIALOG
# ═════════════════════════════════════════════════════════════════════════════
class HealthGoalsDialog(QDialog):
    def __init__(self, user_id, parent=None):
        super().__init__(parent)
        
        self.user_id = user_id
        self.setWindowTitle(tr("health_target_title"))
        self.setMinimumWidth(400)

        layout = QVBoxLayout(self)
        layout.setSpacing(12)

        self.steps_goal_input = QSpinBox()
        self.steps_goal_input.setRange(1000, 50000)
        self.steps_goal_input.setSingleStep(500)
        self.steps_goal_input.setSuffix(tr("unit_steps"))
        layout.addWidget(_lbl(tr("health_steps_goal_label"), size=12))
        layout.addWidget(self.steps_goal_input)

        self.sleep_goal_input = QDoubleSpinBox()
        self.sleep_goal_input.setRange(4.0, 12.0)
        self.sleep_goal_input.setSingleStep(0.5)
        self.sleep_goal_input.setSuffix(tr("unit_hours"))
        layout.addWidget(_lbl(tr("health_sleep_goal_label"), size=12))
        layout.addWidget(self.sleep_goal_input)

        self.water_goal_input = QSpinBox()
        self.water_goal_input.setRange(1000, 10000)
        self.water_goal_input.setSingleStep(250)
        self.water_goal_input.setSuffix(tr("unit_ml"))
        layout.addWidget(_lbl(tr("health_water_goal_label"), size=12))
        layout.addWidget(self.water_goal_input)

        self.calorie_goal_input = QSpinBox()
        self.calorie_goal_input.setRange(1000, 5000)
        self.calorie_goal_input.setSingleStep(100)
        self.calorie_goal_input.setSuffix(tr("unit_kcal"))
        layout.addWidget(_lbl(tr("health_calorie_goal_label"), size=12))
        layout.addWidget(self.calorie_goal_input)

        goals = db.get_health_goals(self.user_id)
        nutrition_goals = db.get_nutrition_goals(self.user_id)
        self.steps_goal_input.setValue(goals.get('daily_steps', 10000))
        self.sleep_goal_input.setValue(goals.get('daily_sleep_hours', 7.0))
        self.water_goal_input.setValue(db.get_water_goal(self.user_id))
        self.calorie_goal_input.setValue(nutrition_goals.get('daily_calories', 2000))

        buttons = QHBoxLayout()
        save_btn = _btn(tr("health_save_goals_btn"), "solid", self.accept)
        cancel_btn = _btn(tr("forgot_cancel_btn"), "ghost", self.reject)
        buttons.addWidget(save_btn)
        buttons.addWidget(cancel_btn)
        layout.addLayout(buttons)

    def accept(self):
        db.update_health_goals(
            self.user_id,
            self.steps_goal_input.value(),
            self.sleep_goal_input.value()
        )
        db.set_water_goal(self.user_id, self.water_goal_input.value())
        current_nutrition = db.get_nutrition_goals(self.user_id)
        db.update_nutrition_goals(
            self.user_id,
            self.calorie_goal_input.value(),
            current_nutrition.get('daily_protein', 80),
            current_nutrition.get('daily_carbs', 250),
            current_nutrition.get('daily_fat', 70)
        )
        super().accept()

    def reject(self):
        super().reject()

# ══════════════════════════════════════════════════════════════════════════════
#  HEALTH & FOOD PAGE
# ══════════════════════════════════════════════════════════════════════════════
class HealthFoodPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self.current_date = date.today()
        self._build()
        AppState.register(self.load)

    def _build(self):
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        container = QWidget()
        self.main_layout = QVBoxLayout(container)
        self.main_layout.setSpacing(16)
        self.main_layout.setContentsMargins(20, 20, 20, 20)

        # ── Header / Judul Halaman ──
        title_widget = QWidget()
        title_layout = QHBoxLayout(title_widget)
        title_layout.setContentsMargins(0, 0, 0, 0)
        self.title_label = _lbl(tr("health_title"), "section", 14, True)   # <-- simpan sebagai atribut
        title_layout.addWidget(self.title_label)
        title_layout.addStretch()
        self.main_layout.addWidget(title_widget)
        self.main_layout.addWidget(_sep())

        # Date selector
        self._build_date_selector()
        # Nutrition summary
        self._build_nutrition_summary()
        # BMI & auto goals
        self._build_bmi_section()
        # Action buttons
        self._build_action_buttons()
        # Food log
        self._build_food_log_section()
        # Water section
        self._build_water_section()
        # Health status cards
        self._build_health_status_section()
        # Health input
        self._build_health_input_section()
        # History charts & tips
        self._build_history_charts()

        scroll.setWidget(container)
        root_layout = QVBoxLayout(self)
        root_layout.addWidget(scroll)
        self.load()

    # ---------- Date selector ----------
    def _build_date_selector(self):
        date_widget = QWidget()
        date_layout = QHBoxLayout(date_widget)
        date_layout.setContentsMargins(0, 0, 0, 0)
        self.prev_date_btn = _btn(tr("nav_prev"), h=32)
        self.prev_date_btn.setMinimumWidth(40)
        self.prev_date_btn.clicked.connect(self._prev_date)
        self.date_label = _lbl("", size=13, bold=True)
        self.date_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.next_date_btn = _btn(tr("nav_next"), h=32)
        self.next_date_btn.setMinimumWidth(40)
        self.next_date_btn.clicked.connect(self._next_date)
        self.today_btn = _btn(tr("food_today"), h=32)
        self.today_btn.clicked.connect(self._today)
        date_layout.addWidget(self.prev_date_btn)
        date_layout.addWidget(self.date_label, 1)
        date_layout.addWidget(self.next_date_btn)
        date_layout.addWidget(self.today_btn)
        self.main_layout.addWidget(date_widget)

    def _prev_date(self):
        self.current_date -= timedelta(days=1)
        self.load()
    def _next_date(self):
        self.current_date += timedelta(days=1)
        self.load()
    def _today(self):
        self.current_date = date.today()
        self.load()

    # ---------- Nutrition summary ----------
    def _build_nutrition_summary(self):
        summary_widget = QWidget()
        summary_layout = QHBoxLayout(summary_widget)
        summary_layout.setSpacing(12)
        self.cal_card = self._stat_card(tr("food_cal_stat"), "0 / 0 kcal", "#f0a800")
        self.protein_card = self._stat_card(tr("food_protein_stat"), "0 / 0 g", "#80c000")
        self.carbs_card = self._stat_card(tr("food_carbs_stat"), "0 / 0 g", "#4da6ff")
        self.fat_card = self._stat_card(tr("food_fat_stat"), "0 / 0 g", "#e05050")
        summary_layout.addWidget(self.cal_card)
        summary_layout.addWidget(self.protein_card)
        summary_layout.addWidget(self.carbs_card)
        summary_layout.addWidget(self.fat_card)
        self.main_layout.addWidget(summary_widget)

        self.cal_progress = QProgressBar()
        self.cal_progress.setMinimumHeight(12)
        self.cal_progress.setTextVisible(False)
        self.main_layout.addWidget(self.cal_progress)

        goals_btn = _btn(tr("food_set_goals_btn"), "gold", self._open_goals)
        self.main_layout.addWidget(goals_btn)

    def _stat_card(self, title, value, color):
        card = _card()
        layout = QVBoxLayout(card)
        layout.setContentsMargins(10, 8, 10, 8)
        lbl_title = QLabel(title)
        lbl_title.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        lbl_value = QLabel(value)
        lbl_value.setStyleSheet(f"color:{color}; font-size:16px; font-weight:bold;")
        layout.addWidget(lbl_title)
        layout.addWidget(lbl_value)
        card.value_label = lbl_value
        return card

    def _open_goals(self):
        dlg = SetGoalsDialog(self.user_id, self)
        if dlg.exec():
            self.load()

    # ---------- Water section ----------
    def _build_water_section(self):
        water_group = QGroupBox(tr("food_tab_water"))
        water_layout = QVBoxLayout(water_group)

        goal_widget = QWidget()
        goal_layout = QHBoxLayout(goal_widget)
        self.water_goal_label = QLabel(tr("food_water_goal_default"))
        self.water_goal_label.setStyleSheet(f"color:{_T('light')}; font-weight:bold;")
        set_goal_btn = _btn(tr("food_water_set_goal"), "gold", self._set_water_goal)
        goal_layout.addWidget(self.water_goal_label)
        goal_layout.addStretch()
        goal_layout.addWidget(set_goal_btn)
        water_layout.addWidget(goal_widget)

        self.water_progress = QProgressBar()
        self.water_progress.setMinimumHeight(20)
        self.water_progress.setTextVisible(True)
        self.water_progress.setFormat("%v / %m ml")
        water_layout.addWidget(self.water_progress)

        add_water_widget = QWidget()
        add_water_layout = QHBoxLayout(add_water_widget)
        for amount, label in [(250, tr("food_water_add_250")), (500, tr("food_water_add_500")), (1000, tr("food_water_add_1000"))]:
            btn = _btn(label, "solid")
            btn.clicked.connect(partial(self._add_water, amount))
            add_water_layout.addWidget(btn)
        water_layout.addWidget(add_water_widget)

        custom_row = QHBoxLayout()
        self.water_custom = QSpinBox()
        self.water_custom.setRange(1, 5000)
        self.water_custom.setValue(1000)
        self.water_custom.setSuffix(tr("unit_ml"))
        custom_btn = _btn(tr("dialog_add"), h=36)
        custom_btn.clicked.connect(lambda: self._add_water(self.water_custom.value()))
        custom_row.addWidget(QLabel(tr("food_water_custom_label")))
        custom_row.addWidget(self.water_custom)
        custom_row.addWidget(custom_btn)
        water_layout.addLayout(custom_row)

        self.water_logs_group = QGroupBox(tr("food_water_log_title"))
        self.water_logs_layout = QVBoxLayout(self.water_logs_group)
        water_layout.addWidget(self.water_logs_group)

        self.main_layout.addWidget(water_group)

    def _add_water(self, amount):
        db.add_water_log(self.user_id, amount, self.current_date.isoformat())
        SND.complete()
        self.load()
        total = db.get_water_total(self.user_id, self.current_date.isoformat())
        goal = db.get_water_goal(self.user_id)
        if total >= goal and total - amount < goal:
            db.gain_xp_gold(self.user_id, 10, 2)
            db.add_notification(self.user_id, tr("water_goal_reached"), "success")
            AppState.refresh()

    def _delete_water_log(self, log_id):
        db.delete_water_log(self.user_id, log_id)
        SND.click()
        self.load()

    def _set_water_goal(self):
        from PyQt6.QtWidgets import QInputDialog
        current = db.get_water_goal(self.user_id)
        new_goal, ok = QInputDialog.getInt(self, tr("food_water_goal_dialog_title"), tr("food_water_goal_dialog_label"), current, 500, 10000, 100)
        if ok:
            db.set_water_goal(self.user_id, new_goal)
            SND.notify()
            self.load()

    # ---------- BMI section ----------
    def _build_bmi_section(self):
        bmi_group = QGroupBox(tr("food_bmi_title"))
        bmi_layout = QVBoxLayout(bmi_group)
        form = QFormLayout()
        self.height_input = QDoubleSpinBox()
        self.height_input.setRange(100, 250)
        self.height_input.setSuffix(tr("unit_cm"))
        self.weight_input = QDoubleSpinBox()
        self.weight_input.setRange(30, 300)
        self.weight_input.setSuffix(tr("unit_kg"))
        self.age_input = QSpinBox()
        self.age_input.setRange(15, 100)
        self.gender_combo = QComboBox()
        self.gender_combo.addItems([tr("food_bmi_gender_m"), tr("food_bmi_gender_f")])
        self.activity_combo = QComboBox()
        self.activity_combo.addItems([tr("food_bmi_activity_sedentary"), tr("food_bmi_activity_light"),
                                      tr("food_bmi_activity_moderate"), tr("food_bmi_activity_active"),
                                      tr("food_bmi_activity_very_active")])
        form.addRow(tr("food_bmi_height"), self.height_input)
        form.addRow(tr("food_bmi_weight"), self.weight_input)
        form.addRow(tr("food_bmi_age"), self.age_input)
        form.addRow(tr("food_bmi_gender"), self.gender_combo)
        form.addRow(tr("food_bmi_activity"), self.activity_combo)
        bmi_layout.addLayout(form)

        btn_row = QHBoxLayout()
        calc_btn = _btn(tr("food_bmi_calc"), "solid", self._calculate_bmi)
        set_target_btn = _btn(tr("food_bmi_set_target"), "gold", self._set_auto_goals)
        btn_row.addWidget(calc_btn)
        btn_row.addWidget(set_target_btn)
        bmi_layout.addLayout(btn_row)

        self.bmi_result_label = QLabel(tr("food_bmi_result"))
        self.bmi_result_label.setWordWrap(True)
        self.bmi_result_label.setStyleSheet(f"color:{_T('accent')}; font-weight:bold; padding:8px; background:{_T('panel')}; border-radius:6px;")
        bmi_layout.addWidget(self.bmi_result_label)

        self.main_layout.addWidget(bmi_group)
        self.load_bmi_settings()

    def load_bmi_settings(self):
        settings = db.get_user_bmi_settings(self.user_id)
        self.height_input.setValue(settings["height_cm"])
        self.weight_input.setValue(settings["weight_kg"])
        self.age_input.setValue(settings["age"])
        idx = self.gender_combo.findText(settings["gender"])
        if idx >= 0:
            self.gender_combo.setCurrentIndex(idx)
        activity_factors = [1.2, 1.375, 1.55, 1.725, 1.9]
        if settings["activity_factor"] in activity_factors:
            idx = activity_factors.index(settings["activity_factor"])
            self.activity_combo.setCurrentIndex(idx)

    def _calculate_bmi(self):
        height = self.height_input.value() / 100
        weight = self.weight_input.value()
        bmi = weight / (height * height)
        if bmi < 18.5:
            status = tr("food_bmi_status_underweight")
            color = "#f0a800"
        elif 18.5 <= bmi < 25:
            status = tr("food_bmi_status_normal")
            color = "#80c000"
        elif 25 <= bmi < 30:
            status = tr("food_bmi_status_overweight")
            color = "#e05050"
        else:
            status = tr("food_bmi_status_obese")
            color = "#e05050"
        self.bmi_result_label.setText(tr("food_bmi_result_format", color=color, bmi=bmi, status=status))

    def _set_auto_goals(self):
        height = self.height_input.value()
        weight = self.weight_input.value()
        age = self.age_input.value()
        gender = self.gender_combo.currentText()
        activity_idx = self.activity_combo.currentIndex()
        activity_factors = [1.2, 1.375, 1.55, 1.725, 1.9]
        factor = activity_factors[activity_idx]
        if gender == "Laki-laki":
            bmr = 10 * weight + 6.25 * height - 5 * age + 5
        else:
            bmr = 10 * weight + 6.25 * height - 5 * age - 161
        tdee = bmr * factor
        recommended_cal = int(tdee)
        recommended_protein = int(weight * 1.6)
        recommended_carbs = int(recommended_cal * 0.5 / 4)
        recommended_fat = int(recommended_cal * 0.3 / 9)
        db.update_nutrition_goals(self.user_id, recommended_cal, recommended_protein,
                                 recommended_carbs, recommended_fat)
        db.update_user_bmi_settings(self.user_id, height, weight, age, gender, factor)
        db.log_user_weight(self.user_id, weight)
        SND.complete()
        _show(self, tr("food_target_updated_title"),
              tr("food_target_updated_msg", cal=recommended_cal),
              "success")
        self.load()
        AppState.refresh()

    # ---------- Food log section ----------
    def _build_food_log_section(self):
        self.logs_group = QGroupBox(tr("food_log_group_title"))
        self.logs_group.setObjectName("food_log_group")   # opsional
        self.logs_layout = QVBoxLayout(self.logs_group)
        self.main_layout.addWidget(self.logs_group)

    def _refresh_food_log(self, log_date):
        # Bersihkan layout
        while self.logs_layout.count():
            item = self.logs_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        logs = db.get_food_logs(self.user_id, log_date)
        if not logs:
            empty = QLabel(tr("food_no_logs_today"))
            empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            empty.setStyleSheet(f"color:{_T('muted')}; padding:20px;")
            self.logs_layout.addWidget(empty)
            return

        meal_order = ["breakfast", "lunch", "dinner", "snack"]
        meal_names = {"breakfast": tr("food_meal_breakfast"),
                      "lunch": tr("food_meal_lunch"),
                      "dinner": tr("food_meal_dinner"),
                      "snack": tr("food_meal_snack")}

        for meal in meal_order:
            meal_logs = [l for l in logs if l["meal_type"] == meal]
            if meal_logs:
                # Label jenis makan
                meal_label = QLabel(meal_names[meal])
                meal_label.setStyleSheet(f"color:{_T('light')}; font-weight:bold; margin-top:8px;")
                self.logs_layout.addWidget(meal_label)

                # Tampilkan kartu log
                for log in meal_logs:
                    card = self._make_food_log_card(log)
                    self.logs_layout.addWidget(card)

                # ── Drop area untuk memindahkan log ke meal type ini ──
                drop_area = QFrame()
                drop_area.setAcceptDrops(True)
                drop_area.setMinimumHeight(40)
                drop_area.setStyleSheet(f"""
                    QFrame {{
                        border: 2px dashed {_T('border')};
                        border-radius: 6px;
                        background: {_T('bg')};
                        margin: 4px 0;
                    }}
                    QFrame:hover {{
                        border-color: {_T('accent')};
                        background: {_T('panel')};
                    }}
                """)

                def make_drop_handler(target_meal):
                    def handler(event):
                        if not event.mimeData().hasFormat("application/x-craftlife-card"):
                            event.ignore()
                            return
                        raw = event.mimeData().data("application/x-craftlife-card")
                        info = json.loads(raw.data().decode())
                        if info["mode"] == "food" and info["user_id"] == self.user_id:
                            # Gunakan move_item_to_folder dengan mode "food"
                            db.move_item_to_folder(self.user_id, "food", info["item_id"], target_meal)
                            self.load()
                            event.acceptProposedAction()
                        else:
                            event.ignore()
                    return handler

                def make_drag_enter_handler():
                    def handler(event):
                        if event.mimeData().hasFormat("application/x-craftlife-card"):
                            event.acceptProposedAction()
                        else:
                            event.ignore()
                    return handler

                drop_area.dropEvent = make_drop_handler(meal)
                drop_area.dragEnterEvent = make_drag_enter_handler()

                self.logs_layout.addWidget(drop_area)

    def _make_food_log_card(self, log: dict) -> QFrame:
        content = QWidget()
        row = QHBoxLayout(content)
        row.setContentsMargins(12, 8, 12, 8)
        row.setSpacing(10)

        icon = _emoji_label(log["icon"], ICON_CARD)
        icon.setMinimumWidth(40)
        row.addWidget(icon)

        info = QVBoxLayout()
        info.setSpacing(2)
        lang = AppState.get_language()
        display_name = get_food_name(log['name'], lang)
        name = QLabel(tr("food_log_name_serving", name=display_name, serving=log['serving']))
        name.setStyleSheet(f"font-size:13px; font-weight:bold; color:{_T('text')};")
        info.addWidget(name)

        nut = QLabel(tr("food_nutrition_detail", cal=log['calories'], protein=log['protein'], carbs=log['carbs'], fat=log['fat']))
        nut.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        info.addWidget(nut)

        if log.get("notes"):
            notes = QLabel(f"📝 {log['notes']}")
            notes.setStyleSheet(f"color:{_T('muted')}; font-size:10px; font-style:italic;")
            info.addWidget(notes)

        row.addLayout(info, 1)

        del_btn = _btn("🗑", "danger", h=32)
        del_btn.setMinimumWidth(36)
        del_btn.clicked.connect(lambda _, lid=log["id"]: self._delete_food_log(lid))
        row.addWidget(del_btn)

        card = DraggableCard(
            item_id=log["id"],
            mode="food",
            user_id=self.user_id,
            current_folder_id=log.get("meal_type"),
            content_widget=content,
            parent=self
        )
        return card

    def _delete_food_log(self, log_id):
        db.delete_food_log(self.user_id, log_id)
        SND.click()
        self.load()

    # ---------- Health status section ----------
    def _build_health_status_section(self):
        self.status_grid_layout = QGridLayout()
        self.status_grid_layout.setSpacing(10)
        self.main_layout.addLayout(self.status_grid_layout)

        self.note_display_widget = QWidget()
        note_layout = QHBoxLayout(self.note_display_widget)
        note_layout.setContentsMargins(0, 8, 0, 0)
        
        self.note_display_label = QLabel()
        self.note_display_label.setWordWrap(True)
        self.note_display_label.setStyleSheet(
            f"color:{_T('muted')}; background:{_T('panel')}; "
            f"padding:10px; border-radius:6px; border:1px solid {_T('border')};"
            f"font-style:italic;"
        )
        note_layout.addWidget(self.note_display_label)
        self.main_layout.addWidget(self.note_display_widget)

    def _refresh_health_status(self, log_date):
        # Bersihkan grid
        while self.status_grid_layout.count():
            item = self.status_grid_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        # Ambil data
        conn = db.get_conn()
        today_row = conn.execute(
            "SELECT * FROM health_logs WHERE user_id=? AND log_date=?",
            (self.user_id, log_date)
        ).fetchone()
        conn.close()
        today_log = dict(today_row) if today_row else None

        health_goals = db.get_health_goals(self.user_id)
        water_total = db.get_water_total(self.user_id, log_date)
        water_goal = db.get_water_goal(self.user_id)
        nutrition = db.get_nutrition_summary(self.user_id, log_date)
        nutrition_goals = db.get_nutrition_goals(self.user_id)
        burned = db.get_total_calories_burned_today(self.user_id, log_date)
        net_cal = nutrition['calories'] - burned
        steps_goal = health_goals.get('daily_steps', 10000)
        sleep_goal = health_goals.get('daily_sleep_hours', 7.0)

        # Berat: prioritas historis, fallback ke user_health_goals
        if today_log and today_log.get('weight_kg') is not None:
            weight_kg = today_log['weight_kg']
        else:
            weight_kg = health_goals.get('weight_kg', 70.0)
        height_cm = health_goals.get('height_cm', 170)

        # Buat kartu
        cards = [
            self._stat_card(tr("health_steps"), tr("health_steps_value", steps=today_log['steps'] if today_log else 0, goal=steps_goal), "#80c000"),
            self._stat_card(tr("health_sleep"), tr("health_sleep_value", sleep=today_log['sleep_hours'] if today_log else 0, goal=sleep_goal), "#4da6ff"),
            self._stat_card(tr("health_water"), tr("health_water_value", water=water_total, goal=water_goal), "#38bdf8"),
            self._stat_card(tr("health_mood"), tr("health_mood_value", mood=(today_log['mood'] if today_log else "normal").capitalize()), "#f0a800"),
            self._stat_card(tr("health_weight"), tr("health_weight_value", weight=weight_kg), "#a97fff"),
            self._stat_card(tr("health_height"), tr("health_height_value", height=height_cm), "#a97fff"),
            self._stat_card(tr("health_hr"), tr("health_hr_value", hr=today_log.get('resting_hr', 0) if today_log else 0), "#e05050"),
            self._stat_card(tr("health_stress"), tr("health_stress_value", stress=(today_log.get('stress_level','normal') if today_log else "normal").capitalize()), "#80c000"),
            self._stat_card(tr("health_calories"), tr("health_calories_value", cal=nutrition['calories'], goal=nutrition_goals['daily_calories']), "#ff9f1c"),
            self._stat_card(tr("health_protein"), tr("health_protein_value", protein=nutrition['protein'], goal=nutrition_goals['daily_protein']), "#f4a261"),
            self._stat_card(tr("health_burned"), tr("health_burned_value", burned=burned), "#e76f51"),
            self._stat_card(tr("health_net_calories"), tr("health_net_calories_value", net=net_cal, goal=nutrition_goals['daily_calories']), "#2a9d8f"),
        ]
        for i, card in enumerate(cards):
            row = i // 4
            col = i % 4
            self.status_grid_layout.addWidget(card, row, col)

        # ── Tampilkan Catatan (Note) jika ada ──
        if today_log and today_log.get('notes'):
            self.note_display_label.setText(f"📝  {today_log['notes']}")
            self.note_display_label.setStyleSheet(
                f"color:{_T('text')}; background:{_T('panel')}; "
                f"padding:10px; border-radius:6px; border:1px solid {_T('border')};"
            )
        else:
            self.note_display_label.setText(tr("health_note_placeholder"))
            self.note_display_label.setStyleSheet(
                f"color:{_T('muted')}; background:{_T('panel')}; "
                f"padding:10px; border-radius:6px; border:1px solid {_T('border')};"
                f"font-style:italic;"
            )

    # ---------- Health input section ----------
    def _build_health_input_section(self):
        input_group = QGroupBox(tr("health_tab_input"))
        input_layout = QVBoxLayout(input_group)

        phys_box = QGroupBox(tr("health_activity_group"))
        phys_form = QFormLayout(phys_box)
        self.steps_input = QSpinBox()
        self.steps_input.setRange(0, 50000)
        self.steps_input.setSuffix(tr("unit_steps"))
        self.hr_input = QSpinBox()
        self.hr_input.setRange(0, 220)
        self.hr_input.setSuffix(" bpm")

        self.weight_input_daily = QDoubleSpinBox()
        self.weight_input_daily.setRange(30, 300)
        self.weight_input_daily.setSingleStep(0.5)
        self.weight_input_daily.setSuffix(tr("unit_kg"))
        self.height_input_daily = QDoubleSpinBox()
        self.height_input_daily.setRange(100, 250)
        self.height_input_daily.setSingleStep(0.5)
        self.height_input_daily.setSuffix(tr("unit_cm"))

        phys_form.addRow(tr("health_steps_label"), self.steps_input)
        phys_form.addRow(tr("health_hr_label"), self.hr_input)
        phys_form.addRow(tr("health_weight"), self.weight_input_daily)
        phys_form.addRow(tr("health_height"), self.height_input_daily)
        input_layout.addWidget(phys_box)

        sleep_box = QGroupBox(tr("health_sleep_group"))
        sleep_form = QFormLayout(sleep_box)
        self.sleep_input = QDoubleSpinBox()
        self.sleep_input.setRange(0, 24)
        self.sleep_input.setSingleStep(0.5)
        self.sleep_input.setSuffix(tr("unit_hours"))
        self.stress_combo = QComboBox()
        self.stress_combo.addItem(tr("health_stress_low"), "low")
        self.stress_combo.addItem(tr("health_stress_normal"), "normal")
        self.stress_combo.addItem(tr("health_stress_high"), "high")
        self.mood_combo = QComboBox()
        self.mood_combo.addItem(tr("health_mood_happy"), "happy")
        self.mood_combo.addItem(tr("health_mood_normal"), "normal")
        self.mood_combo.addItem(tr("health_mood_tired"), "tired")
        self.mood_combo.addItem(tr("health_mood_sad"), "sad")
        sleep_form.addRow(tr("health_sleep_label"), self.sleep_input)
        sleep_form.addRow(tr("health_stress_label"), self.stress_combo)
        sleep_form.addRow(tr("health_mood_label"), self.mood_combo)
        input_layout.addWidget(sleep_box)

        notes_box = QGroupBox(tr("health_notes_group"))
        notes_vlay = QVBoxLayout(notes_box)
        self.notes_input = QLineEdit()
        self.notes_input.setPlaceholderText(tr("health_notes_placeholder"))
        notes_vlay.addWidget(self.notes_input)
        input_layout.addWidget(notes_box)

        save_btn = _btn(tr("health_save"), "solid", self._save_health)
        input_layout.addWidget(save_btn)
        self.main_layout.addWidget(input_group)

    def _refresh_health_input(self, log_date):
        conn = db.get_conn()
        row = conn.execute("SELECT * FROM health_logs WHERE user_id=? AND log_date=?", (self.user_id, log_date)).fetchone()
        conn.close()
        today_log = dict(row) if row else None

        # Ambil goals untuk fallback
        health_goals = db.get_health_goals(self.user_id)

        if today_log:
            self.steps_input.setValue(today_log.get('steps', 0))
            self.sleep_input.setValue(today_log.get('sleep_hours', 0))
            self.hr_input.setValue(today_log.get('resting_hr', 0))
            # ── AMBIL BERAT & TINGGI DARI LOG (jika ada) ──
            self.weight_input_daily.setValue(today_log.get('weight_kg', health_goals.get('weight_kg', 70.0)))
            self.height_input_daily.setValue(today_log.get('height_cm', health_goals.get('height_cm', 170.0)))
            self.mood_combo.setCurrentIndex(self.mood_combo.findData(today_log.get('mood', 'normal')))
            self.stress_combo.setCurrentIndex(self.stress_combo.findData(today_log.get('stress_level', 'normal')))
            self.notes_input.setText(today_log.get('notes', ''))
        else:
            self.steps_input.setValue(0)
            self.sleep_input.setValue(0)
            self.hr_input.setValue(0)
            # ── FALLBACK KE GOALS ──
            self.weight_input_daily.setValue(health_goals.get('weight_kg', 70.0))
            self.height_input_daily.setValue(health_goals.get('height_cm', 170.0))
            self.mood_combo.setCurrentIndex(1)
            self.stress_combo.setCurrentIndex(1)
            self.notes_input.clear()

    def _save_health(self):
        log_date = self.current_date.isoformat()
        water_total = db.get_water_total(self.user_id, log_date)
        nutrition = db.get_nutrition_summary(self.user_id, log_date)
        calories_burned = db.get_total_calories_burned_today(self.user_id, log_date)
        net_calories = nutrition['calories'] - calories_burned

        steps = self.steps_input.value()
        sleep = self.sleep_input.value()
        resting_hr = self.hr_input.value()
        weight = self.weight_input_daily.value()
        height = self.height_input_daily.value()
        mood = self.mood_combo.currentData()
        stress = self.stress_combo.currentData()
        notes = self.notes_input.text()

        db.add_health_log(
            self.user_id, log_date,
            steps=steps,
            sleep_hours=sleep,
            water_ml=water_total,
            weight_kg=weight,                   # kirim berat
            resting_hr=resting_hr,
            stress_level=stress,
            mood=mood,
            notes=notes,
            net_calories=net_calories
        )

        # ── UPDATE JUGA KE user_health_goals (agar BMI dan grafik konsisten) ──
        current_goals = db.get_health_goals(self.user_id)
        db.update_health_goals(
            self.user_id,
            current_goals.get('daily_steps', 10000),
            current_goals.get('daily_sleep_hours', 7.0),
            height_cm=height,
            weight_kg=weight
        )

        SND.complete()
        _show(self, tr("saved_title"), tr("health_data_saved"), "success")
        self.load()
        AppState.refresh()

    # ---------- History charts ----------
    def _build_history_charts(self):
        # Trend cards
        trend_group = QGroupBox(tr("health_avg_7days"))
        trend_layout = QHBoxLayout(trend_group)
        self.avg_steps_card = self._trend_card(tr("health_avg_steps"), "0", tr("health_avg_7days_suffix"), "#80c000")
        self.avg_sleep_card = self._trend_card(tr("health_avg_sleep"), "0 jam", tr("health_avg_7days_suffix"), "#4da6ff")
        self.avg_water_card = self._trend_card(tr("health_avg_water"), "0 ml", tr("health_avg_7days_suffix"), "#38bdf8")
        self.avg_hr_card = self._trend_card(tr("health_avg_hr"), "0 bpm", tr("health_avg_7days_suffix"), "#e05050")
        trend_layout.addWidget(self.avg_steps_card)
        trend_layout.addWidget(self.avg_sleep_card)
        trend_layout.addWidget(self.avg_water_card)
        trend_layout.addWidget(self.avg_hr_card)
        self.main_layout.addWidget(trend_group)

        self.weight_chart_label = QLabel()
        self.weight_chart_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        weight_group = QGroupBox(tr("health_weight_trend"))
        weight_group_layout = QVBoxLayout(weight_group)
        weight_group_layout.addWidget(self.weight_chart_label)
        self.main_layout.addWidget(weight_group)

        self.height_chart_label = QLabel()
        self.height_chart_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        height_group = QGroupBox(tr("health_height_trend"))
        height_group_layout = QVBoxLayout(height_group)
        height_group_layout.addWidget(self.height_chart_label)
        self.main_layout.addWidget(height_group)

        tips_group = QGroupBox(tr("health_tips"))
        tips_layout = QVBoxLayout(tips_group)
        self.tips_label = QLabel()
        self.tips_label.setWordWrap(True)
        tips_layout.addWidget(self.tips_label)
        self.main_layout.addWidget(tips_group)

    def _trend_card(self, title, value, note, color=_T("light")):
        card = _card()
        layout = QVBoxLayout(card)
        layout.setContentsMargins(12, 10, 12, 10)
        t = QLabel(title)
        t.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        v = QLabel(value)
        v.setStyleSheet(f"color:{color}; font-size:16px; font-weight:bold;")
        n = QLabel(note)
        n.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        layout.addWidget(t)
        layout.addWidget(v)
        layout.addWidget(n)
        card.value_label = v
        return card

    def _refresh_history_charts(self):
        logs_7days = db.get_health_logs(self.user_id, days=7)
        health_goals = db.get_health_goals(self.user_id)
        water_goal = db.get_water_goal(self.user_id)
        # Rata-rata
        if logs_7days:
            avg_steps = sum(l['steps'] for l in logs_7days) // len(logs_7days)
            avg_sleep = round(sum(l['sleep_hours'] for l in logs_7days) / len(logs_7days), 1)
            avg_water = sum(l.get('water_ml', 0) for l in logs_7days) // len(logs_7days)
            avg_hr = sum(l.get('resting_hr', 0) for l in logs_7days) // len(logs_7days)
        else:
            avg_steps = avg_sleep = avg_water = avg_hr = 0
        self.avg_steps_card.value_label.setText(f"{avg_steps}")
        self.avg_sleep_card.value_label.setText(f"{avg_sleep} {tr('health_unit_hour')}")
        self.avg_water_card.value_label.setText(f"{avg_water} {tr('health_unit_ml')}")
        self.avg_hr_card.value_label.setText(f"{avg_hr} {tr('health_unit_bpm')}")

        # Grafik berat
        self._update_weight_chart(logs_7days)
        # Grafik tinggi
        self._update_height_chart(logs_7days, health_goals)
        # Tips
        self._update_tips(logs_7days)

    def _update_weight_chart(self, logs_7days):
        if not EXPORT_IMPORTS_OK:
            return
        try:
            import matplotlib.pyplot as plt
            from io import BytesIO
            from PyQt6.QtGui import QPixmap
            days = []
            weights = []
            for i in range(7):
                d = (date.today() - timedelta(days=6-i)).isoformat()
                days.append(d[5:])
                log = next((l for l in logs_7days if l['log_date'] == d), None)
                if log and log.get('weight_kg') is not None:
                    weights.append(float(log['weight_kg']))
                else:
                    weights.append(0.0)
            fig, ax = plt.subplots(figsize=(6, 3))
            ax.plot(days, weights, marker='o', color='#80c000', linewidth=2)
            ax.fill_between(days, weights, color='#80c000', alpha=0.2)
            ax.set_ylabel('Berat (kg)' if AppState.get_language()=="id" else 'Weight (kg)', color='#e8e8e8')
            ax.set_title(tr("health_chart_weight"), color='#7bbf3e')
            ax.set_facecolor('#2d2d2d')
            fig.patch.set_facecolor('#2d2d2d')
            ax.tick_params(colors='#e8e8e8')
            plt.setp(ax.get_xticklabels(), rotation=20, ha='right')
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=80, bbox_inches='tight')
            buf.seek(0)
            plt.close(fig)
            pixmap = QPixmap()
            pixmap.loadFromData(buf.read())
            self.weight_chart_label.setPixmap(pixmap.scaledToWidth(500, Qt.TransformationMode.SmoothTransformation))
        except Exception as e:
            print(f"Weight chart error: {e}")

    def _update_height_chart(self, logs_7days, health_goals):
        if not EXPORT_IMPORTS_OK:
            return
        try:
            import matplotlib.pyplot as plt
            from io import BytesIO
            from PyQt6.QtGui import QPixmap
            height_cm = health_goals.get('height_cm', 170)
            days = []
            heights = []
            for i in range(7):
                d = (date.today() - timedelta(days=6-i)).isoformat()
                days.append(d[5:])
                heights.append(height_cm)
            fig, ax = plt.subplots(figsize=(6, 3))
            ax.plot(days, heights, marker='o', color='#4da6ff', linewidth=2, linestyle='-')
            ax.fill_between(days, heights, color='#4da6ff', alpha=0.2)
            ax.set_ylabel('Tinggi (cm)' if AppState.get_language()=="id" else 'Height (cm)', color='#e8e8e8')
            ax.set_title(tr("health_chart_height"), color='#7bbf3e')
            ax.set_facecolor('#2d2d2d')
            fig.patch.set_facecolor('#2d2d2d')
            ax.tick_params(colors='#e8e8e8')
            plt.setp(ax.get_xticklabels(), rotation=20, ha='right')
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=80, bbox_inches='tight')
            buf.seek(0)
            plt.close(fig)
            pixmap = QPixmap()
            pixmap.loadFromData(buf.read())
            self.height_chart_label.setPixmap(pixmap.scaledToWidth(500, Qt.TransformationMode.SmoothTransformation))
        except Exception as e:
            print(f"Height chart error: {e}")

    def _update_tips(self, logs_7days):
        today = date.today().isoformat()
        nutrition = db.get_nutrition_summary(self.user_id, today)
        burned = db.get_total_calories_burned_today(self.user_id, today)
        net_cal = nutrition['calories'] - burned
        goals = db.get_nutrition_goals(self.user_id)
        dynamic_tip = ""
        if net_cal < 0:
            dynamic_tip = tr("health_tip_calorie_deficit")
        elif net_cal > goals['daily_calories'] * 1.1:
            dynamic_tip = tr("health_tip_calorie_surplus")
        else:
            dynamic_tip = tr("health_tip_calorie_normal")
        import random
        tip_keys = [f"health_tip_static_{i}" for i in range(1, 8)]
        random_tip = tr(random.choice(tip_keys))
        self.tips_label.setText(f"{random_tip} {dynamic_tip}")

    # ---------- Action buttons ----------
    def _build_action_buttons(self):
        btn_layout = QHBoxLayout()
        add_food_btn = _btn(tr("food_add_custom"), h=36)
        add_food_btn.clicked.connect(self._open_add_food)
        log_food_btn = _btn(tr("food_log"), "solid", self._open_log_food)
        recipe_btn = _btn(tr("food_recipes"), h=36)
        recipe_btn.clicked.connect(self._open_recipe_manager)
        export_btn = _btn(tr("food_export"), h=36)
        export_btn.clicked.connect(self._export_nutrition)
        btn_layout.addWidget(add_food_btn)
        btn_layout.addWidget(log_food_btn)
        btn_layout.addWidget(recipe_btn)
        btn_layout.addWidget(export_btn)
        self.main_layout.addLayout(btn_layout)

    def _open_add_food(self):
        dlg = AddFoodDialog(self.user_id, mode="add", parent=self)
        if dlg.exec():
            self.load()
    def _open_log_food(self):
        dlg = AddFoodDialog(self.user_id, mode="log", parent=self)
        if dlg.exec():
            self.load()
    def _open_recipe_manager(self):
        dlg = RecipeManagerDialog(self.user_id, self)
        if dlg.exec():
            self.load()
    def _export_nutrition(self):
        # Salin dari FoodPage._export_nutrition
        from PyQt6.QtWidgets import QFileDialog, QComboBox, QDialog, QVBoxLayout, QLabel, QDialogButtonBox
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("food_export_format_title"))
        dlg.setModal(True)
        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel(tr("economy_export_label")))
        combo = QComboBox()
        combo.addItems([tr("export_csv_option"), "Excel (.xlsx)", "Word (.docx)", "PDF (.pdf)"])
        layout.addWidget(combo)
        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(dlg.accept)
        buttons.rejected.connect(dlg.reject)
        layout.addWidget(buttons)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        fmt = combo.currentText()
        food_data = db.get_food_export_data(self.user_id, days=30)
        if fmt == "CSV (.csv)":
            path, _ = QFileDialog.getSaveFileName(self, tr("food_save_csv"), "", tr("food_csv_filter"))
            if path:
                if not path.endswith(".csv"): path += ".csv"
                self._export_food_csv(food_data, path)
        elif fmt == "Excel (.xlsx)":
            path, _ = QFileDialog.getSaveFileName(self, tr("food_save_excel"), "", tr("food_xlsx_filter"))
            if path:
                if not path.endswith(".xlsx"): path += ".xlsx"
                self._export_food_excel(food_data, path)
        elif fmt == "Word (.docx)":
            path, _ = QFileDialog.getSaveFileName(self, tr("food_save_word"), "", tr("food_docx_filter"))
            if path:
                if not path.endswith(".docx"): path += ".docx"
                self._export_food_word(food_data, path)
        else:  # PDF
            path, _ = QFileDialog.getSaveFileName(self, tr("food_save_pdf"), "", tr("food_pdf_filter"))
            if path:
                if not path.endswith(".pdf"): path += ".pdf"
                self._export_food_pdf(food_data, path)

    # Ekspor (copy dari FoodPage)
    def _export_food_csv(self, food_data, filepath):
        import csv
        try:
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow([tr("export_date"), tr("export_calories"), tr("export_protein"), tr("export_carbs"), tr("export_fat"), tr("export_water_ml"), tr("export_calories_burned"), tr("export_net_calories")])
                for fd in food_data:
                    writer.writerow([fd['date'], fd['calories'], fd['protein'], fd['carbs'], fd['fat'], fd['water_ml'], fd['calories_burned'], fd['net_calories']])
            _show(self, tr("berhasil_title"), f"Data nutrisi diekspor ke {filepath}", "success")
        except Exception as e:
            _show(self, tr("msg_error"), f"Gagal mengekspor CSV: {str(e)}", "error")

    def _export_food_excel(self, food_data, filepath):
        if not EXPORT_IMPORTS_OK:
            _show(self, tr("msg_error"), tr("export_lib_openpyxl"), "error")
            return
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill
            from openpyxl.utils import get_column_letter
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = tr("excel_sheet_nutrition")
            headers = [tr("export_date"), tr("export_calories"), tr("export_protein"), tr("export_carbs"), tr("export_fat"), tr("export_water_ml"), tr("export_calories_burned"), tr("export_net_calories")]
            ws.append(headers)
            for col in range(1, len(headers)+1):
                cell = ws.cell(row=1, column=col)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="5a8a2e", end_color="5a8a2e", fill_type="solid")
                cell.alignment = Alignment(horizontal="center")
            for fd in food_data:
                ws.append([fd['date'], fd['calories'], fd['protein'], fd['carbs'], fd['fat'], fd['water_ml'], fd['calories_burned'], fd['net_calories']])
            for col in ws.columns:
                max_len = 0
                col_letter = get_column_letter(col[0].column)
                for cell in col:
                    try: max_len = max(max_len, len(str(cell.value)))
                    except Exception: pass
                ws.column_dimensions[col_letter].width = min(max_len + 2, 20)
            wb.save(filepath)
            _show(self, tr("berhasil_title"), f"Data nutrisi diekspor ke {filepath}", "success")
        except Exception as e:
            _show(self, tr("msg_error"), f"Gagal mengekspor: {str(e)}", "error")

    def _export_food_word(self, food_data, filepath):
        if not EXPORT_IMPORTS_OK:
            _show(self, tr("msg_error"), tr("export_lib_docx"), "error") 
            return
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = Document()
            title = doc.add_heading("CraftLife - Data Nutrisi & Air", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            total_cal = sum(fd['calories'] for fd in food_data)
            avg_cal = total_cal / len(food_data) if food_data else 0
            total_water = sum(fd['water_ml'] for fd in food_data)
            avg_water = total_water / len(food_data) if food_data else 0
            doc.add_heading("Ringkasan 30 Hari", level=1)
            doc.add_paragraph(f"📊 Total Kalori: {total_cal:.0f} kcal")
            doc.add_paragraph(f"📈 Rata-rata Kalori Harian: {avg_cal:.0f} kcal")
            doc.add_paragraph(f"💧 Total Air: {total_water:.0f} ml")
            doc.add_paragraph(f"🚰 Rata-rata Air Harian: {avg_water:.0f} ml")
            doc.add_heading("Data Harian", level=1)
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = [tr("export_date"), tr("export_calories"), tr("export_protein"), tr("export_carbs"), tr("export_fat"), tr("export_water_ml"), tr("export_calories_burned"), tr("export_net_calories")]
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
            for fd in food_data:
                row_cells = table.add_row().cells
                row_cells[0].text = fd['date']
                row_cells[1].text = str(fd['calories'])
                row_cells[2].text = str(fd['protein'])
                row_cells[3].text = str(fd['carbs'])
                row_cells[4].text = str(fd['fat'])
                row_cells[5].text = str(fd['water_ml'])
                row_cells[6].text = str(fd['calories_burned'])
                row_cells[7].text = str(fd['net_calories'])
            doc.save(filepath)
            _show(self, tr("berhasil_title"), f"Data nutrisi diekspor ke {filepath}", "success")
        except Exception as e:
            _show(self, tr("msg_error"), f"Gagal mengekspor: {str(e)}", "error")

    def _export_food_pdf(self, food_data, filepath):
        if not EXPORT_IMPORTS_OK:
            _show(self, tr("msg_error"), tr("export_lib_reportlab"), "error") 
            return
        try:
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            doc = SimpleDocTemplate(filepath, pagesize=landscape(A4))
            story = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(name='Title', parent=styles['Title'], alignment=1, fontSize=16)
            story.append(Paragraph("CraftLife - Data Nutrisi & Air", title_style))
            story.append(Spacer(1, 0.2*inch))
            total_cal = sum(fd['calories'] for fd in food_data)
            avg_cal = total_cal / len(food_data) if food_data else 0
            total_water = sum(fd['water_ml'] for fd in food_data)
            avg_water = total_water / len(food_data) if food_data else 0
            summary_data = [
                ["Total Kalori 30 Hari", f"{total_cal:.0f} kcal"],
                ["Rata-rata Kalori Harian", f"{avg_cal:.0f} kcal"],
                ["Total Air 30 Hari", f"{total_water:.0f} ml"],
                ["Rata-rata Air Harian", f"{avg_water:.0f} ml"],
            ]
            sum_table = Table(summary_data, colWidths=[3*inch, 2*inch])
            sum_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#5a8a2e")),
                ('TEXTCOLOR', (0,0), (0,-1), colors.white),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ]))
            story.append(sum_table)
            story.append(Spacer(1, 0.2*inch))
            data = [[tr("export_date"), tr("export_calories"), tr("export_protein"), tr("export_carbs"), tr("export_fat"), tr("export_water_ml"), tr("export_calories_burned"), tr("export_net_calories")]]
            for fd in food_data:
                data.append([fd['date'], str(fd['calories']), str(fd['protein']), str(fd['carbs']), str(fd['fat']), str(fd['water_ml']), str(fd['calories_burned']), str(fd['net_calories'])])
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#5a8a2e")),
                ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 8),
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ]))
            story.append(table)
            doc.build(story)
            _show(self, tr("berhasil_title"), f"Data nutrisi diekspor ke {filepath}", "success")
        except Exception as e:
            _show(self, tr("msg_error"), f"Gagal mengekspor: {str(e)}", "error")

    # ---------- Bonus ----------
    def _check_nutrition_bonus(self, log_date):
        if self.current_date == date.today():
            r = db.check_daily_nutrition_bonus(self.user_id, log_date)
            if r.get("ok"):
                SND.level_up()
                msg = tr("health_nutrition_bonus_msg", xp=r['xp_gained'], gold=r['gold_gained'])
                if r.get("leveled_up"): 
                    msg += tr("level_up_bonus")
                _show(self, tr("bonus_title"), msg, "success")
                AppState.refresh()

    # ---------- LOAD ----------
    def load(self):
        if not AppState.user_id:
            return
        log_date = self.current_date.isoformat()
        today = date.today()
        self.date_label.setText(self.current_date.strftime("%A, %d %B %Y"))
        self.next_date_btn.setEnabled(self.current_date < today)

        # ── UPDATE TERJEMAHAN ──
        self.title_label.setText(tr("health_title"))
        self.logs_group.setTitle(tr("food_log_group_title"))

        # Refresh semua bagian
        self._refresh_nutrition_summary(log_date)
        self._refresh_water_section(log_date)
        self._refresh_food_log(log_date)
        self._refresh_health_status(log_date)
        self._refresh_health_input(log_date)
        self._refresh_history_charts()
        self._check_nutrition_bonus(log_date)

    def _refresh_nutrition_summary(self, log_date):
        summary = db.get_nutrition_summary(self.user_id, log_date)
        goals = db.get_nutrition_goals(self.user_id)
        self.cal_card.value_label.setText(f"{summary['calories']:.0f} / {goals['daily_calories']} kcal")
        self.protein_card.value_label.setText(f"{summary['protein']:.0f} / {goals['daily_protein']} g")
        self.carbs_card.value_label.setText(f"{summary['carbs']:.0f} / {goals['daily_carbs']} g")
        self.fat_card.value_label.setText(f"{summary['fat']:.0f} / {goals['daily_fat']} g")
        cal_percent = int((summary['calories'] / goals['daily_calories']) * 100) if goals['daily_calories'] > 0 else 0
        self.cal_progress.setValue(min(100, cal_percent))
        if cal_percent >= 100:
            self.cal_progress.setStyleSheet("QProgressBar::chunk { background: #e05050; }")
        elif cal_percent >= 80:
            self.cal_progress.setStyleSheet("QProgressBar::chunk { background: #f0a800; }")
        else:
            self.cal_progress.setStyleSheet("QProgressBar::chunk { background: #80c000; }")

    def _refresh_water_section(self, log_date):
        water_goal = db.get_water_goal(self.user_id)
        water_total = db.get_water_total(self.user_id, log_date)
        self.water_goal_label.setText(f"Target: {water_goal} ml")
        self.water_progress.setMaximum(water_goal)
        self.water_progress.setValue(water_total)
        self.water_progress.setFormat(tr("water_progress_format", current=water_total, goal=water_goal))
        # Refresh water logs
        while self.water_logs_layout.count():
            item = self.water_logs_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        water_logs = db.get_water_logs(self.user_id, log_date)
        if not water_logs:
            empty = QLabel(tr("food_no_water_today"))
            empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            empty.setStyleSheet(f"color:{_T('muted')}; padding:20px;")
            self.water_logs_layout.addWidget(empty)
        else:
            for wlog in water_logs:
                card = _card()
                row = QHBoxLayout(card)
                row.setContentsMargins(12, 8, 12, 8)
                row.addWidget(QLabel("💧"))
                row.addWidget(QLabel(f"+{wlog['amount_ml']} ml"), 1)
                from datetime import datetime
                try:
                    dt = datetime.fromisoformat(wlog['created_at'])
                    time_str = dt.strftime("%H:%M")
                except Exception:
                    time_str = wlog['created_at'][11:16]
                row.addWidget(QLabel(time_str))
                del_btn = _btn("🗑", "danger", h=28)
                del_btn.setMinimumWidth(36)
                del_btn.clicked.connect(lambda _, lid=wlog["id"]: self._delete_water_log(lid))
                row.addWidget(del_btn)
                self.water_logs_layout.addWidget(card)

    def closeEvent(self, e):
        AppState.unregister(self.load)
        super().closeEvent(e)

# ══════════════════════════════════════════════════════════════════════════════
#  Calendar Page
# ══════════════════════════════════════════════════════════════════════════════
class CalendarPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self.current_year = datetime.now().year
        self.current_month = datetime.now().month
        self.holidays = {}
        self.notes = {}
        self._holidays_fetched = False
        self._loading_holidays = False
        self._build()
        AppState.register(self.load)
        AppState.register_lang_cb(self.load)

    def _build(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(20, 16, 20, 16)
        main_layout.setSpacing(10)

        # ── Header ──
        hdr = QHBoxLayout()
        hdr.addWidget(_lbl(tr("calendar_title"), "section", 14, True))
        hdr.addStretch()

        self.prev_year_btn = _btn("◀", h=34)
        self.prev_year_btn.clicked.connect(self._prev_year)
        self.year_label = _lbl(str(self.current_year), size=14, bold=True)
        self.next_year_btn = _btn("▶", h=34)
        self.next_year_btn.clicked.connect(self._next_year)
        self.today_btn = _btn(tr("food_today"), h=34)
        self.today_btn.clicked.connect(self._goto_today)

        hdr.addWidget(self.prev_year_btn)
        hdr.addWidget(self.year_label)
        hdr.addWidget(self.next_year_btn)
        hdr.addWidget(self.today_btn)
        main_layout.addLayout(hdr)
        main_layout.addWidget(_sep())

        # ── Scroll Area ──
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.scroll_area.setStyleSheet(f"""
            QScrollArea {{
                border: none;
                background: transparent;
            }}
        """)

        self.months_container = QWidget()
        self.months_grid = QGridLayout(self.months_container)
        self.months_grid.setSpacing(20)               # lebih lega
        self.months_grid.setContentsMargins(10, 10, 10, 10)

        # Agar 3 kolom meregang rata
        for col in range(3):
            self.months_grid.setColumnStretch(col, 1)

        self.scroll_area.setWidget(self.months_container)
        main_layout.addWidget(self.scroll_area, 1)    # stretch agar memenuhi sisa ruang

        self.load()

    def load(self):
        if not AppState.user_id:
            return
        self._fetch_notes()
        if not self._holidays_fetched and not self._loading_holidays:
            self._loading_holidays = True
            QTimer.singleShot(100, self._fetch_holidays)
        self._render()

    def _fetch_notes(self):
        """Ambil catatan untuk tahun ini + tahun lalu + tahun depan (3 tahun)."""
        all_notes = {}
        for y in [self.current_year - 1, self.current_year, self.current_year + 1]:
            notes = db.get_calendar_notes(self.user_id, year=y)
            all_notes.update(notes)
        self.notes = all_notes

    def _fetch_holidays(self):
        holidays = {}
        for year in [self.current_year - 1, self.current_year, self.current_year + 1]:
            year_data = get_holidays_for_year(year)
            holidays.update(year_data)
        self.holidays = holidays
        self._holidays_fetched = True
        self._loading_holidays = False
        print(f"[Calendar] Loaded {len(holidays)} holidays for {self.current_year-1}-{self.current_year+1}")
        QTimer.singleShot(0, self._render)

    def _render(self):
        # Kosongkan grid
        while self.months_grid.count():
            item = self.months_grid.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Render 12 bulan
        for month in range(1, 13):
            month_widget = self._make_month(month)
            row = (month - 1) // 3
            col = (month - 1) % 3
            self.months_grid.addWidget(month_widget, row, col)

        self.year_label.setText(str(self.current_year))

    def _make_month(self, month):
        widget = QWidget()
        widget.setObjectName("card")
        widget.setMinimumHeight(420)   # cukup untuk 6 baris + header
        widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)

        layout = QVBoxLayout(widget)
        layout.setContentsMargins(8, 8, 8, 8)   # margin minimal
        layout.setSpacing(6)

        # Nama bulan
        month_name = tr(f"month_{month:02d}")
        lbl_month = QLabel(month_name)
        lbl_month.setStyleSheet(f"font-size:16px; font-weight:bold; color:{_T('light')}; padding-bottom:4px;")
        lbl_month.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(lbl_month)

        # Grid hari
        grid = QGridLayout()
        grid.setHorizontalSpacing(4)    # jarak horizontal (kiri-kanan) tetap rapat
        grid.setVerticalSpacing(10)     # jarak vertikal (atas-bawah) lebih longgar
        grid.setContentsMargins(0, 0, 0, 0)

        # Header hari (Senin–Minggu)
        days_abbr = [tr(f"day_{i}") for i in range(7)]
        for col, d in enumerate(days_abbr):
            lbl = QLabel(d)
            lbl.setStyleSheet(f"font-size:11px; color:{_T('muted')}; font-weight:bold;")
            lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            grid.addWidget(lbl, 0, col)

        # Isi tanggal
        first_day, num_days = calmod.monthrange(self.current_year, month)
        row = 1
        col = first_day
        for day in range(1, num_days + 1):
            date_str = f"{self.current_year:04d}-{month:02d}-{day:02d}"
            is_today = (date_str == datetime.now().strftime("%Y-%m-%d"))
            is_holiday = date_str in self.holidays
            has_note = date_str in self.notes and self.notes[date_str].strip() != ""

            btn = QPushButton(str(day))
            btn.setMinimumSize(40, 44)          # ukuran lebih proporsional
            btn.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {_T('bg')};
                    border: 1px solid {_T('border')};
                    border-radius: 6px;
                    font-size: 13px;
                    font-weight: bold;
                }}
                QPushButton:hover {{
                    background: {_T('primary')};
                    color: #fff;
                }}
            """)
            if is_today:
                btn.setStyleSheet(btn.styleSheet() + f"""
                    QPushButton {{
                        background: {_T('primary')};
                        color: #fff;
                        border: 2px solid {_T('accent')};
                    }}
                """)
            if is_holiday:
                holiday_name = get_holiday_name(date_str, AppState.get_language())
                btn.setToolTip(f"🏷️ {holiday_name}" if holiday_name else "Libur")
                btn.setStyleSheet(btn.styleSheet() + "QPushButton { color: #e05050; }")
            if has_note:
                btn.setText(f"{day}\n📝")
                btn.setStyleSheet(btn.styleSheet() + f"""
                    QPushButton {{
                        border-color: {_T('accent')};
                    }}
                """)

            btn.clicked.connect(lambda checked, d=date_str, n=has_note: self._open_note_dialog(d, n))
            grid.addWidget(btn, row, col)
            col += 1
            if col >= 7:
                col = 0
                row += 1

        # ⭐ Kunci: Beri bobot peregangan pada setiap baris agar tinggi terdistribusi rata
        for r in range(1, row):   # row terakhir adalah jumlah baris + 1
            grid.setRowStretch(r, 1)

        layout.addLayout(grid)
        return widget

    def _open_note_dialog(self, date_str, has_existing):
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("calendar_note_title", date=date_str))
        dlg.setMinimumSize(400, 300)
        dlg.setStyleSheet(build_ss())

        layout = QVBoxLayout(dlg)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)

        # Tampilkan hari libur jika ada
        holiday_name = get_holiday_name(date_str, AppState.get_language())
        if holiday_name:
            info = QLabel(tr("calendar_holiday_info", name=holiday_name))
            info.setStyleSheet(f"color:{_T('accent')}; font-weight:bold;")
            layout.addWidget(info)

        layout.addWidget(QLabel(tr("calendar_note_label")))

        note_edit = QTextEdit()
        note_edit.setPlaceholderText(tr("calendar_note_placeholder"))
        if has_existing:
            note_edit.setText(self.notes.get(date_str, ""))
        layout.addWidget(note_edit)

        # Tombol
        button_box = QDialogButtonBox()
        save_btn = button_box.addButton(tr("dialog_save"), QDialogButtonBox.ButtonRole.AcceptRole)
        delete_btn = button_box.addButton(tr("calendar_delete"), QDialogButtonBox.ButtonRole.DestructiveRole)
        cancel_btn = button_box.addButton(tr("btn_cancel"), QDialogButtonBox.ButtonRole.RejectRole)

        save_btn.clicked.connect(lambda: self._save_note(date_str, note_edit.toPlainText(), dlg))
        delete_btn.clicked.connect(lambda: self._delete_note(date_str, dlg))
        cancel_btn.clicked.connect(dlg.reject)

        layout.addWidget(button_box)
        dlg.exec()

    def _save_note(self, date_str, note, dialog):
        if note.strip():
            db.save_calendar_note(self.user_id, date_str, note.strip())
            SND.complete()
        else:
            db.delete_calendar_note(self.user_id, date_str)
            SND.click()
        dialog.accept()
        QTimer.singleShot(0, self.load)

    def _delete_note(self, date_str, dialog):
        db.delete_calendar_note(self.user_id, date_str)
        SND.click()
        dialog.accept()
        QTimer.singleShot(0, self.load)

    def _prev_year(self):
        self.current_year -= 1
        self._fetch_holidays()
        QTimer.singleShot(0, self.load)

    def _next_year(self):
        self.current_year += 1
        self._fetch_holidays()
        QTimer.singleShot(0, self.load)

    def _goto_today(self):
        self.current_year = datetime.now().year
        self.current_month = datetime.now().month
        self._fetch_holidays()
        QTimer.singleShot(0, self.load)

    def _toggle_left_panel(self, collapsed):
        # Cari left widget (index 0 di splitter)
        try:
            left = self.findChild(QWidget, "learning_left")
            if not left:
                # Fallback: ambil widget pertama di splitter
                left = self.sender().parent().findChild(QSplitter).widget(0) if hasattr(self, 'sender') else None
            # Simpan sizes
            if not hasattr(self, '_splitter'):
                # Cari splitter
                for child in self.findChildren(QSplitter):
                    self._splitter = child
                    break
            if hasattr(self, '_splitter'):
                sizes = self._splitter.sizes()
                if collapsed:
                    self._splitter.setSizes([0, 600, 320])
                    self.btn_collapse_left.setText("Sources ▶")
                else:
                    self._splitter.setSizes([280, 400, 320])
                    self.btn_collapse_left.setText("◀ Sources")
        except Exception as e:
            print(f"Toggle left failed: {e}")

    def _toggle_right_panel(self, collapsed):
        try:
            if not hasattr(self, '_splitter'):
                for child in self.findChildren(QSplitter):
                    self._splitter = child
                    break
            if hasattr(self, '_splitter'):
                if collapsed:
                    self._splitter.setSizes([280, 400, 0])
                    self.btn_collapse_right.setText("◀ Studio")
                else:
                    self._splitter.setSizes([280, 400, 320])
                    self.btn_collapse_right.setText("Studio ▶")
        except Exception as e:
            print(f"Toggle right failed: {e}")

    def closeEvent(self, e):
        AppState.unregister(self.load)
        AppState.unregister_lang_cb(self.load)
        super().closeEvent(e)

# ══════════════════════════════════════════════════════════════════════════════
#  Notes Page
# ══════════════════════════════════════════════════════════════════════════════
class NotesTextEdit(QTextEdit):
    """Custom QTextEdit untuk paste dengan format yang disesuaikan."""
    def __init__(self, parent):
        super().__init__(parent)
        self.notes_page = parent

    def insertFromMimeData(self, source):
        # FIX PRESERVE: keep original HTML formatting (font/size/color) from source
        if source.hasHtml():
            html = source.html()
            self.insertHtml(html)
        elif source.hasText():
            txt = source.text()
            # Auto-konversi LaTeX → simbol matematika saat paste, supaya
            # hasil copy seperti "\frac{2^5\cdot2^{-3}}{2^2}" langsung tampil
            # sebagai kalimat matematika, bukan teks mentah.
            if mathtools.has_latex(txt):
                txt = mathtools.latex_to_unicode(txt)
            self.insertPlainText(txt)
        else:
            super().insertFromMimeData(source)


class MathPreviewDialog(QDialog):
    """Render ekspresi LaTeX yang terdeteksi di catatan via matplotlib mathtext
    — tampil seperti rumus di buku, plus sumber teksnya."""
    def __init__(self, chunks: list, parent=None):
        super().__init__(parent)
        self.setWindowTitle(tr("notes_math_preview_title"))
        self.setMinimumSize(540, 400)
        self.setStyleSheet(build_ss())
        lay = QVBoxLayout(self)
        content = QWidget()
        v = QVBoxLayout(content)
        v.setSpacing(10)
        if not chunks:
            none_lbl = _lbl(tr("notes_math_none"), "sub", 12)
            none_lbl.setWordWrap(True)
            v.addWidget(none_lbl)
        else:
            color = _T("text")
            for expr in chunks:
                card = _card()
                cv = QVBoxLayout(card)
                pm = self._render_expr(expr, color)
                if pm is not None:
                    img = QLabel()
                    img.setPixmap(pm)
                    img.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    cv.addWidget(img)
                else:
                    fail_lbl = _lbl(tr("notes_math_render_fail"), "sub", 10)
                    cv.addWidget(fail_lbl)
                src = QLabel(expr)
                src.setWordWrap(True)
                src.setAlignment(Qt.AlignmentFlag.AlignCenter)
                src.setStyleSheet(f"color:{_T('muted')}; font-size:10px;")
                cv.addWidget(src)
                v.addWidget(card)
        v.addStretch()
        lay.addWidget(_scrolled(content))

    def _render_expr(self, expr, color):
        if not MPL_QT_OK:
            return None
        try:
            fig = Figure(figsize=(0.01, 0.01), tight_layout=False)
            fig.patch.set_alpha(0.0)
            fig.text(0.5, 0.5, f"${expr}$", fontsize=17, color=color,
                     ha="center", va="center")
            buf = BytesIO()
            fig.savefig(buf, dpi=110, transparent=True,
                        bbox_inches="tight", pad_inches=0.15)
            buf.seek(0)
            pm = QPixmap()
            pm.loadFromData(buf.getvalue(), "PNG")
            return pm if not pm.isNull() else None
        except Exception:
            return None

class NotesPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self.current_folder_id = None
        self.current_note_id = None
        self.show_archived = False
        self.search_text = ""
        self._is_dirty = False
        self._note_font_settings = {}
        self._default_font_size = 12
        self._default_color = QColor("#e8e8e8")
        self._current_bold = False
        self._current_italic = False
        self._current_underline = False
        self._current_font_size = 12
        self._current_color = QColor("#e8e8e8")
        self._current_zoom = 100
        self._base_html = ""
        self._is_scaling = False
        
        self._build()
        AppState.register(self.load)
        AppState.register_lang_cb(self.load)

    def _build(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(20, 16, 20, 16)
        main_layout.setSpacing(10)

        # === HEADER ===
        hdr = QHBoxLayout()
        hdr.addWidget(_lbl(tr("notes_title"), "section", 14, True))
        hdr.addStretch()
        self.archive_btn = _btn(tr("notes_archive"), h=34)
        self.archive_btn.clicked.connect(self._toggle_archive)
        hdr.addWidget(self.archive_btn)
        self.archive_toggle_btn = _btn(tr("notes_show_archived"), h=34)
        self.archive_toggle_btn.clicked.connect(self._toggle_show_archived)
        hdr.addWidget(self.archive_toggle_btn)
        main_layout.addLayout(hdr)
        main_layout.addWidget(_sep())

        # === SEARCH ===
        search_widget = QWidget()
        search_layout = QHBoxLayout(search_widget)
        search_layout.setContentsMargins(0, 0, 0, 0)
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText(tr("notes_search_placeholder"))
        self.search_input.textChanged.connect(self._on_search)
        search_layout.addWidget(self.search_input)
        main_layout.addWidget(search_widget)

        # === TOOLBAR ===
        toolbar = QHBoxLayout()
        toolbar.setSpacing(6)
        add_folder_btn = _btn(tr("notes_add_folder"), h=34)
        add_folder_btn.clicked.connect(self._add_folder)
        add_note_btn = _btn(tr("notes_add_note"), "solid", h=34)
        add_note_btn.clicked.connect(self._add_note)
        dup_note_btn = _btn(tr("notes_duplicate_btn"), h=34)
        dup_note_btn.clicked.connect(self._duplicate_note)
        dup_note_btn.setToolTip(tr("notes_duplicate_tooltip"))
        delete_btn = _btn(tr("notes_delete"), "danger", h=34)
        delete_btn.clicked.connect(self._delete_selected)
        toolbar.addWidget(add_folder_btn)
        toolbar.addWidget(add_note_btn)
        toolbar.addWidget(dup_note_btn)
        toolbar.addWidget(delete_btn)
        toolbar.addStretch()
        main_layout.addLayout(toolbar)

        # === SPLITTER: LEFT (Folder Tree + Notes List) | RIGHT (Editor) ===
        splitter = QSplitter(Qt.Orientation.Horizontal)

        # ----- LEFT PANEL -----
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(6)

        # Folder Tree
        folder_row = QHBoxLayout()
        folder_row.addWidget(_lbl(tr("notes_folder_label"), size=12))
        folder_row.addStretch()
        self.expand_btn = _btn(tr("expand_all"), h=28)
        self.expand_btn.clicked.connect(lambda: self.folder_tree.expandAll())
        self.collapse_btn = _btn(tr("collapse_all"), h=28)
        self.collapse_btn.clicked.connect(lambda: self.folder_tree.collapseAll())
        folder_row.addWidget(self.expand_btn)
        folder_row.addWidget(self.collapse_btn)

        self.folder_tree = QTreeWidget()
        self.folder_tree.setHeaderHidden(True)
        self.folder_tree.setMinimumHeight(200)
        self.folder_tree.setIndentation(20)
        self.folder_tree.itemClicked.connect(self._on_folder_selected)
        self.folder_tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.folder_tree.customContextMenuRequested.connect(self._show_folder_context_menu)

        left_layout.addLayout(folder_row)
        left_layout.addWidget(self.folder_tree)

        # Notes List
        left_layout.addWidget(_lbl(tr("notes_list_label"), size=12))
        self.notes_list = QListWidget()
        self.notes_list.setMinimumHeight(200)
        self.notes_list.itemClicked.connect(self._on_note_selected)
        left_layout.addWidget(self.notes_list)

        # Note Count
        self.note_count_label = _lbl("", "sub", 11)
        left_layout.addWidget(self.note_count_label)

        splitter.addWidget(left_widget)

        # ----- RIGHT PANEL (Editor) -----
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(10, 0, 0, 0)
        right_layout.setSpacing(6)

        # Title
        right_layout.addWidget(_lbl(tr("notes_title_label"), size=12))
        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText(tr("notes_note_title_ph"))
        self.title_edit.setMinimumHeight(34)
        self.title_edit.textChanged.connect(self._mark_dirty)
        right_layout.addWidget(self.title_edit)

        # ---- MATH TOOLBAR ----
        math_scroll = QScrollArea()
        math_scroll.setWidgetResizable(True)
        math_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        math_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        math_scroll.setFixedHeight(80)
        math_scroll.setStyleSheet(f"""
            QScrollArea {{
                background: transparent;
                border: none;
            }}
            QScrollBar:horizontal {{
                background: {_T('bg')};
                height: 10px;
                border-radius: 5px;
            }}
            QScrollBar::handle:horizontal {{
                background: {_T('border')};
                border-radius: 5px;
                min-width: 30px;
            }}
            QScrollBar::handle:horizontal:hover {{
                background: {_T('primary')};
            }}
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{
                width: 0px;
            }}
        """)
        math_container = QWidget()
        math_container.setStyleSheet(f"background: transparent; padding: 8px 0px;")
        math_layout = QHBoxLayout(math_container)
        math_layout.setContentsMargins(4, 4, 4, 4)
        math_layout.setSpacing(6)

        symbols = [
            ("²", "²"), ("³", "³"), ("√", "√"), ("π", "π"),
            ("×", "×"), ("÷", "÷"), ("±", "±"), ("∑", "∑"),
            ("∫", "∫"), ("∞", "∞"), ("α", "α"), ("β", "β"),
            ("γ", "γ"), ("θ", "θ"), ("λ", "λ"), ("μ", "μ"),
            ("σ", "σ"), ("τ", "τ"), ("φ", "φ"), ("ω", "ω"),
            ("∂", "∂"), ("∇", "∇"), ("∆", "∆"), ("ℵ", "ℵ"),
            ("ℜ", "ℜ"), ("ℑ", "ℑ"), ("℘", "℘"),
        ]
        for label, sym in symbols:
            btn = QPushButton(label)
            btn.setFixedSize(40, 40)
            btn.setStyleSheet(f"""
                font-size: 18px;
                font-weight: bold;
                background: {_T('panel')};
                border: 1px solid {_T('border')};
                border-radius: 6px;
            """)
            btn.clicked.connect(lambda checked, s=sym: self._insert_symbol(s))
            math_layout.addWidget(btn)

        sup_btn = QPushButton("x²")
        sup_btn.setFixedSize(40, 40)
        sup_btn.setStyleSheet(f"""
            font-size: 15px;
            background: {_T('panel')};
            border: 1px solid {_T('border')};
            border-radius: 6px;
        """)
        sup_btn.clicked.connect(self._insert_superscript)
        math_layout.addWidget(sup_btn)

        sub_btn = QPushButton("x₂")
        sub_btn.setFixedSize(40, 40)
        sub_btn.setStyleSheet(f"""
            font-size: 15px;
            background: {_T('panel')};
            border: 1px solid {_T('border')};
            border-radius: 6px;
        """)
        sub_btn.clicked.connect(self._insert_subscript)
        math_layout.addWidget(sub_btn)

        frac_btn = QPushButton("a/b")
        frac_btn.setFixedSize(40, 40)
        frac_btn.setStyleSheet(f"""
            font-size: 15px;
            background: {_T('panel')};
            border: 1px solid {_T('border')};
            border-radius: 6px;
        """)
        frac_btn.clicked.connect(self._insert_fraction)
        math_layout.addWidget(frac_btn)

        # Tombol ∑ LaTeX: menu konversi & pratinjau render (matplotlib)
        self._latex_btn = QPushButton("∑ ƒ(x)")
        self._latex_btn.setFixedSize(64, 40)
        self._latex_btn.setStyleSheet(f"""
            font-size: 13px; font-weight: bold;
            background: {_T('panel')};
            border: 1px solid {_T('border')};
            border-radius: 6px;
        """)
        self._latex_btn.setToolTip(tr("notes_math_btn_tooltip"))
        self._latex_btn.clicked.connect(self._latex_menu)
        math_layout.addWidget(self._latex_btn)

        math_layout.addStretch()
        math_scroll.setWidget(math_container)
        right_layout.addWidget(math_scroll)

        # ---- FONT TOOLBAR ----
        font_toolbar = QWidget()
        font_toolbar.setStyleSheet(f"background: transparent; border: none; padding: 4px 0px;")
        font_layout = QHBoxLayout(font_toolbar)
        font_layout.setContentsMargins(8, 4, 8, 4)
        font_layout.setSpacing(6)

        font_label = QLabel("Font:")
        font_label.setStyleSheet(f"color: {_T('muted')}; font-size: 12px;")
        font_layout.addWidget(font_label)

        self.font_size_combo = QComboBox()
        self.font_size_combo.setMinimumHeight(32)
        self.font_size_combo.setFixedWidth(70)
        self.font_size_combo.setStyleSheet(f"""
            QComboBox {{
                background: {_T('panel')};
                color: {_T('text')};
                border: 1px solid {_T('border')};
                border-radius: 4px;
                padding: 4px 8px;
                font-size: 12px;
            }}
            QComboBox::drop-down {{
                border: none;
            }}
            QComboBox QAbstractItemView {{
                background: {_T('panel')};
                color: {_T('text')};
                selection-background-color: {_T('primary')};
            }}
        """)
        font_sizes = [8, 9, 10, 11, 12, 14, 16, 18, 20, 24, 28, 32, 36, 48, 72]
        for s in font_sizes:
            self.font_size_combo.addItem(f"{s}px", s)
        self.font_size_combo.setCurrentIndex(4)
        self.font_size_combo.currentIndexChanged.connect(self._change_font_size)
        font_layout.addWidget(self.font_size_combo)

        sep1 = QFrame()
        sep1.setFrameShape(QFrame.Shape.VLine)
        sep1.setStyleSheet(f"color: {_T('border')};")
        sep1.setFixedHeight(28)
        font_layout.addWidget(sep1)

        self.color_btn = QPushButton("🎨")
        self.color_btn.setFixedSize(36, 32)
        self.color_btn.setStyleSheet(f"""
            QPushButton {{
                font-size: 16px;
                background: {_T('panel')};
                border: 1px solid {_T('border')};
                border-radius: 4px;
            }}
            QPushButton:hover {{
                background: {_T('border')};
            }}
        """)
        self.color_btn.clicked.connect(self._choose_font_color)
        font_layout.addWidget(self.color_btn)

        self.color_indicator = QLabel()
        self.color_indicator.setFixedSize(28, 28)
        self.color_indicator.setStyleSheet(f"""
            background: {_T('text')};
            border: 1px solid {_T('border')};
            border-radius: 4px;
        """)
        font_layout.addWidget(self.color_indicator)

        sep2 = QFrame()
        sep2.setFrameShape(QFrame.Shape.VLine)
        sep2.setStyleSheet(f"color: {_T('border')};")
        sep2.setFixedHeight(28)
        font_layout.addWidget(sep2)

        self.bold_btn = QPushButton("B")
        self.bold_btn.setFixedSize(36, 32)
        self.bold_btn.setCheckable(True)
        self.bold_btn.setStyleSheet(f"""
            QPushButton {{
                font-weight: bold;
                font-size: 14px;
                background: {_T('panel')};
                border: 1px solid {_T('border')};
                border-radius: 4px;
            }}
            QPushButton:checked {{
                background: {_T('primary')};
                color: #fff;
                border-color: {_T('light')};
            }}
            QPushButton:hover {{
                background: {_T('border')};
            }}
        """)
        self.bold_btn.clicked.connect(lambda: self._toggle_format('bold'))
        font_layout.addWidget(self.bold_btn)

        self.italic_btn = QPushButton("I")
        self.italic_btn.setFixedSize(36, 32)
        self.italic_btn.setCheckable(True)
        self.italic_btn.setStyleSheet(f"""
            QPushButton {{
                font-style: italic;
                font-size: 14px;
                background: {_T('panel')};
                border: 1px solid {_T('border')};
                border-radius: 4px;
            }}
            QPushButton:checked {{
                background: {_T('primary')};
                color: #fff;
                border-color: {_T('light')};
            }}
            QPushButton:hover {{
                background: {_T('border')};
            }}
        """)
        self.italic_btn.clicked.connect(lambda: self._toggle_format('italic'))
        font_layout.addWidget(self.italic_btn)

        self.underline_btn = QPushButton("U")
        self.underline_btn.setFixedSize(36, 32)
        self.underline_btn.setCheckable(True)
        self.underline_btn.setStyleSheet(f"""
            QPushButton {{
                text-decoration: underline;
                font-size: 14px;
                background: {_T('panel')};
                border: 1px solid {_T('border')};
                border-radius: 4px;
            }}
            QPushButton:checked {{
                background: {_T('primary')};
                color: #fff;
                border-color: {_T('light')};
            }}
            QPushButton:hover {{
                background: {_T('border')};
            }}
        """)
        self.underline_btn.clicked.connect(lambda: self._toggle_format('underline'))
        font_layout.addWidget(self.underline_btn)

        sep_zoom = QFrame()
        sep_zoom.setFrameShape(QFrame.Shape.VLine)
        sep_zoom.setStyleSheet(f"color: {_T('border')};")
        sep_zoom.setFixedHeight(28)
        font_layout.addWidget(sep_zoom)

        zoom_label = QLabel(tr("notes_zoom_label"))
        zoom_label.setStyleSheet(f"color: {_T('muted')}; font-size: 12px;")
        font_layout.addWidget(zoom_label)

        self.zoom_slider = QSlider(Qt.Orientation.Horizontal)
        self.zoom_slider.setRange(50, 200)
        self.zoom_slider.setValue(100)
        self.zoom_slider.setTickInterval(10)
        self.zoom_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.zoom_slider.setFixedWidth(120)
        self.zoom_slider.setStyleSheet(f"""
            QSlider::groove:horizontal {{
                height: 6px;
                background: {_T('border')};
                border-radius: 3px;
            }}
            QSlider::handle:horizontal {{
                background: {_T('primary')};
                width: 16px;
                margin: -5px 0;
                border-radius: 8px;
            }}
        """)
        self.zoom_slider.valueChanged.connect(self._on_zoom_changed)
        font_layout.addWidget(self.zoom_slider)

        self.zoom_value_label = QLabel("100%")
        self.zoom_value_label.setStyleSheet(f"color: {_T('text')}; font-size: 12px; font-weight: bold; min-width: 40px;")
        font_layout.addWidget(self.zoom_value_label)

        font_layout.addStretch()
        right_layout.addWidget(font_toolbar)

        # ---- EDITOR (QTextEdit LANGSUNG, tanpa QGraphicsView) ----
        self.content_edit = QTextEdit()
        self.content_edit.setAcceptRichText(True)
        self.content_edit.setPlaceholderText(tr("notes_note_content_ph"))
        self.content_edit.cursorPositionChanged.connect(self._apply_default_format)
        self.content_edit.textChanged.connect(self._mark_dirty)
        self.content_edit.setMinimumHeight(200)
        right_layout.addWidget(self.content_edit, 1)  # stretch

        # ---- SAVE BUTTON ----
        save_btn = _btn(tr("notes_save"), "solid", self._save_note, 40)
        right_layout.addWidget(save_btn)

        splitter.addWidget(right_widget)
        splitter.setSizes([350, 700])

        main_layout.addWidget(splitter, 1)

        self._is_dirty = False
        self.load()

    # ========== LOAD / REFRESH ==========
    def load(self):
        if not AppState.user_id:
            return
        # Preserve current folder selection across reloads (FIX 2: don't jump to All Notes)
        preserve_id = self.current_folder_id
        # Block signals to avoid double _load_notes during tree rebuild
        self.folder_tree.blockSignals(True)
        self._load_folder_tree(preserve_id=preserve_id)
        self.folder_tree.blockSignals(False)
        # Now load notes for preserved folder
        self._load_notes()

    def _load_folder_tree(self, preserve_id=None):
        # preserve_id: None means -1 (All), 0 means No Folder, else folder id
        if preserve_id is None and hasattr(self, 'current_folder_id'):
            preserve_id = self.current_folder_id
        # Map None -> -1 for tree data value
        if preserve_id is None:
            target_data = -1
        elif preserve_id == 0:
            target_data = 0
        else:
            target_data = preserve_id
        self.folder_tree.clear()
        root_item = QTreeWidgetItem(self.folder_tree)
        root_item.setText(0, "📂 " + tr("notes_all"))
        root_item.setData(0, Qt.ItemDataRole.UserRole, -1)
        root_item.setExpanded(True)

        no_folder_item = QTreeWidgetItem(root_item)
        no_folder_item.setText(0, "📭 " + tr("notes_no_folder"))
        no_folder_item.setData(0, Qt.ItemDataRole.UserRole, 0)
        no_folder_item.setExpanded(True)

        tree_data = db.get_note_folders_tree(self.user_id)
        self._populate_tree_items(root_item, tree_data)

        self.folder_tree.expandAll()
        # Try to restore previous selection without triggering _on_folder_selected
        if not self._select_tree_item_silent(target_data):
            # Fallback to All
            self._select_tree_item_silent(-1)

    def _select_tree_item_silent(self, folder_data):
        """Select tree item without emitting signal / calling _load_notes. Returns True if found."""
        iterator = QTreeWidgetItemIterator(self.folder_tree)
        while iterator.value():
            item = iterator.value()
            data = item.data(0, Qt.ItemDataRole.UserRole)
            if data == folder_data:
                self.folder_tree.setCurrentItem(item)
                # Update current_folder_id directly without triggering load
                if folder_data == -1:
                    self.current_folder_id = None
                elif folder_data == 0:
                    self.current_folder_id = 0
                else:
                    self.current_folder_id = folder_data
                return True
            iterator += 1
        return False

    def _populate_tree_items(self, parent_item, folders):
        for folder in folders:
            item = QTreeWidgetItem(parent_item)
            item.setText(0, f"{folder['icon']} {folder['name']}")
            item.setData(0, Qt.ItemDataRole.UserRole, folder["id"])
            item.setExpanded(True)
            if folder.get("children"):
                self._populate_tree_items(item, folder["children"])

    def _select_tree_item_by_id(self, folder_id):
        # Legacy method - now uses silent + explicit load
        if self._select_tree_item_silent(folder_id):
            self._load_notes()
            return
        if self.folder_tree.topLevelItemCount() > 0:
            root = self.folder_tree.topLevelItem(0)
            self.folder_tree.setCurrentItem(root)
            self._on_folder_selected(root, 0)

    def _on_folder_selected(self, item, column):
        folder_id = item.data(0, Qt.ItemDataRole.UserRole)
        if folder_id == -1:
            self.current_folder_id = None
        elif folder_id == 0:
            self.current_folder_id = 0
        else:
            self.current_folder_id = folder_id
        self._load_notes()

    def _on_note_selected(self, item):
        note_id = item.data(Qt.ItemDataRole.UserRole)
        if note_id:
            self._load_note(note_id)

    def _load_notes(self):
        self.notes_list.clear()

        folder_id = self.current_folder_id

        # --- Jika sedang search, kita akan ambil notes dengan filter subfolder ---
        if self.search_text:
            # Tentukan daftar folder_id yang diizinkan
            allowed_folder_ids = None  # None berarti semua folder

            if folder_id is None or folder_id == -1:
                # "Semua Catatan" -> semua notes
                all_notes = db.get_notes(self.user_id, None, include_archived=self.show_archived)
                allowed_folder_ids = None  # tidak perlu filter
            elif folder_id == 0:
                # "Tanpa Folder" -> hanya notes dengan folder_id IS NULL
                all_notes = db.get_notes(self.user_id, -1, include_archived=self.show_archived)
                # Karena kita ambil khusus tanpa folder, tidak perlu filter lagi
                allowed_folder_ids = None
            else:
                # Folder spesifik -> ambil semua notes, lalu filter by folder ids termasuk subfolder
                all_notes = db.get_notes(self.user_id, None, include_archived=self.show_archived)
                # Dapatkan semua id folder yang diizinkan (termasuk subfolder)
                allowed_folder_ids = self._get_all_subfolder_ids(folder_id)

            # Filter berdasarkan search text
            if self.search_text:
                all_notes = [n for n in all_notes if self.search_text in n["title"].lower() or self.search_text in n["content"].lower()]

            # Jika allowed_folder_ids tidak None, filter notes yang folder_id-nya ada di list
            if allowed_folder_ids is not None:
                all_notes = [n for n in all_notes if n["folder_id"] in allowed_folder_ids]

        else:
            # --- Tidak search: gunakan logika normal (hanya folder yang dipilih, tanpa subfolder) ---
            if folder_id is None:
                all_notes = db.get_notes(self.user_id, None, include_archived=self.show_archived)
            elif folder_id == 0:
                all_notes = db.get_notes(self.user_id, -1, include_archived=self.show_archived)
            else:
                all_notes = db.get_notes(self.user_id, folder_id, include_archived=self.show_archived)

        if self.search_text:
            all_notes = [n for n in all_notes if self.search_text in n["title"].lower() or self.search_text in n["content"].lower()]

        if not all_notes:
            self.note_count_label.setText(tr("notes_archive_empty") if self.show_archived else tr("notes_count_format", count=0))
            self._clear_editor()
            return

        self.note_count_label.setText(tr("notes_count_format", count=len(all_notes)))
        for n in all_notes:
            item = QListWidgetItem(f"{'📦 ' if n['is_archived'] else ''}{n['title']}")
            item.setData(Qt.ItemDataRole.UserRole, n["id"])
            self.notes_list.addItem(item)

        if self.search_text:
            self.note_count_label.setText(tr("notes_search_result", count=len(all_notes)))

        if all_notes:
            self.notes_list.setCurrentRow(0)
            self._load_note(all_notes[0]["id"])
        else:
            self._clear_editor()
            self.current_note_id = None

    def _get_all_subfolder_ids(self, folder_id):
        """Rekursif ambil semua ID subfolder dari folder_id (termasuk folder_id sendiri)."""
        ids = [folder_id]
        conn = db.get_conn()
        sub = conn.execute("SELECT id FROM note_folders WHERE parent_id=?", (folder_id,)).fetchall()
        conn.close()
        for row in sub:
            ids.extend(self._get_all_subfolder_ids(row["id"]))
        return ids

    def _clear_editor(self):
        if hasattr(self, 'title_edit') and self.title_edit is not None:
            try:
                self.title_edit.clear()
            except RuntimeError:
                pass
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                self.content_edit.clear()
            except RuntimeError:
                pass
        self.current_note_id = None
        self._is_dirty = False

        if hasattr(self, 'zoom_slider') and self.zoom_slider is not None:
            try:
                self.zoom_slider.setValue(100)
            except RuntimeError:
                pass
        if hasattr(self, 'zoom_value_label') and self.zoom_value_label is not None:
            try:
                self.zoom_value_label.setText("100%")
            except RuntimeError:
                pass
        self._current_zoom = 100
        self._base_html = ""
        self._is_scaling = False
        # Reset visual tanpa mengubah base (base sudah kosong)
        try:
            self.content_edit.blockSignals(True)
            self.content_edit.clear()
            self.content_edit.blockSignals(False)
        except:
            pass

    def _scale_html(self, html: str, zoom: int) -> str:
        """Scale semua font-size di HTML untuk tampilan zoom (visual only)."""
        if not html or zoom == 100:
            return html
        try:
            scale = zoom / 100.0
            def repl(m):
                val = float(m.group(1))
                unit = m.group(2)
                new_val = val * scale
                # Format: keep 1 decimal if needed, else int
                if new_val.is_integer():
                    return f"font-size:{int(new_val)}{unit}"
                else:
                    return f"font-size:{new_val:.1f}{unit}"
            # Handle pt, px, %
            html = re.sub(r'font-size\s*:\s*(\d+(?:\.\d+)?)\s*(pt|px|%)', repl, html)
            return html
        except:
            return html

    def _unscale_html(self, html: str, zoom: int) -> str:
        """Kembalikan HTML dari zoom ke 100% untuk save permanen."""
        if not html or zoom == 100:
            return html
        try:
            scale = 100.0 / zoom
            def repl(m):
                val = float(m.group(1))
                unit = m.group(2)
                new_val = val * scale
                if new_val.is_integer():
                    return f"font-size:{int(new_val)}{unit}"
                else:
                    return f"font-size:{new_val:.1f}{unit}"
            html = re.sub(r'font-size\s*:\s*(\d+(?:\.\d+)?)\s*(pt|px|%)', repl, html)
            return html
        except:
            return html

    def _apply_zoom_to_editor(self, zoom_percent):
        """Zoom visual via HTML scaling — tidak mengubah save permanen (FIX font 16px)."""
        if not hasattr(self, 'content_edit') or self.content_edit is None:
            return
        try:
            zoom_percent = int(zoom_percent)
            self._current_zoom = zoom_percent
            if not self._base_html:
                # Jika belum ada base, ambil dari current content (unscaled)
                try:
                    cur = self.content_edit.toHtml()
                    self._base_html = self._unscale_html(cur, self._current_zoom) if hasattr(self, '_current_zoom') else cur
                except:
                    self._base_html = ""
            if self._base_html:
                self._is_scaling = True
                scaled = self._scale_html(self._base_html, zoom_percent)
                # Block signals agar tidak trigger _mark_dirty
                self.content_edit.blockSignals(True)
                self.content_edit.setHtml(scaled)
                self.content_edit.blockSignals(False)
                self._is_scaling = False
            # Update label
            if hasattr(self, 'zoom_value_label') and self.zoom_value_label is not None:
                try:
                    self.zoom_value_label.setText(f"{zoom_percent}%")
                except:
                    pass
        except Exception as e:
            self._is_scaling = False
            pass

    def _load_note(self, note_id):
        note = db.get_note(note_id, self.user_id)
        if not note:
            return

        if hasattr(self, 'title_edit') and self.title_edit is not None:
            try:
                self.title_edit.setText(note["title"])
            except RuntimeError:
                pass
        # Simpan base HTML (100%) untuk save permanen
        base = note["content"] or ""
        self._base_html = base
        self._current_zoom = 100
        if hasattr(self, 'zoom_slider') and self.zoom_slider is not None:
            try:
                self.zoom_slider.blockSignals(True)
                self.zoom_slider.setValue(100)
                self.zoom_slider.blockSignals(False)
            except RuntimeError:
                pass
        if hasattr(self, 'zoom_value_label') and self.zoom_value_label is not None:
            try:
                self.zoom_value_label.setText("100%")
            except RuntimeError:
                pass
        # Tampilkan scaled version sesuai zoom saat ini (100% = base)
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                self._is_scaling = True
                self.content_edit.blockSignals(True)
                scaled = self._scale_html(base, self._current_zoom)
                self.content_edit.setHtml(scaled)
                self.content_edit.blockSignals(False)
                self._is_scaling = False
            except RuntimeError:
                self._is_scaling = False
                pass
        self.current_note_id = note_id
        self._is_dirty = False

        # Ambil format
        cursor = None
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                cursor = self.content_edit.textCursor()
                cursor.movePosition(QTextCursor.MoveOperation.Start)
                fmt = cursor.charFormat()
            except RuntimeError:
                fmt = None
        else:
            fmt = None

        if note_id in self._note_font_settings:
            size, color = self._note_font_settings[note_id]
            self._current_font_size = size
            self._current_color = color
        else:
            if fmt is not None:
                if fmt.fontPointSize() > 0:
                    size = fmt.fontPointSize()
                else:
                    size = self._default_font_size
                if fmt.foreground().color().isValid():
                    color = fmt.foreground().color()
                else:
                    color = self._default_color
            else:
                size = self._default_font_size
                color = self._default_color
            self._note_font_settings[note_id] = (size, color)
            self._current_font_size = size
            self._current_color = color

        if fmt is not None:
            self._current_bold = (fmt.fontWeight() == QFont.Weight.Bold)
            self._current_italic = fmt.fontItalic()
            self._current_underline = fmt.fontUnderline()
        else:
            self._current_bold = False
            self._current_italic = False
            self._current_underline = False

        # Update toolbar UI
        if hasattr(self, 'font_size_combo') and self.font_size_combo is not None:
            try:
                index = self.font_size_combo.findData(self._current_font_size)
                if index >= 0:
                    self.font_size_combo.setCurrentIndex(index)
            except RuntimeError:
                pass

        if hasattr(self, 'color_indicator') and self.color_indicator is not None:
            try:
                self.color_indicator.setStyleSheet(f"""
                    background: {self._current_color.name()};
                    border: 1px solid {_T('border')};
                    border-radius: 4px;
                """)
            except RuntimeError:
                pass

        for btn, state in [(self.bold_btn, self._current_bold),
                           (self.italic_btn, self._current_italic),
                           (self.underline_btn, self._current_underline)]:
            if btn is not None:
                try:
                    btn.setChecked(state)
                except RuntimeError:
                    pass

        default_fmt = QTextCharFormat()
        default_fmt.setFontPointSize(self._current_font_size)
        default_fmt.setForeground(self._current_color)
        default_fmt.setFontWeight(QFont.Weight.Bold if self._current_bold else QFont.Weight.Normal)
        default_fmt.setFontItalic(self._current_italic)
        default_fmt.setFontUnderline(self._current_underline)
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                self.content_edit.setCurrentCharFormat(default_fmt)
            except RuntimeError:
                pass

        if cursor is not None and hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                cursor.movePosition(QTextCursor.MoveOperation.Start)
                self.content_edit.setTextCursor(cursor)
            except RuntimeError:
                pass

    def _on_content_changed(self):
        self._mark_dirty()

    def _mark_dirty(self):
        self._is_dirty = True

    # ========== FOLDER CONTEXT MENU ==========
    def _show_folder_context_menu(self, position):
        item = self.folder_tree.itemAt(position)
        if not item:
            return
        menu = QMenu(self)
        folder_id = item.data(0, Qt.ItemDataRole.UserRole)

        if folder_id is not None and folder_id > 0:
            add_sub_act = menu.addAction(tr("notes_add_subfolder"))
            edit_name_act = menu.addAction(tr("notes_edit_folder"))
            edit_icon_act = menu.addAction(tr("notes_edit_icon"))
            duplicate_act = menu.addAction(tr("notes_duplicate_folder"))
            menu.addSeparator()
            delete_act = menu.addAction(tr("notes_delete"))
            menu.addSeparator()
            child_count = self._count_children(folder_id)
            info_act = menu.addAction(tr("notes_folder_info", count=child_count))
            info_act.setEnabled(False)

            action = menu.exec(self.folder_tree.viewport().mapToGlobal(position))
            if action == add_sub_act:
                self._add_subfolder(folder_id)
            elif action == edit_name_act:
                self._edit_folder_name(folder_id)
            elif action == edit_icon_act:
                self._edit_folder_icon(folder_id)
            elif action == duplicate_act:
                self._duplicate_folder(folder_id)
            elif action == delete_act:
                self._delete_folder_item(folder_id)

        elif folder_id == 0 or folder_id == -1:
            add_act = menu.addAction(tr("notes_add_folder"))
            if menu.exec(self.folder_tree.viewport().mapToGlobal(position)) == add_act:
                self._add_folder()

    # ========== FOLDER OPERATIONS ==========
    def _add_folder(self, parent_id=None):
        name, ok = QInputDialog.getText(self, tr("notes_folder_name"), tr("notes_folder_name_ph"))
        if ok and name.strip():
            r = db.add_note_folder(self.user_id, name.strip(), parent_id=parent_id)
            if r["ok"]:
                SND.notify()
                cur = self.folder_tree.currentItem()
                pd = cur.data(0, Qt.ItemDataRole.UserRole) if cur else -1
                self.folder_tree.blockSignals(True)
                self._load_folder_tree(preserve_id=pd)
                self.folder_tree.blockSignals(False)
                self._load_notes()

    def _add_subfolder(self, parent_id):
        name, ok = QInputDialog.getText(self, tr("notes_subfolder_title"), tr("notes_subfolder_name"))
        if ok and name.strip():
            r = db.add_note_folder(self.user_id, name.strip(), parent_id=parent_id)
            if r["ok"]:
                SND.notify()
                cur = self.folder_tree.currentItem()
                pd = cur.data(0, Qt.ItemDataRole.UserRole) if cur else -1
                self.folder_tree.blockSignals(True)
                self._load_folder_tree(preserve_id=pd)
                self.folder_tree.blockSignals(False)
                self._load_notes()
                self._select_tree_item_by_id(r["folder_id"])
                _show(self, tr("berhasil_title"), tr("notes_subfolder_added", name=name.strip()), "success")
            else:
                SND.error()
                _show(self, tr("gagal_title"), r.get("msg", "Gagal menambah subfolder"), "error")

    def _edit_folder_name(self, folder_id):
        current = db.get_conn().execute(
            "SELECT name FROM note_folders WHERE id=? AND user_id=?",
            (folder_id, self.user_id)
        ).fetchone()
        if not current:
            return
        old_name = current["name"]
        new_name, ok = QInputDialog.getText(self, tr("notes_edit_folder_title"), tr("notes_folder_name_ph"), text=old_name)
        if ok and new_name.strip():
            conn = db.get_conn()
            conn.execute("UPDATE note_folders SET name=? WHERE id=? AND user_id=?", (new_name.strip(), folder_id, self.user_id))
            conn.commit()
            conn.close()
            SND.notify()
            pd = self.current_folder_id
            if pd is None:
                pd=-1
            elif pd==0:
                pd=0
            self.folder_tree.blockSignals(True)
            self._load_folder_tree(preserve_id=pd)
            self.folder_tree.blockSignals(False)
            self._load_notes()
            _show(self, tr("berhasil_title"), tr("notes_folder_renamed"), "success")

    def _edit_folder_icon(self, folder_id):
        icons = ["📁", "📂", "⭐", "❤️", "🔥", "💼", "🏠", "🎯", "📖", "🎮", "💡", "🌿", "📚", "🏷️", "📎", "📌"]
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("notes_edit_icon_title"))
        dlg.setMinimumWidth(400)
        dlg.setStyleSheet(build_ss())
        layout = QVBoxLayout(dlg)
        layout.addWidget(_lbl(tr("notes_select_icon"), size=12))
        grid = QGridLayout()
        row, col = 0, 0
        for icon in icons:
            btn = QPushButton(icon)
            btn.setFixedSize(50, 50)
            btn.setStyleSheet("font-size: 24px;")
            btn.clicked.connect(lambda checked, i=icon: self._save_folder_icon(folder_id, i, dlg))
            grid.addWidget(btn, row, col)
            col += 1
            if col > 5:
                col = 0
                row += 1
        layout.addLayout(grid)
        cancel_btn = _btn(tr("btn_cancel"), "flat", dlg.reject)
        layout.addWidget(cancel_btn)
        dlg.exec()

    def _save_folder_icon(self, folder_id, new_icon, dialog):
        db.update_note_folder_icon(folder_id, self.user_id, new_icon)
        SND.notify()
        dialog.accept()
        self._load_folder_tree()
        _show(self, tr("berhasil_title"), tr("notes_icon_updated"), "success")

    def _duplicate_folder(self, folder_id):
        r = db.duplicate_note_folder(self.user_id, folder_id)
        if r["ok"]:
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            self._load_folder_tree()
            if "new_folder_id" in r:
                self._select_tree_item_by_id(r["new_folder_id"])
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")

    def _delete_folder_item(self, folder_id):
        reply = QMessageBox.question(self, tr("confirm_title"), tr("notes_folder_delete_confirm"),
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            conn = db.get_conn()
            conn.execute("DELETE FROM note_folders WHERE id=? AND user_id=?", (folder_id, self.user_id))
            conn.commit()
            conn.close()
            SND.click()
            self._load_folder_tree()

    def _count_children(self, folder_id):
        conn = db.get_conn()
        count = conn.execute(
            "SELECT COUNT(*) FROM notes WHERE user_id=? AND folder_id=?",
            (self.user_id, folder_id)
        ).fetchone()[0]
        subfolders = conn.execute(
            "SELECT id FROM note_folders WHERE user_id=? AND parent_id=?",
            (self.user_id, folder_id)
        ).fetchall()
        conn.close()
        for sub in subfolders:
            count += self._count_children(sub["id"])
        return count

    # ========== NOTES OPERATIONS ==========
    def _add_note(self):
        current_item = self.folder_tree.currentItem()
        if current_item:
            folder_id = current_item.data(0, Qt.ItemDataRole.UserRole)
            if folder_id == -1:
                folder_id = None
            elif folder_id == 0:
                folder_id = None
            # Remember current folder data for preserve
            preserve_data = current_item.data(0, Qt.ItemDataRole.UserRole)
        else:
            folder_id = None
            preserve_data = -1

        title = "Untitled"
        r = db.add_note(self.user_id, folder_id, title)
        if r["ok"]:
            SND.complete()
            # Preserve folder selection - don't jump to All (FIX 2)
            self.folder_tree.blockSignals(True)
            self._load_folder_tree(preserve_id=preserve_data)
            self.folder_tree.blockSignals(False)
            self._load_notes()
            # Select the newly created note
            for i in range(self.notes_list.count()):
                it = self.notes_list.item(i)
                if it.data(Qt.ItemDataRole.UserRole) == r["note_id"]:
                    self.notes_list.setCurrentItem(it)
                    self._load_note(r["note_id"])
                    break

    def _save_note(self):
        if self.current_note_id is None:
            return
        title = ""
        if hasattr(self, 'title_edit') and self.title_edit is not None:
            try:
                title = self.title_edit.text().strip()
            except RuntimeError:
                pass
        if not title:
            title = "Untitled"
        content = ""
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                # Ambil HTML display (mungkin scaled) lalu kembalikan ke 100% untuk save permanen
                scaled_html = self.content_edit.toHtml()
                content = self._unscale_html(scaled_html, self._current_zoom)
                # Update base agar konsisten
                self._base_html = content
            except RuntimeError:
                content = self._base_html
        db.update_note(self.current_note_id, self.user_id, title=title, content=content)
        self._is_dirty = False
        SND.complete()
        _show(self, tr("berhasil_title"), tr("notes_saved"), "success")
        # FIX 2: Preserve folder and current note selection instead of jumping to All
        saved_folder = self.current_folder_id
        saved_note = self.current_note_id
        self._load_notes()
        # Re-select saved note if still visible
        for i in range(self.notes_list.count()):
            it = self.notes_list.item(i)
            if it.data(Qt.ItemDataRole.UserRole) == saved_note:
                self.notes_list.setCurrentItem(it)
                break
        # Ensure folder tree still shows correct folder
        if saved_folder is None:
            self._select_tree_item_silent(-1)
        elif saved_folder == 0:
            self._select_tree_item_silent(0)
        else:
            self._select_tree_item_silent(saved_folder)

    def _duplicate_note(self):
        """FIX 3: Duplicate note to chosen folder/subfolder"""
        if self.current_note_id is None:
            _show(self, tr("msg_error"), tr("notes_select_note_first"), "error")
            return
        note = db.get_note(self.current_note_id, self.user_id)
        if not note:
            return
        # Build folder list for dialog: All folders including hierarchy
        folders = db.get_note_folders(self.user_id)
        # Build display names with hierarchy
        # Use tree to get full paths
        from collections import defaultdict
        # Simple flat list with indentation via parent
        tree = db.get_note_folders_tree(self.user_id)
        flat = []
        def flatten(nodes, depth=0):
            for n in nodes:
                prefix = "  " * depth + ("└ " if depth else "")
                flat.append((f"{n['icon']} {prefix}{n['name']}", n['id']))
                if n.get("children"):
                    flatten(n["children"], depth+1)
        flatten(tree)
        # Add "No Folder" and "Current" options
        options = [(tr("notes_no_folder"), 0)] + flat
        # Dialog to choose destination
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("notes_duplicate_title"))
        dlg.setMinimumWidth(420)
        dlg.setStyleSheet(build_ss())
        lay = QVBoxLayout(dlg)
        lay.setSpacing(12)
        lay.addWidget(_lbl(tr("notes_duplicate_choose_folder"), size=12))
        combo = QComboBox()
        for name, fid in options:
            combo.addItem(name, fid)
        # Preselect current folder
        cur_fid = note.get("folder_id")
        if cur_fid is None:
            idx = 0
        else:
            idx = next((i for i,(n,f) in enumerate(options) if f==cur_fid), 0)
        combo.setCurrentIndex(idx)
        lay.addWidget(combo)
        btn_row = QHBoxLayout()
        btn_row.addStretch()
        btn_row.addWidget(_btn(tr("msg_cancel"), "", dlg.reject))
        btn_row.addWidget(_btn(tr("notes_duplicate_btn"), "solid", dlg.accept))
        lay.addLayout(btn_row)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        dest_id = combo.currentData()
        dest_folder = None if dest_id == 0 else dest_id
        r = db.duplicate_note(self.user_id, self.current_note_id, dest_folder)
        if r.get("ok"):
            SND.complete()
            _show(self, tr("berhasil_title"), r.get("msg", "Catatan diduplikat!"), "success")
            # Preserve and refresh
            preserve = dest_id if dest_id != 0 else (-1 if dest_folder is None else dest_id)
            # Actually preserve destination folder view
            self.folder_tree.blockSignals(True)
            self._load_folder_tree(preserve_id=dest_id if dest_id!=0 else -1)
            self.folder_tree.blockSignals(False)
            self._load_notes()
            # Select new note
            new_id = r.get("new_note_id")
            if new_id:
                for i in range(self.notes_list.count()):
                    it = self.notes_list.item(i)
                    if it.data(Qt.ItemDataRole.UserRole) == new_id:
                        self.notes_list.setCurrentItem(it)
                        self._load_note(new_id)
                        break
        else:
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg","Gagal duplikat"), "error")

    def _delete_selected(self):
        if self.current_note_id is None:
            return
        reply = QMessageBox.question(self, tr("confirm_title"), tr("notes_delete_confirm"),
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            note = db.get_note(self.current_note_id, self.user_id)
            nm = (note or {}).get("title", "") or "Untitled"
            r = db.delete_note(self.current_note_id, self.user_id)
            SND.click()
            self._clear_editor()
            self._load_notes()
            # Tawarkan undo via toast; fallback ke dialog klasik
            tid = r.get("trash_id") if isinstance(r, dict) else None
            mw = self.window()
            if tid and hasattr(mw, "show_undo_toast"):
                mw.show_undo_toast(nm, tid)
            else:
                _show(self, tr("berhasil_title"), tr("notes_deleted"), "success")

    def _toggle_archive(self):
        if self.current_note_id is None:
            return
        note = db.get_note(self.current_note_id, self.user_id)
        if not note:
            return
        new_state = 0 if note["is_archived"] else 1
        db.archive_note(self.current_note_id, self.user_id, new_state)
        SND.notify()
        if new_state:
            _show(self, tr("info_title"), tr("notes_archived"), "info")
        else:
            _show(self, tr("info_title"), tr("notes_unarchived"), "info")
        self._clear_editor()
        self._load_notes()

    def _toggle_show_archived(self):
        self.show_archived = not self.show_archived
        if self.show_archived:
            self.archive_toggle_btn.setText(tr("notes_hide_archived"))
        else:
            self.archive_toggle_btn.setText(tr("notes_show_archived"))
        self._load_notes()

    def _on_search(self, text):
        self.search_text = text.strip().lower()
        self._load_notes()

    # ========== ZOOM (visual only, tidak mengubah save) ==========
    def _on_zoom_changed(self, value):
        if hasattr(self, 'zoom_value_label') and self.zoom_value_label is not None:
            try:
                self.zoom_value_label.setText(f"{value}%")
            except RuntimeError:
                pass
        # Simpan base terbaru dari display saat ini sebelum ganti zoom (agar ketikan terbaru tidak hilang)
        try:
            if not self._is_scaling and self._base_html is not None:
                cur_html = self.content_edit.toHtml()
                # Unscale current display ke base (100%)
                self._base_html = self._unscale_html(cur_html, self._current_zoom)
        except:
            pass
        self._current_zoom = int(value)
        # Scale dari base_html ke display
        if self._base_html:
            self._is_scaling = True
            try:
                scaled = self._scale_html(self._base_html, self._current_zoom)
                self.content_edit.blockSignals(True)
                self.content_edit.setHtml(scaled)
                self.content_edit.blockSignals(False)
            except:
                pass
            self._is_scaling = False
        # Tidak menyimpan ke database (zoom tidak permanen)

    # ========== FONT TOOLBAR ==========
    def _change_font_size(self, index):
        if not hasattr(self, 'font_size_combo') or self.font_size_combo is None:
            return
        try:
            size = self.font_size_combo.currentData()
            if not size:
                return
        except RuntimeError:
            return

        cursor = None
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                cursor = self.content_edit.textCursor()
            except RuntimeError:
                return
        if cursor is None:
            return

        # Jika sedang zoom, sesuaikan display size agar save tetap 16pt
        display_size = size * self._current_zoom / 100 if hasattr(self, '_current_zoom') else size
        fmt = QTextCharFormat()
        fmt.setFontPointSize(display_size)
        fmt.setForeground(self._current_color)
        if cursor.hasSelection():
            cursor.mergeCharFormat(fmt)
        else:
            try:
                self.content_edit.setCurrentCharFormat(fmt)
            except RuntimeError:
                return

        # Simpan base size (100%) untuk save
        if self.current_note_id:
            self._note_font_settings[self.current_note_id] = (size, self._current_color)
        self._current_font_size = size
        # Update base_html agar konsisten (ambil current display lalu unscale)
        try:
            if not self._is_scaling:
                cur_html = self.content_edit.toHtml()
                self._base_html = self._unscale_html(cur_html, self._current_zoom)
        except:
            pass
        try:
            self.content_edit.setFocus()
        except RuntimeError:
            pass

    def _choose_font_color(self):
        if not hasattr(self, 'color_indicator') or self.color_indicator is None:
            return
        try:
            current_color = self.color_indicator.palette().window().color()
        except RuntimeError:
            return

        color = QColorDialog.getColor(current_color, self, "Pilih Warna Teks")
        if color.isValid():
            if hasattr(self, 'color_indicator') and self.color_indicator is not None:
                try:
                    self.color_indicator.setStyleSheet(f"""
                        background: {color.name()};
                        border: 1px solid {_T('border')};
                        border-radius: 4px;
                    """)
                except RuntimeError:
                    pass

            cursor = None
            if hasattr(self, 'content_edit') and self.content_edit is not None:
                try:
                    cursor = self.content_edit.textCursor()
                except RuntimeError:
                    pass
            if cursor is None:
                return

            display_size = self._current_font_size * self._current_zoom / 100 if hasattr(self, '_current_zoom') else self._current_font_size
            fmt = QTextCharFormat()
            fmt.setFontPointSize(display_size)
            fmt.setForeground(color)
            if cursor.hasSelection():
                cursor.mergeCharFormat(fmt)
            else:
                try:
                    self.content_edit.setCurrentCharFormat(fmt)
                except RuntimeError:
                    pass

            if self.current_note_id:
                self._note_font_settings[self.current_note_id] = (self._current_font_size, color)
            self._current_color = color
            # Update base agar save konsisten
            try:
                if not self._is_scaling:
                    cur_html = self.content_edit.toHtml()
                    self._base_html = self._unscale_html(cur_html, self._current_zoom)
            except:
                pass
            try:
                self.content_edit.setFocus()
            except RuntimeError:
                pass

    def _toggle_format(self, format_type):
        cursor = None
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                cursor = self.content_edit.textCursor()
            except RuntimeError:
                return
        if cursor is None:
            return

        fmt = QTextCharFormat()
        if format_type == 'bold':
            current = self._current_bold
            new_state = not current
            fmt.setFontWeight(QFont.Weight.Bold if new_state else QFont.Weight.Normal)
            self._current_bold = new_state
            if hasattr(self, 'bold_btn') and self.bold_btn is not None:
                try:
                    self.bold_btn.setChecked(new_state)
                except RuntimeError:
                    pass
        elif format_type == 'italic':
            current = self._current_italic
            new_state = not current
            fmt.setFontItalic(new_state)
            self._current_italic = new_state
            if hasattr(self, 'italic_btn') and self.italic_btn is not None:
                try:
                    self.italic_btn.setChecked(new_state)
                except RuntimeError:
                    pass
        elif format_type == 'underline':
            current = self._current_underline
            new_state = not current
            fmt.setFontUnderline(new_state)
            self._current_underline = new_state
            if hasattr(self, 'underline_btn') and self.underline_btn is not None:
                try:
                    self.underline_btn.setChecked(new_state)
                except RuntimeError:
                    pass
        else:
            return

        display_size = self._current_font_size * self._current_zoom / 100 if hasattr(self, '_current_zoom') else self._current_font_size
        fmt.setFontPointSize(display_size)
        fmt.setForeground(self._current_color)

        if cursor.hasSelection():
            cursor.mergeCharFormat(fmt)
        else:
            try:
                self.content_edit.setCurrentCharFormat(fmt)
            except RuntimeError:
                pass
        # Update base for save
        try:
            if not self._is_scaling:
                cur_html = self.content_edit.toHtml()
                self._base_html = self._unscale_html(cur_html, self._current_zoom)
        except:
            pass
        try:
            self.content_edit.setFocus()
        except RuntimeError:
            pass

    def _apply_default_format(self):
        if not self.current_note_id:
            return
        if self.current_note_id in self._note_font_settings:
            size, color = self._note_font_settings[self.current_note_id]
        else:
            size = self._default_font_size
            color = self._default_color

        cursor = None
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                cursor = self.content_edit.textCursor()
            except RuntimeError:
                return
        if cursor is None:
            return

        if not cursor.hasSelection():
            fmt = QTextCharFormat()
            display_size = size * self._current_zoom / 100 if hasattr(self, '_current_zoom') else size
            fmt.setFontPointSize(display_size)
            fmt.setForeground(color)
            fmt.setFontWeight(QFont.Weight.Bold if self._current_bold else QFont.Weight.Normal)
            fmt.setFontItalic(self._current_italic)
            fmt.setFontUnderline(self._current_underline)
            try:
                self.content_edit.setCurrentCharFormat(fmt)
            except RuntimeError:
                pass

    # ========== SYMBOL / MATH INSERT ==========
    def _insert_symbol(self, symbol):
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                self.content_edit.insertPlainText(symbol)
                self.content_edit.setFocus()
            except RuntimeError:
                pass

    def _insert_superscript(self):
        cursor = None
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                cursor = self.content_edit.textCursor()
            except RuntimeError:
                return
        if cursor is None:
            return
        if cursor.hasSelection():
            selected = cursor.selectedText()
            html = f"<sup>{selected}</sup>"
            cursor.removeSelectedText()
            try:
                self.content_edit.insertHtml(html)
            except RuntimeError:
                pass
        else:
            try:
                cursor.insertText("x²")
            except RuntimeError:
                pass
        try:
            self.content_edit.setFocus()
        except RuntimeError:
            pass

    def _insert_subscript(self):
        cursor = None
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                cursor = self.content_edit.textCursor()
            except RuntimeError:
                return
        if cursor is None:
            return
        if cursor.hasSelection():
            selected = cursor.selectedText()
            html = f"<sub>{selected}</sub>"
            cursor.removeSelectedText()
            try:
                self.content_edit.insertHtml(html)
            except RuntimeError:
                pass
        else:
            try:
                cursor.insertText("x₂")
            except RuntimeError:
                pass
        try:
            self.content_edit.setFocus()
        except RuntimeError:
            pass

    def _latex_menu(self):
        """Menu tombol ∑ LaTeX: konversi seleksi/semua + pratinjau render."""
        menu = QMenu(self)
        menu.addAction(tr("notes_math_convert_sel"), self._latex_convert_selection)
        menu.addAction(tr("notes_math_convert_all"), self._latex_convert_all)
        menu.addSeparator()
        menu.addAction(tr("notes_math_preview"), self._latex_preview)
        btn = self._latex_btn
        menu.exec(btn.mapToGlobal(btn.rect().bottomLeft()))

    def _latex_convert_selection(self):
        cursor = self.content_edit.textCursor()
        sel = cursor.selectedText()
        if not sel or not mathtools.has_latex(sel):
            SND.error()
            _show(self, tr("notes_math_preview_title"), tr("notes_math_none"), "info")
            return
        cursor.insertText(mathtools.latex_to_unicode(sel))
        SND.complete()
        _show(self, tr("notes_math_preview_title"), tr("notes_math_converted"), "success")

    def _latex_convert_all(self):
        full = self.content_edit.toPlainText()
        if not mathtools.has_latex(full):
            SND.error()
            _show(self, tr("notes_math_preview_title"), tr("notes_math_none"), "info")
            return
        pos = self.content_edit.verticalScrollBar().value()
        self.content_edit.setPlainText(mathtools.latex_to_unicode(full))
        self.content_edit.verticalScrollBar().setValue(pos)
        SND.complete()
        _show(self, tr("notes_math_preview_title"), tr("notes_math_converted"), "success")

    def _latex_preview(self):
        chunks = mathtools.find_math_chunks(self.content_edit.toPlainText())
        MathPreviewDialog(chunks, self).exec()

    def _insert_fraction(self):
        cursor = None
        if hasattr(self, 'content_edit') and self.content_edit is not None:
            try:
                cursor = self.content_edit.textCursor()
            except RuntimeError:
                return
        if cursor is None:
            return

        # Use display size for visual, base size is stored
        base_size = self._current_font_size
        display_size = base_size * self._current_zoom / 100 if hasattr(self, '_current_zoom') else base_size
        size = int(display_size)
        color = self._current_color.name()

        if cursor.hasSelection():
            numerator = cursor.selectedText()
            cursor.removeSelectedText()
            denominator = "b"
            self._insert_fraction_html(numerator, denominator, size, color, cursor)
            return

        cursor.movePosition(QTextCursor.MoveOperation.Left, QTextCursor.MoveMode.KeepAnchor, 20)
        text_left = cursor.selectedText()
        cursor.clearSelection()
        cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.MoveAnchor, 20)
        import re
        match = re.search(r'(\d+)\s*/\s*(\d+)$', text_left)
        if match:
            numerator = match.group(1)
            denominator = match.group(2)
            start_pos = len(text_left) - len(match.group(0))
            cursor.movePosition(QTextCursor.MoveOperation.Left, QTextCursor.MoveMode.KeepAnchor, len(match.group(0)))
            cursor.removeSelectedText()
            self._insert_fraction_html(numerator, denominator, size, color, cursor)
            return

        self._insert_fraction_html("a", "b", size, color, cursor)

    def _insert_fraction_html(self, numerator, denominator, size, color, cursor):
        html = (
            '<span style="font-size:{}px;color:{};">'
            '<sup>{}</sup>/<sub>{}</sub>'
            '</span>'
        ).format(size, color, numerator, denominator)
        try:
            cursor.insertHtml(html)
            cursor.insertText(" ")
        except RuntimeError:
            return

        fmt = QTextCharFormat()
        fmt.setFontPointSize(size)
        fmt.setForeground(QColor(color))
        try:
            self.content_edit.setCurrentCharFormat(fmt)
            self.content_edit.setFocus()
        except RuntimeError:
            pass

    # ========== CLOSE ==========
    def _toggle_left_panel(self, collapsed):
        # Cari left widget (index 0 di splitter)
        try:
            left = self.findChild(QWidget, "learning_left")
            if not left:
                # Fallback: ambil widget pertama di splitter
                left = self.sender().parent().findChild(QSplitter).widget(0) if hasattr(self, 'sender') else None
            # Simpan sizes
            if not hasattr(self, '_splitter'):
                # Cari splitter
                for child in self.findChildren(QSplitter):
                    self._splitter = child
                    break
            if hasattr(self, '_splitter'):
                sizes = self._splitter.sizes()
                if collapsed:
                    self._splitter.setSizes([0, 600, 320])
                    self.btn_collapse_left.setText("Sources ▶")
                else:
                    self._splitter.setSizes([280, 400, 320])
                    self.btn_collapse_left.setText("◀ Sources")
        except Exception as e:
            print(f"Toggle left failed: {e}")

    def _toggle_right_panel(self, collapsed):
        try:
            if not hasattr(self, '_splitter'):
                for child in self.findChildren(QSplitter):
                    self._splitter = child
                    break
            if hasattr(self, '_splitter'):
                if collapsed:
                    self._splitter.setSizes([280, 400, 0])
                    self.btn_collapse_right.setText("◀ Studio")
                else:
                    self._splitter.setSizes([280, 400, 320])
                    self.btn_collapse_right.setText("Studio ▶")
        except Exception as e:
            print(f"Toggle right failed: {e}")

    def closeEvent(self, e):
        AppState.unregister(self.load)
        AppState.unregister_lang_cb(self.load)
        super().closeEvent(e)

# ══════════════════════════════════════════════════════════════════════════════
#  Reminders Page 
# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
#  📚 LEARNING PAGE — NotebookLM ala CraftLife
# ══════════════════════════════════════════════════════════════════════════════
class LearningPage(QWidget):
    """NotebookLM clone: Sources + Chat + Studio"""
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self.current_notebook_id = None
        self._chat_history = []
        self._build()
        AppState.register(self.load)
        AppState.register_lang_cb(self.load)
        self.load()

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(20, 16, 20, 16)
        root.setSpacing(10)

        # Header
        hdr = QHBoxLayout()
        hdr.addWidget(_lbl(tr("learning_title"), "section", 15, True))
        hdr.addWidget(_lbl(tr("learning_subtitle"), "sub", 11))
        hdr.addStretch()
        # API status
        self.api_status = _lbl("", "sub", 10)
        hdr.addWidget(self.api_status)
        api_btn = _btn("🔑 API Key", h=32)
        api_btn.clicked.connect(self._manage_api_key)
        hdr.addWidget(api_btn)
        root.addLayout(hdr)
        root.addWidget(_sep())

        # Notebook bar
        nb_bar = QHBoxLayout()
        nb_bar.addWidget(_lbl(tr("learning_notebook"), size=12))
        self.nb_combo = QComboBox()
        self.nb_combo.setMinimumWidth(200)
        self.nb_combo.currentIndexChanged.connect(self._on_notebook_changed)
        nb_bar.addWidget(self.nb_combo, 1)
        nb_bar.addWidget(_btn(tr("learning_new_notebook"), h=32, slot=self._create_notebook))
        nb_bar.addWidget(_btn(tr("learning_rename_notebook"), h=32, slot=self._rename_notebook))
        nb_bar.addWidget(_btn(tr("learning_delete_notebook"), "danger", h=32, slot=self._delete_notebook))
        root.addLayout(nb_bar)

        # Collapse buttons row
        collapse_row = QHBoxLayout()
        collapse_row.setContentsMargins(0,0,0,0)
        self.btn_collapse_left = QToolButton()
        self.btn_collapse_left.setText("◀ Sources")
        self.btn_collapse_left.setCheckable(True)
        self.btn_collapse_left.setChecked(False)
        self.btn_collapse_left.setToolTip("Collapse/Expand Sources")
        self.btn_collapse_left.toggled.connect(self._toggle_left_panel)
        self.btn_collapse_right = QToolButton()
        self.btn_collapse_right.setText("Studio ▶")
        self.btn_collapse_right.setCheckable(True)
        self.btn_collapse_right.setChecked(False)
        self.btn_collapse_right.setToolTip("Collapse/Expand Studio")
        self.btn_collapse_right.toggled.connect(self._toggle_right_panel)
        collapse_row.addWidget(self.btn_collapse_left)
        collapse_row.addStretch()
        collapse_row.addWidget(QLabel("📚 Learning Studio"))
        collapse_row.addStretch()
        collapse_row.addWidget(self.btn_collapse_right)
        root.addLayout(collapse_row)

        # Main splitter 3 kolom (collapsible)
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setChildrenCollapsible(False)

        # === KIRI: Sources ===
        left_w = QWidget()
        left_w.setObjectName("learning_left")
        left_lay = QVBoxLayout(left_w)
        left_lay.setContentsMargins(0,0,0,0)
        left_lay.setSpacing(6)
        left_lay.addWidget(_lbl(tr("learning_sources"), size=12, bold=True))
        src_btn_row = QHBoxLayout()
        add_src_btn = _btn("＋ Add Source", "solid", h=32)
        add_src_btn.clicked.connect(self._add_source_menu)
        src_btn_row.addWidget(add_src_btn)
        left_lay.addLayout(src_btn_row)
        self.src_list = QListWidget()
        self.src_list.setMinimumWidth(220)
        self.src_list.itemDoubleClicked.connect(self._view_source)
        self.src_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.src_list.customContextMenuRequested.connect(self._src_context_menu)
        left_lay.addWidget(self.src_list)
        del_src_btn = _btn(tr("learning_delete_source"), "danger", h=32, slot=self._delete_source)
        left_lay.addWidget(del_src_btn)
        self.src_count_lbl = _lbl("", "sub", 10)
        left_lay.addWidget(self.src_count_lbl)
        splitter.addWidget(left_w)

        # === TENGAH: Chat ===
        mid_w = QWidget()
        mid_lay = QVBoxLayout(mid_w)
        mid_lay.setContentsMargins(0,0,0,0)
        mid_lay.setSpacing(6)
        mid_lay.addWidget(_lbl(tr("learning_chat_title"), size=12, bold=True))
        self.chat_area = QTextEdit()
        self.chat_area.setReadOnly(True)
        self.chat_area.setPlaceholderText("Tanya sesuatu tentang sources... (contoh: Ringkas semua PDF)")
        self.chat_area.setMinimumHeight(240)
        self.chat_area.setStyleSheet(f"QTextEdit {{ background: {_T('bg')}; border: 1px solid {_T('border')}; border-radius: 10px; padding: 6px; }}")
        mid_lay.addWidget(self.chat_area)
        # Citations
        self.citation_lbl = _lbl("", "sub", 10)
        self.citation_lbl.setWordWrap(True)
        mid_lay.addWidget(self.citation_lbl)
        chat_input_row = QHBoxLayout()
        self.chat_input = QLineEdit()
        self.chat_input.setPlaceholderText("Ketik pertanyaan... Enter untuk kirim")
        self.chat_input.returnPressed.connect(self._send_chat)
        chat_input_row.addWidget(self.chat_input)
        send_btn = _btn("📤 Kirim", "solid", h=36, slot=self._send_chat)
        chat_input_row.addWidget(send_btn)
        mid_lay.addLayout(chat_input_row)
        clear_chat_btn = _btn("🧹 Clear Chat", h=32, slot=self._clear_chat)
        mid_lay.addWidget(clear_chat_btn)
        splitter.addWidget(mid_w)

        # === KANAN: Studio ===
        right_w = QWidget()
        right_lay = QVBoxLayout(right_w)
        right_lay.setContentsMargins(0,0,0,0)
        right_lay.setSpacing(6)
        right_lay.addWidget(_lbl(tr("learning_studio"), size=12, bold=True))
        # Studio buttons grid
        grid = QGridLayout()
        grid.setSpacing(6)
        self.studio_btns = {}
        studio_defs = [
            ("audio_overview", tr("learning_studio_audio"), 0,0),
            ("mind_map", tr("learning_studio_mindmap"), 0,1),
            ("study_guide", tr("learning_studio_guide"), 1,0),
            ("briefing", tr("learning_studio_briefing"), 1,1),
            ("faq", tr("learning_studio_faq"), 2,0),
            ("timeline", tr("learning_studio_timeline"), 2,1),
            ("flashcards", tr("learning_studio_flashcards"), 3,0),
            ("summary", tr("learning_studio_summary"), 3,1),
        ]
        for key, label, r,c in studio_defs:
            btn = _btn(label, h=36)
            btn.clicked.connect(lambda _, k=key: self._generate_studio(k))
            grid.addWidget(btn, r, c)
            self.studio_btns[key] = btn
        right_lay.addLayout(grid)
        # Studio output
        self.studio_tabs = QTabWidget()
        self.studio_text = QTextEdit()
        self.studio_text.setReadOnly(True)
        self.studio_text.setPlaceholderText(tr("learning_chat_placeholder"))
        self.studio_tabs.addTab(_scrolled(self.studio_text), tr("learning_result"))
        # Mind map view
        self.mindmap_view = QGraphicsView()
        self.mindmap_scene = QGraphicsScene()
        self.mindmap_view.setScene(self.mindmap_scene)
        self.mindmap_view.setMinimumHeight(200)
        self.studio_tabs.addTab(self.mindmap_view, tr("learning_mindmap_tab"))
        right_lay.addWidget(self.studio_tabs, 1)
        # Save/Export
        exp_row = QHBoxLayout()
        exp_row.addWidget(_btn(tr("learning_save"), h=32, slot=self._save_studio))
        exp_row.addWidget(_btn(tr("learning_export"), h=32, slot=self._export_studio))
        right_lay.addLayout(exp_row)
        splitter.addWidget(right_w)

        splitter.setSizes([280, 400, 320])
        self._splitter = splitter
        root.addWidget(splitter, 1)
        self._refresh_api_status()

    # ── Notebook ─────────────────────────────────────────────────────────
    def load(self):
        if not AppState.user_id:
            return
        self._load_notebooks()
        self._refresh_api_status()

    def _refresh_api_status(self):
        try:
            key = db.get_gemini_api_key(self.user_id) if hasattr(db, 'get_gemini_api_key') else ""
            if key and (key.startswith("AIza") or key.startswith("AQ.")):
                self.api_status.setText(tr("learning_api_ready"))
                self.api_status.setStyleSheet("color:#80c000; font-weight:bold;")
            else:
                self.api_status.setText(tr("learning_api_missing"))
                self.api_status.setStyleSheet("color:#f0a800;")
        except:
            self.api_status.setText("")

    def _manage_api_key(self):
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("learning_api_dialog_title"))
        dlg.setMinimumWidth(500)
        dlg.setStyleSheet(build_ss())
        lay = QVBoxLayout(dlg)
        lay.setSpacing(12)
        lay.addWidget(_lbl(tr("learning_api_dialog_desc"), "sub", 11))
        inp = QLineEdit()
        inp.setPlaceholderText("AIza... / AQ...")
        try:
            cur = db.get_gemini_api_key(self.user_id)
            inp.setText(cur)
        except:
            pass
        inp.setEchoMode(QLineEdit.EchoMode.Password)
        lay.addWidget(inp)
        lay.addWidget(_lbl(tr("learning_api_key_info"), "sub", 10))
        btn_row = QHBoxLayout()
        btn_row.addStretch()
        btn_row.addWidget(_btn(tr("msg_cancel"), "", dlg.reject))
        btn_row.addWidget(_btn(tr("msg_ok"), "solid", dlg.accept))
        lay.addLayout(btn_row)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            key = inp.text().strip()
            db.set_gemini_api_key(self.user_id, key)
            SND.complete()
            _show(self, tr("berhasil_title"), tr("learning_api_saved"), "success")
            self._refresh_api_status()

    def _load_notebooks(self):
        # Preserve current selection (FIX: generate studio jangan pindah notebook)
        prev_id = self.current_notebook_id
        self.nb_combo.blockSignals(True)
        self.nb_combo.clear()
        try:
            nbs = db.get_learning_notebooks(self.user_id)
        except:
            nbs = []
        if not nbs:
            try:
                db.create_learning_notebook(self.user_id, "Notebook Pertama")
                nbs = db.get_learning_notebooks(self.user_id)
            except:
                pass
        for nb in nbs:
            self.nb_combo.addItem(f"📓 {nb['title']}", nb['id'])
        # Restore previous selection if still exists
        if prev_id is not None:
            idx = self.nb_combo.findData(prev_id)
            if idx >= 0:
                self.nb_combo.setCurrentIndex(idx)
                self.current_notebook_id = prev_id
            else:
                self.current_notebook_id = self.nb_combo.currentData() if nbs else None
        else:
            self.current_notebook_id = self.nb_combo.currentData() if nbs else None
        self.nb_combo.blockSignals(False)
        if self.current_notebook_id:
            self._load_sources()
            self._load_chats()
            self._load_generations()
        else:
            self.current_notebook_id = None

    def _on_notebook_changed(self, idx):
        if idx < 0:
            return
        self.current_notebook_id = self.nb_combo.currentData()
        self._load_sources()
        self._load_chats()
        self._load_generations()

    def _create_notebook(self):
        name, ok = QInputDialog.getText(self, tr("learning_new_notebook_title"), tr("learning_new_notebook_prompt"))
        if ok and name.strip():
            db.create_learning_notebook(self.user_id, name.strip())
            self._load_notebooks()
            SND.complete()

    def _rename_notebook(self):
        if not self.current_notebook_id:
            return
        cur = self.nb_combo.currentText().replace("📓 ","")
        name, ok = QInputDialog.getText(self, tr("learning_rename_title"), tr("learning_rename_prompt"), text=cur)
        if ok and name.strip():
            db.update_learning_notebook(self.current_notebook_id, self.user_id, name.strip())
            self._load_notebooks()

    def _delete_notebook(self):
        if not self.current_notebook_id:
            return
        reply = QMessageBox.question(self, tr("learning_delete_notebook_title"), tr("learning_delete_notebook_msg"),
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_learning_notebook(self.current_notebook_id, self.user_id)
            self._load_notebooks()

    # ── Sources ──────────────────────────────────────────────────────────
    def _load_sources(self):
        self.src_list.clear()
        if not self.current_notebook_id:
            self.src_count_lbl.setText("0 sources")
            return
        try:
            srcs = db.get_learning_sources(self.current_notebook_id)
        except:
            srcs = []
        for s in srcs:
            icon = {"pdf":"📄","docx":"📃","txt":"📝","paste":"📋","website":"🌐","youtube":"▶️"}.get(s['type'], "📄")
            item = QListWidgetItem(f"{icon} {s['title']} ({s['type']})")
            item.setData(Qt.ItemDataRole.UserRole, s['id'])
            item.setToolTip(s['title'][:100])
            self.src_list.addItem(item)
        self.src_count_lbl.setText(f"{len(srcs)} sources • {sum(len(s['content']) for s in srcs)//1000}k chars")

    def _add_source_menu(self):
        if not self.current_notebook_id:
            _show(self, tr("gagal_title"), tr("learning_need_notebook"), "error")
            return
        menu = QMenu(self)
        menu.addAction("📄 Upload PDF", lambda: self._add_source_file("pdf"))
        menu.addAction("📃 Upload DOCX", lambda: self._add_source_file("docx"))
        menu.addAction("📝 Upload TXT", lambda: self._add_source_file("txt"))
        menu.addAction("📋 Paste Text", self._add_source_paste)
        menu.addAction("🌐 Website URL", lambda: self._add_source_url("website"))
        menu.addAction("▶️ Youtube URL", lambda: self._add_source_url("youtube"))
        menu.exec(QCursor.pos())

    def _add_source_file(self, ftype):
        filt = {"pdf":"PDF (*.pdf)", "docx":"Word (*.docx)", "txt":"Text (*.txt *.md)"}[ftype]
        path, _ = QFileDialog.getOpenFileName(self, f"Pilih {ftype.upper()}", "", filt)
        if not path:
            return
        loading = LoadingDialog(f"Membaca {os.path.basename(path)}...", self)
        loading.show()
        QApplication.processEvents()
        try:
            if ftype == "pdf":
                content = lh.extract_from_pdf(path) if LEARNING_AVAILABLE and lh else open(path,'rb').read().decode(errors='ignore')[:50000]
            elif ftype == "docx":
                content = lh.extract_from_docx(path) if LEARNING_AVAILABLE else "Gagal baca docx"
            else:
                content = lh.extract_from_txt(path) if LEARNING_AVAILABLE else open(path,'r',encoding='utf-8',errors='ignore').read()[:50000]
            if not content.strip():
                content = f"[File kosong: {path}]"
            title = os.path.basename(path)
            db.add_learning_source(self.current_notebook_id, self.user_id, ftype, title, path, content[:80000])
            SND.complete()
            _show(self, tr("berhasil_title"), tr("learning_source_added_detail", title=title, chars=len(content)//1000), "success")
        except Exception as e:
            SND.error()
            _show(self, "Gagal", str(e), "error")
        finally:
            loading.accept()
            self._load_sources()

    def _add_source_paste(self):
        dlg = QDialog(self)
        dlg.setWindowTitle("Paste Text")
        dlg.setMinimumSize(600, 400)
        dlg.setStyleSheet(build_ss())
        lay = QVBoxLayout(dlg)
        lay.addWidget(_lbl(tr("learning_paste_title_label"), size=12))
        title_in = _input("Contoh: Materi Biologi Bab 1")
        lay.addWidget(title_in)
        lay.addWidget(_lbl(tr("learning_paste_desc"), size=12))
        text_edit = QTextEdit()
        text_edit.setPlaceholderText("Paste panjang di sini...")
        lay.addWidget(text_edit)
        btn_row = QHBoxLayout()
        btn_row.addStretch()
        btn_row.addWidget(_btn("Batal", "", dlg.reject))
        btn_row.addWidget(_btn("Simpan", "solid", dlg.accept))
        lay.addLayout(btn_row)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            title = title_in.text().strip() or "Paste Text"
            content = text_edit.toPlainText().strip()
            if not content:
                _show(self, tr("gagal_title"), tr("learning_text_empty"), "error")
                return
            db.add_learning_source(self.current_notebook_id, self.user_id, "paste", title, "", content[:80000])
            self._load_sources()
            SND.complete()

    def _add_source_url(self, utype):
        url, ok = QInputDialog.getText(self, "Tambah URL", f"Masukkan {utype} URL:")
        if not ok or not url.strip():
            return
        url = url.strip()
        loading = LoadingDialog(f"Mengambil {utype}...", self)
        loading.show()
        QApplication.processEvents()
        try:
            if utype == "website":
                content = lh.fetch_website_text(url) if LEARNING_AVAILABLE else f"[Website: {url}]"
                title = url[:50]
            else:
                content = lh.fetch_youtube_transcript(url) if LEARNING_AVAILABLE else f"[Youtube: {url}]"
                title = f"Youtube: {url[:40]}"
            if "Gagal" in content or "tidak ditemukan" in content:
                _show(self, "Gagal", content, "error")
                loading.accept()
                return
            db.add_learning_source(self.current_notebook_id, self.user_id, utype, title, url, content[:80000])
            SND.complete()
            _show(self, tr("berhasil_title"), tr("learning_source_added_generic", type=utype), "success")
        except Exception as e:
            _show(self, "Gagal", str(e), "error")
        finally:
            loading.accept()
            self._load_sources()

    def _view_source(self, item):
        sid = item.data(Qt.ItemDataRole.UserRole)
        # Cari source
        try:
            srcs = db.get_learning_sources(self.current_notebook_id)
            src = next((s for s in srcs if s['id']==sid), None)
            if not src:
                return
            dlg = QDialog(self)
            dlg.setWindowTitle(f"Source: {src['title']}")
            dlg.setMinimumSize(700, 500)
            dlg.setStyleSheet(build_ss())
            lay = QVBoxLayout(dlg)
            lay.addWidget(_lbl(src['title'], "section", 13, True))
            lay.addWidget(_lbl(f"Tipe: {src['type']} • {len(src['content'])//1000}k chars", "sub", 10))
            txt = QTextEdit()
            txt.setReadOnly(True)
            txt.setPlainText(src['content'][:20000])
            lay.addWidget(txt)
            lay.addWidget(_btn("Tutup", "", dlg.accept))
            dlg.exec()
        except Exception as e:
            _show(self, "Error", str(e), "error")

    def _src_context_menu(self, pos):
        item = self.src_list.itemAt(pos)
        if not item:
            return
        menu = QMenu(self)
        menu.addAction(tr("learning_view_source"), lambda: self._view_source(item))
        menu.addAction(tr("learning_delete_source"), self._delete_source)
        menu.exec(self.src_list.viewport().mapToGlobal(pos))

    def _delete_source(self):
        item = self.src_list.currentItem()
        if not item:
            return
        sid = item.data(Qt.ItemDataRole.UserRole)
        reply = QMessageBox.question(self, tr("learning_delete_source_title"), tr("learning_delete_source_msg"),
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_learning_source(sid, self.user_id)
            self._load_sources()

    # ── Chat ─────────────────────────────────────────────────────────────
    def _load_chats(self):
        self.chat_area.clear()
        self._chat_history = []
        if not self.current_notebook_id:
            return
        try:
            chats = db.get_learning_chats(self.current_notebook_id)
            for c in chats:
                self._append_chat(c['role'], c['content'], c.get('citations'))
                self._chat_history.append({"role": c['role'], "content": c['content']})
        except:
            pass

    def _format_chat_html(self, text: str) -> str:
        """Convert markdown ke HTML modern - tanpa kotak kosong, rapi seperti ChatGPT."""
        import html
        # Bersihkan teks aneh: hapus baris yang cuma "□" atau "D:" terpotong, dan empty markers
        text = re.sub(r'\n\s*□\s*\n', '\n\n', text)
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        # Simpan code blocks
        code_blocks = []
        def code_repl(m):
            code_blocks.append(f"<pre style='background:{_T('panel')};border:1px solid {_T('border')};border-radius:8px;padding:10px;margin:10px 0;white-space:pre-wrap;font-family:monospace;font-size:11px;line-height:1.4;'>{html.escape(m.group(1).strip())}</pre>")
            return f"__CODE_{len(code_blocks)-1}__"
        text = re.sub(r'```(.*?)```', code_repl, text, flags=re.DOTALL)
        # Escape
        text = html.escape(text)
        # Bold **text** -> <b>
        text = re.sub(r'\*\*(.*?)\*\*', r'<b style="color:' + _T('light') + ';">\1</b>', text)
        # Italic - hanya untuk _text_ atau *text* yang tidak di bold
        text = re.sub(r'(?<!\*)\*(?!\*)([^\*\n]+?)\*(?!\*)', r'<i style="color:' + _T('muted') + ';">\1</i>', text)
        # Headers # ## ###
        text = re.sub(r'^###\s*(.+)$', r'<div style="font-size:13px;font-weight:bold;color:' + _T('light') + ';margin:8px 0 4px 0;">\1</div>', text, flags=re.MULTILINE)
        text = re.sub(r'^##\s*(.+)$', r'<div style="font-size:14px;font-weight:bold;color:' + _T('accent') + ';margin:10px 0 6px 0;border-bottom:1px solid ' + _T('border') + ';padding-bottom:4px;">\1</div>', text, flags=re.MULTILINE)
        text = re.sub(r'^#\s*(.+)$', r'<div style="font-size:15px;font-weight:bold;color:' + _T('text') + ';margin:12px 0 6px 0;">\1</div>', text, flags=re.MULTILINE)
        # Bersihkan baris kosong berlebih dan tanda aneh
        text = re.sub(r'\n{3,}', '\n\n', text)
        lines = text.split('\n')
        html_lines = []
        in_list = False
        for line in lines:
            stripped = line.strip()
            # Skip baris yang cuma kotak kosong atau "□" atau "D:" tanpa konteks
            if stripped in ["□", "D:", "D", ""] and len(stripped) < 3:
                # Jangan buat kotak kosong, skip saja kecuali butuh spacing
                if stripped == "":
                    # Hanya 1 spacer per 2 baris kosong
                    if html_lines and not html_lines[-1].startswith('<div style="height'):
                        html_lines.append('')
                continue
            if re.match(r'^[\-\•\*]\s+', stripped):
                if not in_list:
                    html_lines.append(f'<ul style="margin:6px 0 6px 0;padding:0;list-style:none;color:{_T("text")};">')
                    in_list = True
                content = re.sub(r'^[\-\•\*]\s+', '', stripped)
                # Hapus "D:" terpotong di awal (bug dari AI yang kepotong "Definisi")
                if content.startswith("D:"):
                    content = "Definisi:" + content[2:]
                if content.startswith("D "):
                    content = "Definisi " + content[2:]
                # Pisahkan judul bold
                if ":" in content and "<b" not in content:
                    parts = content.split(":", 1)
                    if len(parts[0]) < 50 and len(parts[1]) > 5:
                        content = f'<b style="color:{_T("light")};">{parts[0]}:</b>{parts[1]}'
                html_lines.append(f'<li style="margin:5px 0;padding:8px 12px;background:{_T("panel")};border:1px solid {_T("border")};border-radius:8px;line-height:1.5;font-size:12px;">• {content}</li>')
            else:
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                if stripped == "":
                    continue  # Skip empty, sudah ada spacing dari li/div
                elif stripped.startswith("Ada yang bisa") or "bisa saya bantu" in stripped.lower() or "berdasarkan sumber" in stripped.lower():
                    html_lines.append(f'<div style="margin:8px 0;padding:0;color:{_T("text")};line-height:1.6;">{line}</div>')
                elif stripped.startswith("Semangat belajar"):
                    html_lines.append(f'<div style="margin:12px 0 4px 0;padding:10px 14px;background:qlineargradient(x1:0,y1:0,x2:1,y2:0,stop:0 rgba(124,92,255,15), stop:1 rgba(34,211,238,10));border:1px solid {_T("accent")};border-radius:10px;color:{_T("light")};font-style:italic;text-align:center;">{line}</div>')
                else:
                    html_lines.append(f'<div style="margin:5px 0;line-height:1.65;">{line}</div>')
        if in_list:
            html_lines.append('</ul>')
        text = "\n".join(html_lines)
        # Restore code blocks
        for i, cb in enumerate(code_blocks):
            text = text.replace(f"__CODE_{i}__", cb)
        return text

    def _append_chat(self, role, content, citations=None):
        is_ai = role == "assistant"
        color = _T("accent") if is_ai else _T("light")
        bg = _T("panel") if is_ai else "rgba(124,92,255,12)"
        name = "🤖 AI" if is_ai else "🧑 Kamu"
        # Format content jika AI
        html_content = self._format_chat_html(content) if is_ai else f'<div style="margin:4px 0;">{content}</div>'
        bubble = f"""
        <div style="margin:10px 0;padding:0;background:transparent;">
            <div style="display:inline-block;max-width:100%;padding:14px 16px;background:{bg};border:1px solid {_T('border')};border-radius:12px;box-shadow:0 1px 3px rgba(0,0,0,12);">
                <div style="font-weight:bold;color:{color};font-size:11px;margin-bottom:8px;letter-spacing:0.3px;">{name} <span style="color:{_T('muted')};font-weight:normal;font-size:9px;">• {self._get_time_str()}</span></div>
                <div style="color:{_T('text')};font-size:12.5px;line-height:1.7;">{html_content}</div>
            </div>
        </div>
        """
        self.chat_area.append(bubble)
        if citations:
            try:
                cites = json.loads(citations) if isinstance(citations, str) else citations
                if cites:
                    self.citation_lbl.setText("📚 Sources: " + ", ".join([str(c)[:30] for c in cites][:3]))
            except:
                pass
        # Auto scroll
        QTimer.singleShot(50, lambda: self.chat_area.verticalScrollBar().setValue(self.chat_area.verticalScrollBar().maximum()))
    
    def _get_time_str(self):
        from datetime import datetime
        return datetime.now().strftime("%H:%M")

    def _send_chat(self):
        q = self.chat_input.text().strip()
        if not q or not self.current_notebook_id:
            return
        # Cek sources
        try:
            srcs = db.get_learning_sources(self.current_notebook_id)
            if not srcs:
                _show(self, tr("info_title"), tr("learning_need_source_chat"), "info")
                return
        except:
            pass
        self._append_chat("user", q)
        db.add_learning_chat(self.current_notebook_id, "user", q)
        self.chat_input.clear()
        # Loading
        self.chat_area.append('<i style="color:#888;">🤖 Mengetik...</i>')
        QApplication.processEvents()
        try:
            api_key = db.get_gemini_api_key(self.user_id) if hasattr(db, 'get_gemini_api_key') else ""
            # Ambil relevant chunks
            chunks = db.get_relevant_chunks(self.current_notebook_id, q, top_k=4) if hasattr(db, 'get_relevant_chunks') else []
            chunk_texts = [c['chunk_text'] if isinstance(c, dict) else str(c) for c in chunks]
            if not chunk_texts:
                # Fallback: ambil semua sources content dipotong
                srcs = db.get_learning_sources(self.current_notebook_id)
                chunk_texts = [s['content'][:1000] for s in srcs[:2]]
            # Panggil AI
            if LEARNING_AVAILABLE and lh:
                ans = lh.chat_with_sources(q, chunk_texts, self._chat_history, api_key)
            else:
                ans = f"[Mock] Jawaban untuk: {q}\n\nBerdasarkan sources, ini ringkasan... (Isi API Key untuk jawaban real)"
            # Hapus loading
            cur = self.chat_area.toPlainText()
            # Remove last typing
            self.chat_area.clear()
            # Re-render chats
            self._load_chats()
            # Append new answer
            self._append_chat("assistant", ans)
            db.add_learning_chat(self.current_notebook_id, "assistant", ans)
            self._chat_history.append({"role":"user","content":q})
            self._chat_history.append({"role":"assistant","content":ans})
            SND.complete()
            # Reward XP
            try:
                db.gain_xp_gold(self.user_id, 10, 5)
                AppState.refresh()
            except:
                pass
        except Exception as e:
            self.chat_area.append(f'<div style="color:#e05050;">Error: {e}</div>')
        finally:
            # Hapus typing indicator dengan reload
            pass

    def _clear_chat(self):
        if not self.current_notebook_id:
            return
        reply = QMessageBox.question(self, tr("learning_clear_chat_title"), tr("learning_clear_chat_msg"), QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.clear_learning_chats(self.current_notebook_id)
            self._load_chats()
            SND.click()

    # ── Studio ───────────────────────────────────────────────────────────
    def _load_generations(self):
        # Untuk sekarang hanya tampilkan yang terbaru di text
        pass

    def _generate_studio(self, stype):
        if not self.current_notebook_id:
            _show(self, tr("gagal_title"), tr("learning_need_notebook2"), "error")
            return
        # Cek sources
        try:
            srcs = db.get_learning_sources(self.current_notebook_id)
            if not srcs:
                _show(self, tr("info_title"), tr("learning_need_source_studio"), "info")
                return
        except:
            pass
        # Prompt untuk studio
        query, ok = QInputDialog.getText(self, tr("learning_studio"), tr("learning_topic_prompt", stype=stype))
        if not ok:
            return
        query = query.strip()
        loading = LoadingDialog(f"{tr('learning_generating')} {stype}...", self)
        loading.show()
        QApplication.processEvents()
        try:
            api_key = db.get_gemini_api_key(self.user_id) if hasattr(db, 'get_gemini_api_key') else ""
            # Ambil chunks relevan
            chunks = db.get_relevant_chunks(self.current_notebook_id, query or "ringkasan", top_k=5) if hasattr(db, 'get_relevant_chunks') else []
            chunk_texts = [c['chunk_text'] if isinstance(c, dict) else str(c) for c in chunks]
            if not chunk_texts:
                srcs = db.get_learning_sources(self.current_notebook_id)
                chunk_texts = [s['content'][:1200] for s in srcs[:3]]
            if LEARNING_AVAILABLE and lh:
                content = lh.generate_studio_content(stype, query, chunk_texts, api_key)
            else:
                content = f"[Mock {stype}] untuk query: {query}\n\n" + "\n\n".join(chunk_texts)[:2000]
            # Simpan
            db.add_learning_generation(self.current_notebook_id, stype, f"{stype} - {query[:30] or 'General'}", content)
            # Tampilkan dengan format modern
            if stype == "mind_map":
                self._render_mindmap(content)
                self.studio_tabs.setCurrentIndex(1)
            elif stype == "flashcards":
                self._render_flashcards(content)
                self.studio_tabs.setCurrentIndex(0)
            elif stype in ["study_guide", "briefing", "faq", "timeline", "summary", "audio_overview"]:
                html = self._format_studio_html(content, stype)
                self.studio_text.setHtml(html)
                self.studio_tabs.setCurrentIndex(0)
            else:
                self.studio_text.setPlainText(content)
                self.studio_tabs.setCurrentIndex(0)
            SND.complete()
            _show(self, "Berhasil", f"{stype} berhasil digenerate!", "success")
            try:
                db.gain_xp_gold(self.user_id, 15, 8)
                AppState.refresh()
            except:
                pass
        except Exception as e:
            SND.error()
            _show(self, "Gagal", str(e), "error")
        finally:
            loading.accept()

    def _render_mindmap(self, json_text):
        self.mindmap_scene.clear()
        if "Quota" in json_text or "Error Gemini" in json_text or "429" in json_text:
            self.studio_text.setPlainText(json_text)
            err_item = self.mindmap_scene.addText("Gagal generate mind map\nQuota habis / JSON invalid")
            err_item.setDefaultTextColor(QColor("#e05050"))
            self.mindmap_scene.addRect(err_item.boundingRect(), QPen(QColor("#e05050")), QColor("transparent"))
            return
        try:
            data = None
            m = re.search(r'\{[\s\S]*\}', json_text)
            raw = m.group(0) if m else json_text
            for attempt in range(3):
                try:
                    if attempt == 0:
                        data = json.loads(raw)
                    elif attempt == 1:
                        fixed = raw.replace("'", '"')
                        fixed = re.sub(r',\s*}', '}', fixed)
                        fixed = re.sub(r',\s*]', ']', fixed)
                        # Fix missing quotes around keys
                        fixed = re.sub(r'(\w+):', r'"\1":', fixed)
                        fixed = re.sub(r'""(\w+)""', r'"\1"', fixed)
                        data = json.loads(fixed)
                    elif attempt == 2:
                        import ast
                        data = ast.literal_eval(raw)
                    break
                except:
                    continue
            if data is None:
                raise ValueError("JSON tidak valid")
            if "central" not in data or "branches" not in data:
                if "topic" in data:
                    data = {"central": data.get("topic", "Central"), "branches": data.get("branches", [])}
                elif "title" in data:
                    data = {"central": data.get("title", "Central"), "branches": []}
                else:
                    raise ValueError("Struktur JSON tidak sesuai")
            central = data.get("central", "Central")
            branches = data.get("branches", [])
            # Layout rapi: central di atas, branches di bawah terdistribusi horizontal, children di bawah branch
            import math
            # Hitung ukuran central
            central_text = str(central)[:40]
            # Central node di atas tengah
            cw, ch = 160, 40
            central_item = self.mindmap_scene.addRect(-cw/2, -140, cw, ch, QPen(QColor(_T("accent")), 2), QColor(_T("panel")))
            central_item.setZValue(10)
            t = self.mindmap_scene.addText(central_text)
            t.setPos(-t.boundingRect().width()/2, -140 + (ch - t.boundingRect().height())/2)
            t.setDefaultTextColor(QColor(_T("accent")))
            # Buat font bold untuk central
            f = t.font()
            f.setBold(True)
            f.setPointSize(10)
            t.setFont(f)
            
            n = len(branches) if branches else 0
            if n == 0:
                self.mindmap_view.setSceneRect(self.mindmap_scene.itemsBoundingRect().adjusted(-20,-20,20,20))
                self.studio_text.setPlainText(json_text)
                return
            
            # Hitung posisi branches: distribusi horizontal
            total_width = min(600, n * 180)
            start_x = -total_width/2 + 80
            spacing = total_width / max(1, n)
            branch_y = -40
            child_y_offset = 70
            
            for i, br in enumerate(branches):
                bx = start_x + i * spacing
                # Branch node
                label = str(br.get("label", f"Branch {i}"))[:30]
                # Ukuran rect menyesuaikan panjang label
                bw = max(100, min(160, len(label)*7 + 20))
                bh = 32
                # Shadow effect
                shadow = self.mindmap_scene.addRect(bx - bw/2 + 2, branch_y + 2, bw, bh, QPen(Qt.PenStyle.NoPen), QColor(0,0,0,30))
                rect = self.mindmap_scene.addRect(bx - bw/2, branch_y, bw, bh, QPen(QColor(_T("primary")), 1.5), QColor(_T("panel")))
                rect.setZValue(5)
                # Line dari central ke branch
                self.mindmap_scene.addLine(0, -100, bx, branch_y, QPen(QColor(_T("accent")), 1.5, Qt.PenStyle.SolidLine))
                # Text branch
                bt = self.mindmap_scene.addText(label)
                bt.setPos(bx - bt.boundingRect().width()/2, branch_y + (bh - bt.boundingRect().height())/2)
                bt.setDefaultTextColor(QColor(_T("text")))
                bf = bt.font()
                bf.setBold(True)
                bf.setPointSize(9)
                bt.setFont(bf)
                bt.setZValue(6)
                
                # Children di bawah branch - vertical list
                children = br.get("children", [])[:4]
                for j, ch in enumerate(children):
                    ch_text = str(ch)[:35]
                    cw2, ch2 = 140, 24
                    cx = bx
                    cy = branch_y + child_y_offset + j * (ch2 + 8)
                    # Line branch -> child
                    self.mindmap_scene.addLine(bx, branch_y + bh, cx, cy, QPen(QColor(_T("border")), 1, Qt.PenStyle.DashLine))
                    # Child rect
                    cr = self.mindmap_scene.addRect(cx - cw2/2, cy, cw2, ch2, QPen(QColor(_T("border"))), QColor(_T("bg")))
                    cr.setZValue(3)
                    ct = self.mindmap_scene.addText(ch_text)
                    # Wrap jika terlalu panjang
                    if ct.boundingRect().width() > cw2 - 10:
                        # Potong dan tambah ...
                        while ct.boundingRect().width() > cw2 - 10 and len(ch_text) > 10:
                            ch_text = ch_text[:-4] + "..."
                            ct.setPlainText(ch_text)
                    ct.setPos(cx - ct.boundingRect().width()/2, cy + (ch2 - ct.boundingRect().height())/2)
                    ct.setDefaultTextColor(QColor(_T("muted")))
                    f2 = ct.font()
                    f2.setPointSize(8)
                    ct.setFont(f2)
                    ct.setZValue(4)
            
            # Set scene rect dengan padding
            rect = self.mindmap_scene.itemsBoundingRect()
            self.mindmap_view.setSceneRect(rect.adjusted(-30, -30, 30, 30))
            self.mindmap_view.centerOn(0, -20)
            # Juga tampilkan JSON mentah di Result tab untuk copy
            self.studio_text.setPlainText(json_text)
        except Exception as e:
            self.studio_text.setPlainText(f"[Gagal render mindmap: {e}]\n\nRaw:\n{json_text}")

    def _format_studio_html(self, text: str, stype: str) -> str:
        """Format studio text (markdown) jadi HTML modern untuk QTextEdit."""
        import html as html_lib
        text = html_lib.escape(text)
        text = re.sub(r'\*\*(.*?)\*\*', r'<b style="color:' + _T('light') + ';">\\1</b>', text)
        lines = text.split('\n')
        out = []
        in_list = False
        for line in lines:
            s = line.strip()
            if re.match(r'^[\-\•\*]\s+', s):
                if not in_list:
                    out.append(f'<ul style="margin:8px 0 8px 16px;padding:0;list-style:none;">')
                    in_list = True
                c = re.sub(r'^[\-\•\*]\s+', '', s)
                if ":" in c and "<b" not in c and len(c.split(":")[0]) < 50:
                    p = c.split(":",1)
                    c = f'<b style="color:{_T("light")};">{p[0]}:</b>{p[1]}'
                out.append(f'<li style="margin:6px 0;padding:8px 12px;background:{_T("panel")};border:1px solid {_T("border")};border-radius:8px;">• {c}</li>')
            else:
                if in_list:
                    out.append('</ul>')
                    in_list = False
                if s == "":
                    continue
                else:
                    if s.startswith("### "):
                        s = s[4:]
                        out.append(f'<div style="font-size:13px;font-weight:bold;color:{_T("light")};margin:10px 0 6px 0;">{s}</div>')
                    elif s.startswith("## "):
                        s = s[3:]
                        out.append(f'<div style="font-size:14px;font-weight:bold;color:{_T("accent")};margin:12px 0 8px 0;border-bottom:1px solid {_T("border")};padding-bottom:4px;">{s}</div>')
                    elif s.startswith("# "):
                        s = s[2:]
                        out.append(f'<div style="font-size:16px;font-weight:bold;color:{_T("text")};margin:14px 0 8px 0;">{s}</div>')
                    else:
                        out.append(f'<div style="margin:6px 0;line-height:1.6;">{line}</div>')
        if in_list:
            out.append('</ul>')
        body = "\n".join(out)
        return f'<div style="font-family: Segoe UI, sans-serif; font-size:12.5px; color:{_T("text")}; line-height:1.6;">{body}</div>'

    def _render_flashcards(self, json_text: str):
        """Render flashcards JSON jadi HTML modern flip-card grid."""
        try:
            import html as html_lib
            m = re.search(r'\[.*\]', json_text, re.DOTALL)
            raw = m.group(0) if m else json_text
            fixed = raw.replace("'", '"')
            fixed = re.sub(r',\s*}', '}', fixed)
            fixed = re.sub(r',\s*]', ']', fixed)
            data = __import__('json').loads(fixed)
            if not isinstance(data, list):
                raise ValueError("Not a list")
            html = f'<div style="display:flex;flex-wrap:wrap;gap:10px;padding:10px;">'
            for i, card in enumerate(data[:12]):
                front = html_lib.escape(str(card.get("front", f"Q{i+1}")))
                back = html_lib.escape(str(card.get("back", "")))
                html += f'<div style="flex:1 1 45%;min-width:200px;background:{_T("panel")};border:1px solid {_T("border")};border-radius:10px;overflow:hidden;"><div style="background:{_T("accent")};color:white;padding:8px 12px;font-weight:bold;font-size:11px;">Card {i+1}</div><div style="padding:10px 12px;"><div style="font-weight:bold;color:{_T("light")};margin-bottom:6px;">Q: {front}</div><div style="height:1px;background:{_T("border")};margin:6px 0;"></div><div style="color:{_T("muted")};font-size:11px;">A: {back}</div></div></div>'
            html += '</div>'
            self.studio_text.setHtml(html)
        except Exception as e:
            self.studio_text.setHtml(self._format_studio_html(json_text, "flashcards"))


    def _save_studio(self):
        content = self.studio_text.toPlainText().strip()
        if not content:
            _show(self, tr("info_title"), tr("learning_no_result_save"), "info")
            return
        # Sudah auto disimpan saat generate, cuma notifikasi
        _show(self, tr("berhasil_title"), tr("learning_studio_saved"), "success")

    def _export_studio(self):
        content = self.studio_text.toPlainText().strip()
        if not content:
            return
        path, _ = QFileDialog.getSaveFileName(self, "Export Studio", "", "Text (*.txt);;Word (*.docx);;PDF (*.pdf)")
        if not path:
            return
        try:
            if path.endswith(".txt"):
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(content)
            elif path.endswith(".docx"):
                from docx import Document
                doc = Document()
                doc.add_paragraph(content)
                doc.save(path)
            elif path.endswith(".pdf"):
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                from reportlab.lib.styles import getSampleStyleSheet
                doc = SimpleDocTemplate(path)
                styles = getSampleStyleSheet()
                story = [Paragraph(content.replace("\n","<br/>"), styles['Normal'])]
                doc.build(story)
            _show(self, "Berhasil", f"Export ke {path}", "success")
        except Exception as e:
            _show(self, "Gagal", str(e), "error")

    def _toggle_left_panel(self, collapsed):
        # Cari left widget (index 0 di splitter)
        try:
            left = self.findChild(QWidget, "learning_left")
            if not left:
                # Fallback: ambil widget pertama di splitter
                left = self.sender().parent().findChild(QSplitter).widget(0) if hasattr(self, 'sender') else None
            # Simpan sizes
            if not hasattr(self, '_splitter'):
                # Cari splitter
                for child in self.findChildren(QSplitter):
                    self._splitter = child
                    break
            if hasattr(self, '_splitter'):
                sizes = self._splitter.sizes()
                if collapsed:
                    self._splitter.setSizes([0, 600, 320])
                    self.btn_collapse_left.setText("Sources ▶")
                else:
                    self._splitter.setSizes([280, 400, 320])
                    self.btn_collapse_left.setText("◀ Sources")
        except Exception as e:
            print(f"Toggle left failed: {e}")

    def _toggle_right_panel(self, collapsed):
        try:
            if not hasattr(self, '_splitter'):
                for child in self.findChildren(QSplitter):
                    self._splitter = child
                    break
            if hasattr(self, '_splitter'):
                if collapsed:
                    self._splitter.setSizes([280, 400, 0])
                    self.btn_collapse_right.setText("◀ Studio")
                else:
                    self._splitter.setSizes([280, 400, 320])
                    self.btn_collapse_right.setText("Studio ▶")
        except Exception as e:
            print(f"Toggle right failed: {e}")

    def closeEvent(self, e):
        AppState.unregister(self.load)
        AppState.unregister_lang_cb(self.load)
        super().closeEvent(e)


class RemindersPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self.player = None
        self.audio_output = None
        self._test_sound_playing = False
        self._build()
        AppState.register(self.load)
        AppState.register_lang_cb(self.load)

    def _build(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(20, 16, 20, 16)
        main_layout.setSpacing(10)

        # Header
        hdr = QHBoxLayout()
        hdr.addWidget(_lbl(tr("reminders_title"), "section", 14, True))
        hdr.addStretch()
        main_layout.addLayout(hdr)
        main_layout.addWidget(_sep())

        # Toolbar
        toolbar = QHBoxLayout()
        add_btn = _btn(tr("reminders_add"), "solid", self._add_reminder)
        refresh_btn = _btn(tr("reminders_refresh"), slot=self.load)
        toolbar.addWidget(add_btn)
        toolbar.addWidget(refresh_btn)
        toolbar.addStretch()
        main_layout.addLayout(toolbar)

        # Daftar reminder
        self.reminder_list = QListWidget()
        self.reminder_list.setMinimumHeight(300)
        self.reminder_list.itemClicked.connect(self._on_item_clicked)
        main_layout.addWidget(self.reminder_list)

        # Tombol aksi (untuk item yang dipilih)
        action_row = QHBoxLayout()
        self.edit_btn = _btn(tr("reminders_edit"), slot=self._edit_selected)
        self.delete_btn = _btn(tr("reminders_delete"), "danger", self._delete_selected)
        self.toggle_btn = _btn(tr("reminders_toggle"), slot=self._toggle_selected) 
        self.test_btn = _btn(tr("reminders_test"), "gold", self._test_selected)
        action_row.addWidget(self.edit_btn)
        action_row.addWidget(self.delete_btn)
        action_row.addWidget(self.toggle_btn)
        action_row.addWidget(self.test_btn)
        action_row.addStretch()
        main_layout.addLayout(action_row)

        self._selected_id = None
        self.load()

    def load(self):
        if not AppState.user_id:
            return
        self.reminder_list.clear()
        reminders = db.get_reminders(self.user_id)
        if not reminders:
            empty = QListWidgetItem(tr("reminders_empty"))
            empty.setFlags(Qt.ItemFlag.NoItemFlags)
            self.reminder_list.addItem(empty)
            return
        for r in reminders:
            status = "🔔" if r["is_active"] else "🔕"
            triggered = " ✅" if r["triggered"] else ""
            time_str = r["reminder_datetime"][:16].replace("T", " ")
            item = QListWidgetItem(f"{status} {r['title']} - {time_str}{triggered}")
            item.setData(Qt.ItemDataRole.UserRole, r["id"])
            if not r["is_active"]:
                item.setForeground(QColor(_T("muted")))
            self.reminder_list.addItem(item)

    def _add_reminder(self):
        dlg = ReminderDialog(self.user_id, parent=self)
        if dlg.exec():
            self.load()

    def _edit_selected(self):
        if self._selected_id is None:
            return
        r = db.get_reminder(self._selected_id, self.user_id)
        if r:
            dlg = ReminderDialog(self.user_id, r, parent=self)
            if dlg.exec():
                self.load()

    def _delete_selected(self):
        if self._selected_id is None:
            return
        self._stop_sound()
        reply = QMessageBox.question(self, tr("confirm_title"), tr("reminders_delete_confirm"),
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_reminder(self._selected_id, self.user_id)
            SND.click()
            self.load()

    def _toggle_selected(self):
        if self._selected_id is None:
            return
        self._stop_sound() 
        r = db.get_reminder(self._selected_id, self.user_id)
        if r:
            new_state = 0 if r["is_active"] else 1
            db.update_reminder(self._selected_id, self.user_id, is_active=new_state)
            if new_state:
                # reset triggered jika diaktifkan kembali
                db.reset_reminder_triggered(self._selected_id, self.user_id)
            SND.notify()
            self.load()

    def _test_selected(self):
        if self._selected_id is None:
            return
        r = db.get_reminder(self._selected_id, self.user_id)
        if r:
            self._stop_sound()  # hentikan suara sebelumnya jika ada
            self._test_sound_playing = True
            self._play_sound(r["sound_type"], r["sound_file"])
            # Tampilkan dialog dengan tombol OK, setelah OK suara dihentikan
            dlg = QDialog(self)
            dlg.setWindowTitle(tr("reminders_test_title"))
            dlg.setMinimumWidth(350)
            dlg.setMinimumHeight(120)
            dlg.setStyleSheet(build_ss())
            layout = QVBoxLayout(dlg)
            layout.setContentsMargins(20, 20, 20, 20)
            layout.setSpacing(12)
            layout.addWidget(QLabel(tr("reminders_test_msg", title=r["title"])))
            ok_btn = _btn(tr("msg_ok"), "solid", dlg.accept)
            layout.addWidget(ok_btn)
            dlg.exec()
            # Setelah dialog ditutup, hentikan suara
            self._stop_sound()
            self._test_sound_playing = False

    def _play_sound(self, sound_type, sound_file=None):
        if sound_type == "default":
            SND.notify()
        elif sound_type == "beep1":
            SND._beep(600, 200)
            SND._beep(800, 200)
        elif sound_type == "beep2":
            SND._beep(400, 300)
            SND._beep(600, 300)
            SND._beep(800, 300)
        elif sound_type == "custom" and sound_file and os.path.exists(sound_file):
            self._play_mp3(sound_file)
        else:
            SND.notify()

    def _play_mp3(self, filepath):
        try:
            if not os.path.exists(filepath):
                print(f"File MP3 tidak ditemukan: {filepath}")
                return
            
            if self.player is None:
                self.player = QMediaPlayer()
                self.audio_output = QAudioOutput()
                self.player.setAudioOutput(self.audio_output)
                
                # Tambahkan debug error
                self.player.errorOccurred.connect(self._handle_media_error)
            
            url = QUrl.fromLocalFile(filepath)
            self.player.setSource(url)
            self.audio_output.setVolume(1.0)
            self.player.play()
            print(f"Memutar MP3: {filepath}")
            
        except Exception as e:
            print(f"Error playing MP3: {e}")
            import traceback
            traceback.print_exc()

    def _stop_sound(self):
        try:
            if self.player and self.player.isPlaying():
                self.player.stop()
                print("Suara dihentikan")
            self._test_sound_playing = False
        except Exception as e:
            print(f"Error stopping sound: {e}")

    def _handle_media_error(self, error):
        print(f"Media error: {error}")

    def _on_item_clicked(self, item):
        self._selected_id = item.data(Qt.ItemDataRole.UserRole)

    def closeEvent(self, e):
        self._stop_sound()
        AppState.unregister(self.load)
        AppState.unregister_lang_cb(self.load)
        if self.player:
            self.player.stop()
        super().closeEvent(e)

class ReminderDialog(QDialog):
    def __init__(self, user_id, reminder=None, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        self.reminder = reminder
        self.setWindowTitle(tr("reminders_edit_title") if reminder else tr("reminders_add_title"))
        self.setMinimumWidth(500)
        self.setMinimumHeight(450)
        self.setStyleSheet(build_ss())
        self._build()
        if reminder:
            self._load_data()

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(12)

        lay.addWidget(_lbl(self.windowTitle(), "section", 14, True))
        lay.addWidget(_sep())

        # Judul
        lay.addWidget(_lbl(tr("reminders_title_label"), size=12))
        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText(tr("reminders_title_ph"))
        lay.addWidget(self.title_edit)

        # Deskripsi
        lay.addWidget(_lbl(tr("reminders_desc_label"), size=12))
        self.desc_edit = QTextEdit()
        self.desc_edit.setPlaceholderText(tr("reminders_desc_ph"))
        self.desc_edit.setMaximumHeight(80)
        lay.addWidget(self.desc_edit)

        # Waktu
        lay.addWidget(_lbl(tr("reminders_datetime_label"), size=12))
        self.datetime_edit = QDateTimeEdit()
        self.datetime_edit.setCalendarPopup(True)
        self.datetime_edit.setDateTime(QDateTime.currentDateTime())
        self.datetime_edit.setDisplayFormat("yyyy-MM-dd HH:mm")
        lay.addWidget(self.datetime_edit)

        # ========== Repeat Options ==========
        lay.addWidget(_lbl(tr("reminders_repeat_label"), size=12))
        self.repeat_combo = QComboBox()
        self.repeat_combo.addItem(tr("reminders_repeat_none"), "none")
        self.repeat_combo.addItem(tr("reminders_repeat_daily"), "daily")
        self.repeat_combo.addItem(tr("reminders_repeat_weekly"), "weekly")
        self.repeat_combo.addItem(tr("reminders_repeat_custom"), "custom")
        self.repeat_combo.currentIndexChanged.connect(self._on_repeat_changed)
        lay.addWidget(self.repeat_combo)

        # Container untuk custom days (checkbox)
        self.custom_days_widget = QWidget()
        custom_days_layout = QHBoxLayout(self.custom_days_widget)
        self.day_checkboxes = []
        day_names = [tr("day_mon_short"), tr("day_tue_short"), tr("day_wed_short"),
                     tr("day_thu_short"), tr("day_fri_short"), tr("day_sat_short"), tr("day_sun_short")]
        for i, name in enumerate(day_names):
            cb = QCheckBox(name)
            cb.setStyleSheet(f"color:{_T('text')};")
            self.day_checkboxes.append(cb)
            custom_days_layout.addWidget(cb)
        self.custom_days_widget.setVisible(False)
        lay.addWidget(self.custom_days_widget)

        # Suara
        lay.addWidget(_lbl(tr("reminders_sound_label"), size=12))
        sound_row = QHBoxLayout()
        self.sound_combo = QComboBox()
        self.sound_combo.addItem(tr("reminders_sound_default"), "default")
        self.sound_combo.addItem(tr("reminders_sound_beep1"), "beep1")
        self.sound_combo.addItem(tr("reminders_sound_beep2"), "beep2")
        self.sound_combo.addItem(tr("reminders_sound_custom"), "custom")
        self.sound_combo.currentIndexChanged.connect(self._on_sound_changed)
        sound_row.addWidget(self.sound_combo, 1)

        self.browse_btn = _btn(tr("reminders_browse"), slot=self._browse_file)  # ← perbaiki
        self.browse_btn.setEnabled(False)
        sound_row.addWidget(self.browse_btn)

        self.sound_file_label = QLabel("")
        self.sound_file_label.setStyleSheet(f"color:{_T('muted')}; font-size:11px;")
        lay.addLayout(sound_row)
        lay.addWidget(self.sound_file_label)

        lay.addSpacing(8)
        save_btn = _btn(tr("dialog_save"), "solid", self._save)  # ← sudah benar (slot di posisi ketiga)
        lay.addWidget(save_btn)

        root.addWidget(_scrolled(content))

    def _on_sound_changed(self, idx):
        sound_type = self.sound_combo.currentData()
        self.browse_btn.setEnabled(sound_type == "custom")
        if sound_type != "custom":
            self.sound_file_label.setText("")

    def _browse_file(self):
        filepath, _ = QFileDialog.getOpenFileName(
            self,
            tr("reminders_select_mp3"),
            "",
            "MP3 Files (*.mp3);;All Files (*.*)"
        )
        if filepath:
            self.sound_file_label.setText(os.path.basename(filepath))
            self.sound_file_label.setStyleSheet(f"color:{_T('accent')}; font-size:11px;")
            self._selected_file = filepath

    def _load_data(self):
        r = self.reminder
        self.title_edit.setText(r["title"])
        self.desc_edit.setText(r["description"] or "")
        dt = QDateTime.fromString(r["reminder_datetime"], "yyyy-MM-dd HH:mm:ss")
        if dt.isValid():
            self.datetime_edit.setDateTime(dt)
        # Load repeat
        idx = self.repeat_combo.findData(r.get("repeat_type", "none"))
        if idx >= 0:
            self.repeat_combo.setCurrentIndex(idx)
        # Load repeat days
        days = r.get("repeat_days", "")
        if days:
            day_list = [int(d.strip()) for d in days.split(',') if d.strip()]
            for i, cb in enumerate(self.day_checkboxes):
                cb.setChecked(i in day_list)
        if r["sound_file"]:
            self.sound_file_label.setText(os.path.basename(r["sound_file"]))
            self._selected_file = r["sound_file"]

    def _on_repeat_changed(self, idx):
        repeat_type = self.repeat_combo.currentData()
        self.custom_days_widget.setVisible(repeat_type == "custom")

    def _save(self):
        if hasattr(self.parent(), '_stop_sound'):
            self.parent()._stop_sound()
        title = self.title_edit.text().strip()
        if not title:
            _show(self, tr("msg_error"), tr("reminders_title_required"), "error")
            return
        description = self.desc_edit.toPlainText().strip()
        dt = self.datetime_edit.dateTime()
        if not dt.isValid():
            _show(self, tr("msg_error"), tr("reminders_invalid_datetime"), "error")
            return
        now = QDateTime.currentDateTime()
        if dt < now:
            reply = QMessageBox.question(self, tr("confirm_title"), tr("reminders_past_datetime_confirm"),
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply != QMessageBox.StandardButton.Yes:
                return

        sound_type = self.sound_combo.currentData()
        sound_file = None
        if sound_type == "custom":
            if hasattr(self, '_selected_file') and self._selected_file:
                sound_file = self._selected_file
            else:
                _show(self, tr("msg_error"), tr("reminders_custom_file_required"), "error")
                return

        datetime_str = dt.toString("yyyy-MM-dd HH:mm:ss")

        repeat_type = self.repeat_combo.currentData()
        repeat_days = ""
        if repeat_type == "custom":
            selected = [str(i) for i, cb in enumerate(self.day_checkboxes) if cb.isChecked()]
            if not selected:
                _show(self, tr("msg_error"), tr("reminders_custom_days_required"), "error")
                return
            repeat_days = ",".join(selected)

        if self.reminder:
            db.update_reminder(
                self.reminder["id"],
                self.user_id,
                title=title,
                description=description,
                reminder_datetime=datetime_str,
                sound_type=sound_type,
                sound_file=sound_file,
                triggered=0,
                repeat_type=repeat_type,
                repeat_days=repeat_days,
            )
            SND.complete()
            _show(self, tr("berhasil_title"), tr("reminders_updated"), "success")
        else:
            r = db.add_reminder(self.user_id, title, description, datetime_str, sound_type, sound_file, repeat_type, repeat_days)
            if r["ok"]:
                SND.complete()
                _show(self, tr("berhasil_title"), tr("reminders_added"), "success")
            else:
                SND.error()
                _show(self, tr("gagal_title"), r.get("msg", "Gagal menambah reminder"), "error")
                return
        self.accept()

# ══════════════════════════════════════════════════════════════════════════════
#  Music Page
# ══════════════════════════════════════════════════════════════════════════════
class MusicPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self.player = QMediaPlayer()
        self.audio_output = QAudioOutput()
        devices = QMediaDevices.audioOutputs()
        if devices:
            self.audio_output.setDevice(devices[0])
        self.current_playlist_id = -1
        self.playlists = []
        self.player.setAudioOutput(self.audio_output)
        self.current_playlist = []
        self.current_index = -1
        self.shuffle_mode = False
        self.repeat_mode = False
        self.current_playlist_id = None
        self.playlist_name = ""
        self._build()
        self._ensure_favorite_playlist()
        AppState.register(self._reload_playlists)
        self._reload_playlists()
        self._update_ui_style()

    def _update_ui_style(self):
        # Update style tombol shuffle/repeat
        if self.shuffle_mode:
            self.shuffle_btn.setStyleSheet(f"background: {_T('primary')}; color: #fff;")
        else:
            self.shuffle_btn.setStyleSheet("")
        if self.repeat_mode:
            self.repeat_btn.setStyleSheet(f"background: {_T('primary')}; color: #fff;")
        else:
            self.repeat_btn.setStyleSheet("")

    def _build(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(20, 16, 20, 16)
        main_layout.setSpacing(10)

        # === Header ===
        hdr = QHBoxLayout()
        hdr.addWidget(_lbl(tr("music_title"), "section", 14, True))
        hdr.addStretch()
        self.folder_btn = _btn("📁 " + tr("music_select_folder"), slot=self._select_folder)
        self.folder_btn.setMinimumWidth(160)
        hdr.addWidget(self.folder_btn)
        main_layout.addLayout(hdr)
        main_layout.addWidget(_sep())

        # Playlist selector
        self.playlist_combo = QComboBox()
        self.playlist_combo.setMinimumWidth(200)
        self.playlist_combo.currentIndexChanged.connect(self._on_playlist_changed)
        hdr.addWidget(self.playlist_combo)

        # Tombol aksi playlist
        self.new_playlist_btn = _btn(tr("music_new_playlist"), slot=self._create_new_playlist)
        hdr.addWidget(self.new_playlist_btn)

        self.add_song_btn = _btn(tr("music_add_song"), slot=self._add_song_to_playlist)
        hdr.addWidget(self.add_song_btn)

        # === Playlist Management ===
        playlist_mgmt = QHBoxLayout()
        self.playlist_combo = QComboBox()
        self.playlist_combo.addItem(tr("music_select_playlist"), -1)
        self.playlist_combo.currentIndexChanged.connect(self._load_selected_playlist)
        playlist_mgmt.addWidget(QLabel(tr("music_playlist_label")))
        playlist_mgmt.addWidget(self.playlist_combo, 1)

        self.save_playlist_btn = _btn(tr("music_save_playlist"), slot=self._save_current_playlist)
        self.save_playlist_btn.setMinimumWidth(120)
        self.delete_playlist_btn = _btn(tr("music_delete_playlist"), "danger", slot=self._delete_current_playlist)
        self.delete_playlist_btn.setMinimumWidth(120)
        playlist_mgmt.addWidget(self.save_playlist_btn)
        playlist_mgmt.addWidget(self.delete_playlist_btn)
        main_layout.addLayout(playlist_mgmt)

        # === Playlist (list widget) ===
        self.playlist_widget = QListWidget()
        self.playlist_widget.setMinimumHeight(200)
        self.playlist_widget.itemDoubleClicked.connect(self._play_selected)
        main_layout.addWidget(self.playlist_widget)
        self.playlist_widget.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.playlist_widget.customContextMenuRequested.connect(self._show_context_menu)

        # === Now Playing ===
        now_playing_widget = QWidget()
        now_layout = QHBoxLayout(now_playing_widget)
        now_layout.setContentsMargins(0, 4, 0, 4)
        self.now_playing_label = QLabel(tr("music_now_playing"))
        self.now_playing_label.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
        self.now_playing_label.setStyleSheet(f"font-size:14px; color:{_T('accent')}; font-weight:bold;")
        now_layout.addWidget(self.now_playing_label, 1)
        now_layout.addWidget(QLabel(tr("music_playlist_playing")))
        self.playlist_name_label = QLabel("")
        self.playlist_name_label.setStyleSheet(f"color:{_T('muted')}; font-size:12px;")
        now_layout.addWidget(self.playlist_name_label)
        main_layout.addWidget(now_playing_widget)

        # === Progress Bar ===
        self.progress_slider = QSlider(Qt.Orientation.Horizontal)
        self.progress_slider.setRange(0, 100)
        self.progress_slider.sliderMoved.connect(self._seek)
        self.progress_slider.setMinimumHeight(16)
        main_layout.addWidget(self.progress_slider)

        # Lyrics area (collapsible)
        lyrics_widget = QWidget()
        lyrics_layout = QVBoxLayout(lyrics_widget)
        lyrics_layout.setContentsMargins(0, 0, 0, 0)

        self.lyrics_toggle_btn = _btn("📜 Lirik", slot=self._toggle_lyrics, h=30)
        self.lyrics_toggle_btn.setCheckable(True)
        lyrics_layout.addWidget(self.lyrics_toggle_btn)

        self.lyrics_text = QTextEdit()
        self.lyrics_text.setReadOnly(True)
        self.lyrics_text.setMinimumHeight(150)
        self.lyrics_text.setStyleSheet(f"background: {_T('panel')}; border: 1px solid {_T('border')}; border-radius: 6px; padding: 8px;")
        self.lyrics_text.setVisible(False)
        lyrics_layout.addWidget(self.lyrics_text)

        main_layout.addWidget(lyrics_widget)

        # === Time Labels ===
        time_layout = QHBoxLayout()
        self.current_time_label = QLabel("0:00")
        self.total_time_label = QLabel("0:00")
        time_layout.addWidget(self.current_time_label)
        time_layout.addStretch()
        time_layout.addWidget(self.total_time_label)
        main_layout.addLayout(time_layout)

        # Controls and Volume dalam satu baris
        controls_widget = QWidget()
        controls_layout = QHBoxLayout(controls_widget)
        controls_layout.setContentsMargins(0, 0, 0, 0)
        controls_layout.setSpacing(8)

        self.shuffle_btn = _btn("🔀", slot=self._toggle_shuffle, h=40)
        self.shuffle_btn.setCheckable(True)
        self.shuffle_btn.setFixedWidth(50)
        controls_layout.addWidget(self.shuffle_btn)

        self.prev_btn = _btn("⏮", slot=self._prev, h=40)
        self.prev_btn.setFixedWidth(50)
        controls_layout.addWidget(self.prev_btn)

        self.play_btn = _btn("▶️", slot=self._play_pause, h=40)
        self.play_btn.setFixedWidth(70)
        controls_layout.addWidget(self.play_btn)

        self.next_btn = _btn("⏭", slot=self._next, h=40)
        self.next_btn.setFixedWidth(50)
        controls_layout.addWidget(self.next_btn)

        self.repeat_btn = _btn("🔁", slot=self._toggle_repeat, h=40)
        self.repeat_btn.setCheckable(True)
        self.repeat_btn.setFixedWidth(50)
        controls_layout.addWidget(self.repeat_btn)

        controls_layout.addStretch()

        # ===== STYLING ELEMEN =====
        btn_style = """
            QPushButton {
                font-size: 20px;
                background: transparent;
                border: none;
                padding: 4px 8px;
            }
            QPushButton:hover {
                background: #3a3a3a;
                border-radius: 8px;
            }
            QPushButton:checked {
                background: #5a8a2e;
                color: #fff;
                border-radius: 8px;
            }
        """
        for btn in [self.shuffle_btn, self.prev_btn, self.play_btn, self.next_btn, self.repeat_btn]:
            btn.setStyleSheet(btn_style)

        # Volume
        volume_label = QLabel("🔊")
        controls_layout.addWidget(volume_label)
        self.volume_slider = QSlider(Qt.Orientation.Horizontal)
        self.volume_slider.setRange(0, 100)
        self.volume_slider.setValue(70)
        self.volume_slider.setFixedWidth(100)  # <-- batasi lebar slider
        self.volume_slider.valueChanged.connect(self._set_volume)
        self.audio_output.setVolume(0.7)
        controls_layout.addWidget(self.volume_slider)

        main_layout.addWidget(controls_widget)

        # Timer untuk update progress
        self.timer = QTimer()
        self.timer.timeout.connect(self._update_progress)
        self.timer.start(500)

        # Koneksi sinyal player
        self.player.positionChanged.connect(self._on_position_changed)
        self.player.mediaStatusChanged.connect(self._on_media_status_changed)
        self.player.errorOccurred.connect(self._on_player_error)

        # Load data awal
        self._reload_playlists()

    def _ensure_favorite_playlist(self):
        """Buat playlist Favorite otomatis jika belum ada."""
        import json
        playlists = db.get_all_playlists(self.user_id)
        has_fav = any(p['is_favorite'] for p in playlists)
        if not has_fav:
            db.create_playlist(self.user_id, "Favorite", is_favorite=1)
            # Refresh daftar playlist
            self._reload_playlists()

    def _show_context_menu(self, pos):
        item = self.playlist_widget.itemAt(pos)
        if not item:
            return
        index = self.playlist_widget.row(item)
        menu = QMenu(self)
        # Move to other playlist
        move_menu = menu.addMenu(tr("music_move_to_playlist"))
        copy_menu = menu.addMenu(tr("music_copy_to_playlist"))
        remove_action = menu.addAction(tr("music_remove_from_playlist"))
        # Tambahkan daftar playlist tujuan (kecuali dirinya sendiri)
        for pl in self.playlists:
            if pl['id'] == self.current_playlist_id:
                continue
            move_action = move_menu.addAction(pl['name'])
            move_action.triggered.connect(lambda checked, pl_id=pl['id'], idx=index: self._move_song(idx, pl_id))
            copy_action = copy_menu.addAction(pl['name'])
            copy_action.triggered.connect(lambda checked, pl_id=pl['id'], idx=index: self._copy_song(idx, pl_id))
        menu.addSeparator()
        remove_action.triggered.connect(lambda: self._remove_song(index))
        menu.exec(self.playlist_widget.viewport().mapToGlobal(pos))

    def _move_song(self, index, target_playlist_id):
        if db.move_song_to_playlist(self.user_id, self.current_playlist_id, target_playlist_id, index):
            self._load_current_playlist()

    def _copy_song(self, index, target_playlist_id):
        if db.copy_song_to_playlist(self.user_id, self.current_playlist_id, target_playlist_id, index):
            self._load_current_playlist()

    def _remove_song(self, index):
        if db.remove_song_from_playlist(self.user_id, self.current_playlist_id, index):
            self._load_current_playlist()

    # ===== PLAYLIST MANAGEMENT =====
    def _save_playlist(self):
        if not self.current_playlist:
            _show(self, tr("msg_error"), "Tidak ada lagu dalam playlist!", "error")
            return
        name, ok = QInputDialog.getText(self, tr("music_save_playlist"), "Nama playlist:")
        if ok and name.strip():
            db.save_playlist(self.user_id, name.strip(), self.current_playlist)
            _show(self, tr("berhasil_title"), f"Playlist '{name}' berhasil disimpan!", "success")

    def _load_playlist(self):
        playlists = db.get_all_playlists(self.user_id)
        if not playlists:
            _show(self, tr("msg_info"), "Belum ada playlist tersimpan.", "info")
            return
        # Tampilkan dialog pilih playlist
        items = [f"{p['name']} ({len(json.loads(p['files']))} lagu)" for p in playlists]
        item, ok = QInputDialog.getItem(self, tr("music_load_playlist"), "Pilih playlist:", items, 0, False)
        if ok and item:
            # Ambil playlist yang dipilih
            for p in playlists:
                if f"{p['name']} ({len(json.loads(p['files']))} lagu)" == item:
                    self.current_playlist = json.loads(p['files'])
                    self.current_playlist_name = p['name']
                    self._refresh_playlist_widget()
                    self.current_index = 0 if self.current_playlist else -1
                    if self.current_index >= 0:
                        self._play_current()
                    break

    def _refresh_playlist_widget(self):
        self.playlist_widget.clear()
        for f in self.current_playlist:
            self.playlist_widget.addItem(os.path.basename(f))

    def _reload_playlists(self):
        """Muat semua playlist dari database dan refresh combo."""
        self.playlists = db.get_all_playlists(self.user_id)
        self.playlist_combo.blockSignals(True)
        self.playlist_combo.clear()
        for pl in self.playlists:
            name = f"⭐ {pl['name']}" if pl['is_favorite'] else pl['name']
            self.playlist_combo.addItem(name, pl['id'])
        self.playlist_combo.blockSignals(False)
        # Pilih playlist yang sedang aktif (jika ada)
        if self.current_playlist_id != -1:
            idx = self.playlist_combo.findData(self.current_playlist_id)
            if idx >= 0:
                self.playlist_combo.setCurrentIndex(idx)
            else:
                self.playlist_combo.setCurrentIndex(0)
                self.current_playlist_id = self.playlist_combo.currentData()
        elif self.playlist_combo.count() > 0:
            self.playlist_combo.setCurrentIndex(0)
            self.current_playlist_id = self.playlist_combo.currentData()
        self._load_current_playlist()

    def _load_current_playlist(self):
        """Muat lagu dari playlist yang sedang dipilih ke list widget."""
        if self.current_playlist_id == -1:
            return
        playlist = db.get_playlist(self.user_id, self.current_playlist_id)
        if not playlist:
            return
        self.current_playlist = json.loads(playlist['tracks'])
        self.current_playlist_name = playlist['name']
        self._refresh_playlist_widget()
        # Jika ada lagu yang sedang diputar tapi tidak ada di playlist baru, hentikan
        if self.current_index >= len(self.current_playlist):
            self.player.stop()
            self.current_index = -1
            self.play_btn.setText("▶️")
            self.now_playing_label.setText(tr("music_now_playing"))

    def _on_playlist_changed(self, index):
        if index < 0:
            return
        self.current_playlist_id = self.playlist_combo.currentData()
        self._load_current_playlist()

    def _create_new_playlist(self):
        name, ok = QInputDialog.getText(self, tr("music_new_playlist"), "Nama playlist baru:")
        if ok and name.strip():
            pl_id = db.create_playlist(self.user_id, name.strip())
            self._reload_playlists()
            # Pilih playlist baru
            idx = self.playlist_combo.findData(pl_id)
            if idx >= 0:
                self.playlist_combo.setCurrentIndex(idx)
                self.current_playlist_id = pl_id
                self._load_current_playlist()

    def _add_song_to_playlist(self):
        """Tambahkan file audio ke playlist aktif."""
        if self.current_playlist_id == -1:
            _show(self, tr("msg_error"), "Pilih playlist terlebih dahulu!", "error")
            return
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            tr("music_add_song"),
            "",
            "Audio Files (*.mp3 *.wav *.flac *.m4a *.ogg)"
        )
        if file_paths:
            for f in file_paths:
                db.add_song_to_playlist(self.user_id, self.current_playlist_id, f)
            self._load_current_playlist()

    def _delete_current_playlist(self):
        """Hapus playlist (kecuali Favorite)."""
        if self.current_playlist_id == -1:
            return
        playlist = db.get_playlist(self.user_id, self.current_playlist_id)
        if playlist.get('is_favorite'):
            _show(self, tr("msg_error"), "Tidak bisa menghapus playlist Favorite!", "error")
            return
        reply = QMessageBox.question(self, tr("confirm_title"), f"Hapus playlist '{playlist['name']}'?",
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_playlist(self.user_id, self.current_playlist_id)
            self.current_playlist_id = -1
            self._reload_playlists()

    def _load_selected_playlist(self, idx):
        """Load tracks dari playlist yang dipilih"""
        playlist_id = self.playlist_combo.currentData()
        if playlist_id == -1:
            return
        r = db.load_playlist(playlist_id, self.user_id)
        if r["ok"]:
            self.current_playlist = r["tracks"]
            self.playlist_name = r["name"]
            self.current_playlist_id = playlist_id
            self.playlist_name_label.setText(r["name"])
            self._update_playlist_widget()
            if self.current_playlist:
                self.current_index = 0
                self.now_playing_label.setText(os.path.basename(self.current_playlist[0]))

    def _save_current_playlist(self):
        if not self.current_playlist:
            _show(self, tr("msg_error"), "Tidak ada lagu dalam playlist!", "error")
            return
        name, ok = QInputDialog.getText(self, tr("music_save_playlist"), "Nama playlist:")
        if ok and name.strip():
            r = db.save_playlist(self.user_id, name.strip(), self.current_playlist)
            if r["ok"]:
                self.current_playlist_id = r["id"]
                _show(self, tr("berhasil_title"), r["msg"], "success")
                self._reload_playlists()
                # Pilih playlist yang baru disimpan
                idx = self.playlist_combo.findData(self.current_playlist_id)
                if idx >= 0:
                    self.playlist_combo.setCurrentIndex(idx)
                    self._load_current_playlist()
            else:
                _show(self, tr("gagal_title"), r["msg"], "error")

    def _toggle_lyrics(self):
        if self.lyrics_toggle_btn.isChecked():
            self.lyrics_text.setVisible(True)
            # Muat lirik untuk lagu saat ini
            if self.current_index >= 0 and self.current_index < len(self.current_playlist):
                self._show_lyrics(self.current_playlist[self.current_index])
        else:
            self.lyrics_text.setVisible(False)

    # ===== FILE LOADING =====
    def _select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, tr("music_select_folder"))
        if folder:
            files = []
            for root, dirs, files_in_dir in os.walk(folder):
                for f in files_in_dir:
                    if f.lower().endswith(('.mp3', '.wav', '.flac', '.m4a', '.ogg')):
                        files.append(os.path.join(root, f))
            if files:
                # Tambahkan ke playlist aktif
                for f in files:
                    db.add_song_to_playlist(self.user_id, self.current_playlist_id, f)
                self._reload_playlists()

    def _update_playlist_widget(self):
        self.playlist_widget.clear()
        for path in self.current_playlist:
            self.playlist_widget.addItem(os.path.basename(path))
        # Highlight lagu yang sedang diputar
        if 0 <= self.current_index < len(self.current_playlist):
            self.playlist_widget.setCurrentRow(self.current_index)

    def _get_lyrics_from_file(self, filepath):
        """Baca lirik dari metadata file (ID3 tag USLT/SYLT)"""
        try:
            if filepath.lower().endswith('.mp3'):
                audio = MP3(filepath)
                # Cari tag USLT (Unsychronized lyrics) atau SYLT (Synchronized lyrics)
                for tag in audio.keys():
                    if tag.startswith('USLT') or tag.startswith('SYLT'):
                        return str(audio[tag])
            elif filepath.lower().endswith('.flac'):
                audio = FLAC(filepath)
                if 'lyrics' in audio:
                    return audio['lyrics'][0]
            elif filepath.lower().endswith('.m4a'):
                audio = MP4(filepath)
                # M4A biasanya pakai tag '©lyr'
                if '\xa9lyr' in audio:
                    return audio['\xa9lyr'][0]
            elif filepath.lower().endswith('.ogg'):
                audio = OggVorbis(filepath)
                if 'lyrics' in audio:
                    return audio['lyrics'][0]
        except Exception as e:
            print(f"Gagal baca lirik dari file: {e}")
        return None

    def _get_lyrics_from_api(self, artist, title):
        """Cari lirik dari API lyrics.ovh (gratis, tanpa API key)"""
        try:
            url = f"https://api.lyrics.ovh/v1/{artist}/{title}"
            response = requests.get(url, timeout=5)
            if response.status_code == 200:
                data = response.json()
                if 'lyrics' in data:
                    return data['lyrics']
        except Exception as e:
            print(f"Gagal ambil lirik dari API: {e}")
        return None

    def _extract_artist_title(self, filename):
        """Ekstrak artist dan title dari nama file (format: Artist - Title.mp3)"""
        name = os.path.splitext(filename)[0]
        if ' - ' in name:
            artist, title = name.split(' - ', 1)
            return artist.strip(), title.strip()
        # fallback: gunakan nama file sebagai title
        return "", name

    def _get_lyrics(self, filepath):
        """Utama: coba dari file, lalu dari API"""
        lyrics = self._get_lyrics_from_file(filepath)
        if lyrics:
            return lyrics

        # Coba dari API
        filename = os.path.basename(filepath)
        artist, title = self._extract_artist_title(filename)
        if artist and title:
            lyrics = self._get_lyrics_from_api(artist, title)
            if lyrics:
                return lyrics

        return None

    def _show_lyrics(self, filepath):
        """Tampilkan lirik di panel"""
        self.lyrics_text.clear()
        if not filepath:
            self.lyrics_text.setPlainText(tr("music_no_lyrics"))
            return

        lyrics = self._get_lyrics(filepath)
        if lyrics:
            self.lyrics_text.setPlainText(lyrics)
            self.lyrics_text.setStyleSheet(f"color: {_T('text')}; font-size: 13px;")
        else:
            self.lyrics_text.setPlainText(tr("music_lyrics_not_found"))
            self.lyrics_text.setStyleSheet(f"color: {_T('muted')}; font-size: 13px; font-style: italic;")

    # ===== PLAYBACK =====
    def _play_selected(self, item):
        idx = self.playlist_widget.row(item)
        if 0 <= idx < len(self.current_playlist):
            self.current_index = idx
            self._play_current()

    def _play_current(self):
        if self.current_index < 0 or self.current_index >= len(self.current_playlist):
            return
        filepath = self.current_playlist[self.current_index]
        url = QUrl.fromLocalFile(filepath)
        self.player.setSource(url)
        self.player.play()
        self.play_btn.setText("⏸️")
        self.now_playing_label.setText(os.path.basename(filepath))
        self._update_playlist_widget()
        if self.lyrics_toggle_btn.isChecked():
            self._show_lyrics(filepath)

    def _play_pause(self):
        if self.player.isPlaying():
            self.player.pause()
            self.play_btn.setText("▶️")
        else:
            self.player.play()
            self.play_btn.setText("⏸️")

    def _next(self):
        if not self.current_playlist:
            return
        if self.shuffle_mode:
            import random
            new_idx = random.randint(0, len(self.current_playlist)-1)
            while new_idx == self.current_index and len(self.current_playlist) > 1:
                new_idx = random.randint(0, len(self.current_playlist)-1)
            self.current_index = new_idx
        else:
            self.current_index = (self.current_index + 1) % len(self.current_playlist)
        self._play_current()

    def _prev(self):
        if not self.current_playlist:
            return
        if self.shuffle_mode:
            import random
            self.current_index = random.randint(0, len(self.current_playlist)-1)
        else:
            self.current_index = (self.current_index - 1) % len(self.current_playlist)
        self._play_current()

    def _toggle_shuffle(self, checked):
        self.shuffle_mode = checked
        self._update_ui_style()

    def _toggle_repeat(self, checked):
        self.repeat_mode = checked
        self._update_ui_style()

    def _seek(self, pos):
        if self.player.isSeekable():
            self.player.setPosition(int(pos * self.player.duration() / 100))

    def _set_volume(self, val):
        self.audio_output.setVolume(val / 100.0)

    def _update_progress(self):
        if self.player.duration() > 0:
            progress = int(self.player.position() / self.player.duration() * 100)
            self.progress_slider.setValue(progress)
            self.current_time_label.setText(self._format_time(self.player.position()))
            self.total_time_label.setText(self._format_time(self.player.duration()))

    def _format_time(self, ms):
        s = ms // 1000
        m = s // 60
        s = s % 60
        return f"{m}:{s:02d}"

    def _on_position_changed(self, pos):
        pass

    def _on_media_status_changed(self, status):
        if status == QMediaPlayer.MediaStatus.EndOfMedia:
            if self.repeat_mode:
                self.player.setPosition(0)
                self.player.play()
            else:
                self._next()

    def _on_player_error(self, error):
        print(f"Player error: {error}")
        if error == QMediaPlayer.Error.FormatError:
            _show(self, tr("msg_error"), tr("music_format_error"), "error")
        elif error == QMediaPlayer.Error.ResourceError:
            _show(self, tr("msg_error"), tr("music_resource_error"), "error")

    def closeEvent(self, e):
        self.timer.stop()
        self.player.stop()
        AppState.unregister(self._reload_playlists)
        super().closeEvent(e)

# ══════════════════════════════════════════════════════════════════════════════
#  v1.3.0 TAHAP 2 — WeekdaySelector, PomodoroPage, HeatmapWidget,
#                   UndoToast, HabitTemplateDialog
# ══════════════════════════════════════════════════════════════════════════════
class WeekdaySelector(QWidget):
    """7 tombol hari (Sen–Min) untuk recurrence. Kosong = tiap hari."""

    _DAY_KEYS = ["day_mon", "day_tue", "day_wed", "day_thu",
                 "day_fri", "day_sat", "day_sun"]

    def __init__(self, preset: str = "", parent=None):
        super().__init__(parent)
        lay = QHBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0)
        lay.setSpacing(5)
        selected = db.parse_repeat_days(preset)
        self._btns = []
        for i, key in enumerate(self._DAY_KEYS):
            b = QPushButton(tr(key))
            b.setCheckable(True)
            b.setFixedSize(46, 32)
            b.setChecked(i in selected)
            b.setCursor(Qt.CursorShape.PointingHandCursor)
            lay.addWidget(b)
            self._btns.append(b)
        lay.addStretch()
        self.setStyleSheet(
            f"QPushButton {{ background: {_T('panel')}; color: {_T('text')};"
            f" border: 1px solid {_T('border')}; border-radius: 6px; font-size: 11px; }}"
            f"QPushButton:checked {{ background: {_T('accent')}; color: white;"
            f" border: 1px solid {_T('accent')}; font-weight: bold; }}")

    def get_days_str(self) -> str:
        return db.repeat_days_to_str(
            {i for i, b in enumerate(self._btns) if b.isChecked()})


# ══════════════════════════════════════════════════════════════════════════════
#  POMODORO PAGE  🍅  — Focus Mode dengan hadiah XP & Gold
# ══════════════════════════════════════════════════════════════════════════════
class PomodoroPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self._state = "idle"        # 'idle' | 'focus' | 'break'
        self._remaining = 0
        self._total = 0
        self._paused = False
        self._timer = QTimer(self)
        self._timer.setInterval(1000)
        self._timer.timeout.connect(self._tick)
        self._build()
        self.load()
        AppState.register_lang_cb(self.load)

    # ── UI ────────────────────────────────────────────────────────────────────
    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(20, 16, 20, 16)
        root.setSpacing(14)

        root.addWidget(_lbl(tr("pomodoro_title"), "section", 16, True))
        root.addWidget(_lbl(tr("pomodoro_subtitle"), "sub", 12))
        root.addWidget(_sep())

        body = QHBoxLayout()
        body.setSpacing(14)

        # ── Kartu kiri: TIMER ──
        timer_card = _card()
        tc = QVBoxLayout(timer_card)
        tc.setContentsMargins(24, 24, 24, 24)
        tc.setSpacing(14)

        self._state_lbl = _lbl(tr("pomodoro_state_idle"), size=14, bold=True)
        self._state_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        tc.addWidget(self._state_lbl)

        self._time_lbl = QLabel("25:00")
        self._time_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._time_lbl.setStyleSheet(
            f"font-size: 64px; font-weight: bold; color: {_T('light')};")
        tc.addWidget(self._time_lbl)

        self._progress = QProgressBar()
        self._progress.setRange(0, 100)
        self._progress.setValue(0)
        self._progress.setTextVisible(False)
        self._progress.setFixedHeight(8)
        tc.addWidget(self._progress)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)
        self._btn_start = _btn(tr("pomodoro_start"), "solid", self._start, 44)
        self._btn_pause = _btn(tr("pomodoro_pause"), "", self._pause_toggle, 44)
        self._btn_reset = _btn(tr("pomodoro_reset"), "", self._reset, 44)
        self._btn_giveup = _btn(tr("pomodoro_give_up"), "danger", self._give_up, 44)
        for b in (self._btn_start, self._btn_pause, self._btn_reset, self._btn_giveup):
            btn_row.addWidget(b)
        tc.addLayout(btn_row)
        tc.addStretch()
        body.addWidget(timer_card, 3)

        # ── Kolom kanan: SETTINGS + STATS ──
        right = QVBoxLayout()
        right.setSpacing(14)

        set_card = _card()
        sc = QVBoxLayout(set_card)
        sc.setContentsMargins(18, 18, 18, 18)
        sc.setSpacing(10)
        sc.addWidget(_lbl(tr("pomodoro_task_label"), "section", 12, True))
        self._task_input = _input(tr("pomodoro_task_placeholder"))
        sc.addWidget(self._task_input)

        dur_row = QHBoxLayout()
        dur_row.setSpacing(10)
        f_col = QVBoxLayout()
        f_col.addWidget(_lbl(tr("pomodoro_focus_label"), "sub", 11))
        self._focus_spin = QSpinBox()
        self._focus_spin.setRange(5, 120)
        self._focus_spin.setValue(25)
        self._focus_spin.setSuffix(" " + tr("pomodoro_minutes_unit"))
        f_col.addWidget(self._focus_spin)
        b_col = QVBoxLayout()
        b_col.addWidget(_lbl(tr("pomodoro_break_label"), "sub", 11))
        self._break_spin = QSpinBox()
        self._break_spin.setRange(1, 30)
        self._break_spin.setValue(5)
        self._break_spin.setSuffix(" " + tr("pomodoro_minutes_unit"))
        b_col.addWidget(self._break_spin)
        dur_row.addLayout(f_col)
        dur_row.addLayout(b_col)
        sc.addLayout(dur_row)
        right.addWidget(set_card)

        stat_card = _card()
        st = QVBoxLayout(stat_card)
        st.setContentsMargins(18, 18, 18, 18)
        st.setSpacing(8)
        self._stat_today = _lbl("", size=12)
        self._stat_total = _lbl("", size=12)
        st.addWidget(self._stat_today)
        st.addWidget(self._stat_total)
        st.addWidget(_sep())
        st.addWidget(_lbl(tr("pomodoro_recent"), "section", 12, True))
        self._recent_lbl = _lbl("", "sub", 11)
        self._recent_lbl.setWordWrap(True)
        st.addWidget(self._recent_lbl)
        right.addWidget(stat_card)
        right.addStretch()

        body.addLayout(right, 2)
        root.addLayout(body, 1)
        self._update_controls()

    # ── State machine ─────────────────────────────────────────────────────────
    def _start(self):
        SND.click()
        self._begin_phase("focus")
        self._timer.start()

    def _begin_phase(self, phase: str):
        minutes = self._focus_spin.value() if phase == "focus" else self._break_spin.value()
        self._state = phase
        self._total = minutes * 60
        self._remaining = self._total
        self._paused = False
        self._refresh_display()
        self._update_controls()

    def _pause_toggle(self):
        if self._state == "idle":
            return
        self._paused = not self._paused
        SND.click()
        self._update_controls()

    def _reset(self):
        self._timer.stop()
        self._state = "idle"
        self._paused = False
        self._remaining = self._focus_spin.value() * 60
        self._total = self._remaining
        self._refresh_display()
        self._update_controls()

    def _give_up(self):
        if self._state == "idle":
            return
        SND.click()
        self._reset()

    def _tick(self):
        if self._paused:
            return
        self._remaining -= 1
        if self._remaining <= 0:
            self._timer.stop()
            self._phase_finished()
        else:
            self._refresh_display()

    def _phase_finished(self):
        if self._state == "focus":
            minutes = self._focus_spin.value()
            task = self._task_input.text().strip()
            r = db.complete_pomodoro(self.user_id, minutes, task)
            SND.level_up()
            self.load()
            msg = tr("pomodoro_complete_msg",
                     xp=r.get("xp_gained", 0), gold=r.get("gold_gained", 0),
                     mins=minutes)
            if r.get("leveled_up"):
                celebrate_levelup(self)
                msg += f"\n🎉 {tr('level_up_msg', lvl=r['new_level'])}"
            _show(self, tr("pomodoro_complete_title"), msg, "success")
            # Otomatis lanjut ke istirahat
            self._begin_phase("break")
            self._timer.start()
        else:  # break selesai
            SND.notify()
            self._state = "idle"
            self._remaining = self._focus_spin.value() * 60
            self._total = self._remaining
            _show(self, tr("info_title"), tr("pomodoro_break_done"), "info")
        self._refresh_display()
        self._update_controls()

    # ── Display helpers ───────────────────────────────────────────────────────
    def _refresh_display(self):
        m, s = divmod(max(0, self._remaining), 60)
        self._time_lbl.setText(f"{m:02d}:{s:02d}")
        if self._total > 0:
            done = (self._total - self._remaining) / self._total
            self._progress.setValue(int(done * 100))
        if self._state == "focus":
            self._state_lbl.setText(tr("pomodoro_state_focus"))
            self._state_lbl.setStyleSheet(f"color: {_T('accent')}; font-size: 14px; font-weight: bold;")
        elif self._state == "break":
            self._state_lbl.setText(tr("pomodoro_state_break"))
            self._state_lbl.setStyleSheet("color: #7ac74c; font-size: 14px; font-weight: bold;")
        else:
            self._state_lbl.setText(tr("pomodoro_state_idle"))
            self._state_lbl.setStyleSheet(f"color: {_T('sub')}; font-size: 14px;")

    def _update_controls(self):
        running = self._state != "idle"
        self._btn_start.setVisible(not running)
        self._btn_pause.setVisible(running)
        self._btn_pause.setText(tr("pomodoro_resume") if self._paused
                                else tr("pomodoro_pause"))
        self._btn_giveup.setVisible(running)
        self._btn_reset.setVisible(True)
        self._focus_spin.setEnabled(not running)
        self._break_spin.setEnabled(not running)
        self._task_input.setEnabled(not running)

    def load(self):
        """Refresh statistik (dipanggil juga saat bahasa berganti)."""
        try:
            s = db.get_pomodoro_stats(self.user_id)
            self._stat_today.setText(
                f"📅 {tr('pomodoro_today')}: "
                + tr("pomodoro_stat_sessions", n=s["today_sessions"])
                + " · " + tr("pomodoro_stat_minutes", n=s["today_minutes"]))
            self._stat_total.setText(
                f"🏆 {tr('pomodoro_total')}: "
                + tr("pomodoro_stat_sessions", n=s["total_sessions"])
                + " · " + tr("pomodoro_stat_minutes", n=s["total_minutes"]))
            recent = db.get_recent_pomodoros(self.user_id, 5)
            if recent:
                lines = []
                for it in recent:
                    when = (it.get("completed_at") or "")[5:16].replace("T", " ")
                    name = it.get("task_name") or "—"
                    lines.append(f"🍅 {it['duration_minutes']}′ {name} · {when}")
                self._recent_lbl.setText("\n".join(lines))
            else:
                self._recent_lbl.setText(tr("pomodoro_no_recent"))
        except Exception as e:
            log.error(f"Pomodoro load gagal: {e}")

    def closeEvent(self, e):
        self._timer.stop()
        super().closeEvent(e)


# ══════════════════════════════════════════════════════════════════════════════
#  HEATMAP WIDGET  —  Peta aktivitas ala GitHub contribution graph
# ══════════════════════════════════════════════════════════════════════════════
class HeatmapWidget(QWidget):
    """Grid minggu×7 hari; intensitas warna = jumlah aktivitas sukses."""

    def __init__(self, weeks=17, cell=13, gap=3, parent=None):
        super().__init__(parent)
        self.weeks = weeks
        self.cell = cell
        self.gap = gap
        self._data = {}
        self._cells = []          # [(QRect, date_str, count)]
        self.setMouseTracking(True)
        w = weeks * (cell + gap) + gap
        h = 7 * (cell + gap) + gap
        self.setFixedSize(w, h)

    def set_data(self, data: dict):
        self._data = data or {}
        self.update()

    def set_compact(self, compact: bool):
        """Mode ringkas: sel lebih kecil (dipakai konfigurasi widget dashboard)."""
        self.cell = 8 if compact else 13
        self.gap = 2 if compact else 3
        w = self.weeks * (self.cell + self.gap) + self.gap
        h = 7 * (self.cell + self.gap) + self.gap
        self.setFixedSize(w, h)
        self.update()

    def _colors(self):
        return [
            _T("panel"),     # 0
            "#1d5c2e",       # 1–2
            "#2e8b44",       # 3–5
            "#45b45e",       # 6–9
            "#7dde8a",       # 10+
        ]

    def _color_for(self, n: int) -> str:
        if n <= 0:
            return self._colors()[0]
        if n <= 2:
            return self._colors()[1]
        if n <= 5:
            return self._colors()[2]
        if n <= 9:
            return self._colors()[3]
        return self._colors()[4]

    def paintEvent(self, event):
        from PyQt6.QtCore import QRectF
        p = QPainter(self)
        p.setRenderHint(QPainter.RenderHint.Antialiasing)
        self._cells = []
        today = date.today()
        start = today - timedelta(days=self.weeks * 7 - 1)
        start -= timedelta(days=start.weekday())   # mulai dari hari Senin
        border = QColor(_T("border"))
        for col in range(self.weeks):
            for row in range(7):
                d = start + timedelta(days=col * 7 + row)
                if d > today:
                    continue
                iso = d.isoformat()
                n = self._data.get(iso, 0)
                x = self.gap + col * (self.cell + self.gap)
                y = self.gap + row * (self.cell + self.gap)
                rect = QRectF(x, y, self.cell, self.cell)
                if n <= 0:
                    p.setPen(border)
                else:
                    p.setPen(Qt.PenStyle.NoPen)
                p.setBrush(QColor(self._color_for(n)))
                p.drawRoundedRect(rect, 2.5, 2.5)
                self._cells.append((rect, iso, n))
        p.end()

    def mouseMoveEvent(self, event):
        pos = event.position()
        for rect, iso, n in self._cells:
            if rect.contains(pos):
                key = "heatmap_tooltip" if n > 0 else "heatmap_tooltip_zero"
                self.setToolTip(tr(key, date=iso, n=n))
                return
        self.setToolTip("")


# ══════════════════════════════════════════════════════════════════════════════
#  UNDO TOAST  —  popup kecil 10 detik dengan tombol Urungkan
# ══════════════════════════════════════════════════════════════════════════════
class UndoToast(QFrame):
    TTL_SECONDS = 10

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("undotoast")
        self.setStyleSheet(
            f"QFrame#undotoast {{ background: {_T('panel')};"
            f" border: 1px solid {_T('accent')}; border-radius: 10px; }}"
            f"QFrame#undotoast QLabel {{ background: transparent; border: none; }}")
        lay = QHBoxLayout(self)
        lay.setContentsMargins(14, 10, 14, 10)
        lay.setSpacing(10)
        self._msg = QLabel("")
        self._count = QLabel("")
        self._count.setStyleSheet(f"color: {_T('sub')}; font-size: 11px;")
        undo = QPushButton(tr("undo_btn"))
        undo.setCursor(Qt.CursorShape.PointingHandCursor)
        undo.setStyleSheet(
            f"QPushButton {{ background: {_T('accent')}; color: white; border: none;"
            f" border-radius: 6px; padding: 6px 12px; font-weight: bold; }}"
            f"QPushButton:hover {{ background: {_T('light')}; }}")
        undo.clicked.connect(self._do_undo)
        close = QPushButton("✕")
        close.setFixedSize(24, 24)
        close.setCursor(Qt.CursorShape.PointingHandCursor)
        close.setStyleSheet(
            f"QPushButton {{ background: transparent; color: {_T('sub')}; border: none; }}"
            f"QPushButton:hover {{ color: {_T('text')}; }}")
        close.clicked.connect(self._expire)
        lay.addWidget(self._msg)
        lay.addWidget(self._count)
        lay.addWidget(undo)
        lay.addWidget(close)

        self._on_undo = None
        self._remaining = 0
        self._timer = QTimer(self)
        self._timer.setInterval(1000)
        self._timer.timeout.connect(self._tick)

    def show_message(self, text: str, on_undo):
        self._on_undo = on_undo
        self._msg.setText(text)
        self._remaining = self.TTL_SECONDS
        self._count.setText(f"{self._remaining}s")
        self.adjustSize()
        self._reposition()
        self.show()
        self.raise_()
        self._timer.start()

    def _reposition(self):
        if self.parent():
            x = (self.parent().width() - self.width()) // 2
            y = self.parent().height() - self.height() - 28
            self.move(max(8, x), max(8, y))

    def _tick(self):
        self._remaining -= 1
        if self._remaining <= 0:
            self._expire()
        else:
            self._count.setText(f"{self._remaining}s")

    def _do_undo(self):
        self._timer.stop()
        self.hide()
        if callable(self._on_undo):
            self._on_undo()

    def _expire(self):
        self._timer.stop()
        self.hide()


# ══════════════════════════════════════════════════════════════════════════════
#  HABIT TEMPLATE DIALOG  —  paket habit siap pakai
# ══════════════════════════════════════════════════════════════════════════════
class HabitTemplateDialog(QDialog):
    def __init__(self, user_id: int, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        self.setWindowTitle(tr("template_title"))
        self.setMinimumWidth(540)
        self.setStyleSheet(build_ss())

        root = QVBoxLayout(self)
        root.setContentsMargins(24, 24, 24, 24)
        root.setSpacing(12)
        root.addWidget(_lbl(tr("template_title"), "section", 15, True))
        root.addWidget(_lbl(tr("template_subtitle"), "sub", 12))
        root.addWidget(_sep())

        u = db.get_user(user_id)
        lang = u.get("language", "id")

        content = QWidget()
        lay = QVBoxLayout(content)
        lay.setSpacing(10)
        lay.setContentsMargins(0, 0, 0, 0)

        for t in db.get_habit_templates(lang):
            card = _card()
            row = QHBoxLayout(card)
            row.setContentsMargins(16, 14, 16, 14)
            row.setSpacing(14)
            icon = _emoji_label(t["icon"], ICON_CARD)
            row.addWidget(icon)
            col = QVBoxLayout()
            col.setSpacing(3)
            name = _lbl(t["name"], size=13, bold=True)
            desc = _lbl(t["desc"], "sub", 11)
            desc.setWordWrap(True)
            cnt = _lbl(tr("template_count", n=t["count"]), "sub", 11)
            col.addWidget(name)
            col.addWidget(desc)
            col.addWidget(cnt)
            row.addLayout(col, 1)
            apply_btn = _btn(tr("template_apply"), "solid", h=36)
            apply_btn.setMinimumWidth(110)
            apply_btn.clicked.connect(partial(self._apply, t))
            row.addWidget(apply_btn)
            lay.addWidget(card)
        lay.addStretch()

        sa = _scrolled(content)
        root.addWidget(sa, 1)

    def _apply(self, template: dict):
        n = db.apply_habit_template(self.user_id, template["key"])
        SND.level_up()
        _show(self, tr("berhasil_title"),
              tr("template_applied", n=n, name=template["name"]), "success")
        self.accept()


# ══════════════════════════════════════════════════════════════════════════════
#  CRAFTING PAGE  🔨  — gabungkan equipment jadi item langka
# ══════════════════════════════════════════════════════════════════════════════
class CraftingPage(QWidget):
    def __init__(self, user_id: int):
        super().__init__()
        self.user_id = user_id
        self._build()
        AppState.register(self.load)
        AppState.register_lang_cb(self.load)

    def _build(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(20, 16, 20, 16)
        root.setSpacing(12)
        root.addWidget(_lbl(tr("crafting_title"), "section", 15, True))
        root.addWidget(_lbl(tr("crafting_subtitle"), "sub", 12))
        root.addWidget(_sep())

        self._inner = QWidget()
        self._lay = QVBoxLayout(self._inner)
        self._lay.setSpacing(12)
        self._lay.setContentsMargins(0, 0, 0, 0)
        root.addWidget(_scrolled(self._inner), 1)
        self.load()

    def load(self):
        if not AppState.user_id:
            return
        while self._lay.count():
            it = self._lay.takeAt(0)
            if it.widget():
                it.widget().deleteLater()

        for r in db.get_crafting_recipes():
            chk = db.can_craft(self.user_id, r["id"])
            owned_already = any(i["item_id"] == r["output_id"]
                                for i in db.get_inventory(self.user_id))

            card = _card()
            lay = QVBoxLayout(card)
            lay.setContentsMargins(18, 16, 18, 16)
            lay.setSpacing(8)

            # Header: icon + nama output + buff
            head = QHBoxLayout()
            head.addWidget(_emoji_label(r["output_icon"], ICON_CARD))
            col = QVBoxLayout()
            col.setSpacing(2)
            col.addWidget(_lbl(r["output_name"], size=14, bold=True))
            buff = _lbl(r["output_buff"], size=11)
            buff.setStyleSheet(f"color: {_T('accent')};")
            col.addWidget(buff)
            u = db.get_user(self.user_id)
            lang = u.get("language", "id")
            desc = _lbl(r["desc"][0] if lang == "id" else r["desc"][1], "sub", 11)
            col.addWidget(desc)
            head.addLayout(col, 1)
            lay.addLayout(head)

            # Materials
            mat_row = QHBoxLayout()
            mat_row.setSpacing(10)
            mat_row.addWidget(_lbl(tr("crafting_needs"), "sub", 11))
            inv_ids = {i["item_id"] for i in db.get_inventory(self.user_id)}
            for m in r["inputs"]:
                have = m["id"] in inv_ids
                tag = tr("crafting_have_tag") if have else tr("crafting_missing_tag")
                mlbl = _lbl(f"{tag} {m['icon']} {m['name']}", size=11)
                mlbl.setStyleSheet(
                    f"color: {_T('light') if have else '#e05050'};")
                mat_row.addWidget(mlbl)
            mat_row.addStretch()
            lay.addLayout(mat_row)

            # Gold + tombol
            foot = QHBoxLayout()
            gl = _lbl(tr("crafting_gold_cost", gold=r["gold"]), size=11)
            gl.setStyleSheet(
                f"color: {'#f0c040' if chk['gold_ok'] else '#e05050'};")
            foot.addWidget(gl)
            if not chk["gold_ok"]:
                foot.addWidget(_lbl(
                    tr("crafting_gold_short", have=int(chk["have_gold"]),
                       need=r["gold"]), "sub", 10))
            foot.addStretch()
            if owned_already:
                done = _lbl(tr("crafting_owned"), size=11, bold=True)
                done.setStyleSheet(f"color: {_T('light')};")
                foot.addWidget(done)
            else:
                btn = _btn(tr("crafting_btn"), "solid", h=36)
                btn.setMinimumWidth(130)
                btn.setEnabled(chk["ok"])
                btn.clicked.connect(partial(self._craft, r))
                foot.addWidget(btn)
            lay.addLayout(foot)
            self._lay.addWidget(card)
        self._lay.addStretch()

    def _craft(self, recipe):
        r = db.craft_item(self.user_id, recipe["id"])
        if r["ok"]:
            SND.level_up()
            _show(self, tr("berhasil_title"), r["msg"], "success")
        else:
            SND.error()
            _show(self, tr("gagal_title"), r["msg"], "error")
        AppState.refresh()
        self.load()


# ══════════════════════════════════════════════════════════════════════════════
#  🔍 COMMAND PALETTE (Ctrl+K) — lompat ke halaman / aksi cepat
# ══════════════════════════════════════════════════════════════════════════════
class CommandPalette(QDialog):
    """Palette ala VSCode: ketik untuk mencari halaman/aksi, Enter untuk buka."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.result_kind = None          # "page" | "action"
        self.result_key = None
        self.setWindowTitle(tr("palette_title"))
        self.setMinimumSize(540, 420)
        self.setStyleSheet(build_ss())
        lay = QVBoxLayout(self)
        lay.setContentsMargins(14, 14, 14, 12)
        lay.setSpacing(8)

        self._search = _input(tr("palette_placeholder"))
        self._search.textChanged.connect(self._refresh_list)
        lay.addWidget(self._search)

        self._list = QListWidget()
        self._list.itemActivated.connect(self._pick)
        lay.addWidget(self._list, 1)

        footer = _lbl(tr("palette_footer"), "sub", 9)
        footer.setWordWrap(True)
        lay.addWidget(footer)

        # Entri: (kind, key, label)
        self._pages = []
        for tkey, pkey in NavBar._TABS:
            icon = NavBar.ICON_MAP.get(pkey, "📄")
            self._pages.append(("page", pkey, f"{icon} {tr(tkey)}"))
        self._acts = [
            ("action", "habit", tr("palette_action_add_habit")),
            ("action", "daily", tr("palette_action_add_daily")),
            ("action", "todo",  tr("palette_action_add_todo")),
        ]
        self._refresh_list("")

    def _add_header(self, text: str):
        it = QListWidgetItem(text)
        it.setFlags(Qt.ItemFlag.NoItemFlags)
        it.setForeground(QColor(_T("muted")))
        self._list.addItem(it)

    def _add_entry(self, entry):
        it = QListWidgetItem(entry[2])
        it.setData(Qt.ItemDataRole.UserRole, (entry[0], entry[1]))
        self._list.addItem(it)

    def _select_first(self):
        for i in range(self._list.count()):
            if self._list.item(i).flags() & Qt.ItemFlag.ItemIsSelectable:
                self._list.setCurrentRow(i)
                return

    def _refresh_list(self, text: str):
        q = (text or "").strip().lower()
        self._list.clear()
        pages = [e for e in self._pages if not q or q in e[2].lower()]
        acts = [e for e in self._acts if not q or q in e[2].lower()]
        if pages:
            self._add_header(tr("palette_section_pages"))
            for e in pages:
                self._add_entry(e)
        if acts:
            self._add_header(tr("palette_section_actions"))
            for e in acts:
                self._add_entry(e)
        if self._list.count() == 0:
            it = QListWidgetItem(tr("palette_no_results"))
            it.setFlags(Qt.ItemFlag.NoItemFlags)
            self._list.addItem(it)
        else:
            self._select_first()

    def _pick(self, item):
        data = item.data(Qt.ItemDataRole.UserRole)
        if not data:
            return
        self.result_kind, self.result_key = data
        self.accept()

    def keyPressEvent(self, e):
        if e.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            item = self._list.currentItem()
            if item is not None:
                self._pick(item)
            return
        if e.key() == Qt.Key.Key_Down:
            self._list.setFocus()
            return
        super().keyPressEvent(e)


# ══════════════════════════════════════════════════════════════════════════════
#  ➕ QUICK ADD DIALOG — tambah habit/daily/quest kilat (tray / Ctrl+N)
# ══════════════════════════════════════════════════════════════════════════════
class QuickAddDialog(QDialog):
    def __init__(self, user_id: int, parent=None, preset_type="habit"):
        super().__init__(parent)
        self.user_id = user_id
        self.setWindowTitle(tr("quickadd_title"))
        self.setMinimumWidth(400)
        self.setStyleSheet(build_ss())
        lay = QVBoxLayout(self)
        lay.setSpacing(10)

        row = QHBoxLayout()
        row.addWidget(_lbl(tr("quickadd_type"), size=12))
        self._type = QComboBox()
        self._type.addItem("⛏ Habit", "habit")
        self._type.addItem("📅 Daily", "daily")
        self._type.addItem("📜 Quest", "todo")
        self._type.setCurrentIndex({"habit": 0, "daily": 1, "todo": 2}.get(preset_type, 0))
        row.addWidget(self._type, 1)
        lay.addLayout(row)

        self._name = _input(tr("quickadd_name_ph"))
        self._name.returnPressed.connect(self._add)
        lay.addWidget(self._name)

        btns = QHBoxLayout()
        btns.addStretch()
        btns.addWidget(_btn(tr("msg_cancel"), "", self.reject))
        btns.addWidget(_btn(tr("quickadd_add"), "solid", self._add))
        lay.addLayout(btns)
        self._name.setFocus()

    def _add(self):
        name = self._name.text().strip()
        if not name:
            SND.error()
            self._name.setFocus()
            return
        mode = self._type.currentData()
        if mode == "habit":
            r = db.add_habit(self.user_id, name)
        elif mode == "daily":
            r = db.add_daily(self.user_id, name)
        else:
            r = db.add_todo(self.user_id, name)
        if isinstance(r, dict) and not r.get("ok", True):
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", ""), "error")
            return
        SND.complete()
        self.created_name = name
        self.accept()


# ══════════════════════════════════════════════════════════════════════════════
#  🧭 ONBOARDING WIZARD — 4 layar untuk user baru (sekali saja)
# ══════════════════════════════════════════════════════════════════════════════
class OnboardingWizard(QDialog):
    _SUGGESTIONS = {
        "id": [("💧", "Minum 8 gelas air"), ("📚", "Baca buku 10 menit"),
               ("🏃", "Olahraga 15 menit")],
        "en": [("💧", "Drink 8 glasses of water"), ("📚", "Read 10 minutes"),
               ("🏃", "Exercise 15 minutes")],
    }

    def __init__(self, user_id: int, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        u = db.get_user(user_id) or {}
        self._name = u.get("display_name") or u.get("username", "")
        self.setWindowTitle(tr("onboard_welcome_title"))
        self.setMinimumSize(580, 440)
        self.setStyleSheet(build_ss())

        root = QVBoxLayout(self)
        root.setContentsMargins(20, 18, 20, 14)
        root.setSpacing(12)

        self._stack = QStackedWidget()
        root.addWidget(self._stack, 1)

        # ── Halaman 0: Sambutan ──
        p0 = QWidget()
        v0 = QVBoxLayout(p0)
        v0.setSpacing(12)
        t0 = _lbl(tr("onboard_welcome_title"), "section", 17, True)
        t0.setAlignment(Qt.AlignmentFlag.AlignCenter)
        b0 = _lbl(tr("onboard_welcome_body").replace("**", "").replace("\\n", "\n"),
                  size=12)
        b0.setWordWrap(True)
        b0.setAlignment(Qt.AlignmentFlag.AlignCenter)
        v0.addStretch()
        v0.addWidget(_emoji_label("⛏️", ICON_HERO))
        v0.addWidget(t0)
        v0.addWidget(b0)
        v0.addStretch()
        self._stack.addWidget(p0)

        # ── Halaman 1: Pilih class ──
        p1 = QWidget()
        v1 = QVBoxLayout(p1)
        v1.setSpacing(8)
        t1 = _lbl(tr("onboard_class_title"), "section", 15, True)
        b1 = _lbl(tr("onboard_class_body"), "sub", 11)
        b1.setWordWrap(True)
        v1.addWidget(t1)
        v1.addWidget(b1)
        self._class_radios = {}
        cur_cls = u.get("avatar_class", "warrior")
        emojis = {"warrior": "⚔️", "mage": "🧙", "archer": "🏹",
                  "healer": "❤️", "rogue": "🗡️"}
        for cls in ("warrior", "mage", "archer", "healer", "rogue"):
            nm = tr(f"class_{cls}_name")
            bonus = tr(f"class_{cls}_bonus")
            rb = QRadioButton(f"{emojis.get(cls, '')} {nm} — {bonus}")
            rb.setStyleSheet(f"font-size:12px; color:{_T('text')};")
            rb.setChecked(cls == cur_cls)
            self._class_radios[cls] = rb
            v1.addWidget(rb)
        v1.addStretch()
        self._stack.addWidget(p1)

        # ── Halaman 2: 3 habit pertama ──
        p2 = QWidget()
        v2 = QVBoxLayout(p2)
        v2.setSpacing(8)
        t2 = _lbl(tr("onboard_habits_title"), "section", 15, True)
        b2 = _lbl(tr("onboard_habits_body"), "sub", 11)
        b2.setWordWrap(True)
        v2.addWidget(t2)
        v2.addWidget(b2)
        self._habit_rows = []
        lang = AppState.get_language()
        for icon, sugg in self._SUGGESTIONS.get(lang, self._SUGGESTIONS["en"]):
            le = _input("")
            le.setText(sugg)
            self._habit_rows.append((icon, le))
            v2.addWidget(le)
        v2.addStretch()
        self._stack.addWidget(p2)

        # ── Halaman 3: Selesai ──
        p3 = QWidget()
        v3 = QVBoxLayout(p3)
        v3.setSpacing(12)
        t3 = _lbl(tr("onboard_done_title"), "section", 17, True)
        t3.setAlignment(Qt.AlignmentFlag.AlignCenter)
        b3 = _lbl(tr("onboard_done_body", name=self._name).replace("\\n", "\n"), size=12)
        b3.setWordWrap(True)
        b3.setAlignment(Qt.AlignmentFlag.AlignCenter)
        v3.addStretch()
        v3.addWidget(_emoji_label("🎉", ICON_HERO))
        v3.addWidget(t3)
        v3.addWidget(b3)
        v3.addStretch()
        self._stack.addWidget(p3)

        # ── Tombol navigasi ──
        nav = QHBoxLayout()
        self._skip_btn = _btn(tr("onboard_skip"), "", self._skip)
        self._back_btn = _btn(tr("onboard_back"), "", self._back)
        self._next_btn = _btn(tr("onboard_next"), "solid", self._next)
        self._back_btn.setVisible(False)
        nav.addWidget(self._skip_btn)
        nav.addStretch()
        nav.addWidget(self._back_btn)
        nav.addWidget(self._next_btn)
        root.addLayout(nav)

        self._step = 0

    def _refresh_nav(self):
        self._back_btn.setVisible(self._step in (1, 2))
        self._skip_btn.setVisible(self._step in (0, 1))
        if self._step == 2:
            self._next_btn.setText(tr("onboard_finish"))
        elif self._step == 3:
            self._next_btn.setText(tr("wrapped_close"))
            self._back_btn.setVisible(False)
            self._skip_btn.setVisible(False)
        else:
            self._next_btn.setText(tr("onboard_next"))

    def _back(self):
        self._step = max(0, self._step - 1)
        self._stack.setCurrentIndex(self._step)
        self._refresh_nav()

    def _next(self):
        if self._step == 2:
            self._finish_setup()
        elif self._step == 3:
            self.accept()
            return
        self._step = min(3, self._step + 1)
        self._stack.setCurrentIndex(self._step)
        self._refresh_nav()

    def _selected_class(self):
        for cls, rb in self._class_radios.items():
            if rb.isChecked():
                return cls
        return None

    def _finish_setup(self):
        uid = self.user_id
        # 1. Terapkan class (abaikan jika gagal / tidak berubah)
        try:
            cls = self._selected_class()
            u = db.get_user(uid) or {}
            if cls and cls != u.get("avatar_class"):
                db.change_class(uid, cls)
        except Exception as e:
            log.warning(f"Onboarding set class gagal: {e}")
        # 2. Buat habit dari input yang terisi
        try:
            for icon, le in self._habit_rows:
                name = le.text().strip()
                if name:
                    db.add_habit(uid, name, icon=icon)
        except Exception as e:
            log.warning(f"Onboarding buat habit gagal: {e}")
        # 3. Bonus awal petualangan 💰
        try:
            db.gain_xp_gold(uid, 0, 25)
        except Exception:
            pass

    def _skip(self):
        db.mark_onboarding_done(self.user_id)
        self.accept()

    def accept(self):
        # Selesai / lewati → tandai onboarding selesai agar tidak muncul lagi
        try:
            db.mark_onboarding_done(self.user_id)
        except Exception:
            pass
        super().accept()

    def reject(self):
        try:
            db.mark_onboarding_done(self.user_id)
        except Exception:
            pass
        super().reject()


# ══════════════════════════════════════════════════════════════════════════════
#  😴 HEALTH ↔ PRODUCTIVITY CHART — matplotlib tertanam di Qt
# ══════════════════════════════════════════════════════════════════════════════
class HealthChartWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._lay = QVBoxLayout(self)
        self._lay.setContentsMargins(0, 0, 0, 0)
        self._lay.setSpacing(4)
        self.setMinimumHeight(190)

    def set_compact(self, compact: bool):
        self.setMinimumHeight(140 if compact else 190)

    def _clear(self):
        while self._lay.count():
            it = self._lay.takeAt(0)
            if it.widget():
                it.widget().deleteLater()

    def refresh(self, data: dict):
        self._clear()
        enough = data.get("days_with_sleep_data", 0) >= 3
        if not MPL_QT_OK or not enough:
            lbl = _lbl(tr("healthchart_need_data"), "sub", 11)
            lbl.setWordWrap(True)
            self._lay.addWidget(lbl)
            return

        series = data.get("series", [])
        r = data.get("correlation", 0.0)
        xs = list(range(len(series)))
        sleeps = [s for _, s, _ in series]
        tasks = [t for _, _, t in series]
        dates = [d for d, _, _ in series]

        fig = Figure(figsize=(6, 2.3), tight_layout=True)
        fig.patch.set_alpha(0.0)
        ax1 = fig.add_subplot(111)
        ax1.set_facecolor("none")
        ax2 = ax1.twinx()
        ax2.set_facecolor("none")

        sleep_c = _T("accent")
        task_c = _T("glow") if _T("glow") != "#888" else "#f5b133"
        muted = _T("muted")

        ax1.bar(xs, sleeps, width=0.7, color=sleep_c, alpha=0.55, label="😴")
        ax2.plot(xs, tasks, color=task_c, linewidth=1.8, marker="o",
                 markersize=2.5, label="✅")

        ax1.set_ylim(0, max(9, max(sleeps, default=0) + 1))
        ax2.set_ylim(0, max(5, max(tasks, default=0) + 1))
        tick_idx = list(range(0, len(xs), 7)) + [len(xs) - 1]
        tick_idx = sorted(set(i for i in tick_idx if 0 <= i < len(xs)))
        ax1.set_xticks(tick_idx)
        ax1.set_xticklabels([dates[i][5:] for i in tick_idx], fontsize=7,
                            color=muted)
        for ax in (ax1, ax2):
            ax.tick_params(colors=muted, labelsize=7)
            for spine in ax.spines.values():
                spine.set_visible(False)

        canvas = FigureCanvas(fig)
        canvas.setStyleSheet("background: transparent;")
        self._lay.addWidget(canvas)

        # Verdict korelasi
        if r >= 0.3:
            verdict = tr("healthchart_verdict_pos")
        elif r <= -0.3:
            verdict = tr("healthchart_verdict_neg")
        else:
            verdict = tr("healthchart_verdict_neutral")
        lbl = _lbl(tr("healthchart_corr", r=f"{r:+.2f}", verdict=verdict),
                   "sub", 10)
        lbl.setWordWrap(True)
        self._lay.addWidget(lbl)


# ══════════════════════════════════════════════════════════════════════════════
#  🎁 YEAR WRAPPED DIALOG — ringkasan setahun ala Spotify Wrapped
# ══════════════════════════════════════════════════════════════════════════════
class YearWrappedDialog(QDialog):
    def __init__(self, user_id: int, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        w = db.get_year_wrapped(user_id)
        u = AppState.user() or {}
        name = u.get("display_name") or u.get("username", "")
        self.setWindowTitle(tr("wrapped_title", year=w["year"]))
        self.setMinimumSize(520, 460)
        self.setStyleSheet(build_ss())

        root = QVBoxLayout(self)
        root.setContentsMargins(20, 18, 20, 16)
        root.setSpacing(12)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")
        inner = QWidget()
        lay = QVBoxLayout(inner)
        lay.setSpacing(10)
        scroll.setWidget(inner)
        root.addWidget(scroll, 1)

        if not w.get("total_done"):
            hero = _emoji_label("🎁", ICON_HERO)
            msg = _lbl(tr("wrapped_empty"), "sub", 13)
            msg.setWordWrap(True)
            msg.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lay.addStretch()
            lay.addWidget(hero)
            lay.addWidget(msg)
            lay.addStretch()
        else:
            hero = _lbl(tr("wrapped_hero", name=name, year=w["year"]),
                        "section", 15, True)
            hero.setWordWrap(True)
            lay.addWidget(hero)

            def stat(text, big=False):
                l = _lbl(text, size=13 if not big else 15, bold=big)
                l.setWordWrap(True)
                return l

            card1 = _card()
            c1 = QVBoxLayout(card1)
            c1.addWidget(stat(tr("wrapped_total", n=w["total_done"]), big=True))
            c1.addWidget(stat(tr("wrapped_active_days", n=w["active_days"])))
            if w.get("best_day"):
                c1.addWidget(stat(tr("wrapped_best", date=w["best_day"],
                                     n=w["best_day_count"])))
            lay.addWidget(card1)

            if w.get("top_habits"):
                card2 = _card()
                c2 = QVBoxLayout(card2)
                c2.addWidget(stat(tr("wrapped_habits_top")))
                for h in w["top_habits"]:
                    c2.addWidget(stat(f"   {h.get('icon', '⛏')} {h.get('name', '?')}  ×{h.get('count', 0)}"))
                lay.addWidget(card2)

            card3 = _card()
            c3 = QVBoxLayout(card3)
            c3.addWidget(stat(tr("wrapped_focus", sessions=w["focus_sessions"],
                                 mins=w["focus_minutes"])))
            c3.addWidget(stat(tr("wrapped_economy",
                                 income=f"{w['income']:,.0f}",
                                 expense=f"{w['expense']:,.0f}")))
            c3.addWidget(stat(tr("wrapped_level", lvl=w["level"],
                                 streak=w["longest_streak"])))
            lay.addWidget(card3)
            lay.addStretch()

        close = _btn(tr("wrapped_close"), "solid", self.accept, 40)
        root.addWidget(close)


# ══════════════════════════════════════════════════════════════════════════════
#  ✨ PARTICLE OVERLAY — konfeti perayaan level-up
# ══════════════════════════════════════════════════════════════════════════════
class ParticleOverlay(QWidget):
    COLORS = ("#ffd166", "#7c5cff", "#22d3ee", "#ff6b6b", "#80c000", "#f5b133")

    def __init__(self, parent):
        super().__init__(parent)
        self.setAttribute(Qt.WidgetAttribute.WA_TransparentForMouseEvents, True)
        self.setStyleSheet("background: transparent;")
        self.setGeometry(parent.rect())
        w = max(1, self.width())
        h = max(1, self.height())
        self._parts = []
        for _ in range(70):
            self._parts.append({
                "x": random.uniform(0, w),
                "y": h + random.uniform(0, 30),
                "vx": random.uniform(-1.6, 1.6),
                "vy": random.uniform(-8.5, -4.0),
                "size": random.uniform(3, 7),
                "color": random.choice(self.COLORS),
                "life": 1.0,
                "decay": random.uniform(0.012, 0.028),
            })
        self._timer = QTimer(self)
        self._timer.timeout.connect(self._tick)
        self._timer.start(28)
        self.show()
        self.raise_()

    def _tick(self):
        alive = False
        for p in self._parts:
            p["x"] += p["vx"]
            p["y"] += p["vy"]
            p["vy"] += 0.22            # gravitasi
            p["life"] -= p["decay"]
            if p["life"] > 0:
                alive = True
        if not alive:
            self._timer.stop()
            self.hide()
            self.deleteLater()
            return
        self.update()

    def paintEvent(self, event):
        qp = QPainter(self)
        for p in self._parts:
            if p["life"] <= 0:
                continue
            c = QColor(p["color"])
            c.setAlphaF(max(0.0, min(1.0, p["life"])))
            qp.setPen(Qt.PenStyle.NoPen)
            qp.setBrush(c)
            s = int(p["size"])
            qp.drawEllipse(int(p["x"]), int(p["y"]), s, s)
        qp.end()


def celebrate_levelup(anchor: QWidget):
    """Ledakkan konfeti di atas jendela utama (panggil saat level-up)."""
    try:
        win = anchor.window() if anchor is not None else None
        if win is None:
            return
        ParticleOverlay(win)
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════════════════════
#  🌳 TALENT TREE DIALOG — pohon talent per class (3 tier × 2 node)
# ══════════════════════════════════════════════════════════════════════════════
class TalentTreeDialog(QDialog):
    def __init__(self, user_id: int, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        u = db.get_user(user_id) or {}
        cls = u.get("avatar_class", "warrior")
        self.setWindowTitle(tr("talents_title", **{"class": cls.title()}))
        self.setMinimumSize(660, 500)
        self.setStyleSheet(build_ss())
        root = QVBoxLayout(self)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")
        self._container = QWidget()
        self._vbox = QVBoxLayout(self._container)
        self._vbox.setSpacing(10)
        scroll.setWidget(self._container)
        root.addWidget(scroll, 1)
        root.addWidget(_btn(tr("wrapped_close"), "", self.accept, 36))
        self._reload()

    def _reload(self):
        while self._vbox.count():
            it = self._vbox.takeAt(0)
            if it.widget():
                it.widget().deleteLater()
            elif it.layout():
                sub = it.layout()
                while sub.count():
                    si = sub.takeAt(0)
                    if si.widget():
                        si.widget().deleteLater()
        st = db.get_talent_state(self.user_id)
        pts = _lbl(tr("talents_points", n=st["points"]), "sub", 12, True)
        self._vbox.addWidget(pts)
        for tier in (1, 2, 3):
            gb = QGroupBox(tr("talents_tier", n=tier,
                              lvl=tier * db.TALENT_POINT_PER_LEVELS))
            gl = QHBoxLayout(gb)
            gl.setSpacing(10)
            for node in st["tiers"][tier]:
                gl.addWidget(self._node_card(node))
            self._vbox.addWidget(gb)
        self._vbox.addStretch()

    def _node_card(self, node: dict) -> QFrame:
        card = _card()
        v = QVBoxLayout(card)
        v.setSpacing(4)
        head = _lbl(f"{node['icon']} {node['name']}", size=13, bold=True)
        head.setWordWrap(True)
        v.addWidget(head)
        desc = _lbl(node["desc"], "sub", 10)
        desc.setWordWrap(True)
        v.addWidget(desc)
        if node["unlocked"]:
            ok = _lbl(tr("talents_unlocked"), size=11, bold=True)
            ok.setStyleSheet(f"color:{_T('accent')};")
            v.addWidget(ok)
        else:
            chk = db.can_unlock_talent(self.user_id, node["key"])
            if chk.get("ok"):
                v.addWidget(_btn(tr("talents_unlock"), "gold",
                                 lambda: self._unlock(node["key"]), 32))
            else:
                reason = chk.get("reason")
                if reason == "level":
                    txt = tr("talents_locked_level", lvl=chk.get("need_level", 0))
                elif reason == "prereq":
                    txt = tr("talents_locked_prereq", t=node["tier"] - 1)
                elif reason == "points":
                    txt = tr("talents_no_points")
                else:
                    txt = "🔒"
                lbl = _lbl(txt, "sub", 10)
                lbl.setWordWrap(True)
                v.addWidget(lbl)
        return card

    def _unlock(self, key: str):
        r = db.unlock_talent(self.user_id, key)
        if r.get("ok"):
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            AppState.refresh()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", ""), "error")
        self._reload()


# ══════════════════════════════════════════════════════════════════════════════
#  👾 CUSTOM BOSS DIALOG — boss buatan user untuk guild
# ══════════════════════════════════════════════════════════════════════════════
class CustomBossDialog(QDialog):
    def __init__(self, user_id: int, guild_id, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        self.guild_id = guild_id
        self.created = False
        self.setWindowTitle(tr("cboss_title"))
        self.setMinimumWidth(420)
        self.setStyleSheet(build_ss())
        lay = QVBoxLayout(self)
        lay.setSpacing(10)

        lay.addWidget(_lbl(tr("cboss_name"), size=12))
        self._name = _input("👾 ...")
        lay.addWidget(self._name)

        row = QHBoxLayout()
        row.addWidget(_lbl(tr("cboss_icon"), size=12))
        self._icon = QComboBox()
        self._icon.addItems(db.CUSTOM_BOSS_ICONS)
        row.addWidget(self._icon, 1)
        lay.addLayout(row)

        form = QFormLayout()
        self._hp = QSpinBox()
        self._hp.setRange(100, 10000)
        self._hp.setValue(1000)
        self._hp.setSingleStep(100)
        self._atk = QSpinBox()
        self._atk.setRange(1, 150)
        self._atk.setValue(20)
        self._lvl = QSpinBox()
        self._lvl.setRange(1, 99)
        self._lvl.setValue(10)
        form.addRow(tr("cboss_hp"), self._hp)
        form.addRow(tr("cboss_atk"), self._atk)
        form.addRow(tr("cboss_minlvl"), self._lvl)
        lay.addLayout(form)

        lay.addWidget(_btn(tr("cboss_create"), "solid", self._create, 40))

    def _create(self):
        r = db.create_custom_boss(self.user_id, self.guild_id,
                                  self._name.text(), self._icon.currentText(),
                                  self._hp.value(), self._atk.value(),
                                  self._lvl.value())
        if r.get("ok"):
            self.created = True
            SND.complete()
            _show(self, tr("berhasil_title"), r["msg"], "success")
            self.accept()
        else:
            SND.error()
            _show(self, tr("gagal_title"), r.get("msg", ""), "error")


# ══════════════════════════════════════════════════════════════════════════════
#  ⚙️ DASHBOARD WIDGET DIALOG — visibilitas, urutan & kepadatan widget
# ══════════════════════════════════════════════════════════════════════════════
class DashboardWidgetDialog(QDialog):

    def __init__(self, user_id: int, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        self.setWindowTitle(tr("dashwidgets_title"))
        self.setMinimumWidth(460)
        self.setStyleSheet(build_ss())
        self._cfg = [dict(w) for w in db.get_dashboard_widgets(user_id)]
        lay = QVBoxLayout(self)
        lay.setSpacing(8)
        hint = _lbl(tr("dashwidgets_hint"), "sub", 11)
        hint.setWordWrap(True)
        lay.addWidget(hint)
        self._rows = QVBoxLayout()
        self._rows.setSpacing(6)
        lay.addLayout(self._rows)
        lay.addStretch()
        btnrow = QHBoxLayout()
        btnrow.addStretch()
        btnrow.addWidget(_btn(tr("msg_cancel"), "", self.reject))
        btnrow.addWidget(_btn(tr("msg_ok"), "solid", self._save))
        lay.addLayout(btnrow)
        self._render()

    def _render(self):
        while self._rows.count():
            it = self._rows.takeAt(0)
            if it.layout():
                sub = it.layout()
                while sub.count():
                    si = sub.takeAt(0)
                    if si.widget():
                        si.widget().deleteLater()
            elif it.widget():
                it.widget().deleteLater()
        n = len(self._cfg)
        for i, w in enumerate(self._cfg):
            row = QHBoxLayout()
            up = _btn("▲", h=28)
            up.setFixedWidth(34)
            up.setEnabled(i > 0)
            up.clicked.connect(lambda _, i=i: self._move(i, -1))
            dn = _btn("▼", h=28)
            dn.setFixedWidth(34)
            dn.setEnabled(i < n - 1)
            dn.clicked.connect(lambda _, i=i: self._move(i, 1))
            vis = QCheckBox(tr(f"widget_{w['key']}"))
            vis.setChecked(bool(w.get("visible", True)))
            vis.stateChanged.connect(
                lambda _s, i=i, cb=vis: self._set(i, "visible", cb.isChecked()))
            comp = QCheckBox(tr("widget_compact"))
            comp.setChecked(bool(w.get("compact", False)))
            comp.stateChanged.connect(
                lambda _s, i=i, cb=comp: self._set(i, "compact", cb.isChecked()))
            row.addWidget(up)
            row.addWidget(dn)
            row.addWidget(vis, 1)
            row.addWidget(comp)
            self._rows.addLayout(row)

    def _move(self, i: int, d: int):
        j = i + d
        self._cfg[i], self._cfg[j] = self._cfg[j], self._cfg[i]
        self._render()

    def _set(self, i: int, key: str, val):
        self._cfg[i][key] = val

    def _save(self):
        db.set_dashboard_widgets(self.user_id, self._cfg)
        self.accept()


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN WINDOW
# ══════════════════════════════════════════════════════════════════════════════
class MainWindow(QMainWindow):
    logout_signal = pyqtSignal()
    def __init__(self, user: dict):
        super().__init__()
        
        self.reminder_player = None
        self.reminder_audio_output = None
        self.reminder_beep_timer = None

        self.setWindowIcon(QIcon(get_icon_path('craftlife.ico')))
        AppState.set_user(user["id"])
        if db.is_account_locked(user["id"]):
            QTimer.singleShot(500, lambda: _show(
                self, tr("info_title"), tr("profile_account_locked"), "warning"
            ))
        QTimer.singleShot(300, lambda: TimeSync.sync())
        QTimer.singleShot(500, lambda: db.reset_daily_tasks(user["id"]))
        QTimer.singleShot(1200, self._post_login_popups)

        dn = user.get("display_name") or user.get("username", "")
        self.setWindowTitle(tr("app_title_format", name=dn))
        self.setMinimumSize(1080, 700)
        self.setStyleSheet(build_ss())
        self._pages: dict = {}
        self._build()
        AppState.register(self._topbar.refresh)

        self._timer = QTimer()
        self._timer.timeout.connect(self._topbar.refresh)
        self._timer.start(20000)
        self._reminder_timer = QTimer()
        self._reminder_timer.timeout.connect(self._check_reminders)
        self._reminder_timer.start(5000)  # cek setiap 5 detik
        self._setup_tray_icon()
        self._setup_shortcuts()
        self._pages["settings"].language_changed.connect(self._refresh_language_only)
        self._active_reminder_id = None
        self._pages["settings"].language_changed.connect(self._refresh_language_only)

    def _play_reminder_beep_loop(self, beep_func):
        """Looping beep menggunakan QTimer (setiap 2 detik)."""
        if self.reminder_beep_timer:
            self.reminder_beep_timer.stop()
            self.reminder_beep_timer.deleteLater()
            self.reminder_beep_timer = None
        self.reminder_beep_timer = QTimer()
        self.reminder_beep_timer.timeout.connect(beep_func)
        self.reminder_beep_timer.start(2000)  # 2 detik

    def _play_reminder_mp3_loop(self, filepath):
        """Looping MP3 menggunakan QMediaPlayer (ulang otomatis saat selesai)."""
        if not os.path.exists(filepath):
            return
        if self.reminder_player is None:
            self.reminder_player = QMediaPlayer()
            self.reminder_audio_output = QAudioOutput()
            self.reminder_player.setAudioOutput(self.reminder_audio_output)
        # Hapus koneksi sebelumnya jika ada
        try:
            self.reminder_player.mediaStatusChanged.disconnect()
        except Exception:
            pass
        # Fungsi untuk memutar ulang saat selesai
        def on_status_changed(status):
            if status == QMediaPlayer.MediaStatus.EndOfMedia:
                self.reminder_player.play()
        self.reminder_player.mediaStatusChanged.connect(on_status_changed)
        url = QUrl.fromLocalFile(filepath)
        self.reminder_player.setSource(url)
        self.reminder_audio_output.setVolume(1.0)
        self.reminder_player.play()

    def _stop_reminder_sounds(self):
        """Hentikan semua suara reminder (beep dan MP3)."""
        if self.reminder_beep_timer:
            self.reminder_beep_timer.stop()
            self.reminder_beep_timer.deleteLater()
            self.reminder_beep_timer = None
        if self.reminder_player:
            try:
                self.reminder_player.mediaStatusChanged.disconnect()
            except Exception:
                pass
            if self.reminder_player.isPlaying():
                self.reminder_player.stop()

    def load(self):
        """Called when language changes (if registered)"""
        self._refresh_language_only()

    def _refresh_language_only(self):
        for p in self._pages.values():
            if hasattr(p, "load"):
                p.load()
        if hasattr(self, '_topbar'):
            self._topbar.load()
        if hasattr(self, '_nav'):
            self._nav.load()

    def _refresh_language_only(self):
        # Hanya refresh teks tanpa reload semua halaman (biar cepat)
        for p in self._pages.values():
            if hasattr(p, "load"):
                p.load()
        # Update top bar
        if hasattr(self, '_topbar'):
            self._topbar.load()

    def _setup_tray_icon(self):
        """Buat System Tray Icon dengan menu konteks."""
        try:
            self.tray_icon = QSystemTrayIcon(self)
            self.tray_icon.setIcon(QIcon(get_icon_path('craftlife.ico')))
            self.tray_icon.setToolTip("⛏ CraftLife")
            
            tray_menu = QMenu()

            show_action = tray_menu.addAction(tr("restore_window"))
            show_action.triggered.connect(self.showNormal)

            quick_action = tray_menu.addAction(tr("tray_quick_add"))
            quick_action.triggered.connect(self._open_quick_add)

            tray_menu.addSeparator()

            quit_action = tray_menu.addAction(tr("exit_application"))
            quit_action.triggered.connect(QApplication.instance().quit)
            
            self.tray_icon.setContextMenu(tray_menu)
            self.tray_icon.show()
        except Exception as e:
            print(f"Tray icon tidak tersedia: {e}")

    # ── ⌨️ KEYBOARD SHORTCUTS ────────────────────────────────────────────────
    def _setup_shortcuts(self):
        """Ctrl+K palette · Ctrl+1-9 lompat halaman · Ctrl+N item baru · F5 muat ulang."""
        try:
            QShortcut(QKeySequence("Ctrl+K"), self, activated=self._open_palette)
            nav_keys = [k for _, k in NavBar._TABS[:9]]
            for i, key in enumerate(nav_keys):
                QShortcut(QKeySequence(f"Ctrl+{i + 1}"), self,
                          activated=lambda k=key: self._shortcut_goto(k))
            QShortcut(QKeySequence("Ctrl+N"), self,
                      activated=self._shortcut_new_item)
            for seq in ("F5", "Ctrl+R"):
                QShortcut(QKeySequence(seq), self, activated=self._shortcut_reload)
        except Exception as e:
            log.error(f"Shortcut gagal dipasang: {e}")

    def _shortcut_goto(self, key: str):
        try:
            self._nav._select(key)
        except Exception:
            pass
        self._switch(key)

    def _shortcut_new_item(self):
        """Ctrl+N: item baru sesuai konteks halaman aktif (habit/daily/todo)."""
        mode = "habit"
        cur = self._stack.currentWidget()
        if isinstance(cur, TaskPage) and cur.mode in ("habit", "daily", "todo"):
            mode = cur.mode
        QuickAddDialog(AppState.user_id, self, preset_type=mode).exec()
        AppState.refresh()

    def _shortcut_reload(self):
        page = self._stack.currentWidget()
        if hasattr(page, "load"):
            try:
                page.load()
            except Exception:
                pass

    def _open_palette(self):
        dlg = CommandPalette(self)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        if dlg.result_kind == "page":
            self._shortcut_goto(str(dlg.result_key))
        elif dlg.result_kind == "action":
            mode = str(dlg.result_key) if dlg.result_key in ("habit", "daily", "todo") else "habit"
            AddTaskDialog(mode, AppState.user_id, self).exec()
            AppState.refresh()

    def _open_quick_add(self):
        """Quick-add dari system tray: angkat jendela lalu tampilkan dialog."""
        try:
            self.showNormal()
            self.raise_()
            self.activateWindow()
        except Exception:
            pass
        QuickAddDialog(AppState.user_id, self).exec()
        AppState.refresh()

    def _post_login_popups(self):
        """Urutan popup setelah login: onboarding (sekali) → hadiah login harian."""
        try:
            u = db.get_user(AppState.user_id) or {}
            if not u.get("onboarding_done", 0):
                OnboardingWizard(AppState.user_id, self).exec()
                AppState.refresh()
        except Exception as e:
            log.error(f"Onboarding gagal: {e}")
        self._check_daily_login()

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_F11:
            if self.isFullScreen():
                self.showNormal()
            else:
                self.showFullScreen()
        super().keyPressEvent(event)

    def _build(self):
        central = QWidget()
        self.setCentralWidget(central)
        root = QVBoxLayout(central)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Top bar
        self._topbar = TopBar(self._show_notif, self._goto_profile, self._toggle_theme)
        root.addWidget(self._topbar)

        # Body
        body = QHBoxLayout()
        body.setContentsMargins(0, 0, 0, 0)
        body.setSpacing(0)

        self._nav = NavBar()
        self._nav.tab_changed.connect(self._switch)

        # Bungkus NavBar dengan scroll area
        nav_scroll = QScrollArea()
        nav_scroll.setWidgetResizable(True)
        nav_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        nav_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        nav_scroll.setWidget(self._nav)
        nav_scroll.setMinimumWidth(70)
        body.addWidget(nav_scroll)

        uid = AppState.user_id
        
        # ===== BUAT HANYA SETTINGS SECARA LANGSUNG (tidak lazy) =====
        # Settings diperlukan untuk koneksi sinyal sejak awal
        settings_page = SettingsPage(uid)
        self._pages = {"settings": settings_page}   # inisialisasi dengan settings
        self._stack = QStackedWidget()
        self._stack.addWidget(settings_page)
        self._pages["settings"].theme_changed.connect(self._retheme)

        # ===== LAZY LOAD UNTUK HALAMAN LAINNYA =====
        self._page_classes = {
            "dashboard":    lambda: DashboardPage(uid),
            "habits":       lambda: TaskPage(uid, "habit"),
            "dailies":      lambda: TaskPage(uid, "daily"),
            "todos":        lambda: TaskPage(uid, "todo"),
            "pomodoro":     lambda: PomodoroPage(uid),
            "sport":        lambda: SportTrackPage(uid),
            "crafting":     lambda: CraftingPage(uid),
            "economy":      lambda: EconomyPage(uid),
            "health_food":  lambda: HealthFoodPage(uid),
            "calendar":     lambda: CalendarPage(uid),
            "notes":        lambda: NotesPage(uid),
            "learning":     lambda: LearningPage(uid),
            "reminders":    lambda: RemindersPage(uid),
            "music":        lambda: MusicPage(uid),
            "shop":         lambda: ShopPage(uid),
            "pets":         lambda: PetsPage(uid),
            "friends":      lambda: FriendsPage(uid),
            "guild":        lambda: GuildPage(uid),
            "achievements": lambda: AchievementPage(uid),
            "profile":      lambda: ProfilePage(uid),
            "leaderboard":  lambda: LeaderboardPage(),
        }

        body.addWidget(self._stack, 1)
        root.addLayout(body, 1)

        # Halaman awal: dashboard
        self._switch("dashboard")

    def _switch(self, key: str):
        # Buat halaman jika belum ada
        if key not in self._pages:
            if key in self._page_classes:
                page = self._page_classes[key]()
                self._pages[key] = page
                self._stack.addWidget(page)
            else:
                return
        
        page = self._pages.get(key)
        if not page:
            return
        self._stack.setCurrentWidget(page)
        # Tampilkan & animasikan dulu (instan), baru load data di tick berikutnya
        # agar pindah tab responsif dan animasi tidak tersendat.
        if hasattr(page, "_inner"):
            fade_in(page._inner, 180)
        if hasattr(page, "load"):
            QTimer.singleShot(0, lambda: self._safe_load(page, key))

    def _safe_load(self, page, key):
        try:
            page.load()
        except Exception as e:
            import traceback
            err_text = traceback.format_exc()
            try:
                with open("crash.log", "a", encoding="utf-8") as f:
                    f.write(f"{datetime.now()} - Error loading {key}: {e}\n{err_text}\n")
            except Exception:
                pass
            _show(self, tr("msg_error"), tr("error_loading_page", page=key, error=e), "error")

    def _check_daily_login(self):
        """Klaim hadiah login harian (1×/hari) + banner event musiman."""
        try:
            r = db.claim_daily_login(AppState.user_id)
            if not r.get("claimed"):
                return
            SND.level_up()
            txt = tr("daily_login_msg", day=r["streak"], xp=r["xp"], gold=r["gold"])
            if r.get("item"):
                it = db.SHOP_ITEMS.get(r["item"], {})
                txt += "\n" + tr("daily_login_item_bonus",
                                 icon=it.get("icon", ""), name=it.get("name", ""))
            if r.get("leveled_up"):
                celebrate_levelup(self)
                txt += f"\n🎉 {tr('level_up_msg', lvl=r['new_level'])}"
            # Banner event musiman (boss & item limited)
            lang = AppState.get_language()
            for ev in db.get_active_seasonal_events():
                nm = ev["name"][0] if lang == "id" else ev["name"][1]
                txt += "\n\n" + tr("seasonal_active_banner",
                                   icon=ev["icon"], name=nm)
            _show(self, tr("daily_login_title"), txt, "success")
            AppState.refresh()
        except Exception as e:
            log.error(f"Daily login reward gagal: {e}")

    def show_undo_toast(self, name: str, trash_id):
        """Tampilkan toast 'Item dihapus' dengan tombol Urungkan (10 detik)."""
        if not trash_id:
            return
        if getattr(self, "_undo_toast", None) is None:
            self._undo_toast = UndoToast(self)
        uid = AppState.user_id

        def do_restore():
            r = db.restore_from_trash(uid, trash_id)
            if r.get("ok"):
                SND.complete()
                AppState.refresh()
            else:
                SND.error()
                _show(self, tr("gagal_title"), r.get("msg", ""), "error")

        self._undo_toast.show_message(tr("undo_item_deleted", name=name), do_restore)

    def _show_notif(self):
        NotifPopup(AppState.user_id, self).exec()

    def _goto_profile(self):
        self._nav._select("profile")
        self._switch("profile")

    def _retheme(self, show_msg=True):
        """Called when user changes theme in Settings — updates everything."""
        th  = db.get_user_theme(AppState.user_id)
        apply_theme(th)
        apply_accessibility(AppState.user())
        app = QApplication.instance()
        if app is not None:
            app.setFont(QFont("Segoe UI", max(6, round(10 * _FONT_SCALE))))
        SoundEngine.enabled = bool(
            AppState.user().get("sound_enabled", 1))

        new_ss = build_ss()
        QApplication.instance().setStyleSheet(new_ss)
        self.setStyleSheet(new_ss)

        self._topbar.retheme()
        self._nav.retheme()
        self._nav.reload_texts() 

        # Reload every page so colours update
        for p in self._pages.values():
            p.setStyleSheet(new_ss)
            if hasattr(p, "load"):
                try:
                    p.load()
                except Exception:
                    pass

        if show_msg:
            _show(self, tr("theme_changed_title"), tr("theme_change_success"), "success")

    def _toggle_theme(self):
        """Quick switch between Modern Dark / Modern Light from the top bar."""
        cur = (db.get_user(AppState.user_id) or {}).get("theme", "modern_dark")
        new_key = "modern_light" if cur == "modern_dark" else "modern_dark"
        db.set_user_theme(AppState.user_id, new_key)
        self._retheme(show_msg=False)

    def _play_mp3(self, filepath):
        try:
            if not os.path.exists(filepath):
                print(f"File MP3 tidak ditemukan: {filepath}")
                return
            
            if self.player is None:
                self.player = QMediaPlayer()
                self.audio_output = QAudioOutput()
                self.player.setAudioOutput(self.audio_output)
                
                # Tambahkan debug error
                self.player.errorOccurred.connect(self._handle_media_error)
            
            url = QUrl.fromLocalFile(filepath)
            self.player.setSource(url)
            self.audio_output.setVolume(1.0)
            self.player.play()
            print(f"Memutar MP3: {filepath}")
            
        except Exception as e:
            print(f"Error playing MP3: {e}")
            import traceback
            traceback.print_exc()

    def _handle_media_error(self, error):
        print(f"Media error: {error}")

    def _check_reminders(self):
        if not AppState.user_id:
            return
        reminders = db.get_pending_reminders(AppState.user_id)
        for r in reminders:
            self._show_reminder_alert(r)
            repeat_type = r.get("repeat_type", "none")
            if repeat_type != "none":
                next_dt = db.get_next_reminder_datetime(
                    r["reminder_datetime"],
                    repeat_type,
                    r.get("repeat_days", "")
                )
                if next_dt:
                    db.update_reminder(
                        r["id"],
                        AppState.user_id,
                        reminder_datetime=next_dt,
                        triggered=0,
                        is_active=1
                    )
                else:
                    db.update_reminder(
                        r["id"],
                        AppState.user_id,
                        is_active=0,
                        triggered=1
                    )
            else:
                db.mark_reminder_triggered(r["id"], AppState.user_id)

    def _show_reminder_alert(self, reminder):
        """Tampilkan dialog reminder dengan suara looping sampai OK diklik."""
        # Tampilkan notifikasi tray jika tersedia
        if hasattr(self, 'tray_icon') and self.tray_icon.isVisible():
            self.tray_icon.showMessage(
                tr("reminders_alert_title"),
                f"⏰ {reminder['title']}",
                QSystemTrayIcon.MessageIcon.Information,
                5000
            )

        # Mulai suara looping sesuai jenis
        sound_type = reminder["sound_type"]
        sound_file = reminder["sound_file"]

        if sound_type == "custom" and sound_file and os.path.exists(sound_file):
            self._play_reminder_mp3_loop(sound_file)
        else:
            # Tentukan fungsi beep
            if sound_type == "beep1":
                beep_func = lambda: SND._beep(600, 200)
            elif sound_type == "beep2":
                beep_func = lambda: SND._beep(400, 300)
            else:
                beep_func = lambda: SND.notify()
            self._play_reminder_beep_loop(beep_func)

        # Buat dialog
        dlg = QDialog(self)
        dlg.setWindowTitle(tr("reminders_alert_title"))
        dlg.setMinimumWidth(400)
        dlg.setMinimumHeight(180)
        dlg.setStyleSheet(build_ss())

        layout = QVBoxLayout(dlg)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(14)

        msg = f"⏰ {reminder['title']}\n{reminder['description'] or ''}\n\nWaktu: {reminder['reminder_datetime']}"
        lbl = QLabel(msg)
        lbl.setWordWrap(True)
        lbl.setStyleSheet(f"color: {_T('text')}; font-size:13px;")
        layout.addWidget(lbl)

        ok_btn = _btn(tr("msg_ok"), "solid", dlg.accept, 40)
        layout.addWidget(ok_btn)

        # Hentikan suara saat dialog ditutup (OK atau close)
        def on_finished():
            self._stop_reminder_sounds()
        dlg.finished.connect(on_finished)

        dlg.exec()

        # Tambahkan notifikasi ke database
        db.add_notification(
            AppState.user_id,
            db.tr_db(AppState.user_id, "reminders_notification", title=reminder['title']),
            "warning"
        )

    def _play_mp3_reminder(self, filepath):
        try:
            if not os.path.exists(filepath):
                print(f"File MP3 reminder tidak ditemukan: {filepath}")
                return
            
            if not hasattr(self, '_player') or self._player is None:
                self._player = QMediaPlayer()
                self._audio_output = QAudioOutput()
                self._player.setAudioOutput(self._audio_output)
                self._player.errorOccurred.connect(self._handle_media_error)
            
            url = QUrl.fromLocalFile(filepath)
            self._player.setSource(url)
            self._audio_output.setVolume(1.0)
            self._player.play()
            print(f"Memutar MP3 reminder: {filepath}")
            
        except Exception as e:
            print(f"Error playing MP3 reminder: {e}")

    def _handle_media_error(self, error):
        print(f"Media error: {error}")

    def _stop_reminder_sound(self):
        try:
            if hasattr(self, '_player') and self._player is not None:
                if self._player.isPlaying():
                    self._player.stop()
                    print("Reminder sound stopped")
        except Exception as e:
            print(f"Error stopping reminder sound: {e}")

    def closeEvent(self, e):
        # Hentikan semua timer
        if hasattr(self, '_timer') and self._timer.isActive():
            self._timer.stop()
        if hasattr(self, '_reminder_timer') and self._reminder_timer.isActive():
            self._reminder_timer.stop()
        
        # Hentikan suara reminder
        self._stop_reminder_sounds()
        
        # ── FORCE CHECKPOINT SEBELUM TUTUP ──
        db.stop_periodic_checkpoint()
        db.force_checkpoint()
        
        # ── (OPSIONAL) BUAT BACKUP OTOMATIS ──
        # db.backup_database()  # Buka komentar jika ingin backup setiap tutup
        
        AppState.unregister(self._topbar.refresh)
        super().closeEvent(e)


# ══════════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════
def main():
    setup_error_handling()
    try:
        db.init_db()
    except Exception:
        import traceback
        with open("init_db_error.log", "w") as f:
            traceback.print_exc(file=f)
        print("Error di init_db(), lihat init_db_error.log")
        input("Tekan Enter untuk keluar...")   # pause
        sys.exit(1)

    import os
    os.environ["QT_LOGGING_RULES"] = "*.debug=false;qt.qpa.*=false"

    app = QApplication(sys.argv)
    font = QFont("Segoe UI", 10)
    app.setFont(font)
    app.setApplicationName("CraftLife")
    app.setStyleSheet(build_ss())

    login = LoginWindow()
    main_win = None

    def on_login(user):
        nonlocal main_win
        AppState.set_user(user["id"])
        db.start_periodic_checkpoint(10)

        th = db.get_user_theme(user["id"])
        apply_theme(th)
        apply_accessibility(db.get_user(user["id"]) or {})
        app.setFont(QFont("Segoe UI", max(6, round(10 * _FONT_SCALE))))
        app.setStyleSheet(build_ss())
        main_win = MainWindow(user)
        main_win.logout_signal.connect(lambda: login.show()) 
        main_win.show()
        login.hide()

    def on_main_window_closed():
        nonlocal main_win
        if main_win is not None:
            main_win.deleteLater() 
        login.show()

    # ── Hubungkan sinyal login ──
    login.logged_in.connect(on_login)

    # ── Jalankan login dialog ──
    result = login.exec()
    if result == QDialog.DialogCode.Accepted:
        sys.exit(app.exec())
    else:
        sys.exit(0)


if __name__ == "__main__":
    main()