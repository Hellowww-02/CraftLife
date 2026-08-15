"""
CraftLife — database.py  v1.0
Works as PyInstaller .exe (Desktop)
"""

# ══════════════════════════════════════════════════════════════════════════════
# STRUKTUR FILE  (Table of Contents) — database.py
# ══════════════════════════════════════════════════════════════════════════════
# 1.  KONEKSI & MIGRASI DB : get_conn (pooled), init_db, _safe_alter, migrate_*
# 2.  AUTH & USER          : register_user, login_user, change_password, get_user,
#                            update_user, set_avatar, delete_account
# 3.  XP / GOLD / HP/SKILL : gain_xp_gold, lose_hp, penalize_*, recalculate_all_buffs,
#                            use_class_skill, apply_skill_effect
# 4.  TASKS                : habits, dailies, todos (CRUD + complete + folders + reorder)
# 5.  SPORT                : sport_activities (CRUD + complete + sport points)
# 6.  ECONOMY              : economy_items, debts, savings, investments, subscriptions
# 7.  SHOP / PETS          : SHOP_ITEMS, PETS_DATA, buy/use/sell, adopt/feed/train pet
# 8.  GUILD & BOSS         : create_guild, guild_members, BOSSES, start/attack_boss
# 9.  ACHIEVEMENTS         : check_achievements, claim, ACHIEVEMENTS_REBALANCED
# 10. HEALTH / FOOD        : food_logs, nutrition goals, water, health_logs, recipes
# 11. NOTES / REMINDERS    : notes (nested folders), reminders (repeat), calendar_notes
# 12. MUSIC / PLAYLIST     : playlists CRUD, song move/copy
# 13. LOCK PROFILE         : lock_account, unlock_account, is_account_locked
# 14. RANK                 : calculate_rank
# 15. THEME & SETTINGS     : THEMES (7 tema), get/set_user_theme, language, currency
# 16. BACKUP & CHECKPOINT  : force_checkpoint, backup_database, periodic checkpoint
# ══════════════════════════════════════════════════════════════════════════════

from food_data import DEFAULT_FOODS, get_food_name
from applog import get_logger

import functools
import hashlib
import json
import os
import secrets
import sqlite3
import sys
import threading
import time
import traceback
import uuid
from datetime import date, datetime, timedelta
from pathlib import Path

log = get_logger(__name__)

try:
    from PyQt6.QtWidgets import QApplication
except ImportError:
    QApplication = None

# ========== TERJEMAHAN UNTUK DATABASE ==========
def tr_db(user_id=None, key=None, lang=None, **kwargs):
    """Ambil terjemahan untuk user tertentu berdasarkan bahasa di database."""
    if user_id is not None:
        try:
            conn = get_conn()
            row = conn.execute("SELECT language FROM users WHERE id=?", (user_id,)).fetchone()
            conn.close()
            lang = row["language"] if row and row["language"] else "id"
        except:
            lang = "id"
    elif lang is None:
        lang = "id"  # default
    
    from translations import get_text
    text = get_text(key, lang)
    if kwargs:
        return text.format(**kwargs)
    return text

def log_crash(error_msg):
    """Catat error ke file crash.log di folder yang sama dengan database."""
    log_path = os.path.join(os.path.dirname(DB_PATH), "crash.log")
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(f"{datetime.now()} - ERROR: {error_msg}\n")
        traceback.print_exc(file=f)
        f.write("\n")

def retry_on_lock(func):
    """Coba ulang jika database locked (maks 3 kali)"""
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        for attempt in range(3):
            try:
                return func(*args, **kwargs)
            except sqlite3.OperationalError as e:
                if "database is locked" in str(e) and attempt < 2:
                    time.sleep(0.2 * (attempt + 1))
                    continue
                raise
        return func(*args, **kwargs)
    return wrapper

def local_now():
    try:
        import tzlocal
        tz = tzlocal.get_localzone()
        return datetime.now(tz)
    except ImportError:
        # fallback ke waktu lokal sistem
        return datetime.now()

# ── Path auto-detect (VSCode + .exe) ─────────────────────────────────────────

def log_db_error(e, query=""):
    with open("db_error.log", "a", encoding="utf-8") as f:
        f.write(f"{datetime.now()} - Error: {e}\nQuery: {query}\n\n")

def get_db_path():
    if getattr(sys, 'frozen', False):
        # Jika dijalankan sebagai .exe, simpan DB di %APPDATA%\CraftLife
        appdata = os.getenv('APPDATA')
        if not appdata:
            appdata = os.path.expanduser('~')
        app_dir = Path(appdata) / 'CraftLife'
        app_dir.mkdir(parents=True, exist_ok=True)
        return str(app_dir / 'craftlife.db')
    else:
        # Jika dijalankan sebagai skrip Python biasa, simpan di folder yang sama dengan skrip
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), 'craftlife.db')

DB_PATH = get_db_path()

import threading as _db_threading
_db_conn_local = _db_threading.local()

class _PooledConnection:
    """Wrapper koneksi: .close() jadi no-op supaya koneksi bisa di-reuse
    (connection pooling) -> jauh lebih cepat & bebas 'database is locked'.
    Akses lain didelegasikan ke koneksi asli."""
    def __init__(self, conn):
        object.__setattr__(self, "_conn", conn)
    def close(self):
        pass  # no-op: jangan benar-benar tutup
    def __getattr__(self, name):
        return getattr(object.__getattribute__(self, "_conn"), name)
    def __setattr__(self, name, value):
        if name == "_conn":
            object.__setattr__(self, name, value)
        else:
            setattr(object.__getattribute__(self, "_conn"), name, value)

def get_conn():
    """Koneksi DB yang di-reuse per thread (pooled). .close() adalah no-op."""
    c = getattr(_db_conn_local, "conn", None)
    if c is None:
        real = sqlite3.connect(DB_PATH, timeout=60.0, check_same_thread=False, cached_statements=256)
        real.row_factory = sqlite3.Row
        real.execute("PRAGMA foreign_keys = ON")
        real.execute("PRAGMA journal_mode = WAL")
        real.execute("PRAGMA busy_timeout = 60000")
        real.execute("PRAGMA synchronous = NORMAL")
        real.execute("PRAGMA cache_size = -20000")
        real.execute("PRAGMA wal_autocheckpoint = 0")
        c = _PooledConnection(real)
        _db_conn_local.conn = c
    return c


def reset_connection():
    """Tutup & buang koneksi pooled milik thread ini.

    Dipakai saat file DB berganti (mis. unit test yang membuat DB sementara
    baru); tanpa ini, koneksi lama masih menunjuk ke file DB sebelumnya.
    """
    global_conn = getattr(_db_conn_local, "conn", None)
    if global_conn is not None:
        try:
            object.__getattribute__(global_conn, "_conn").close()
        except Exception:
            pass
        _db_conn_local.conn = None

def _safe_alter(cur, table, col, defn):
    try:
        cur.execute(f"ALTER TABLE {table} ADD COLUMN {col} {defn}") 
    except Exception:
        pass
    
def update_guild(guild_id, **kwargs):
    conn = get_conn()
    fields = ", ".join(f"{k}=?" for k in kwargs)
    conn.execute(f"UPDATE guilds SET {fields} WHERE id=?", list(kwargs.values()) + [guild_id])
    conn.commit()
    conn.close()

def get_friend_profile_details(friend_id):
    """Data lengkap untuk dialog profil teman (lebih detail dari versi lama):
    dasar user, progres XP, guild, achievement (jumlah + 6 terbaru), statistik
    tugas, sport level, rebirth, dan tanggal bergabung."""
    u = get_user(friend_id)
    if not u:
        return None
    conn = get_conn()
    ach_total = conn.execute("SELECT COUNT(*) FROM achievements").fetchone()[0]
    ach_done = conn.execute(
        "SELECT COUNT(*) FROM user_achievements WHERE user_id=? AND unlocked_at IS NOT NULL",
        (friend_id,)
    ).fetchone()[0]
    latest = conn.execute("""
        SELECT a.name, a.icon, a.category, ua.unlocked_at
        FROM user_achievements ua JOIN achievements a ON a.id = ua.achievement_id
        WHERE ua.user_id=? AND ua.unlocked_at IS NOT NULL
        ORDER BY ua.unlocked_at DESC LIMIT 6
    """, (friend_id,)).fetchall()
    tasks_done = conn.execute(
        "SELECT COUNT(*) FROM task_history WHERE user_id=? AND action='success'",
        (friend_id,)
    ).fetchone()[0]
    pomodoro_min = conn.execute(
        "SELECT COALESCE(SUM(duration_minutes),0) FROM pomodoro_sessions WHERE user_id=?",
        (friend_id,)
    ).fetchone()[0]
    guild_name = None
    if u.get("guild_id"):
        g = conn.execute("SELECT name FROM guilds WHERE id=?", (u["guild_id"],)).fetchone()
        guild_name = g["name"] if g else None
    conn.close()

    level = u.get("level", 1) or 1
    xp = u.get("xp", 0) or 0
    return {
        "user": u,
        "level": level, "xp": xp, "xp_needed": level * 150,
        "total_xp_earned": u.get("total_xp_earned", 0) or 0,
        "sport_level": u.get("sport_level", 1) or 1,
        "rebirth_count": u.get("rebirth_count", 0) or 0,
        "selected_title": u.get("selected_title", "") or "",
        "avatar_class": u.get("avatar_class", "") or "",
        "guild_name": guild_name,
        "join_date": (u.get("created_at") or "")[:10],
        "achievements_done": ach_done, "achievements_total": ach_total,
        "latest_achievements": [dict(r) for r in latest],
        "tasks_done": tasks_done,
        "pomodoro_minutes": pomodoro_min,
        "stats": get_stats(friend_id),
    }


def get_leaderboard(limit=50):
    conn = get_conn()
    rows = conn.execute("""
        SELECT username, display_name, level, total_xp_earned, gold,
               COALESCE(selected_title, '') as selected_title,
               COALESCE(sport_level, 1) as sport_level,
               (SELECT COUNT(*) FROM user_pets WHERE user_id=users.id) as pet_count,
               COALESCE(rebirth_count, 0) as rebirth_count
        FROM users
        WHERE is_admin = 0
        ORDER BY level DESC, total_xp_earned DESC
        LIMIT ?
    """, (limit,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def migrate_sort_order():
    """Isi sort_order untuk data yang sudah ada berdasarkan created_at."""
    max_retries = 3
    for attempt in range(max_retries):
        try:
            conn = get_conn()
            tables = ["habits", "dailies", "todos", "sport_activities", "economy_items"]
            for table in tables:
                users = conn.execute(f"SELECT DISTINCT user_id FROM {table}").fetchall()
                for u in users:
                    uid = u["user_id"]
                    rows = conn.execute(f"""
                        SELECT id, folder_id, created_at FROM {table}
                        WHERE user_id = ?
                        ORDER BY COALESCE(folder_id, -1), created_at
                    """, (uid,)).fetchall()
                    current_folder = None
                    order = 0
                    for row in rows:
                        if row["folder_id"] != current_folder:
                            current_folder = row["folder_id"]
                            order = 0
                        order += 1
                        conn.execute(
                            f"UPDATE {table} SET sort_order = ? WHERE id = ?",
                            (order, row["id"])
                        )
            conn.commit()
            conn.close()
            return  # sukses
        except sqlite3.OperationalError as e:
            if "database is locked" in str(e) and attempt < max_retries - 1:
                time.sleep(1 * (attempt + 1))
                continue
            else:
                raise
    conn.close()

def init_db():
    conn = get_conn()
    c = conn.cursor()

    # ========== SCHEMA MIGRATION (selalu dijalankan) ==========
    # Matikan foreign key sementara agar migrasi aman
    c.execute("PRAGMA foreign_keys = OFF")

    # --- Buat semua tabel dengan struktur baru (IF NOT EXISTS) ---
    c.execute("""CREATE TABLE IF NOT EXISTS users(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL,
        display_name TEXT DEFAULT '',
        bio TEXT DEFAULT '',
        avatar_class TEXT DEFAULT 'warrior',
        avatar_color TEXT DEFAULT '#5a8a2e',
        avatar_emoji TEXT DEFAULT '⚔️',
        level INTEGER DEFAULT 1,
        xp INTEGER DEFAULT 0,
        hp INTEGER DEFAULT 50,
        max_hp INTEGER DEFAULT 50,
        mp INTEGER DEFAULT 30,
        max_mp INTEGER DEFAULT 30,
        gold REAL DEFAULT 0,
        gems INTEGER DEFAULT 10,
        longest_streak INTEGER DEFAULT 0,
        total_habits_done INTEGER DEFAULT 0,
        total_dailies_done INTEGER DEFAULT 0,
        total_todos_done INTEGER DEFAULT 0,
        total_xp_earned INTEGER DEFAULT 0,
        total_gold_earned REAL DEFAULT 0,
        boss_damage_bonus REAL DEFAULT 0,
        xp_multiplier REAL DEFAULT 1.0,
        gold_multiplier REAL DEFAULT 1.0,
        hp_damage_reduction REAL DEFAULT 0,
        has_revive INTEGER DEFAULT 0,
        mp_bonus INTEGER DEFAULT 0,
        guild_id INTEGER,
        last_login TEXT,
        theme TEXT DEFAULT 'modern_dark', 
        sound_enabled INTEGER DEFAULT 1,
        created_at TEXT DEFAULT(datetime('now'))
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS guilds(
        id INTEGER PRIMARY KEY,   -- tanpa AUTOINCREMENT
        name TEXT NOT NULL,
        description TEXT DEFAULT '',
        leader_id INTEGER,
        quest_id TEXT,
        created_at TEXT DEFAULT(datetime('now')),
        level INTEGER DEFAULT 1,
        exp INTEGER DEFAULT 0,
        buff_xp REAL DEFAULT 0,
        buff_gold REAL DEFAULT 0,
        buff_damage REAL DEFAULT 0
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS guild_members(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        guild_id INTEGER,
        user_id INTEGER,
        joined_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(guild_id) REFERENCES guilds(id),
        FOREIGN KEY(user_id) REFERENCES users(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS habits(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '⚔️',
        difficulty TEXT DEFAULT 'medium',
        xp_reward INTEGER DEFAULT 25,
        gold_reward REAL DEFAULT 5,
        positive INTEGER DEFAULT 1,
        negative INTEGER DEFAULT 0,
        counter_up INTEGER DEFAULT 0,
        counter_down INTEGER DEFAULT 0,
        streak INTEGER DEFAULT 0,
        done_today INTEGER DEFAULT 0,
        last_done TEXT,
        notes TEXT DEFAULT '',
        last_action TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS dailies(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '📅',
        difficulty TEXT DEFAULT 'medium',
        xp_reward INTEGER DEFAULT 30,
        gold_reward REAL DEFAULT 6,
        streak INTEGER DEFAULT 0,
        done_today INTEGER DEFAULT 0,
        last_done TEXT,
        notes TEXT DEFAULT '',
        last_action TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS todos(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '📜',
        priority TEXT DEFAULT 'medium',
        xp_reward INTEGER DEFAULT 40,
        gold_reward REAL DEFAULT 8,
        done INTEGER DEFAULT 0,
        due_date TEXT,
        notes TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS economy_items(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '💰',
        type TEXT NOT NULL,      -- 'income' atau 'expense'
        amount REAL NOT NULL,
        category TEXT DEFAULT 'other',
        date TEXT NOT NULL,      -- format YYYY-MM-DD
        notes TEXT DEFAULT '',
        folder_id INTEGER,
        created_at TEXT DEFAULT(datetime('now')),
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS inventory(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        item_id TEXT NOT NULL,
        item_type TEXT NOT NULL,
        quantity INTEGER DEFAULT 1,
        equipped INTEGER DEFAULT 0,
        obtained_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS user_pets(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        pet_id TEXT NOT NULL,
        is_active INTEGER DEFAULT 0,
        happiness INTEGER DEFAULT 50,
        level INTEGER DEFAULT 1,
        exp INTEGER DEFAULT 0,
        hunger INTEGER DEFAULT 100,
        last_fed TEXT,
        adopted_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS boss_battles(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        guild_id INTEGER,
        boss_id TEXT NOT NULL,
        boss_name TEXT NOT NULL,
        boss_icon TEXT DEFAULT '🐉',
        boss_tier TEXT DEFAULT 'normal',
        boss_hp REAL DEFAULT 100,
        boss_max_hp REAL DEFAULT 100,
        boss_attack REAL DEFAULT 5,
        status TEXT DEFAULT 'active',
        started_at TEXT DEFAULT(datetime('now')),
        ended_at TEXT,
        FOREIGN KEY(guild_id) REFERENCES guilds(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS boss_rewards(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        guild_id INTEGER NOT NULL,
        boss_name TEXT NOT NULL,
        boss_tier TEXT DEFAULT 'normal',
        xp_reward INTEGER DEFAULT 0,
        gold_reward REAL DEFAULT 0,
        is_claimed INTEGER DEFAULT 0,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY(guild_id) REFERENCES guilds(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS debts(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        amount REAL NOT NULL,
        due_date TEXT,
        is_paid INTEGER DEFAULT 0,
        notes TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        paid_at TEXT,
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    # Catatan Hutang (piutang): orang lain berhutang KE user.
    # Saat dibuat → auto-catat EXPENSE; saat lunas → auto-catat INCOME.
    # expense_item_id/income_item_id menunjuk entri economy_items terkait.
    c.execute("""CREATE TABLE IF NOT EXISTS debt_notes(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        person_name TEXT NOT NULL,
        amount REAL NOT NULL DEFAULT 0,
        date TEXT,
        notes TEXT DEFAULT '',
        status TEXT DEFAULT 'unpaid',
        expense_item_id INTEGER,
        income_item_id INTEGER,
        created_at TEXT DEFAULT(datetime('now')),
        paid_at TEXT,
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS activity_log(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        action TEXT NOT NULL,
        detail TEXT,
        xp_gained INTEGER DEFAULT 0,
        gold_gained REAL DEFAULT 0,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS notifications(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        message TEXT NOT NULL,
        type TEXT DEFAULT 'info',
        is_read INTEGER DEFAULT 0,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS guild_invites(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        guild_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        status TEXT DEFAULT 'pending',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(guild_id) REFERENCES guilds(id),
        FOREIGN KEY(user_id) REFERENCES users(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS guild_requests(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        guild_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        status TEXT DEFAULT 'pending',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(guild_id) REFERENCES guilds(id),
        FOREIGN KEY(user_id) REFERENCES users(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS friends(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        friend_id INTEGER NOT NULL,
        status TEXT DEFAULT 'pending',
        action_user_id INTEGER,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id),
        FOREIGN KEY(friend_id) REFERENCES users(id),
        UNIQUE(user_id, friend_id)
    )""")

    # Profile photos and consent-based couple relationships use the same local DB.
    c.execute("""CREATE TABLE IF NOT EXISTS user_profile_photos(
        user_id INTEGER PRIMARY KEY,
        image_data BLOB NOT NULL,
        mime_type TEXT NOT NULL,
        width INTEGER NOT NULL,
        height INTEGER NOT NULL,
        size_bytes INTEGER NOT NULL,
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS couple_relationships(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_a_id INTEGER NOT NULL,
        user_b_id INTEGER NOT NULL,
        requested_by INTEGER NOT NULL,
        status TEXT NOT NULL DEFAULT 'pending'
            CHECK(status IN ('pending','accepted','rejected','cancelled')),
        created_at TEXT DEFAULT(datetime('now')),
        responded_at TEXT,
        UNIQUE(user_a_id, user_b_id),
        CHECK(user_a_id < user_b_id),
        FOREIGN KEY(user_a_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY(user_b_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY(requested_by) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS love_spaces(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        couple_relationship_id INTEGER NOT NULL UNIQUE,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(couple_relationship_id) REFERENCES couple_relationships(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS love_space_members(
        love_space_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        role TEXT DEFAULT 'member',
        joined_at TEXT DEFAULT(datetime('now')),
        PRIMARY KEY(love_space_id, user_id),
        FOREIGN KEY(love_space_id) REFERENCES love_spaces(id) ON DELETE CASCADE,
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS love_space_photos(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        love_space_id INTEGER,
        owner_user_id INTEGER NOT NULL,
        visibility TEXT NOT NULL DEFAULT 'private'
            CHECK(visibility IN ('private','shared')),
        image_data BLOB NOT NULL,
        mime_type TEXT NOT NULL,
        width INTEGER NOT NULL,
        height INTEGER NOT NULL,
        size_bytes INTEGER NOT NULL,
        caption TEXT DEFAULT '',
        photo_date TEXT,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(love_space_id) REFERENCES love_spaces(id) ON DELETE CASCADE,
        FOREIGN KEY(owner_user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("CREATE INDEX IF NOT EXISTS idx_couple_status_a ON couple_relationships(user_a_id,status)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_couple_status_b ON couple_relationships(user_b_id,status)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_love_members_user ON love_space_members(user_id,love_space_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_love_photos_space ON love_space_photos(love_space_id,created_at)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_love_photos_owner ON love_space_photos(owner_user_id,created_at)")

    c.execute("""CREATE TABLE IF NOT EXISTS guild_leader_transfers(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        guild_id INTEGER NOT NULL,
        old_leader_id INTEGER NOT NULL,
        created_at TEXT DEFAULT(datetime('now')),
        status TEXT DEFAULT 'pending'
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS messages(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        sender_id INTEGER NOT NULL,
        receiver_id INTEGER NOT NULL,
        message TEXT NOT NULL,
        is_read INTEGER DEFAULT 0,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(sender_id) REFERENCES users(id),
        FOREIGN KEY(receiver_id) REFERENCES users(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS guild_messages(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        guild_id INTEGER NOT NULL,
        sender_id INTEGER NOT NULL,
        message TEXT NOT NULL,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(guild_id) REFERENCES guilds(id),
        FOREIGN KEY(sender_id) REFERENCES users(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS user_sessions(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        token_hash TEXT NOT NULL,
        created_at TEXT DEFAULT(datetime('now')),
        expires_at TEXT NOT NULL,
        FOREIGN KEY(user_id) REFERENCES users(id)
    )""")
    c.execute("CREATE INDEX IF NOT EXISTS idx_user_sessions_user ON user_sessions(user_id)")

    # Optional Supabase offline-first metadata. No credentials are stored here.
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_user_links(
        local_user_id INTEGER PRIMARY KEY,
        cloud_user_id TEXT NOT NULL UNIQUE,
        email TEXT NOT NULL,
        status TEXT DEFAULT 'linked',
        linked_at TEXT DEFAULT(datetime('now')),
        last_sync_at TEXT,
        FOREIGN KEY(local_user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS sync_queue(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        local_user_id INTEGER NOT NULL,
        entity_type TEXT NOT NULL,
        entity_local_id TEXT,
        operation TEXT NOT NULL,
        payload TEXT DEFAULT '{}',
        status TEXT DEFAULT 'pending',
        retry_count INTEGER DEFAULT 0,
        next_retry_at TEXT,
        last_error TEXT,
        created_at TEXT DEFAULT(datetime('now')),
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(local_user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_entity_map(
        local_user_id INTEGER NOT NULL,
        entity_type TEXT NOT NULL,
        local_id TEXT NOT NULL,
        cloud_id TEXT NOT NULL,
        updated_at TEXT DEFAULT(datetime('now')),
        PRIMARY KEY(local_user_id,entity_type,local_id),
        UNIQUE(entity_type,cloud_id),
        FOREIGN KEY(local_user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_conversations(
        cloud_id TEXT PRIMARY KEY,
        local_user_id INTEGER NOT NULL,
        other_local_user_id INTEGER NOT NULL,
        updated_at TEXT,
        unread_count INTEGER DEFAULT 0,
        last_message_at TEXT,
        last_message_preview TEXT DEFAULT '',
        FOREIGN KEY(local_user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY(other_local_user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_messages(
        cloud_id TEXT PRIMARY KEY,
        conversation_cloud_id TEXT NOT NULL,
        sender_cloud_id TEXT NOT NULL,
        client_message_id TEXT,
        reply_to_cloud_id TEXT,
        body TEXT NOT NULL,
        created_at TEXT NOT NULL,
        edited_at TEXT,
        deleted_at TEXT,
        sync_status TEXT DEFAULT 'synced'
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_message_reactions(
        message_cloud_id TEXT NOT NULL,
        user_cloud_id TEXT NOT NULL,
        reaction TEXT NOT NULL,
        updated_at TEXT,
        PRIMARY KEY(message_cloud_id,user_cloud_id)
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS local_message_reactions(
        message_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        reaction TEXT NOT NULL,
        updated_at TEXT DEFAULT(datetime('now')),
        PRIMARY KEY(message_id,user_id),
        FOREIGN KEY(message_id) REFERENCES messages(id) ON DELETE CASCADE,
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_guild_messages_cache(
        cloud_id TEXT PRIMARY KEY,
        guild_cloud_id TEXT NOT NULL,
        sender_cloud_id TEXT NOT NULL,
        client_message_id TEXT,
        reply_to_cloud_id TEXT,
        body TEXT NOT NULL,
        created_at TEXT NOT NULL,
        edited_at TEXT,
        deleted_at TEXT
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_guild_reactions_cache(
        message_cloud_id TEXT NOT NULL,
        user_cloud_id TEXT NOT NULL,
        reaction TEXT NOT NULL,
        updated_at TEXT,
        PRIMARY KEY(message_cloud_id,user_cloud_id)
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS chat_attachments_cache(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        cloud_id TEXT UNIQUE,
        local_user_id INTEGER NOT NULL,
        conversation_cloud_id TEXT,
        message_cloud_id TEXT,
        local_message_id INTEGER,
        uploader_cloud_id TEXT,
        storage_path TEXT,
        original_filename TEXT NOT NULL,
        mime_type TEXT NOT NULL,
        size_bytes INTEGER NOT NULL,
        width INTEGER,
        height INTEGER,
        sha256 TEXT NOT NULL,
        file_data BLOB,
        thumbnail_data BLOB,
        sync_status TEXT DEFAULT 'pending',
        created_at TEXT DEFAULT(datetime('now')),
        deleted_at TEXT,
        FOREIGN KEY(local_user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY(local_message_id) REFERENCES messages(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_presence_cache(
        cloud_user_id TEXT PRIMARY KEY,
        status TEXT DEFAULT 'offline',
        device_name TEXT DEFAULT '',
        last_seen_at TEXT,
        updated_at TEXT
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_notifications_cache(
        cloud_id TEXT PRIMARY KEY,
        local_user_id INTEGER NOT NULL,
        notification_type TEXT NOT NULL,
        actor_cloud_id TEXT,
        entity_type TEXT,
        entity_id TEXT,
        payload TEXT DEFAULT '{}',
        is_read INTEGER DEFAULT 0,
        created_at TEXT,
        FOREIGN KEY(local_user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    # Phase 3: stable non-secret device identity and conflict-safe personal snapshots.
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_device_state(
        local_user_id INTEGER PRIMARY KEY,
        device_id TEXT NOT NULL UNIQUE,
        device_name TEXT NOT NULL DEFAULT 'CraftLife Desktop',
        platform TEXT NOT NULL DEFAULT 'desktop',
        app_version TEXT NOT NULL DEFAULT '',
        registered_at TEXT,
        last_seen_at TEXT,
        FOREIGN KEY(local_user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS cloud_personal_sync_state(
        local_user_id INTEGER NOT NULL,
        document_key TEXT NOT NULL DEFAULT 'tracker_v1',
        last_local_hash TEXT,
        remote_revision INTEGER NOT NULL DEFAULT 0,
        remote_hash TEXT,
        remote_updated_at TEXT,
        conflict_status TEXT NOT NULL DEFAULT 'none',
        remote_payload TEXT,
        local_backup_payload TEXT,
        last_pushed_at TEXT,
        last_pulled_at TEXT,
        PRIMARY KEY(local_user_id,document_key),
        FOREIGN KEY(local_user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("CREATE INDEX IF NOT EXISTS idx_sync_queue_pending ON sync_queue(local_user_id,status,next_retry_at)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_cloud_messages_conversation ON cloud_messages(conversation_cloud_id,created_at)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_cloud_reactions_message ON cloud_message_reactions(message_cloud_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_local_reactions_message ON local_message_reactions(message_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_cloud_guild_messages_time ON cloud_guild_messages_cache(guild_cloud_id,created_at)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_chat_attachments_message_cloud ON chat_attachments_cache(message_cloud_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_chat_attachments_local_message ON chat_attachments_cache(local_message_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_chat_attachments_pending ON chat_attachments_cache(local_user_id,sync_status)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_cloud_notifications_unread ON cloud_notifications_cache(local_user_id,is_read,created_at)")

    c.execute("""CREATE TABLE IF NOT EXISTS pomodoro_sessions(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        task_name TEXT DEFAULT '',
        duration_minutes INTEGER DEFAULT 25,
        xp_earned INTEGER DEFAULT 0,
        gold_earned REAL DEFAULT 0,
        completed_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS trash_bin(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        item_type TEXT NOT NULL,
        item_name TEXT DEFAULT '',
        payload TEXT NOT NULL,
        deleted_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id)
    )""")
    c.execute("CREATE INDEX IF NOT EXISTS idx_trash_user ON trash_bin(user_id)")

    c.execute("""CREATE TABLE IF NOT EXISTS user_talents(
        user_id INTEGER NOT NULL,
        talent_key TEXT NOT NULL,
        unlocked_at TEXT DEFAULT(datetime('now')),
        PRIMARY KEY(user_id, talent_key),
        FOREIGN KEY(user_id) REFERENCES users(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS custom_bosses(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        guild_id INTEGER,
        creator_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '👾',
        hp INTEGER DEFAULT 500,
        atk INTEGER DEFAULT 10,
        min_level INTEGER DEFAULT 1,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(creator_id) REFERENCES users(id)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS pvp_challenges(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        challenger_id INTEGER NOT NULL,
        opponent_id INTEGER NOT NULL,
        start_date TEXT,
        end_date TEXT,
        status TEXT DEFAULT 'pending',
        winner_id INTEGER,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(challenger_id) REFERENCES users(id),
        FOREIGN KEY(opponent_id) REFERENCES users(id)
    )""")
    c.execute("CREATE INDEX IF NOT EXISTS idx_pvp_users ON pvp_challenges(challenger_id, opponent_id)")

    c.execute("""CREATE TABLE IF NOT EXISTS backup_codes(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        code_hash TEXT NOT NULL,
        is_used INTEGER DEFAULT 0,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS sport_activities(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        sport_type TEXT DEFAULT 'running',
        icon TEXT DEFAULT '🏃',
        difficulty TEXT DEFAULT 'medium',
        xp_reward INTEGER DEFAULT 25,
        gold_reward REAL DEFAULT 5,
        sport_points_reward INTEGER DEFAULT 15,
        done_today INTEGER DEFAULT 0,
        streak INTEGER DEFAULT 0,
        last_done TEXT,
        notes TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    # Log reps per aktivitas olahraga (push-up, pull-up, squat, dst.) —
    # menggerakkan sistem RANK per variasi yang dibuat user.
    c.execute("""CREATE TABLE IF NOT EXISTS sport_rep_logs(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        activity_id INTEGER NOT NULL,
        reps INTEGER NOT NULL DEFAULT 0,
        sets INTEGER DEFAULT 1,
        log_date TEXT NOT NULL,
        note TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY(activity_id) REFERENCES sport_activities(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS food_items(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '🍎',
        calories REAL DEFAULT 0,
        protein REAL DEFAULT 0,
        carbs REAL DEFAULT 0,
        fat REAL DEFAULT 0,
        is_custom INTEGER DEFAULT 0,
        user_id INTEGER,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS food_logs(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        food_id INTEGER NOT NULL,
        serving REAL DEFAULT 1.0,
        calories REAL DEFAULT 0,
        protein REAL DEFAULT 0,
        carbs REAL DEFAULT 0,
        fat REAL DEFAULT 0,
        meal_type TEXT DEFAULT 'snack',  -- breakfast, lunch, dinner, snack
        log_date TEXT NOT NULL,  -- YYYY-MM-DD
        notes TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY(food_id) REFERENCES food_items(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS user_nutrition_goals(
        user_id INTEGER PRIMARY KEY,
        daily_calories INTEGER DEFAULT 2000,
        daily_protein INTEGER DEFAULT 50,
        daily_carbs INTEGER DEFAULT 250,
        daily_fat INTEGER DEFAULT 70,
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS food_achievements(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        achievement_date TEXT NOT NULL,  -- YYYY-MM-DD
        bonus_claimed INTEGER DEFAULT 0,
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,
        UNIQUE(user_id, achievement_date)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS water_logs(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        amount_ml INTEGER DEFAULT 0,
        log_date TEXT NOT NULL,  -- YYYY-MM-DD
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS user_water_goals(
        user_id INTEGER PRIMARY KEY,
        daily_ml INTEGER DEFAULT 2000,
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS user_health_goals(
        user_id INTEGER PRIMARY KEY,
        daily_steps INTEGER DEFAULT 10000,
        daily_sleep_hours REAL DEFAULT 7.0,
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS recipes(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '🍲',
        serving_size REAL DEFAULT 1,
        notes TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS recipe_items(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        recipe_id INTEGER NOT NULL,
        food_id INTEGER NOT NULL,
        quantity REAL DEFAULT 1,
        FOREIGN KEY(recipe_id) REFERENCES recipes(id) ON DELETE CASCADE,
        FOREIGN KEY(food_id) REFERENCES food_items(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS savings(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '🏦',
        target_amount REAL DEFAULT 0,
        current_amount REAL DEFAULT 0,
        target_date TEXT,
        notes TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""
        CREATE TABLE IF NOT EXISTS health_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            log_date TEXT NOT NULL,
            steps INTEGER DEFAULT 0,
            sleep_hours REAL DEFAULT 0,
            water_ml INTEGER DEFAULT 0,
            weight_kg REAL DEFAULT 0,
            resting_hr INTEGER DEFAULT 0,
            stress_level TEXT DEFAULT 'normal',
            mood TEXT DEFAULT 'normal',
            notes TEXT DEFAULT '',
            created_at TEXT DEFAULT(datetime('now')),
            net_calories REAL DEFAULT 0,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,
            UNIQUE(user_id, log_date)
        )
    """)   

    c.execute("""
        CREATE TABLE IF NOT EXISTS achievements (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            description TEXT NOT NULL,
            icon TEXT DEFAULT '🏆',
            category TEXT NOT NULL,
            requirement_type TEXT NOT NULL,
            requirement_value INTEGER NOT NULL,
            xp_reward INTEGER DEFAULT 0,
            gold_reward REAL DEFAULT 0,
            is_hidden INTEGER DEFAULT 0
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS user_achievements (
            user_id INTEGER NOT NULL,
            achievement_id INTEGER NOT NULL,
            progress INTEGER DEFAULT 0,
            unlocked_at TEXT,
            claimed INTEGER DEFAULT 0,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,
            FOREIGN KEY(achievement_id) REFERENCES achievements(id) ON DELETE CASCADE,
            PRIMARY KEY(user_id, achievement_id)
        )
    """)

    c.execute("""CREATE TABLE IF NOT EXISTS investments(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '📈',
        amount REAL DEFAULT 0,
        invested_date TEXT DEFAULT(datetime('now')),
        last_update TEXT DEFAULT(datetime('now')),
        target_return REAL DEFAULT 0,
        notes TEXT DEFAULT '',
        is_active INTEGER DEFAULT 1,
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS subscriptions(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '📅',
        amount REAL DEFAULT 0,
        due_date TEXT NOT NULL,
        period TEXT DEFAULT 'monthly',
        is_recurring INTEGER DEFAULT 1,
        last_charged TEXT,
        notes TEXT DEFAULT '',
        is_active INTEGER DEFAULT 1,
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS redeem_codes(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        code TEXT UNIQUE NOT NULL,
        reward_type TEXT NOT NULL,   -- 'admin', 'xp', 'gold', 'item', etc.
        reward_value INTEGER DEFAULT 0,
        reward_item TEXT,            -- untuk item
        is_one_time INTEGER DEFAULT 1,
        used_by INTEGER,             -- user_id yang memakai (jika one-time)
        used_at TEXT,
        created_at TEXT DEFAULT(datetime('now'))
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS task_history(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        task_type TEXT NOT NULL,
        task_id INTEGER NOT NULL,
        action_date TEXT NOT NULL,
        action TEXT NOT NULL,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("CREATE INDEX IF NOT EXISTS idx_task_history_user_task_date ON task_history(user_id, task_type, task_id, action_date)")

    c.execute("""CREATE TABLE IF NOT EXISTS note_folders(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '📁',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS notes(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        folder_id INTEGER,
        title TEXT NOT NULL,
        content TEXT DEFAULT '',
        is_archived INTEGER DEFAULT 0,
        created_at TEXT DEFAULT(datetime('now')),
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY(folder_id) REFERENCES note_folders(id) ON DELETE SET NULL
    )""")

    # ========== REMINDERS ==========
    c.execute("""CREATE TABLE IF NOT EXISTS reminders(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        description TEXT,
        reminder_datetime TEXT NOT NULL,
        sound_type TEXT DEFAULT 'default',
        sound_file TEXT,
        is_active INTEGER DEFAULT 1,
        triggered INTEGER DEFAULT 0,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    c.execute("CREATE INDEX IF NOT EXISTS idx_reminders_user_datetime ON reminders(user_id, reminder_datetime)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_reminders_active ON reminders(user_id, is_active, triggered)")

    # ========== CALENDAR NOTES ==========
    c.execute("""CREATE TABLE IF NOT EXISTS calendar_notes(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        note_date TEXT NOT NULL,  -- YYYY-MM-DD
        note TEXT NOT NULL,
        created_at TEXT DEFAULT(datetime('now')),
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,
        UNIQUE(user_id, note_date)
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS playlists(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        tracks TEXT NOT NULL DEFAULT '[]',
        is_favorite INTEGER DEFAULT 0,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    # ========== LEARNING PAGE (NotebookLM) ==========
    c.execute("""CREATE TABLE IF NOT EXISTS learning_notebooks(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS learning_sources(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        notebook_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        type TEXT NOT NULL,
        title TEXT NOT NULL,
        path TEXT,
        content TEXT NOT NULL,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(notebook_id) REFERENCES learning_notebooks(id) ON DELETE CASCADE,
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS learning_chunks(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        source_id INTEGER NOT NULL,
        chunk_text TEXT NOT NULL,
        chunk_index INTEGER NOT NULL,
        FOREIGN KEY(source_id) REFERENCES learning_sources(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS learning_chats(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        notebook_id INTEGER NOT NULL,
        role TEXT NOT NULL,
        content TEXT NOT NULL,
        citations TEXT,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(notebook_id) REFERENCES learning_notebooks(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS learning_generations(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        notebook_id INTEGER NOT NULL,
        type TEXT NOT NULL,
        title TEXT NOT NULL,
        content TEXT NOT NULL,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(notebook_id) REFERENCES learning_notebooks(id) ON DELETE CASCADE
    )""")

    # ========== LOVE PAGE ==========
    c.execute("""CREATE TABLE IF NOT EXISTS relationship_profiles(
        user_id INTEGER PRIMARY KEY,
        partner_name TEXT DEFAULT '',
        partner_gender TEXT DEFAULT 'female',
        partner_age INTEGER DEFAULT 25,
        relationship_type TEXT DEFAULT 'dating',
        start_date TEXT,
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS relationship_events(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        event_date TEXT NOT NULL,
        category TEXT DEFAULT 'date',
        notes TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS relationship_memories(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        memory_date TEXT NOT NULL,
        notes TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS relationship_checkins(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        checkin_date TEXT NOT NULL,
        my_mood INTEGER DEFAULT 3,
        partner_mood INTEGER DEFAULT 3,
        connection_score INTEGER DEFAULT 3,
        note TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        UNIQUE(user_id, checkin_date),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS relationship_prompt_responses(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        prompt_key TEXT NOT NULL,
        category TEXT DEFAULT 'connection',
        prompt_text TEXT NOT NULL,
        my_answer TEXT DEFAULT '',
        partner_answer TEXT DEFAULT '',
        response_date TEXT DEFAULT(date('now')),
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS relationship_prompt_favorites(
        user_id INTEGER NOT NULL,
        prompt_key TEXT NOT NULL,
        created_at TEXT DEFAULT(datetime('now')),
        PRIMARY KEY(user_id, prompt_key),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS relationship_weekly_reviews(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        week_start TEXT NOT NULL,
        appreciation TEXT DEFAULT '',
        wins TEXT DEFAULT '',
        support_needed TEXT DEFAULT '',
        shared_intention TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        updated_at TEXT DEFAULT(datetime('now')),
        UNIQUE(user_id, week_start),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS relationship_bucket_items(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        category TEXT DEFAULT 'dream',
        target_date TEXT,
        is_done INTEGER DEFAULT 0,
        created_at TEXT DEFAULT(datetime('now')),
        completed_at TEXT,
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS menstrual_settings(
        user_id INTEGER PRIMARY KEY,
        tracked_person TEXT DEFAULT 'partner',
        last_period_start TEXT,
        cycle_length INTEGER DEFAULT 28,
        period_length INTEGER DEFAULT 5,
        updated_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS menstrual_cycles(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        start_date TEXT NOT NULL,
        end_date TEXT,
        notes TEXT DEFAULT '',
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    c.execute("CREATE INDEX IF NOT EXISTS idx_relationship_events_user_date ON relationship_events(user_id, event_date)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_relationship_memories_user_date ON relationship_memories(user_id, memory_date)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_relationship_checkins_user_date ON relationship_checkins(user_id, checkin_date)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_relationship_prompts_user_date ON relationship_prompt_responses(user_id, response_date)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_relationship_reviews_user_week ON relationship_weekly_reviews(user_id, week_start)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_menstrual_cycles_user_start ON menstrual_cycles(user_id, start_date)")

    # ========== MIGRASI PLAYLISTS KE SKEMA BARU ==========
    c.execute("PRAGMA table_info(playlists)")
    cols = [row[1] for row in c.fetchall()]
    
    if 'tracks' not in cols:
        log.info("Migrasi playlists ke skema baru...")
        
        # 1. Buat tabel baru dengan skema yang benar
        c.execute("""
            CREATE TABLE IF NOT EXISTS playlists_new (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                name TEXT NOT NULL,
                tracks TEXT NOT NULL DEFAULT '[]',
                is_favorite INTEGER DEFAULT 0,
                created_at TEXT DEFAULT(datetime('now'))
            )
        """)
        
        # 2. Copy data dari tabel lama
        if 'files' in cols:
            # Ada kolom files, ambil datanya
            c.execute("""
                INSERT INTO playlists_new (id, user_id, name, tracks, is_favorite, created_at)
                SELECT id, user_id, name, files, is_favorite, created_at FROM playlists
            """)
        else:
            # Tidak ada files, gunakan tracks kosong
            c.execute("""
                INSERT INTO playlists_new (id, user_id, name, tracks, is_favorite, created_at)
                SELECT id, user_id, name, '[]', is_favorite, created_at FROM playlists
            """)
        
        # 3. Hapus tabel lama dan rename tabel baru
        c.execute("DROP TABLE playlists")
        c.execute("ALTER TABLE playlists_new RENAME TO playlists")
        log.info("Migrasi playlists selesai.")

    # ========== HAPUS KOLOM 'files' LAMA JIKA MASIH ADA ==========
    c.execute("PRAGMA table_info(playlists)")
    cols = [row[1] for row in c.fetchall()]
    if 'files' in cols:
        log.warning("Menghapus kolom 'files' usang dari playlists...")
        # Buat tabel baru tanpa kolom files
        c.execute("""
            CREATE TABLE IF NOT EXISTS playlists_new (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                name TEXT NOT NULL,
                tracks TEXT NOT NULL DEFAULT '[]',
                is_favorite INTEGER DEFAULT 0,
                created_at TEXT DEFAULT(datetime('now'))
            )
        """)
        # Salin data dari tabel lama, abaikan kolom files
        c.execute("""
            INSERT INTO playlists_new (id, user_id, name, tracks, is_favorite, created_at)
            SELECT id, user_id, name, tracks, is_favorite, created_at FROM playlists
        """)
        c.execute("DROP TABLE playlists")
        c.execute("ALTER TABLE playlists_new RENAME TO playlists")
        log.warning("Kolom 'files' berhasil dihapus.")

    # --- MIGRASI: Perbaiki tabel boss_battles jika masih pakai party_id ---
    try:
        c.execute("PRAGMA table_info(boss_battles)")
        columns = [row[1] for row in c.fetchall()]
        if 'party_id' in columns:
            log.info("Migrasi boss_battles: party_id -> guild_id")
            c.execute("""CREATE TABLE boss_battles_new (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                guild_id INTEGER,
                boss_id TEXT NOT NULL,
                boss_name TEXT NOT NULL,
                boss_icon TEXT DEFAULT '🐉',
                boss_tier TEXT DEFAULT 'normal',
                boss_hp REAL DEFAULT 100,
                boss_max_hp REAL DEFAULT 100,
                boss_attack REAL DEFAULT 5,
                status TEXT DEFAULT 'active',
                started_at TEXT DEFAULT(datetime('now')),
                ended_at TEXT
            )""")
            c.execute("""INSERT INTO boss_battles_new
                (id, guild_id, boss_id, boss_name, boss_icon, boss_tier, boss_hp, boss_max_hp, boss_attack, status, started_at, ended_at)
                SELECT id, party_id, boss_id, boss_name, boss_icon, boss_tier, boss_hp, boss_max_hp, boss_attack, status, started_at, ended_at
                FROM boss_battles""")
            c.execute("DROP TABLE boss_battles")
            c.execute("ALTER TABLE boss_battles_new RENAME TO boss_battles")
            log.info("Migrasi boss_battles selesai.")
    except Exception as e:
        log.error(f"Migrasi boss_battles gagal (bisa diabaikan jika tidak perlu): {e}")
        log_crash(f"Migrasi boss_battles gagal: {e}")

    # --- Safe migration untuk kolom tambahan (schema) ---
    migrate_cols = [
        ("bio","TEXT DEFAULT ''"),
        ("avatar_emoji","TEXT DEFAULT '⚔️'"),
        ("gems","INTEGER DEFAULT 10"),
        ("longest_streak","INTEGER DEFAULT 0"),
        ("total_habits_done","INTEGER DEFAULT 0"),
        ("total_dailies_done","INTEGER DEFAULT 0"),
        ("total_todos_done","INTEGER DEFAULT 0"),
        ("total_xp_earned","INTEGER DEFAULT 0"),
        ("total_gold_earned","REAL DEFAULT 0"),
        ("boss_damage_bonus","REAL DEFAULT 0"),
        ("xp_multiplier","REAL DEFAULT 1.0"),
        ("gold_multiplier","REAL DEFAULT 1.0"),
        ("hp_damage_reduction","REAL DEFAULT 0"),
        ("has_revive","INTEGER DEFAULT 0"),
        ("mp_bonus","INTEGER DEFAULT 0"),
        ("theme","TEXT DEFAULT 'modern_dark'"),
        ("sound_enabled","INTEGER DEFAULT 1"),
        ("last_class_change", "TEXT"),
        ("reset_code", "TEXT"),
        ("reset_expiry", "TEXT"),
        ("sport_level", "INTEGER DEFAULT 1"),
        ("sport_xp", "INTEGER DEFAULT 0"),
        ("total_sport_points_earned", "INTEGER DEFAULT 0"),
        ("skill_buff_data", "TEXT DEFAULT '{}'"),
        ("class_passive_buffs", "TEXT DEFAULT '{}'"),
        ("is_admin", "INTEGER DEFAULT 0"),
        ("language", "TEXT DEFAULT 'en'")
    ]

    for col, defn in migrate_cols:
        _safe_alter(c, "users", col, defn)

    # Migrasi kolom untuk tabel lain
    _safe_alter(c, "habits", "last_action", "TEXT DEFAULT ''")
    _safe_alter(c, "dailies", "last_action", "TEXT DEFAULT ''")
    _safe_alter(c, "users", "security_question", "TEXT")
    _safe_alter(c, "users", "security_answer_hash", "TEXT")
    # ── Anti brute-force login (lockout sementara) ──
    _safe_alter(c, "users", "failed_attempts", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "locked_until", "TEXT")
    # ── Recurrence fleksibel (jadwal hari tertentu) ──
    _safe_alter(c, "habits", "repeat_days", "TEXT DEFAULT ''")
    _safe_alter(c, "dailies", "repeat_days", "TEXT DEFAULT ''")
    # ── Gamifikasi v1.3.0 Tahap 3 ──
    _safe_alter(c, "inventory", "enchant_level", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "last_login_reward", "TEXT")
    _safe_alter(c, "users", "login_streak", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "selected_title", "TEXT DEFAULT ''")
    _safe_alter(c, "users", "total_crafts", "INTEGER DEFAULT 0")
    # ── Akses & Fitur Besar Tahap 4 ──
    _safe_alter(c, "users", "font_scale", "INTEGER DEFAULT 100")
    _safe_alter(c, "users", "high_contrast", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "onboarding_done", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "talent_points_spent", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "dashboard_widgets", "TEXT DEFAULT ''")
    _safe_alter(c, "users", "gemini_api_key", "TEXT DEFAULT ''")
    _safe_alter(c, "users", "cloud_user_id", "TEXT")
    _safe_alter(c, "users", "is_cloud_shadow", "INTEGER DEFAULT 0")
    _safe_alter(c, "friends", "cloud_id", "TEXT")
    _safe_alter(c, "notifications", "cloud_id", "TEXT")
    _safe_alter(c, "couple_relationships", "cloud_id", "TEXT")
    _safe_alter(c, "couple_relationships", "cloud_status", "TEXT")
    _safe_alter(c, "couple_relationships", "cloud_ended_at", "TEXT")
    _safe_alter(c, "couple_relationships", "cloud_grace_ends_at", "TEXT")
    _safe_alter(c, "love_spaces", "cloud_id", "TEXT")
    _safe_alter(c, "love_space_photos", "cloud_id", "TEXT")
    _safe_alter(c, "love_space_photos", "cloud_storage_path", "TEXT")
    _safe_alter(c, "love_space_photos", "sync_status", "TEXT DEFAULT 'local'")
    c.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_users_cloud_id ON users(cloud_user_id) WHERE cloud_user_id IS NOT NULL")
    c.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_friends_cloud_id ON friends(cloud_id) WHERE cloud_id IS NOT NULL")
    c.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_notifications_cloud_id ON notifications(cloud_id) WHERE cloud_id IS NOT NULL")
    c.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_couple_cloud_id ON couple_relationships(cloud_id) WHERE cloud_id IS NOT NULL")
    c.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_love_space_cloud_id ON love_spaces(cloud_id) WHERE cloud_id IS NOT NULL")
    c.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_love_photo_cloud_id ON love_space_photos(cloud_id) WHERE cloud_id IS NOT NULL")
    # Phase 4A cloud-native shared Love Space mirrors.
    for love_table in (
        "relationship_profiles","relationship_events","relationship_memories",
        "relationship_checkins","relationship_prompt_responses",
        "relationship_prompt_favorites","relationship_weekly_reviews",
        "relationship_bucket_items","menstrual_settings","menstrual_cycles"
    ):
        _safe_alter(c,love_table,"cloud_id","TEXT")
        _safe_alter(c,love_table,"cloud_updated_at","TEXT")
        c.execute(f"CREATE UNIQUE INDEX IF NOT EXISTS idx_{love_table}_cloud_id ON {love_table}(cloud_id) WHERE cloud_id IS NOT NULL")
    _safe_alter(c, "sport_activities", "calories_burned", "INTEGER DEFAULT 0")
    _safe_alter(c, "sport_activities", "duration_minutes", "INTEGER DEFAULT 30")
    _safe_alter(c, "habits", "folder_id", "INTEGER")
    _safe_alter(c, "dailies", "folder_id", "INTEGER")
    _safe_alter(c, "todos", "folder_id", "INTEGER")
    _safe_alter(c, "sport_activities", "folder_id", "INTEGER")
    _safe_alter(c, "economy_items", "folder_id", "INTEGER")
    _safe_alter(c, "debts", "penalty_applied", "INTEGER DEFAULT 0")
    _safe_alter(c, "debts", "penalty_amount", "INTEGER DEFAULT 0")
    _safe_alter(c, "debts", "name", "TEXT NOT NULL")
    _safe_alter(c, 'health_logs', 'water_ml', 'INTEGER DEFAULT 0')
    _safe_alter(c, 'health_logs', 'weight_kg', 'REAL DEFAULT 0')
    _safe_alter(c, 'health_logs', 'resting_hr', 'INTEGER DEFAULT 0')
    _safe_alter(c, 'health_logs', 'stress_level', "TEXT DEFAULT 'normal'")
    _safe_alter(c, "users", "total_tasks_completed", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "total_gold_spent", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "currency", "TEXT DEFAULT 'IDR'")
    _safe_alter(c, "health_logs", "net_calories", "REAL DEFAULT 0")
    _safe_alter(c, "user_health_goals", "height_cm", "INTEGER DEFAULT 170")
    _safe_alter(c, "user_health_goals", "weight_kg", "REAL DEFAULT 70")
    _safe_alter(c, "user_health_goals", "age", "INTEGER DEFAULT 25")
    _safe_alter(c, "user_health_goals", "gender", "TEXT DEFAULT 'Laki-laki'")
    _safe_alter(c, "user_health_goals", "activity_factor", "REAL DEFAULT 1.55")
    _safe_alter(c, "users", "is_admin", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "rebirth_count", "INTEGER DEFAULT 0")
    _safe_alter(c, "economy_items", "sort_order", "INTEGER DEFAULT 0")
    _safe_alter(c, "food_logs", "sort_order", "INTEGER DEFAULT 0")
    _safe_alter(c, "dailies", "fail_streak", "INTEGER DEFAULT 0")
    _safe_alter(c, "messages", "deleted_by", "TEXT DEFAULT ''")
    # Phase 4B.1 direct-chat reply/edit/delete/reaction cache.
    _safe_alter(c,"messages","reply_to_id","INTEGER")
    _safe_alter(c,"messages","edited_at","TEXT")
    _safe_alter(c,"messages","deleted_at","TEXT")
    _safe_alter(c,"cloud_messages","reply_to_cloud_id","TEXT")
    _safe_alter(c,"cloud_messages","edited_at","TEXT")
    _safe_alter(c,"cloud_messages","deleted_at","TEXT")
    _safe_alter(c,"cloud_conversations","unread_count","INTEGER DEFAULT 0")
    _safe_alter(c,"cloud_conversations","last_message_at","TEXT")
    _safe_alter(c,"cloud_conversations","last_message_preview","TEXT DEFAULT ''")
    _safe_alter(c,"chat_attachments_cache","thumbnail_data","BLOB")
    _safe_alter(c, "note_folders", "parent_id", "INTEGER DEFAULT NULL")
    _safe_alter(c, "reminders", "repeat_type", "TEXT DEFAULT 'none'")
    _safe_alter(c, "reminders", "repeat_days", "TEXT DEFAULT ''")
    _safe_alter(c, "users", "has_spyglass", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "crit_chance", "INTEGER DEFAULT 10")
    _safe_alter(c, "guilds", "crit_chance", "INTEGER DEFAULT 0")
    _safe_alter(c, "boss_battles", "boss_crit_chance", "INTEGER DEFAULT 15")
    _safe_alter(c, "users", "block_chance", "INTEGER DEFAULT 20")
    _safe_alter(c, "users", "block_strength", "INTEGER DEFAULT 10")
    _safe_alter(c, "users", "ultimate_last_used", "TEXT")
    _safe_alter(c, "users", "boss_attacks_today", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "boss_attacks_date", "TEXT")
    _safe_alter(c, "boss_battles", "participants", "TEXT")
    _safe_alter(c, "boss_battles", "attack_counts", "TEXT")
    _safe_alter(c, "boss_battles", "raid_leader_id", "INTEGER")
    _safe_alter(c, "guilds", "boss_defeated_today", "INTEGER DEFAULT 0")
    _safe_alter(c, "guilds", "boss_defeated_date", "TEXT")
    _safe_alter(c, "dailies", "freeze_slots", "INTEGER DEFAULT 0")
    _safe_alter(c, "playlists", "is_favorite", "INTEGER DEFAULT 0")
    _safe_alter(c, "playlists", "tracks", "TEXT NOT NULL DEFAULT '[]'")
    _safe_alter(c, "notes", "zoom_level", "INTEGER DEFAULT 100")
    _safe_alter(c, "users", "is_locked", "INTEGER DEFAULT 0")
    _safe_alter(c, "users", "locked_at", "TEXT")
    _safe_alter(c, "users", "last_tracking_date", "TEXT")

    for col, defn in [("level", "INTEGER DEFAULT 1"),
                      ("exp", "INTEGER DEFAULT 0"),
                      ("buff_xp", "REAL DEFAULT 0"),
                      ("buff_gold", "REAL DEFAULT 0"),
                      ("buff_damage", "REAL DEFAULT 0")]:
        try: c.execute(f"ALTER TABLE guilds ADD COLUMN {col} {defn}")
        except: pass

    for col, defn in [("level", "INTEGER DEFAULT 1"),
                      ("exp", "INTEGER DEFAULT 0"),
                      ("hunger", "INTEGER DEFAULT 100"),
                      ("last_fed", "TEXT")]:
        try: c.execute(f"ALTER TABLE user_pets ADD COLUMN {col} {defn}")
        except: pass

    # Tambahkan sort_order untuk drag & drop reorder
    for table in ["habits", "dailies", "todos", "sport_activities"]:
        _safe_alter(c, table, "sort_order", "INTEGER DEFAULT 0")

    # ── Task Folders (for organize habits/dailies/todos/sport) ────────────────
    c.execute("""CREATE TABLE IF NOT EXISTS task_folders(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        mode TEXT NOT NULL,
        name TEXT NOT NULL,
        icon TEXT DEFAULT '📁',
        collapsed INTEGER DEFAULT 0,
        created_at TEXT DEFAULT(datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")

    # ========== DATA MIGRATION (hanya untuk database baru/kosong) ==========
    # CEK: Apakah database sudah memiliki data?
    table_check = c.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name='users'"
    ).fetchone()
    is_new_db = True
    if table_check:
        user_count = c.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        if user_count > 0:
            is_new_db = False
            log.warning(f"Database sudah terisi ({user_count} user), skip initial data insertion.")

    # ========== HANYA UNTUK DATABASE BARU ==========
    log.info("Database baru, menjalankan migrasi data...")

    # Inisialisasi achievements
    init_achievements()

    # ========== MIGRASI ACHIEVEMENT ==========
    migrate_achievements()

    # ========== MIGRASI REDEEM CODES ==========
    migrate_redeem_codes()

    # Tambah makanan default
    cur = c.execute("SELECT name FROM food_items WHERE is_custom = 0")
    existing_default_names = {row[0] for row in cur.fetchall()}
    new_foods_added = 0
    for food in DEFAULT_FOODS:
        name_id = food[0]   # ambil nama Indonesia
        if name_id not in existing_default_names:
            c.execute(
                "INSERT INTO food_items(name, icon, calories, protein, carbs, fat, is_custom) "
                "VALUES(?,?,?,?,?,?,?)",
                (name_id, food[2], food[3], food[4], food[5], food[6], 0)
            )
            new_foods_added += 1
    if new_foods_added > 0:
        log.info(f"Menambahkan {new_foods_added} makanan default baru.")
    else:
        log.warning("Tidak ada makanan default baru, skip.")

    # Insert default redeem codes
    cur = c.execute("SELECT COUNT(*) FROM redeem_codes")
    if cur.fetchone()[0] == 0:
        default_codes = [
            ("ADMINADMINADMIN", "admin", 0, None, 1),
            ("WELCOME100", "xp", 100, None, 1),
            ("STARTGOLD", "gold", 500, None, 1),
            ("WOODSWORD", "item", 0, "wooden_sword", 1),
            ("GOLDENAPPLE", "item", 0, "golden_apple", 1),
        ]
        for code, rtype, rval, ritem, onetime in default_codes:
            try:
                c.execute(
                    "INSERT INTO redeem_codes(code, reward_type, reward_value, reward_item, is_one_time) VALUES(?,?,?,?,?)",
                    (code, rtype, rval, ritem, onetime)
                )
            except:
                pass

    # ── OPTIMASI: Buat indeks untuk mempercepat query ─────────────────────────
    c.execute("CREATE INDEX IF NOT EXISTS idx_activity_log_user_id ON activity_log(user_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_activity_log_created_at ON activity_log(created_at)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_food_logs_user_log ON food_logs(user_id, log_date)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_food_logs_user_id ON food_logs(user_id)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_health_logs_user_log ON health_logs(user_id, log_date)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_health_logs_user_id ON health_logs(user_id)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_water_logs_user_log ON water_logs(user_id, log_date)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_sport_activities_user_done ON sport_activities(user_id, done_today, last_done)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_sport_activities_user_id ON sport_activities(user_id)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_economy_items_user_date ON economy_items(user_id, date)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_economy_items_user_type ON economy_items(user_id, type)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_habits_user_id ON habits(user_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_dailies_user_id ON dailies(user_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_todos_user_id ON todos(user_id)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_notifications_user_read ON notifications(user_id, is_read)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_user_pets_user_active ON user_pets(user_id, is_active)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_guild_members_guild ON guild_members(guild_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_guild_members_user ON guild_members(user_id)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_messages_users ON messages(sender_id, receiver_id)")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_debts_user_paid ON debts(user_id, is_paid)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_debt_notes_user_status ON debt_notes(user_id, status)")

    c.execute("CREATE INDEX IF NOT EXISTS idx_habits_sort_order ON habits(user_id, sort_order)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_dailies_sort_order ON dailies(user_id, sort_order)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_todos_sort_order ON todos(user_id, sort_order)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_sport_activities_sort_order ON sport_activities(user_id, sort_order)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_sport_rep_logs_user_act_date ON sport_rep_logs(user_id, activity_id, log_date)")

    c.execute("CREATE INDEX IF NOT EXISTS idx_economy_items_sort_order ON economy_items(user_id, sort_order)")

    c.execute("CREATE INDEX IF NOT EXISTS idx_food_logs_sort_order ON food_logs(user_id, sort_order)")

    c.execute("CREATE INDEX IF NOT EXISTS idx_notes_user_folder ON notes(user_id, folder_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_notes_user_archived ON notes(user_id, is_archived)")

    c.execute("CREATE INDEX IF NOT EXISTS idx_note_folders_parent ON note_folders(parent_id)")

    c.execute("CREATE INDEX IF NOT EXISTS idx_habits_last_done ON habits(user_id, last_done)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_dailies_last_done ON dailies(user_id, last_done)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_task_history_user_date ON task_history(user_id, task_type, action_date)")

    # Nyalakan kembali foreign key
    c.execute("PRAGMA foreign_keys = ON")
    conn.commit()
    conn.close()
    log.info(f"Ready: {DB_PATH}")

    # ── Migrasi SATU KALI (ditandai PRAGMA user_version=1) ──
    # 1) migrate_sort_order: backfill sort_order hanya SEKALI — bukan tiap
    #    startup! (Dulu dipanggil setiap init sehingga urutan manual hasil
    #    tombol ⬆/⬇ selalu tertimpa oleh urutan created_at saat restart.)
    # 2) Urutkan ulang economy_items per user+folder: TERBARU di paling atas
    #    (date DESC, created_at DESC). Ekspektasi user: entri terakhir
    #    ditambahkan tampil paling atas, bukan tanggal paling lama.
    try:
        uv = c.execute("PRAGMA user_version").fetchone()[0]
        if uv < 1:
            migrate_sort_order()
            rows = c.execute(
                """SELECT id, user_id, folder_id FROM economy_items
                   ORDER BY user_id, folder_id, date DESC, created_at DESC, id DESC"""
            ).fetchall()
            counters = {}
            for r in rows:
                key = (r["user_id"], r["folder_id"])
                pos = counters.get(key, 0)
                c.execute("UPDATE economy_items SET sort_order=? WHERE id=?", (pos, r["id"]))
                counters[key] = pos + 1
            c.execute("PRAGMA user_version = 1")
            log.info("Migrasi urutan economy_items (terbaru di atas) selesai.")
    except Exception as e:
        log.error(f"Migrasi urutan economy_items gagal (aman diabaikan): {e}")

# ========== KODE REDEEM BARU ==========
NEW_REDEEM_CODES = [
    ("FREEXP200", "xp", 200, None, 1),
    ("FREEGOLD500", "gold", 500, None, 1),
    ("FREEPOTION", "item", 0, "health_potion", 1),
    ("FREESWORD", "item", 0, "iron_sword", 1),
    ("BIGBONUS", "gold", 1000, None, 1),
    # ── 🆕 v1.4.0 — Redeem baru (shop items baru & bonus) ──
    ("BRONZE70", "item", 0, "bronze_sword", 1),
    ("STEEL250", "item", 0, "steel_helm", 1),
    ("TRAVEL190", "item", 0, "travelers_boots", 1),
    ("ARCANE450", "item", 0, "arcane_ring", 1),
    ("BERRYPIE", "item", 0, "berry_pie", 1),
    ("MANACOOKIE", "item", 0, "mana_cookie", 1),
    ("FROST550", "item", 0, "frost_guard", 1),
    ("VOID800", "item", 0, "void_core", 1),
    ("SCHOLAR500", "item", 0, "scholar_crown", 1),
    ("EMBER400", "item", 0, "ember_charm", 1),
]

def migrate_redeem_codes():
    """Tambahkan kode redeem baru jika belum ada."""
    conn = get_conn()
    cur = conn.cursor()
    for code, rtype, rval, ritem, onetime in NEW_REDEEM_CODES:
        existing = cur.execute("SELECT id FROM redeem_codes WHERE code = ?", (code,)).fetchone()
        if not existing:
            try:
                cur.execute(
                    "INSERT INTO redeem_codes(code, reward_type, reward_value, reward_item, is_one_time) VALUES(?,?,?,?,?)",
                    (code, rtype, rval, ritem, onetime)
                )
            except:
                pass
    conn.commit()
    conn.close()
    log.info(f"Migrated {len(NEW_REDEEM_CODES)} new redeem codes.")

# ========== KONSTANTA TABEL YANG DI-EKSPOR ==========
EXPORT_TABLES = [
    'task_folders',
    'note_folders',
    'food_items',
    'recipes',
    'habits',
    'dailies',
    'todos',
    'sport_activities',
    'notes',
    'recipe_items',
    'food_logs',
    'economy_items',
    'debts',
    'savings',
    'investments',
    'subscriptions',
    'water_logs',
    'health_logs',
    'user_nutrition_goals',
    'user_water_goals',
    'user_health_goals',
    'task_history',
    'pomodoro_sessions',
    'sport_rep_logs',
    'calendar_notes',
    'reminders',
    'relationship_profiles',
    'relationship_events',
    'relationship_memories',
    'relationship_checkins',
    'relationship_prompt_responses',
    'relationship_prompt_favorites',
    'relationship_weekly_reviews',
    'relationship_bucket_items',
    'menstrual_settings',
    'menstrual_cycles',
]

# ========== EKSPOR ==========
def export_tracker_data(user_id):
    conn = get_conn()
    # Cek admin
    u = conn.execute("SELECT is_admin FROM users WHERE id=?", (user_id,)).fetchone()
    if u and u["is_admin"]:
        conn.close()
        raise PermissionError("Admin tidak diperbolehkan mengekspor data tracker.")
    
    data = {"version": 1, "tables": {}}
    
    for table in EXPORT_TABLES:
        if table == "recipe_items":
            # Ambil recipe_items yang terkait dengan recipes milik user, sertakan food_name
            rows = conn.execute("""
                SELECT ri.*, fi.name as food_name
                FROM recipe_items ri
                JOIN recipes r ON ri.recipe_id = r.id
                JOIN food_items fi ON ri.food_id = fi.id
                WHERE r.user_id = ?
            """, (user_id,)).fetchall()
        elif table == "food_items":
            # Hanya custom food
            rows = conn.execute(
                "SELECT * FROM food_items WHERE user_id=? AND is_custom=1",
                (user_id,)
            ).fetchall()
        elif table == "food_logs":
            # Ambil food_logs beserta nama makanan (untuk mapping default food nanti)
            rows = conn.execute("""
                SELECT fl.*, fi.name as food_name
                FROM food_logs fl
                LEFT JOIN food_items fi ON fl.food_id = fi.id
                WHERE fl.user_id = ?
            """, (user_id,)).fetchall()
        else:
            # Tabel lain yang memiliki kolom user_id
            rows = conn.execute(f"SELECT * FROM {table} WHERE user_id=?", (user_id,)).fetchall()
        
        list_rows = [dict(r) for r in rows]
        # Hapus kolom user_id dari setiap baris (jika ada)
        for row in list_rows:
            row.pop('user_id', None)
        
        data["tables"][table] = list_rows
    
    conn.close()
    return data

# ========== BERSIHKAN DATA TRACKER ==========
def _clear_tracker_data_conn(conn, user_id):
    """Hapus tracker memakai koneksi aktif agar import dapat benar-benar atomik."""
    conn.execute(
        "DELETE FROM recipe_items WHERE recipe_id IN (SELECT id FROM recipes WHERE user_id=?)",
        (user_id,)
    )
    for table in EXPORT_TABLES:
        if table == 'recipe_items':
            continue
        conn.execute(f"DELETE FROM {table} WHERE user_id=?", (user_id,))


def clear_tracker_data(user_id):
    """Hapus semua data tracker user (tanpa menyentuh data akun)."""
    conn = get_conn()
    try:
        conn.execute("PRAGMA foreign_keys = OFF")
        _clear_tracker_data_conn(conn, user_id)
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.execute("PRAGMA foreign_keys = ON")
        conn.close()

# ========== HELPER: DAPATKAN ID DEFAULT FOOD BERDASARKAN NAMA ==========
def get_default_food_id_by_name(name):
    """Cari ID makanan default (is_custom=0) berdasarkan nama."""
    conn = get_conn()
    row = conn.execute(
        "SELECT id FROM food_items WHERE is_custom=0 AND name=?",
        (name,)
    ).fetchone()
    conn.close()
    return row["id"] if row else None

# ========== IMPOR ==========
def import_tracker_data(user_id, data, preserve_progress=False):
    """Impor data tracker ke ``user_id``.

    Impor manual lama tetap mereset progres. Sinkronisasi cloud memakai
    ``preserve_progress=True`` agar streak, status selesai, dan history tetap utuh.
    """
    conn = get_conn()
    # Cek admin
    u = conn.execute("SELECT is_admin FROM users WHERE id=?", (user_id,)).fetchone()
    if u and u["is_admin"]:
        conn.close()
        raise PermissionError("Admin tidak diperbolehkan mengimpor data tracker.")
    
    conn.execute("PRAGMA foreign_keys = OFF")
    try:
        # 1. Hapus data lama pada transaksi yang sama. Jika import gagal,
        # rollback menjaga data lokal lama tetap utuh.
        _clear_tracker_data_conn(conn, user_id)
        
        # 2. Mapping untuk foreign key
        mapping = {
            'task_folders': {},
            'note_folders': {},
            'food_items': {},      # hanya custom yang di-insert
            'recipes': {},
            'habits': {},
            'dailies': {},
            'todos': {},
            'sport_activities': {},
            'notes': {},
        }
        # Mapping khusus task_history
        task_mapping = {
            'habit': {},
            'daily': {},
            'todo': {},
            'sport': {},
        }
        
        # 3. Insert parent tables (tanpa FK ke tabel lain dalam export)
        parent_tables = ['task_folders', 'note_folders', 'food_items', 'recipes']
        note_folder_parents = {}
        for table in parent_tables:
            rows = data['tables'].get(table, [])
            for row in rows:
                old_parent = row.get('parent_id')  
                insert_row = {k: v for k, v in row.items() if k != 'id'}
                insert_row['user_id'] = user_id
                # Jika ada kolom created_at/updated_at, tetap pakai nilai dari file
                cols = ','.join(insert_row.keys())
                placeholders = ','.join(['?'] * len(insert_row))
                cur = conn.execute(
                    f"INSERT INTO {table} ({cols}) VALUES ({placeholders})",
                    list(insert_row.values())
                )
                new_id = cur.lastrowid                 # ← PERBAIKI TYPO
                mapping[table][row['id']] = new_id
                if table == 'note_folders':            # ← HANYA UNTUK note_folders
                    note_folder_parents[row['id']] = old_parent
        
        # Update parent_id untuk note_folders
        for old_id, old_parent in note_folder_parents.items():
            if old_parent is not None and old_parent in mapping['note_folders']:
                new_parent = mapping['note_folders'][old_parent]
                new_id = mapping['note_folders'][old_id]
                conn.execute(
                    "UPDATE note_folders SET parent_id = ? WHERE id = ?",
                    (new_parent, new_id)
                )

        # 4. Insert child tables dengan foreign key
        # 4a. habits, dailies, todos, sport
        for table, fk_table, fk_col in [
            ('habits', 'task_folders', 'folder_id'),
            ('dailies', 'task_folders', 'folder_id'),
            ('todos', 'task_folders', 'folder_id'),
            ('sport_activities', 'task_folders', 'folder_id'),
        ]:
            rows = data['tables'].get(table, [])
            for row in rows:
                insert_row = {k: v for k, v in row.items() if k != 'id'}
                insert_row['user_id'] = user_id

                # Impor manual mempertahankan perilaku lama (template tanpa progres).
                # Cloud restore harus mereplikasi progres secara utuh.
                if not preserve_progress:
                    if table == 'habits':
                        insert_row['streak'] = 0
                        insert_row['done_today'] = 0
                        insert_row['counter_up'] = 0
                        insert_row['counter_down'] = 0
                        insert_row['last_done'] = None
                        insert_row['last_action'] = ''
                    elif table == 'dailies':
                        insert_row['streak'] = 0
                        insert_row['done_today'] = 0
                        insert_row['fail_streak'] = 0
                        insert_row['last_done'] = None
                        insert_row['last_action'] = ''
                    elif table == 'todos':
                        insert_row['done'] = 0
                    elif table == 'sport_activities':
                        insert_row['streak'] = 0
                        insert_row['done_today'] = 0
                        insert_row['last_done'] = None
                    
                # Map folder_id
                old_fid = row.get('folder_id')
                if old_fid is not None and old_fid in mapping['task_folders']:
                    insert_row['folder_id'] = mapping['task_folders'][old_fid]
                else:
                    insert_row['folder_id'] = None
                cols = ','.join(insert_row.keys())
                placeholders = ','.join(['?'] * len(insert_row))
                cur = conn.execute(
                    f"INSERT INTO {table} ({cols}) VALUES ({placeholders})",
                    list(insert_row.values())
                )
                new_id = cur.lastrowid
                mapping[table][row['id']] = new_id
                # Simpan di task_mapping
                ttype = 'habit' if table == 'habits' else \
                        'daily' if table == 'dailies' else \
                        'todo' if table == 'todos' else 'sport'
                task_mapping[ttype][row['id']] = new_id
        
        # 4b. notes
        rows = data['tables'].get('notes', [])
        for row in rows:
            old_parent = row.get('parent_id')
            insert_row = {k: v for k, v in row.items() if k != 'id'}
            insert_row['user_id'] = user_id
            old_fid = row.get('folder_id')
            if old_fid is not None and old_fid in mapping['note_folders']:
                insert_row['folder_id'] = mapping['note_folders'][old_fid]
            else:
                insert_row['folder_id'] = None
            cols = ','.join(insert_row.keys())
            placeholders = ','.join(['?'] * len(insert_row))
            cur = conn.execute(
                f"INSERT INTO notes ({cols}) VALUES ({placeholders})",
                list(insert_row.values())
            )
            mapping['notes'][row['id']] = cur.lastrowid
        
        # 4c. recipe_items
        rows = data['tables'].get('recipe_items', [])
        for row in rows:
            insert_row = {k: v for k, v in row.items() if k != 'id'}
            # Map recipe_id
            old_rid = row.get('recipe_id')
            if old_rid in mapping['recipes']:
                insert_row['recipe_id'] = mapping['recipes'][old_rid]
            else:
                continue  # skip jika resep tidak ada (seharusnya tidak terjadi)
            # Map food_id
            old_fid = row.get('food_id')
            if old_fid in mapping['food_items']:
                insert_row['food_id'] = mapping['food_items'][old_fid]
            else:
                # Cari default food berdasarkan nama (jika tersedia)
                food_name = row.get('food_name')
                if food_name:
                    default_id = get_default_food_id_by_name(food_name)
                    if default_id:
                        insert_row['food_id'] = default_id
                    else:
                        # Jika tidak ditemukan, skip atau raise error
                        continue
                else:
                    continue
            # food_name hanya metadata portabel untuk memetakan default food.
            insert_row.pop('food_name', None)
            cols = ','.join(insert_row.keys())
            placeholders = ','.join(['?'] * len(insert_row))
            conn.execute(
                f"INSERT INTO recipe_items ({cols}) VALUES ({placeholders})",
                list(insert_row.values())
            )
        
        # 4d. food_logs
        rows = data['tables'].get('food_logs', [])
        for row in rows:
            insert_row = {k: v for k, v in row.items() if k != 'id'}
            insert_row['user_id'] = user_id
            old_fid = row.get('food_id')
            if old_fid in mapping['food_items']:
                insert_row['food_id'] = mapping['food_items'][old_fid]
            else:
                food_name = row.get('food_name')
                if food_name:
                    default_id = get_default_food_id_by_name(food_name)
                    if default_id:
                        insert_row['food_id'] = default_id
                    else:
                        continue
                else:
                    continue
            # Hapus food_name agar tidak ikut INSERT
            insert_row.pop('food_name', None)

            cols = ','.join(insert_row.keys())
            placeholders = ','.join(['?'] * len(insert_row))
            conn.execute(
                f"INSERT INTO food_logs ({cols}) VALUES ({placeholders})",
                list(insert_row.values())
            )
        
        # 4e. economy_items (opsional folder_id)
        rows = data['tables'].get('economy_items', [])
        for row in rows:
            insert_row = {k: v for k, v in row.items() if k != 'id'}
            insert_row['user_id'] = user_id
            old_fid = row.get('folder_id')
            if old_fid is not None and old_fid in mapping['task_folders']:
                insert_row['folder_id'] = mapping['task_folders'][old_fid]
            else:
                insert_row['folder_id'] = None
            cols = ','.join(insert_row.keys())
            placeholders = ','.join(['?'] * len(insert_row))
            conn.execute(
                f"INSERT INTO economy_items ({cols}) VALUES ({placeholders})",
                list(insert_row.values())
            )
        
        # 4f. Riwayat task dan log repetisi perlu memetakan ID task/activity baru.
        for row in data['tables'].get('task_history', []):
            task_type = row.get('task_type')
            old_task_id = row.get('task_id')
            new_task_id = task_mapping.get(task_type, {}).get(old_task_id)
            if new_task_id is None:
                continue
            insert_row = {k: v for k, v in row.items() if k != 'id'}
            insert_row['user_id'] = user_id
            insert_row['task_id'] = new_task_id
            cols = ','.join(insert_row.keys())
            placeholders = ','.join(['?'] * len(insert_row))
            conn.execute(f"INSERT INTO task_history ({cols}) VALUES ({placeholders})", list(insert_row.values()))

        for row in data['tables'].get('sport_rep_logs', []):
            new_activity_id = mapping['sport_activities'].get(row.get('activity_id'))
            if new_activity_id is None:
                continue
            insert_row = {k: v for k, v in row.items() if k != 'id'}
            insert_row['user_id'] = user_id
            insert_row['activity_id'] = new_activity_id
            cols = ','.join(insert_row.keys())
            placeholders = ','.join(['?'] * len(insert_row))
            conn.execute(f"INSERT INTO sport_rep_logs ({cols}) VALUES ({placeholders})", list(insert_row.values()))

        # 4g. Tabel tanpa foreign key selain user_id.
        tables_no_fk = [
            'debts', 'savings', 'investments', 'subscriptions',
            'water_logs', 'health_logs', 'user_nutrition_goals',
            'user_water_goals', 'user_health_goals', 'pomodoro_sessions',
            'calendar_notes', 'reminders',
            'relationship_profiles', 'relationship_events', 'relationship_memories',
            'relationship_checkins', 'relationship_prompt_responses',
            'relationship_prompt_favorites', 'relationship_weekly_reviews',
            'relationship_bucket_items', 'menstrual_settings', 'menstrual_cycles'
        ]
        for table in tables_no_fk:
            rows = data['tables'].get(table, [])
            for row in rows:
                insert_row = {k: v for k, v in row.items() if k != 'id'}
                insert_row['user_id'] = user_id
                cols = ','.join(insert_row.keys())
                placeholders = ','.join(['?'] * len(insert_row))
                conn.execute(
                    f"INSERT INTO {table} ({cols}) VALUES ({placeholders})",
                    list(insert_row.values())
                )
        
        conn.commit()
    
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        conn.execute("PRAGMA foreign_keys = ON")
        conn.close()

def get_food_translation(name_id: str, lang: str = "id") -> str:
    """Wrapper untuk memudahkan akses terjemahan nama makanan dari database."""
    from food_data import get_food_name
    return get_food_name(name_id, lang)

# ========== DATABASE BACKUP & FORCE SYNC ==========
def force_checkpoint():
    """Paksa WAL checkpoint agar semua data di memori tertulis ke database utama."""
    try:
        conn = get_conn()
        conn.execute("PRAGMA wal_checkpoint(TRUNCATE)")
        conn.close()
        log.info("Checkpoint berhasil.")
    except Exception as e:
        log.error(f"Checkpoint gagal: {e}")

def prune_backups(backup_dir, keep=7, max_age_days=30):
    """Rotasi backup: simpan `keep` file terbaru; file yang lebih tua dari
    `max_age_days` hari juga dihapus. Return jumlah file yang dihapus."""
    removed = 0
    try:
        files = [f for f in os.listdir(backup_dir)
                 if f.startswith("craftlife_backup_") and f.endswith(".db")]
        files.sort(key=lambda f: os.path.getmtime(os.path.join(backup_dir, f)),
                   reverse=True)
        now = time.time()
        for i, f in enumerate(files):
            path = os.path.join(backup_dir, f)
            too_many = i >= keep                                # lewat kuota terbaru
            too_old = os.path.getmtime(path) < now - (max_age_days * 86400)
            if too_many or too_old:
                try:
                    os.remove(path)
                    removed += 1
                    log.info(f"Backup dipangkas: {f}")
                except OSError as e:
                    log.warning(f"Gagal menghapus backup {f}: {e}")
    except Exception as e:
        log.error(f"Prune backup gagal: {e}")
    return removed


def backup_database():
    """Buat backup database ke folder backups/ dengan timestamp.
    Rotasi otomatis: hanya 7 backup terbaru (maks 30 hari) yang disimpan."""
    try:
        # Force checkpoint dulu agar backup lengkap
        force_checkpoint()

        src = DB_PATH
        if not os.path.exists(src):
            log.warning("Database tidak ditemukan, backup batal.")
            return None

        backup_dir = os.path.join(os.path.dirname(src), "backups")
        os.makedirs(backup_dir, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        dst = os.path.join(backup_dir, f"craftlife_backup_{timestamp}.db")
        # Hindari tabrakan nama jika dua backup terjadi dalam detik yang sama
        n = 2
        while os.path.exists(dst):
            dst = os.path.join(backup_dir, f"craftlife_backup_{timestamp}_{n}.db")
            n += 1

        import shutil
        shutil.copy2(src, dst)
        log.info(f"Backup berhasil: {dst}")

        prune_backups(backup_dir)
        return dst
    except Exception as e:
        log.error(f"Backup gagal: {e}")
        return None

# ── Auth ──────────────────────────────────────────────────────────────────────

_PBKDF2_ITERATIONS = 260_000
_PBKDF2_PREFIX = "pbkdf2_sha256"


def _hash(pw):
    """LEGACY: SHA-256 tanpa salt.

    HANYA dipakai untuk memverifikasi hash lama milik user yang sudah ada
    sebelum migrasi PBKDF2. Jangan gunakan untuk menyimpan hash baru —
    gunakan _hash_password().
    """
    return hashlib.sha256(pw.encode()).hexdigest()


def _hash_password(pw):
    """Hash password baru: PBKDF2-HMAC-SHA256 dengan salt acak per password.

    Format: pbkdf2_sha256$<iterations>$<salt_hex>$<hash_hex>
    """
    salt = secrets.token_hex(16)
    dk = hashlib.pbkdf2_hmac(
        "sha256", pw.encode(), bytes.fromhex(salt), _PBKDF2_ITERATIONS
    )
    return f"{_PBKDF2_PREFIX}${_PBKDF2_ITERATIONS}${salt}${dk.hex()}"


def _verify_password(pw, stored):
    """Verifikasi password terhadap hash tersimpan.

    Mendukung format baru (PBKDF2) dan format legacy (SHA-256 tanpa salt),
    jadi user lama tetap bisa login. Perbandingan memakai compare_digest
    untuk mencegah timing attack.
    """
    if not stored:
        return False
    if stored.startswith(_PBKDF2_PREFIX + "$"):
        try:
            _, iters, salt, expected = stored.split("$")
            dk = hashlib.pbkdf2_hmac(
                "sha256", pw.encode(), bytes.fromhex(salt), int(iters)
            )
            return secrets.compare_digest(dk.hex(), expected)
        except (ValueError, TypeError):
            return False
    # Legacy: SHA-256 tanpa salt
    return secrets.compare_digest(_hash(pw), stored)


def _is_legacy_hash(stored):
    """True jika hash masih format lama (belum PBKDF2)."""
    return bool(stored) and not stored.startswith(_PBKDF2_PREFIX + "$")


# ── Keamanan: aturan password & anti brute-force ─────────────────────────────

PASSWORD_MIN_LEN = 8          # minimal panjang password baru
LOGIN_MAX_ATTEMPTS = 5        # percobaan gagal sebelum lockout
LOGIN_LOCK_MINUTES = 5        # durasi lockout sementara
SESSION_TOKEN_DAYS = 30       # masa berlaku token "Remember Me"


def validate_password_strength(pw: str):
    """Cek aturan minimal password. Return (ok, pesan_key atau None)."""
    if len(pw or "") < PASSWORD_MIN_LEN:
        return False, "db_password_too_short"
    return True, None


# ── Session token "Remember Me" ───────────────────────────────────────────────
# File session di disk HANYA menyimpan token acak (bukan password_hash).
# Di database tersimpan hash SHA-256 dari token (token ber-entropy tinggi,
# jadi SHA-256 tanpa salt sudah memadai). Token dicuri pun hanya berlaku
# sampai expires_at dan ikut hangus saat password diganti / akun di-lock.

def _hash_token(token: str) -> str:
    return hashlib.sha256(token.encode()).hexdigest()


def create_session_token(user_id: int) -> str:
    """Buat token session baru; return token mentah untuk disimpan di client."""
    purge_expired_sessions()
    token = secrets.token_urlsafe(32)
    expires = (datetime.now() + timedelta(days=SESSION_TOKEN_DAYS)).isoformat()
    conn = get_conn()
    conn.execute(
        "INSERT INTO user_sessions(user_id, token_hash, created_at, expires_at)"
        " VALUES(?,?,?,?)",
        (user_id, _hash_token(token), datetime.now().isoformat(), expires)
    )
    conn.commit()
    conn.close()
    log.info(f"Session token dibuat untuk user_id={user_id}")
    return token


def validate_session_token(user_id: int, token: str) -> bool:
    """True jika token valid & belum kedaluwarsa."""
    if not token:
        return False
    conn = get_conn()
    row = conn.execute(
        "SELECT expires_at FROM user_sessions WHERE user_id=? AND token_hash=?",
        (user_id, _hash_token(token))
    ).fetchone()
    conn.close()
    if not row:
        return False
    try:
        if datetime.fromisoformat(row["expires_at"]) < datetime.now():
            return False
    except (ValueError, TypeError):
        return False
    return True


def delete_session_token(user_id: int, token: str) -> None:
    """Hapus satu token (dipakai saat logout)."""
    conn = get_conn()
    conn.execute(
        "DELETE FROM user_sessions WHERE user_id=? AND token_hash=?",
        (user_id, _hash_token(token))
    )
    conn.commit()
    conn.close()


def delete_all_session_tokens(user_id: int) -> None:
    """Hanguskan SEMUA session user (dipakai saat ganti password/lock/hapus akun)."""
    conn = get_conn()
    conn.execute("DELETE FROM user_sessions WHERE user_id=?", (user_id,))
    conn.commit()
    conn.close()
    log.info(f"Semua session token user_id={user_id} dihanguskan")


def purge_expired_sessions() -> None:
    """Bersihkan token kedaluwarsa (hemat storage, dipanggil tiap login)."""
    try:
        conn = get_conn()
        conn.execute("DELETE FROM user_sessions WHERE expires_at < ?",
                     (datetime.now().isoformat(),))
        conn.commit()
        conn.close()
    except Exception as e:
        log.warning(f"Purge session kedaluwarsa gagal: {e}")


def register_user(username, password, display_name="", bio="", avatar_class="warrior"):
    ok_pw, key = validate_password_strength(password)
    if not ok_pw:
        return {"ok": False, "msg": tr_db(lang="id", key=key)}
    conn = get_conn()
    try:
        conn.execute(
            "INSERT INTO users(username,password_hash,display_name,bio,avatar_class)"
            " VALUES(?,?,?,?,?)",
            (username.lower().strip(), _hash_password(password),
             display_name or username, bio, avatar_class)
        )
        conn.commit()
        user_id = conn.execute("SELECT id FROM users WHERE username=?", (username.lower().strip(),)).fetchone()
        uid = user_id["id"] if user_id else None
        msg = tr_db(user_id=uid, key="db_register_success") if uid else "Registrasi berhasil!"
        return {"ok": True, "msg": msg}
    except sqlite3.IntegrityError:
        return {"ok": False, "msg": tr_db(lang="id", key="db_register_username_taken")}
    finally:
        conn.close()


def login_user(username, password):
    """Login dengan proteksi brute-force: setelah LOGIN_MAX_ATTEMPTS kali gagal,
    akun di-lock sementara selama LOGIN_LOCK_MINUTES menit."""
    uname = username.lower().strip()
    conn = get_conn()
    row = conn.execute("SELECT * FROM users WHERE username=?", (uname,)).fetchone()

    if row:
        # ── Cek lockout aktif ──
        locked_until = row["locked_until"] if "locked_until" in row.keys() else None
        if locked_until:
            try:
                until = datetime.fromisoformat(locked_until)
            except (ValueError, TypeError):
                until = None
            if until and datetime.now() < until:
                mins = max(1, int((until - datetime.now()).total_seconds() // 60) + 1)
                conn.close()
                log.warning(f"Login ditolak (lockout aktif) untuk '{uname}'")
                return {"ok": False, "locked": True,
                        "msg": tr_db(lang="id", key="db_login_locked", mins=mins)}
            # Lockout sudah lewat → bersihkan
            conn.execute("UPDATE users SET failed_attempts=0, locked_until=NULL WHERE id=?",
                         (row["id"],))
            conn.commit()

        if _verify_password(password, row["password_hash"]):
            # Sukses: reset counter gagal
            conn.execute("UPDATE users SET failed_attempts=0, locked_until=NULL WHERE id=?",
                         (row["id"],))
            # Migrasi transparan: upgrade hash legacy (SHA-256 tanpa salt) ke PBKDF2
            if _is_legacy_hash(row["password_hash"]):
                new_hash = _hash_password(password)
                conn.execute("UPDATE users SET password_hash=? WHERE id=?",
                             (new_hash, row["id"]))
            conn.execute("UPDATE users SET last_login=? WHERE id=?",
                         (local_now().isoformat(), row["id"]))
            conn.commit()
            # Ambil ulang agar user dict membawa data terbaru
            row = conn.execute("SELECT * FROM users WHERE id=?", (row["id"],)).fetchone()
            conn.close()
            log.info(f"Login sukses: '{uname}' (user_id={row['id']})")
            return {"ok": True, "user": dict(row)}

        # ── Password salah: naikkan counter, lock jika melewati batas ──
        attempts = (row["failed_attempts"] or 0) + 1
        if attempts >= LOGIN_MAX_ATTEMPTS:
            lock_until = (datetime.now()
                          + timedelta(minutes=LOGIN_LOCK_MINUTES)).isoformat()
            conn.execute("UPDATE users SET failed_attempts=0, locked_until=? WHERE id=?",
                         (lock_until, row["id"]))
            conn.commit()
            conn.close()
            log.warning(f"Lockout {LOGIN_LOCK_MINUTES} mnt untuk '{uname}' "
                        f"setelah {LOGIN_MAX_ATTEMPTS}x gagal")
            return {"ok": False, "locked": True,
                    "msg": tr_db(lang="id", key="db_login_locked",
                                 mins=LOGIN_LOCK_MINUTES)}
        conn.execute("UPDATE users SET failed_attempts=? WHERE id=?",
                     (attempts, row["id"]))
        conn.commit()
        conn.close()
        log.info(f"Login gagal untuk '{uname}' (percobaan {attempts}/{LOGIN_MAX_ATTEMPTS})")
        return {"ok": False,
                "msg": tr_db(lang="id", key="db_login_failed_attempts",
                             left=LOGIN_MAX_ATTEMPTS - attempts)}

    conn.close()
    log.info(f"Login gagal: username '{uname}' tidak dikenal")
    return {"ok": False, "msg": tr_db(lang="id", key="db_login_failed")}


def change_password(user_id, old_pw, new_pw):
    u = get_user(user_id)
    if not _verify_password(old_pw, u.get("password_hash", "")):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_old_password_wrong")}
    ok_pw, key = validate_password_strength(new_pw)
    if not ok_pw:
        return {"ok": False, "msg": tr_db(user_id=user_id, key=key)}
    conn = get_conn()
    conn.execute("UPDATE users SET password_hash=? WHERE id=?",
                 (_hash_password(new_pw), user_id))
    conn.commit()
    conn.close()
    # Ganti password = semua session lama hangus (termasuk perangkat lain)
    delete_all_session_tokens(user_id)
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_password_changed")}


# ── User ──────────────────────────────────────────────────────────────────────

def get_user(user_id):
    conn = get_conn()
    row = conn.execute("SELECT * FROM users WHERE id=?", (user_id,)).fetchone()
    conn.close()
    return dict(row) if row else {}

@retry_on_lock
def update_user(user_id, **kw):
    if not kw:
        return
    fields = ", ".join(f"{k}=?" for k in kw)
    conn = get_conn()
    conn.execute(f"UPDATE users SET {fields} WHERE id=?",
                 list(kw.values()) + [user_id])
    conn.commit()
    conn.close()

@retry_on_lock
def recalculate_all_buffs(user_id):
    conn = get_conn()
    try:
        # Ambil data user
        u = conn.execute("SELECT * FROM users WHERE id=?", (user_id,)).fetchone()
        if not u:
            return
        u = dict(u)

        # Inventory (sertakan enchant level untuk scaling buff)
        inv_rows = conn.execute("SELECT item_id, COALESCE(enchant_level,0) AS el FROM inventory WHERE user_id=?", (user_id,)).fetchall()
        owned = {row["item_id"] for row in inv_rows}
        enchant_lvl = {}
        for row in inv_rows:
            enchant_lvl[row["item_id"]] = max(enchant_lvl.get(row["item_id"], 0), row["el"])

        # Hitung buff dari item
        dmg = 0.0
        xp_pct = 0.0
        gold_pct = 0.0
        reduc = 0.0
        mp = 0
        revive = 0
        has_spyglass = False
        crit = 0
        block_chance = 20
        block_strength = 10
        for iid in owned:
            b = SHOP_ITEMS.get(iid, {}).get("buff", {})
            # Enchant: +ENCHANT_BUFF_BONUS kekuatan buff per level (item consumable tidak dihitung di sini)
            mult = 1.0 + ENCHANT_BUFF_BONUS * enchant_lvl.get(iid, 0)
            dmg += b.get("boss_dmg", 0) * mult
            xp_pct += b.get("xp_pct", 0) / 100 * mult
            gold_pct += b.get("gold_pct", 0) / 100 * mult
            reduc += b.get("hp_reduc", 0) * mult
            mp += int(b.get("mp_bonus", 0) * mult)
            crit += b.get("crit_chance", 0) * mult
            block_chance += b.get("block_chance", 0) * mult
            block_strength += b.get("block_strength", 0) * mult
            if b.get("revive"):
                revive = 1
            if iid == "spyglass":
                has_spyglass = True

        # Ambil semua pet aktif
        active_pets = conn.execute(
            "SELECT pet_id, level FROM user_pets WHERE user_id=? AND is_active=1",
            (user_id,)
        ).fetchall()
        for pet_row in active_pets:
            pid = pet_row["pet_id"]
            base = PETS_DATA.get(pid, {}).get("base_buff", {})
            level = pet_row["level"]
            scale = 1 + (level - 1) * 0.05   # scaling 5% per level
            xp_pct += base.get("xp_pct", 0) * scale / 100   
            gold_pct += base.get("gold_pct", 0) * scale / 100
            dmg += base.get("boss_dmg", 0) * scale
            reduc += base.get("hp_reduc", 0) * scale

        # Guild buff
        gid = u.get("guild_id")
        if gid:
            guild = conn.execute("SELECT buff_xp, buff_gold, buff_damage FROM guilds WHERE id=?", (gid,)).fetchone()
            if guild:
                xp_pct += guild["buff_xp"] / 100
                gold_pct += guild["buff_gold"] / 100
                dmg += guild["buff_damage"]

        # Rebirth buffs
        rebirth_count = u.get("rebirth_count", 0)
        xp_pct += rebirth_count * 0.10
        gold_pct += rebirth_count * 0.05

        # Class passive buffs
        class_row = conn.execute("SELECT class_passive_buffs FROM users WHERE id=?", (user_id,)).fetchone()
        class_buffs = {}
        if class_row and class_row["class_passive_buffs"]:
            import json
            class_buffs = json.loads(class_row["class_passive_buffs"])
        xp_pct += class_buffs.get("xp_multiplier", 1.0) - 1.0
        gold_pct += class_buffs.get("gold_multiplier", 1.0) - 1.0

        # Talent tree (node pasif yang di-unlock user)
        try:
            unlocked = conn.execute(
                "SELECT talent_key FROM user_talents WHERE user_id=?", (user_id,)
            ).fetchall()
            for trow in unlocked:
                teff = TALENTS.get(trow["talent_key"], {}).get("buff", {})
                dmg += teff.get("boss_dmg", 0)
                xp_pct += teff.get("xp_pct", 0) / 100
                gold_pct += teff.get("gold_pct", 0) / 100
                reduc += teff.get("hp_reduc", 0)
                mp += teff.get("mp_bonus", 0)
        except Exception:
            pass  # tabel user_talents mungkin belum ada pada DB lama pra-migrasi

        # Hitung MP baru
        base_mp = 30 + (u["level"] - 1) * 5
        new_max_mp = base_mp + mp
        current_mp = u["mp"]
        if current_mp > new_max_mp:
            current_mp = new_max_mp

        # Update user dalam satu pernyataan
        conn.execute("""UPDATE users SET
            boss_damage_bonus=?,
            xp_multiplier=?,
            gold_multiplier=?,
            hp_damage_reduction=?,
            mp_bonus=?,
            max_mp=?,
            mp=?,
            has_revive=?,
            has_spyglass=?,
            crit_chance=?,
            block_chance=?,
            block_strength=?
            WHERE id=?""",
            (dmg, 
            round(1.0 + xp_pct, 4), 
            round(1.0 + gold_pct, 4),
            reduc, 
            mp, 
            new_max_mp, 
            current_mp, 
            revive, 
            1 if has_spyglass else 0, 
            crit + 10, 
            block_chance, 
            block_strength, 
            user_id))
        conn.commit()
    finally:
        conn.close()


# ── XP / Gold / HP ────────────────────────────────────────────────────────────

@retry_on_lock
def gain_xp_gold(user_id, xp_base, gold_base, skip_achievements=False):
    u = get_user(user_id)
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    if not u:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_user_not_found"), "leveled_up": False}
    
    buffs = get_skill_buffs(user_id)
    
    # Hitung XP multiplier dari skill
    skill_xp_mult = buffs.get("xp_multiplier", 1.0)
    if buffs.get("xp_remaining", 0) > 0:
        buffs["xp_remaining"] -= 1
        if buffs["xp_remaining"] == 0:
            buffs.pop("xp_multiplier", None)
            buffs.pop("xp_remaining", None)
        set_skill_buffs(user_id, buffs)
    else:
        skill_xp_mult = 1.0
    
    # Hitung Gold multiplier dari skill (Archer)
    skill_gold_mult = buffs.get("gold_multiplier", 1.0)
    if buffs.get("gold_remaining", 0) > 0:
        buffs["gold_remaining"] -= 1
        if buffs["gold_remaining"] == 0:
            buffs.pop("gold_multiplier", None)
            buffs.pop("gold_remaining", None)
        set_skill_buffs(user_id, buffs)
    else:
        skill_gold_mult = 1.0
    
    # Hitung multiplier dasar
    total_xp_mult = u.get("xp_multiplier", 1.0) * skill_xp_mult
    total_gold_mult = u.get("gold_multiplier", 1.0) * skill_gold_mult

    # Jika admin, kalikan 5
    if u.get("is_admin", 0):
        total_xp_mult *= 5
        total_gold_mult *= 5

    xp = int(xp_base * total_xp_mult)
    # Gold SELALU integer — multiplier (pet/skill/admin) bisa menghasilkan
    # pecahan; tanpa round, kolom gold menyimpan float dan tampil "1234.56".
    gold = int(round(gold_base * total_gold_mult))
    new_xp   = u["xp"] + xp
    new_gold = u["gold"] + gold
    new_lvl  = u["level"]
    leveled  = False
    needed   = new_lvl * 150
    while new_xp >= needed:
        new_xp -= needed
        new_lvl += 1
        leveled = True
        mhp = 50 + (new_lvl - 1) * 10
        mmp = 30 + (new_lvl - 1) * 5
        class_buffs = get_class_passive_buffs(user_id)
        hp_mult = class_buffs.get("hp_multiplier", 1.0)
        mhp = int(mhp * hp_mult)
        mp_bonus = u.get("mp_bonus", 0)
        mmp += mp_bonus
        # BUGFIX: level harus ikut ditulis SEKARANG — jika tidak, apply_hp_multiplier
        # membaca level lama dari DB dan menimpa max_hp yang baru saja diset.
        update_user(user_id, level=new_lvl, max_hp=mhp, hp=mhp, max_mp=mmp, mp=mmp)
        add_notification(user_id, tr_db(user_id=user_id, key="db_level_up", lvl=new_lvl), "levelup")
        apply_hp_multiplier(user_id)
        needed = new_lvl * 150
    if leveled and not skip_achievements:
        check_achievements(user_id, "level_reach", new_lvl)
    update_user(user_id,
                xp=new_xp, level=new_lvl, gold=new_gold,
                total_xp_earned=u.get("total_xp_earned", 0) + xp,
                total_gold_earned=u.get("total_gold_earned", 0.0) + gold)
    if not skip_achievements:
        check_achievements(user_id, "total_gold", u.get("total_gold_earned", 0) + gold) 
    log_activity(user_id, "reward", tr_db(user_id=user_id, key="log_reward", xp=xp, gold=gold), xp, gold)
    u2 = get_user(user_id)
    if u2.get("guild_id"):
        add_guild_exp(u2["guild_id"], xp_base // 5) 
    return {"ok": True, "leveled_up": leveled, "new_level": new_lvl,
            "new_xp": new_xp, "xp_gained": xp, "gold_gained": gold}


def lose_hp(user_id, base_amount, ignore_reduction=False):
    u = get_user(user_id)
    if is_account_locked(user_id):
        return {"revived": False, "new_hp": 0}
    if not u:
        return {"revived": False, "new_hp": 0}
    level = u.get("level", 1)
    additional = min(25, level * 0.5)
    total_damage = base_amount + additional
    if ignore_reduction:
        reduc = 0
    else:
        reduc = u.get("hp_damage_reduction", 0)
    actual = max(0.0, total_damage - reduc)
    actual_int = int(actual)                     # bulatkan ke bawah
    new_hp = max(0, u["hp"] - actual_int)
    new_hp_int = int(new_hp)                     # pastikan integer

    if new_hp_int == 0 and u.get("has_revive"):
        new_hp_int = int(u["max_hp"] * 0.3)
        update_user(user_id, hp=new_hp_int, has_revive=0)
        recalculate_all_buffs(user_id)
        add_notification(user_id, tr_db(user_id=user_id, key="db_totem_revive"), "success")
        return {"revived": True, "new_hp": new_hp_int, "damage_taken": actual_int}

    update_user(user_id, hp=new_hp_int)
    if new_hp_int == 0:
        add_notification(user_id, tr_db(user_id=user_id, key="db_hp_zero"), "danger")
    return {"revived": False, "new_hp": new_hp_int, "damage_taken": actual_int}

def penalize_gold(user_id, base_amount=5):
    u = get_user(user_id)
    if is_account_locked(user_id):
        return {"gold_lost": 0, "remaining_gold": 0}
    if not u:
        return {"gold_lost": 0, "remaining_gold": 0}
    level = u.get("level", 1)
    additional = min(25, level * 0.5)
    total_penalty = int(base_amount + additional)
    current_gold = u.get("gold", 0)
    if current_gold <= 0:
        return {"gold_lost": 0, "remaining_gold": 0}
    gold_lost = min(total_penalty, current_gold)
    new_gold = current_gold - gold_lost
    update_user(user_id, gold=new_gold)
    log_activity(user_id, "penalty_gold", tr_db(user_id=user_id, key="log_penalty_gold", gold=gold_lost), 0, -gold_lost)
    return {"gold_lost": gold_lost, "remaining_gold": new_gold}


def penalize_xp(user_id):
    """
    Kurangi XP user sebesar 10% dari XP yang dimiliki saat itu.
    Minimal 1 XP (jika XP > 0).
    Bisa menyebabkan turun level jika XP tidak cukup untuk menahan penalti.
    """
    u = get_user(user_id)
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    if not u:
        return {"xp_lost": 0, "new_level": 1}
    
    current_xp = u.get("xp", 0)
    current_level = u.get("level", 1)
    
    # Hitung 10% dari XP yang dimiliki, minimal 1 XP
    if current_xp <= 0:
        return {"xp_lost": 0, "new_level": current_level}
    
    xp_lost = max(1, int(current_xp * 0.1))  # 10% dari XP, minimal 1
    
    new_xp = current_xp - xp_lost
    new_level = current_level
    
    # Turunkan level jika XP negatif
    while new_xp < 0 and new_level > 1:
        new_level -= 1
        new_xp += new_level * 150
    if new_xp < 0:
        new_xp = 0
    
    update_user(user_id, xp=new_xp, level=new_level)
    log_activity(
        user_id,
        "penalty_xp",
        tr_db(user_id=user_id, key="log_penalty_xp", xp=xp_lost, level=new_level),
        -xp_lost,
        0
    )
    add_notification(
        user_id,
        tr_db(user_id=user_id, key="penalty_xp_loss_percent", xp=xp_lost, new_level=new_level),
        "danger"
    )
    
    return {"xp_lost": xp_lost, "new_level": new_level, "new_xp": new_xp}

def restore_mp(user_id, amount=5):
    u = get_user(user_id)
    cap = u["max_mp"] + u.get("mp_bonus", 0)
    update_user(user_id, mp=min(cap, u["mp"] + amount))


# ── Class Skills ──────────────────────────────────────────────────────────────

CLASS_SKILLS = {
    "warrior": {"name": "Shield Bash",  "icon": "🛡️",
                "mp_cost": 10,
                "desc": "Kurangi 50% damage dari boss satu kali"},
    "mage":    {"name": "Arcane Surge", "icon": "✨",
                "mp_cost": 15,
                "desc": "+30% XP untuk 3 habit berikutnya"},
    "archer":  {"name": "Gold Shot",    "icon": "🏹",
                "mp_cost": 10,
                "desc": "+50% Gold dari habit berikutnya"},
    "healer":  {"name": "Regenerate",   "icon": "💚",
                "mp_cost": 20,
                "desc": "Pulihkan 20 HP sekarang juga"},
    "rogue":   {"name": "Shadow Step",  "icon": "🗡️",
                "mp_cost": 15,
                "desc": "+1 streak untuk daily berikutnya"},
}


def use_class_skill(user_id):
    u = get_user(user_id)
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    cls = u.get("avatar_class", "warrior")
    skill = CLASS_SKILLS.get(cls, {})
    cost = skill.get("mp_cost", 10)
    
    # Admin: gratis, tanpa cek MP
    if u.get("is_admin", 0):
        cost = 0
    elif u["mp"] < cost:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_mp_insufficient", cost=cost, mp=u['mp'])}
    
    # Kurangi MP (hanya jika bukan admin atau admin dengan cost 0)
    update_user(user_id, mp=u["mp"] - cost)
    
    # Handle skill Healer
    if cls == "healer":
        new_hp = min(u["max_hp"], u["hp"] + 30)
        update_user(user_id, hp=new_hp)
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_heal_regenerate")}
    
    # Skill lain
    result = apply_skill_effect(user_id, cls)
    if result["ok"]:
        add_notification(user_id, result["msg"], "info")
    return result

def get_skill_buffs(user_id):
    """Ambil data buff skill user dari database (JSON)"""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    row = conn.execute("SELECT skill_buff_data FROM users WHERE id=?", (user_id,)).fetchone()
    conn.close()
    if row and row["skill_buff_data"]:
        return json.loads(row["skill_buff_data"])
    return {}

def set_skill_buffs(user_id, buffs):
    """Simpan data buff skill user ke database"""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("UPDATE users SET skill_buff_data=? WHERE id=?", (json.dumps(buffs), user_id))
    conn.commit()
    conn.close()

def get_class_passive_buffs(user_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    row = conn.execute("SELECT class_passive_buffs FROM users WHERE id=?", (user_id,)).fetchone()
    conn.close()
    if row and row["class_passive_buffs"]:
        return json.loads(row["class_passive_buffs"])
    return {}

def set_class_passive_buffs(user_id, buffs):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("UPDATE users SET class_passive_buffs=? WHERE id=?", (json.dumps(buffs), user_id))
    conn.commit()
    conn.close()

def update_class_passive_buffs(user_id):
    u = get_user(user_id)
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    cls = u.get("avatar_class", "warrior")
    buffs = {}
    if cls == "warrior":
        buffs = {"hp_multiplier": 1.10, "hp_bonus": 0}   # +10% HP (dari 20%)
    elif cls == "mage":
        buffs = {"xp_multiplier": 1.08}                  # +8% XP (dari 15%)
    elif cls == "archer":
        buffs = {"gold_multiplier": 1.06}                # +6% Gold (dari 10%)
    elif cls == "rogue":
        buffs = {"streak_bonus": 1}                      # +1 streak awal untuk daily/habit
    # Healer tidak punya buff pasif (sudah punya skill penyembuhan)
    set_class_passive_buffs(user_id, buffs)
    recalculate_all_buffs(user_id)

def apply_skill_effect(user_id, skill_name):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    """Terapkan efek skill ke user (simpan state)"""
    buffs = get_skill_buffs(user_id)
    if skill_name == "warrior":
        buffs["shield_active"] = True
        msg = tr_db(user_id=user_id, key="db_skill_shield")
    elif skill_name == "mage":
        buffs["xp_multiplier"] = 1.3
        buffs["xp_remaining"] = 3
        msg = tr_db(user_id=user_id, key="db_skill_arcane")
    elif skill_name == "archer":
        buffs["gold_multiplier"] = 1.5
        buffs["gold_remaining"] = 1
        msg = tr_db(user_id=user_id, key="db_skill_gold_shot")
    elif skill_name == "rogue":
        buffs["double_streak"] = 1
        msg = tr_db(user_id=user_id, key="db_skill_shadow")
    else:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_skill_unknown")}
    
    set_skill_buffs(user_id, buffs)
    return {"ok": True, "msg": msg}

# ── Habits ────────────────────────────────────────────────────────────────────

_XP  = {"trivial": 8,  "easy": 15, "medium": 25, "hard": 40, "epic": 60}
_GLD = {"trivial": 2,  "easy": 3,  "medium": 5,  "hard": 8,  "epic": 12}

# Daily rewards
_DAILY_XP = {"easy": 20, "medium": 30, "hard": 50, "epic": 75}
_DAILY_GLD = {"easy": 4, "medium": 6, "hard": 10, "epic": 15}

# Todo rewards
_TODO_XP = {"trivial": 10, "easy": 20, "medium": 40, "hard": 60}
_TODO_GLD = {"trivial": 2, "easy": 4, "medium": 8, "hard": 14}


def get_habits(user_id):
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM habits WHERE user_id=? ORDER BY sort_order",
        (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def add_habit(user_id, name, icon="⚔️", difficulty="medium", positive=1, negative=0, notes="", repeat_days=""):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    max_order = conn.execute(
        "SELECT COALESCE(MAX(sort_order), 0) FROM habits WHERE user_id=? AND folder_id IS NULL",
        (user_id,)
    ).fetchone()[0]
    new_order = max_order + 1

    conn.execute(
        "INSERT INTO habits(user_id,name,icon,difficulty,"
        "xp_reward,gold_reward,positive,negative,notes,sort_order,repeat_days)"
        " VALUES(?,?,?,?,?,?,?,?,?,?,?)",
        (user_id, name, icon, difficulty,
         _XP.get(difficulty, 25), _GLD.get(difficulty, 5),
         positive, negative, notes, new_order, repeat_days or "")
    )
    conn.commit()
    conn.close()

def update_habit(habit_id, user_id, **kwargs):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    if "difficulty" in kwargs:
        diff = kwargs["difficulty"]
        kwargs.setdefault("xp_reward", _XP.get(diff, 25))
        kwargs.setdefault("gold_reward", _GLD.get(diff, 5))
    
    if not kwargs:
        return
    fields = ", ".join(f"{k}=?" for k in kwargs)
    conn = get_conn()
    conn.execute(f"UPDATE habits SET {fields} WHERE id=? AND user_id=?", list(kwargs.values()) + [habit_id, user_id])
    conn.commit()
    conn.close()

@retry_on_lock
def complete_habit(user_id, habit_id, direction="up"):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    try:
        h = conn.execute("SELECT * FROM habits WHERE id=? AND user_id=?", (habit_id, user_id)).fetchone()
        if not h:
            return {"ok": False}
        today = date.today().isoformat()
        if h["done_today"]:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_habit_already_done")}
        new_streak = h["streak"] + 1 if direction == "up" else 0
        conn.execute("""UPDATE habits SET done_today=1, streak=?, last_done=?, last_action=?, 
                        counter_up=counter_up+?, counter_down=counter_down+? WHERE id=?""",
                     (new_streak, today, direction, 1 if direction=="up" else 0, 1 if direction=="down" else 0, habit_id))
        conn.execute("UPDATE users SET total_habits_done=total_habits_done+1 WHERE id=?", (user_id,))
        conn.execute("UPDATE users SET total_tasks_completed = total_tasks_completed + 1 WHERE id=?", (user_id,))
        new_total = conn.execute("SELECT total_tasks_completed FROM users WHERE id=?", (user_id,)).fetchone()[0]
        conn.commit()
    finally:
        conn.close()

    if direction == "up":
        restore_mp(user_id, 3)
        u = get_user(user_id)
        if new_streak > u.get("longest_streak", 0):
            update_user(user_id, longest_streak=new_streak)
        # Panggil achievements
        check_achievements(user_id, "habit_complete", 1)
        check_achievements(user_id, "habit_streak", new_streak)
        check_achievements(user_id, "total_tasks", new_total)
        today = date.today().isoformat()
        log_task_history(user_id, "habit", habit_id, "success", today)
        return gain_xp_gold(user_id, h["xp_reward"], h["gold_reward"])
    else:
        today = date.today().isoformat()
        log_task_history(user_id, "habit", habit_id, "fail", today)
        damage_result = lose_hp(user_id, 5, ignore_reduction=True)
        return {"ok": True, "lost_hp": damage_result.get("damage_taken", 5)}

def add_guild_exp(guild_id, amount):
    # ── Step 1: Hitung level & update guilds, lalu TUTUP koneksi ─────────────
    conn = get_conn()
    g = conn.execute("SELECT * FROM guilds WHERE id=?", (guild_id,)).fetchone()
    if not g:
        conn.close()
        return
    g = dict(g)
    new_exp = g["exp"] + amount
    new_level = g["level"]
    needed = new_level * 500
    leveled = False
    member_ids = []
    while new_exp >= needed:
        new_exp -= needed
        new_level += 1
        leveled = True
        needed = new_level * 500
    if leveled:
        buff_xp = new_level * 2
        buff_gold = new_level * 1
        buff_damage = new_level * 1
        buff_crit = new_level * 1
        conn.execute(
            "UPDATE guilds SET level=?, exp=?, buff_xp=?, buff_gold=?, buff_damage=?, crit_chance=? WHERE id=?",
            (new_level, new_exp, buff_xp, buff_gold, buff_damage, buff_crit, guild_id))
        members = conn.execute("SELECT user_id FROM guild_members WHERE guild_id=?", (guild_id,)).fetchall()
        member_ids = [m["user_id"] for m in members]
    else:
        conn.execute("UPDATE guilds SET exp=? WHERE id=?", (new_exp, guild_id))
    conn.commit()
    conn.close()   # ← TUTUP dulu sebelum panggil add_notification / recalculate_all_buffs

    # ── Step 2: Notifikasi & recalc buff (koneksi terpisah, tidak nested) ─────
    if leveled:
        for uid in member_ids:
            add_notification(uid, tr_db(user_id=uid, key="db_guild_level_up", name=g['name'], lvl=new_level), "levelup")
            recalculate_all_buffs(uid)

def delete_habit(user_id, habit_id):
    purge_trash()
    conn = get_conn()
    trash_id = _stash_before_delete(conn, user_id, "habits", habit_id, "habit")
    conn.execute("DELETE FROM habits WHERE id=? AND user_id=?",
                 (habit_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True, "trash_id": trash_id}


@retry_on_lock
def reset_daily_tasks(user_id):
    today = date.today().isoformat()
    conn = get_conn()
    conn.execute(
        "UPDATE habits SET done_today=0"
        " WHERE user_id=? AND (last_done IS NULL OR last_done!=?)",
        (user_id, today))
    conn.execute(
        "UPDATE dailies SET done_today=0"
        " WHERE user_id=? AND (last_done IS NULL OR last_done!=?)",
        (user_id, today))
    conn.execute(
        "UPDATE sport_activities SET done_today=0"
        " WHERE user_id=? AND (last_done IS NULL OR last_done!=?)",
        (user_id, today))
    conn.commit()
    fill_skipped_history(user_id) 
    conn.close()


# ── Dailies ───────────────────────────────────────────────────────────────────
def get_dailies(user_id):
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM dailies WHERE user_id=? ORDER BY sort_order",
        (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def add_daily(user_id, name, icon="📅", difficulty="medium", notes="", repeat_days=""):
    conn = get_conn()
    # Hitung sort_order terakhir untuk user ini tanpa folder
    max_order = conn.execute(
        "SELECT COALESCE(MAX(sort_order), 0) FROM dailies WHERE user_id=? AND folder_id IS NULL",
        (user_id,)
    ).fetchone()[0]
    new_order = max_order + 1

    conn.execute(
        "INSERT INTO dailies(user_id,name,icon,difficulty,"
        "xp_reward,gold_reward,notes,sort_order,repeat_days) VALUES(?,?,?,?,?,?,?,?,?)",
        (user_id, name, icon, difficulty,
         _DAILY_XP.get(difficulty, 30), _DAILY_GLD.get(difficulty, 6), notes, new_order,
         repeat_days or "")
    )
    conn.commit()
    conn.close()

def update_daily(daily_id, user_id, **kwargs):
    if "difficulty" in kwargs:
        diff = kwargs["difficulty"]
        kwargs.setdefault("xp_reward", _DAILY_XP.get(diff, 30))
        kwargs.setdefault("gold_reward", _DAILY_GLD.get(diff, 6))
    
    if not kwargs:
        return
    fields = ", ".join(f"{k}=?" for k in kwargs)
    conn = get_conn()
    conn.execute(f"UPDATE dailies SET {fields} WHERE id=? AND user_id=?", list(kwargs.values()) + [daily_id, user_id])
    conn.commit()
    conn.close()

@retry_on_lock
def complete_daily(user_id, daily_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    try:
        d = conn.execute("SELECT * FROM dailies WHERE id=? AND user_id=?", (daily_id, user_id)).fetchone()
        if not d or d["done_today"]:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_daily_already_done")}
        d = dict(d)
    finally:
        conn.close()

    today = date.today().isoformat()
    buffs = get_skill_buffs(user_id)
    streak_bonus = 2 if buffs.get("double_streak") else 1
    if buffs.get("double_streak"):
        buffs.pop("double_streak", None)
        set_skill_buffs(user_id, buffs)
    new_streak = d["streak"] + streak_bonus

    conn2 = get_conn()
    try:
        conn2.execute("UPDATE dailies SET done_today=1, streak=?, last_done=?, last_action='up', fail_streak=0 WHERE id=?", (new_streak, today, daily_id))
        conn2.execute("UPDATE users SET total_dailies_done=total_dailies_done+1 WHERE id=?", (user_id,))
        conn2.execute("UPDATE users SET total_tasks_completed = total_tasks_completed + 1 WHERE id=?", (user_id,))
        new_total = conn2.execute("SELECT total_tasks_completed FROM users WHERE id=?", (user_id,)).fetchone()[0]
        conn2.commit()
    finally:
        conn2.close()

    restore_mp(user_id, 5)
    check_achievements(user_id, "daily_complete", 1)
    check_achievements(user_id, "total_tasks", new_total)
    today = date.today().isoformat()
    log_task_history(user_id, "daily", daily_id, "success", today)
    return gain_xp_gold(user_id, d["xp_reward"], d["gold_reward"])

@retry_on_lock
def fail_daily(user_id, daily_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    d = conn.execute("SELECT * FROM dailies WHERE id=? AND user_id=?", (daily_id, user_id)).fetchone()
    if not d or d["done_today"]:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_daily_fail_already")}
    
    today = date.today().isoformat()
    
    # ── CEK FREEZE SLOT ──
    if d["freeze_slots"] > 0:
        new_freeze = d["freeze_slots"] - 1
        conn.execute("""
            UPDATE dailies 
            SET done_today=1, freeze_slots=?, last_done=?, last_action='frozen' 
            WHERE id=?
        """, (new_freeze, today, daily_id))
        conn.execute("UPDATE users SET total_dailies_done=total_dailies_done+1 WHERE id=?", (user_id,))
        conn.commit()
        conn.close()
        log_task_history(user_id, "daily", daily_id, "freeze", today)
        return {
            "ok": True,
            "freeze_used": True,
            "remaining_freezes": new_freeze,
            "msg": tr_db(user_id, "freeze_used_message", remaining=new_freeze)
        }
    
    # ── LOGIKA PENALTI NORMAL (jika tidak ada freeze) ──
    new_fail_streak = (d["fail_streak"] or 0) + 1
    conn.execute("""
        UPDATE dailies 
        SET done_today=1, streak=0, fail_streak=?, last_done=?, last_action='down' 
        WHERE id=?
    """, (new_fail_streak, today, daily_id))
    conn.execute("UPDATE users SET total_dailies_done=total_dailies_done+1 WHERE id=?", (user_id,))
    conn.commit()
    conn.close()

    log_task_history(user_id, "daily", daily_id, "fail", today)

    u = get_user(user_id)
    damage_taken = 0
    penalty_type = "hp"
    penalty_amount = 0

    if u.get("hp", 0) > 0:
        hp_result = lose_hp(user_id, 5, ignore_reduction=True)
        damage_taken = hp_result.get("damage_taken", 0)
        penalty_type = "hp"
        penalty_amount = damage_taken
        u = get_user(user_id)

    if u.get("hp", 0) <= 0:
        gold_result = penalize_gold(user_id, base_amount=5)
        gold_lost = gold_result.get("gold_lost", 0)
        if gold_lost > 0:
            penalty_type = "gold"
            penalty_amount = gold_lost
            add_notification(user_id, tr_db(user_id=user_id, key="penalty_hp_zero_gold", gold=gold_lost), "warning")
        else:
            xp_result = penalize_xp(user_id)
            xp_lost = xp_result.get("xp_lost", 0)
            if xp_lost > 0:
                penalty_type = "xp"
                penalty_amount = xp_lost
            else:
                penalty_type = "none"
                penalty_amount = 0
                add_notification(user_id, tr_db(user_id=user_id, key="penalty_all_failed"), "error")

    return {
        "ok": True,
        "lost_hp": damage_taken,
        "fail_streak": new_fail_streak,
        "penalty_type": penalty_type,
        "penalty_amount": penalty_amount
    }

def delete_daily(user_id, daily_id):
    purge_trash()
    conn = get_conn()
    trash_id = _stash_before_delete(conn, user_id, "dailies", daily_id, "daily")
    conn.execute("DELETE FROM dailies WHERE id=? AND user_id=?",
                 (daily_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True, "trash_id": trash_id}

def add_freeze_to_daily(user_id, daily_id):
    """Gunakan 1 Ice Block untuk menambah 1 freeze slot ke Daily tertentu (max 3)."""
    conn = get_conn()
    try:
        daily = conn.execute(
            "SELECT id, freeze_slots FROM dailies WHERE id=? AND user_id=?",
            (daily_id, user_id)
        ).fetchone()
        if not daily:
            return {"ok": False, "msg": tr_db(user_id, "db_daily_not_found")}
        
        if daily["freeze_slots"] >= 3:
            return {"ok": False, "msg": tr_db(user_id, "freeze_max_reached")}
        
        inv = conn.execute(
            "SELECT id FROM inventory WHERE user_id=? AND item_id='ice_block' AND quantity>0",
            (user_id,)
        ).fetchone()
        if not inv:
            return {"ok": False, "msg": tr_db(user_id, "freeze_no_item")}
        
        conn.execute("UPDATE inventory SET quantity = quantity - 1 WHERE id=?", (inv["id"],))
        new_freeze = daily["freeze_slots"] + 1
        conn.execute(
            "UPDATE dailies SET freeze_slots = freeze_slots + 1 WHERE id=?",
            (daily_id,)
        )
        conn.commit()
        return {
            "ok": True,
            "remaining": new_freeze,
            "msg": tr_db(user_id, "freeze_added_success", remaining=new_freeze)
        }
    except Exception as e:
        conn.rollback()
        return {"ok": False, "msg": tr_db(user_id, "freeze_error", error=str(e))}
    finally:
        conn.close()


# ── Todos ─────────────────────────────────────────────────────────────────────

def get_todos(user_id):
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM todos WHERE user_id=? ORDER BY sort_order",
        (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def add_todo(user_id, name, priority="medium", icon="📜",
             due_date=None, notes=""):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    # Hitung sort_order terakhir untuk user ini tanpa folder
    max_order = conn.execute(
        "SELECT COALESCE(MAX(sort_order), 0) FROM todos WHERE user_id=? AND folder_id IS NULL",
        (user_id,)
    ).fetchone()[0]
    new_order = max_order + 1

    conn.execute(
        "INSERT INTO todos(user_id,name,icon,priority,"
        "xp_reward,gold_reward,due_date,notes,sort_order) VALUES(?,?,?,?,?,?,?,?,?)",
        (user_id, name, icon, priority,
         _TODO_XP.get(priority, 40), _TODO_GLD.get(priority, 8), due_date, notes, new_order)
    )
    conn.commit()
    conn.close()

def update_todo(todo_id, user_id, **kwargs):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    if "priority" in kwargs:
        prio = kwargs["priority"]
        kwargs.setdefault("xp_reward", _TODO_XP.get(prio, 40))
        kwargs.setdefault("gold_reward", _TODO_GLD.get(prio, 8))
    
    if not kwargs:
        return
    fields = ", ".join(f"{k}=?" for k in kwargs)
    conn = get_conn()
    conn.execute(f"UPDATE todos SET {fields} WHERE id=? AND user_id=?", list(kwargs.values()) + [todo_id, user_id])
    conn.commit()
    conn.close()

@retry_on_lock
def complete_todo(user_id, todo_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    try:
        t = conn.execute("SELECT * FROM todos WHERE id=? AND user_id=?", (todo_id, user_id)).fetchone()
        if not t or t["done"]:
            return {"ok": False}
        conn.execute("UPDATE todos SET done=1 WHERE id=?", (todo_id,))
        conn.execute("UPDATE users SET total_todos_done=total_todos_done+1 WHERE id=?", (user_id,))
        conn.execute("UPDATE users SET total_tasks_completed = total_tasks_completed + 1 WHERE id=?", (user_id,))
        new_total = conn.execute("SELECT total_tasks_completed FROM users WHERE id=?", (user_id,)).fetchone()[0]
        xp_reward = t["xp_reward"]
        gold_reward = t["gold_reward"]
        conn.commit()
    finally:
        conn.close()
    
    # Panggil log_task_history setelah transaksi selesai untuk menghindari deadlock
    today = date.today().isoformat()
    log_task_history(user_id, "todo", todo_id, "success", today)
    
    restore_mp(user_id, 4)
    check_achievements(user_id, "todo_complete", 1)
    check_achievements(user_id, "total_tasks", new_total)
    return gain_xp_gold(user_id, xp_reward, gold_reward)

def delete_todo(user_id, todo_id):
    purge_trash()
    conn = get_conn()
    if is_account_locked(user_id):
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    trash_id = _stash_before_delete(conn, user_id, "todos", todo_id, "todo")
    conn.execute("DELETE FROM todos WHERE id=? AND user_id=?",
                 (todo_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True, "trash_id": trash_id}

# ═══════════════════════════════════════════════════════════════════════════
#  TASK FOLDERS  (shared by habits / dailies / todos / sport)
# ═══════════════════════════════════════════════════════════════════════════

def get_task_folders(user_id, mode):
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM task_folders WHERE user_id=? AND mode=? ORDER BY created_at",
        (user_id, mode)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def add_task_folder(user_id, mode, name, icon="📁"):
    conn = get_conn()
    cur = conn.execute(
        "INSERT INTO task_folders(user_id,mode,name,icon) VALUES(?,?,?,?)",
        (user_id, mode, name, icon))
    folder_id = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "folder_id": folder_id}


def update_task_folder(folder_id, user_id, **kwargs):
    if not kwargs:
        return
    conn = get_conn()
    fields = ", ".join(f"{k}=?" for k in kwargs)
    conn.execute(
        f"UPDATE task_folders SET {fields} WHERE id=? AND user_id=?",
        list(kwargs.values()) + [folder_id, user_id])
    conn.commit()
    conn.close()


def delete_task_folder(user_id, folder_id, mode):
    """Hapus folder dan lepas semua item dari folder tersebut (set folder_id=NULL)."""
    conn = get_conn()
    table_map = {"habit": "habits", "daily": "dailies",
                 "todo": "todos", "sport": "sport_activities"}
    tbl = table_map.get(mode)
    if tbl:
        conn.execute(f"UPDATE {tbl} SET folder_id=NULL WHERE folder_id=? AND user_id=?",
                     (folder_id, user_id))
    conn.execute("DELETE FROM task_folders WHERE id=? AND user_id=?",
                 (folder_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True}


def duplicate_task_folder(user_id, folder_id, mode):
    """Duplikasi folder beserta semua item di dalamnya."""
    conn = get_conn()
    folder = conn.execute(
        "SELECT * FROM task_folders WHERE id=? AND user_id=?",
        (folder_id, user_id)).fetchone()
    if not folder:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_folder_not_found")}

    # Buat folder baru
    cur = conn.execute(
        "INSERT INTO task_folders(user_id,mode,name,icon) VALUES(?,?,?,?)",
        (user_id, mode, folder["name"], folder["icon"]))
    new_folder_id = cur.lastrowid
    conn.commit()
    conn.close()

    # Duplikasi semua item dalam folder
    table_map = {"habit": "habits", "daily": "dailies",
                 "todo": "todos", "sport": "sport_activities"}
    tbl = table_map.get(mode)
    if tbl:
        conn2 = get_conn()
        items = conn2.execute(
            f"SELECT * FROM {tbl} WHERE folder_id=? AND user_id=?",
            (folder_id, user_id)).fetchall()
        for item in items:
            d = dict(item)
            d.pop("id", None)
            d["folder_id"] = new_folder_id
            # Reset progress fields
            for field in ("done_today","done","streak","last_done","last_action","counter_up","counter_down"):
                if field in d:
                    d[field] = 0 if field != "last_done" and field != "last_action" else ""
            cols   = ", ".join(d.keys())
            placeholders = ", ".join("?" * len(d))
            conn2.execute(f"INSERT INTO {tbl}({cols}) VALUES({placeholders})",
                          list(d.values()))
        conn2.commit()
        conn2.close()

    return {"ok": True, "new_folder_id": new_folder_id}


def set_item_folder(user_id, mode, item_id, folder_id):
    table_map = {
        "habit": "habits",
        "daily": "dailies",
        "todo": "todos",
        "sport": "sport_activities",
        "economy": "economy_items",
        "food": "food_logs"
    }
    tbl = table_map.get(mode)
    if not tbl:
        return

    conn = get_conn()
    try:
        if mode == "food":
            # folder_id di sini adalah meal_type (string)
            if folder_id not in ("breakfast", "lunch", "dinner", "snack", None):
                folder_id = "snack"  # fallback
            conn.execute(
                f"UPDATE {tbl} SET meal_type=? WHERE id=? AND user_id=?",
                (folder_id, item_id, user_id)
            )
            conn.commit()
            return

        # Mode lain: update folder_id dan sort_order
        if folder_id is not None:
            row = conn.execute(
                f"SELECT MAX(sort_order) as max_order FROM {tbl} WHERE user_id=? AND folder_id=?",
                (user_id, folder_id)
            ).fetchone()
        else:
            row = conn.execute(
                f"SELECT MAX(sort_order) as max_order FROM {tbl} WHERE user_id=? AND folder_id IS NULL",
                (user_id,)
            ).fetchone()
        max_order = row["max_order"] if row and row["max_order"] is not None else -1
        new_order = max_order + 1
        conn.execute(
            f"UPDATE {tbl} SET folder_id=?, sort_order=? WHERE id=? AND user_id=?",
            (folder_id, new_order, item_id, user_id)
        )
        conn.commit()
    finally:
        conn.close()


SPORT_TYPES = {
    "running":      {"name": "Lari",          "icon": "🏃"},
    "gym":          {"name": "Gym",            "icon": "🏋️"},
    "cycling":      {"name": "Bersepeda",      "icon": "🚴"},
    "swimming":     {"name": "Renang",         "icon": "🏊"},
    "yoga":         {"name": "Yoga",           "icon": "🧘"},
    "football":     {"name": "Olahraga Bola",  "icon": "⚽"},
    "calisthenics": {"name": "Kalistenik",     "icon": "🤸"},
    "martial_arts": {"name": "Bela Diri",      "icon": "🥊"},
    "badminton":    {"name": "Badminton",       "icon": "🏸"},
    "other":        {"name": "Lainnya",         "icon": "🏅"},
}

# XP / Gold / Sport Points reward per difficulty
_SP_XP  = {"easy": 15, "medium": 25, "hard": 40, "epic": 60}
_SP_GLD = {"easy": 3,  "medium": 5,  "hard": 8,  "epic": 12}
_SP_PTS = {"easy": 8,  "medium": 15, "hard": 25, "epic": 40}


def get_sport_activities(user_id, sport_type=None):
    conn = get_conn()
    if sport_type:
        rows = conn.execute(
            "SELECT * FROM sport_activities WHERE user_id=? AND sport_type=? ORDER BY sort_order",
            (user_id, sport_type)).fetchall()
    else:
        rows = conn.execute(
            "SELECT * FROM sport_activities WHERE user_id=? ORDER BY sort_order",
            (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def add_sport_activity(user_id, name, sport_type="running", icon="🏃",
                       difficulty="medium", notes="", calories_burned=0, duration_minutes=30):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    # Hitung sort_order terakhir untuk user ini tanpa folder
    max_order = conn.execute(
        "SELECT COALESCE(MAX(sort_order), 0) FROM sport_activities WHERE user_id=? AND folder_id IS NULL",
        (user_id,)
    ).fetchone()[0]
    new_order = max_order + 1

    conn.execute("""
        INSERT INTO sport_activities
        (user_id, name, sport_type, icon, difficulty,
         xp_reward, gold_reward, sport_points_reward, notes,
         calories_burned, duration_minutes, sort_order)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?)
    """, (user_id, name, sport_type, icon, difficulty,
          _SP_XP.get(difficulty, 25), _SP_GLD.get(difficulty, 5),
          _SP_PTS.get(difficulty, 15), notes,
          calories_burned, duration_minutes, new_order))
    conn.commit()
    conn.close()


def update_sport_activity(activity_id, user_id, **kwargs):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    # Jika difficulty berubah, sesuaikan reward otomatis
    if "difficulty" in kwargs:
        diff = kwargs["difficulty"]
        kwargs.setdefault("xp_reward",           _SP_XP.get(diff, 25))
        kwargs.setdefault("gold_reward",          _SP_GLD.get(diff, 5))
        kwargs.setdefault("sport_points_reward",  _SP_PTS.get(diff, 15))
    conn = get_conn()
    fields = ", ".join(f"{k}=?" for k in kwargs)
    conn.execute(
        f"UPDATE sport_activities SET {fields} WHERE id=? AND user_id=?",
        list(kwargs.values()) + [activity_id, user_id])
    conn.commit()
    conn.close()


@retry_on_lock
def delete_sport_activity(user_id, activity_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    row = conn.execute("SELECT last_done FROM sport_activities WHERE id=? AND user_id=?", (activity_id, user_id)).fetchone()
    log_date = row["last_done"] if row else None
    conn.execute("DELETE FROM sport_activities WHERE id=? AND user_id=?", (activity_id, user_id))
    conn.commit()
    conn.close()
    if log_date:
        update_daily_net_calories(user_id, log_date)


@retry_on_lock
def complete_sport_activity(user_id, activity_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    try:
        a = conn.execute("SELECT * FROM sport_activities WHERE id=? AND user_id=?", (activity_id, user_id)).fetchone()
        if not a:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_sport_activity_not_found")}
        if a["done_today"]:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_sport_activity_done_today")}
        today = date.today().isoformat()
        conn.execute("UPDATE sport_activities SET done_today=1, streak=streak+1, last_done=? WHERE id=?", (today, activity_id))
        conn.commit()
    finally:
        conn.close()
    today = date.today().isoformat()
    log_task_history(user_id, "sport", activity_id, "success", today)
    result = gain_xp_gold(user_id, a["xp_reward"], a["gold_reward"])
    sp_result = gain_sport_points(user_id, a["sport_points_reward"])
    result["sport_points_gained"] = a["sport_points_reward"]
    result["sport_leveled_up"] = sp_result.get("leveled_up", False)
    result["new_sport_level"] = sp_result.get("new_sport_level", 1)
    result["new_sport_xp"] = sp_result.get("new_sport_xp", 0)
    check_achievements(user_id, "sport_points", a["sport_points_reward"])
    check_achievements(user_id, "sport_streak", a["streak"] + 1)
    update_daily_net_calories(user_id, today)
    return result


def gain_sport_points(user_id, points):
    """Tambahkan sport points & hitung sport level.
    Sport level TIDAK mempengaruhi main level/XP user sama sekali.
    Formula naik level: butuh level*100 sport points.
    """
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    row = conn.execute(
        "SELECT COALESCE(sport_level,1) as sport_level,"
        " COALESCE(sport_xp,0) as sport_xp,"
        " COALESCE(total_sport_points_earned,0) as total_sport_points_earned"
        " FROM users WHERE id=?", (user_id,)).fetchone()
    if not row:
        conn.close()
        return {"ok": False}

    new_sport_xp   = row["sport_xp"] + points
    new_sport_lvl  = row["sport_level"]
    total_earned   = row["total_sport_points_earned"] + points
    leveled        = False
    needed         = new_sport_lvl * 100
    while new_sport_xp >= needed:
        new_sport_xp -= needed
        new_sport_lvl += 1
        leveled = True
        needed = new_sport_lvl * 100

    conn.execute(
        "UPDATE users SET sport_xp=?, sport_level=?, total_sport_points_earned=?"
        " WHERE id=?",
        (new_sport_xp, new_sport_lvl, total_earned, user_id))
    conn.commit()
    conn.close()

    if leveled:
        add_notification(user_id, tr_db(user_id=user_id, key="db_sport_level_up", lvl=new_sport_lvl), "levelup")

    return {"ok": True, "leveled_up": leveled,
            "new_sport_level": new_sport_lvl,
            "new_sport_xp": new_sport_xp}

# ── Duplikasi ─────────────────────────────────────────
def duplicate_habit(user_id, habit_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    h = conn.execute("SELECT * FROM habits WHERE id=? AND user_id=?", (habit_id, user_id)).fetchone()
    if not h:
        conn.close()
        return {"ok": False}
    new_name = h['name']  # tanpa tambahan (copy)
    conn.execute("""
        INSERT INTO habits(user_id,name,icon,difficulty,xp_reward,gold_reward,positive,negative,notes)
        VALUES(?,?,?,?,?,?,?,?,?)
    """, (user_id, new_name, h['icon'], h['difficulty'], h['xp_reward'], h['gold_reward'],
          h['positive'], h['negative'], h['notes']))
    conn.commit()
    conn.close()
    return {"ok": True}

def duplicate_daily(user_id, daily_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    d = conn.execute("SELECT * FROM dailies WHERE id=? AND user_id=?", (daily_id, user_id)).fetchone()
    if not d:
        conn.close()
        return {"ok": False}
    new_name = d['name']  # tanpa tambahan (copy)
    conn.execute("""
        INSERT INTO dailies(user_id,name,icon,difficulty,xp_reward,gold_reward,notes)
        VALUES(?,?,?,?,?,?,?)
    """, (user_id, new_name, d['icon'], d['difficulty'], d['xp_reward'], d['gold_reward'], d['notes']))
    conn.commit()
    conn.close()
    return {"ok": True}

def duplicate_todo(user_id, todo_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    t = conn.execute("SELECT * FROM todos WHERE id=? AND user_id=?", (todo_id, user_id)).fetchone()
    if not t:
        conn.close()
        return {"ok": False}
    new_name = t['name']  # tanpa tambahan (copy)
    conn.execute("""
        INSERT INTO todos(user_id,name,icon,priority,xp_reward,gold_reward,due_date,notes)
        VALUES(?,?,?,?,?,?,?,?)
    """, (user_id, new_name, t['icon'], t['priority'], t['xp_reward'], t['gold_reward'], t['due_date'], t['notes']))
    conn.commit()
    conn.close()
    return {"ok": True}

def duplicate_sport_activity(user_id, activity_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    a = conn.execute("SELECT * FROM sport_activities WHERE id=? AND user_id=?", (activity_id, user_id)).fetchone()
    if not a:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_sport_activity_not_found")}
    new_name = a['name']  # tanpa tambahan (copy)
    conn.execute("""
        INSERT INTO sport_activities(
            user_id, name, sport_type, icon, difficulty,
            xp_reward, gold_reward, sport_points_reward, notes
        ) VALUES(?,?,?,?,?,?,?,?,?)
    """, (user_id, new_name, a['sport_type'], a['icon'], a['difficulty'],
          a['xp_reward'], a['gold_reward'], a['sport_points_reward'], a['notes']))
    conn.commit()
    conn.close()
    return {"ok": True}

# ── Shop data ─────────────────────────────────────────────────────────────────

SHOP_ITEMS = {
    # ── Weapons ──
    "wooden_sword":    {"name": "Wooden Sword",    "icon": "🗡️", "cost": 50,
                        "type": "weapon",    "desc": "Starter blade",
                        "buff": {"boss_dmg": 2},
                        "buff_desc": "+2 Boss Damage"},
    "enchanted_bow":   {"name": "Enchanted Bow",   "icon": "🏹", "cost": 180,
                        "type": "weapon",    "desc": "Ranged power",
                        "buff": {"boss_dmg": 6},
                        "buff_desc": "+6 Boss Damage"},
    "trident":         {"name": "Trident",         "icon": "🔱", "cost": 350,
                        "type": "weapon",    "desc": "Legendary weapon",
                        "buff": {"boss_dmg": 10},
                        "buff_desc": "+10 Boss Damage"},
    "iron_sword":      {"name": "Iron Sword",      "icon": "🗡️", "cost": 100,
                        "type": "weapon",    "desc": "Sharp blade",
                        "buff": {"crit_chance": 8},
                        "buff_desc": "+8% Critical Chance"},
    "diamond_sword":   {"name": "Diamond Sword",   "icon": "💎", "cost": 320,
                        "type": "weapon",    "desc": "Powerful strike",
                        "buff": {"crit_chance": 15},
                        "buff_desc": "+15% Critical Chance"},
    "netherite_sword": {"name": "Netherite Sword", "icon": "⚔️", "cost": 650,
                        "type": "weapon",    "desc": "Legendary blade",
                        "buff": {"crit_chance": 20, "boss_dmg": 5},
                        "buff_desc": "+20% Critical Chance, +5 Boss DMG"},

    # ── Armor ──
    "shield":          {"name": "Shield",          "icon": "🛡️", "cost": 120,
                        "type": "armor",     "desc": "Reduces HP damage",
                        "buff": {"hp_reduc": 8},
                        "buff_desc": "-8 per HP hit taken"},
    "golden_boots":    {"name": "Golden Boots",    "icon": "👢", "cost": 140,
                        "type": "armor",     "desc": "Swift & rich",
                        "buff": {"gold_pct": 10},
                        "buff_desc": "+10% Gold earned"},
    "diamond_armor":   {"name": "Diamond Armor",   "icon": "💎", "cost": 300,
                        "type": "armor",     "desc": "Max protection",
                        "buff": {"hp_reduc": 15},
                        "buff_desc": "-15 per HP hit taken"},
    "elytra":          {"name": "Elytra Wings",    "icon": "🪽", "cost": 500,
                        "type": "armor",     "desc": "Glide & grow",
                        "buff": {"xp_pct": 10},
                        "buff_desc": "+10% XP all sources"},
    "tower_shield":    {"name": "Tower Shield",    "icon": "🛡️", "cost": 200,
                        "type": "armor",    "desc": "Great defense",
                        "buff": {"block_strength": 15},
                        "buff_desc": "+15 Block Strength"},
    "guardian_chestplate": {"name": "Guardian Chestplate", "icon": "🛡️", "cost": 400,
                        "type": "armor",    "desc": "Guardian armor",
                        "buff": {"block_chance": 10, "block_strength": 15},
                        "buff_desc": "+10% Block Chance, +15 Block Strength"},
    "diamond_chestplate": {"name": "Diamond Chestplate", "icon": "💎", "cost": 600,
                        "type": "armor",    "desc": "Diamond protection",
                        "buff": {"block_chance": 15, "block_strength": 20},
                        "buff_desc": "+15% Block Chance, +20 Block Strength"},

    # ── Tools ──
    "iron_pickaxe":    {"name": "Iron Pickaxe",    "icon": "⛏️", "cost": 100,
                        "type": "tool",      "desc": "Mine habits faster",
                        "buff": {"xp_pct": 8},
                        "buff_desc": "+8% XP from habits"},
    "compass":         {"name": "Compass",         "icon": "🧭", "cost": 80,
                        "type": "tool",      "desc": "Navigate to gold",
                        "buff": {"gold_pct": 6},
                        "buff_desc": "+6% Gold earned"},
    "spyglass":        {"name": "Spyglass",        "icon": "🔭", "cost": 90,
                        "type": "tool",      "desc": "Scout ahead",
                        "buff": {},
                        "buff_desc": "Reveal boss stats"},

    # ── Special ──
    "ice_block":       {"name": "Ice Block",       "icon": "🧊", "cost": 25,
                        "type": "consumable", "desc": "Freeze Daily streak (1 slot)",
                        "buff": {},
                        "buff_desc": "Use: +1 Freeze Slot for a Daily"},
    "blaze_rod":       {"name": "Blaze Rod",       "icon": "🔥", "cost": 160,
                        "type": "special",   "desc": "Nether fire",
                        "buff": {"boss_dmg": 4},
                        "buff_desc": "+4 Boss Damage"},
    "totem":           {"name": "Totem of Life",   "icon": "🗿", "cost": 400,
                        "type": "legendary", "desc": "Auto-revive from death",
                        "buff": {"revive": True},
                        "buff_desc": "Auto-revive once at 30% HP"},

    # ── Consumables (HP & MP) ──
    "golden_apple":    {"name": "Golden Apple",     "icon": "🍎", "cost": 50,
                        "type": "consumable", "desc": "Restore 25 HP",
                        "buff": {},
                        "buff_desc": "Use: +25 HP"},
    "enchanted_apple": {"name": "Enchanted Apple",  "icon": "🍏", "cost": 200,
                        "type": "consumable", "desc": "Restore 100 HP",
                        "buff": {},
                        "buff_desc": "Use: +100 HP"},
    "health_potion":   {"name": "Health Potion",    "icon": "❤️‍🩹", "cost": 100,
                        "type": "consumable", "desc": "Restore 50 HP",
                        "buff": {},
                        "buff_desc": "Use: +50 HP"},
    "greater_health_potion": {"name": "Greater Health Potion", "icon": "❤️", "cost": 150,
                        "type": "consumable", "desc": "Restore 75 HP",
                        "buff": {},
                        "buff_desc": "Use: +75 HP"},
    "mana_potion":     {"name": "Mana Potion",      "icon": "💙", "cost": 80,
                        "type": "consumable", "desc": "Restore 15 MP",
                        "buff": {},
                        "buff_desc": "Use: +15 MP"},
    "greater_mana_potion": {"name": "Greater Mana Potion", "icon": "💎", "cost": 200,
                        "type": "consumable", "desc": "Restore 35 MP",
                        "buff": {},
                        "buff_desc": "Use: +35 MP"},
    "super_mana_potion": {"name": "Super Mana Potion", "icon": "🔮", "cost": 500,
                        "type": "consumable", "desc": "Restore 80 MP",
                        "buff": {},
                        "buff_desc": "Use: +80 MP"},
    "elixir":          {"name": "Elixir of Life",   "icon": "🧪", "cost": 1000,
                        "type": "consumable", "desc": "Restore 80 HP & 40 MP",
                        "buff": {},
                        "buff_desc": "Use: +80 HP & +40 MP"},
    "ender_pearl":     {"name": "Ender Pearl",      "icon": "🔮", "cost": 300,
                        "type": "consumable", "desc": "Permanently increase Max MP",
                        "buff": {},
                        "buff_desc": "Use: +30 Max MP (permanent)"},

    # ── Legendary ──
    "nether_star":     {"name": "Nether Star",     "icon": "⭐", "cost": 700,
                        "type": "legendary",
                        "desc": "Power of the Nether",
                        "buff": {"xp_pct": 10, "gold_pct": 10, "boss_dmg": 5},
                        "buff_desc": "+10% XP, +10% Gold, +5 Boss DMG"},
    "beacon":          {"name": "Beacon",          "icon": "🏮", "cost": 1200,
                        "type": "legendary", "desc": "Strongest relic",
                        "buff": {"xp_pct": 12, "gold_pct": 12,
                                 "boss_dmg": 8, "hp_reduc": 5},
                        "buff_desc": "+12% XP, +12% Gold, +8 DMG, -5 HP taken"},

    # ── Craft-only (hanya bisa didapat lewat Crafting, tidak dijual di Shop) ──
    "bedrock_sword":   {"name": "Bedrock Sword",   "icon": "🪨", "cost": 2500,
                        "type": "weapon",    "desc": "Crafted from two legends",
                        "buff": {"crit_chance": 25, "boss_dmg": 18},
                        "buff_desc": "+25% Critical Chance, +18 Boss DMG",
                        "craft_only": True},
    "phantom_wings":   {"name": "Phantom Wings",   "icon": "🕊️", "cost": 1800,
                        "type": "armor",     "desc": "Crafted wind & riches",
                        "buff": {"xp_pct": 18, "gold_pct": 10},
                        "buff_desc": "+18% XP, +10% Gold earned",
                        "craft_only": True},
    "aegis_of_void":   {"name": "Aegis of the Void", "icon": "🛡️", "cost": 2400,
                        "type": "armor",     "desc": "Crafted ultimate defense",
                        "buff": {"hp_reduc": 25, "block_chance": 20, "block_strength": 30},
                        "buff_desc": "-25 HP taken, +20% Block, +30 Block STR",
                        "craft_only": True},

    # ── Seasonal items (hanya muncul saat event berlangsung) ──
    "ketupat_feast":   {"name": "Ketupat Feast",   "icon": "🍙", "cost": 75,
                        "type": "consumable", "desc": "Lebaran special meal",
                        "buff": {},
                        "buff_desc": "Use: +75 HP", "seasonal": "ramadan"},
    "candy_bag":       {"name": "Candy Bag",       "icon": "🍬", "cost": 40,
                        "type": "consumable", "desc": "Trick or treat!",
                        "buff": {},
                        "buff_desc": "Use: +30 HP", "seasonal": "halloween"},
    "snowball_fight":  {"name": "Snowball Fight",  "icon": "❄️", "cost": 50,
                        "type": "consumable", "desc": "Christmas fun",
                        "buff": {},
                        "buff_desc": "Use: +50 HP", "seasonal": "christmas"},

    # ── 🆕 v1.3.0 — Katalog baru (10 item toko + 3 item craft-only) ──
    # Weapons
    "storm_blade":     {"name": "Storm Blade",     "icon": "🌩️", "cost": 950,
                        "type": "weapon",    "desc": "Slash with thunder speed",
                        "buff": {"boss_dmg": 10, "crit_chance": 8},
                        "buff_desc": "+10 Boss DMG, +8% Crit Chance"},
    "obsidian_dagger": {"name": "Obsidian Dagger", "icon": "🔪", "cost": 420,
                        "type": "weapon",    "desc": "Silent but deadly",
                        "buff": {"crit_chance": 18},
                        "buff_desc": "+18% Critical Chance"},
    # Armor
    "turtle_shell":    {"name": "Turtle Shell Helm","icon": "🐢", "cost": 460,
                        "type": "armor",     "desc": "Slow but unbreakable",
                        "buff": {"hp_reduc": 10, "block_chance": 6},
                        "buff_desc": "-10 HP taken, +6% Block Chance"},
    "wind_cloak":      {"name": "Wind Cloak",      "icon": "🌬️", "cost": 380,
                        "type": "armor",     "desc": "Light as air, rich as wind",
                        "buff": {"xp_pct": 6, "gold_pct": 6},
                        "buff_desc": "+6% XP, +6% Gold"},
    # Tools
    "lucky_charm":     {"name": "Lucky Charm",     "icon": "🍀", "cost": 260,
                        "type": "tool",      "desc": "Fortune favors the grind",
                        "buff": {"xp_pct": 4, "gold_pct": 6},
                        "buff_desc": "+4% XP, +6% Gold"},
    "scholar_tome":    {"name": "Scholar's Tome",  "icon": "📖", "cost": 320,
                        "type": "tool",      "desc": "Knowledge is XP",
                        "buff": {"xp_pct": 9},
                        "buff_desc": "+9% XP earned"},
    # Consumables (terhubung ke hp_map/mp_map di use_consumable)
    "honey_bottle":    {"name": "Honey Bottle",    "icon": "🍯", "cost": 60,
                        "type": "consumable", "desc": "Sweet recovery",
                        "buff": {},
                        "buff_desc": "Use: +35 HP"},
    "sturdy_stew":     {"name": "Sturdy Stew",     "icon": "🍲", "cost": 90,
                        "type": "consumable", "desc": "Warm meal for body & mind",
                        "buff": {},
                        "buff_desc": "Use: +20 HP & +10 MP"},
    "dragon_breath":   {"name": "Dragon's Breath", "icon": "🐉", "cost": 240,
                        "type": "consumable", "desc": "Bottled arcane fire",
                        "buff": {},
                        "buff_desc": "Use: +60 MP"},
    # Legendary
    "dragon_egg":      {"name": "Dragon Egg",      "icon": "🐲", "cost": 1600,
                        "type": "legendary", "desc": "Slumbering ancient power",
                        "buff": {"xp_pct": 12, "gold_pct": 12, "boss_dmg": 10},
                        "buff_desc": "+12% XP, +12% Gold, +10 Boss DMG"},
    # Craft-only (hanya via Crafting)
    "inferno_blade":   {"name": "Inferno Blade",   "icon": "🔥", "cost": 2800,
                        "type": "weapon",    "desc": "Forged in nether fire",
                        "buff": {"boss_dmg": 22, "crit_chance": 12},
                        "buff_desc": "+22 Boss DMG, +12% Crit Chance",
                        "craft_only": True},
    "healers_blessing": {"name": "Healer's Blessing", "icon": "💚", "cost": 2200,
                        "type": "armor",     "desc": "Blessed restorative ward",
                        "buff": {"hp_reduc": 8, "mp_bonus": 25},
                        "buff_desc": "-8 HP taken, +25 Max MP",
                        "craft_only": True},
    "gilded_compass":  {"name": "Gilded Compass",  "icon": "🧭", "cost": 2000,
                        "type": "tool",      "desc": "Points straight to treasure",
                        "buff": {"gold_pct": 15, "xp_pct": 10},
                        "buff_desc": "+15% Gold, +10% XP",
                        "craft_only": True},
    # ── 🆕 v1.4.0 — New Shop Items (balanced grindy, tidak OP) ──
    "bronze_sword":     {"name": "Bronze Sword",    "icon": "🗡️", "cost": 70,
                        "type": "weapon",    "desc": "Balanced starter blade",
                        "buff": {"boss_dmg": 3},
                        "buff_desc": "+3 Boss Damage"},
    "steel_helm":       {"name": "Steel Helm",      "icon": "⛑️", "cost": 250,
                        "type": "armor",     "desc": "Sturdy head protection",
                        "buff": {"hp_reduc": 7, "block_chance": 4},
                        "buff_desc": "-7 HP taken, +4% Block Chance"},
    "travelers_boots": {"name": "Traveler's Boots","icon": "🥾", "cost": 190,
                        "type": "armor",     "desc": "For long journeys",
                        "buff": {"gold_pct": 7, "xp_pct": 3},
                        "buff_desc": "+7% Gold, +3% XP"},
    "arcane_ring":      {"name": "Arcane Ring",     "icon": "💍", "cost": 450,
                        "type": "tool",      "desc": "Whispers of mana",
                        "buff": {"xp_pct": 7, "gold_pct": 4},
                        "buff_desc": "+7% XP, +4% Gold"},
    "berry_pie":        {"name": "Berry Pie",       "icon": "🥧", "cost": 45,
                        "type": "consumable", "desc": "Sweet healing",
                        "buff": {},
                        "buff_desc": "Use: +30 HP"},
    "mana_cookie":      {"name": "Mana Cookie",     "icon": "🍪", "cost": 110,
                        "type": "consumable", "desc": "Crumbly mana boost",
                        "buff": {},
                        "buff_desc": "Use: +20 MP"},
    "frost_guard":      {"name": "Frost Guard",     "icon": "❄️", "cost": 2600,
                        "type": "armor",     "desc": "Crafted ice shield",
                        "buff": {"hp_reduc": 12, "block_chance": 8, "block_strength": 10},
                        "buff_desc": "-12 HP taken, +8% Block, +10 Block STR",
                        "craft_only": True},
    "scholar_crown":    {"name": "Scholar Crown",   "icon": "👑", "cost": 2400,
                        "type": "tool",      "desc": "Crown of wisdom",
                        "buff": {"xp_pct": 14, "gold_pct": 6},
                        "buff_desc": "+14% XP, +6% Gold",
                        "craft_only": True},
    "void_core":        {"name": "Void Core",       "icon": "🌑", "cost": 3000,
                        "type": "legendary", "desc": "Heart of the void",
                        "buff": {"xp_pct": 10, "gold_pct": 10, "boss_dmg": 7, "hp_reduc": 5},
                        "buff_desc": "+10% XP, +10% Gold, +7 DMG, -5 HP",
                        "craft_only": True},
    "ember_charm":      {"name": "Ember Charm",     "icon": "🧿", "cost": 2200,
                        "type": "tool",      "desc": "Warm lucky ember",
                        "buff": {"gold_pct": 9, "xp_pct": 6, "boss_dmg": 2},
                        "buff_desc": "+9% Gold, +6% XP, +2 Boss DMG",
                        "craft_only": True},
}

PETS_DATA = {
    "wolf":     {"name": "Wolf", "icon": "🐺", "cost": 150,
                 "bonus": "+6% XP",                     
                 "base_buff": {"xp_pct": 6}},
    "cat":      {"name": "Cat", "icon": "🐱", "cost": 120,
                 "bonus": "-6 HP loss",                 
                 "base_buff": {"hp_reduc": 6}},
    "parrot":   {"name": "Parrot", "icon": "🦜", "cost": 140,
                 "bonus": "+4% Gold",                   
                 "base_buff": {"gold_pct": 4}},
    "panda":    {"name": "Panda", "icon": "🐼", "cost": 240,
                 "bonus": "+10% XP",                     
                 "base_buff": {"xp_pct": 10}},
    "fox":      {"name": "Fox", "icon": "🦊", "cost": 180,
                 "bonus": "+8% Gold",                   
                 "base_buff": {"gold_pct": 8}},
    "bee":      {"name": "Bee", "icon": "🐝", "cost": 110,
                 "bonus": "-4 HP loss",                 
                 "base_buff": {"hp_reduc": 4}},
    "dragon":   {"name": "Dragon", "icon": "🐉", "cost": 600,
                 "bonus": "+15% XP, +4 boss dmg",        
                 "base_buff": {"xp_pct": 15, "boss_dmg": 4}},
    "turtle":   {"name": "Turtle", "icon": "🐢", "cost": 130,
                 "bonus": "-8 HP loss",                 
                 "base_buff": {"hp_reduc": 8}},
    "axolotl":  {"name": "Axolotl", "icon": "🦎", "cost": 200,
                 "bonus": "-2 HP loss",                 
                 "base_buff": {"hp_reduc": 2}},
    "enderman": {"name": "Enderman", "icon": "👾", "cost": 400,
                 "bonus": "+12% XP, +4% Gold, +2 boss dmg", 
                 "base_buff": {"xp_pct": 12, "gold_pct": 4, "boss_dmg": 2}},
    "phoenix":     {"name": "Phoenix",   "icon": "🐦‍🔥", "cost": 800,
                    "bonus": "+10% XP, +5% Gold, +2 boss dmg",
                    "base_buff": {"xp_pct": 10, "gold_pct": 5, "boss_dmg": 2}},
    "unicorn":     {"name": "Unicorn",   "icon": "🦄", "cost": 650,
                    "bonus": "+8% XP, +8% Gold",
                    "base_buff": {"xp_pct": 8, "gold_pct": 8}},
    "griffin":     {"name": "Griffin",   "icon": "🦅", "cost": 700,
                    "bonus": "+10% XP, -5 HP loss",
                    "base_buff": {"xp_pct": 10, "hp_reduc": 5}},
    "mermaid":     {"name": "Mermaid",   "icon": "🧜‍♀️", "cost": 500,
                    "bonus": "+6% Gold, -8 HP loss",
                    "base_buff": {"gold_pct": 6, "hp_reduc": 8}},
    "slime":       {"name": "Slime",     "icon": "🟢", "cost": 80,
                    "bonus": "+2% Gold, -2 HP loss",
                    "base_buff": {"gold_pct": 2, "hp_reduc": 2}},
    "ghast":       {"name": "Ghost",     "icon": "👻", "cost": 150,
                    "bonus": "+4% XP, -4 HP loss",
                    "base_buff": {"xp_pct": 4, "hp_reduc": 4}},
    "skeleton":    {"name": "Skeleton",  "icon": "💀", "cost": 120,
                    "bonus": "+3% Gold, +1 boss dmg",
                    "base_buff": {"gold_pct": 3, "boss_dmg": 1}},
    "zombie":      {"name": "Zombie",    "icon": "🧟", "cost": 100,
                    "bonus": "+2% XP, +1 boss dmg",
                    "base_buff": {"xp_pct": 2, "boss_dmg": 1}},
    "creeper":     {"name": "Creeper",   "icon": "💥", "cost": 90,
                 "bonus": "+2% Gold, -2 HP loss",
                 "base_buff": {"gold_pct": 2, "hp_reduc": 2}},
    # ── 🆕 v1.4.0 — Pets baru (balanced, grindy) ──
    "owl":       {"name": "Owl",       "icon": "🦉", "cost": 220,
                 "bonus": "+5% XP, +2% Gold",
                 "base_buff": {"xp_pct": 5, "gold_pct": 2}},
    "hamster":   {"name": "Hamster",   "icon": "🐹", "cost": 95,
                 "bonus": "+3% Gold, -2 HP loss",
                 "base_buff": {"gold_pct": 3, "hp_reduc": 2}},
    "dolphin":   {"name": "Dolphin",   "icon": "🐬", "cost": 300,
                 "bonus": "+7% Gold, -4 HP loss",
                 "base_buff": {"gold_pct": 7, "hp_reduc": 4}},
    "bat":       {"name": "Bat",       "icon": "🦇", "cost": 160,
                 "bonus": "-5 HP loss, +2% XP",
                 "base_buff": {"hp_reduc": 5, "xp_pct": 2}},
    "lion":      {"name": "Lion",      "icon": "🦁", "cost": 550,
                 "bonus": "+8% XP, +3 boss dmg",
                 "base_buff": {"xp_pct": 8, "boss_dmg": 3}},
    "capybara":  {"name": "Capybara",  "icon": "🦫", "cost": 180,
                 "bonus": "+4% XP, -4 HP loss",
                 "base_buff": {"xp_pct": 4, "hp_reduc": 4}},
}

BOSSES = {
    # Beginner
    "zombie":         {"name": "Zombie",          "icon": "🧟",
                       "tier": "beginner", "hp": 200,  "atk": 5,
                       "xp": 80,   "gold": 20,  "min_level": 1},
    "skeleton":       {"name": "Skeleton Archer", "icon": "💀",
                       "tier": "beginner", "hp": 300,  "atk": 8,
                       "xp": 120,  "gold": 30,  "min_level": 1},
    "creeper":        {"name": "Creeper",         "icon": "💥",
                       "tier": "beginner", "hp": 250,  "atk": 12,
                       "xp": 100,  "gold": 25,  "min_level": 2},
    # Normal
    "zombie_king":    {"name": "Zombie King",     "icon": "👑",
                       "tier": "normal",   "hp": 600,  "atk": 15,
                       "xp": 250,  "gold": 60,  "min_level": 3},
    "skeleton_lord":  {"name": "Skeleton Lord",   "icon": "☠️",
                       "tier": "normal",   "hp": 800,  "atk": 20,
                       "xp": 350,  "gold": 80,  "min_level": 5},
    "blaze_lord":     {"name": "Blaze Lord",      "icon": "🔥",
                       "tier": "normal",   "hp": 700,  "atk": 18,
                       "xp": 300,  "gold": 70,  "min_level": 4},
    # Hard
    "iron_golem":     {"name": "Iron Golem Boss", "icon": "⚙️",
                       "tier": "hard",    "hp": 1500, "atk": 25,
                       "xp": 600,  "gold": 150, "min_level": 8},
    "creeper_god":    {"name": "Creeper God",     "icon": "💣",
                       "tier": "hard",    "hp": 1200, "atk": 30,
                       "xp": 700,  "gold": 180, "min_level": 10},
    "spider_queen":   {"name": "Spider Queen",    "icon": "🕷️",
                       "tier": "hard",    "hp": 1000, "atk": 22,
                       "xp": 500,  "gold": 120, "min_level": 7},
    # Elite
    "wither":         {"name": "The Wither",      "icon": "💀",
                       "tier": "elite",   "hp": 2500, "atk": 40,
                       "xp": 1200, "gold": 300, "min_level": 15},
    "ender_dragon":   {"name": "Ender Dragon",    "icon": "🐲",
                       "tier": "elite",   "hp": 3000, "atk": 50,
                       "xp": 1500, "gold": 400, "min_level": 20},
    # Legendary
    "elder_guardian": {"name": "Elder Guardian",  "icon": "👁️",
                       "tier": "legendary", "hp": 5000, "atk": 70,
                       "xp": 3000, "gold": 800, "min_level": 30},
    "herobrine":      {"name": "Herobrine",       "icon": "👻",
                       "tier": "legendary", "hp": 8000, "atk": 100,
                       "xp": 5000, "gold": 1500,"min_level": 50},
    # ── Seasonal bosses (hanya saat event aktif) ──
    "ketupat_golem":  {"name": "Ketupat Golem",   "icon": "🌙",
                       "tier": "seasonal", "hp": 900,  "atk": 16,
                       "xp": 400,  "gold": 120, "min_level": 3,
                       "seasonal_event": "ramadan"},
    "pumpkin_king":   {"name": "Pumpkin King",    "icon": "🎃",
                       "tier": "seasonal", "hp": 1100, "atk": 24,
                       "xp": 600,  "gold": 160, "min_level": 6,
                       "seasonal_event": "halloween"},
    "krampus":        {"name": "Krampus",          "icon": "🎄",
                       "tier": "seasonal", "hp": 1800, "atk": 32,
                       "xp": 900,  "gold": 250, "min_level": 10,
                       "seasonal_event": "christmas"},
}

BOSS_TIER_ORDER = ["beginner", "normal", "hard", "elite", "legendary", "seasonal"]
BOSS_TIER_COLOR = {
    "beginner":  "#7bbf3e",
    "normal":    "#f0a800",
    "hard":      "#e05050",
    "elite":     "#a97fff",
    "legendary": "#ff6b00",
    "seasonal":  "#2dd4bf",
}

SECURITY_QUESTIONS = [
    "Apa nama hewan peliharaan pertama Anda?",
    "Apa nama sekolah dasar Anda?",
    "Apa nama ibu kota tempat Anda lahir?",
    "Siapa tokoh idola Anda?",
    "Apa makanan favorit Anda?",
    "Apa merek handphone pertama Anda?",
    "Apa nama jalan tempat Anda tinggal saat kecil?",
]

# ── Inventory / Shop ──────────────────────────────────────────────────────────

def get_inventory(user_id):
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM inventory WHERE user_id=?", (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def buy_item(user_id, item_id):
    item = SHOP_ITEMS.get(item_id)
    if not item:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_item_not_found")}
    u = get_user(user_id)
    if u["gold"] < item["cost"]:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_gold_insufficient", cost=item['cost'])}
    conn = get_conn()
    ex = conn.execute(
        "SELECT * FROM inventory WHERE user_id=? AND item_id=?",
        (user_id, item_id)).fetchone()
    if ex and item["type"] not in ("consumable",):
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_item_already_owned")}
    if ex:
        conn.execute("UPDATE inventory SET quantity=quantity+1 WHERE id=?",
                     (ex["id"],))
    else:
        conn.execute(
            "INSERT INTO inventory(user_id,item_id,item_type) VALUES(?,?,?)",
            (user_id, item_id, item["type"]))
    conn.execute("UPDATE users SET gold=gold-? WHERE id=?",
                 (item["cost"], user_id))
    conn.commit()
    conn.close()
    update_total_spent(user_id, item["cost"])
    recalculate_all_buffs(user_id)
    log_activity(user_id, "buy", tr_db(user_id=user_id, key="log_buy", item=item['name'], gold=item['cost']), 0, -item["cost"])
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_item_bought", icon=item['icon'], name=item['name'], buff=item['buff_desc'])}


def use_item(user_id, item_id):
    item = SHOP_ITEMS.get(item_id)
    if not item or item["type"] != "consumable":
        return {"ok": False}
    conn = get_conn()
    inv = conn.execute(
        "SELECT * FROM inventory WHERE user_id=? AND item_id=? AND quantity>0",
        (user_id, item_id)
    ).fetchone()
    if not inv:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_item_not_found")}
    
    hp_map = {
        "golden_apple": 25,
        "enchanted_apple": 100,
        "health_potion": 50,
        "greater_health_potion": 75,
        "honey_bottle": 35,
        "sturdy_stew": 20,
        "berry_pie": 30,
        # Seasonal consumables
        "ketupat_feast": 75,
        "candy_bag": 30,
        "snowball_fight": 50
    }
    mp_map = {
        "mana_potion": 15,
        "greater_mana_potion": 35,
        "super_mana_potion": 80,
        "elixir": 40,
        "sturdy_stew": 10,
        "dragon_breath": 60,
        "mana_cookie": 20
    }
    max_mp_map = {
        "ender_pearl": 30
    }
    
    restore_hp = hp_map.get(item_id, 0)
    restore_mp = mp_map.get(item_id, 0)
    restore_max_mp = max_mp_map.get(item_id, 0)
    
    if restore_hp == 0 and restore_mp == 0 and restore_max_mp == 0:
        conn.close()
        return {"ok": False, "msg": "Item tidak bisa digunakan."}
    
    msg_parts = []
    
    if restore_hp > 0:
        conn.execute("UPDATE users SET hp = MIN(max_hp, hp + ?) WHERE id=?", (restore_hp, user_id))
        msg_parts.append(f"+{restore_hp} HP")
    
    if restore_mp > 0:
        conn.execute("UPDATE users SET mp = MIN(max_mp, mp + ?) WHERE id=?", (restore_mp, user_id))
        msg_parts.append(f"+{restore_mp} MP")
    
    if restore_max_mp > 0:
        # Tambah Max MP permanen + isi MP sebesar jumlah yang sama agar langsung terasa
        conn.execute(
            "UPDATE users SET max_mp = max_mp + ?, mp = mp + ? WHERE id=?",
            (restore_max_mp, restore_max_mp, user_id)
        )
        msg_parts.append(f"+{restore_max_mp} Max MP (permanen)")
    
    conn.execute("UPDATE inventory SET quantity = quantity - 1 WHERE id=?", (inv["id"],))
    conn.commit()
    conn.close()
    
    msg = " & ".join(msg_parts)
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_item_used", icon=item['icon'], restore=msg)}

@retry_on_lock
def sell_item(user_id, item_id, quantity=1):
    """
    Menjual item kembali ke shop dengan harga 10% dari harga asli.
    - item_id: ID item di inventory (bukan shop item_id)
    - quantity: jumlah yang dijual (default 1)
    """
    conn = get_conn()
    try:
        # Ambil data item dari inventory
        inv = conn.execute(
            "SELECT * FROM inventory WHERE id = ? AND user_id = ?",
            (item_id, user_id)
        ).fetchone()
        if not inv:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_item_not_found")}
        
        # Ambil data shop item
        shop_item = SHOP_ITEMS.get(inv["item_id"])
        if not shop_item:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_item_not_found")}
        
        # Hitung harga jual (10% dari harga beli)
        sell_price = shop_item["cost"] * 0.1
        sell_price = max(1, int(sell_price))  # minimal 1 Gold
        
        # Cek quantity
        if inv["quantity"] < quantity:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_item_insufficient_quantity")}
        
        # Kurangi quantity atau hapus item
        new_quantity = inv["quantity"] - quantity
        if new_quantity <= 0:
            conn.execute("DELETE FROM inventory WHERE id = ?", (item_id,))
        else:
            conn.execute("UPDATE inventory SET quantity = quantity - ? WHERE id = ?", (quantity, item_id))
        
        # Tambahkan gold ke user
        total_gold = sell_price * quantity
        conn.execute("UPDATE users SET gold = gold + ? WHERE id = ?", (total_gold, user_id))
        conn.commit()
    
        try:
            recalculate_all_buffs(user_id)
        except Exception as e:
            log_crash(f"recalculate_all_buffs failed after selling item: {e}")

        # Catat aktivitas
        log_activity(
            user_id,
            "sell_item",
            tr_db(user_id=user_id, key="log_sell_item", name=shop_item['name'], qty=quantity, gold=total_gold),
            0,
            total_gold
        )
        
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_item_sold", name=shop_item['name'], qty=quantity, gold=total_gold)}
    except Exception as e:
        conn.rollback()
        return {"ok": False, "msg": str(e)}
    finally:
        conn.close()

# ── Pets ──────────────────────────────────────────────────────────────────────
def get_user_pets(user_id):
    conn = get_conn()
    rows = conn.execute("SELECT * FROM user_pets WHERE user_id=?", (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def adopt_pet(user_id, pet_id):
    pet = PETS_DATA.get(pet_id)
    if not pet:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_not_found")}
    u = get_user(user_id)
    if u["gold"] < pet["cost"]:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_gold_insufficient", cost=pet['cost'])}
    conn = get_conn()
    try:
        if conn.execute("SELECT 1 FROM user_pets WHERE user_id=? AND pet_id=?", (user_id, pet_id)).fetchone():
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_already_owned")}
        conn.execute("UPDATE users SET gold=gold-? WHERE id=?", (pet["cost"], user_id))
        conn.execute("INSERT INTO user_pets(user_id, pet_id, hunger, happiness) VALUES(?,?,100,50)", (user_id, pet_id))
        conn.commit()
    finally:
        conn.close()
    check_achievements(user_id, "pet_adopt", 1)
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_pet_adopted", icon=pet['icon'], name=pet['name'])}

def equip_pet(user_id, pet_id):
    """Equip pet (aktifkan). Maksimal 2 pet jika user level >=25, selain itu maksimal 1.
       Jika sudah mencapai batas, akan mengganti pet aktif yang paling lama (ID terkecil)."""
    conn = get_conn()
    try:
        # Ambil level user
        user = conn.execute("SELECT level FROM users WHERE id=?", (user_id,)).fetchone()
        if not user:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_user_not_found")}
        user_level = user["level"]
        max_pets = 2 if user_level >= 25 else 1
        
        # Cek apakah pet yang akan di-equip sudah aktif?
        pet = conn.execute(
            "SELECT is_active FROM user_pets WHERE user_id=? AND pet_id=?",
            (user_id, pet_id)
        ).fetchone()
        if not pet:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_not_found")}
        if pet["is_active"]:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_already_active")}
        
        # Hitung jumlah pet yang sudah aktif
        active_count = conn.execute(
            "SELECT COUNT(*) as cnt FROM user_pets WHERE user_id=? AND is_active=1",
            (user_id,)
        ).fetchone()["cnt"]
        
        if active_count >= max_pets:
            # Cari pet aktif yang paling lama (ID terkecil) untuk diganti
            oldest_active = conn.execute(
                "SELECT id FROM user_pets WHERE user_id=? AND is_active=1 ORDER BY id LIMIT 1",
                (user_id,)
            ).fetchone()
            if oldest_active:
                # Nonaktifkan pet lama
                conn.execute("UPDATE user_pets SET is_active=0 WHERE id=?", (oldest_active["id"],))
            else:
                # Seharusnya tidak terjadi
                return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_equip_fail")}
        
        # Aktifkan pet baru
        conn.execute("UPDATE user_pets SET is_active=1 WHERE user_id=? AND pet_id=?", (user_id, pet_id))
        conn.commit()
    finally:
        conn.close()
    
    recalculate_all_buffs(user_id)
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_pet_equipped", name=pet_id)}

def unequip_pet(user_id, pet_id):
    """Nonaktifkan pet (unequip). Hanya bisa dilakukan jika pet sedang aktif."""
    conn = get_conn()
    try:
        # Cek apakah pet ada dan aktif
        pet = conn.execute(
            "SELECT is_active FROM user_pets WHERE user_id=? AND pet_id=?",
            (user_id, pet_id)
        ).fetchone()
        if not pet:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_not_found")}
        if not pet["is_active"]:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_not_active")}
        
        # Nonaktifkan pet
        conn.execute(
            "UPDATE user_pets SET is_active=0 WHERE user_id=? AND pet_id=?",
            (user_id, pet_id)
        )
        conn.commit()
    finally:
        conn.close()
    
    recalculate_all_buffs(user_id)
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_pet_unequipped", name=pet_id)}

@retry_on_lock
def feed_pet(user_id, pet_id):
    pet = get_user_pet_by_id(user_id, pet_id)
    if not pet:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_not_found")}
    
    u = get_user(user_id)
    cost = 30
    if u["gold"] < cost:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_gold_insufficient", cost=cost)}
    
    # Ambil nama pet dengan aman
    pet_data = PETS_DATA.get(pet_id, {})
    pet_name = pet_data.get('name', f"Pet-{pet_id}")
    
    conn = get_conn()
    try:
        conn.execute("UPDATE user_pets SET hunger=100, last_fed=? WHERE id=?", (local_now().isoformat(), pet["id"]))
        conn.execute("UPDATE users SET gold=gold-? WHERE id=?", (cost, user_id))
        conn.commit()
    finally:
        conn.close()
    
    # Kembalikan data mentah, BUKAN string terjemahan
    return {
        "ok": True,
        "name": pet_name,
        "cost": cost
    }

@retry_on_lock
def train_pet(user_id, pet_id):
    MAX_PET_LEVEL = 20
    pet = get_user_pet_by_id(user_id, pet_id)
    if not pet:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_not_found")}
    
    level = pet["level"]  # ← WAJIB, ambil level dari pet
    
    if level >= MAX_PET_LEVEL:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_max_level", max=MAX_PET_LEVEL)}
    
    if pet["hunger"] < 20:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pet_hungry")}
    
    u = get_user(user_id)
    pet_data = PETS_DATA.get(pet_id, {})
    pet_name = pet_data.get('name', f"Pet-{pet_id}")
    
    cost = 25 + (level - 1) * 5
    if u["gold"] < cost:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_gold_insufficient", cost=cost)}
    
    needed = level * 100
    # FIX 8: Random EXP - 30% chance 25% missing XP, 70% chance <25% missing XP
    current_exp = pet["exp"] or 0
    missing = max(1, needed - current_exp)
    import random as _rnd
    if _rnd.random() < 0.30:
        exp_gain = int(missing * 0.25)
    else:
        low = int(missing * 0.05)
        high = int(missing * 0.24)
        low = max(1, low)
        high = max(low, high)
        exp_gain = _rnd.randint(low, high)
    exp_gain = max(5, exp_gain)
    exp_gain = min(exp_gain, missing)
    
    conn = get_conn()
    leveled = False
    try:
        new_hunger = max(0, pet["hunger"] - 20)
        conn.execute("UPDATE user_pets SET hunger=? WHERE id=?", (new_hunger, pet["id"]))
        leveled = add_pet_exp(conn, pet["id"], exp_gain)
        conn.execute("UPDATE users SET gold=gold-? WHERE id=?", (cost, user_id))
        conn.commit()
    finally:
        conn.close()
    
    new_level = level
    if leveled:
        updated_pet = get_user_pet_by_id(user_id, pet_id)
        new_level = updated_pet["level"] if updated_pet else level + 1
        add_notification(user_id, tr_db(user_id=user_id, key="db_pet_level_up", name=pet_name, level=new_level), "levelup")
        recalculate_all_buffs(user_id)
    
    return {
        "ok": True,
        "name": pet_name,
        "exp_gained": exp_gain,
        "cost": cost,
        "leveled_up": leveled,
        "new_level": new_level
    }

def add_pet_exp(conn, pet_row_id, amount):
    """Return True jika level naik, False jika tidak"""
    pet = conn.execute("SELECT user_id, exp, level FROM user_pets WHERE id=?", (pet_row_id,)).fetchone()
    if not pet:
        return False
    new_exp = pet["exp"] + amount
    new_level = pet["level"]
    needed = pet["level"] * 100
    leveled = False
    while new_exp >= needed:
        new_exp -= needed
        new_level += 1
        needed = new_level * 100
        leveled = True
    conn.execute("UPDATE user_pets SET exp=?, level=? WHERE id=?", (new_exp, new_level, pet_row_id))
    return leveled

def get_user_pet_by_id(user_id, pet_id):
    conn = get_conn()
    pet = conn.execute("SELECT * FROM user_pets WHERE user_id=? AND pet_id=?", (user_id, pet_id)).fetchone()
    conn.close()
    return pet

# ── Party / Boss ──────────────────────────────────────────────────────────────

def create_guild(leader_id, name, description=""):
    u = get_user(leader_id)
    if u.get("is_admin", 0):
        return {"ok": False, "msg": tr_db(user_id=leader_id, key="db_admin_cannot_guild")}
    conn = get_conn()
    gid = get_next_guild_id()
    conn.execute("INSERT INTO guilds(id, name, description, leader_id) VALUES(?,?,?,?)",
                 (gid, name, description, leader_id))
    conn.execute("INSERT INTO guild_members(guild_id, user_id) VALUES(?,?)", (gid, leader_id))
    conn.execute("UPDATE users SET guild_id=? WHERE id=?", (gid, leader_id))
    conn.commit()
    conn.close()
    check_achievements(leader_id, "join_guild", 1)
    return {"ok": True, "guild_id": gid, "msg": tr_db(user_id=leader_id, key="db_guild_created", name=name, id=gid)}


# ── Guild Request Join ───────────────────────────────────
def send_guild_request(user_id, guild_id):
    u = get_user(user_id)
    if u.get("is_admin", 0):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_admin_cannot_guild")}
    conn = get_conn()
    guild = conn.execute("SELECT id, name FROM guilds WHERE id=?", (guild_id,)).fetchone()
    if not guild:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_guild_not_found")}
    # Cek apakah sudah member
    existing = conn.execute("SELECT 1 FROM guild_members WHERE guild_id=? AND user_id=?", (guild_id, user_id)).fetchone()
    if existing:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_already_in_guild")}
    # Cek apakah sudah ada request pending
    pending = conn.execute("SELECT 1 FROM guild_requests WHERE guild_id=? AND user_id=? AND status='pending'", (guild_id, user_id)).fetchone()
    if pending:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_guild_request_pending")}
    conn.execute("INSERT INTO guild_requests(guild_id, user_id) VALUES(?,?)", (guild_id, user_id))
    conn.commit()
    # Ambil leader_id sebelum conn ditutup
    leader = conn.execute("SELECT leader_id FROM guilds WHERE id=?", (guild_id,)).fetchone()
    leader_id = leader["leader_id"] if leader else None
    conn.close()
    # Notifikasi ke leader (setelah conn ditutup agar tidak nested)
    if leader_id:
        add_notification(leader_id, tr_db(user_id=leader_id, key="db_guild_request_notif", name=get_user(user_id)['display_name'], guild=guild['name']), "info")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_guild_request_sent")}

def get_guild_requests(guild_id):
    conn = get_conn()
    rows = conn.execute("""
        SELECT gr.*, u.display_name, u.username
        FROM guild_requests gr
        JOIN users u ON gr.user_id = u.id
        WHERE gr.guild_id=? AND gr.status='pending'
    """, (guild_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def accept_guild_request(guild_id, leader_id, request_id):
    conn = get_conn()
    guild = conn.execute("SELECT leader_id, name FROM guilds WHERE id=?", (guild_id,)).fetchone()
    if not guild or guild["leader_id"] != leader_id:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=leader_id, key="db_guild_request_accept_only_leader")}
    
    # Cek jumlah member
    count = conn.execute("SELECT COUNT(*) FROM guild_members WHERE guild_id=?", (guild_id,)).fetchone()[0]
    if count >= 20:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=leader_id, key="db_guild_full")}
    
    req = conn.execute("SELECT * FROM guild_requests WHERE id=? AND guild_id=? AND status='pending'", (request_id, guild_id)).fetchone()
    if not req:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=leader_id, key="db_guild_request_not_found")}
    conn.execute("UPDATE guild_requests SET status='accepted' WHERE id=?", (request_id,))
    conn.execute("INSERT INTO guild_members(guild_id, user_id) VALUES(?,?)", (guild_id, req["user_id"]))
    conn.execute("UPDATE users SET guild_id=? WHERE id=?", (guild_id, req["user_id"]))
    conn.commit()
    conn.close()
    check_achievements(req["user_id"], "join_guild", 1)
    add_notification(req["user_id"], tr_db(user_id=req["user_id"], key="db_guild_request_accepted_notif", name=guild['name']), "success")
    return {"ok": True, "msg": tr_db(user_id=leader_id, key="db_guild_member_accepted")}

def reject_guild_request(guild_id, leader_id, request_id):
    conn = get_conn()
    guild = conn.execute("SELECT leader_id FROM guilds WHERE id=?", (guild_id,)).fetchone()
    if not guild or guild["leader_id"] != leader_id:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=leader_id, key="db_guild_request_reject_only_leader")}
    conn.execute("UPDATE guild_requests SET status='rejected' WHERE id=? AND guild_id=?", (request_id, guild_id))
    conn.commit()
    conn.close()
    return {"ok": True, "msg": tr_db(user_id=leader_id, key="db_guild_request_rejected")}

def get_guild(guild_id):
    conn = get_conn()
    g = conn.execute("SELECT * FROM guilds WHERE id=?",
                     (guild_id,)).fetchone()
    if not g:
        conn.close()
        return {}
    members = conn.execute(
        "SELECT u.id,u.display_name,u.level,u.avatar_class,"
        "u.avatar_emoji,u.hp,u.max_hp"
        " FROM users u JOIN guild_members gm ON u.id=gm.user_id"
        " WHERE gm.guild_id=?",
        (guild_id,)).fetchall()
    boss = conn.execute("SELECT * FROM boss_battles WHERE guild_id=? AND status='active'", (guild_id,)).fetchone()
    conn.close()
    return {
        "guild":   dict(g),
        "members": [dict(m) for m in members],
        "boss":    dict(boss) if boss else None,
    }

def get_next_guild_id():
    conn = get_conn()
    # Cari ID terkecil yang tidak terpakai
    used = conn.execute("SELECT id FROM guilds ORDER BY id").fetchall()
    used_ids = {row[0] for row in used}
    next_id = 1
    while next_id in used_ids:
        next_id += 1
    conn.close()
    return next_id

def start_boss(guild_id, boss_id, user_data, participant_ids=None):
    """
    Memulai boss battle dengan peserta tertentu.
    user_data: dict user leader (harus berisi id, level, guild_id)
    participant_ids: list of user_id (termasuk leader). Jika None, hanya leader.
    """
    import json

    # ── FALLBACK: Jika user_data bukan dict, coba ambil dari database ──
    if not isinstance(user_data, dict):
        try:
            uid = int(user_data)
            conn = get_conn()
            row = conn.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
            conn.close()
            if row:
                user_data = dict(row)
            else:
                return {"ok": False, "msg": "User tidak ditemukan."}
        except:
            return {"ok": False, "msg": "Data user tidak valid."}
    
    # ── Validasi user_data ──
    if not user_data:
        return {"ok": False, "msg": "Data user kosong."}
    if not isinstance(user_data, dict):
        return {"ok": False, "msg": f"Data user bukan dict: {type(user_data)}"}
    
    user_id = user_data.get("id")
    if not user_id:
        return {"ok": False, "msg": "ID user tidak ditemukan dalam data."}
    try:
        user_id = int(user_id)
    except:
        return {"ok": False, "msg": "ID user tidak valid."}
    if user_id <= 0:
        return {"ok": False, "msg": "ID user harus positif."}
    
    # ── Ambil data boss (bawaan atau custom buatan user/guild) ──
    boss = get_effective_boss(boss_id)
    if not boss:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_boss_not_found")}
    # ── Boss seasonal hanya saat event aktif (custom boss tidak dibatasi) ──
    if not is_boss_available(boss_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_boss_seasonal_inactive")}
    
    conn = get_conn()
    try:
        # ── CEK LIMIT BOSS PER HARI (GUILD) ──
        from datetime import date
        today = date.today().isoformat()
        guild = conn.execute(
            "SELECT boss_defeated_today, boss_defeated_date FROM guilds WHERE id=?",
            (guild_id,)
        ).fetchone()
        if guild:
            if guild["boss_defeated_date"] == today:
                if guild["boss_defeated_today"] >= 3:
                    conn.close()
                    return {"ok": False, "msg": "Guild sudah mengalahkan 3 boss hari ini. Tunggu besok!"}
            else:
                # Reset counter jika hari berbeda
                conn.execute(
                    "UPDATE guilds SET boss_defeated_today=0, boss_defeated_date=? WHERE id=?",
                    (today, guild_id)
                )
                conn.commit()
        else:
            # Guild tidak ditemukan (seharusnya tidak terjadi)
            conn.close()
            return {"ok": False, "msg": "Guild tidak ditemukan."}

        # ── Cek apakah leader berada di guild yang benar ──
        if user_data.get("guild_id") != guild_id:
            conn.close()
            return {"ok": False, "msg": "Kamu tidak berada di guild ini."}
        
        # ── Cek level leader ──
        if user_data["level"] < boss.get("min_level", 1):
            conn.close()
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_boss_level_too_low", min_lvl=boss['min_level'])}
        
        # ── Cek boss aktif ──
        if conn.execute("SELECT 1 FROM boss_battles WHERE guild_id=? AND status='active'", (guild_id,)).fetchone():
            conn.close()
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_boss_already_active")}
        
        # ── Proses peserta ──
        if participant_ids is None:
            participant_ids = [user_id]
        
        valid_participants = []
        for pid in participant_ids:
            try:
                pid = int(pid)
            except:
                continue
            if pid <= 0:
                continue
            p_row = conn.execute("SELECT id, level, guild_id FROM users WHERE id=?", (pid,)).fetchone()
            if p_row and p_row["guild_id"] == guild_id:
                p_level = p_row["level"]
                if p_level >= boss.get("min_level", 1) and p_level <= boss.get("max_level", 999):
                    valid_participants.append(pid)
        
        # Leader wajib ada
        if user_id not in valid_participants:
            valid_participants.insert(0, user_id)
        
        # Maksimal 5
        valid_participants = valid_participants[:5]
        
        # ── Insert boss battle ──
        tier_crit = {"beginner": 10, "normal": 15, "hard": 25, "elite": 35, "legendary": 50}
        crit_chance = tier_crit.get(boss["tier"], 15)
        
        participants_str = [str(uid) for uid in valid_participants]
        participants_json = json.dumps(participants_str)
        attack_counts_json = json.dumps({str(uid): 0 for uid in valid_participants})
        
        conn.execute(
            "INSERT INTO boss_battles(guild_id,boss_id,boss_name,boss_icon,boss_tier,"
            "boss_hp,boss_max_hp,boss_attack,boss_crit_chance,participants,attack_counts,raid_leader_id)"
            " VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
            (guild_id, boss_id, boss["name"], boss["icon"], boss["tier"],
             boss["hp"], boss["hp"], boss["atk"], crit_chance,
             participants_json, attack_counts_json, user_id)
        )
        conn.execute("UPDATE guilds SET quest_id=? WHERE id=?", (boss_id, guild_id))
        conn.commit()
        conn.close()
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_boss_started", icon=boss['icon'], name=boss['name'], tier=boss['tier'].upper())}
    
    except Exception as e:
        conn.rollback()
        conn.close()
        return {"ok": False, "msg": f"Error: {str(e)}"}

@retry_on_lock
def attack_boss(user_id, guild_id, action="light"):
    import random
    import json
    from datetime import datetime, timedelta, date

    conn = None
    try:
        conn = get_conn()
        
        # Ambil data user
        u = conn.execute("SELECT * FROM users WHERE id=?", (user_id,)).fetchone()
        if not u:
            return {"ok": False, "msg": "User not found"}
        u = dict(u)
        if u["hp"] <= 0:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_boss_hp_zero")}
        if u["mp"] < 0:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_mp_invalid")}

        # Ambil data boss
        boss_row = conn.execute(
            "SELECT * FROM boss_battles WHERE guild_id=? AND status='active'", (guild_id,)
        ).fetchone()
        if not boss_row:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_boss_no_active")}
        boss = dict(boss_row)

        # Cek peserta
        participants_data = boss.get("participants")
        participants = json.loads(participants_data) if participants_data else []
        participants = [str(p) for p in participants]
        if str(user_id) not in participants:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_raid_not_participant")}

        # Ambil buff skill
        buff_row = conn.execute(
            "SELECT skill_buff_data FROM users WHERE id=?", (user_id,)
        ).fetchone()
        buffs = json.loads(buff_row["skill_buff_data"]) if buff_row and buff_row["skill_buff_data"] else {}
        shield_active = buffs.get("shield_active", False)

        # Hitung damage maksimum user
        base_damage = 25
        user_max_dmg = int(base_damage + u.get("boss_damage_bonus", 0))
        if u.get("is_admin", 0):
            user_max_dmg *= 10
        if user_max_dmg < 1:
            user_max_dmg = 1

        ultimate_multipliers = {
            "warrior": 3,
            "mage": 4,
            "archer": 3.5,
            "healer": 2.5,
            "rogue": 5
        }
        cls = u.get("avatar_class", "warrior")

        # ── Action handling ──
        if action == "block":
            # ── CEK MP ──
            if u["mp"] < 5:
                return {"ok": False, "msg": tr_db(user_id=user_id, key="db_mp_insufficient_msg", cost=5, mp=u['mp'])}
            # ── KURANGI MP ──
            conn.execute("UPDATE users SET mp = mp - 5 WHERE id=?", (user_id,))
            
            block_strength = u.get("block_strength", 10)
            roll = random.randint(1, 100)
            if roll <= 50:
                reduction = block_strength
            else:
                reduction = random.randint(0, block_strength)
            buffs["block_active"] = True
            buffs["block_reduction"] = reduction
            conn.execute(
                "UPDATE users SET skill_buff_data=? WHERE id=?",
                (json.dumps(buffs), user_id)
            )
            conn.commit()
            conn.close()
            return {
                "ok": True,
                "action": "block",
                "block_reduction": reduction,
                "msg": tr_db(user_id=user_id, key="db_boss_block_success", reduction=reduction)
            }

        if action == "light":
            max_dmg_action = user_max_dmg // 2
            if max_dmg_action < 1:
                max_dmg_action = 1
            mp_cost = 0
        elif action == "heavy":
            max_dmg_action = user_max_dmg
            mp_cost = 5
            if u["mp"] < mp_cost:
                return {"ok": False, "msg": tr_db(user_id=user_id, key="db_mp_insufficient_msg", cost=mp_cost, mp=u['mp'])}
        elif action == "ultimate":
            max_dmg_action = int(user_max_dmg * ultimate_multipliers.get(cls, 3))
            mp_cost = 50
            if u["mp"] < mp_cost:
                return {"ok": False, "msg": tr_db(user_id=user_id, key="db_mp_insufficient_msg", cost=mp_cost, mp=u['mp'])}
            last_used = u.get("ultimate_last_used")
            if last_used:
                try:
                    last_dt = datetime.fromisoformat(last_used)
                    if datetime.now() - last_dt < timedelta(minutes=5):
                        remaining_sec = int((timedelta(minutes=5) - (datetime.now() - last_dt)).total_seconds())
                        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_ultimate_cooldown", seconds=remaining_sec)}
                except:
                    pass
            conn.execute("UPDATE users SET ultimate_last_used=? WHERE id=?", (datetime.now().isoformat(), user_id))
            extra_effect = ""
            if cls == "warrior":
                extra_effect = tr_db(user_id=user_id, key="db_ultimate_warrior_effect")
            elif cls == "mage":
                extra_effect = tr_db(user_id=user_id, key="db_ultimate_mage_effect")
            elif cls == "archer":
                extra_effect = tr_db(user_id=user_id, key="db_ultimate_archer_effect")
            elif cls == "healer":
                new_hp = min(u["max_hp"], u["hp"] + 20)
                conn.execute("UPDATE users SET hp=? WHERE id=?", (new_hp, user_id))
                extra_effect = tr_db(user_id=user_id, key="db_ultimate_healer_effect", heal=20)
            elif cls == "rogue":
                extra_effect = tr_db(user_id=user_id, key="db_ultimate_rogue_effect")
        else:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_unknown_action")}

        if mp_cost > 0:
            conn.execute("UPDATE users SET mp = mp - ? WHERE id=?", (mp_cost, user_id))

        # Critical user
        user_crit = u.get("crit_chance", 10)
        guild = conn.execute("SELECT crit_chance FROM guilds WHERE id=?", (guild_id,)).fetchone()
        if guild:
            user_crit += int(guild["crit_chance"])
        if action == "ultimate" and cls == "rogue":
            user_crit += 20

        roll_user = random.randint(1, 100)
        if roll_user <= user_crit:
            user_damage = max_dmg_action
            user_critical = True
        else:
            user_damage = random.randint(1, max(1, max_dmg_action - 1))
            user_critical = False

        # Update attack_counts
        attack_counts = json.loads(boss["attack_counts"]) if boss.get("attack_counts") else {}
        key = str(user_id)
        attack_counts[key] = attack_counts.get(key, 0) + 1
        conn.execute(
            "UPDATE boss_battles SET attack_counts=? WHERE id=?",
            (json.dumps(attack_counts), boss["id"])
        )

        # Terapkan damage ke boss
        new_boss_hp = max(0.0, boss["boss_hp"] - user_damage)

        # ── Jika boss mati ──
        if new_boss_hp <= 0:
            conn.execute(
                "UPDATE boss_battles SET boss_hp=0, status='defeated', ended_at=? WHERE id=?",
                (local_now().isoformat(), boss["id"])
            )

            # ── Tambah counter guild ──
            today = date.today().isoformat()
            conn.execute(
                """UPDATE guilds 
                   SET boss_defeated_today = boss_defeated_today + 1,
                       boss_defeated_date = ?
                   WHERE id = ?""",
                (today, guild_id)
            )

            bdata = BOSSES.get(boss["boss_id"], {})
            members = conn.execute(
                "SELECT user_id FROM guild_members WHERE guild_id=?", (guild_id,)
            ).fetchall()
            cnt = max(1, len(members))
            xp_reward = bdata.get("xp", 200) // cnt
            gold_reward = bdata.get("gold", 50) // cnt

            rewarded_members = []
            for m in members:
                uid = str(m["user_id"])
                if attack_counts.get(uid, 0) > 0:
                    rewarded_members.append(m["user_id"])
                    conn.execute("""
                        INSERT INTO boss_rewards(user_id, guild_id, boss_name, boss_tier, xp_reward, gold_reward)
                        VALUES(?,?,?,?,?,?)
                    """, (m["user_id"], guild_id, boss["boss_name"], boss["boss_tier"], xp_reward, gold_reward))

            conn.commit()
            conn.close()
            conn = None

            # Berikan reward setelah koneksi ditutup
            for uid in rewarded_members:
                gain_xp_gold(uid, xp_reward, gold_reward)
                check_achievements(uid, "boss_kill", 1)
                add_notification(uid, tr_db(user_id=uid, key="db_boss_defeated_notif", name=boss['boss_name']), "success")

            return {
                "ok": True,
                "defeated": True,
                "action": action,
                "user_damage": user_damage,
                "user_critical": user_critical,
                "boss_hp_left": 0,
                "extra_effect": extra_effect if action == "ultimate" else "",
                "msg": tr_db(user_id=user_id, key="db_boss_defeated", name=boss['boss_name'])
            }

        # ── Boss serang balik ──
        boss_max_dmg = int(boss["boss_attack"])
        if boss_max_dmg < 1:
            boss_max_dmg = 1
        boss_crit = int(boss.get("boss_crit_chance", 15))
        roll_boss = random.randint(1, 100)
        if roll_boss <= boss_crit:
            boss_damage_raw = boss_max_dmg
            boss_critical = True
        else:
            boss_damage_raw = random.randint(1, max(1, boss_max_dmg - 1))
            boss_critical = False

        # Block
        block_reduction = 0
        if buffs.get("block_active"):
            block_reduction = buffs.get("block_reduction", 0)
            boss_damage_raw = max(0, boss_damage_raw - block_reduction)
            buffs.pop("block_active", None)
            buffs.pop("block_reduction", None)
            conn.execute(
                "UPDATE users SET skill_buff_data=? WHERE id=?",
                (json.dumps(buffs), user_id)
            )

        # Shield Bash
        if shield_active:
            boss_damage_raw = boss_damage_raw // 2
            buffs.pop("shield_active", None)
            conn.execute(
                "UPDATE users SET skill_buff_data=? WHERE id=?",
                (json.dumps(buffs), user_id)
            )

        # Damage ke user
        reduc = u.get("hp_damage_reduction", 0)
        actual = max(0.0, boss_damage_raw - reduc)
        actual_int = int(actual)
        actual_damage = actual_int
        new_hp = max(0, u["hp"] - actual_int)
        new_hp_int = int(new_hp)

        revived = False
        if new_hp_int == 0 and u.get("has_revive"):
            new_hp_int = int(u["max_hp"] * 0.3)
            conn.execute("UPDATE users SET hp=?, has_revive=0 WHERE id=?", (new_hp_int, user_id))
            revived = True
        else:
            conn.execute("UPDATE users SET hp=? WHERE id=?", (new_hp_int, user_id))

        conn.execute("UPDATE boss_battles SET boss_hp=? WHERE id=?", (new_boss_hp, boss["id"]))

        conn.commit()
        conn.close()
        conn = None

        if revived:
            recalculate_all_buffs(user_id)
            add_notification(user_id, tr_db(user_id=user_id, key="db_totem_revive"), "success")

        return {
            "ok": True,
            "defeated": False,
            "action": action,
            "user_damage": user_damage,
            "user_critical": user_critical,
            "boss_hp_left": new_boss_hp,
            "boss_max_hp": boss["boss_max_hp"],
            "boss_damage": boss_damage_raw,
            "boss_critical": boss_critical,
            "actual_damage": actual_damage,
            "shield_used": shield_active,
            "block_reduction": block_reduction,
            "revived": revived,
            "extra_effect": extra_effect if action == "ultimate" else "",
            "mp_cost": mp_cost
        }

    except Exception as e:
        if conn:
            conn.rollback()
            conn.close()
        log_crash(f"attack_boss error: {e}")
        return {"ok": False, "msg": f"Error: {str(e)}"}

def get_unclaimed_boss_rewards(user_id):
    """Ambil daftar reward boss yang belum diklaim oleh user"""
    conn = get_conn()
    rows = conn.execute("""
        SELECT * FROM boss_rewards 
        WHERE user_id=? AND is_claimed=0 
        ORDER BY created_at DESC
    """, (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def claim_boss_reward(reward_id, user_id):
    """Klaim reward, berikan XP dan Gold ke user, tandai sebagai sudah diklaim.
    PENTING: Koneksi harus ditutup sebelum memanggil gain_xp_gold agar tidak terjadi
    'database is locked' karena nested connections.
    """
    # ── Step 1: Baca data reward, lalu TUTUP koneksi ──────────────────────────
    conn = get_conn()
    try:
        reward_row = conn.execute(
            "SELECT * FROM boss_rewards WHERE id=? AND user_id=? AND is_claimed=0",
            (reward_id, user_id)
        ).fetchone()
        if not reward_row:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_boss_reward_invalid")}
        reward = dict(reward_row)   # konversi ke dict sebelum close
    except Exception as e:
        log_crash(f"claim_boss_reward SELECT error: {e}")
        return {"ok": False, "msg": f"Error membaca reward: {e}"}
    finally:
        conn.close()                # ← TUTUP sebelum panggil gain_xp_gold

    # ── Step 2: Berikan reward (gain_xp_gold membuka/menutup koneksinya sendiri) ─
    try:
        result = gain_xp_gold(user_id, reward["xp_reward"], reward["gold_reward"])
    except Exception as e:
        log_crash(f"claim_boss_reward gain_xp_gold error: {e}")
        return {"ok": False, "msg": f"Error saat memberikan reward XP/Gold: {e}"}

    # ── Step 3: Tandai sebagai diklaim (koneksi baru, tidak ada yg nested) ────
    conn2 = get_conn()
    try:
        conn2.execute("UPDATE boss_rewards SET is_claimed=1 WHERE id=?", (reward_id,))
        conn2.commit()
    except Exception as e:
        log_crash(f"claim_boss_reward UPDATE is_claimed error: {e}")
        return {"ok": False, "msg": f"Error menandai reward sebagai diklaim: {e}"}
    finally:
        conn2.close()

    return {
        "ok": True,
        "msg": tr_db(user_id=user_id, key="db_boss_reward_claimed", xp=reward['xp_reward'], gold=reward['gold_reward']),
        "xp_gained": reward["xp_reward"],
        "gold_gained": reward["gold_reward"],
        "leveled_up": result.get("leveled_up", False),
    }

# ── Stats ─────────────────────────────────────────────────────────────────────

def get_stats(user_id):
    conn = get_conn()
    u  = dict(conn.execute(
        "SELECT * FROM users WHERE id=?", (user_id,)).fetchone())
    hd = conn.execute(
        "SELECT COUNT(*) c FROM habits WHERE user_id=? AND done_today=1",
        (user_id,)).fetchone()["c"]
    ht = conn.execute(
        "SELECT COUNT(*) c FROM habits WHERE user_id=?",
        (user_id,)).fetchone()["c"]
    dd = conn.execute(
        "SELECT COUNT(*) c FROM dailies WHERE user_id=? AND done_today=1",
        (user_id,)).fetchone()["c"]
    dt = conn.execute(
        "SELECT COUNT(*) c FROM dailies WHERE user_id=?",
        (user_id,)).fetchone()["c"]
    td = conn.execute(
        "SELECT COUNT(*) c FROM todos WHERE user_id=? AND done=1",
        (user_id,)).fetchone()["c"]
    tt = conn.execute(
        "SELECT COUNT(*) c FROM todos WHERE user_id=?",
        (user_id,)).fetchone()["c"]
    ms = conn.execute(
        "SELECT MAX(streak) s FROM habits WHERE user_id=?",
        (user_id,)).fetchone()["s"] or 0
    ic = conn.execute(
        "SELECT COUNT(*) c FROM inventory WHERE user_id=?",
        (user_id,)).fetchone()["c"]
    pc = conn.execute(
        "SELECT COUNT(*) c FROM user_pets WHERE user_id=?",
        (user_id,)).fetchone()["c"]
    bk = conn.execute(
        "SELECT COUNT(*) c FROM boss_battles bb"
        " JOIN guild_members gm ON bb.guild_id=gm.guild_id"
        " WHERE gm.user_id=? AND bb.status='defeated'",
        (user_id,)).fetchone()["c"]
    log = conn.execute(
        "SELECT * FROM activity_log WHERE user_id=?"
        " ORDER BY created_at DESC LIMIT 30",
        (user_id,)).fetchall()
    wk = conn.execute(
        "SELECT date(created_at) day,"
        "SUM(xp_gained) xp, SUM(gold_gained) gold"
        " FROM activity_log WHERE user_id=?"
        " AND created_at>=date('now','-7 days')"
        " GROUP BY day ORDER BY day",
        (user_id,)).fetchall()
    conn.close()
    return {
        "user": u,
        "habits_done_today": hd, "habits_total": ht,
        "dailies_done_today": dd, "dailies_total": dt,
        "todos_done": td, "todos_total": tt,
        "max_streak": ms, "inv_count": ic, "pet_count": pc,
        "bosses_killed": bk,
        "recent_log": [dict(r) for r in log],
        "weekly": [dict(r) for r in wk],
    }


# ── Sport Stats helper (dipakai get_stats & StatsPage) ───────────────────────
# ============================================================
# SPORT REP TRACKING & RANKS — push-up/pull-up/dll per variasi
# ============================================================

# Tier rank berdasarkan TOTAL reps kumulatif per aktivitas.
# (min_reps, key) — nama terjemahan via key "sport_rank_<key>".
SPORT_REP_RANKS = [
    (0,    "rookie"),
    (50,   "bronze"),
    (150,  "silver"),
    (300,  "gold"),
    (600,  "platinum"),
    (1200, "diamond"),
    (2500, "master"),
    (5000, "mythic"),
]

_RANK_ICONS = {
    "rookie": "🎽", "bronze": "🥉", "silver": "🥈", "gold": "🥇",
    "platinum": "💠", "diamond": "💎", "master": "👑", "mythic": "🌌",
}


def get_rep_rank(total_reps):
    """Info tier rank untuk total reps tertentu.

    Return: {index, key, icon, min_reps, next_reps(None=jika MAX), progress_frac}
    — nama tampil diterjemahkan di UI via tr("sport_rank_<key>")."""
    total = max(0, int(total_reps or 0))
    idx = 0
    for i, (min_r, _key) in enumerate(SPORT_REP_RANKS):
        if total >= min_r:
            idx = i
        else:
            break
    min_reps, key = SPORT_REP_RANKS[idx]
    next_reps = SPORT_REP_RANKS[idx + 1][0] if idx + 1 < len(SPORT_REP_RANKS) else None
    if next_reps:
        frac = min(1.0, max(0.0, (total - min_reps) / (next_reps - min_reps)))
    else:
        frac = 1.0
    return {
        "index": idx, "key": key, "icon": _RANK_ICONS.get(key, "🎽"),
        "min_reps": min_reps, "next_reps": next_reps, "progress_frac": frac,
    }


def add_sport_rep_log(user_id, activity_id, reps, sets=1, log_date=None, note=""):
    """Catat sesi reps untuk satu aktivitas. Return info rank before/after
    supaya UI bisa merayakan kenaikan rank."""
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    try:
        reps = int(reps)
        sets = max(1, int(sets))
    except (TypeError, ValueError):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="msg_invalid_amount")}
    if reps <= 0:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="msg_invalid_amount")}
    if not log_date:
        log_date = date.today().isoformat()

    conn = get_conn()
    a = conn.execute(
        "SELECT id FROM sport_activities WHERE id=? AND user_id=?",
        (activity_id, user_id)
    ).fetchone()
    if not a:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_sport_activity_not_found")}
    total_before = conn.execute(
        "SELECT COALESCE(SUM(reps),0) AS t FROM sport_rep_logs WHERE user_id=? AND activity_id=?",
        (user_id, activity_id)
    ).fetchone()["t"]
    conn.execute("""
        INSERT INTO sport_rep_logs(user_id, activity_id, reps, sets, log_date, note, created_at)
        VALUES(?,?,?,?,?,?,?)
    """, (user_id, activity_id, reps, sets, log_date, (note or "").strip(),
          local_now().isoformat()))
    conn.commit()
    total_after = total_before + reps
    conn.close()

    before, after = get_rep_rank(total_before), get_rep_rank(total_after)
    return {
        "ok": True, "total_reps": total_after,
        "rank_before": before, "rank_after": after,
        "rank_up": after["index"] > before["index"],
    }


def get_sport_rep_total(user_id, activity_id):
    """Total reps kumulatif satu aktivitas (penggerak rank)."""
    conn = get_conn()
    row = conn.execute(
        "SELECT COALESCE(SUM(reps),0) AS t FROM sport_rep_logs WHERE user_id=? AND activity_id=?",
        (user_id, activity_id)
    ).fetchone()
    conn.close()
    return row["t"]


def get_sport_reps_between(user_id, start_date, end_date, activity_id=None):
    """Total reps pada rentang tanggal (aktivitas tertentu atau semua)."""
    conn = get_conn()
    q = """SELECT COALESCE(SUM(reps),0) AS t FROM sport_rep_logs
           WHERE user_id=? AND log_date BETWEEN ? AND ?"""
    params = [user_id, start_date, end_date]
    if activity_id is not None:
        q += " AND activity_id=?"
        params.append(activity_id)
    row = conn.execute(q, params).fetchone()
    conn.close()
    return row["t"]


def get_sport_rep_series(user_id, end_date=None, days=7, activity_id=None):
    """Deret reps harian untuk chart: [{date, reps}] dengan zero-fill."""
    end_d = date.fromisoformat(end_date) if end_date else date.today()
    start_d = end_d - timedelta(days=days - 1)
    conn = get_conn()
    q = """SELECT log_date, COALESCE(SUM(reps),0) AS reps FROM sport_rep_logs
           WHERE user_id=? AND log_date BETWEEN ? AND ?"""
    params = [user_id, start_d.isoformat(), end_d.isoformat()]
    if activity_id is not None:
        q += " AND activity_id=?"
        params.append(activity_id)
    q += " GROUP BY log_date"
    rows = conn.execute(q, params).fetchall()
    conn.close()
    by_date = {r["log_date"]: r["reps"] for r in rows}
    out = []
    for i in range(days):
        d = (start_d + timedelta(days=i)).isoformat()
        out.append({"date": d, "reps": by_date.get(d, 0)})
    return out


def get_sport_rep_history(user_id, activity_id, limit=5):
    """Sesi reps terakhir untuk satu aktivitas (terbaru dulu)."""
    conn = get_conn()
    rows = conn.execute("""
        SELECT * FROM sport_rep_logs
        WHERE user_id=? AND activity_id=?
        ORDER BY log_date DESC, id DESC LIMIT ?
    """, (user_id, activity_id, limit)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def get_user_sport_rep_stats(user_id):
    """Ringkasan reps per aktivitas: {activity_id: {"total": t, "today": x,
    "week": y, "rank": get_rep_rank(t)}} — dipakai kartu & header Sport page."""
    today = date.today().isoformat()
    week_start = (date.today() - timedelta(days=6)).isoformat()
    conn = get_conn()
    rows = conn.execute("""
        SELECT activity_id,
               SUM(reps) AS total,
               SUM(CASE WHEN log_date=? THEN reps ELSE 0 END) AS today,
               SUM(CASE WHEN log_date>=? THEN reps ELSE 0 END) AS week
        FROM sport_rep_logs WHERE user_id=? GROUP BY activity_id
    """, (today, week_start, user_id)).fetchall()
    conn.close()
    out = {}
    for r in rows:
        total = r["total"] or 0
        out[r["activity_id"]] = {
            "total": total, "today": r["today"] or 0, "week": r["week"] or 0,
            "rank": get_rep_rank(total),
        }
    return out


def get_sport_stats(user_id):
    conn = get_conn()
    total_sport = conn.execute(
        "SELECT COUNT(*) c FROM sport_activities WHERE user_id=?",
        (user_id,)).fetchone()["c"]
    done_sport_today = conn.execute(
        "SELECT COUNT(*) c FROM sport_activities WHERE user_id=? AND done_today=1",
        (user_id,)).fetchone()["c"]
    max_sport_streak = conn.execute(
        "SELECT MAX(streak) s FROM sport_activities WHERE user_id=?",
        (user_id,)).fetchone()["s"] or 0
    u = conn.execute(
        "SELECT COALESCE(sport_level,1) as sport_level,"
        " COALESCE(sport_xp,0) as sport_xp,"
        " COALESCE(total_sport_points_earned,0) as total_sport_points_earned"
        " FROM users WHERE id=?", (user_id,)).fetchone()
    conn.close()
    return {
        "total_sport": total_sport,
        "done_sport_today": done_sport_today,
        "max_sport_streak": max_sport_streak,
        "sport_level": u["sport_level"] if u else 1,
        "sport_xp": u["sport_xp"] if u else 0,
        "total_sport_points_earned": u["total_sport_points_earned"] if u else 0,
    }


# ── Notifications ─────────────────────────────────────────────────────────────

@retry_on_lock
def add_notification(user_id, message, type_="info"):
    conn = get_conn()
    try:
        conn.execute(
            "INSERT INTO notifications(user_id, message, type, created_at) VALUES(?,?,?,?)",
            (user_id, message, type_, local_now().isoformat())
        )
        conn.commit()
    finally:
        conn.close()


def get_notifications(user_id, unread_only=True):
    conn = get_conn()
    q = ("SELECT * FROM notifications WHERE user_id=?"
         + (" AND is_read=0" if unread_only else "")
         + " ORDER BY created_at DESC LIMIT 20")
    rows = conn.execute(q, (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def mark_read(user_id):
    conn = get_conn()
    conn.execute("UPDATE notifications SET is_read=1 WHERE user_id=?",(user_id,))
    conn.execute("UPDATE cloud_notifications_cache SET is_read=1 WHERE local_user_id=?",(user_id,))
    conn.commit();conn.close()


def get_notification_center(user_id,limit=50,offset=0,type_filter=None):
    conn=get_conn();params=[user_id];where="n.user_id=?"
    if type_filter and type_filter!="all":
        where+=" AND (COALESCE(c.notification_type,n.type)=? OR COALESCE(c.entity_type,'')=?)";params.extend([type_filter,type_filter])
    params.extend([max(1,min(500,int(limit))),max(0,int(offset))])
    rows=conn.execute(f"""SELECT n.*,c.notification_type,c.actor_cloud_id,c.entity_type,c.entity_id,c.payload AS cloud_payload
      FROM notifications n LEFT JOIN cloud_notifications_cache c ON c.cloud_id=n.cloud_id
      WHERE {where} ORDER BY n.created_at DESC,n.id DESC LIMIT ? OFFSET ?""",tuple(params)).fetchall();conn.close()
    result=[]
    for row in rows:
        item=dict(row)
        try:item["payload"]=json.loads(item.get("cloud_payload") or "{}")
        except (TypeError,ValueError):item["payload"]={}
        result.append(item)
    return result


def mark_notification_read(user_id,notification_id):
    conn=get_conn();row=conn.execute("SELECT cloud_id FROM notifications WHERE id=? AND user_id=?",(notification_id,user_id)).fetchone()
    conn.execute("UPDATE notifications SET is_read=1 WHERE id=? AND user_id=?",(notification_id,user_id))
    if row and row["cloud_id"]:conn.execute("UPDATE cloud_notifications_cache SET is_read=1 WHERE cloud_id=? AND local_user_id=?",(row["cloud_id"],user_id))
    conn.commit();conn.close();return row["cloud_id"] if row else None


@retry_on_lock
def log_activity(user_id, action, detail, xp, gold):
    conn = get_conn()
    conn.execute(
        "INSERT INTO activity_log(user_id, action, detail, xp_gained, gold_gained, created_at) VALUES(?,?,?,?,?,?)",
        (user_id, action, detail, xp, gold, local_now().isoformat())
    )
    conn.commit()
    conn.close()


# ── Avatar / Settings ─────────────────────────────────────────────────────────

AVATAR_CLASSES = {
    "warrior": {"name": "Warrior", "icon": "⚔️",
                "bonus": "HP+20%, Skill: Shield Bash (10 MP)"},
    "mage":    {"name": "Mage",    "icon": "🧙",
                "bonus": "XP+15%, Skill: Arcane Surge (15 MP)"},
    "archer":  {"name": "Archer",  "icon": "🏹",
                "bonus": "Gold+10%, Skill: Gold Shot (10 MP)"},
    "healer":  {"name": "Healer",  "icon": "💊",
                "bonus": "Skill: Regenerate +30 HP (20 MP)"},
    "rogue":   {"name": "Rogue",   "icon": "🗡️",
                "bonus": "Streak bonus, Skill: Shadow Step (15 MP)"},
}

THEMES = {
    # ── Modern Aurora (default, toggle target) ──
    "modern_dark": {
        "label": "🔮 Aurora Dark",
        "primary": "#8b5cf6", "light": "#a78bfa",
        "bg": "#0b0b16", "bg2": "#14122c", "bg3": "#0c1a26",
        "panel": "rgba(26,28,50,0.90)", "border": "rgba(150,140,255,0.22)",
        "accent": "#22d3ee", "accent2": "#8b5cf6", "accent3": "#22d3ee", "glow": "#a78bfa",
        "text": "#eef0ff", "muted": "#9b9fc4",
    },
    "modern_light": {
        "label": "🌤️ Aurora Light",
        "primary": "#7c3aed", "light": "#6d28d9",
        "bg": "#eceffb", "bg2": "#e7e0f8", "bg3": "#e2eff7",
        "panel": "rgba(255,255,255,0.92)", "border": "rgba(124,108,210,0.30)",
        "accent": "#0891b2", "accent2": "#7c3aed", "accent3": "#0891b2", "glow": "#8b5cf6",
        "text": "#1b1633", "muted": "#5c5570",
    },
    "overworld": {
        "label": "🌿 Overworld",
        "primary": "#34d399", "light": "#6ee7b7",
        "bg": "#06120c", "bg2": "#0a2114", "bg3": "#0c1a14",
        "panel": "rgba(20,42,28,0.90)", "border": "rgba(80,255,150,0.22)",
        "accent": "#a3e635", "accent2": "#34d399", "accent3": "#a3e635", "glow": "#6ee7b7",
        "text": "#e6f7ec", "muted": "#8aaa9a",
    },
    "nether": {
        "label": "🔥 Nether",
        "primary": "#f43f5e", "light": "#fb7185",
        "bg": "#13060a", "bg2": "#240a10", "bg3": "#1a0a06",
        "panel": "rgba(44,18,22,0.90)", "border": "rgba(255,96,80,0.22)",
        "accent": "#fb923c", "accent2": "#f43f5e", "accent3": "#fb923c", "glow": "#fb7185",
        "text": "#ffe9ea", "muted": "#b08a8a",
    },
    "the_end": {
        "label": "🌌 The End",
        "primary": "#c026d3", "light": "#e879f9",
        "bg": "#0a0614", "bg2": "#170a28", "bg3": "#10081c",
        "panel": "rgba(30,20,50,0.90)", "border": "rgba(205,125,255,0.22)",
        "accent": "#818cf8", "accent2": "#c026d3", "accent3": "#818cf8", "glow": "#e879f9",
        "text": "#f3eaff", "muted": "#a896b8",
    },
    "ocean": {
        "label": "🌊 Ocean",
        "primary": "#0ea5e9", "light": "#38bdf8",
        "bg": "#04101a", "bg2": "#082033", "bg3": "#06141f",
        "panel": "rgba(16,36,54,0.90)", "border": "rgba(80,200,255,0.22)",
        "accent": "#22d3ee", "accent2": "#0ea5e9", "accent3": "#22d3ee", "glow": "#38bdf8",
        "text": "#e3f4ff", "muted": "#8aa6b8",
    },
    "ancient_city": {
        "label": "🏚️ Ancient City",
        "primary": "#14b8a6", "light": "#2dd4bf",
        "bg": "#06141a", "bg2": "#0a2329", "bg3": "#06181c",
        "panel": "rgba(18,40,44,0.90)", "border": "rgba(60,230,210,0.22)",
        "accent": "#5eead4", "accent2": "#14b8a6", "accent3": "#5eead4", "glow": "#2dd4bf",
        "text": "#e3f8f4", "muted": "#88a8a0",
    },
}

def get_user_theme(user_id):
    u = get_user(user_id)
    return THEMES.get(u.get("theme", "modern_dark"), THEMES["modern_dark"])


def set_user_theme(user_id, key):
    if key in THEMES:
        update_user(user_id, theme=key)
        return {"ok": True}
    return {"ok": False}


def set_avatar(user_id, avatar_class=None, color=None,
               emoji=None, bio=None, display_name=None):
    kw = {}
    if avatar_class and avatar_class in AVATAR_CLASSES:
        kw["avatar_class"] = avatar_class
    if color:
        kw["avatar_color"] = color
    if emoji:
        kw["avatar_emoji"] = emoji
    if bio is not None:
        kw["bio"] = bio
    if display_name:
        kw["display_name"] = display_name
    if kw:
        update_user(user_id, **kw)
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_avatar_updated")}

def change_class(user_id, new_class):
    """Ganti class avatar — maksimal sekali sehari."""
    if new_class not in AVATAR_CLASSES:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_class_unknown")}
    u = get_user(user_id)
    last_change = u.get("last_class_change", "")
    today = date.today().isoformat()
    if last_change == today:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_class_change_cooldown")}
    update_user(user_id, avatar_class=new_class, last_class_change=today)
    update_class_passive_buffs(user_id)
    recalculate_all_buffs(user_id)
    apply_hp_multiplier(user_id)
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_class_changed", name=AVATAR_CLASSES[new_class]['name'])}

def apply_hp_multiplier(user_id):
    u = get_user(user_id)
    class_buffs = get_class_passive_buffs(user_id)
    hp_mult = class_buffs.get("hp_multiplier", 1.0)
    base_max_hp = 50 + (u["level"] - 1) * 10
    new_max_hp = int(base_max_hp * hp_mult)
    new_hp = min(u["hp"], new_max_hp)
    update_user(user_id, max_hp=new_max_hp, hp=new_hp)

def set_user_settings(user_id, sound_enabled=None):
    kw = {}
    if sound_enabled is not None:
        kw["sound_enabled"] = int(sound_enabled)
    if kw:
        update_user(user_id, **kw)

# ── Ekspor ke Excel ─────────────────────────────────────────
def export_habits_excel(user_id, filepath):
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    habits = get_habits(user_id)
    dailies = get_dailies(user_id)
    todos = get_todos(user_id)
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "CraftLife Data"
    
    # Header
    headers = ["Jenis", "Nama", "Kesulitan/Prioritas", "Streak", "Terakhir", "Catatan"]
    ws.append(headers)
    for col in range(1, 7):
        cell = ws.cell(row=1, column=col)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="5a8a2e", end_color="5a8a2e", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
    
    for h in habits:
        ws.append(["Habit", h["name"], h["difficulty"], h["streak"], h["last_done"], h["notes"]])
    for d in dailies:
        ws.append(["Daily", d["name"], d["difficulty"], d["streak"], d["last_done"], d["notes"]])
    for t in todos:
        ws.append(["Quest", t["name"], t["priority"], "", "", t["notes"]])
    
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[col_letter].width = adjusted_width
    
    wb.save(filepath)
    return filepath

# ── Ekspor ke Word ─────────────────────────────────────────
def export_habits_word(user_id, filepath):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    habits = get_habits(user_id)
    dailies = get_dailies(user_id)
    todos = get_todos(user_id)
    
    doc = Document()
    title = doc.add_heading("CraftLife - Data Habits", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Habits
    doc.add_heading("Habits", level=1)
    for h in habits:
        p = doc.add_paragraph()
        p.add_run(f"{h['icon']} {h['name']}  ").bold = True
        p.add_run(f"[{h['difficulty']}]  ")
        p.add_run(f"+{h['xp_reward']} XP, +{h['gold_reward']} Gold  ")
        if h['streak'] > 0:
            p.add_run(f"🔥 Streak: {h['streak']} hari")
        if h['notes']:
            doc.add_paragraph(f"📝 Catatan: {h['notes']}", style="Intense Quote")
    
    # Dailies
    doc.add_heading("Dailies", level=1)
    for d in dailies:
        p = doc.add_paragraph()
        p.add_run(f"{d['icon']} {d['name']}  ").bold = True
        p.add_run(f"[{d['difficulty']}]  ")
        p.add_run(f"+{d['xp_reward']} XP, +{d['gold_reward']} Gold  ")
        if d['streak'] > 0:
            p.add_run(f"🔥 Streak: {d['streak']} hari")
        if d['notes']:
            doc.add_paragraph(f"📝 Catatan: {d['notes']}", style="Intense Quote")
    
    # Todos / Quests
    doc.add_heading("Quests", level=1)
    for t in todos:
        p = doc.add_paragraph()
        p.add_run(f"{t['icon']} {t['name']}  ").bold = True
        p.add_run(f"[{t['priority']}]  ")
        p.add_run(f"+{t['xp_reward']} XP, +{t['gold_reward']} Gold")
        if t['notes']:
            doc.add_paragraph(f"📝 Catatan: {t['notes']}", style="Intense Quote")
    
    doc.save(filepath)
    return filepath

# ── Ekspor ke PDF ─────────────────────────────────────────
def export_habits_pdf(user_id, filepath):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    habits = get_habits(user_id)
    dailies = get_dailies(user_id)
    todos = get_todos(user_id)
    
    doc = SimpleDocTemplate(filepath, pagesize=landscape(A4))
    story = []
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name='Title', parent=styles['Title'], alignment=1, fontSize=16)
    story.append(Paragraph("CraftLife - Data Habits", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Data tabel
    data = [["Jenis", "Nama", "Difficulty/Priority", "Streak", "Last Done", "Notes"]]
    for h in habits:
        data.append(["Habit", h['name'], h['difficulty'], str(h['streak']), h['last_done'] or "-", h['notes'] or "-"])
    for d in dailies:
        data.append(["Daily", d['name'], d['difficulty'], str(d['streak']), d['last_done'] or "-", d['notes'] or "-"])
    for t in todos:
        data.append(["Quest", t['name'], t['priority'], "-", "-", t['notes'] or "-"])
    
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#5a8a2e")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 10),
        ('BOTTOMPADDING', (0,0), (-1,0), 8),
        ('BACKGROUND', (0,1), (-1,-1), colors.beige),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(table)
    doc.build(story)
    return filepath

def export_habits_csv(user_id, filepath):
    import csv
    habits = get_habits(user_id)
    dailies = get_dailies(user_id)
    todos = get_todos(user_id)
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(["Jenis", "Nama", "Kesulitan", "Streak", "Terakhir", "Catatan"])
        for h in habits:
            writer.writerow(["Habit", h["name"], h["difficulty"], h["streak"], h["last_done"], h["notes"]])
        for d in dailies:
            writer.writerow(["Daily", d["name"], d["difficulty"], d["streak"], d["last_done"], d["notes"]])
        for t in todos:
            writer.writerow(["Quest", t["name"], t["priority"], "", "", t["notes"]])
    return filepath

# ── Friend System ─────────────────────────────────────────
def send_friend_request(user_id, target_username):
    conn = get_conn()
    # Cek apakah pengirim admin
    sender = conn.execute("SELECT is_admin FROM users WHERE id=?", (user_id,)).fetchone()
    if sender and sender["is_admin"]:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_friend_admin_cannot_send")}
    
    target = conn.execute("SELECT id, is_admin FROM users WHERE username=?", (target_username.lower(),)).fetchone()
    if not target:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_user_not_found")}
    if target["is_admin"]:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_friend_admin_target")}
    target = conn.execute("SELECT id FROM users WHERE username=?", (target_username.lower(),)).fetchone()
    if not target:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_user_not_found")}
    if target["id"] == user_id:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_friend_self")}
    existing = conn.execute("SELECT * FROM friends WHERE (user_id=? AND friend_id=?) OR (user_id=? AND friend_id=?)", 
                            (user_id, target["id"], target["id"], user_id)).fetchone()
    if existing:
        if existing["status"] == "accepted":
            conn.close()
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_friend_already")}
        else:
            conn.close()
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_friend_request_pending")}
    conn.execute("INSERT INTO friends(user_id, friend_id, status, action_user_id) VALUES(?,?,?,?)",
                 (user_id, target["id"], "pending", user_id))
    conn.commit()
    conn.close()
    add_notification(target["id"], tr_db(user_id=target["id"], key="db_friend_request_notif", name=get_user(user_id)['display_name']), "info")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_friend_request_sent")}

def accept_friend_request(user_id, request_id):
    conn = get_conn()
    # Cek apakah penerima admin
    user = conn.execute("SELECT is_admin FROM users WHERE id=?", (user_id,)).fetchone()
    if user and user["is_admin"]:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_friend_accept_admin_cannot")}
    req = conn.execute("SELECT * FROM friends WHERE id=? AND friend_id=? AND status='pending'", (request_id, user_id)).fetchone()
    if not req:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_friend_request_invalid")}
    conn.execute("UPDATE friends SET status='accepted' WHERE id=?", (request_id,))
    conn.commit()
    conn.close()
    check_achievements(req["user_id"], "friend_count", 1)
    check_achievements(user_id, "friend_count", 1)
    add_notification(req["user_id"], tr_db(user_id=req["user_id"], key="db_friend_accepted_notif", name=get_user(user_id)['display_name']), "success")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_friend_accepted")}

def reject_friend_request(user_id, request_id):
    conn = get_conn()
    conn.execute("DELETE FROM friends WHERE id=? AND friend_id=? AND status='pending'", (request_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_friend_rejected")}

def get_friends(user_id):
    conn = get_conn()
    # Cek jika user admin, langsung return []
    u = conn.execute("SELECT is_admin FROM users WHERE id=?", (user_id,)).fetchone()
    if u and u["is_admin"]:
        conn.close()
        return []
    rows = conn.execute("""
        SELECT u.id, u.display_name, u.username, u.avatar_emoji, u.avatar_color, u.level, u.cloud_user_id
        FROM friends f
        JOIN users u ON (f.user_id = u.id OR f.friend_id = u.id)
        WHERE (f.user_id=? OR f.friend_id=?) AND f.status='accepted' AND u.id != ?
        GROUP BY u.id
    """, (user_id, user_id, user_id)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_friendship_between(user_id,friend_id):
    conn=get_conn();row=conn.execute("""SELECT * FROM friends WHERE status='accepted'
        AND ((user_id=? AND friend_id=?) OR (user_id=? AND friend_id=?)) LIMIT 1""",
        (user_id,friend_id,friend_id,user_id)).fetchone();conn.close();return dict(row) if row else {}


def get_friendship_record(friendship_id,user_id=None):
    conn=get_conn()
    if user_id is None:row=conn.execute("SELECT * FROM friends WHERE id=?",(friendship_id,)).fetchone()
    else:row=conn.execute("SELECT * FROM friends WHERE id=? AND (user_id=? OR friend_id=?)",(friendship_id,user_id,user_id)).fetchone()
    conn.close();return dict(row) if row else {}


def get_couple_relationship_record(relationship_id,user_id=None):
    conn=get_conn()
    if user_id is None:row=conn.execute("SELECT * FROM couple_relationships WHERE id=?",(relationship_id,)).fetchone()
    else:row=conn.execute("SELECT * FROM couple_relationships WHERE id=? AND (user_a_id=? OR user_b_id=?)",(relationship_id,user_id,user_id)).fetchone()
    conn.close();return dict(row) if row else {}


def get_pending_friend_requests(user_id):
    conn = get_conn()
    u = conn.execute("SELECT is_admin FROM users WHERE id=?", (user_id,)).fetchone()
    if u and u["is_admin"]:
        conn.close()
        return []
    rows = conn.execute("""
        SELECT f.id, u.id AS sender_id, u.display_name, u.username
        FROM friends f
        JOIN users u ON f.user_id = u.id
        WHERE f.friend_id=? AND f.status='pending'
    """, (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


# ── Profile Photo / Couple Relationship / Shared Love Space ─────────────
PROFILE_PHOTO_MAX_BYTES = 3 * 1024 * 1024
LOVE_PHOTO_MAX_BYTES = 4 * 1024 * 1024


def _decoded_image_dimensions(blob):
    if blob.startswith(b"\x89PNG\r\n\x1a\n") and len(blob) >= 24 and blob[12:16] == b"IHDR":
        return int.from_bytes(blob[16:20], "big"), int.from_bytes(blob[20:24], "big"), "image/png"
    if blob.startswith(b"\xff\xd8\xff") and blob.endswith(b"\xff\xd9"):
        pos = 2; sof = {0xC0,0xC1,0xC2,0xC3,0xC5,0xC6,0xC7,0xC9,0xCA,0xCB,0xCD,0xCE,0xCF}
        while pos + 4 <= len(blob):
            if blob[pos] != 0xFF:
                pos += 1; continue
            while pos < len(blob) and blob[pos] == 0xFF: pos += 1
            if pos >= len(blob): break
            marker = blob[pos]; pos += 1
            if marker in (0xD8,0xD9) or 0xD0 <= marker <= 0xD7: continue
            if pos + 2 > len(blob): break
            length = int.from_bytes(blob[pos:pos+2], "big")
            if length < 2 or pos + length > len(blob): break
            if marker in sof and length >= 7:
                height = int.from_bytes(blob[pos+3:pos+5], "big")
                width = int.from_bytes(blob[pos+5:pos+7], "big")
                return width, height, "image/jpeg"
            pos += length
    return 0, 0, ""


def _validate_stored_image(image_data, mime_type, width, height, max_bytes):
    if not isinstance(image_data, (bytes, bytearray, memoryview)):
        return False, "invalid_data"
    blob = bytes(image_data)
    if not blob or len(blob) > max_bytes:
        return False, "file_size"
    actual_width, actual_height, detected_mime = _decoded_image_dimensions(blob)
    if not detected_mime or detected_mime != mime_type:
        return False, "file_type"
    if ((actual_width, actual_height) != (int(width), int(height))
            or not (32 <= actual_width <= 4096 and 32 <= actual_height <= 4096)):
        return False, "dimensions"
    return True, ""


def set_profile_photo(actor_user_id, target_user_id, image_data, mime_type, width, height):
    if int(actor_user_id) != int(target_user_id):
        return {"ok": False, "code": "forbidden", "msg": tr_db(user_id=actor_user_id, key="photo_forbidden")}
    ok, code = _validate_stored_image(image_data, mime_type, width, height, PROFILE_PHOTO_MAX_BYTES)
    if not ok:
        return {"ok": False, "code": code, "msg": tr_db(user_id=actor_user_id, key=f"photo_error_{code}")}
    blob = bytes(image_data)
    conn = get_conn(); conn.execute("""
        INSERT INTO user_profile_photos(user_id,image_data,mime_type,width,height,size_bytes,updated_at)
        VALUES(?,?,?,?,?,?,datetime('now'))
        ON CONFLICT(user_id) DO UPDATE SET image_data=excluded.image_data,
            mime_type=excluded.mime_type,width=excluded.width,height=excluded.height,
            size_bytes=excluded.size_bytes,updated_at=datetime('now')
    """, (target_user_id, blob, mime_type, int(width), int(height), len(blob)))
    conn.commit(); conn.close()
    return {"ok": True, "code": "saved"}


def remove_profile_photo(actor_user_id, target_user_id):
    if int(actor_user_id) != int(target_user_id):
        return {"ok": False, "code": "forbidden", "msg": tr_db(user_id=actor_user_id, key="photo_forbidden")}
    conn = get_conn(); conn.execute("DELETE FROM user_profile_photos WHERE user_id=?", (target_user_id,))
    conn.commit(); conn.close(); return {"ok": True}


def get_profile_photo(user_id):
    conn = get_conn(); row = conn.execute(
        "SELECT image_data,mime_type,width,height,size_bytes,updated_at FROM user_profile_photos WHERE user_id=?",
        (user_id,)).fetchone(); conn.close()
    if not row:
        return None
    data = dict(row); data["image_data"] = bytes(data["image_data"]); return data


def _canonical_couple_pair(user_id, other_user_id):
    a, b = sorted((int(user_id), int(other_user_id)))
    return a, b


def _friendship_is_accepted(conn, user_id, other_user_id):
    return conn.execute("""
        SELECT 1 FROM friends WHERE status='accepted'
          AND ((user_id=? AND friend_id=?) OR (user_id=? AND friend_id=?))
    """, (user_id, other_user_id, other_user_id, user_id)).fetchone() is not None


def get_active_couple_relationship(user_id):
    conn = get_conn(); row = conn.execute("""
        SELECT cr.*, ls.id AS love_space_id
        FROM couple_relationships cr
        LEFT JOIN love_spaces ls ON ls.couple_relationship_id=cr.id
        WHERE (cr.status='accepted' OR (cr.cloud_status='ended' AND datetime(cr.cloud_grace_ends_at)>datetime('now')))
          AND (cr.user_a_id=? OR cr.user_b_id=?)
        ORDER BY CASE WHEN cr.status='accepted' THEN 0 ELSE 1 END LIMIT 1
    """, (user_id, user_id)).fetchone(); conn.close()
    if not row:return None
    result=dict(row);result["read_only_grace"]=result.get("cloud_status")=="ended";return result


def get_couple_context(user_id):
    relationship = get_active_couple_relationship(user_id)
    owner = get_user(user_id)
    if not relationship:
        return {"active": False, "user": owner, "partner": None, "relationship": None, "love_space_id": None}
    partner_id = relationship["user_b_id"] if relationship["user_a_id"] == user_id else relationship["user_a_id"]
    return {"active": True, "user": owner, "partner": get_user(partner_id),
            "relationship": relationship, "love_space_id": relationship.get("love_space_id")}


def get_couple_status_between(user_id, other_user_id):
    a, b = _canonical_couple_pair(user_id, other_user_id)
    conn = get_conn(); row = conn.execute(
        "SELECT * FROM couple_relationships WHERE user_a_id=? AND user_b_id=?", (a, b)).fetchone(); conn.close()
    if not row:
        return {"status": "friend"}
    data = dict(row)
    if data.get("cloud_status")=="ended":data["status"]="ended"
    if data["status"] == "pending":
        data["direction"] = "outgoing" if data["requested_by"] == user_id else "incoming"
    return data


def send_couple_request(user_id, friend_id):
    if int(user_id) == int(friend_id):
        return {"ok": False, "code": "self"}
    a, b = _canonical_couple_pair(user_id, friend_id)
    conn = get_conn()
    if not _friendship_is_accepted(conn, user_id, friend_id):
        conn.close(); return {"ok": False, "code": "not_friends"}
    busy = conn.execute("""
        SELECT id FROM couple_relationships WHERE status='accepted'
          AND (user_a_id IN (?,?) OR user_b_id IN (?,?)) LIMIT 1
    """, (user_id, friend_id, user_id, friend_id)).fetchone()
    if busy:
        conn.close(); return {"ok": False, "code": "partner_exists"}
    row = conn.execute("SELECT * FROM couple_relationships WHERE user_a_id=? AND user_b_id=?", (a, b)).fetchone()
    if row and row["status"] == "accepted":
        conn.close(); return {"ok": False, "code": "already_couple"}
    if row and row["status"] == "pending":
        conn.close(); return {"ok": False, "code": "pending"}
    if row:
        conn.execute("""UPDATE couple_relationships SET requested_by=?,status='pending',
                        created_at=datetime('now'),responded_at=NULL WHERE id=?""", (user_id, row["id"]))
        relationship_id = row["id"]
    else:
        cur = conn.execute("""INSERT INTO couple_relationships
            (user_a_id,user_b_id,requested_by,status) VALUES(?,?,?,'pending')""", (a, b, user_id))
        relationship_id = cur.lastrowid
    conn.commit(); conn.close()
    add_notification(friend_id, tr_db(user_id=friend_id, key="couple_request_notification",
                                     name=get_user(user_id).get("display_name", "")), "info")
    return {"ok": True, "code": "pending", "relationship_id": relationship_id}


def get_pending_couple_requests(user_id):
    conn = get_conn(); rows = conn.execute("""
        SELECT cr.*,
          CASE WHEN cr.user_a_id=? THEN cr.user_b_id ELSE cr.user_a_id END AS other_user_id,
          u.display_name AS other_display_name,u.username AS other_username
        FROM couple_relationships cr
        JOIN users u ON u.id=CASE WHEN cr.user_a_id=? THEN cr.user_b_id ELSE cr.user_a_id END
        WHERE cr.status='pending' AND (cr.user_a_id=? OR cr.user_b_id=?)
        ORDER BY cr.created_at DESC
    """, (user_id, user_id, user_id, user_id)).fetchall(); conn.close()
    out=[]
    for row in rows:
        item=dict(row); item["direction"]="outgoing" if item["requested_by"]==user_id else "incoming"; out.append(item)
    return out


def respond_couple_request(user_id, relationship_id, accept):
    conn = get_conn(); conn.execute("BEGIN IMMEDIATE")
    try:
        row = conn.execute("SELECT * FROM couple_relationships WHERE id=? AND status='pending'", (relationship_id,)).fetchone()
        if not row or user_id not in (row["user_a_id"], row["user_b_id"]) or row["requested_by"] == user_id:
            conn.rollback(); return {"ok": False, "code": "invalid_request"}
        if not accept:
            conn.execute("UPDATE couple_relationships SET status='rejected',responded_at=datetime('now') WHERE id=?", (relationship_id,))
            conn.commit(); return {"ok": True, "code": "rejected"}
        busy = conn.execute("""SELECT id FROM couple_relationships WHERE status='accepted' AND id!=?
            AND (user_a_id IN (?,?) OR user_b_id IN (?,?)) LIMIT 1""",
            (relationship_id, row["user_a_id"], row["user_b_id"], row["user_a_id"], row["user_b_id"])).fetchone()
        if busy:
            conn.rollback(); return {"ok": False, "code": "partner_exists"}
        if not _friendship_is_accepted(conn, row["user_a_id"], row["user_b_id"]):
            conn.rollback(); return {"ok": False, "code": "not_friends"}
        conn.execute("UPDATE couple_relationships SET status='accepted',responded_at=datetime('now') WHERE id=?", (relationship_id,))
        space = conn.execute("SELECT id FROM love_spaces WHERE couple_relationship_id=?", (relationship_id,)).fetchone()
        if space:
            love_space_id = space["id"]
        else:
            love_space_id = conn.execute("INSERT INTO love_spaces(couple_relationship_id) VALUES(?)", (relationship_id,)).lastrowid
        conn.executemany("INSERT OR IGNORE INTO love_space_members(love_space_id,user_id) VALUES(?,?)",
                         [(love_space_id,row["user_a_id"]),(love_space_id,row["user_b_id"])])
        # Keep a single shared profile/settings owner while preserving either user's old data.
        primary_id, secondary_id = min(row["user_a_id"], row["user_b_id"]), max(row["user_a_id"], row["user_b_id"])
        primary_profile = conn.execute("SELECT 1 FROM relationship_profiles WHERE user_id=?", (primary_id,)).fetchone()
        if not primary_profile:
            old_profile = conn.execute("SELECT * FROM relationship_profiles WHERE user_id=?", (secondary_id,)).fetchone()
            if old_profile:
                conn.execute("""INSERT OR IGNORE INTO relationship_profiles
                    (user_id,partner_name,partner_gender,partner_age,relationship_type,start_date,updated_at)
                    VALUES(?,?,?,?,?,?,?)""", (primary_id,old_profile["partner_name"],old_profile["partner_gender"],
                    old_profile["partner_age"],old_profile["relationship_type"],old_profile["start_date"],old_profile["updated_at"]))
        primary_cycle = conn.execute("SELECT 1 FROM menstrual_settings WHERE user_id=?", (primary_id,)).fetchone()
        if not primary_cycle:
            old_cycle = conn.execute("SELECT * FROM menstrual_settings WHERE user_id=?", (secondary_id,)).fetchone()
            if old_cycle:
                conn.execute("""INSERT OR IGNORE INTO menstrual_settings
                    (user_id,tracked_person,last_period_start,cycle_length,period_length,updated_at)
                    VALUES(?,?,?,?,?,?)""", (primary_id,old_cycle["tracked_person"],old_cycle["last_period_start"],
                    old_cycle["cycle_length"],old_cycle["period_length"],old_cycle["updated_at"]))
        # Other user-owned Love rows remain intact and become visible through shared scope reads.
        conn.commit()
        partner_id = row["user_b_id"] if row["user_a_id"] == user_id else row["user_a_id"]
        add_notification(partner_id, tr_db(user_id=partner_id, key="couple_request_accepted_notification",
                                           name=get_user(user_id).get("display_name", "")), "success")
        return {"ok": True, "code": "accepted", "love_space_id": love_space_id}
    except Exception:
        conn.rollback(); raise
    finally:
        conn.close()


def cancel_couple_request(user_id, relationship_id):
    conn = get_conn(); cur = conn.execute("""UPDATE couple_relationships
        SET status='cancelled',responded_at=datetime('now')
        WHERE id=? AND requested_by=? AND status='pending'""", (relationship_id, user_id))
    conn.commit(); conn.close()
    return {"ok": cur.rowcount > 0, "code": "cancelled" if cur.rowcount else "invalid_request"}


def end_local_couple_relationship(user_id,relationship_id):
    """End a local-only Couple immediately; cloud relationships use the server RPC/grace policy."""
    conn=get_conn();cur=conn.execute("""UPDATE couple_relationships SET status='cancelled',responded_at=datetime('now')
      WHERE id=? AND status='accepted' AND (user_a_id=? OR user_b_id=?)""",(relationship_id,user_id,user_id));conn.commit();changed=cur.rowcount;conn.close()
    return {"ok":bool(changed),"code":"ended" if changed else "invalid_relationship"}


def _love_scope_user_ids(user_id):
    relationship = get_active_couple_relationship(user_id)
    if not relationship:
        return [int(user_id)]
    return [int(relationship["user_a_id"]), int(relationship["user_b_id"])]


def _love_primary_user_id(user_id):
    return min(_love_scope_user_ids(user_id))


def _love_space_membership(conn, user_id, love_space_id):
    return conn.execute("SELECT 1 FROM love_space_members WHERE love_space_id=? AND user_id=?",
                        (love_space_id, user_id)).fetchone() is not None


def add_love_space_photo(requester_user_id, image_data, mime_type, width, height,
                         caption="", photo_date=None, visibility="private"):
    visibility = visibility if visibility in ("private", "shared") else "private"
    ok, code = _validate_stored_image(image_data, mime_type, width, height, LOVE_PHOTO_MAX_BYTES)
    if not ok:
        return {"ok": False, "code": code, "msg": tr_db(user_id=requester_user_id, key=f"photo_error_{code}")}
    love_space_id = None
    if visibility == "shared":
        relationship = get_active_couple_relationship(requester_user_id)
        if not relationship or not relationship.get("love_space_id"):
            return {"ok": False, "code": "no_couple"}
        love_space_id = relationship["love_space_id"]
        conn = get_conn()
        if not _love_space_membership(conn, requester_user_id, love_space_id):
            conn.close(); return {"ok": False, "code": "forbidden"}
        conn.close()
    blob=bytes(image_data); conn=get_conn(); cur=conn.execute("""
        INSERT INTO love_space_photos(love_space_id,owner_user_id,visibility,image_data,mime_type,
                                      width,height,size_bytes,caption,photo_date)
        VALUES(?,?,?,?,?,?,?,?,?,?)
    """, (love_space_id,requester_user_id,visibility,blob,mime_type,int(width),int(height),
          len(blob),caption.strip(),photo_date or date.today().isoformat()))
    conn.commit(); conn.close(); return {"ok": True, "photo_id": cur.lastrowid}


def get_love_space_photos(requester_user_id, limit=100):
    relationship = get_active_couple_relationship(requester_user_id)
    love_space_id = relationship.get("love_space_id") if relationship else None
    conn=get_conn()
    if love_space_id and _love_space_membership(conn, requester_user_id, love_space_id):
        rows=conn.execute("""SELECT p.*,u.display_name AS uploader_name FROM love_space_photos p
            JOIN users u ON u.id=p.owner_user_id
            WHERE (p.visibility='private' AND p.owner_user_id=?)
               OR (p.visibility='shared' AND p.love_space_id=?
                   AND EXISTS(SELECT 1 FROM love_space_members m WHERE m.love_space_id=p.love_space_id AND m.user_id=?))
            ORDER BY p.photo_date DESC,p.created_at DESC LIMIT ?""",
            (requester_user_id,love_space_id,requester_user_id,limit)).fetchall()
    else:
        rows=conn.execute("""SELECT p.*,u.display_name AS uploader_name FROM love_space_photos p
            JOIN users u ON u.id=p.owner_user_id
            WHERE p.visibility='private' AND p.owner_user_id=?
            ORDER BY p.photo_date DESC,p.created_at DESC LIMIT ?""", (requester_user_id,limit)).fetchall()
    conn.close(); out=[]
    for row in rows:
        item=dict(row); item["image_data"]=bytes(item["image_data"]); out.append(item)
    return out


def get_love_space_photo(requester_user_id, photo_id):
    photos=get_love_space_photos(requester_user_id, 1000)
    return next((photo for photo in photos if photo["id"]==photo_id), None)


def delete_love_space_photo(requester_user_id, photo_id):
    conn=get_conn(); cur=conn.execute(
        "DELETE FROM love_space_photos WHERE id=? AND owner_user_id=?", (photo_id,requester_user_id))
    conn.commit(); conn.close(); return {"ok": cur.rowcount>0, "code": "deleted" if cur.rowcount else "forbidden"}


# ========== GUILD KICK MEMBER ==========
def kick_guild_member(guild_id, leader_id, target_user_id):
    conn = get_conn()
    guild = conn.execute("SELECT leader_id, name FROM guilds WHERE id=?", (guild_id,)).fetchone()
    if not guild or guild["leader_id"] != leader_id:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=leader_id, key="db_guild_kick_only_leader")}
    if target_user_id == leader_id:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=leader_id, key="db_guild_kick_self")}
    member = conn.execute("SELECT 1 FROM guild_members WHERE guild_id=? AND user_id=?", (guild_id, target_user_id)).fetchone()
    if not member:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=leader_id, key="db_guild_kick_not_member")}
    conn.execute("DELETE FROM guild_members WHERE guild_id=? AND user_id=?", (guild_id, target_user_id))
    conn.execute("UPDATE users SET guild_id=NULL WHERE id=?", (target_user_id,))
    conn.commit()
    conn.close()
    add_notification(target_user_id, tr_db(user_id=target_user_id, key="db_guild_kicked_notif", name=guild['name']), "danger")
    return {"ok": True, "msg": tr_db(user_id=leader_id, key="db_guild_kicked")}

def transfer_guild_leadership(guild_id, current_leader_id, new_leader_id):
    """Transfer kepemimpinan guild dari current_leader ke new_leader."""
    conn = get_conn()
    try:
        # Ambil data guild (termasuk name)
        guild = conn.execute(
            "SELECT leader_id, name FROM guilds WHERE id = ?", (guild_id,)
        ).fetchone()
        if not guild or guild["leader_id"] != current_leader_id:
            return {"ok": False, "msg": tr_db(user_id=current_leader_id, key="db_guild_transfer_not_leader")}
        
        # Cek apakah new_leader adalah member guild
        member = conn.execute(
            "SELECT 1 FROM guild_members WHERE guild_id = ? AND user_id = ?",
            (guild_id, new_leader_id)
        ).fetchone()
        if not member:
            return {"ok": False, "msg": tr_db(user_id=current_leader_id, key="db_guild_transfer_not_member")}
        
        # Transfer leadership
        conn.execute(
            "UPDATE guilds SET leader_id = ? WHERE id = ?",
            (new_leader_id, guild_id)
        )
        conn.commit()
        
        # Kirim notifikasi
        add_notification(new_leader_id, tr_db(user_id=new_leader_id, key="db_guild_leader_transfer_notif", name=guild['name']), "success")
        add_notification(current_leader_id, tr_db(user_id=current_leader_id, key="db_guild_transfer_old_notif", name=guild['name'], id=new_leader_id), "info")
    except Exception as e:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=current_leader_id, key="db_guild_transfer_failed", error=str(e))}
    finally:
        conn.close()
    return {"ok": True, "msg": tr_db(user_id=current_leader_id, key="db_guild_transfer_success")}

# ========== GUILD LEADER TRANSFER (saat leader keluar) ==========
def leave_guild_with_transfer(user_id):
    u = get_user(user_id)
    if u.get("is_admin", 0):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_guild_leave_admin")}
    u = get_user(user_id)
    gid = u.get("guild_id")
    if not gid:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_guild_leave_not_in")}
    conn = get_conn()
    guild = conn.execute("SELECT leader_id, name FROM guilds WHERE id=?", (gid,)).fetchone()
    if not guild:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_guild_not_found")}
    
    # Jika user adalah leader
    if guild["leader_id"] == user_id:
        # Ambil member lain (kecuali dirinya sendiri)
        members = conn.execute("SELECT user_id FROM guild_members WHERE guild_id=? AND user_id!=?", (gid, user_id)).fetchall()
        if members:
            # Ada member lain: lakukan transfer leader
            conn.execute("INSERT INTO guild_leader_transfers(guild_id, old_leader_id) VALUES(?,?)", (gid, user_id))
            conn.commit()
            conn.execute("DELETE FROM guild_members WHERE user_id=? AND guild_id=?", (user_id, gid))
            conn.execute("UPDATE users SET guild_id=NULL WHERE id=?", (user_id,))
            conn.commit()
            for m in members:
                add_notification(m["user_id"], tr_db(user_id=m["user_id"], key="db_guild_leader_left", name=guild['name']), "warning")
            conn.close()
            return {"ok": True, "msg": tr_db(user_id=user_id, key="db_guild_leave_transfer")}
        else:
            # Tidak ada anggota lain: hapus semua data terkait guild, lalu guild
            conn.execute("DELETE FROM guild_members WHERE guild_id=?", (gid,))
            conn.execute("DELETE FROM guild_invites WHERE guild_id=?", (gid,))
            conn.execute("DELETE FROM guild_requests WHERE guild_id=?", (gid,))
            conn.execute("DELETE FROM guild_leader_transfers WHERE guild_id=?", (gid,))
            conn.execute("DELETE FROM boss_battles WHERE guild_id=?", (gid,))
            conn.execute("DELETE FROM guild_messages WHERE guild_id=?", (gid,))
            conn.execute("DELETE FROM boss_rewards WHERE guild_id=?", (gid,))   # <-- TAMBAHKAN INI
            conn.execute("DELETE FROM guilds WHERE id=?", (gid,))
            conn.execute("UPDATE users SET guild_id=NULL WHERE id=?", (user_id,))
            conn.commit()
            conn.close()
            return {"ok": True, "msg": tr_db(user_id=user_id, key="db_guild_leave_disband")}
    else:
        # Bukan leader: keluar biasa
        conn.execute("DELETE FROM guild_members WHERE user_id=? AND guild_id=?", (user_id, gid))
        conn.execute("UPDATE users SET guild_id=NULL WHERE id=?", (user_id,))
        conn.commit()
        conn.close()
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_guild_leave_success")}

def accept_leader_transfer(user_id, transfer_id):
    conn = get_conn()
    trans = conn.execute("SELECT * FROM guild_leader_transfers WHERE id=? AND status='pending'", (transfer_id,)).fetchone()
    if not trans:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_guild_transfer_invalid")}
    member = conn.execute("SELECT 1 FROM guild_members WHERE guild_id=? AND user_id=?", (trans["guild_id"], user_id)).fetchone()
    if not member:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_guild_transfer_not_member_accept")}
    conn.execute("UPDATE guilds SET leader_id=? WHERE id=?", (user_id, trans["guild_id"]))
    conn.execute("UPDATE guild_leader_transfers SET status='accepted' WHERE id=?", (transfer_id,))
    conn.commit()
    conn.close()
    add_notification(user_id, tr_db(user_id=user_id, key="db_guild_leader_accepted"), "success")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_guild_transfer_accepted")}

# ========== KICK FRIEND ==========
def remove_friend(user_id, friend_id):
    u = get_user(user_id)
    if u.get("is_admin", 0):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_friend_remove_admin")}
    conn = get_conn()
    conn.execute("DELETE FROM friends WHERE (user_id=? AND friend_id=?) OR (user_id=? AND friend_id=?)",
                 (user_id, friend_id, friend_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_friend_removed")}

# ========== PRIVATE CHAT ==========
def send_message(sender_id, receiver_id, message, created_at=None, reply_to_id=None):
    sender = get_user(sender_id)
    if sender.get("is_admin", 0):
        return {"ok": False, "msg": tr_db(user_id=sender_id, key="db_chat_admin_cannot_send")}
    conn = get_conn()
    if created_at is None:
        created_at = datetime.now().isoformat()
    if reply_to_id is not None:
        target=conn.execute("""SELECT 1 FROM messages WHERE id=? AND
          ((sender_id=? AND receiver_id=?) OR (sender_id=? AND receiver_id=?))""",
          (reply_to_id,sender_id,receiver_id,receiver_id,sender_id)).fetchone()
        if not target:reply_to_id=None
    cur=conn.execute(
        "INSERT INTO messages(sender_id,receiver_id,message,created_at,reply_to_id) VALUES(?,?,?,?,?)",
        (sender_id,receiver_id,message,created_at,reply_to_id)
    )
    conn.commit();message_id=cur.lastrowid;conn.close()
    add_notification(receiver_id, tr_db(user_id=receiver_id, key="db_chat_new_message", name=get_user(sender_id)['display_name']), "info")
    return {"ok": True,"message_id":message_id}

def get_messages(user_id, other_id, limit=50):
    conn = get_conn()
    rows = conn.execute("""
        SELECT * FROM messages
        WHERE ((sender_id=? AND receiver_id=?) OR (sender_id=? AND receiver_id=?))
        AND (deleted_by IS NULL OR deleted_by NOT LIKE ?)
        ORDER BY created_at DESC LIMIT ?
    """, (user_id, other_id, other_id, user_id, f'%{user_id}%', limit)).fetchall()
    result=[dict(r) for r in reversed(rows)]
    for message in result:
        reactions=conn.execute("SELECT user_id,reaction FROM local_message_reactions WHERE message_id=?",(message["id"],)).fetchall()
        message["reactions"]={str(row["user_id"]):row["reaction"] for row in reactions}
        attachments=conn.execute("SELECT * FROM chat_attachments_cache WHERE local_message_id=? AND deleted_at IS NULL ORDER BY id",(message["id"],)).fetchall()
        message["attachments"]=[dict(row) for row in attachments]
    conn.close();return result


def edit_local_message(user_id,message_id,body):
    body=(body or "").strip()
    if not body:return {"ok":False,"code":"empty"}
    conn=get_conn();cur=conn.execute("""UPDATE messages SET message=?,edited_at=datetime('now')
      WHERE id=? AND sender_id=? AND deleted_at IS NULL""",(body[:4000],message_id,user_id));conn.commit();changed=cur.rowcount;conn.close()
    return {"ok":bool(changed),"code":"ok" if changed else "forbidden"}


def delete_local_message(user_id,message_id):
    conn=get_conn();cur=conn.execute("""UPDATE messages SET message='',deleted_at=COALESCE(deleted_at,datetime('now')),edited_at=NULL
      WHERE id=? AND sender_id=?""",(message_id,user_id));changed=cur.rowcount
    if changed:
        conn.execute("DELETE FROM local_message_reactions WHERE message_id=?",(message_id,))
        conn.execute("UPDATE chat_attachments_cache SET deleted_at=datetime('now') WHERE local_message_id=?",(message_id,))
    conn.commit();conn.close();return {"ok":bool(changed)}


def set_local_message_reaction(user_id,message_id,reaction):
    conn=get_conn();exists=conn.execute("SELECT 1 FROM messages WHERE id=? AND deleted_at IS NULL",(message_id,)).fetchone()
    if not exists:conn.close();return {"ok":False}
    if reaction:
        conn.execute("""INSERT INTO local_message_reactions(message_id,user_id,reaction,updated_at) VALUES(?,?,?,datetime('now'))
          ON CONFLICT(message_id,user_id) DO UPDATE SET reaction=excluded.reaction,updated_at=datetime('now')""",
          (message_id,user_id,str(reaction)[:16]))
    else:conn.execute("DELETE FROM local_message_reactions WHERE message_id=? AND user_id=?",(message_id,user_id))
    conn.commit();conn.close();return {"ok":True,"active":bool(reaction)}


def get_unread_count_between(user_id,other_id):
    conn=get_conn();count=conn.execute("""SELECT COUNT(*) FROM messages WHERE receiver_id=? AND sender_id=?
      AND is_read=0 AND deleted_at IS NULL""",(user_id,other_id)).fetchone()[0];conn.close();return count


def mark_messages_read(user_id, other_id):
    conn = get_conn()
    conn.execute("UPDATE messages SET is_read=1 WHERE receiver_id=? AND sender_id=?", (user_id, other_id))
    conn.commit()
    conn.close()

def get_unread_count(user_id):
    conn = get_conn()
    count = conn.execute("SELECT COUNT(*) FROM messages WHERE receiver_id=? AND is_read=0 AND deleted_at IS NULL", (user_id,)).fetchone()[0]
    conn.close()
    return count

def clear_friend_chat(user_id, friend_id):
    """Tandai semua pesan antara user_id dan friend_id sebagai dihapus oleh user_id (soft delete)."""
    conn = get_conn()
    rows = conn.execute(
        "SELECT id, deleted_by FROM messages WHERE (sender_id=? AND receiver_id=?) OR (sender_id=? AND receiver_id=?)",
        (user_id, friend_id, friend_id, user_id)
    ).fetchall()
    for row in rows:
        deleted_by = row["deleted_by"] or ""
        ids = set(deleted_by.split(',')) if deleted_by else set()
        if str(user_id) not in ids:
            ids.add(str(user_id))
            new_deleted_by = ','.join(sorted(ids, key=int))
            conn.execute("UPDATE messages SET deleted_by=? WHERE id=?", (new_deleted_by, row["id"]))
    conn.commit()
    conn.close()

def clear_guild_chat(guild_id):
    """Hapus semua pesan guild (hard delete)."""
    conn = get_conn()
    conn.execute("DELETE FROM guild_messages WHERE guild_id=?", (guild_id,))
    conn.commit()
    conn.close()

# ========== GUILD CHAT ==========
def send_guild_message(guild_id, sender_id, message, created_at=None):
    sender = get_user(sender_id)
    if sender.get("is_admin", 0):
        return {"ok": False, "msg": tr_db(user_id=sender_id, key="db_guild_chat_admin_cannot")}
    if created_at is None:
        created_at = datetime.now().isoformat()
    # ── Step 1: Insert pesan & ambil member list, lalu TUTUP koneksi ─────────
    conn = get_conn()
    conn.execute(
        "INSERT INTO guild_messages(guild_id, sender_id, message, created_at) VALUES(?,?,?,?)",
        (guild_id, sender_id, message, created_at)
    )
    conn.commit()
    members = conn.execute(
        "SELECT user_id FROM guild_members WHERE guild_id=? AND user_id!=?",
        (guild_id, sender_id)
    ).fetchall()
    member_ids = [m["user_id"] for m in members]
    conn.close()   # ← TUTUP sebelum panggil get_user / add_notification
    # ── Step 2: Notifikasi (koneksi terpisah, tidak nested) ──────────────────
    sender_name = get_user(sender_id)["display_name"]
    for uid in member_ids:
        add_notification(uid, tr_db(user_id=uid, key="db_guild_chat_message", name=sender_name, msg=message[:50]), "info")
    return {"ok": True}

def get_guild_messages(guild_id, limit=100):
    conn = get_conn()
    rows = conn.execute("""
        SELECT gm.*, u.display_name
        FROM guild_messages gm
        JOIN users u ON gm.sender_id = u.id
        WHERE gm.guild_id=?
        ORDER BY gm.created_at DESC LIMIT ?
    """, (guild_id, limit)).fetchall()
    conn.close()
    return [dict(r) for r in reversed(rows)]

# ── Guild Invites (Undangan dari Leader) ─────────────────────────────────
def get_guild_invites(user_id):
    conn = get_conn()
    rows = conn.execute("""
        SELECT gi.*, g.name as guild_name 
        FROM guild_invites gi 
        JOIN guilds g ON gi.guild_id = g.id 
        WHERE gi.user_id=? AND gi.status='pending'
    """, (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def accept_invite(user_id, invite_id):
    u = get_user(user_id)
    if u.get("is_admin", 0):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_admin_cannot_guild")}
    conn = get_conn()
    inv = conn.execute("SELECT * FROM guild_invites WHERE id=? AND user_id=? AND status='pending'", (invite_id, user_id)).fetchone()
    if not inv:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_invite_invalid")}
    
    # Cek jumlah member
    count = conn.execute("SELECT COUNT(*) FROM guild_members WHERE guild_id=?", (inv["guild_id"],)).fetchone()[0]
    if count >= 20:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_guild_full")}
    
    conn.execute("UPDATE guild_invites SET status='accepted' WHERE id=?", (invite_id,))
    conn.execute("INSERT INTO guild_members(guild_id, user_id) VALUES(?,?)", (inv["guild_id"], user_id))
    conn.execute("UPDATE users SET guild_id=? WHERE id=?", (inv["guild_id"], user_id))
    conn.commit()
    conn.close()
    add_notification(user_id, tr_db(user_id=user_id, key="db_guild_joined"), "success")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_accept_invite_welcome")}

def reject_invite(user_id, invite_id):
    conn = get_conn()
    conn.execute("UPDATE guild_invites SET status='rejected' WHERE id=? AND user_id=?", (invite_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_invite_rejected")}


# ─────────────────────────────────────────────────────────────────────────────── #
def set_security_question(user_id, question, answer):
    """Simpan pertanyaan & jawaban keamanan (jawaban di-hash PBKDF2)"""
    conn = get_conn()
    conn.execute(
        "UPDATE users SET security_question=?, security_answer_hash=? WHERE id=?",
        (question, _hash_password(answer.strip().lower()), user_id)
    )
    conn.commit()
    conn.close()

def verify_security_answer(user_id, answer):
    """Cek apakah jawaban cocok"""
    conn = get_conn()
    row = conn.execute(
        "SELECT security_answer_hash FROM users WHERE id=?",
        (user_id,)
    ).fetchone()
    conn.close()
    if not row or not row["security_answer_hash"]:
        return False
    return _verify_password(answer.strip().lower(), row["security_answer_hash"])

def reset_password_by_security(user_id, new_password):
    """Reset password tanpa perlu email, asumsi sudah verifikasi jawaban"""
    ok_pw, key = validate_password_strength(new_password)
    if not ok_pw:
        return {"ok": False, "msg": tr_db(user_id=user_id, key=key)}
    conn = get_conn()
    conn.execute(
        "UPDATE users SET password_hash=? WHERE id=?",
        (_hash_password(new_password), user_id)
    )
    conn.commit()
    conn.close()
    delete_all_session_tokens(user_id)
    return {"ok": True}


def generate_backup_codes(user_id, num_codes=5):
    """Generate backup codes untuk user, simpan hash-nya."""
    conn = get_conn()
    # Hapus kode lama yang belum dipakai
    conn.execute("DELETE FROM backup_codes WHERE user_id=? AND is_used=0", (user_id,))
    codes = []
    for _ in range(num_codes):
        # Generate kode 8 karakter alfanumerik
        raw_code = secrets.token_hex(4).upper()  # 8 karakter
        code_hash = hashlib.sha256(raw_code.encode()).hexdigest()
        conn.execute("INSERT INTO backup_codes(user_id, code_hash) VALUES(?,?)", (user_id, code_hash))
        codes.append(raw_code)
    conn.commit()
    conn.close()
    return codes  # kembalikan kode asli (hanya sekali ini, untuk ditampilkan ke user)

def get_user_backup_codes(user_id, only_unused=True):
    """Ambil daftar backup codes (hash) user."""
    conn = get_conn()
    if only_unused:
        rows = conn.execute("SELECT id, code_hash FROM backup_codes WHERE user_id=? AND is_used=0", (user_id,)).fetchall()
    else:
        rows = conn.execute("SELECT id, code_hash, is_used FROM backup_codes WHERE user_id=?", (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def verify_backup_code(user_id, raw_code):
    """Cek apakah kode cadangan valid dan belum dipakai."""
    code_hash = hashlib.sha256(raw_code.upper().strip().encode()).hexdigest()
    conn = get_conn()
    row = conn.execute("SELECT id FROM backup_codes WHERE user_id=? AND code_hash=? AND is_used=0", (user_id, code_hash)).fetchone()
    if row:
        # Tandai sebagai terpakai
        conn.execute("UPDATE backup_codes SET is_used=1 WHERE id=?", (row["id"],))
        conn.commit()
        conn.close()
        return True
    conn.close()
    return False

def reset_password_with_backup_code(user_id, new_password):
    """Reset password setelah verifikasi backup code (fungsi terpisah agar aman)."""
    conn = get_conn()
    ok_pw, key = validate_password_strength(new_password)
    if not ok_pw:
        return {"ok": False, "msg": tr_db(user_id=user_id, key=key)}
    conn.execute("UPDATE users SET password_hash=? WHERE id=?", (_hash_password(new_password), user_id))
    conn.commit()
    conn.close()
    delete_all_session_tokens(user_id)
    return {"ok": True}

def regenerate_backup_codes(user_id):
    """Generate ulang semua backup codes (menghapus yang lama)."""
    return generate_backup_codes(user_id, num_codes=5)


# --- Economy CRUD --- #
# ============================================================
# CATATAN HUTANG (PIUTANG) — orang lain berhutang KE user.
# Membuat catatan → otomatis tercatat EXPENSE (uang keluar);
# saat ditandai lunas → otomatis tercatat INCOME (uang kembali).
# ============================================================

def add_debt_note(user_id, person_name, amount_idr, date_str=None, notes=''):
    """Buat catatan hutang orang lain + catat nominal sebagai expense."""
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    name = (person_name or "").strip()
    if not name:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="debnote_person_empty")}
    try:
        amt = float(amount_idr)
    except (TypeError, ValueError):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="msg_invalid_amount")}
    if amt <= 0:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="msg_invalid_amount")}
    if not date_str:
        date_str = date.today().isoformat()
    else:
        try:
            datetime.strptime(date_str, "%Y-%m-%d")
        except ValueError:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="economy_debt_date_invalid")}

    # 1) Catat expense lebih dulu (uang keluar ke peminjam)
    cat = tr_db(user_id=user_id, key="economy_debnote_category")
    exp_name = tr_db(user_id=user_id, key="economy_debnote_expense_name", name=name)
    r = add_economy_item(user_id, exp_name, "📤", "expense", amt, cat, date_str, notes)
    if not r.get("ok"):
        return {"ok": False, "msg": r.get("msg", tr_db(user_id=user_id, key="msg_invalid_amount"))}

    # 2) Simpan catatan dengan tautan ke entri expense
    conn = get_conn()
    cur = conn.execute("""
        INSERT INTO debt_notes(user_id, person_name, amount, date, notes, status, expense_item_id, created_at)
        VALUES(?,?,?,?,?,'unpaid',?,?)
    """, (user_id, name, amt, date_str, notes, r["id"], local_now().isoformat()))
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "id": new_id}


def get_debt_notes(user_id, status=None):
    """Daftar catatan hutang: yang belum lunas di atas, lalu tanggal terbaru."""
    conn = get_conn()
    q = "SELECT * FROM debt_notes WHERE user_id=?"
    params = [user_id]
    if status in ("unpaid", "paid"):
        q += " AND status=?"
        params.append(status)
    q += " ORDER BY CASE status WHEN 'unpaid' THEN 0 ELSE 1 END, date DESC, id DESC"
    rows = conn.execute(q, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def get_total_unpaid_debt_notes(user_id):
    """Total nominal yang belum dilunasi orang lain ke user (IDR)."""
    conn = get_conn()
    row = conn.execute(
        "SELECT COALESCE(SUM(amount),0) AS t FROM debt_notes WHERE user_id=? AND status='unpaid'",
        (user_id,)
    ).fetchone()
    conn.close()
    return row["t"]


def settle_debt_note(user_id, note_id):
    """Tandai lunas + catat nominal sebagai INCOME (uang kembali ke user)."""
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn = get_conn()
    row = conn.execute(
        "SELECT * FROM debt_notes WHERE id=? AND user_id=?", (note_id, user_id)
    ).fetchone()
    if not row:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="debnote_not_found")}
    if row["status"] == "paid":
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="debnote_already_paid")}
    row = dict(row)
    conn.close()

    # Catat income (pelunasan = uang masuk kembali)
    cat = tr_db(user_id=user_id, key="economy_debnote_category")
    inc_name = tr_db(user_id=user_id, key="economy_debnote_income_name", name=row["person_name"])
    r = add_economy_item(user_id, inc_name, "📥", "income", row["amount"], cat,
                         date.today().isoformat(), row["notes"])
    if not r.get("ok"):
        return {"ok": False, "msg": r.get("msg", tr_db(user_id=user_id, key="msg_invalid_amount"))}

    conn = get_conn()
    conn.execute(
        "UPDATE debt_notes SET status='paid', paid_at=?, income_item_id=? WHERE id=?",
        (local_now().isoformat(), r["id"], note_id)
    )
    conn.commit()
    conn.close()
    return {"ok": True, "msg": tr_db(user_id=user_id, key="debnote_settled", name=row["person_name"])}


def delete_debt_note(user_id, note_id):
    """Hapus catatan hutang BESERTA entri ekonomi terkait (expense/income)."""
    conn = get_conn()
    row = conn.execute(
        "SELECT expense_item_id, income_item_id FROM debt_notes WHERE id=? AND user_id=?",
        (note_id, user_id)
    ).fetchone()
    if not row:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="debnote_not_found")}
    linked_ids = [row["expense_item_id"], row["income_item_id"]]
    conn.execute("DELETE FROM debt_notes WHERE id=? AND user_id=?", (note_id, user_id))
    for eid in linked_ids:
        if eid:
            conn.execute("DELETE FROM economy_items WHERE id=? AND user_id=?", (eid, user_id))
    conn.commit()
    conn.close()
    return {"ok": True}


def get_economy_items(user_id, type_filter=None, category_filter=None, search=None):
    """Ambil semua item ekonomi user, dengan filter opsional."""
    conn = get_conn()
    query = "SELECT * FROM economy_items WHERE user_id = ?"
    params = [user_id]
    if type_filter and type_filter in ('income', 'expense'):
        query += " AND type = ?"
        params.append(type_filter)
    if category_filter and category_filter != 'all':
        query += " AND category = ?"
        params.append(category_filter)
    if search:
        query += " AND name LIKE ?"
        params.append(f'%{search}%')
    query += " ORDER BY sort_order, date DESC, created_at DESC"   # sort_order utama
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def add_economy_item(user_id, name, icon, type_, amount, category, date_str, notes='', folder_id=None):
    """Tambah pemasukan/pengeluaran."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    # Item baru tampil PALING ATAS (terbaru → terlama) per folder (termasuk NULL).
    # sort_order ASC = posisi teratas, jadi item baru mengambil min-1.
    if folder_id is None:
        min_order = conn.execute(
            "SELECT COALESCE(MIN(sort_order), 1) FROM economy_items WHERE user_id=? AND folder_id IS NULL",
            (user_id,)
        ).fetchone()[0]
    else:
        min_order = conn.execute(
            "SELECT COALESCE(MIN(sort_order), 1) FROM economy_items WHERE user_id=? AND folder_id=?",
            (user_id, folder_id)
        ).fetchone()[0]
    new_order = min_order - 1

    cur = conn.execute("""
        INSERT INTO economy_items(user_id, name, icon, type, amount, category, date, notes, folder_id, sort_order, updated_at)
        VALUES(?,?,?,?,?,?,?,?,?,?, datetime('now'))
    """, (user_id, name, icon, type_, amount, category, date_str, notes, folder_id, new_order))
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "id": new_id}

def update_economy_item(item_id, user_id, **kwargs):
    """Update item ekonomi (nama, icon, type, amount, category, date, notes, folder_id)."""
    if not kwargs:
        return
    kwargs['updated_at'] = 'datetime("now")'  # special case
    fields = []
    values = []
    for k, v in kwargs.items():
        if k == 'updated_at':
            fields.append("updated_at = datetime('now')")
        else:
            fields.append(f"{k} = ?")
            values.append(v)
    values.append(item_id)
    values.append(user_id)
    query = f"UPDATE economy_items SET {', '.join(fields)} WHERE id = ? AND user_id = ?"
    conn = get_conn()
    conn.execute(query, values)
    conn.commit()
    conn.close()

def delete_economy_item(user_id, item_id):
    conn = get_conn()
    conn.execute("DELETE FROM economy_items WHERE id = ? AND user_id = ?", (item_id, user_id))
    conn.commit()
    conn.close()

def duplicate_economy_item(user_id, item_id):
    """Duplikasi item ekonomi (tanpa folder_id biar user pilih sendiri)."""
    conn = get_conn()
    original = conn.execute("SELECT * FROM economy_items WHERE id = ? AND user_id = ?", (item_id, user_id)).fetchone()
    if not original:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_redeem_item_fail")}
    d = dict(original)
    d.pop('id')
    d.pop('created_at')
    d.pop('updated_at')
    d['name'] = d['name']  # biarkan sama, user bisa edit nanti
    d['folder_id'] = None
    conn.execute("""
        INSERT INTO economy_items(user_id, name, icon, type, amount, category, date, notes, folder_id)
        VALUES(?,?,?,?,?,?,?,?,?)
    """, (d['user_id'], d['name'], d['icon'], d['type'], d['amount'], d['category'], d['date'], d['notes'], d['folder_id']))
    conn.commit()
    conn.close()
    return {"ok": True}

def get_economy_summary(user_id, year=None, month=None):
    """Ringkasan pemasukan, pengeluaran, saldo per bulan.
       Jika year/month None, ambil bulan berjalan berdasarkan date item.
    """
    conn = get_conn()
    # Ambil data semua item
    rows = conn.execute("SELECT type, amount FROM economy_items WHERE user_id = ?", (user_id,)).fetchall()
    conn.close()
    total_income = sum(r['amount'] for r in rows if r['type'] == 'income')
    total_expense = sum(r['amount'] for r in rows if r['type'] == 'expense')
    balance = total_income - total_expense
    return {
        'total_income': total_income,
        'total_expense': total_expense,
        'balance': balance,
    }

def get_economy_count(user_id):
    """Ambil jumlah total transaksi ekonomi user."""
    conn = get_conn()
    row = conn.execute("SELECT COUNT(*) as cnt FROM economy_items WHERE user_id = ?", (user_id,)).fetchone()
    conn.close()
    return row['cnt'] if row else 0

def get_economy_weekly(user_id):
    """Return list of daily income/expense for last 7 days."""
    conn = get_conn()
    rows = conn.execute("""
        SELECT date(date) as day,
               SUM(CASE WHEN type='income' THEN amount ELSE 0 END) as income,
               SUM(CASE WHEN type='expense' THEN amount ELSE 0 END) as expense
        FROM economy_items
        WHERE user_id = ? AND date >= date('now', '-6 days')
        GROUP BY date(date)
        ORDER BY day
    """, (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_economy_daily_totals(user_id, date_str):
    """Total income/expense/net untuk SATU tanggal (kolom date, format ISO).
    Dipakai kartu harian di tab Harian halaman Economy."""
    conn = get_conn()
    row = conn.execute("""
        SELECT
            COALESCE(SUM(CASE WHEN type='income' THEN amount END), 0) AS income,
            COALESCE(SUM(CASE WHEN type='expense' THEN amount END), 0) AS expense
        FROM economy_items
        WHERE user_id = ? AND date = ?
    """, (user_id, date_str)).fetchone()
    conn.close()
    inc, exp = row["income"], row["expense"]
    return {"income": inc, "expense": exp, "net": inc - exp}


def get_economy_balance_until(user_id, date_str):
    """Saldo kumulatif (income - expense) untuk SEMUA transaksi s.d. akhir
    tanggal date_str — 'saldo per hari itu', melengkapi saldo all-time."""
    conn = get_conn()
    row = conn.execute("""
        SELECT
            COALESCE(SUM(CASE WHEN type='income' THEN amount END), 0) AS income,
            COALESCE(SUM(CASE WHEN type='expense' THEN amount END), 0) AS expense
        FROM economy_items
        WHERE user_id = ? AND date <= ?
    """, (user_id, date_str)).fetchone()
    conn.close()
    return row["income"] - row["expense"]


def get_economy_items_by_date(user_id, date_str):
    """Item ekonomi pada satu tanggal — urutan terbaru di atas (konsisten
    dengan urutan tab Semua)."""
    conn = get_conn()
    rows = conn.execute("""
        SELECT * FROM economy_items
        WHERE user_id = ? AND date = ?
        ORDER BY sort_order, created_at DESC, id DESC
    """, (user_id, date_str)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def get_economy_daily_series(user_id, end_date=None, days=7):
    """Deret harian untuk chart: [{date, income, expense}] sebanyak `days`
    hari BERURUTAN berakhir di end_date (default: hari ini). Hari tanpa
    transaksi tetap muncul dengan nilai 0 supaya chart tidak bolong."""
    end_d = date.fromisoformat(end_date) if end_date else date.today()
    start_d = end_d - timedelta(days=days - 1)
    conn = get_conn()
    rows = conn.execute("""
        SELECT date,
               COALESCE(SUM(CASE WHEN type='income' THEN amount END), 0) AS income,
               COALESCE(SUM(CASE WHEN type='expense' THEN amount END), 0) AS expense
        FROM economy_items
        WHERE user_id = ? AND date BETWEEN ? AND ?
        GROUP BY date
    """, (user_id, start_d.isoformat(), end_d.isoformat())).fetchall()
    conn.close()
    by_date = {r["date"]: (r["income"], r["expense"]) for r in rows}
    out = []
    for i in range(days):
        d = (start_d + timedelta(days=i)).isoformat()
        inc, exp = by_date.get(d, (0, 0))
        out.append({"date": d, "income": inc, "expense": exp})
    return out


def get_economy_categories(user_id):
    """Ambil daftar kategori unik yang pernah dipakai user."""
    conn = get_conn()
    rows = conn.execute("SELECT DISTINCT category FROM economy_items WHERE user_id = ? ORDER BY category", (user_id,)).fetchall()
    conn.close()
    return [r['category'] for r in rows]


def get_full_export_data(user_id):
    """Ambil semua data user untuk keperluan ekspor."""
    user = get_user(user_id)
    stats = get_stats(user_id)
    sport_stats = get_sport_stats(user_id)
    eco_summary = get_economy_summary(user_id)
    eco_weekly = get_economy_weekly(user_id)
    habits = get_habits(user_id)
    dailies = get_dailies(user_id)
    todos = get_todos(user_id)
    sport_activities = get_sport_activities(user_id)
    economy_items = get_economy_items(user_id)
    health_summary = get_health_summary(user_id, days=30)
    health_logs = get_health_logs(user_id, days=30)
    return {
        "user": user,
        "stats": stats,
        "sport_stats": sport_stats,
        "eco_summary": eco_summary,
        "eco_weekly": eco_weekly,
        "habits": habits,
        "dailies": dailies,
        "todos": todos,
        "sport_activities": sport_activities,
        "economy_items": economy_items,
        "health_summary": health_summary,
        "health_logs": health_logs,
    }


# ========== DEBT FUNCTIONS ==========
@retry_on_lock
def add_debt(user_id, name, amount, due_date=None, notes=''):
    conn = get_conn()
    try:
        cur = conn.execute(
            "INSERT INTO debts(user_id, name, amount, due_date, notes) VALUES(?,?,?,?,?)",
            (user_id, name, amount, due_date, notes))
        debt_id = cur.lastrowid
        conn.commit()
        return {"ok": True, "debt_id": debt_id}
    except Exception as e:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_general_error", error=str(e))}
    finally:
        conn.close()

@retry_on_lock
def get_debts(user_id, include_paid=False):
    conn = get_conn()
    if include_paid:
        rows = conn.execute("""
            SELECT *, 
                CASE 
                    WHEN due_date IS NOT NULL AND due_date < date('now') THEN 1 
                    ELSE 0 
                END as is_overdue
            FROM debts 
            WHERE user_id=?
            ORDER BY is_paid ASC, due_date ASC, created_at DESC
        """, (user_id,)).fetchall()
    else:
        rows = conn.execute(
            "SELECT * FROM debts WHERE user_id=? AND is_paid=0 ORDER BY due_date ASC, created_at DESC",
            (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

# ========== SAVINGS FUNCTIONS ==========
@retry_on_lock
def add_saving(user_id, name, icon, target_amount, current_amount, target_date, notes):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    try:
        cur = conn.execute("""
            INSERT INTO savings(user_id, name, icon, target_amount, current_amount, target_date, notes)
            VALUES(?,?,?,?,?,?,?)
        """, (user_id, name, icon, target_amount, current_amount, target_date, notes))
        saving_id = cur.lastrowid
        conn.commit()
        return {"ok": True, "saving_id": saving_id}
    except Exception as e:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_general_error", error=str(e))}
    finally:
        conn.close()

@retry_on_lock
def get_savings(user_id):
    conn = get_conn()
    rows = conn.execute("""
        SELECT * FROM savings 
        WHERE user_id=?
        ORDER BY target_date ASC, created_at DESC
    """, (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def update_saving(saving_id, user_id, **kwargs):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    if not kwargs:
        return
    fields = ", ".join(f"{k}=?" for k in kwargs)
    conn = get_conn()
    conn.execute(
        f"UPDATE savings SET {fields}, updated_at=datetime('now') WHERE id=? AND user_id=?",
        list(kwargs.values()) + [saving_id, user_id])
    conn.commit()
    conn.close()

@retry_on_lock
def delete_saving(saving_id, user_id):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn = get_conn()
    conn.execute("DELETE FROM savings WHERE id=? AND user_id=?", (saving_id, user_id))
    conn.commit()
    conn.close()

@retry_on_lock
def add_to_saving(saving_id, user_id, amount):
    """Tambahkan dana ke tabungan (update current_amount)"""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    try:
        cur = conn.execute("SELECT current_amount FROM savings WHERE id=? AND user_id=?", (saving_id, user_id))
        row = cur.fetchone()
        if not row:
            return {"ok": False, "msg": "Tabungan tidak ditemukan"}
        new_amount = row["current_amount"] + amount
        conn.execute("UPDATE savings SET current_amount=? WHERE id=?", (new_amount, saving_id))
        conn.commit()
        return {"ok": True, "new_amount": new_amount}
    except Exception as e:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_general_error", error=str(e))}
    finally:
        conn.close()

@retry_on_lock
def withdraw_from_saving(saving_id, user_id, amount):
    """Kurangi current_amount tabungan, return new amount"""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    try:
        cur = conn.execute("SELECT current_amount FROM savings WHERE id=? AND user_id=?", (saving_id, user_id))
        row = cur.fetchone()
        if not row:
            return {"ok": False, "msg": "Tabungan tidak ditemukan"}
        if row["current_amount"] < amount:
            return {"ok": False, "msg": "Saldo tabungan tidak cukup"}
        new_amount = row["current_amount"] - amount
        conn.execute("UPDATE savings SET current_amount=? WHERE id=?", (new_amount, saving_id))
        conn.commit()
        return {"ok": True, "new_amount": new_amount}
    except Exception as e:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_general_error", error=str(e))}
    finally:
        conn.close()

@retry_on_lock
def update_debt(debt_id, user_id, **kwargs):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    if not kwargs:
        return
    fields = ", ".join(f"{k}=?" for k in kwargs)
    conn = get_conn()
    conn.execute(
        f"UPDATE debts SET {fields} WHERE id=? AND user_id=?",
        list(kwargs.values()) + [debt_id, user_id])
    conn.commit()
    conn.close()

@retry_on_lock
def delete_debt(debt_id, user_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("DELETE FROM debts WHERE id=? AND user_id=?", (debt_id, user_id))
    conn.commit()
    conn.close()

@retry_on_lock
def mark_debt_paid(debt_id, user_id):
    """
    Tandai hutang sebagai lunas, kurangi saldo ekonomi (pemasukan - pengeluaran),
    buat transaksi expense otomatis, TIDAK mempengaruhi gold.
    """
    # Step 1: Ambil data hutang
    conn = get_conn()
    debt = conn.execute("SELECT * FROM debts WHERE id=? AND user_id=? AND is_paid=0", (debt_id, user_id)).fetchone()
    if not debt:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_debt_not_found")}
    debt = dict(debt)
    conn.close()

    # Step 2: Hitung saldo ekonomi saat ini (total pemasukan - total pengeluaran)
    eco_items = get_economy_items(user_id)
    total_income = sum(i['amount'] for i in eco_items if i['type'] == 'income')
    total_expense = sum(i['amount'] for i in eco_items if i['type'] == 'expense')
    current_balance = total_income - total_expense

    # Step 3: Validasi kecukupan saldo
    if current_balance < debt['amount']:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_debt_insufficient_balance", balance=current_balance)}

    # Step 4: Buat transaksi expense otomatis
    from datetime import date
    today = date.today().isoformat()
    expense_name = f"PELUNASAN HUTANG: {debt['name']}"
    add_economy_item(
        user_id,
        name=expense_name,
        icon="💸",
        type_="expense",
        amount=debt["amount"],
        category="Hutang",
        date_str=today,
        notes=f"Pelunasan hutang {debt['name']} (ID hutang {debt_id})"
    )

    # Step 5: Catat activity log (opsional)
    log_activity(user_id, "debt_paid", tr_db(user_id=user_id, key="log_debt_paid", name=debt['name'], amount=debt['amount']), 0, 0)

    # Step 6: Tandai hutang sebagai lunas
    conn2 = get_conn()
    conn2.execute(
        "UPDATE debts SET is_paid=1, paid_at=? WHERE id=?",
        (datetime.now().isoformat(), debt_id))
    conn2.commit()
    conn2.close()

    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_debt_paid", name=debt['name'], amount=debt['amount'])}

def apply_late_debt_penalty(user_id, debt_id, days_late):
    """Terapkan penalti untuk hutang yang terlambat >3 hari.
       Penalti = 10 * (level/5) gold, minimal 10, maksimal 100.
       Penalti hanya diterapkan sekali per hutang.
       Mengembalikan dict dengan jumlah gold yang hilang.
    """
    conn = get_conn()
    try:
        # Cek hutang
        debt = conn.execute(
            "SELECT is_paid, penalty_applied, name FROM debts WHERE id=? AND user_id=?",
            (debt_id, user_id)
        ).fetchone()
        if not debt or debt["is_paid"] == 1 or debt.get("penalty_applied", 0) == 1:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_debt_penalty_already")}
        
        # Ambil data user dari koneksi yang sama
        u = conn.execute("SELECT id, gold, level FROM users WHERE id=?", (user_id,)).fetchone()
        if not u:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_debt_penalty_user_not_found")}
        
        level = u["level"]
        penalty = max(10, min(100, int(10 * (level / 5))))
        current_gold = u["gold"]
        new_gold = max(0, current_gold - penalty)
        
        # Update gold user dan tandai hutang (pakai koneksi yang sama)
        conn.execute("UPDATE users SET gold=? WHERE id=?", (new_gold, user_id))
        conn.execute("UPDATE debts SET penalty_applied=1, penalty_amount=? WHERE id=?", (penalty, debt_id))
        conn.commit()
    finally:
        conn.close()
    
    # Notifikasi & log (koneksi terpisah, setelah conn ditutup)
    log_activity(user_id, "debt_penalty", tr_db(user_id=user_id, key="log_debt_penalty", gold=penalty), 0, -penalty)
    add_notification(user_id, tr_db(user_id=user_id, key="db_debt_penalty_notif", name=debt['name'], days=days_late, gold=penalty), "danger")
    
    return {"ok": True, "gold_lost": penalty}

def get_overdue_debts_count(user_id):
    conn = get_conn()
    row = conn.execute("""
        SELECT COUNT(*) as cnt FROM debts 
        WHERE user_id=? AND is_paid=0 
        AND due_date IS NOT NULL AND due_date < date('now')
    """, (user_id,)).fetchone()
    conn.close()
    return row["cnt"] if row else 0

def get_total_unpaid_debt(user_id):
    conn = get_conn()
    row = conn.execute("SELECT COALESCE(SUM(amount),0) as total FROM debts WHERE user_id=? AND is_paid=0", (user_id,)).fetchone()
    conn.close()
    return row["total"] if row else 0.0


# ========== FOOD & NUTRITION TRACKER ==========

@retry_on_lock
def get_food_items(user_id, include_default=True):
    """Ambil semua makanan (default + custom milik user)."""
    conn = get_conn()
    if include_default:
        rows = conn.execute("""
            SELECT * FROM food_items 
            WHERE is_custom=0 OR (is_custom=1 AND user_id=?)
            ORDER BY name
        """, (user_id,)).fetchall()
    else:
        rows = conn.execute("""
            SELECT * FROM food_items 
            WHERE user_id=? AND is_custom=1
            ORDER BY name
        """, (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def add_custom_food(user_id, name, icon, calories, protein, carbs, fat):
    """Tambah makanan custom user."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    cur = conn.execute("""
        INSERT INTO food_items(user_id, name, icon, calories, protein, carbs, fat, is_custom)
        VALUES(?,?,?,?,?,?,?,1)
    """, (user_id, name, icon, calories, protein, carbs, fat))
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "id": new_id}

@retry_on_lock
def delete_food_item(user_id, food_id):
    """Hapus makanan custom (hanya milik sendiri)."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("DELETE FROM food_items WHERE id=? AND user_id=? AND is_custom=1", (food_id, user_id))
    conn.commit()
    conn.close()

@retry_on_lock
def log_food(user_id, food_id, serving, meal_type, log_date, notes=""):
    """Catat konsumsi makanan."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    food = conn.execute("SELECT calories, protein, carbs, fat FROM food_items WHERE id=?", (food_id,)).fetchone()
    if not food:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_food_not_found")}
    calories = food["calories"] * serving
    protein = food["protein"] * serving
    carbs = food["carbs"] * serving
    fat = food["fat"] * serving
    conn.execute("""
        INSERT INTO food_logs(user_id, food_id, serving, calories, protein, carbs, fat, meal_type, log_date, notes)
        VALUES(?,?,?,?,?,?,?,?,?,?)
    """, (user_id, food_id, serving, calories, protein, carbs, fat, meal_type, log_date, notes))
    conn.commit()
    conn.close()
    update_daily_net_calories(user_id, log_date)
    return {"ok": True, "calories": calories, "protein": protein, "carbs": carbs, "fat": fat}

@retry_on_lock
def get_food_logs(user_id, log_date=None):
    """Ambil catatan makanan untuk tanggal tertentu (default: hari ini)."""
    if log_date is None:
        from datetime import date
        log_date = date.today().isoformat()
    conn = get_conn()
    rows = conn.execute("""
        SELECT fl.*, fi.name, fi.icon 
        FROM food_logs fl
        JOIN food_items fi ON fl.food_id = fi.id
        WHERE fl.user_id=? AND fl.log_date=?
        ORDER BY fl.meal_type, fl.created_at
    """, (user_id, log_date)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def delete_food_log(user_id, log_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    row = conn.execute("SELECT log_date FROM food_logs WHERE id=? AND user_id=?", (log_id, user_id)).fetchone()
    if not row:
        conn.close()
        return
    log_date = row["log_date"]
    conn.execute("DELETE FROM food_logs WHERE id=? AND user_id=?", (log_id, user_id))
    conn.commit()
    conn.close()
    update_daily_net_calories(user_id, log_date)

@retry_on_lock
def get_nutrition_summary(user_id, log_date=None):
    """Hitung total kalori, protein, carbs, fat untuk suatu tanggal."""
    logs = get_food_logs(user_id, log_date)
    # Jika get_food_logs mengembalikan dict error (akun terkunci)
    if isinstance(logs, dict) and logs.get("ok") is False:
        return {"calories": 0, "protein": 0, "carbs": 0, "fat": 0}
    # Jika logs bukan list (misal None atau tipe lain), amankan juga
    if not isinstance(logs, list):
        return {"calories": 0, "protein": 0, "carbs": 0, "fat": 0}
    total = {"calories": 0, "protein": 0, "carbs": 0, "fat": 0}
    for log in logs:
        total["calories"] += log["calories"]
        total["protein"] += log["protein"]
        total["carbs"] += log["carbs"]
        total["fat"] += log["fat"]
    return total

@retry_on_lock
def get_nutrition_goals(user_id):
    """Ambil target nutrisi harian user."""
    conn = get_conn()
    row = conn.execute("SELECT * FROM user_nutrition_goals WHERE user_id=?", (user_id,)).fetchone()
    conn.close()
    if row:
        return dict(row)
    # Default goals
    return {"user_id": user_id, "daily_calories": 2000, "daily_protein": 50, 
            "daily_carbs": 250, "daily_fat": 70}

@retry_on_lock
def update_nutrition_goals(user_id, calories, protein, carbs, fat):
    """Update target nutrisi harian."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("""
        INSERT OR REPLACE INTO user_nutrition_goals(user_id, daily_calories, daily_protein, daily_carbs, daily_fat, updated_at)
        VALUES(?,?,?,?,?,datetime('now'))
    """, (user_id, calories, protein, carbs, fat))
    conn.commit()
    conn.close()
    return {"ok": True}

@retry_on_lock
def get_health_goals(user_id):
    conn = get_conn()
    row = conn.execute("SELECT * FROM user_health_goals WHERE user_id=?", (user_id,)).fetchone()
    conn.close()
    if row:
        return dict(row)
    return {"user_id": user_id, "daily_steps": 10000, "daily_sleep_hours": 7.0, "height_cm": 170, "weight_kg": 70}


@retry_on_lock
def update_health_goals(user_id, daily_steps, daily_sleep_hours, height_cm=None, weight_kg=None):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    current = get_health_goals(user_id)
    height_cm = height_cm if height_cm is not None else current.get("height_cm", 170)
    weight_kg = weight_kg if weight_kg is not None else current.get("weight_kg", 70)
    conn.execute("""
        INSERT OR REPLACE INTO user_health_goals(user_id, daily_steps, daily_sleep_hours, height_cm, weight_kg, updated_at)
        VALUES(?,?,?,?,?,datetime('now'))
    """, (user_id, daily_steps, daily_sleep_hours, height_cm, weight_kg))
    conn.commit()
    conn.close()
    return {"ok": True}

@retry_on_lock
def check_daily_nutrition_bonus(user_id, log_date=None):
    """Jika user mencapai minimal 90% target kalori, beri bonus XP 50 dan Gold 10 (sekali sehari)."""
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    if log_date is None:
        from datetime import date
        log_date = date.today().isoformat()
    
    # Step 1: Cek apakah sudah dapat bonus hari ini, lalu TUTUP koneksi
    conn = get_conn()
    try:
        existing = conn.execute(
            "SELECT id FROM food_achievements WHERE user_id=? AND achievement_date=? AND bonus_claimed=1",
            (user_id, log_date)
        ).fetchone()
    finally:
        conn.close()
    
    if existing:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_nutrition_bonus_claimed")}
    
    # Step 2: Ambil data nutrisi dan target (koneksi terpisah via fungsi masing-masing)
    summary = get_nutrition_summary(user_id, log_date)
    goals = get_nutrition_goals(user_id)
    if goals["daily_calories"] == 0:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_nutrition_goal_not_set")}
    
    percent = (summary["calories"] / goals["daily_calories"]) * 100
    if percent >= 90:
        xp_bonus = 50
        gold_bonus = 10
        # Step 3: Update user (via update_user yang sudah punya koneksi sendiri)
        user = get_user(user_id)
        new_xp = user["xp"] + xp_bonus
        new_gold = user["gold"] + gold_bonus
        new_lvl = user["level"]
        leveled = False
        needed = new_lvl * 150
        while new_xp >= needed:
            new_xp -= needed
            new_lvl += 1
            leveled = True
            needed = new_lvl * 150
        update_user(user_id, xp=new_xp, level=new_lvl, gold=new_gold,
                    total_xp_earned=user.get("total_xp_earned",0) + xp_bonus,
                    total_gold_earned=user.get("total_gold_earned",0.0) + gold_bonus)
        check_achievements(user_id, "calorie_goal", 1)
        # Step 4: Tandai sudah klaim (koneksi baru setelah semua write selesai)
        conn2 = get_conn()
        try:
            conn2.execute(
                "INSERT OR IGNORE INTO food_achievements(user_id, achievement_date, bonus_claimed) VALUES(?,?,1)",
                (user_id, log_date)
            )
            conn2.commit()
        finally:
            conn2.close()
        add_notification(user_id, tr_db(user_id=user_id, key="db_nutrition_bonus_notif", xp=xp_bonus, gold=gold_bonus), "success")
        return {"ok": True, "xp_gained": xp_bonus, "gold_gained": gold_bonus, "leveled_up": leveled}
    return {"ok": False, "msg": tr_db(user_id=user_id, key="db_nutrition_goal_not_reached", percent=percent)}

@retry_on_lock
def get_weekly_calories(user_id):
    """Ambil total kalori per hari untuk 7 hari terakhir (termasuk hari ini)."""
    conn = get_conn()
    rows = conn.execute("""
        SELECT log_date, SUM(calories) as total_calories
        FROM food_logs
        WHERE user_id=? AND log_date >= date('now', '-6 days')
        GROUP BY log_date
        ORDER BY log_date
    """, (user_id,)).fetchall()
    conn.close()
    # Buat dictionary untuk semua 7 hari
    from datetime import date, timedelta
    result = {}
    today = date.today()
    for i in range(7):
        d = (today - timedelta(days=6-i)).isoformat()
        result[d] = 0
    for row in rows:
        result[row["log_date"]] = row["total_calories"] or 0
    return result

@retry_on_lock
def get_water_goal(user_id):
    conn = get_conn()
    row = conn.execute("SELECT daily_ml FROM user_water_goals WHERE user_id=?", (user_id,)).fetchone()
    conn.close()
    return row["daily_ml"] if row else 2000

@retry_on_lock
def set_water_goal(user_id, daily_ml):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("""
        INSERT OR REPLACE INTO user_water_goals(user_id, daily_ml, updated_at)
        VALUES(?,?,datetime('now'))
    """, (user_id, daily_ml))
    conn.commit()
    conn.close()

@retry_on_lock
def add_water_log(user_id, amount_ml, log_date=None):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    if log_date is None:
        from datetime import date
        log_date = date.today().isoformat()
    conn = get_conn()
    from database import local_now
    # Total SEBELUM insert — water_goal hanya terhitung SEKALI saat target
    # baru tercapai hari itu (anti dobel-count di setiap log berikutnya).
    prev_total = get_water_total(user_id, log_date)
    conn.execute("""
        INSERT INTO water_logs(user_id, amount_ml, log_date, created_at)
        VALUES(?,?,?,?)
    """, (user_id, amount_ml, log_date, local_now().isoformat()))
    conn.commit()
    total = get_water_total(user_id, log_date)
    goal = get_water_goal(user_id)
    if prev_total < goal <= total:
        check_achievements(user_id, "water_goal", 1)
    conn.close()

@retry_on_lock
def get_water_total(user_id, log_date=None):
    if log_date is None:
        from datetime import date
        log_date = date.today().isoformat()
    conn = get_conn()
    row = conn.execute("SELECT COALESCE(SUM(amount_ml),0) as total FROM water_logs WHERE user_id=? AND log_date=?", (user_id, log_date)).fetchone()
    conn.close()
    return row["total"]

@retry_on_lock
def delete_water_log(user_id, log_id):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn = get_conn()
    conn.execute("DELETE FROM water_logs WHERE id=? AND user_id=?", (log_id, user_id))
    conn.commit()
    conn.close()

@retry_on_lock
def get_water_logs(user_id, log_date=None):
    if log_date is None:
        from datetime import date
        log_date = date.today().isoformat()
    conn = get_conn()
    rows = conn.execute("SELECT * FROM water_logs WHERE user_id=? AND log_date=? ORDER BY created_at DESC", (user_id, log_date)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def get_recipes(user_id):
    conn = get_conn()
    rows = conn.execute("SELECT * FROM recipes WHERE user_id=? ORDER BY name", (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def get_recipe_details(recipe_id):
    conn = get_conn()
    recipe = conn.execute("SELECT * FROM recipes WHERE id=?", (recipe_id,)).fetchone()
    if not recipe:
        conn.close()
        return None
    items = conn.execute("""
        SELECT ri.*, fi.name, fi.icon, fi.calories, fi.protein, fi.carbs, fi.fat
        FROM recipe_items ri
        JOIN food_items fi ON ri.food_id = fi.id
        WHERE ri.recipe_id=?
    """, (recipe_id,)).fetchall()
    conn.close()
    return {"recipe": dict(recipe), "items": [dict(i) for i in items]}

@retry_on_lock
def add_recipe(user_id, name, icon, serving_size, notes, food_items):
    """food_items: list of (food_id, quantity)"""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    cur = conn.execute("""
        INSERT INTO recipes(user_id, name, icon, serving_size, notes)
        VALUES(?,?,?,?,?)
    """, (user_id, name, icon, serving_size, notes))
    recipe_id = cur.lastrowid
    for food_id, qty in food_items:
        conn.execute("INSERT INTO recipe_items(recipe_id, food_id, quantity) VALUES(?,?,?)",
                     (recipe_id, food_id, qty))
    conn.commit()
    conn.close()
    return {"ok": True, "recipe_id": recipe_id}

@retry_on_lock
def delete_recipe(user_id, recipe_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("DELETE FROM recipes WHERE id=? AND user_id=?", (recipe_id, user_id))
    conn.commit()
    conn.close()

@retry_on_lock
def log_recipe(user_id, recipe_id, serving_multiplier, meal_type, log_date, notes=""):
    """Catat seluruh bahan resep ke food_logs sekaligus."""
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    details = get_recipe_details(recipe_id)
    if not details:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_recipe_not_found")}
    total_cal = 0
    total_pro = 0
    total_carb = 0
    total_fat = 0
    for item in details["items"]:
        qty = item["quantity"] * serving_multiplier
        total_cal += item["calories"] * qty
        total_pro += item["protein"] * qty
        total_carb += item["carbs"] * qty
        total_fat += item["fat"] * qty
    conn = get_conn()
    for item in details["items"]:
        qty = item["quantity"] * serving_multiplier
        conn.execute("""
            INSERT INTO food_logs(user_id, food_id, serving, calories, protein, carbs, fat, meal_type, log_date, notes)
            VALUES(?,?,?,?,?,?,?,?,?,?)
        """, (user_id, item["food_id"], qty, item["calories"]*qty, item["protein"]*qty, item["carbs"]*qty, item["fat"]*qty, meal_type, log_date, notes))
    conn.commit()
    conn.close()
    return {"ok": True, "calories": total_cal, "protein": total_pro, "carbs": total_carb, "fat": total_fat}

@retry_on_lock
def update_sport_calories(activity_id, user_id, calories_burned, duration_minutes):
    """Update data kalori terbakar untuk aktivitas sport."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("""
        UPDATE sport_activities 
        SET calories_burned=?, duration_minutes=?
        WHERE id=? AND user_id=?
    """, (calories_burned, duration_minutes, activity_id, user_id))
    conn.commit()
    conn.close()

@retry_on_lock
def get_total_calories_burned_today(user_id, log_date=None):
    """Total kalori terbakar dari sport hari ini."""
    if log_date is None:
        from datetime import date
        log_date = date.today().isoformat()
    conn = get_conn()
    row = conn.execute("""
        SELECT COALESCE(SUM(calories_burned), 0) as total
        FROM sport_activities
        WHERE user_id=? AND done_today=1 AND last_done=?
    """, (user_id, log_date)).fetchone()
    conn.close()
    return row["total"] if row else 0

def get_food_export_data(user_id, days=30):
    """Ambil data nutrisi dan air per hari untuk export (default 30 hari terakhir)."""
    from datetime import date, timedelta
    end_date = date.today()
    start_date = end_date - timedelta(days=days-1)
    data = []
    for i in range(days):
        d = start_date + timedelta(days=i)
        d_str = d.isoformat()
        nutri = get_nutrition_summary(user_id, d_str)
        water = get_water_total(user_id, d_str)
        burned = get_total_calories_burned_today(user_id, d_str)
        data.append({
            "date": d_str,
            "calories": nutri['calories'],
            "protein": nutri['protein'],
            "carbs": nutri['carbs'],
            "fat": nutri['fat'],
            "water_ml": water,
            "calories_burned": burned,
            "net_calories": nutri['calories'] - burned
        })
    return data

def get_food_summary_stats(user_id):
    """Ringkasan statistik nutrisi 30 hari terakhir."""
    from datetime import date, timedelta
    end_date = date.today()
    start_date = end_date - timedelta(days=29)
    total_calories = 0
    total_water = 0
    days_with_data = 0
    for i in range(30):
        d = (start_date + timedelta(days=i)).isoformat()
        nutri = get_nutrition_summary(user_id, d)
        water = get_water_total(user_id, d)
        if nutri['calories'] > 0 or water > 0:
            days_with_data += 1
        total_calories += nutri['calories']
        total_water += water
    return {
        "avg_calories": total_calories / 30,
        "total_calories_30d": total_calories,
        "avg_water": total_water / 30,
        "total_water_30d": total_water,
        "days_tracked": days_with_data
    }

# ========== HEALTH TRACKER (steps, sleep, etc.) ==========
@retry_on_lock
def add_health_log(user_id, log_date, steps=0, sleep_hours=0, water_ml=0, weight_kg=None, resting_hr=0, stress_level="normal", mood="normal", notes="", net_calories=None):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}

    try:
        # Hitung net_calories jika tidak diberikan
        if net_calories is None:
            food = conn.execute(
                "SELECT COALESCE(SUM(calories),0) as total FROM food_logs WHERE user_id=? AND log_date=?",
                (user_id, log_date)
            ).fetchone()
            total_food = food["total"] if food else 0
            sport = conn.execute(
                "SELECT COALESCE(SUM(calories_burned),0) as total FROM sport_activities "
                "WHERE user_id=? AND done_today=1 AND last_done=?",
                (user_id, log_date)
            ).fetchone()
            total_burned = sport["total"] if sport else 0
            net_calories = total_food - total_burned

        # Cek apakah baris sudah ada
        existing = conn.execute(
            "SELECT id, weight_kg FROM health_logs WHERE user_id=? AND log_date=?",
            (user_id, log_date)
        ).fetchone()

        if existing:
            # Jika weight_kg tidak diberikan (None), gunakan nilai yang sudah ada
            if weight_kg is None:
                weight_kg = existing["weight_kg"] if existing["weight_kg"] is not None else 0.0
            # Update hanya kolom yang diberikan
            conn.execute("""
                UPDATE health_logs 
                SET steps=?, sleep_hours=?, water_ml=?, weight_kg=?, resting_hr=?, 
                    stress_level=?, mood=?, notes=?, net_calories=?
                WHERE id=?
            """, (steps, sleep_hours, water_ml, weight_kg, resting_hr, stress_level, mood, notes, net_calories, existing["id"]))
        else:
            # Insert baru, default weight_kg 0 jika None
            if weight_kg is None:
                weight_kg = 0.0
            conn.execute("""
                INSERT INTO health_logs(user_id, log_date, steps, sleep_hours, water_ml, weight_kg, resting_hr, stress_level, mood, notes, net_calories)
                VALUES(?,?,?,?,?,?,?,?,?,?,?)
            """, (user_id, log_date, steps, sleep_hours, water_ml, weight_kg, resting_hr, stress_level, mood, notes, net_calories))
        conn.commit()
    finally:
        conn.close()
    # ── Tracking streak kesehatan (achievement Health Tracker / Health Nut) ──
    # Sebelumnya TIDAK ada yang memanggil event "health_streak" → achievement
    # kategori health tidak pernah berprogres. Dihitung dari hari-hari
    # berurutan yang punya log, berakhir di log_date.
    try:
        conn2 = get_conn()
        rows = conn2.execute(
            """SELECT DISTINCT log_date FROM health_logs
               WHERE user_id=? AND log_date<=? ORDER BY log_date DESC""",
            (user_id, log_date)
        ).fetchall()
        conn2.close()
        streak = 0
        expected = date.fromisoformat(log_date)
        for r in rows:
            d = date.fromisoformat(r["log_date"])
            if d == expected:
                streak += 1
                expected -= timedelta(days=1)
            elif d < expected:
                break
        if streak > 0:
            check_achievements(user_id, "health_streak", streak)
    except Exception as e:
        log.error(f"Health streak tracking gagal (non-fatal): {e}")

@retry_on_lock
def get_health_logs(user_id, days=7):
    """Ambil health logs untuk `days` terakhir (termasuk hari ini)"""
    from datetime import date, timedelta
    end_date = date.today()
    start_date = end_date - timedelta(days=days-1)
    conn = get_conn()
    rows = conn.execute("""
        SELECT * FROM health_logs
        WHERE user_id=? AND log_date BETWEEN ? AND ?
        ORDER BY log_date
    """, (user_id, start_date.isoformat(), end_date.isoformat())).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def update_daily_net_calories(user_id, log_date):
    """Hitung net kalori = total makanan - total kalori terbakar, lalu simpan ke health_logs.
       Hanya perbarui kolom net_calories, tanpa menghapus data lain.
    """
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    try:
        # Total kalori dari makanan
        food = conn.execute(
            "SELECT COALESCE(SUM(calories),0) as total FROM food_logs WHERE user_id=? AND log_date=?",
            (user_id, log_date)
        ).fetchone()
        total_food = food["total"] if food else 0

        # Total kalori terbakar dari sport (hanya yang sudah done_today)
        sport = conn.execute(
            "SELECT COALESCE(SUM(calories_burned),0) as total FROM sport_activities "
            "WHERE user_id=? AND done_today=1 AND last_done=?",
            (user_id, log_date)
        ).fetchone()
        total_burned = sport["total"] if sport else 0

        net_cal = total_food - total_burned

        # Cek apakah sudah ada baris untuk tanggal ini
        existing = conn.execute(
            "SELECT id FROM health_logs WHERE user_id=? AND log_date=?",
            (user_id, log_date)
        ).fetchone()

        if existing:
            # UPDATE, hanya kolom net_calories
            conn.execute(
                "UPDATE health_logs SET net_calories=? WHERE id=?",
                (net_cal, existing["id"])
            )
        else:
            # INSERT baru, kolom lain akan bernilai default (0/null)
            conn.execute(
                "INSERT INTO health_logs(user_id, log_date, net_calories) VALUES(?,?,?)",
                (user_id, log_date, net_cal)
            )
        conn.commit()
    finally:
        conn.close()

@retry_on_lock
def get_health_summary(user_id, days=30):
    """Rata-rata health metrics selama `days` hari terakhir, dengan avg water dari water_logs."""
    from datetime import date, timedelta
    start_date = date.today() - timedelta(days=days-1)
    conn = get_conn()

    # 1. Ambil rata-rata steps, sleep, weight, hr dari health_logs
    row = conn.execute("""
        SELECT AVG(steps) as avg_steps,
               AVG(sleep_hours) as avg_sleep,
               AVG(weight_kg) as avg_weight,
               AVG(resting_hr) as avg_hr,
               COUNT(*) as days_recorded
        FROM health_logs
        WHERE user_id=? AND log_date >= ?
    """, (user_id, start_date.isoformat())).fetchone()

    # 2. Ambil rata-rata air per hari dari water_logs (hanya hari yang ada catatan)
    water_row = conn.execute("""
        SELECT AVG(daily_total) as avg_water
        FROM (
            SELECT SUM(amount_ml) as daily_total
            FROM water_logs
            WHERE user_id=? AND log_date >= ?
            GROUP BY log_date
        )
    """, (user_id, start_date.isoformat())).fetchone()
    conn.close()

    avg_water = int(water_row['avg_water'] or 0) if water_row else 0

    if row and row['days_recorded']:
        return {
            'avg_steps': int(row['avg_steps'] or 0),
            'avg_sleep': round(row['avg_sleep'] or 0, 1),
            'avg_water': avg_water,
            'avg_weight': round(row['avg_weight'] or 0, 1),
            'avg_hr': int(row['avg_hr'] or 0),
            'days_recorded': row['days_recorded']
        }
    return {'avg_steps': 0, 'avg_sleep': 0, 'avg_water': 0, 'avg_weight': 0.0, 'avg_hr': 0, 'days_recorded': 0}

# ========== DEBT INSTALLMENT ==========
@retry_on_lock
def pay_debt_installment(debt_id, user_id, amount):
    """
    Membayar sebagian hutang.
    - amount: jumlah yang dibayar (akan dikurangi dari hutang)
    - Membuat transaksi expense di economy_items
    - Jika hutang lunas, tandai is_paid=1
    - Tidak mempengaruhi gold (hanya catatan ekonomi)
    """
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    # Step 1: Ambil data hutang
    conn = get_conn()
    debt = conn.execute(
        "SELECT id, name, amount, is_paid FROM debts WHERE id=? AND user_id=? AND is_paid=0",
        (debt_id, user_id)
    ).fetchone()
    if not debt:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_debt_not_found")}
    debt = dict(debt)
    conn.close()

    if amount <= 0:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_installment_zero")}
    if amount > debt['amount']:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_installment_exceed", amount=debt['amount'])}

    # Step 2: Catat transaksi expense di economy_items
    from datetime import date
    today = date.today().isoformat()
    expense_name = f"CICILAN HUTANG: {debt['name']} (Rp{amount:.0f})"
    add_economy_item(
        user_id,
        name=expense_name,
        icon="💸",
        type_="expense",
        amount=amount,
        category="Hutang",
        date_str=today,
        notes=f"Cicilan hutang {debt['name']} (ID {debt_id})"
    )

    # Step 3: Update sisa hutang
    remaining = debt['amount'] - amount
    conn2 = get_conn()
    if remaining <= 0.01:
        # Lunas
        conn2.execute(
            "UPDATE debts SET amount=0, is_paid=1, paid_at=? WHERE id=?",
            (datetime.now().isoformat(), debt_id)
        )
    else:
        conn2.execute(
            "UPDATE debts SET amount=? WHERE id=?",
            (remaining, debt_id)
        )
    conn2.commit()
    conn2.close()

    # Step 4: Log aktivitas
    log_activity(user_id, "debt_installment", tr_db(user_id=user_id, key="log_debt_installment", name=debt['name'], amount=amount), 0, 0)

    return {"ok": True, "remaining": remaining, "msg": tr_db(user_id=user_id, key="db_debt_installment", amount=amount, remaining=remaining)}

# ========== ACHIEVEMENT SYSTEM ==========
ACHIEVEMENTS_REBALANCED = [
    # ---- Level ----
    ("Pemula", "Mencapai level 5", "🎯", "level", "level_reach", 5, 150, 75),
    ("Prajurit", "Mencapai level 10", "⚔️", "level", "level_reach", 10, 300, 150),
    ("Ksatria", "Mencapai level 20", "🛡️", "level", "level_reach", 20, 600, 300),
    ("Legenda", "Mencapai level 50", "👑", "level", "level_reach", 50, 2000, 800),

    # ---- Habits ----
    ("Habit Starter", "Selesaikan 10 habit (positif)", "✅", "habit", "habit_complete", 10, 80, 30),
    ("Habit Enthusiast", "Selesaikan 50 habit", "🔥", "habit", "habit_complete", 50, 250, 100),
    ("Habit Master", "Selesaikan 200 habit", "💪", "habit", "habit_complete", 200, 600, 250),
    ("Streak Warrior", "Pertahankan streak habit 7 hari", "⚡", "habit", "habit_streak", 7, 150, 50),
    ("Streak Legend", "Pertahankan streak habit 30 hari", "🏅", "habit", "habit_streak", 30, 600, 200),

    # ---- Dailies ----
    ("Daily Doer", "Selesaikan 20 daily", "📅", "daily", "daily_complete", 20, 100, 40),
    ("Daily Devotee", "Selesaikan 100 daily", "🌟", "daily", "daily_complete", 100, 400, 150),

    # ---- Todos ----
    ("Quest Beginner", "Selesaikan 10 quest", "📜", "todo", "todo_complete", 10, 80, 30),
    ("Quest Champion", "Selesaikan 50 quest", "🏆", "todo", "todo_complete", 50, 300, 100),

    # ---- Sport ----
    ("Sport Rookie", "Kumpulkan 100 sport points", "🏃", "sport", "sport_points", 100, 150, 50),
    ("Sport Athlete", "Kumpulkan 500 sport points", "🏅", "sport", "sport_points", 500, 400, 150),
    ("Sport Legend", "Kumpulkan 2000 sport points", "🏆", "sport", "sport_points", 2000, 1000, 400),
    ("Sport Streak 7", "Sport streak 7 hari", "🔥", "sport", "sport_streak", 7, 150, 60),

    # ---- Economy ----
    ("Saver", "Kumpulkan total 1000 Gold", "💰", "economy", "total_gold", 1000, 150, 75),
    ("Rich Player", "Kumpulkan total 10000 Gold", "💎", "economy", "total_gold", 10000, 600, 300),
    ("Big Spender", "Belanjakan total 5000 Gold di shop", "🛒", "economy", "total_spent", 5000, 300, 100),

    # ---- Pets ----
    ("Pet Lover", "Adopsi 1 pet", "🐾", "pet", "pet_adopt", 1, 80, 30),
    ("Pet Collector", "Adopsi 3 pets", "🐕", "pet", "pet_adopt", 3, 200, 80),
    ("Pet Master", "Adopsi 5 pets", "🐉", "pet", "pet_adopt", 5, 400, 150),

    # ---- Guild & Boss ----
    ("Guild Member", "Bergabung ke guild", "👥", "guild", "join_guild", 1, 150, 75),
    ("Boss Slayer", "Mengalahkan 1 boss", "⚔️", "boss", "boss_kill", 1, 300, 100),
    ("Boss Hunter", "Mengalahkan 5 boss", "🏹", "boss", "boss_kill", 5, 600, 250),

    # ---- Social ----
    ("Friendly", "Miliki 1 teman", "👫", "social", "friend_count", 1, 80, 30),
    ("Popular", "Miliki 5 teman", "🎉", "social", "friend_count", 5, 200, 80),

    # ---- Health & Food ----
    ("Health Tracker", "Catat kesehatan 7 hari berturut-turut", "💚", "health", "health_streak", 7, 150, 60),
    ("Calorie Master", "Mencapai target kalori 10 kali", "🍎", "nutrition", "calorie_goal", 10, 200, 80),
    ("Hydration Hero", "Mencapai target air 20 kali", "💧", "nutrition", "water_goal", 20, 250, 100),

    # ---- Special ----
    ("Completionist", "Selesaikan 1000 tugas total (habit+daily+todo)", "🏆", "special", "total_tasks", 1000, 2500, 1000),
    ("All-Rounder", "Dapatkan minimal 1 achievement dari setiap kategori", "🌟", "special", "category_mastery", 1, 1500, 600),

    # ========== 10 ACHIEVEMENT BARU ==========
    ("Explorer", "Selesaikan 100 quest", "🗺️", "todo", "todo_complete", 100, 300, 100),
    ("Gym Rat", "Kumpulkan 1000 sport points", "🏋️", "sport", "sport_points", 1000, 400, 150),
    ("Millionaire", "Kumpulkan total 50000 Gold", "💰", "economy", "total_gold", 50000, 1000, 500),
    ("Shopaholic", "Belanjakan total 20000 Gold di shop", "🛍️", "economy", "total_spent", 20000, 500, 200),
    ("Pet Breeder", "Adopsi 10 pets", "🐶", "pet", "pet_adopt", 10, 500, 250),
    ("Boss Slayer Elite", "Mengalahkan 20 boss", "⚔️", "boss", "boss_kill", 20, 1000, 400),
    ("Social Butterfly", "Miliki 10 teman", "🦋", "social", "friend_count", 10, 300, 150),
    ("Health Nut", "Catat kesehatan 30 hari berturut-turut", "🥦", "health", "health_streak", 30, 500, 200),
    ("Calorie Crusher", "Mencapai target kalori 50 kali", "🔥", "nutrition", "calorie_goal", 50, 600, 250),
    ("Hydration Champion", "Mencapai target air 50 kali", "💪", "nutrition", "water_goal", 50, 600, 250),

    # ========== 12 ACHIEVEMENT BARU — Fokus, Crafting & tier lanjutan (v1.3.0) ==========
    # ---- Focus (Pomodoro) ----
    ("Focus Sprout", "Selesaikan 5 sesi pomodoro", "🍅", "focus", "pomodoro_sessions", 5, 80, 30),
    ("Focus Warrior", "Selesaikan 50 sesi pomodoro", "🧠", "focus", "pomodoro_sessions", 50, 300, 120),
    ("Time Bender", "Kumpulkan 1000 menit fokus total", "⏳", "focus", "pomodoro_minutes", 1000, 500, 200),
    # ---- Crafting ----
    ("Smith Apprentice", "Tempa 1 item di halaman Crafting", "🔨", "crafting", "craft_count", 1, 150, 60),
    ("Master Forger", "Tempa 3 item di halaman Crafting", "⚒️", "crafting", "craft_count", 3, 400, 180),
    # ---- Tier lanjutan ----
    ("Demigod", "Mencapai level 99", "🌌", "level", "level_reach", 99, 3000, 1200),
    ("Streak Immortal", "Pertahankan streak habit 100 hari", "🌟", "habit", "habit_streak", 100, 1500, 500),
    ("Daily Legend", "Selesaikan 250 daily", "📆", "daily", "daily_complete", 250, 600, 250),
    ("Quest Overlord", "Selesaikan 200 quest", "👑", "todo", "todo_complete", 200, 500, 200),
    ("Boss Terminator", "Mengalahkan 50 boss", "💀", "boss", "boss_kill", 50, 1800, 700),
    ("Gold Emperor", "Kumpulkan total 100000 Gold", "💰", "economy", "total_gold", 100000, 2000, 800),
    ("Task Machine", "Selesaikan 5000 tugas total (habit+daily+todo)", "🤖", "special", "total_tasks", 5000, 3000, 1500),
    # ========== 🆕 8 ACHIEVEMENT BARU v1.4.0 — tier lanjutan grindy ==========
    ("Habit God", "Selesaikan 500 habit", "🔱", "habit", "habit_complete", 500, 1200, 500),
    ("Daily God", "Selesaikan 500 daily", "📆", "daily", "daily_complete", 500, 1200, 500),
    ("Sport God", "Kumpulkan 5000 sport points", "🏆", "sport", "sport_points", 5000, 1500, 600),
    ("Gold Titan", "Kumpulkan total 200000 Gold", "👑", "economy", "total_gold", 200000, 3000, 1200),
    ("Pet God", "Adopsi 15 pets", "🐾", "pet", "pet_adopt", 15, 800, 400),
    ("Boss Emperor", "Mengalahkan 100 boss", "👹", "boss", "boss_kill", 100, 2500, 1000),
    ("Social King", "Miliki 20 teman", "👑", "social", "friend_count", 20, 600, 300),
    ("Focus Overlord", "Selesaikan 200 sesi pomodoro", "⏰", "focus", "pomodoro_sessions", 200, 800, 400),
    ("Craft Legend", "Tempa 10 item di halaman Crafting", "⚒️", "crafting", "craft_count", 10, 800, 350),
]

def init_achievements():
    """Insert data achievement default jika belum ada."""
    conn = get_conn()
    cur = conn.cursor()
    # Cek apakah sudah ada data
    cur.execute("SELECT COUNT(*) FROM achievements")
    if cur.fetchone()[0] > 0:
        conn.close()
        return
    
    achievements = [
        # ---- Level & XP ----
        ("Pemula", "Mencapai level 5", "🎯", "level", "level_reach", 5, 100, 50),
        ("Prajurit", "Mencapai level 10", "⚔️", "level", "level_reach", 10, 250, 100),
        ("Ksatria", "Mencapai level 20", "🛡️", "level", "level_reach", 20, 500, 200),
        ("Legenda", "Mencapai level 50", "👑", "level", "level_reach", 50, 1500, 500),
        # ---- Habits ----
        ("Habit Starter", "Selesaikan 10 habit (positif)", "✅", "habit", "habit_complete", 10, 50, 20),
        ("Habit Enthusiast", "Selesaikan 50 habit", "🔥", "habit", "habit_complete", 50, 200, 80),
        ("Habit Master", "Selesaikan 200 habit", "💪", "habit", "habit_complete", 200, 500, 200),
        ("Streak Warrior", "Pertahankan streak habit 7 hari", "⚡", "habit", "habit_streak", 7, 100, 30),
        ("Streak Legend", "Pertahankan streak habit 30 hari", "🏅", "habit", "habit_streak", 30, 500, 150),
        # ---- Dailies ----
        ("Daily Doer", "Selesaikan 20 daily", "📅", "daily", "daily_complete", 20, 80, 30),
        ("Daily Devotee", "Selesaikan 100 daily", "🌟", "daily", "daily_complete", 100, 300, 100),
        # ---- Todos ----
        ("Quest Beginner", "Selesaikan 10 quest", "📜", "todo", "todo_complete", 10, 60, 25),
        ("Quest Champion", "Selesaikan 50 quest", "🏆", "todo", "todo_complete", 50, 250, 80),
        # ---- Sport ----
        ("Sport Rookie", "Kumpulkan 100 sport points", "🏃", "sport", "sport_points", 100, 100, 30),
        ("Sport Athlete", "Kumpulkan 500 sport points", "🏅", "sport", "sport_points", 500, 300, 100),
        ("Sport Legend", "Kumpulkan 2000 sport points", "🏆", "sport", "sport_points", 2000, 800, 300),
        ("Sport Streak 7", "Sport streak 7 hari", "🔥", "sport", "sport_streak", 7, 100, 40),
        # ---- Economy ----
        ("Saver", "Kumpulkan total 1000 Gold", "💰", "economy", "total_gold", 1000, 100, 50),
        ("Rich Player", "Kumpulkan total 10000 Gold", "💎", "economy", "total_gold", 10000, 500, 200),
        ("Big Spender", "Belanjakan total 5000 Gold di shop", "🛒", "economy", "total_spent", 5000, 200, 80),
        # ---- Pets ----
        ("Pet Lover", "Adopsi 1 pet", "🐾", "pet", "pet_adopt", 1, 50, 20),
        ("Pet Collector", "Adopsi 3 pets", "🐕", "pet", "pet_adopt", 3, 150, 60),
        ("Pet Master", "Adopsi 5 pets", "🐉", "pet", "pet_adopt", 5, 300, 120),
        # ---- Guild & Boss ----
        ("Guild Member", "Bergabung ke guild", "👥", "guild", "join_guild", 1, 100, 50),
        ("Boss Slayer", "Mengalahkan 1 boss", "⚔️", "boss", "boss_kill", 1, 200, 80),
        ("Boss Hunter", "Mengalahkan 5 boss", "🏹", "boss", "boss_kill", 5, 500, 200),
        # ---- Social ----
        ("Friendly", "Miliki 1 teman", "👫", "social", "friend_count", 1, 50, 20),
        ("Popular", "Miliki 5 teman", "🎉", "social", "friend_count", 5, 150, 60),
        # ---- Health & Food ----
        ("Health Tracker", "Catat kesehatan 7 hari berturut-turut", "💚", "health", "health_streak", 7, 100, 40),
        ("Calorie Master", "Mencapai target kalori 10 kali", "🍎", "nutrition", "calorie_goal", 10, 150, 60),
        ("Hydration Hero", "Mencapai target air 20 kali", "💧", "nutrition", "water_goal", 20, 200, 80),
        # ---- Special ----
        ("Completionist", "Selesaikan 1000 tugas total (habit+daily+todo)", "🏆", "special", "total_tasks", 1000, 2000, 800),
        ("All-Rounder", "Dapatkan minimal 1 achievement dari setiap kategori", "🌟", "special", "category_mastery", 1, 1000, 400),
    ]
    
    for ach in achievements:
        cur.execute("""
            INSERT INTO achievements(name, description, icon, category, requirement_type, requirement_value, xp_reward, gold_reward)
            VALUES(?,?,?,?,?,?,?,?)
        """, ach)
    conn.commit()
    conn.close()
    log.info(f"[Achievements] {len(achievements)} achievements initialized.")

def migrate_achievements():
    """Update/rebalance achievement yang sudah ada dan tambah yang baru."""
    conn = get_conn()
    cur = conn.cursor()
    for ach in ACHIEVEMENTS_REBALANCED:
        name, desc, icon, category, req_type, req_val, xp, gold = ach
        # Cek apakah achievement sudah ada
        existing = cur.execute(
            "SELECT id, xp_reward, gold_reward, requirement_value FROM achievements WHERE name = ?",
            (name,)
        ).fetchone()
        if existing:
            # Update jika ada perubahan.
            # Khusus 'category_mastery': requirement_value disinkronkan DINAMIS
            # oleh _eval_category_mastery (= jumlah kategori aktual), jadi
            # migrasi tidak boleh menimpanya kembali ke nilai di list.
            is_mastery = (req_type == "category_mastery")
            req_changed = (existing["requirement_value"] != req_val) and not is_mastery
            reward_changed = (existing["xp_reward"] != xp or existing["gold_reward"] != gold)
            if is_mastery and reward_changed:
                cur.execute("""
                    UPDATE achievements
                    SET description=?, icon=?, category=?, requirement_type=?,
                        xp_reward=?, gold_reward=?
                    WHERE id=?
                """, (desc, icon, category, req_type, xp, gold, existing["id"]))
            elif (not is_mastery) and (req_changed or reward_changed):
                cur.execute("""
                    UPDATE achievements
                    SET description=?, icon=?, category=?, requirement_type=?,
                        requirement_value=?, xp_reward=?, gold_reward=?
                    WHERE id=?
                """, (desc, icon, category, req_type, req_val, xp, gold, existing["id"]))
        else:
            # Insert baru
            cur.execute("""
                INSERT INTO achievements(name, description, icon, category, requirement_type,
                                         requirement_value, xp_reward, gold_reward)
                VALUES(?,?,?,?,?,?,?,?)
            """, (name, desc, icon, category, req_type, req_val, xp, gold))
    conn.commit()
    conn.close()
    log.info(f"[Achievements] Migrated {len(ACHIEVEMENTS_REBALANCED)} achievements.")

def get_achievements_list():
    """Ambil semua data achievement."""
    conn = get_conn()
    rows = conn.execute("SELECT * FROM achievements ORDER BY category, requirement_value").fetchall()
    conn.close()
    return [dict(r) for r in rows]


def tr_achievement(ach_or_name, lang=None):
    """Pusat teks achievement → SELALU dari translations.py (bukan kolom DB).

    Return (name_text, desc_text) sesuai bahasa; kolom name/description di DB
    hanya dipakai sebagai fallback bila key terjemahan belum ada.
    """
    from translations import TRANSLATIONS
    if isinstance(ach_or_name, dict):
        raw_name = ach_or_name.get("name", "")
        raw_desc = ach_or_name.get("description", "")
    else:
        raw_name, raw_desc = str(ach_or_name), ""
    lang = lang or "id"
    idx = 0 if lang == "id" else 1
    nk, dk = f"ach_name_{raw_name}", f"ach_desc_{raw_name}"
    name_text = TRANSLATIONS.get(nk, (raw_name, raw_name))[idx]
    desc_text = TRANSLATIONS.get(dk, (raw_desc, raw_desc))[idx]
    return name_text, desc_text

def _eval_category_mastery(user_id):
    """Evaluasi achievement 'category_mastery' (mis. All-Rounder) setiap ada
    achievement baru yang ter-unlock.

    - requirement_value disinkronkan dinamis = jumlah kategori yang benar-benar
      ada (tidak termasuk baris mastery itu sendiri, agar tidak sirkular).
    - progress = jumlah kategori yang sudah punya >=1 unlock.
    - Aman dari rekursi: reward diberikan dengan skip_achievements=True.
    """
    conn = get_conn()
    mastery = [dict(r) for r in conn.execute(
        "SELECT * FROM achievements WHERE requirement_type='category_mastery'"
    ).fetchall()]
    if not mastery:
        conn.close()
        return
    total_cats = conn.execute(
        "SELECT COUNT(DISTINCT category) FROM achievements WHERE requirement_type!='category_mastery'"
    ).fetchone()[0]
    have_cats = conn.execute("""
        SELECT COUNT(DISTINCT a.category) FROM achievements a
        JOIN user_achievements ua ON ua.achievement_id = a.id AND ua.unlocked_at IS NOT NULL
        WHERE ua.user_id=? AND a.requirement_type!='category_mastery'
    """, (user_id,)).fetchone()[0]

    unlocked = []
    for m in mastery:
        if m["requirement_value"] != total_cats:
            conn.execute("UPDATE achievements SET requirement_value=? WHERE id=?",
                         (total_cats, m["id"]))
            m["requirement_value"] = total_cats
        cur = conn.execute(
            "SELECT progress, unlocked_at FROM user_achievements WHERE user_id=? AND achievement_id=?",
            (user_id, m["id"])
        ).fetchone()
        if cur:
            conn.execute("UPDATE user_achievements SET progress=? WHERE user_id=? AND achievement_id=?",
                         (have_cats, user_id, m["id"]))
            already = cur["unlocked_at"] is not None
        else:
            conn.execute("INSERT INTO user_achievements(user_id, achievement_id, progress) VALUES(?,?,?)",
                         (user_id, m["id"], have_cats))
            already = False
        if not already and have_cats >= total_cats:
            ts = datetime.now().isoformat()
            conn.execute("UPDATE user_achievements SET unlocked_at=?, progress=? WHERE user_id=? AND achievement_id=?",
                         (ts, total_cats, user_id, m["id"]))
            unlocked.append(m)
    conn.commit()
    conn.close()

    # Reward di luar koneksi (gain_xp_gold membuka koneksinya sendiri)
    for ach in unlocked:
        gain_xp_gold(user_id, ach["xp_reward"], ach["gold_reward"], skip_achievements=True)
        ach_name, _desc = tr_achievement(ach, get_user_language(user_id))
        add_notification(user_id, tr_db(user_id=user_id, key="db_achievement_unlocked", icon=ach['icon'], name=ach_name, xp=ach['xp_reward'], gold=ach['gold_reward']), "success")
        log_activity(user_id, "achievement", tr_db(user_id=user_id, key="log_achievement", name=ach_name), ach["xp_reward"], ach["gold_reward"])

def get_user_achievements(user_id):
    """Ambil data progress user untuk semua achievement."""
    conn = get_conn()
    rows = conn.execute("""
        SELECT a.*, 
               COALESCE(ua.progress, 0) as progress,
               ua.unlocked_at,
               ua.claimed
        FROM achievements a
        LEFT JOIN user_achievements ua ON a.id = ua.achievement_id AND ua.user_id = ?
        ORDER BY a.category, a.requirement_value
    """, (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def check_achievements(user_id, event_type, value=1, extra_data=None):
    """Dipanggil setiap kali ada aksi. Update progress achievement.
    Sekarang lebih aman: reward hanya sekali, progress tidak regresif.
    """
    conn = get_conn()
    # Ambil semua achievement yang requirement_type-nya sesuai
    rows = conn.execute("SELECT * FROM achievements WHERE requirement_type = ?", (event_type,)).fetchall()
    if not rows:
        conn.close()
        return []
    
    unlocked_list = []
    for ach in rows:
        ach_id = ach["id"]
        req_val = ach["requirement_value"]
        
        # Ambil progress saat ini
        cur = conn.execute(
            "SELECT progress, unlocked_at, claimed FROM user_achievements WHERE user_id=? AND achievement_id=?",
            (user_id, ach_id)
        ).fetchone()
        progress = cur["progress"] if cur else 0
        already_unlocked = cur and cur["unlocked_at"] is not None
        
        # Hitung progress baru berdasarkan tipe event
        new_progress = progress
        if event_type in ("habit_streak", "sport_streak", "health_streak", "level_reach"):
            # Untuk streak/level, kita simpan nilai maksimum yang pernah dicapai
            new_progress = max(progress, value)
        elif event_type in ("total_tasks", "total_gold", "total_spent",
                            "pomodoro_minutes", "craft_count"):
            # value adalah total absolut saat ini (bukan increment)
            new_progress = value
        else:
            # Default: increment
            new_progress = progress + value
        
        # Simpan progress (selalu update)
        if cur:
            conn.execute(
                "UPDATE user_achievements SET progress=? WHERE user_id=? AND achievement_id=?",
                (new_progress, user_id, ach_id)
            )
        else:
            conn.execute(
                "INSERT INTO user_achievements(user_id, achievement_id, progress) VALUES(?,?,?)",
                (user_id, ach_id, new_progress)
            )
        
        # Cek apakah baru tercapai (belum pernah unlock)
        if not already_unlocked and new_progress >= req_val:
            unlocked_time = datetime.now().isoformat()
            conn.execute(
                "UPDATE user_achievements SET unlocked_at=?, progress=? WHERE user_id=? AND achievement_id=?",
                (unlocked_time, req_val, user_id, ach_id)
            )
            unlocked_list.append(dict(ach))
    
    conn.commit()
    conn.close()
    
    for ach in unlocked_list:
        gain_xp_gold(user_id, ach["xp_reward"], ach["gold_reward"], skip_achievements=True)
        # Nama achievement SELALU dari translations.py (bukan teks mentah DB)
        ach_name, _desc = tr_achievement(ach, get_user_language(user_id))
        add_notification(user_id, tr_db(user_id=user_id, key="db_achievement_unlocked", icon=ach['icon'], name=ach_name, xp=ach['xp_reward'], gold=ach['gold_reward']), "success")
        log_activity(user_id, "achievement", tr_db(user_id=user_id, key="log_achievement", name=ach_name), ach["xp_reward"], ach["gold_reward"])

    # Setiap ada unlock baru → evaluasi ulang achievement 'category_mastery'
    # (mis. All-Rounder): sebelumnya tipe event ini tidak punya call site.
    if unlocked_list:
        _eval_category_mastery(user_id)
    return unlocked_list

def claim_achievement_reward(user_id, achievement_id):
    """Klaim reward jika achievement sudah unlocked tapi belum diklaim (opsional)."""
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn = get_conn()
    ua = conn.execute(
        "SELECT claimed, unlocked_at FROM user_achievements WHERE user_id=? AND achievement_id=? AND unlocked_at IS NOT NULL",
        (user_id, achievement_id)
    ).fetchone()
    if not ua or ua["claimed"]:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_achievement_reward_unavailable")}
    
    ach = conn.execute("SELECT xp_reward, gold_reward FROM achievements WHERE id=?", (achievement_id,)).fetchone()
    if not ach:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_achievement_not_found")}
    
    conn.execute("UPDATE user_achievements SET claimed=1 WHERE user_id=? AND achievement_id=?", (user_id, achievement_id))
    conn.commit()
    conn.close()
    
    gain_xp_gold(user_id, ach["xp_reward"], ach["gold_reward"])
    add_notification(user_id, tr_db(user_id=user_id, key="db_achievement_claimed", xp=ach['xp_reward'], gold=ach['gold_reward']), "info")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_achievement_reward_claimed")}

def update_total_tasks(user_id):
    conn = get_conn()
    conn.execute("UPDATE users SET total_tasks_completed = total_tasks_completed + 1 WHERE id=?", (user_id,))
    new_total = conn.execute("SELECT total_tasks_completed FROM users WHERE id=?", (user_id,)).fetchone()[0]
    conn.commit()
    conn.close()
    check_achievements(user_id, "total_tasks", new_total)

def update_total_spent(user_id, amount):
    conn = get_conn()
    conn.execute("UPDATE users SET total_gold_spent = COALESCE(total_gold_spent,0) + ? WHERE id=?", (amount, user_id))
    new_total = conn.execute("SELECT total_gold_spent FROM users WHERE id=?", (user_id,)).fetchone()[0]
    conn.commit()
    conn.close()
    check_achievements(user_id, "total_spent", new_total)

def update_total_gold_earned(user_id, amount):
    conn = get_conn()
    conn.execute("UPDATE users SET total_gold_earned = total_gold_earned + ? WHERE id=?", (amount, user_id))
    new_total = conn.execute("SELECT total_gold_earned FROM users WHERE id=?", (user_id,)).fetchone()[0]
    conn.commit()
    conn.close()
    check_achievements(user_id, "total_gold", new_total)

def delete_account(user_id, password):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    u = get_user(user_id)
    if not u or not _verify_password(password, u.get("password_hash", "")):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_delete_account_wrong_password")}
    
    conn = get_conn()
    try:
        # Matikan foreign key sementara agar tidak ada constraint error
        conn.execute("PRAGMA foreign_keys = OFF")

        # 1. Hapus semua data yang terkait dengan user_id
        conn.execute("DELETE FROM user_sessions WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM user_pets WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM inventory WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM economy_items WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM debts WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM savings WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM food_logs WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM water_logs WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM health_logs WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM habits WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM dailies WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM todos WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM sport_activities WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM task_folders WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM notifications WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM activity_log WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM backup_codes WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM messages WHERE sender_id = ? OR receiver_id = ?", (user_id, user_id))
        conn.execute("DELETE FROM friends WHERE user_id = ? OR friend_id = ?", (user_id, user_id))
        conn.execute("DELETE FROM guild_members WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM guild_invites WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM guild_requests WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM guild_messages WHERE sender_id = ?", (user_id,))
        conn.execute("DELETE FROM boss_rewards WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM user_achievements WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM user_nutrition_goals WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM food_achievements WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM user_water_goals WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM user_health_goals WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM recipes WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM guild_leader_transfers WHERE old_leader_id = ?", (user_id,))
        
        # 2. Handle jika user adalah leader guild
        guilds = conn.execute("SELECT id FROM guilds WHERE leader_id = ?", (user_id,)).fetchall()
        for guild in guilds:
            other = conn.execute("SELECT user_id FROM guild_members WHERE guild_id = ? LIMIT 1", (guild["id"],)).fetchone()
            if other:
                conn.execute("UPDATE guilds SET leader_id = ? WHERE id = ?", (other["user_id"], guild["id"]))
            else:
                conn.execute("DELETE FROM guilds WHERE id = ?", (guild["id"],))
        
        # 3. Terakhir, hapus user
        conn.execute("DELETE FROM users WHERE id = ?", (user_id,))
        
        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.execute("PRAGMA foreign_keys = ON")
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_delete_account_failed", error=str(e))}
    finally:
        conn.execute("PRAGMA foreign_keys = ON")
        conn.close()
    
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_delete_account_success")}

# ── INVESTING ──────────────────────────────────────────────────────────────
@retry_on_lock
def add_investment(user_id, name, icon, amount, notes=''):
    """Tambahkan kartu investasi baru. Amount diambil dari saldo ekonomi."""
    # Validasi saldo (maksimal 10% dari saldo)
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    try:
        # Hitung saldo ekonomi saat ini
        rows = conn.execute("SELECT type, amount FROM economy_items WHERE user_id=?", (user_id,)).fetchall()
        total_income = sum(r['amount'] for r in rows if r['type'] == 'income')
        total_expense = sum(r['amount'] for r in rows if r['type'] == 'expense')
        balance = total_income - total_expense
        if amount <= 0 or amount > balance:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_investment_exceeds_balance", balance=balance)}
        if amount > balance * 0.1:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_investment_max_10", max=balance*0.1)}
    finally:
        conn.close()

    # Kurangi saldo melalui economy_items (buat expense)
    from datetime import date
    add_economy_item(user_id, name=f"Investasi: {name}", icon=icon,
                     type_="expense", amount=amount, category="Investasi",
                     date_str=date.today().isoformat(), notes=notes)
    # Tambahkan ke tabel investments
    conn2 = get_conn()
    cur = conn2.execute("""
        INSERT INTO investments(user_id, name, icon, amount, notes)
        VALUES(?,?,?,?,?)
    """, (user_id, name, icon, amount, notes))
    inv_id = cur.lastrowid
    conn2.commit()
    conn2.close()
    log_activity(user_id, "invest", tr_db(user_id=user_id, key="log_invest", name=name, amount=amount), 0, -amount)
    return {"ok": True, "invest_id": inv_id, "msg": tr_db(user_id=user_id, key="db_investment_success", name=name, amount=amount)}

@retry_on_lock
def add_investment_return(invest_id, user_id, amount):
    """Tambah nilai investasi secara manual sebesar amount (IDR)."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    inv = conn.execute("SELECT amount FROM investments WHERE id=? AND user_id=? AND is_active=1", (invest_id, user_id)).fetchone()
    if not inv:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_investment_not_found")}
    new_amount = inv['amount'] + amount
    conn.execute("UPDATE investments SET amount=?, last_update=datetime('now') WHERE id=?", (new_amount, invest_id))
    conn.commit()
    conn.close()
    log_activity(user_id, "invest_return", tr_db(user_id=user_id, key="log_invest_return", amount=amount), 0, 0)
    return {"ok": True, "gain": amount, "new_amount": new_amount}

@retry_on_lock
def get_investments(user_id):
    conn = get_conn()
    rows = conn.execute("SELECT * FROM investments WHERE user_id=? AND is_active=1 ORDER BY invested_date DESC", (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def delete_investment(invest_id, user_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("UPDATE investments SET is_active=0 WHERE id=? AND user_id=?", (invest_id, user_id))
    conn.commit()
    conn.close()

@retry_on_lock
def update_investment(invest_id, user_id, **kwargs):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    if not kwargs:
        return
    fields = ", ".join(f"{k}=?" for k in kwargs)
    conn = get_conn()
    conn.execute(f"UPDATE investments SET {fields}, last_update=datetime('now') WHERE id=? AND user_id=?", 
                 list(kwargs.values()) + [invest_id, user_id])
    conn.commit()
    conn.close()

@retry_on_lock
def collect_investment_return(invest_id, user_id, percent=5):
    """Tambah nilai investasi secara manual sebesar persen tertentu (default 5% dari amount saat ini)."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    inv = conn.execute("SELECT amount FROM investments WHERE id=? AND user_id=? AND is_active=1", (invest_id, user_id)).fetchone()
    if not inv:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_investment_not_found")}
    gain = inv['amount'] * percent / 100
    if gain <= 0:
        gain = 100  # minimal gain 100 jika amount kecil
    new_amount = inv['amount'] + gain
    conn.execute("UPDATE investments SET amount=?, last_update=datetime('now') WHERE id=?", (new_amount, invest_id))
    conn.commit()
    conn.close()
    log_activity(user_id, "invest_return", f"Return investasi +{gain:.0f}", 0, 0)
    return {"ok": True, "gain": gain, "new_amount": new_amount}

@retry_on_lock
def withdraw_investment(invest_id, user_id):
    """Tarik semua dana investasi, masukkan ke income, lalu nonaktifkan kartu."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    inv = conn.execute("SELECT name, amount FROM investments WHERE id=? AND user_id=? AND is_active=1", (invest_id, user_id)).fetchone()
    if not inv:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_investment_not_found")}
    amount = inv['amount']
    if amount <= 0:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_investment_no_funds")}
    # Catat sebagai income
    from datetime import date
    add_economy_item(user_id, name=f"Penarikan Investasi: {inv['name']}", icon="📈",
                     type_="income", amount=amount, category="Investasi",
                     date_str=date.today().isoformat(), notes="Penarikan dana investasi")
    # Nonaktifkan investasi
    conn.execute("UPDATE investments SET is_active=0 WHERE id=?", (invest_id,))
    conn.commit()
    conn.close()
    log_activity(user_id, "invest_withdraw", tr_db(user_id=user_id, key="log_invest_withdraw", name=inv['name'], amount=amount), 0, amount)
    return {"ok": True, "amount": amount, "msg": tr_db(user_id=user_id, key="db_investment_withdrawn", amount=amount)}

# ── SUBSCRIPTION ───────────────────────────────────────────────────────────
@retry_on_lock
def add_subscription(user_id, name, icon, amount, due_date, period, is_recurring, notes=''):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    cur = conn.execute("""
        INSERT INTO subscriptions(user_id, name, icon, amount, due_date, period, is_recurring, notes)
        VALUES(?,?,?,?,?,?,?,?)
    """, (user_id, name, icon, amount, due_date, period, 1 if is_recurring else 0, notes))
    sub_id = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "subscription_id": sub_id}

@retry_on_lock
def get_subscriptions(user_id, active_only=True):
    conn = get_conn()
    if active_only:
        rows = conn.execute("SELECT * FROM subscriptions WHERE user_id=? AND is_active=1 ORDER BY due_date ASC", (user_id,)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM subscriptions WHERE user_id=? ORDER BY due_date ASC", (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def delete_subscription(sub_id, user_id):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("UPDATE subscriptions SET is_active=0 WHERE id=? AND user_id=?", (sub_id, user_id))
    conn.commit()
    conn.close()

@retry_on_lock
def update_subscription(sub_id, user_id, **kwargs):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    if not kwargs:
        return
    fields = ", ".join(f"{k}=?" for k in kwargs)
    conn = get_conn()
    conn.execute(f"UPDATE subscriptions SET {fields} WHERE id=? AND user_id=?", list(kwargs.values()) + [sub_id, user_id])
    conn.commit()
    conn.close()

@retry_on_lock
def renew_subscription(sub_id, user_id, auto_pay=True):
    """Perpanjang subscription: untuk non-recurring (manual) atau recurring (auto charge)."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    sub = conn.execute("SELECT name, amount, due_date, period, is_recurring FROM subscriptions WHERE id=? AND user_id=? AND is_active=1", (sub_id, user_id)).fetchone()
    if not sub:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_subscription_not_found")}
    sub = dict(sub)
    today = date.today().isoformat()
    if auto_pay and sub['is_recurring']:
        # Cek saldo ekonomi
        rows = conn.execute("SELECT type, amount FROM economy_items WHERE user_id=?", (user_id,)).fetchall()
        total_income = sum(r['amount'] for r in rows if r['type'] == 'income')
        total_expense = sum(r['amount'] for r in rows if r['type'] == 'expense')
        balance = total_income - total_expense
        if balance < sub['amount']:
            conn.close()
            add_notification(user_id, tr_db(user_id=user_id, key="db_subscription_insufficient_notif", name=sub['name'], amount=sub['amount']), "danger")
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_subscription_insufficient_balance")}
        # Catat expense
        add_economy_item(user_id, name=f"Langganan: {sub['name']}", icon="📅",
                         type_="expense", amount=sub['amount'], category="Subscription",
                         date_str=today, notes=f"Perpanjangan otomatis")
        # Update due_date
        new_due = date.today()
        if sub['period'] == 'monthly':
            from dateutil.relativedelta import relativedelta
            new_due += relativedelta(months=1)
        elif sub['period'] == 'yearly':
            new_due += relativedelta(years=1)
        else:
            new_due = None  # one-time tidak auto renew
        if new_due:
            conn.execute("UPDATE subscriptions SET due_date=?, last_charged=? WHERE id=?", (new_due.isoformat(), today, sub_id))
        else:
            # one-time: nonaktifkan setelah renew (manual)
            conn.execute("UPDATE subscriptions SET is_active=0, last_charged=? WHERE id=?", (today, sub_id))
            add_notification(user_id, tr_db(user_id=user_id, key="db_subscription_renew_warning", name=sub['name']), "warning")
        conn.commit()
        conn.close()
        log_activity(user_id, "subscription", tr_db(user_id=user_id, key="log_subscription_auto", name=sub['name'], amount=sub['amount']), 0, -sub['amount'])
        add_notification(user_id, tr_db(user_id=user_id, key="db_subscription_renewed_auto_notif", name=sub['name'], amount=sub['amount']), "success")
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_subscription_renewed_auto")}
    else:
        # Renew manual (pengguna klik tombol Renew)
        if not auto_pay:
            # Cek saldo dulu
            rows = conn.execute("SELECT type, amount FROM economy_items WHERE user_id=?", (user_id,)).fetchall()
            total_income = sum(r['amount'] for r in rows if r['type'] == 'income')
            total_expense = sum(r['amount'] for r in rows if r['type'] == 'expense')
            balance = total_income - total_expense
            if balance < sub['amount']:
                conn.close()
                return {"ok": False, "msg": tr_db(user_id=user_id, key="db_subscription_insufficient_balance")}
            add_economy_item(user_id, name=f"Langganan: {sub['name']}", icon="📅",
                             type_="expense", amount=sub['amount'], category="Subscription",
                             date_str=today, notes="Perpanjangan manual")
        # Update due_date
        new_due = date.today()
        if sub['period'] == 'monthly':
            from dateutil.relativedelta import relativedelta
            new_due += relativedelta(months=1)
        elif sub['period'] == 'yearly':
            new_due += relativedelta(years=1)
        else:
            new_due = None
        if new_due:
            conn.execute("UPDATE subscriptions SET due_date=?, last_charged=?, is_active=1 WHERE id=?", (new_due.isoformat(), today, sub_id))
        else:
            conn.execute("UPDATE subscriptions SET last_charged=?, is_active=1 WHERE id=?", (today, sub_id))
        conn.commit()
        conn.close()
        log_activity(user_id, "subscription", tr_db(user_id=user_id, key="log_subscription_manual", name=sub['name'], amount=sub['amount']), 0, -sub['amount'])
        add_notification(user_id, tr_db(user_id=user_id, key="db_subscription_renew_manual_success", name=sub['name'], date=new_due), "success")
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_subscription_renewed_manual")}

def check_all_subscriptions(user_id):
    """Cek semua subscription user yang melewati due_date. Untuk recurring auto-charge otomatis."""
    subs = get_subscriptions(user_id, active_only=True)
    today = date.today().isoformat()
    for sub in subs:
        if sub['due_date'] < today:
            if sub['is_recurring']:
                renew_subscription(sub['id'], user_id, auto_pay=True)
            else:
                add_notification(user_id, tr_db(user_id=user_id, key="db_subscription_expired_notif", name=sub['name'], date=sub['due_date']), "warning")

# ========== CURRENCY SETTINGS ========== #
def get_user_currency(user_id):
    conn = get_conn()
    row = conn.execute("SELECT currency FROM users WHERE id=?", (user_id,)).fetchone()
    conn.close()
    return row["currency"] if row else "IDR"

def set_user_currency(user_id, currency):
    conn = get_conn()
    conn.execute("UPDATE users SET currency=? WHERE id=?", (currency, user_id))
    conn.commit()
    conn.close()

# Currency conversion rates (hardcoded, bisa diupdate nanti)
CURRENCY_RATES = {"IDR": 1, "USD": 17800, "EUR": 20700}

def convert_to_idr(amount, from_currency):
    """Konversi dari mata uang user ke IDR"""
    rate = CURRENCY_RATES.get(from_currency, 1)
    return amount * rate

def convert_from_idr(amount_idr, to_currency):
    """Konversi dari IDR ke mata uang user"""
    rate = CURRENCY_RATES.get(to_currency, 1)
    return amount_idr / rate

def get_user_bmi_settings(user_id):
    """Ambil data BMI user (tinggi, berat, usia, gender, aktivitas)."""
    conn = get_conn()
    row = conn.execute(
        "SELECT height_cm, weight_kg, age, gender, activity_factor FROM user_health_goals WHERE user_id=?",
        (user_id,)
    ).fetchone()
    conn.close()
    if row:
        return dict(row)
    return {
        "height_cm": 170,
        "weight_kg": 70,
        "age": 25,
        "gender": "male",
        "activity_factor": 1.55
    }

def update_user_bmi_settings(user_id, height_cm, weight_kg, age, gender, activity_factor):
    """Simpan data BMI user."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("""
        INSERT INTO user_health_goals(user_id, height_cm, weight_kg, age, gender, activity_factor, updated_at)
        VALUES(?,?,?,?,?,?,datetime('now'))
        ON CONFLICT(user_id) DO UPDATE SET
            height_cm=excluded.height_cm,
            weight_kg=excluded.weight_kg,
            age=excluded.age,
            gender=excluded.gender,
            activity_factor=excluded.activity_factor,
            updated_at=datetime('now')
    """, (user_id, height_cm, weight_kg, age, gender, activity_factor))
    conn.commit()
    conn.close()
    return {"ok": True}

@retry_on_lock
def log_user_weight(user_id, weight_kg, log_date=None):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    """Catat berat badan user ke health_logs untuk tanggal tertentu (default hari ini)."""
    if log_date is None:
        from datetime import date
        log_date = date.today().isoformat()
    
    conn = get_conn()
    try:
        # Cek apakah sudah ada entry untuk tanggal tersebut
        existing = conn.execute(
            "SELECT id FROM health_logs WHERE user_id=? AND log_date=?",
            (user_id, log_date)
        ).fetchone()
        
        if existing:
            # Update weight_kg pada entry yang sudah ada
            conn.execute(
                "UPDATE health_logs SET weight_kg=? WHERE id=?",
                (weight_kg, existing["id"])
            )
        else:
            # Buat entry baru dengan nilai default lainnya
            conn.execute("""
                INSERT INTO health_logs(user_id, log_date, weight_kg, steps, sleep_hours, water_ml, resting_hr, stress_level, mood, net_calories)
                VALUES(?,?,?,0,0,0,0,'normal','normal',0)
            """, (user_id, log_date, weight_kg))
        conn.commit()
    finally:
        conn.close()

def get_active_pets_info(user_id):
    """Return list of active pets with their details."""
    conn = get_conn()
    rows = conn.execute("""
        SELECT up.pet_id, up.level, up.is_active, p.name, p.icon, p.base_buff
        FROM user_pets up
        JOIN pets_data p ON up.pet_id = p.id
        WHERE up.user_id=? AND up.is_active=1
    """, (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_active_pets(user_id):
    conn = get_conn()
    rows = conn.execute("SELECT pet_id, level FROM user_pets WHERE user_id=? AND is_active=1", (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def redeem_admin_code(user_id, code):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    """Untuk kompatibilitas, gunakan redeem_code."""
    return redeem_code(user_id, code)

def add_redeem_code(code, reward_type, reward_value=0, reward_item=None, is_one_time=1):
    """Tambahkan kode redeem baru."""
    conn = get_conn()
    try:
        conn.execute(
            "INSERT INTO redeem_codes(code, reward_type, reward_value, reward_item, is_one_time) VALUES(?,?,?,?,?)",
            (code, reward_type, reward_value, reward_item, is_one_time)
        )
        conn.commit()
        return {"ok": True}
    except sqlite3.IntegrityError:
        return {"ok": False, "msg": tr_db(lang="id", key="db_redeem_code_exists")}
    finally:
        conn.close()

def redeem_code(user_id, code):
    """Redeem kode dan berikan hadiah."""
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    # Cek apakah kode valid
    row = conn.execute(
        "SELECT * FROM redeem_codes WHERE code = ? AND (used_by IS NULL OR is_one_time = 0)",
        (code.upper().strip(),)
    ).fetchone()
    if not row:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_redeem_invalid")}
    
    code_data = dict(row)
    
    # Jika one-time, tandai sudah dipakai
    if code_data["is_one_time"]:
        conn.execute(
            "UPDATE redeem_codes SET used_by = ?, used_at = ? WHERE id = ?",
            (user_id, datetime.now().isoformat(), code_data["id"])
        )
        conn.commit()
    conn.close()
    
    # Proses hadiah berdasarkan reward_type
    reward_type = code_data["reward_type"]
    reward_value = code_data["reward_value"]
    reward_item = code_data["reward_item"]
    
    if reward_type == "admin":
        # Ubah user menjadi admin (hanya bisa sekali seumur hidup)
        u = get_user(user_id)
        if u.get("is_admin", 0):
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_redeem_admin_already")}
        set_user_admin(user_id, True)
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_redeem_admin_success")}
    
    elif reward_type == "xp":
        gain_xp_gold(user_id, reward_value, 0)
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_redeem_xp", xp=reward_value), "xp": reward_value}
    
    elif reward_type == "gold":
        gain_xp_gold(user_id, 0, reward_value)
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_redeem_gold", gold=reward_value), "gold": reward_value}
    
    elif reward_type == "item":
        # Tambahkan item ke inventory
        item = SHOP_ITEMS.get(reward_item)
        if not item:
            return {"ok": False, "msg": tr_db(user_id=user_id, key="db_redeem_item_fail")}
        conn2 = get_conn()
        ex = conn2.execute(
            "SELECT * FROM inventory WHERE user_id=? AND item_id=?", (user_id, reward_item)
        ).fetchone()
        if ex:
            conn2.execute("UPDATE inventory SET quantity=quantity+1 WHERE id=?", (ex["id"],))
        else:
            conn2.execute(
                "INSERT INTO inventory(user_id, item_id, item_type) VALUES(?,?,?)",
                (user_id, reward_item, item["type"])
            )
        conn2.commit()
        conn2.close()
        recalculate_all_buffs(user_id)
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_redeem_item_success", name=item['name'])}
    
    else:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_redeem_unknown")}
    
def set_user_admin(user_id, is_admin):
    conn = get_conn()
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    conn.execute("UPDATE users SET is_admin = ? WHERE id = ?", (1 if is_admin else 0, user_id))
    conn.commit()
    conn.close()

# ========== ADMIN CHEAT FUNCTIONS FOR PETS ==========
@retry_on_lock
def admin_level_up_all_pets(user_id):
    """Naikkan level semua pet milik user sebanyak 1 level (tanpa mengubah EXP)."""
    conn = get_conn()
    try:
        pets = conn.execute("SELECT id, level, exp FROM user_pets WHERE user_id=?", (user_id,)).fetchall()
        for pet in pets:
            new_level = pet["level"] + 1
            # EXP tetap, tidak perlu diubah
            conn.execute("UPDATE user_pets SET level=? WHERE id=?", (new_level, pet["id"]))
        conn.commit()
        # Recalculate buff setelah perubahan
        recalculate_all_buffs(user_id)
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_admin_cheat_pets_level_up")}
    except Exception as e:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_general_error", error=str(e))}
    finally:
        conn.close()

@retry_on_lock
def admin_add_exp_all_pets(user_id, amount):
    """Tambahkan EXP ke semua pet (amount dalam integer)."""
    conn = get_conn()
    leveled_any = False
    try:
        pets = conn.execute("SELECT id, exp, level FROM user_pets WHERE user_id=?", (user_id,)).fetchall()
        for pet in pets:
            new_exp = pet["exp"] + amount
            new_level = pet["level"]
            needed = new_level * 100
            while new_exp >= needed:
                new_exp -= needed
                new_level += 1
                needed = new_level * 100
                leveled_any = True
            conn.execute("UPDATE user_pets SET exp=?, level=? WHERE id=?", (new_exp, new_level, pet["id"]))
        conn.commit()
        if leveled_any:
            recalculate_all_buffs(user_id)
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_admin_cheat_pets_add_exp", amount=amount)}
    except Exception as e:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_general_error", error=str(e))}
    finally:
        conn.close()

@retry_on_lock
def admin_feed_all_pets(user_id):
    """Isi hunger semua pet menjadi 100 (kenyang)."""
    conn = get_conn()
    try:
        conn.execute("UPDATE user_pets SET hunger=100, last_fed=? WHERE user_id=?", (local_now().isoformat(), user_id))
        conn.commit()
        return {"ok": True, "msg": tr_db(user_id=user_id, key="db_admin_cheat_pets_feed")}
    except Exception as e:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_general_error", error=str(e))}
    finally:
        conn.close()

def get_user_language(user_id):
    conn = get_conn()
    row = conn.execute("SELECT language FROM users WHERE id=?", (user_id,)).fetchone()
    conn.close()
    return row["language"] if row and row["language"] else "id"

def set_user_language(user_id, lang_code):
    conn = get_conn()
    conn.execute("UPDATE users SET language=? WHERE id=?", (lang_code, user_id))
    conn.commit()
    conn.close()

# ── TASK HISTORY ──────────────────────────────────────────────────────────
@retry_on_lock
def log_task_history(user_id, task_type, task_id, action, action_date=None):
    if action_date is None:
        from datetime import date
        action_date = date.today().isoformat()
    conn = get_conn()
    try:
        conn.execute(
            "DELETE FROM task_history WHERE user_id=? AND task_type=? AND task_id=? AND action_date=?",
            (user_id, task_type, task_id, action_date)
        )
        conn.execute(
            "INSERT INTO task_history(user_id, task_type, task_id, action_date, action) VALUES(?,?,?,?,?)",
            (user_id, task_type, task_id, action_date, action)
        )
        conn.commit()
    finally:
        conn.close()
    if action == "success":
        cloud_type = "quest" if task_type == "todo" else task_type
        enqueue_productivity_event(
            user_id, cloud_type, task_id,
            f"{cloud_type}:{task_id}:{action_date}", local_now().isoformat(),
            {"action_date": action_date},
        )

def fill_skipped_history(user_id):
    from datetime import date, timedelta
    conn = get_conn()
    try:
        # Cek apakah ada habits/dailies yang perlu diproses
        for table, ttype in [("habits", "habit"), ("dailies", "daily")]:
            check = conn.execute(
                f"SELECT COUNT(*) FROM {table} WHERE user_id=? AND last_done IS NOT NULL AND last_done < date('now')",
                (user_id,)
            ).fetchone()[0]
            if check == 0:
                continue
            
            # Ambil data
            if table == "habits":
                rows = conn.execute(f"SELECT id, last_done, streak, repeat_days FROM {table} WHERE user_id=?", (user_id,)).fetchall()
            else:
                rows = conn.execute(f"SELECT id, last_done, streak, fail_streak, repeat_days FROM {table} WHERE user_id=?", (user_id,)).fetchall()
            
            for r in rows:
                last = r["last_done"]
                if not last:
                    continue
                try:
                    last_date = datetime.strptime(last, "%Y-%m-%d").date()
                except ValueError:
                    continue
                current_date = date.today()
                if last_date >= current_date:
                    continue
                
                max_days = 30
                cutoff = current_date - timedelta(days=max_days)
                if last_date < cutoff:
                    last_date = cutoff
                
                d = last_date + timedelta(days=1)
                updates_needed = False
                
                while d < current_date:
                    if QApplication:
                        QApplication.processEvents()
                    d_str = d.isoformat()
                    existing = conn.execute(
                        "SELECT id FROM task_history WHERE user_id=? AND task_type=? AND task_id=? AND action_date=?",
                        (user_id, ttype, r["id"], d_str)
                    ).fetchone()
                    if not existing:
                        # Recurrence: hari di luar jadwal tidak dihitung skip
                        if not is_due_on(r["repeat_days"] if "repeat_days" in r.keys() else "", d):
                            d += timedelta(days=1)
                            continue
                        conn.execute(
                            "INSERT INTO task_history(user_id, task_type, task_id, action_date, action) VALUES(?,?,?,?,?)",
                            (user_id, ttype, r["id"], d_str, "skip")
                        )
                        updates_needed = True
                    d += timedelta(days=1)
                
                # Update hanya sekali di akhir
                if updates_needed:
                    if ttype == "daily":
                        new_fail = (r["fail_streak"] or 0) + 1
                        conn.execute(
                            f"UPDATE {table} SET streak = 0, fail_streak = ? WHERE id = ?",
                            (new_fail, r["id"])
                        )
                    else:
                        conn.execute(
                            f"UPDATE {table} SET streak = 0 WHERE id = ?",
                            (r["id"],)
                        )
        conn.commit()
    finally:
        conn.close()

def get_task_last_history(user_id, task_type, task_id):
    conn = get_conn()
    row = conn.execute(
        "SELECT action_date, action FROM task_history WHERE user_id=? AND task_type=? AND task_id=? ORDER BY action_date DESC, created_at DESC LIMIT 1",
        (user_id, task_type, task_id)
    ).fetchone()
    conn.close()
    return dict(row) if row else None

def get_all_task_history(user_id):
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM task_history WHERE user_id=? ORDER BY action_date DESC, created_at DESC",
        (user_id,)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@retry_on_lock
def reset_user_progress(user_id):
    """
    Reset semua progres user, tetapi pertahankan struktur task dan folder.
    """
    conn = get_conn()
    try:
        conn.execute("PRAGMA foreign_keys = OFF")

        # ========== 1. Reset kolom progres di users ==========
        conn.execute("""
            UPDATE users SET
                level = 1,
                xp = 0,
                hp = 50,
                max_hp = 50,
                mp = 30,
                max_mp = 30,
                gold = 0,
                gems = 10,
                longest_streak = 0,
                total_habits_done = 0,
                total_dailies_done = 0,
                total_todos_done = 0,
                total_xp_earned = 0,
                total_gold_earned = 0,
                boss_damage_bonus = 0,
                xp_multiplier = 1.0,
                gold_multiplier = 1.0,
                hp_damage_reduction = 0,
                has_revive = 0,
                mp_bonus = 0,
                sport_level = 1,
                sport_xp = 0,
                total_sport_points_earned = 0,
                total_tasks_completed = 0,
                total_gold_spent = 0,
                skill_buff_data = '{}',
                class_passive_buffs = '{}'
            WHERE id = ?
        """, (user_id,))

        # ========== 2. Reset habits, dailies, todos, sport ==========
        conn.execute("""
            UPDATE habits SET
                streak = 0,
                done_today = 0,
                counter_up = 0,
                counter_down = 0,
                last_done = NULL,
                last_action = ''
            WHERE user_id = ?
        """, (user_id,))

        conn.execute("""
            UPDATE dailies SET
                streak = 0,
                done_today = 0,
                last_done = NULL,
                last_action = ''
            WHERE user_id = ?
        """, (user_id,))

        conn.execute("UPDATE todos SET done = 0 WHERE user_id = ?", (user_id,))

        conn.execute("""
            UPDATE sport_activities SET
                streak = 0,
                done_today = 0,
                last_done = NULL
            WHERE user_id = ?
        """, (user_id,))

        # ========== 3. Hapus semua data progres yang punya user_id ==========
        tables_with_user_id = [
            "activity_log", "notifications", "task_history",
            "health_logs", "food_logs", "water_logs",
            "economy_items", "debts", "savings", "investments", "subscriptions",
            "user_pets", "inventory", "user_achievements",
            "boss_rewards", "food_achievements",
            "backup_codes"
        ]
        for table in tables_with_user_id:
            conn.execute(f"DELETE FROM {table} WHERE user_id = ?", (user_id,))

        # ========== 4. Hapus data dari tabel yang tidak punya user_id ==========
        # guild_messages pakai sender_id
        conn.execute("DELETE FROM guild_messages WHERE sender_id = ?", (user_id,))
        # messages pakai sender_id atau receiver_id
        conn.execute("DELETE FROM messages WHERE sender_id = ? OR receiver_id = ?", (user_id, user_id))
        # guild_invites dan guild_requests punya user_id, tapi kita juga hapus
        conn.execute("DELETE FROM guild_invites WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM guild_requests WHERE user_id = ?", (user_id,))
        # guild_leader_transfers punya old_leader_id
        conn.execute("DELETE FROM guild_leader_transfers WHERE old_leader_id = ?", (user_id,))
        # recipe_items? Kita hapus recipe dulu agar FK tidak error
        conn.execute("DELETE FROM recipe_items WHERE recipe_id IN (SELECT id FROM recipes WHERE user_id = ?)", (user_id,))
        conn.execute("DELETE FROM recipes WHERE user_id = ?", (user_id,))

        # ========== 5. Reset goals ke default ==========
        conn.execute("""
            INSERT OR REPLACE INTO user_nutrition_goals(user_id, daily_calories, daily_protein, daily_carbs, daily_fat)
            VALUES(?, 2000, 50, 250, 70)
        """, (user_id,))

        conn.execute("""
            INSERT OR REPLACE INTO user_water_goals(user_id, daily_ml)
            VALUES(?, 2000)
        """, (user_id,))

        conn.execute("""
            INSERT OR REPLACE INTO user_health_goals(user_id, daily_steps, daily_sleep_hours, height_cm, weight_kg, age, gender, activity_factor)
            VALUES(?, 10000, 7.0, 170, 70, 25, 'Laki-laki', 1.55)
        """, (user_id,))

        # ========== 6. Reset redeem_codes ==========
        conn.execute("UPDATE redeem_codes SET used_by = NULL, used_at = NULL WHERE used_by = ?", (user_id,))

        # ========== 7. Hapus teman? Tidak, biarkan hubungan tetap ==========
        # ========== 8. Security question? Biarkan, tidak direset ==========

        conn.commit()
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        conn.execute("PRAGMA foreign_keys = ON")
        conn.close()

def reset_progress_keep_assets(user_id):
    """Reset progres, tetapi pertahankan inventory, pets, task, dan folder."""
    conn = get_conn()
    try:
        conn.execute("PRAGMA foreign_keys = OFF")

        # Reset statistik user (sama seperti reset_user_progress)
        conn.execute("""
            UPDATE users SET
                level = 1,
                xp = 0,
                hp = 50,
                max_hp = 50,
                mp = 30,
                max_mp = 30,
                gold = 0,
                gems = 10,
                longest_streak = 0,
                total_habits_done = 0,
                total_dailies_done = 0,
                total_todos_done = 0,
                total_xp_earned = 0,
                total_gold_earned = 0,
                boss_damage_bonus = 0,
                xp_multiplier = 1.0,
                gold_multiplier = 1.0,
                hp_damage_reduction = 0,
                has_revive = 0,
                mp_bonus = 0,
                sport_level = 1,
                sport_xp = 0,
                total_sport_points_earned = 0,
                total_tasks_completed = 0,
                total_gold_spent = 0,
                skill_buff_data = '{}',
                class_passive_buffs = '{}'
            WHERE id = ?
        """, (user_id,))

        # Reset habits, dailies, todos, sport (sama)
        conn.execute("""
            UPDATE habits SET
                streak = 0,
                done_today = 0,
                counter_up = 0,
                counter_down = 0,
                last_done = NULL,
                last_action = ''
            WHERE user_id = ?
        """, (user_id,))

        conn.execute("""
            UPDATE dailies SET
                streak = 0,
                done_today = 0,
                last_done = NULL,
                last_action = ''
            WHERE user_id = ?
        """, (user_id,))

        conn.execute("UPDATE todos SET done = 0 WHERE user_id = ?", (user_id,))

        conn.execute("""
            UPDATE sport_activities SET
                streak = 0,
                done_today = 0,
                last_done = NULL
            WHERE user_id = ?
        """, (user_id,))

        # Hapus data progres (kecuali inventory dan user_pets)
        tables_to_delete = [
            "activity_log", "notifications", "task_history",
            "health_logs", "food_logs", "water_logs",
            "economy_items", "debts", "savings", "investments", "subscriptions",
            "user_achievements",  # achievement direset agar bisa diraih lagi
            "boss_rewards", "food_achievements",
            "backup_codes"
        ]
        for table in tables_to_delete:
            conn.execute(f"DELETE FROM {table} WHERE user_id = ?", (user_id,))

        # Hapus data chat, invite, request, recipe
        conn.execute("DELETE FROM guild_messages WHERE sender_id = ?", (user_id,))
        conn.execute("DELETE FROM messages WHERE sender_id = ? OR receiver_id = ?", (user_id, user_id))
        conn.execute("DELETE FROM guild_invites WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM guild_requests WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM guild_leader_transfers WHERE old_leader_id = ?", (user_id,))
        conn.execute("DELETE FROM recipe_items WHERE recipe_id IN (SELECT id FROM recipes WHERE user_id = ?)", (user_id,))
        conn.execute("DELETE FROM recipes WHERE user_id = ?", (user_id,))

        # Reset goals
        conn.execute("""
            INSERT OR REPLACE INTO user_nutrition_goals(user_id, daily_calories, daily_protein, daily_carbs, daily_fat)
            VALUES(?, 2000, 50, 250, 70)
        """, (user_id,))

        conn.execute("""
            INSERT OR REPLACE INTO user_water_goals(user_id, daily_ml)
            VALUES(?, 2000)
        """, (user_id,))

        conn.execute("""
            INSERT OR REPLACE INTO user_health_goals(user_id, daily_steps, daily_sleep_hours, height_cm, weight_kg, age, gender, activity_factor)
            VALUES(?, 10000, 7.0, 170, 70, 25, 'Laki-laki', 1.55)
        """, (user_id,))

        # Reset redeem_codes
        conn.execute("UPDATE redeem_codes SET used_by = NULL, used_at = NULL WHERE used_by = ?", (user_id,))

        conn.commit()
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        conn.execute("PRAGMA foreign_keys = ON")
        conn.close()

def perform_rebirth(user_id):
    """Lakukan rebirth: cek syarat, reset progres (kecuali inventory/pet), beri buff permanen."""
    u = get_user(user_id)
    if not u:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_user_not_found")}

    conn = get_conn()
    # Syarat 1: minimal 10 achievement
    ach_count = conn.execute(
        "SELECT COUNT(*) FROM user_achievements WHERE user_id=? AND unlocked_at IS NOT NULL",
        (user_id,)
    ).fetchone()[0]
    if ach_count < 10:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_rebirth_achievements_required", count=10)}

    # Syarat 2: level >= 25
    if u["level"] < 25:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_rebirth_level_required", level=25)}

    # Syarat 3: minimal 2 pet
    pet_count = conn.execute("SELECT COUNT(*) FROM user_pets WHERE user_id=?", (user_id,)).fetchone()[0]
    if pet_count < 2:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_rebirth_pets_required", pets=2)}

    # Syarat 4: minimal 6 item di inventory
    item_count = conn.execute("SELECT COUNT(*) FROM inventory WHERE user_id=?", (user_id,)).fetchone()[0]
    if item_count < 6:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_rebirth_items_required", items=6)}
    conn.close()

    # Reset progres (pertahankan inventory, pet, task, folder)
    reset_progress_keep_assets(user_id)

    # Tambah rebirth_count
    conn2 = get_conn()
    conn2.execute("UPDATE users SET rebirth_count = rebirth_count + 1 WHERE id = ?", (user_id,))
    conn2.commit()
    conn2.close()

    # Recalculate buffs untuk mengaplikasikan bonus rebirth
    recalculate_all_buffs(user_id)

    new_count = get_user(user_id).get("rebirth_count", 0)
    add_notification(user_id, tr_db(user_id=user_id, key="db_rebirth_success", count=new_count), "levelup")

    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_rebirth_success", count=new_count), "rebirth_count": new_count}

def get_leaderboard_for_user(user_id):
    """Ambil leaderboard yang hanya menampilkan user sendiri, teman, dan anggota guild."""
    conn = get_conn()
    # Ambil ID teman yang sudah accepted
    friends = conn.execute(
        "SELECT friend_id FROM friends WHERE user_id=? AND status='accepted' "
        "UNION SELECT user_id FROM friends WHERE friend_id=? AND status='accepted'",
        (user_id, user_id)
    ).fetchall()
    friend_ids = [row[0] for row in friends]

    # Ambil guild_id user
    guild_row = conn.execute("SELECT guild_id FROM users WHERE id=?", (user_id,)).fetchone()
    guild_id = guild_row["guild_id"] if guild_row else None

    # Ambil anggota guild jika ada
    guild_member_ids = []
    if guild_id:
        members = conn.execute(
            "SELECT user_id FROM guild_members WHERE guild_id=?", (guild_id,)
        ).fetchall()
        guild_member_ids = [row[0] for row in members]

    # Gabungkan semua ID unik: user sendiri, teman, anggota guild
    all_ids = set([user_id] + friend_ids + guild_member_ids)
    if not all_ids:
        return []

    # Buat query dengan placeholders
    placeholders = ",".join("?" * len(all_ids))
    query = f"""
        SELECT username, display_name, level, total_xp_earned, gold,
               COALESCE(selected_title, '') as selected_title,
               COALESCE(sport_level, 1) as sport_level,
               (SELECT COUNT(*) FROM user_pets WHERE user_id=users.id) as pet_count,
               COALESCE(rebirth_count, 0) as rebirth_count
        FROM users
        WHERE id IN ({placeholders})
        AND is_admin = 0
        ORDER BY level DESC, total_xp_earned DESC
    """
    rows = conn.execute(query, tuple(all_ids)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def move_item_up(user_id, mode, item_id):
    """Tukar sort_order dengan item sebelumnya."""
    table_map = {...}  # sama
    tbl = table_map[mode]
    conn = get_conn()
    # Ambil item saat ini dan sebelumnya
    current = conn.execute(f"SELECT sort_order, folder_id FROM {tbl} WHERE id=?", (item_id,)).fetchone()
    if not current:
        return
    folder_id = current["folder_id"]
    prev = conn.execute(
        f"SELECT id, sort_order FROM {tbl} WHERE user_id=? AND folder_id IS ? AND sort_order < ? ORDER BY sort_order DESC LIMIT 1",
        (user_id, folder_id, current["sort_order"])
    ).fetchone()
    if prev:
        # Tukar sort_order
        conn.execute(f"UPDATE {tbl} SET sort_order=? WHERE id=?", (prev["sort_order"], item_id))
        conn.execute(f"UPDATE {tbl} SET sort_order=? WHERE id=?", (current["sort_order"], prev["id"]))
        conn.commit()
    conn.close()

def move_item_down(user_id, mode, item_id):
    """Tukar sort_order dengan item berikutnya."""
    # Analog

def move_item_to_folder(user_id, mode, item_id, target_folder_id):
    """Pindahkan item ke folder tertentu (untuk habit/daily/todo/sport/economy) 
       atau ubah meal_type (untuk food_logs)."""
    table_map = {
        "habit": "habits",
        "daily": "dailies",
        "todo": "todos",
        "sport": "sport_activities",
        "economy": "economy_items",
        "food": "food_logs"
    }
    tbl = table_map.get(mode)
    if not tbl:
        return
    
    conn = get_conn()
    now = datetime.now().isoformat()
    
    if mode == "food":
        # target_folder_id sebenarnya adalah meal_type
        valid_meals = ("breakfast", "lunch", "dinner", "snack", None)
        if target_folder_id not in valid_meals:
            target_folder_id = "snack"  # fallback
        conn.execute(
            f"UPDATE {tbl} SET meal_type=? WHERE id=? AND user_id=?",
            (target_folder_id, item_id, user_id)
        )
    else:
        conn.execute(
            f"UPDATE {tbl} SET folder_id=?, created_at=? WHERE id=? AND user_id=?",
            (target_folder_id, now, item_id, user_id)
        )
    
    conn.commit()
    conn.close()

def reorder_item(user_id, mode, item_id, direction, status=None):
    """Tukar sort_order dengan item di atas/bawah dalam folder dan status yang sama.
    status: untuk todo = done, untuk habit/daily/sport = done_today.
    """
    table_map = {
        "habit": "habits",
        "daily": "dailies",
        "todo": "todos",
        "sport": "sport_activities",
        "economy": "economy_items"
    }
    tbl = table_map.get(mode)
    if not tbl:
        return {"ok": False, "msg": "Mode tidak didukung"}

    # Tentukan kolom status berdasarkan mode
    if mode == "todo":
        status_col = "done"
    elif mode in ("habit", "daily", "sport"):
        status_col = "done_today"
    else:
        status_col = None

    conn = get_conn()
    try:
        # Ambil data item saat ini
        if status_col:
            current = conn.execute(
                f"SELECT folder_id, sort_order, {status_col} as status FROM {tbl} WHERE id=? AND user_id=?",
                (item_id, user_id)
            ).fetchone()
        else:
            current = conn.execute(
                f"SELECT folder_id, sort_order FROM {tbl} WHERE id=? AND user_id=?",
                (item_id, user_id)
            ).fetchone()
        if not current:
            return {"ok": False, "msg": "Item tidak ditemukan"}

        folder_id = current["folder_id"]
        current_order = current["sort_order"]
        current_status = current["status"] if status_col else None

        # Gunakan status dari parameter jika diberikan, atau status item saat ini
        target_status = status if status is not None else current_status

        # Cari item target (atas/bawah) dengan status yang sama (jika ada)
        if status_col is not None and target_status is not None:
            params = [user_id, folder_id, target_status, current_order]
            if direction == "up":
                query = f"""SELECT id, sort_order FROM {tbl}
                            WHERE user_id=? AND folder_id IS ? AND {status_col}=? AND sort_order < ?
                            ORDER BY sort_order DESC LIMIT 1"""
            else:
                query = f"""SELECT id, sort_order FROM {tbl}
                            WHERE user_id=? AND folder_id IS ? AND {status_col}=? AND sort_order > ?
                            ORDER BY sort_order ASC LIMIT 1"""
            target = conn.execute(query, tuple(params)).fetchone()
        else:
            # Tanpa status (misal economy)
            params = [user_id, folder_id, current_order]
            if direction == "up":
                query = f"""SELECT id, sort_order FROM {tbl}
                            WHERE user_id=? AND folder_id IS ? AND sort_order < ?
                            ORDER BY sort_order DESC LIMIT 1"""
            else:
                query = f"""SELECT id, sort_order FROM {tbl}
                            WHERE user_id=? AND folder_id IS ? AND sort_order > ?
                            ORDER BY sort_order ASC LIMIT 1"""
            target = conn.execute(query, tuple(params)).fetchone()

        if not target:
            return {"ok": True, "moved": False, "msg": "Sudah di posisi paling ujung"}

        # Tukar sort_order
        conn.execute(
            f"UPDATE {tbl} SET sort_order = ? WHERE id = ? AND user_id = ?",
            (target["sort_order"], item_id, user_id)
        )
        conn.execute(
            f"UPDATE {tbl} SET sort_order = ? WHERE id = ? AND user_id = ?",
            (current_order, target["id"], user_id)
        )
        conn.commit()
        return {"ok": True, "moved": True}
    except Exception as e:
        conn.rollback()
        return {"ok": False, "msg": str(e)}
    finally:
        conn.close()

def reorder_item_relative(user_id, mode, source_id, target_id):
    """
    Pindahkan source_id ke posisi relatif terhadap target_id dalam folder yang sama.
    Jika source berada di atas target, pindahkan ke bawah target (after).
    Jika source berada di bawah target, pindahkan ke atas target (before).
    """
    table_map = {
        "habit": "habits",
        "daily": "dailies",
        "todo": "todos",
        "sport": "sport_activities",
        "economy": "economy_items"
    }
    tbl = table_map.get(mode)
    if not tbl:
        return {"ok": False}

    conn = get_conn()
    try:
        # Ambil folder_id source dan target
        source = conn.execute(
            f"SELECT folder_id FROM {tbl} WHERE id=? AND user_id=?",
            (source_id, user_id)
        ).fetchone()
        target = conn.execute(
            f"SELECT folder_id FROM {tbl} WHERE id=? AND user_id=?",
            (target_id, user_id)
        ).fetchone()
        if not source or not target or source["folder_id"] != target["folder_id"]:
            return {"ok": False}

        folder_id = source["folder_id"]

        # Ambil semua id item dalam folder, urutkan berdasarkan sort_order
        rows = conn.execute(
            f"SELECT id FROM {tbl} WHERE user_id=? AND folder_id IS ? ORDER BY sort_order",
            (user_id, folder_id)
        ).fetchall()
        item_ids = [r["id"] for r in rows]

        if source_id not in item_ids or target_id not in item_ids:
            return {"ok": False}

        src_idx = item_ids.index(source_id)
        tgt_idx = item_ids.index(target_id)

        if src_idx == tgt_idx:
            return {"ok": True}  # tidak ada perubahan

        # Hapus source dari list
        item_ids.pop(src_idx)

        # Cari indeks target yang baru (setelah source dihapus)
        new_tgt_idx = item_ids.index(target_id)

        if src_idx < tgt_idx:
            # Source berada di atas target → pindahkan ke bawah target
            insert_pos = new_tgt_idx + 1
        else:
            # Source berada di bawah target → pindahkan ke atas target
            insert_pos = new_tgt_idx

        item_ids.insert(insert_pos, source_id)

        # Update ulang sort_order secara berurutan
        for i, item_id in enumerate(item_ids, start=1):
            conn.execute(
                f"UPDATE {tbl} SET sort_order = ? WHERE id = ? AND user_id = ?",
                (i, item_id, user_id)
            )
        conn.commit()
        return {"ok": True}
    except Exception as e:
        conn.rollback()
        return {"ok": False, "msg": str(e)}
    finally:
        conn.close()

# ========== NOTES CRUD ==========
def get_note_folders(user_id):
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM note_folders WHERE user_id=? ORDER BY name",
        (user_id,)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_note_folders_tree(user_id):
    """Ambil semua folder user dan bangun struktur tree (nested)."""
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM note_folders WHERE user_id=? ORDER BY name",
        (user_id,)
    ).fetchall()
    conn.close()

    folders = [dict(r) for r in rows]
    # Bangun mapping id -> node
    node_map = {f["id"]: {**f, "children": []} for f in folders}
    tree = []
    for f in folders:
        node = node_map[f["id"]]
        parent_id = f.get("parent_id")
        if parent_id is None:
            tree.append(node)
        else:
            parent = node_map.get(parent_id)
            if parent:
                parent["children"].append(node)
            else:
                # fallback: jika parent tidak ditemukan, taruh di root
                tree.append(node)
    return tree

def add_note_folder(user_id, name, icon="📁", parent_id=None):
    conn = get_conn()
    cur = conn.execute(
        "INSERT INTO note_folders(user_id, name, icon, parent_id) VALUES(?,?,?,?)",
        (user_id, name, icon, parent_id)
    )
    folder_id = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "folder_id": folder_id}

def update_note_folder(folder_id, user_id, name=None, icon=None):
    updates = []
    values = []
    if name is not None:
        updates.append("name=?")
        values.append(name)
    if icon is not None:
        updates.append("icon=?")
        values.append(icon)
    if not updates:
        return
    values.append(folder_id)
    values.append(user_id)
    conn = get_conn()
    conn.execute(
        f"UPDATE note_folders SET {', '.join(updates)} WHERE id=? AND user_id=?",
        values
    )
    conn.commit()
    conn.close()

def delete_note_folder(folder_id, user_id):
    conn = get_conn()
    conn.execute("DELETE FROM note_folders WHERE id=? AND user_id=?", (folder_id, user_id))
    conn.commit()
    conn.close()

def duplicate_note_folder(user_id, folder_id):
    """Duplikasi folder beserta semua subfolder dan isi notes-nya."""
    conn = get_conn()
    try:
        # Ambil data folder asli
        src = conn.execute(
            "SELECT name, icon, parent_id FROM note_folders WHERE id=? AND user_id=?",
            (folder_id, user_id)
        ).fetchone()
        if not src:
            return {"ok": False, "msg": "Folder tidak ditemukan"}
        
        new_name = src["name"]
        icon = src["icon"]
        parent_id = src["parent_id"]

        cur = conn.execute(
            "INSERT INTO note_folders(user_id, name, icon, parent_id) VALUES(?,?,?,?)",
            (user_id, new_name, icon, parent_id)
        )
        new_folder_id = cur.lastrowid

        # Duplikasi semua note dari folder asli ke folder baru
        notes = conn.execute(
            "SELECT title, content, is_archived FROM notes WHERE user_id=? AND folder_id=?",
            (user_id, folder_id)
        ).fetchall()
        for note in notes:
            conn.execute(
                "INSERT INTO notes(user_id, folder_id, title, content, is_archived) VALUES(?,?,?,?,?)",
                (user_id, new_folder_id, note["title"], note["content"], note["is_archived"])
            )

        # Duplikasi subfolder secara rekursif
        subfolders = conn.execute(
            "SELECT id FROM note_folders WHERE user_id=? AND parent_id=?",
            (user_id, folder_id)
        ).fetchall()
        for sub in subfolders:
            _duplicate_subfolder(conn, user_id, sub["id"], new_folder_id)

        conn.commit()
        return {"ok": True, "new_folder_id": new_folder_id, "msg": f"Folder '{new_name}' berhasil diduplikasi!"}
    except Exception as e:
        conn.rollback()
        return {"ok": False, "msg": str(e)}
    finally:
        conn.close()

def _duplicate_subfolder(conn, user_id, src_folder_id, new_parent_id):
    """Fungsi internal rekursif untuk duplikasi subfolder."""
    src = conn.execute(
        "SELECT name, icon FROM note_folders WHERE id=? AND user_id=?",
        (src_folder_id, user_id)
    ).fetchone()
    if not src:
        return
    cur = conn.execute(
        "INSERT INTO note_folders(user_id, name, icon, parent_id) VALUES(?,?,?,?)",
        (user_id, src["name"], src["icon"], new_parent_id)
    )
    new_sub_id = cur.lastrowid

    # Duplikasi notes di subfolder
    notes = conn.execute(
        "SELECT title, content, is_archived FROM notes WHERE user_id=? AND folder_id=?",
        (user_id, src_folder_id)
    ).fetchall()
    for note in notes:
        conn.execute(
            "INSERT INTO notes(user_id, folder_id, title, content, is_archived) VALUES(?,?,?,?,?)",
            (user_id, new_sub_id, note["title"], note["content"], note["is_archived"])
        )

    # Proses sub-sub folder
    subsub = conn.execute(
        "SELECT id FROM note_folders WHERE user_id=? AND parent_id=?",
        (user_id, src_folder_id)
    ).fetchall()
    for s in subsub:
        _duplicate_subfolder(conn, user_id, s["id"], new_sub_id)

def update_note_folder_icon(folder_id, user_id, new_icon):
    conn = get_conn()
    conn.execute(
        "UPDATE note_folders SET icon=? WHERE id=? AND user_id=?",
        (new_icon, folder_id, user_id)
    )
    conn.commit()
    conn.close()
    return {"ok": True}

def get_notes(user_id, folder_id=None, include_archived=False):
    conn = get_conn()
    query = "SELECT * FROM notes WHERE user_id=?"
    params = [user_id]
    if folder_id == -1:
        query += " AND folder_id IS NULL"
    elif folder_id is not None:
        query += " AND folder_id=?"
        params.append(folder_id)
    # jika folder_id is None, tidak ada filter folder
    if not include_archived:
        query += " AND is_archived=0"
    query += " ORDER BY updated_at DESC"
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_note(note_id, user_id):
    conn = get_conn()
    row = conn.execute(
        "SELECT * FROM notes WHERE id=? AND user_id=?",
        (note_id, user_id)
    ).fetchone()
    conn.close()
    return dict(row) if row else None

def add_note(user_id, folder_id, title, content=""):
    conn = get_conn()
    cur = conn.execute(
        "INSERT INTO notes(user_id, folder_id, title, content, zoom_level) VALUES(?,?,?,?,?)",
        (user_id, folder_id, title, content, 100)
    )
    note_id = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "note_id": note_id}

def update_note(note_id, user_id, title=None, content=None, folder_id=None, zoom_level=None):
    updates = []
    values = []
    if title is not None:
        updates.append("title=?")
        values.append(title)
    if content is not None:
        updates.append("content=?")
        values.append(content)
    if folder_id is not None:
        updates.append("folder_id=?")
        values.append(folder_id)
    if zoom_level is not None:
        updates.append("zoom_level=?")
        values.append(zoom_level)
    if not updates:
        return
    updates.append("updated_at=datetime('now')")
    values.append(note_id)
    values.append(user_id)
    conn = get_conn()
    conn.execute(
        f"UPDATE notes SET {', '.join(updates)} WHERE id=? AND user_id=?",
        values
    )
    conn.commit()
    conn.close()

def duplicate_note(user_id, note_id, dest_folder_id=None):
    """Duplikasi catatan ke folder tujuan (bisa None untuk tanpa folder). FIX 3"""
    conn = get_conn()
    src = conn.execute("SELECT title, content, is_archived FROM notes WHERE id=? AND user_id=?", (note_id, user_id)).fetchone()
    if not src:
        conn.close()
        return {"ok": False, "msg": "Catatan tidak ditemukan"}
    if dest_folder_id is not None:
        chk = conn.execute("SELECT id FROM note_folders WHERE id=? AND user_id=?", (dest_folder_id, user_id)).fetchone()
        if not chk:
            conn.close()
            return {"ok": False, "msg": "Folder tujuan tidak valid"}
    new_title = src["title"] + " (Copy)"
    cur = conn.execute("INSERT INTO notes(user_id, folder_id, title, content, is_archived) VALUES(?,?,?,?,?)",
        (user_id, dest_folder_id, new_title, src["content"], src["is_archived"]))
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "new_note_id": new_id, "msg": f"Catatan '{new_title}' berhasil diduplikasi!"}

def delete_note(note_id, user_id):
    purge_trash()
    conn = get_conn()
    trash_id = _stash_before_delete(conn, user_id, "notes", note_id, "note")
    conn.execute("DELETE FROM notes WHERE id=? AND user_id=?", (note_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True, "trash_id": trash_id}

def archive_note(note_id, user_id, archived):
    conn = get_conn()
    conn.execute(
        "UPDATE notes SET is_archived=?, updated_at=datetime('now') WHERE id=? AND user_id=?",
        (1 if archived else 0, note_id, user_id)
    )
    conn.commit()
    conn.close()

# ========== REMINDERS CRUD ==========
def get_reminders(user_id, active_only=False):
    conn = get_conn()
    query = "SELECT * FROM reminders WHERE user_id=?"
    params = [user_id]
    if active_only:
        query += " AND is_active=1"
    query += " ORDER BY reminder_datetime ASC"
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_reminder(reminder_id, user_id):
    conn = get_conn()
    row = conn.execute(
        "SELECT * FROM reminders WHERE id=? AND user_id=?",
        (reminder_id, user_id)
    ).fetchone()
    conn.close()
    return dict(row) if row else None

def add_reminder(user_id, title, description, reminder_datetime, sound_type="default", sound_file=None, repeat_type="none", repeat_days=""):
    conn = get_conn()
    cur = conn.execute(
        """INSERT INTO reminders(user_id, title, description, reminder_datetime, sound_type, sound_file, repeat_type, repeat_days)
           VALUES(?,?,?,?,?,?,?,?)""",
        (user_id, title, description, reminder_datetime, sound_type, sound_file, repeat_type, repeat_days)
    )
    reminder_id = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "reminder_id": reminder_id}

def update_reminder(reminder_id, user_id, **kwargs):
    fields = []
    values = []
    for key, value in kwargs.items():
        if key in ("title", "description", "reminder_datetime", "sound_type", "sound_file", "is_active", "triggered", "repeat_type", "repeat_days", "repeat_until"):
            fields.append(f"{key}=?")
            values.append(value)
    if not fields:
        return
    values.append(reminder_id)
    values.append(user_id)
    conn = get_conn()
    conn.execute(
        f"UPDATE reminders SET {', '.join(fields)} WHERE id=? AND user_id=?",
        values
    )
    conn.commit()
    conn.close()

def delete_reminder(reminder_id, user_id):
    conn = get_conn()
    conn.execute("DELETE FROM reminders WHERE id=? AND user_id=?", (reminder_id, user_id))
    conn.commit()
    conn.close()

def get_pending_reminders(user_id):
    """Ambil reminder yang harus dipicu (waktu <= sekarang, aktif, belum terpicu)"""
    conn = get_conn()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    rows = conn.execute(
        """SELECT * FROM reminders
           WHERE user_id=? AND is_active=1 AND triggered=0 AND reminder_datetime <= ?
           ORDER BY reminder_datetime ASC""",
        (user_id, now)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def mark_reminder_triggered(reminder_id, user_id):
    conn = get_conn()
    conn.execute(
        "UPDATE reminders SET triggered=1 WHERE id=? AND user_id=?",
        (reminder_id, user_id)
    )
    conn.commit()
    conn.close()

def reset_reminder_triggered(reminder_id, user_id):
    """Reset triggered untuk reminder yang sudah lewat (misal untuk besok)"""
    conn = get_conn()
    conn.execute(
        "UPDATE reminders SET triggered=0 WHERE id=? AND user_id=?",
        (reminder_id, user_id)
    )
    conn.commit()
    conn.close()

def get_next_reminder_datetime(current_dt_str, repeat_type, repeat_days, repeat_until=None):
    """
    Hitung tanggal/waktu berikutnya berdasarkan repeat_type.
    - repeat_type: 'none', 'daily', 'weekly', 'custom'
    - repeat_days: string angka hari (0=Senin, 6=Minggu) dipisah koma, misal "0,2,4"
    - repeat_until: tanggal akhir (YYYY-MM-DD) atau None
    Mengembalikan string datetime baru (YYYY-MM-DD HH:MM:SS) atau None jika tidak ada (sudah melewati repeat_until)
    """
    if repeat_type == 'none':
        return None

    current_dt = datetime.strptime(current_dt_str, "%Y-%m-%d %H:%M:%S")
    time_part = current_dt.strftime("%H:%M:%S")

    # Hitung next date
    if repeat_type == 'daily':
        next_dt = current_dt + timedelta(days=1)
    elif repeat_type == 'weekly':
        next_dt = current_dt + timedelta(days=7)
    elif repeat_type == 'custom':
        if not repeat_days:
            return None
        days_list = [int(d.strip()) for d in repeat_days.split(',') if d.strip()]
        if not days_list:
            return None
        # Cari hari berikutnya yang ada di days_list
        current_weekday = current_dt.weekday()  # 0=Senin, 6=Minggu
        for offset in range(1, 8):
            candidate = current_dt + timedelta(days=offset)
            if candidate.weekday() in days_list:
                next_dt = candidate
                break
        else:
            # Tidak ada hari yang cocok dalam 7 hari ke depan (seharusnya tidak terjadi)
            return None
    else:
        return None

    # Jika repeat_until, cek apakah next_dt melewati batas
    if repeat_until:
        until_date = datetime.strptime(repeat_until, "%Y-%m-%d").date()
        if next_dt.date() > until_date:
            return None

    # Gabungkan dengan waktu dari reminder asli
    next_dt_str = next_dt.strftime("%Y-%m-%d") + " " + time_part
    return next_dt_str

# ═══════════════════════════════════════════════════════════════════
#  CALENDAR NOTES
# ═══════════════════════════════════════════════════════════════════
@retry_on_lock
def get_calendar_notes(user_id, year=None, month=None):
    """Ambil catatan kalender untuk user, opsional filter tahun/bulan."""
    conn = get_conn()
    query = "SELECT note_date, note FROM calendar_notes WHERE user_id = ?"
    params = [user_id]
    if year is not None and month is not None:
        start_date = f"{year:04d}-{month:02d}-01"
        # akhir bulan
        from calendar import monthrange
        last_day = monthrange(year, month)[1]
        end_date = f"{year:04d}-{month:02d}-{last_day:02d}"
        query += " AND note_date BETWEEN ? AND ?"
        params.extend([start_date, end_date])
    elif year is not None:
        start_date = f"{year:04d}-01-01"
        end_date = f"{year:04d}-12-31"
        query += " AND note_date BETWEEN ? AND ?"
        params.extend([start_date, end_date])
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return {row["note_date"]: row["note"] for row in rows}

@retry_on_lock
def get_calendar_note(user_id, date_str):
    """Ambil catatan untuk satu tanggal."""
    conn = get_conn()
    row = conn.execute("SELECT note FROM calendar_notes WHERE user_id = ? AND note_date = ?", (user_id, date_str)).fetchone()
    conn.close()
    return row["note"] if row else ""

@retry_on_lock
def save_calendar_note(user_id, date_str, note):
    """Simpan atau update catatan untuk tanggal tertentu."""
    conn = get_conn()
    conn.execute(
        """INSERT INTO calendar_notes(user_id, note_date, note, updated_at)
           VALUES(?,?,?,datetime('now'))
           ON CONFLICT(user_id, note_date) DO UPDATE SET note=excluded.note, updated_at=datetime('now')""",
        (user_id, date_str, note)
    )
    conn.commit()
    conn.close()

@retry_on_lock
def delete_calendar_note(user_id, date_str):
    """Hapus catatan untuk tanggal tertentu."""
    conn = get_conn()
    conn.execute("DELETE FROM calendar_notes WHERE user_id = ? AND note_date = ?", (user_id, date_str))
    conn.commit()
    conn.close()

_checkpoint_timer = None

def start_periodic_checkpoint(interval=10):
    global _checkpoint_timer
    if _checkpoint_timer:
        return
    _checkpoint_timer = threading.Timer(interval, _do_checkpoint)
    _checkpoint_timer.daemon = True
    _checkpoint_timer.start()

def _do_checkpoint():
    global _checkpoint_timer
    try:
        _c = sqlite3.connect(DB_PATH, timeout=60.0)
        _c.execute("PRAGMA wal_checkpoint(PASSIVE)")
        _c.close()  # koneksi langsung (bukan pooled) -> hindari bocor di thread Timer
    except:
        pass
    _checkpoint_timer = threading.Timer(10, _do_checkpoint)
    _checkpoint_timer.daemon = True
    _checkpoint_timer.start()

def stop_periodic_checkpoint():
    global _checkpoint_timer
    if _checkpoint_timer:
        _checkpoint_timer.cancel()
        _checkpoint_timer = None
    force_checkpoint()

# ========== GET ALL ACTIVE BUFFS ==========
def get_all_active_buffs(user_id):
    """Ambil semua buff aktif user dalam bentuk list of string (Indonesia)."""
    u = get_user(user_id)
    if not u:
        return []
    buffs = []
    
    # 1. XP Multiplier
    if u.get("xp_multiplier", 1.0) > 1.001:
        buffs.append(f"📈 XP Multiplier: x{u['xp_multiplier']:.2f}")
    # 2. Gold Multiplier
    if u.get("gold_multiplier", 1.0) > 1.001:
        buffs.append(f"💰 Gold Multiplier: x{u['gold_multiplier']:.2f}")
    # 3. Boss Damage Bonus
    if u.get("boss_damage_bonus", 0) > 0:
        buffs.append(f"⚔️ Boss Damage: +{u['boss_damage_bonus']:.0f}")
    # 4. HP Damage Reduction
    if u.get("hp_damage_reduction", 0) > 0:
        buffs.append(f"🛡️ HP Reduction: -{u['hp_damage_reduction']:.0f}")
    # 5. MP Bonus
    if u.get("mp_bonus", 0) > 0:
        buffs.append(f"💙 Max MP: +{u['mp_bonus']}")
    # 6. Revive
    if u.get("has_revive", 0):
        buffs.append("🗿 Totem of Life (revive once)")
    # 7. Critical Chance
    if u.get("crit_chance", 10) > 10:
        buffs.append(f"⚡ Critical Chance: +{u['crit_chance']-10}%")
    # 8. Block Chance
    if u.get("block_chance", 20) > 20:
        buffs.append(f"🛡️ Block Chance: +{u['block_chance']-20}%")
    # 9. Block Strength
    if u.get("block_strength", 10) > 10:
        buffs.append(f"🛡️ Block Strength: +{u['block_strength']-10}")
    # 10. Spyglass
    if u.get("has_spyglass", 0):
        buffs.append("🔭 Spyglass (reveal boss stats)")
    
    # 11. Guild buffs
    guild_id = u.get("guild_id")
    if guild_id:
        conn = get_conn()
        g = conn.execute("SELECT buff_xp, buff_gold, buff_damage, crit_chance FROM guilds WHERE id=?", (guild_id,)).fetchone()
        conn.close()
        if g:
            if g["buff_xp"] > 0:
                buffs.append(f"🏰 Guild XP Bonus: +{g['buff_xp']:.0f}%")
            if g["buff_gold"] > 0:
                buffs.append(f"🏰 Guild Gold Bonus: +{g['buff_gold']:.0f}%")
            if g["buff_damage"] > 0:
                buffs.append(f"🏰 Guild Damage: +{g['buff_damage']:.0f}")
            if g["crit_chance"] > 0:
                buffs.append(f"🏰 Guild Crit Chance: +{g['crit_chance']}%")
    
    # 12. Pet aktif
    conn = get_conn()
    active_pet = conn.execute(
        "SELECT pet_id, level FROM user_pets WHERE user_id=? AND is_active=1",
        (user_id,)
    ).fetchone()
    conn.close()
    if active_pet:
        pet = PETS_DATA.get(active_pet["pet_id"], {})
        pet_name = pet.get('name', active_pet["pet_id"])
        bonus = pet.get('bonus', '')
        level = active_pet["level"]
        buffs.append(f"🐾 {pet_name} (Lv.{level}): {bonus}")
    
    # 13. Class passive
    class_buffs = get_class_passive_buffs(user_id)
    if class_buffs.get("hp_multiplier", 1.0) > 1.001:
        hp_bonus = (class_buffs['hp_multiplier']-1)*100
        buffs.append(f"❤️ Class HP Bonus: +{hp_bonus:.0f}%")
    if class_buffs.get("streak_bonus", 0):
        buffs.append("🔥 Class Streak Bonus: +1 streak")
    
    # 14. Skill buffs aktif
    skill_buffs = get_skill_buffs(user_id)
    if skill_buffs.get("shield_active"):
        buffs.append("🛡️ Shield Bash (next boss damage -50%)")
    if skill_buffs.get("xp_multiplier") and skill_buffs.get("xp_remaining", 0) > 0:
        buffs.append(f"✨ Arcane Surge: +30% XP for {skill_buffs['xp_remaining']} habits")
    if skill_buffs.get("gold_multiplier") and skill_buffs.get("gold_remaining", 0) > 0:
        buffs.append(f"🏹 Gold Shot: +50% Gold for {skill_buffs['gold_remaining']} habit")
    if skill_buffs.get("double_streak"):
        buffs.append("🗡️ Shadow Step: +1 streak for next daily")
    
    # 15. Rebirth buff
    rebirth_count = u.get("rebirth_count", 0)
    if rebirth_count > 0:
        buffs.append(f"🌀 Rebirth Bonus: +{rebirth_count*10}% XP, +{rebirth_count*5}% Gold")
    
    return buffs

# ========== PLAYLIST FUNCTIONS ==========
def save_playlist(user_id, name, files):
    """Simpan playlist ke database, kembalikan ID playlist."""
    try:
        conn = get_conn()
        cur = conn.execute(
            "INSERT INTO playlists(user_id, name, tracks, is_favorite) VALUES(?,?,?,?)",
            (user_id, name, json.dumps(files), 0)
        )
        playlist_id = cur.lastrowid
        conn.commit()
        conn.close()
        return {"ok": True, "id": playlist_id, "msg": "Playlist berhasil disimpan!"}
    except Exception as e:
        return {"ok": False, "msg": f"Gagal menyimpan playlist: {str(e)}"}

def load_playlist(playlist_id, user_id):
    """Ambil tracks dari playlist."""
    conn = get_conn()
    row = conn.execute(
        "SELECT tracks, name FROM playlists WHERE id=? AND user_id=?",
        (playlist_id, user_id)
    ).fetchone()
    conn.close()
    if row:
        return {"ok": True, "tracks": json.loads(row["tracks"]), "name": row["name"]}
    return {"ok": False}

def create_playlist(user_id, name, is_favorite=0):
    conn = get_conn()
    cur = conn.execute(
        "INSERT INTO playlists(user_id, name, tracks, is_favorite) VALUES(?,?,?,?)",
        (user_id, name, json.dumps([]), is_favorite)
    )
    playlist_id = cur.lastrowid
    conn.commit()
    conn.close()
    return playlist_id

def get_all_playlists(user_id):
    conn = get_conn()
    rows = conn.execute(
        "SELECT id, name, tracks, is_favorite FROM playlists WHERE user_id=? ORDER BY is_favorite DESC, name",
        (user_id,)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_playlist(user_id, playlist_id):
    conn = get_conn()
    row = conn.execute(
        "SELECT id, name, tracks, is_favorite FROM playlists WHERE id=? AND user_id=?",
        (playlist_id, user_id)
    ).fetchone()
    conn.close()
    return dict(row) if row else None

def update_playlist_tracks(user_id, playlist_id, tracks):
    conn = get_conn()
    conn.execute(
        "UPDATE playlists SET tracks=? WHERE id=? AND user_id=?",
        (json.dumps(tracks), playlist_id, user_id)
    )
    conn.commit()
    conn.close()

def rename_playlist(user_id, playlist_id, new_name):
    """Ganti nama playlist."""
    conn = get_conn()
    conn.execute(
        "UPDATE playlists SET name=? WHERE id=? AND user_id=?",
        (new_name, playlist_id, user_id)
    )
    conn.commit()
    conn.close()

def delete_playlist(user_id, playlist_id):
    """Hapus playlist (tidak bisa hapus Favorite)."""
    # Cek apakah ini favorite
    playlist = get_playlist(user_id, playlist_id)
    if playlist and playlist['is_favorite']:
        return False  # tidak bisa hapus favorite
    conn = get_conn()
    conn.execute("DELETE FROM playlists WHERE id=? AND user_id=?", (playlist_id, user_id))
    conn.commit()
    conn.close()
    return True

def add_song_to_playlist(user_id, playlist_id, file_path):
    playlist = get_playlist(user_id, playlist_id)
    if not playlist:
        return False
    tracks = json.loads(playlist['tracks'])
    if file_path not in tracks:
        tracks.append(file_path)
        update_playlist_tracks(user_id, playlist_id, tracks)
        return True
    return False

def remove_song_from_playlist(user_id, playlist_id, index):
    playlist = get_playlist(user_id, playlist_id)
    if not playlist:
        return False
    tracks = json.loads(playlist['tracks'])
    if 0 <= index < len(tracks):
        del tracks[index]
        update_playlist_tracks(user_id, playlist_id, tracks)
        return True
    return False

def move_song_to_playlist(user_id, from_playlist_id, to_playlist_id, index):
    from_pl = get_playlist(user_id, from_playlist_id)
    to_pl = get_playlist(user_id, to_playlist_id)
    if not from_pl or not to_pl:
        return False
    from_tracks = json.loads(from_pl['tracks'])
    if index < 0 or index >= len(from_tracks):
        return False
    song = from_tracks.pop(index)
    update_playlist_tracks(user_id, from_playlist_id, from_tracks)
    to_tracks = json.loads(to_pl['tracks'])
    to_tracks.append(song)
    update_playlist_tracks(user_id, to_playlist_id, to_tracks)
    return True

def copy_song_to_playlist(user_id, from_playlist_id, to_playlist_id, index):
    from_pl = get_playlist(user_id, from_playlist_id)
    to_pl = get_playlist(user_id, to_playlist_id)
    if not from_pl or not to_pl:
        return False
    from_tracks = json.loads(from_pl['tracks'])
    if index < 0 or index >= len(from_tracks):
        return False
    song = from_tracks[index]
    to_tracks = json.loads(to_pl['tracks'])
    to_tracks.append(song)
    update_playlist_tracks(user_id, to_playlist_id, to_tracks)
    return True

def lock_account(user_id, password):
    """Lock akun, hanya jika password benar. Kembalikan dict."""
    u = get_user(user_id)
    if not u or not _verify_password(password, u.get("password_hash", "")):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_old_password_wrong")}
    
    conn = get_conn()
    now = datetime.now().isoformat()
    conn.execute(
        "UPDATE users SET is_locked=1, locked_at=?, last_tracking_date=? WHERE id=?",
        (now, now, user_id)
    )
    conn.commit()
    conn.close()
    delete_all_session_tokens(user_id)
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_account_locked")}

def unlock_account(user_id, password):
    """Unlock akun, hanya jika password benar."""
    u = get_user(user_id)
    if not u or not _verify_password(password, u.get("password_hash", "")):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_old_password_wrong")}
    
    conn = get_conn()
    conn.execute("UPDATE users SET is_locked=0, locked_at=NULL WHERE id=?", (user_id,))
    conn.commit()
    conn.close()
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_account_unlocked")}

def is_account_locked(user_id):
    conn = get_conn()
    row = conn.execute("SELECT is_locked FROM users WHERE id=?", (user_id,)).fetchone()
    conn.close()
    return row["is_locked"] == 1 if row else False

def calculate_rank(user_id):
    """Hitung rank user berdasarkan overall progress."""
    u = get_user(user_id)
    if not u:
        return {
            "rank": 0,
            "rank_name_key": "rank_pemula",
            "rank_icon": "🥚",
            "rank_desc_key": "rank_desc_pemula",
            "score": 0,
            "progress": 0,
            "next_rank_name_key": "rank_penambang",
            "next_rank_icon": "⛏️",
            "has_next": True,
            "max_score": 100,
            "scores": {},
        }
    
    # Ambil data tambahan
    conn = get_conn()
    
    # Pet count
    pet_count = conn.execute("SELECT COUNT(*) FROM user_pets WHERE user_id=?", (user_id,)).fetchone()[0]
    
    # Item count
    item_count = conn.execute("SELECT COUNT(*) FROM inventory WHERE user_id=?", (user_id,)).fetchone()[0]
    
    # Achievement unlocked
    ach_count = conn.execute(
        "SELECT COUNT(*) FROM user_achievements WHERE user_id=? AND unlocked_at IS NOT NULL",
        (user_id,)
    ).fetchone()[0]
    
    # Boss killed
    boss_killed = conn.execute(
        "SELECT COUNT(*) FROM boss_battles bb "
        "JOIN guild_members gm ON bb.guild_id = gm.guild_id "
        "WHERE gm.user_id=? AND bb.status='defeated'",
        (user_id,)
    ).fetchone()[0]
    
    conn.close()
    
    # Hitung score (masing-masing max 100 poin)
    scores = {}
    
    # Level (max 100 poin, level 50 = 100 poin)
    scores['level'] = min(100, int(u.get('level', 1) * 2))
    
    # Total XP (max 100 poin, 100000 XP = 100 poin)
    total_xp = u.get('total_xp_earned', 0)
    scores['xp'] = min(100, int(total_xp / 1000))
    
    # Gold (max 100 poin, 100000 Gold = 100 poin)
    gold = u.get('gold', 0)
    scores['gold'] = min(100, int(gold / 1000))
    
    # Pet (max 100 poin, 10 pet = 100 poin)
    scores['pet'] = min(100, pet_count * 10)
    
    # Item (max 100 poin, 20 item = 100 poin)
    scores['item'] = min(100, item_count * 5)
    
    # Achievement (max 100 poin, 30 achievement = 100 poin)
    scores['achievement'] = min(100, int(ach_count * 3.33))
    
    # Rebirth (max 100 poin, 10 rebirth = 100 poin)
    rebirth = u.get('rebirth_count', 0)
    scores['rebirth'] = min(100, rebirth * 10)
    
    # Boss (max 100 poin, 10 boss = 100 poin)
    scores['boss'] = min(100, boss_killed * 10)
    
    # Sport Level (max 100 poin, level 20 = 100 poin)
    sport_level = u.get('sport_level', 1)
    scores['sport'] = min(100, sport_level * 5)
    
    # Total score (rata-rata dari semua kategori)
    total_score = int(sum(scores.values()) / len(scores))
    
    # --- 10 Rank ---
    RANKS = [
        {"min_score": 0,  "name_key": "rank_pemula",          "icon": "🥚", "desc_key": "rank_desc_pemula"},
        {"min_score": 10, "name_key": "rank_penambang",       "icon": "⛏️", "desc_key": "rank_desc_penambang"},
        {"min_score": 20, "name_key": "rank_penjelajah",      "icon": "🪓", "desc_key": "rank_desc_penjelajah"},
        {"min_score": 30, "name_key": "rank_petualang",       "icon": "⚔️", "desc_key": "rank_desc_petualang"},
        {"min_score": 40, "name_key": "rank_ksatria",         "icon": "🛡️", "desc_key": "rank_desc_ksatria"},
        {"min_score": 50, "name_key": "rank_veteran",         "icon": "⭐", "desc_key": "rank_desc_veteran"},
        {"min_score": 60, "name_key": "rank_legenda",         "icon": "🌟", "desc_key": "rank_desc_legenda"},
        {"min_score": 70, "name_key": "rank_raja",            "icon": "👑", "desc_key": "rank_desc_raja"},
        {"min_score": 80, "name_key": "rank_penguasa_naga",   "icon": "🐉", "desc_key": "rank_desc_penguasa_naga"},
        {"min_score": 90, "name_key": "rank_dewa",            "icon": "🌌", "desc_key": "rank_desc_dewa"},
    ]

    # Hitung rank saat ini
    current_rank = 0
    for i, rank in enumerate(RANKS):
        if total_score >= rank["min_score"]:
            current_rank = i

    rank_data = RANKS[current_rank]
    
    # Hitung progress ke rank berikutnya
    next_rank = current_rank + 1
    if next_rank < len(RANKS):
        min_score = RANKS[current_rank]["min_score"]
        next_min_score = RANKS[next_rank]["min_score"]
        progress = int(((total_score - min_score) / (next_min_score - min_score)) * 100) if next_min_score > min_score else 100
        next_rank_name_key = RANKS[next_rank]["name_key"]
        next_rank_icon = RANKS[next_rank]["icon"]
        has_next = True
    else:
        progress = 100
        next_rank_name_key = None
        next_rank_icon = "👑"
        has_next = False

    return {
        "rank": current_rank,
        "rank_name_key": rank_data["name_key"],
        "rank_icon": rank_data["icon"],
        "rank_desc_key": rank_data["desc_key"],
        "score": total_score,
        "progress": progress,
        "next_rank_name_key": next_rank_name_key,
        "next_rank_icon": next_rank_icon,
        "has_next": has_next,
        "max_score": 100,
        "scores": scores
    }


# ══════════════════════════════════════════════════════════════════════════════
#  v1.3.0 TAHAP 2 — Pomodoro, Heatmap, Recurrence, Trash (Undo), Template
# ══════════════════════════════════════════════════════════════════════════════

# ── Pomodoro / Focus Mode ─────────────────────────────────────────────────────

POMODORO_XP_PER_MIN = 1     # 1 XP per menit fokus
POMODORO_GOLD_DIV = 5       # 1 Gold per 5 menit fokus (minimal 2)


def complete_pomodoro(user_id, duration_minutes, task_name=""):
    """Selesaikan 1 sesi fokus → catat sesi + beri XP/Gold bonus."""
    try:
        minutes = int(duration_minutes)
    except (TypeError, ValueError):
        minutes = 0
    if minutes <= 0:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pomodoro_invalid")}

    xp = max(1, minutes * POMODORO_XP_PER_MIN)
    gold = max(2, minutes // POMODORO_GOLD_DIV)

    conn = get_conn()
    completed_at = local_now().isoformat()
    cur = conn.execute(
        "INSERT INTO pomodoro_sessions(user_id, task_name, duration_minutes,"
        " xp_earned, gold_earned, completed_at) VALUES(?,?,?,?,?,?)",
        (user_id, (task_name or "").strip(), minutes, xp, gold, completed_at)
    )
    session_id = cur.lastrowid
    conn.commit()
    conn.close()
    enqueue_productivity_event(user_id,"pomodoro",session_id,f"pomodoro:{session_id}",completed_at,
                               {"duration_minutes":minutes,"task_name":(task_name or "")[:120]})

    r = gain_xp_gold(user_id, xp, gold)
    r["session_minutes"] = minutes
    # Achievement: jumlah sesi (increment) & total menit fokus (absolut)
    check_achievements(user_id, "pomodoro_sessions", 1)
    check_achievements(user_id, "pomodoro_minutes", get_pomodoro_stats(user_id)["total_minutes"])
    log.info(f"Pomodoro selesai user_id={user_id}: {minutes} mnt, +{r['xp_gained']} XP, +{r['gold_gained']} gold")
    return r


def get_pomodoro_stats(user_id):
    """Statistik pomodoro: sesi & menit hari ini + total."""
    conn = get_conn()
    today = date.today().isoformat()
    row_today = conn.execute(
        "SELECT COUNT(*) c, COALESCE(SUM(duration_minutes),0) m FROM pomodoro_sessions"
        " WHERE user_id=? AND substr(completed_at,1,10)=?",
        (user_id, today)
    ).fetchone()
    row_total = conn.execute(
        "SELECT COUNT(*) c, COALESCE(SUM(duration_minutes),0) m FROM pomodoro_sessions"
        " WHERE user_id=?",
        (user_id,)
    ).fetchone()
    conn.close()
    return {
        "today_sessions": row_today["c"], "today_minutes": row_today["m"],
        "total_sessions": row_total["c"], "total_minutes": row_total["m"],
    }


def get_recent_pomodoros(user_id, limit=5):
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM pomodoro_sessions WHERE user_id=? ORDER BY id DESC LIMIT ?",
        (user_id, limit)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


# ── Heatmap aktivitas (GitHub-style) ──────────────────────────────────────────

def get_activity_heatmap(user_id, days=119):
    """Jumlah aktivitas 'success' per hari, `days` hari ke belakang.
    Return: {'YYYY-MM-DD': count, ...}"""
    cutoff = (date.today() - timedelta(days=days - 1)).isoformat()
    conn = get_conn()
    rows = conn.execute(
        "SELECT action_date, COUNT(*) c FROM task_history"
        " WHERE user_id=? AND action='success' AND action_date>=?"
        " GROUP BY action_date",
        (user_id, cutoff)
    ).fetchall()
    conn.close()
    return {r["action_date"]: r["c"] for r in rows}


# ── Recurrence fleksibel (jadwalkan habit/daily per hari) ────────────────────
# repeat_days: string "0,2,4" (0=Senin .. 6=Minggu); kosong = setiap hari.

_DAY_NAMES = {
    "id": ["Sen", "Sel", "Rab", "Kam", "Jum", "Sab", "Min"],
    "en": ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"],
}


def parse_repeat_days(s):
    """'0,2,4' -> {0, 2, 4}; kosong/invalid -> set() (artinya: tiap hari)."""
    out = set()
    for part in str(s or "").split(","):
        part = part.strip()
        if part.isdigit():
            d = int(part)
            if 0 <= d <= 6:
                out.add(d)
    return out


def repeat_days_to_str(days):
    """{0,2,4} -> '0,2,4' (terurut)."""
    return ",".join(str(d) for d in sorted(int(d) for d in days if 0 <= int(d) <= 6))


def is_due_on(repeat_days_str, day_date):
    """True jika item terjadwal pada tanggal `day_date` (datetime.date)."""
    days = parse_repeat_days(repeat_days_str)
    if not days:
        return True                      # kosong = setiap hari
    return day_date.weekday() in days


def is_due_today(repeat_days_str):
    return is_due_on(repeat_days_str, date.today())


def describe_repeat_days(repeat_days_str, lang="id"):
    """Untuk tampilan: '' → 'Tiap hari'; '0,2,4' → 'Sen, Rab, Jum'."""
    days = parse_repeat_days(repeat_days_str)
    names = _DAY_NAMES.get(lang, _DAY_NAMES["en"])
    if not days:
        return tr_db(lang=lang, key="db_recur_every_day")
    return ", ".join(names[d] for d in sorted(days))


# ── Trash bin (Undo delete) ───────────────────────────────────────────────────

TRASH_MAX_AGE_HOURS = 48

# Pemetaan tipe item → (tabel, kolom label nama)
_TRASH_TABLES = {
    "habit": ("habits", "name"),
    "daily": ("dailies", "name"),
    "todo":  ("todos", "name"),
    "note":  ("notes", "title"),
}


def _table_columns(conn, table):
    return {r[1] for r in conn.execute(f"PRAGMA table_info({table})").fetchall()}


def push_to_trash(user_id, item_type, payload: dict, item_name=""):
    """Simpan snapshot baris yang dihapus; return trash_id."""
    if item_type not in _TRASH_TABLES:
        return None
    purge_trash()
    conn = get_conn()
    cur = conn.execute(
        "INSERT INTO trash_bin(user_id, item_type, item_name, payload, deleted_at)"
        " VALUES(?,?,?,?,?)",
        (user_id, item_type, item_name or payload.get("name") or payload.get("title") or "",
         json.dumps(payload, ensure_ascii=False), local_now().isoformat())
    )
    trash_id = cur.lastrowid
    conn.commit()
    conn.close()
    log.info(f"Item {item_type} '{item_name}' masuk trash (id={trash_id})")
    return trash_id


def restore_from_trash(user_id, trash_id):
    """Kembalikan item dari trash ke tabel asalnya (dengan id baru)."""
    conn = get_conn()
    row = conn.execute(
        "SELECT * FROM trash_bin WHERE id=? AND user_id=?", (trash_id, user_id)
    ).fetchone()
    if not row:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_item_not_found")}

    item_type = row["item_type"]
    table = _TRASH_TABLES.get(item_type, (None,))[0]
    if not table:
        conn.close()
        return {"ok": False}

    try:
        data = json.loads(row["payload"])
    except (ValueError, TypeError):
        data = {}

    # Hanya kolom yang memang ada di tabel (id dibuang → auto-increment baru)
    cols = _table_columns(conn, table)
    data = {k: v for k, v in data.items() if k in cols and k != "id" and v is not None}
    if not data:
        conn.close()
        return {"ok": False}

    fields = ", ".join(data.keys())
    placeholders = ", ".join("?" for _ in data)
    conn.execute(f"INSERT INTO {table}({fields}) VALUES({placeholders})",
                 list(data.values()))
    conn.execute("DELETE FROM trash_bin WHERE id=?", (trash_id,))
    conn.commit()
    conn.close()
    log.info(f"Item {item_type} dikembalikan dari trash (trash_id={trash_id})")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_restore_success"),
            "item_type": item_type}


def delete_trash_item(user_id, trash_id):
    conn = get_conn()
    conn.execute("DELETE FROM trash_bin WHERE id=? AND user_id=?", (trash_id, user_id))
    conn.commit()
    conn.close()


def purge_trash(older_than_hours=TRASH_MAX_AGE_HOURS):
    """Hapus item trash yang lebih tua dari `older_than_hours` jam."""
    try:
        cutoff = (datetime.now() - timedelta(hours=older_than_hours)).isoformat()
        conn = get_conn()
        conn.execute("DELETE FROM trash_bin WHERE deleted_at < ?", (cutoff,))
        conn.commit()
        conn.close()
    except Exception as e:
        log.warning(f"Purge trash gagal: {e}")


def _stash_before_delete(conn, user_id, table, row_id, item_type):
    """Serialisasi baris ke trash_bin TEPAT sebelum dihapus (pakai koneksi aktif)."""
    row = conn.execute(f"SELECT * FROM {table} WHERE id=? AND user_id=?",
                       (row_id, user_id)).fetchone()
    if not row:
        return None
    payload = dict(row)
    name = payload.get(_TRASH_TABLES[item_type][1], "")
    cur = conn.execute(
        "INSERT INTO trash_bin(user_id, item_type, item_name, payload, deleted_at)"
        " VALUES(?,?,?,?,?)",
        (user_id, item_type, name, json.dumps(payload, ensure_ascii=False),
         local_now().isoformat())
    )
    return cur.lastrowid


# ── Template habit siap pakai ────────────────────────────────────────────────
# items: (nama_id, nama_en, ikon, difficulty)

HABIT_TEMPLATES = [
    {
        "key": "morning_routine", "icon": "🌅",
        "name": ("Rutinitas Pagi", "Morning Routine"),
        "desc": ("Bangun pagi penuh energi dan mulai hari dengan benar",
                 "Wake up energized and start the day right"),
        "items": [
            ("Bangun sebelum jam 6", "Wake up before 6 AM", "⏰", "hard"),
            ("Minum air putih segelas", "Drink a glass of water", "💧", "easy"),
            ("Olahraga ringan 10 menit", "10 minutes light exercise", "🏃", "medium"),
            ("Sarapan sehat", "Healthy breakfast", "🍳", "easy"),
            ("Rencanakan 3 prioritas hari ini", "Plan today's top 3 priorities", "🎯", "medium"),
        ],
    },
    {
        "key": "healthy_life", "icon": "💪",
        "name": ("Hidup Sehat", "Healthy Life"),
        "desc": ("Kebiasaan kecil untuk tubuh yang lebih sehat",
                 "Small habits for a healthier body"),
        "items": [
            ("Minum 8 gelas air", "Drink 8 glasses of water", "💧", "medium"),
            ("Makan sayur & buah", "Eat vegetables & fruits", "🥗", "easy"),
            ("Jalan 5.000 langkah", "Walk 5,000 steps", "👟", "medium"),
            ("Tidur sebelum jam 11 malam", "Sleep before 11 PM", "💤", "hard"),
            ("Tanpa gula berlebih", "No excessive sugar", "🚫", "hard"),
        ],
    },
    {
        "key": "coding_study", "icon": "💻",
        "name": ("Belajar Coding", "Coding Study"),
        "desc": ("Konsisten belajar ngoding setiap hari",
                 "Consistent daily coding practice"),
        "items": [
            ("Ngoding minimal 1 jam", "Code at least 1 hour", "💻", "medium"),
            ("Baca dokumentasi / artikel teknis", "Read docs / tech articles", "📚", "easy"),
            ("Kerjakan 1 soal latihan", "Solve 1 practice problem", "🧩", "medium"),
            ("Review & refactor kode kemarin", "Review & refactor yesterday's code", "🔍", "easy"),
            ("Push/commit proyek pribadi", "Push/commit personal project", "🚀", "medium"),
        ],
    },
    {
        "key": "mindfulness", "icon": "🧘",
        "name": ("Ketenangan Pikiran", "Mindfulness"),
        "desc": ("Rawat kesehatan mental dengan momen tenang setiap hari",
                 "Care for your mental health with daily calm moments"),
        "items": [
            ("Meditasi 5 menit", "5 minutes of meditation", "🧘", "easy"),
            ("Jurnal syukur (3 hal)", "Gratitude journal (3 things)", "📓", "easy"),
            ("Tanpa media sosial 1 jam", "1 hour without social media", "📵", "hard"),
            ("Tarik napas dalam saat stres", "Deep breaths when stressed", "🌬️", "easy"),
        ],
    },
    {
        "key": "student", "icon": "🎓",
        "name": ("Pelajar Rajin", "Diligent Student"),
        "desc": ("Bangun kebiasaan belajar yang konsisten",
                 "Build consistent study habits"),
        "items": [
            ("Belajar fokus 30 menit", "30 minutes focused study", "📖", "medium"),
            ("Catat materi penting", "Take notes on key material", "✏️", "easy"),
            ("Review pelajaran hari ini", "Review today's lessons", "🔁", "easy"),
            ("Kerjakan PR/tugas tepat waktu", "Finish homework on time", "✅", "medium"),
        ],
    },
]


def get_habit_templates(lang="id"):
    """Daftar template untuk ditampilkan di UI."""
    out = []
    for t in HABIT_TEMPLATES:
        out.append({
            "key": t["key"], "icon": t["icon"],
            "name": t["name"][0] if lang == "id" else t["name"][1],
            "desc": t["desc"][0] if lang == "id" else t["desc"][1],
            "count": len(t["items"]),
        })
    return out


def apply_habit_template(user_id, template_key):
    """Buat semua habit dalam template. Return jumlah habit yang dibuat."""
    template = next((t for t in HABIT_TEMPLATES if t["key"] == template_key), None)
    if not template:
        return 0
    u = get_user(user_id)
    lang = u.get("language", "id")
    idx = 0 if lang == "id" else 1
    for item in template["items"]:
        add_habit(user_id, item[idx], item[2], item[3])
    log.info(f"Template '{template_key}' diterapkan user_id={user_id} ({len(template['items'])} habit)")
    return len(template["items"])


# ══════════════════════════════════════════════════════════════════════════════
#  v1.3.0 TAHAP 3 — Crafting, Enchanting, Daily Login, Titles, Seasonal Events
# ══════════════════════════════════════════════════════════════════════════════

# ── Helper internal: beri item ke inventory ──────────────────────────────────

def _give_item(user_id, item_id, quantity=1):
    """Tambah item ke inventory (stack untuk consumable, unik untuk equipment)."""
    item = SHOP_ITEMS.get(item_id)
    if not item:
        return False
    conn = get_conn()
    ex = conn.execute(
        "SELECT * FROM inventory WHERE user_id=? AND item_id=?",
        (user_id, item_id)).fetchone()
    if ex:
        conn.execute("UPDATE inventory SET quantity=quantity+? WHERE id=?",
                     (quantity, ex["id"]))
    else:
        conn.execute(
            "INSERT INTO inventory(user_id, item_id, item_type, quantity) VALUES(?,?,?,?)",
            (user_id, item_id, item["type"], quantity))
    conn.commit()
    conn.close()
    return True


# ── 🔨 CRAFTING ───────────────────────────────────────────────────────────────

CRAFTING_RECIPES = {
    "bedrock_sword": {
        "output": "bedrock_sword",
        "inputs": ["netherite_sword", "diamond_sword"],
        "gold": 500,
        "desc": ("Lebur dua pedang legendaris menjadi satu",
                 "Fuse two legendary swords into one"),
    },
    "phantom_wings": {
        "output": "phantom_wings",
        "inputs": ["elytra", "golden_boots"],
        "gold": 400,
        "desc": ("Sayap angin yang juga membawa keberuntungan",
                 "Wings of wind that also bring fortune"),
    },
    "aegis_of_void": {
        "output": "aegis_of_void",
        "inputs": ["tower_shield", "diamond_chestplate"],
        "gold": 600,
        "desc": ("Pertahanan pamungkas hasil tempaan",
                 "Ultimate defense, forged"),
    },

    # ── 🆕 v1.3.0 — resep baru ──
    "inferno_blade": {
        "output": "inferno_blade",
        "inputs": ["netherite_sword", "blaze_rod"],
        "gold": 700,
        "desc": ("Tempa pedang legendaris dalam api nether",
                 "Forge a legendary blade in nether fire"),
    },
    "healers_blessing": {
        "output": "healers_blessing",
        "inputs": ["golden_apple", "greater_health_potion", "elixir"],
        "gold": 300,
        "desc": ("Suling ramuan kehidupan menjadi pelindung suci",
                 "Distill life potions into a sacred ward"),
    },
    "gilded_compass": {
        "output": "gilded_compass",
        "inputs": ["compass", "golden_apple"],
        "gold": 450,
        "desc": ("Kompas yang disepuh emas keberuntungan",
                 "A compass gilded with fortune's gold"),
    },
    # ── 🆕 v1.4.0 — Resep baru (grindy but rewarding) ──
    "frost_guard": {
        "output": "frost_guard",
        "inputs": ["turtle_shell", "ice_block", "wind_cloak"],
        "gold": 550,
        "desc": ("Perisai es yang kokoh untuk pertahanan sempurna",
                 "Sturdy ice shield for perfect defense"),
    },
    "scholar_crown": {
        "output": "scholar_crown",
        "inputs": ["scholar_tome", "golden_boots", "compass"],
        "gold": 500,
        "desc": ("Mahkota kebijaksanaan penambah pengalaman",
                 "Crown of wisdom boosting experience"),
    },
    "void_core": {
        "output": "void_core",
        "inputs": ["ender_pearl", "nether_star", "blaze_rod"],
        "gold": 800,
        "desc": ("Inti kehampaan yang menyerap kekuatan",
                 "Void heart absorbing all power"),
    },
    "ember_charm": {
        "output": "ember_charm",
        "inputs": ["blaze_rod", "honey_bottle", "lucky_charm"],
        "gold": 400,
        "desc": ("Jimat bara hangat pembawa keberuntungan",
                 "Warm ember charm of fortune"),
    },
}


def get_crafting_recipes():
    """Daftar resep untuk UI: output + inputs + detail item."""
    out = []
    for rid, r in CRAFTING_RECIPES.items():
        o = SHOP_ITEMS.get(r["output"], {})
        out.append({
            "id": rid,
            "output_id": r["output"],
            "output_name": o.get("name", r["output"]),
            "output_icon": o.get("icon", "🔨"),
            "output_buff": o.get("buff_desc", ""),
            "inputs": [{"id": iid,
                        "name": SHOP_ITEMS.get(iid, {}).get("name", iid),
                        "icon": SHOP_ITEMS.get(iid, {}).get("icon", "❔")}
                       for iid in r["inputs"]],
            "gold": r["gold"],
            "desc": r["desc"],
        })
    return out


def can_craft(user_id, recipe_id):
    """Cek kelengkapan: {ok, missing:[item_id...], need_gold, have_gold}."""
    r = CRAFTING_RECIPES.get(recipe_id)
    if not r:
        return {"ok": False, "missing": [], "need_gold": 0, "have_gold": 0}
    inv = {row["item_id"]: row["quantity"] for row in get_inventory(user_id)}
    missing = [iid for iid in r["inputs"] if inv.get(iid, 0) < 1]
    u = get_user(user_id)
    have_gold = u.get("gold", 0)
    gold_ok = have_gold >= r["gold"]
    return {"ok": not missing and gold_ok, "missing": missing,
            "need_gold": r["gold"], "have_gold": have_gold, "gold_ok": gold_ok}


@retry_on_lock
def craft_item(user_id, recipe_id):
    """Craft: konsumsi input + gold → hasilkan output. Max 1 per resep."""
    r = CRAFTING_RECIPES.get(recipe_id)
    if not r:
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_item_not_found")}
    output = r["output"]

    # Sudah punya output-nya? (equipment unik)
    conn = get_conn()
    ex = conn.execute(
        "SELECT id FROM inventory WHERE user_id=? AND item_id=?",
        (user_id, output)).fetchone()
    if ex:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_craft_already_owned")}

    chk = can_craft(user_id, recipe_id)
    if not chk["ok"]:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_craft_missing")}

    # Konsumsi 1 dari tiap input
    for iid in r["inputs"]:
        row = conn.execute(
            "SELECT * FROM inventory WHERE user_id=? AND item_id=?",
            (user_id, iid)).fetchone()
        if row["quantity"] <= 1:
            conn.execute("DELETE FROM inventory WHERE id=?", (row["id"],))
        else:
            conn.execute("UPDATE inventory SET quantity=quantity-1 WHERE id=?",
                         (row["id"],))

    # Potong gold + beri output + counter untuk title Blacksmith
    u = get_user(user_id)
    conn.execute("UPDATE users SET gold=gold-?, total_crafts=total_crafts+1 WHERE id=?",
                 (r["gold"], user_id))
    o = SHOP_ITEMS[output]
    conn.execute(
        "INSERT INTO inventory(user_id, item_id, item_type, quantity) VALUES(?,?,?,1)",
        (user_id, output, o["type"]))
    conn.commit()
    conn.close()

    recalculate_all_buffs(user_id)
    add_notification(user_id, tr_db(user_id=user_id, key="db_craft_success",
                                    icon=o["icon"], name=o["name"]), "success")
    # Achievement: jumlah item yang sudah ditempa (absolut)
    check_achievements(user_id, "craft_count",
                       (get_user(user_id) or {}).get("total_crafts", 0))
    log.info(f"Craft sukses user_id={user_id}: {output}")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_craft_success",
                                     icon=o["icon"], name=o["name"])}


# ── ✨ ENCHANTING (upgrade equipment memakai XP, ala Minecraft) ──────────────

ENCHANT_MAX_LEVEL = 5
ENCHANT_BASE_XP = 50          # biaya = (level_sekarang+1) × 50 XP
ENCHANT_BUFF_BONUS = 0.12     # +12% kekuatan buff per level enchant


def enchant_cost(current_level):
    """Biaya XP untuk enchant ke level berikutnya."""
    return (current_level + 1) * ENCHANT_BASE_XP


def get_enchant_level(user_id, item_id):
    conn = get_conn()
    row = conn.execute(
        "SELECT enchant_level FROM inventory WHERE user_id=? AND item_id=?",
        (user_id, item_id)).fetchone()
    conn.close()
    return row["enchant_level"] if row else 0


@retry_on_lock
def enchant_item(user_id, item_id):
    """Naikkan enchant level item equipment sebesar 1, bayar dengan XP."""
    item = SHOP_ITEMS.get(item_id)
    if not item or item["type"] == "consumable":
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_enchant_not_allowed")}

    conn = get_conn()
    inv = conn.execute(
        "SELECT * FROM inventory WHERE user_id=? AND item_id=?",
        (user_id, item_id)).fetchone()
    if not inv:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_item_not_found")}

    lvl = inv["enchant_level"] or 0
    if lvl >= ENCHANT_MAX_LEVEL:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_enchant_max",
                                          max=ENCHANT_MAX_LEVEL)}

    cost = enchant_cost(lvl)
    u = get_user(user_id)
    if u.get("xp", 0) < cost:
        conn.close()
        return {"ok": False, "cost": cost,
                "msg": tr_db(user_id=user_id, key="db_enchant_no_xp", cost=cost)}

    conn.execute("UPDATE users SET xp=xp-? WHERE id=?", (cost, user_id))
    conn.execute("UPDATE inventory SET enchant_level=? WHERE id=?", (lvl + 1, inv["id"]))
    conn.commit()
    conn.close()

    recalculate_all_buffs(user_id)
    log.info(f"Enchant user_id={user_id}: {item_id} → +{lvl+1} (biaya {cost} XP)")
    return {"ok": True, "new_level": lvl + 1, "cost": cost,
            "msg": tr_db(user_id=user_id, key="db_enchant_success",
                         icon=item["icon"], name=item["name"], lvl=lvl + 1)}


def get_enchant_map(user_id):
    """{item_id: enchant_level} — dipakai recalculate_all_buffs."""
    conn = get_conn()
    rows = conn.execute(
        "SELECT item_id, enchant_level FROM inventory WHERE user_id=?",
        (user_id,)).fetchall()
    conn.close()
    out = {}
    for r in rows:
        out[r["item_id"]] = max(out.get(r["item_id"], 0), r["enchant_level"] or 0)
    return out


# ── 🎁 DAILY LOGIN REWARD ────────────────────────────────────────────────────

DAILY_REWARDS = [
    {"xp": 20,  "gold": 10},                          # hari 1
    {"xp": 30,  "gold": 15},                          # hari 2
    {"xp": 40,  "gold": 20},                          # hari 3
    {"xp": 50,  "gold": 30},                          # hari 4
    {"xp": 60,  "gold": 40, "item": "health_potion"}, # hari 5
    {"xp": 75,  "gold": 50},                          # hari 6
    {"xp": 100, "gold": 75, "item": "golden_apple"},  # hari 7 (siklus ulang)
]


@retry_on_lock
def claim_daily_login(user_id):
    """Klaim hadiah login harian (1×/hari). Streak naik jika berurutan,
    reset ke 1 jika putus, berputar 7→1 setelah seminggu penuh."""
    today = date.today()
    today_str = today.isoformat()
    u = get_user(user_id)
    last = u.get("last_login_reward") or ""
    if last == today_str:
        return {"claimed": False, "streak": u.get("login_streak", 0)}

    streak = u.get("login_streak", 0) or 0
    if last:
        try:
            last_date = date.fromisoformat(last)
            if last_date == today - timedelta(days=1):
                streak = streak + 1 if streak < 7 else 1   # wrap setelah hari ke-7
            else:
                streak = 1
        except ValueError:
            streak = 1
    else:
        streak = 1

    reward = DAILY_REWARDS[streak - 1]
    conn = get_conn()
    conn.execute("UPDATE users SET last_login_reward=?, login_streak=? WHERE id=?",
                 (today_str, streak, user_id))
    conn.commit()
    conn.close()

    r = gain_xp_gold(user_id, reward["xp"], reward["gold"])
    item_given = reward.get("item")
    if item_given:
        _give_item(user_id, item_given, 1)

    log.info(f"Daily login user_id={user_id}: hari ke-{streak}, "
             f"+{reward['xp']} XP, +{reward['gold']} G, item={item_given}")
    return {"claimed": True, "streak": streak,
            "xp": r.get("xp_gained", reward["xp"]),
            "gold": r.get("gold_gained", reward["gold"]),
            "item": item_given, "leveled_up": r.get("leveled_up", False),
            "new_level": r.get("new_level", 0)}


# ── 🎖️ TITLES / BADGES (terbuka otomatis dari statistik) ────────────────────
# req: (field, nilai_minimal). Field khusus dihitung di get_unlocked_titles.

TITLES = [
    {"key": "novice",    "name": ("🌱 Pemula", "🌱 Novice"),
     "req": None},
    {"key": "streak7",   "name": ("🔥 Pejuang Konsisten", "🔥 Consistent Fighter"),
     "req": ("longest_streak", 7)},
    {"key": "pomodoro",  "name": ("🍅 Ninja Fokus", "🍅 Focus Ninja"),
     "req": ("pomodoro_sessions", 10)},
    {"key": "habit50",   "name": ("💪 Master Habit", "💪 Habit Master"),
     "req": ("total_habits_done", 50)},
    {"key": "level10",   "name": ("⭐ Petarung Berpengalaman", "⭐ Seasoned Fighter"),
     "req": ("level", 10)},
    {"key": "blacksmith","name": ("🔨 Pandai Besi", "🔨 Blacksmith"),
     "req": ("total_crafts", 1)},
    {"key": "gold10k",   "name": ("💰 Sultan", "💰 Tycoon"),
     "req": ("total_gold_earned", 10000)},
    {"key": "habit200",  "name": ("⚡ Legenda Habit", "⚡ Habit Legend"),
     "req": ("total_habits_done", 200)},
    {"key": "level25",   "name": ("🌟 Master Dunia", "🌟 World Master"),
     "req": ("level", 25)},
    {"key": "tasks500",  "name": ("👑 Raja Produktivitas", "👑 Productivity King"),
     "req": ("total_tasks_completed", 500)},
]


def get_unlocked_titles(user_id):
    """Daftar title yang sudah terbuka: [{key, name, unlocked, current, target}]."""
    u = get_user(user_id)
    lang = u.get("language", "id")
    stats = {
        "longest_streak": u.get("longest_streak", 0) or 0,
        "total_habits_done": u.get("total_habits_done", 0) or 0,
        "total_gold_earned": u.get("total_gold_earned", 0) or 0,
        "total_tasks_completed": u.get("total_tasks_completed", 0) or 0,
        "total_crafts": u.get("total_crafts", 0) or 0,
        "level": u.get("level", 1) or 1,
        "pomodoro_sessions": get_pomodoro_stats(user_id)["total_sessions"],
    }
    out = []
    for t in TITLES:
        idx = 0 if lang == "id" else 1
        req = t["req"]
        if req is None:
            unlocked, current, target = True, 1, 0
        else:
            field, target = req
            current = stats.get(field, 0)
            unlocked = current >= target
        out.append({"key": t["key"], "name": t["name"][idx],
                    "unlocked": unlocked, "current": current, "target": target})
    return out


def set_title(user_id, title_key):
    """Pilih title yang tampil di profil & leaderboard."""
    keys = {t["key"] for t in get_unlocked_titles(user_id) if t["unlocked"]}
    if title_key not in keys and title_key != "":
        return {"ok": False}
    update_user(user_id, selected_title=title_key)
    return {"ok": True}


def get_title_display(user_id, lang="id"):
    """Nama title terpilih (untuk profil/leaderboard)."""
    u = get_user(user_id)
    key = u.get("selected_title") or "novice"
    for t in TITLES:
        if t["key"] == key:
            return t["name"][0] if lang == "id" else t["name"][1]
    return ""


# ── 🌙 SEASONAL EVENTS ───────────────────────────────────────────────────────

SEASONAL_EVENTS = {
    "ramadan": {
        "name": ("Ramadan & Lebaran", "Ramadan & Eid"), "icon": "🌙",
        "start": "02-15", "end": "04-15",
        "bosses": ["ketupat_golem"], "items": ["ketupat_feast"],
    },
    "halloween": {
        "name": ("Halloween", "Halloween"), "icon": "🎃",
        "start": "10-25", "end": "11-02",
        "bosses": ["pumpkin_king"], "items": ["candy_bag"],
    },
    "christmas": {
        "name": ("Natal", "Christmas"), "icon": "🎄",
        "start": "12-18", "end": "12-31",
        "bosses": ["krampus"], "items": ["snowball_fight"],
    },
}


def _md_in_window(md, start, end):
    """Cek 'MM-DD' dalam rentang (mendukung rentang yang melewati tahun baru)."""
    if start <= end:
        return start <= md <= end
    return md >= start or md <= end      # wrap, mis. 12-18 s/d 01-05


def get_active_seasonal_events(on_date=None):
    """Event yang aktif pada tanggal tertentu (default: hari ini)."""
    d = on_date or date.today()
    md = d.strftime("%m-%d")
    out = []
    for eid, e in SEASONAL_EVENTS.items():
        if _md_in_window(md, e["start"], e["end"]):
            out.append({"id": eid, "name": e["name"], "icon": e["icon"],
                        "bosses": e["bosses"], "items": e["items"]})
    return out


def _event_window_active(event_id, on_date=None):
    if not event_id or event_id not in SEASONAL_EVENTS:
        return False
    d = on_date or date.today()
    e = SEASONAL_EVENTS[event_id]
    return _md_in_window(d.strftime("%m-%d"), e["start"], e["end"])


def is_boss_available(boss_id, on_date=None):
    """Boss seasonal hanya bisa dilawan saat event-nya aktif; custom bebas."""
    if str(boss_id).startswith("custom_"):
        return get_custom_boss(boss_id) is not None
    boss = BOSSES.get(boss_id)
    if not boss:
        return False
    ev = boss.get("seasonal_event")
    if not ev:
        return True
    return _event_window_active(ev, on_date)


def is_shop_item_visible(item_id, on_date=None):
    """Item craft_only TIDAK tampil di shop; item seasonal hanya saat event."""
    item = SHOP_ITEMS.get(item_id, {})
    if item.get("craft_only"):
        return False
    ev = item.get("seasonal")
    if ev and not _event_window_active(ev, on_date):
        return False
    return True


# ══════════════════════════════════════════════════════════════════════════════
#  v1.3.0 TAHAP 4 — Insights, Korelasi, Wrapped, Talents, Custom Boss, PvP
# ══════════════════════════════════════════════════════════════════════════════

# ── 💡 INSIGHTS OTOMATIS ─────────────────────────────────────────────────────

_WEEKDAY_FULL = {
    "id": ["Senin", "Selasa", "Rabu", "Kamis", "Jumat", "Sabtu", "Minggu"],
    "en": ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"],
}


def get_insights(user_id):
    """Statistik insight otomatis untuk dashboard."""
    conn = get_conn()
    u = get_user(user_id)
    lang = u.get("language", "id")

    # Hari paling produktif (jumlah success per weekday, 90 hari terakhir)
    cutoff = (date.today() - timedelta(days=90)).isoformat()
    rows = conn.execute(
        "SELECT action_date, COUNT(*) c FROM task_history"
        " WHERE user_id=? AND action='success' AND action_date>=? GROUP BY action_date",
        (user_id, cutoff)
    ).fetchall()
    conn.close()

    weekday_counts = [0] * 7
    best_day, best_count = None, 0
    for r in rows:
        try:
            d = date.fromisoformat(r["action_date"])
        except ValueError:
            continue
        weekday_counts[d.weekday()] += r["c"]
        if r["c"] > best_count:
            best_day, best_count = r["action_date"], r["c"]

    top_wd = max(range(7), key=lambda i: weekday_counts[i]) if any(weekday_counts) else None
    names = _WEEKDAY_FULL.get(lang, _WEEKDAY_FULL["en"])

    # Daily yang sedang ber-streak
    conn = get_conn()
    active_streaks = conn.execute(
        "SELECT COUNT(*) c FROM dailies WHERE user_id=? AND streak>0", (user_id,)
    ).fetchone()["c"]
    conn.close()

    return {
        "top_weekday": names[top_wd] if top_wd is not None else None,
        "top_weekday_count": weekday_counts[top_wd] if top_wd is not None else 0,
        "best_day": best_day, "best_day_count": best_count,
        "longest_streak": u.get("longest_streak", 0) or 0,
        "active_streaks": active_streaks,
        "focus_minutes": get_pomodoro_stats(user_id)["total_minutes"],
        "has_data": any(weekday_counts),
    }


# ── 📊 KORELASI KESEHATAN ↔ PRODUKTIVITAS ────────────────────────────────────

def pearson(xs, ys):
    """Koefisien korelasi Pearson sederhana; 0 jika data tidak cukup/konstan."""
    n = len(xs)
    if n < 3 or n != len(ys):
        return 0.0
    mx, my = sum(xs) / n, sum(ys) / n
    num = sum((x - mx) * (y - my) for x, y in zip(xs, ys))
    den_x = sum((x - mx) ** 2 for x in xs) ** 0.5
    den_y = sum((y - my) ** 2 for y in ys) ** 0.5
    if den_x == 0 or den_y == 0:
        return 0.0
    return num / (den_x * den_y)


def get_health_productivity_series(user_id, days=30):
    """Seri harian (tanggal, jam tidur, task selesai) + korelasi Pearson."""
    cutoff = (date.today() - timedelta(days=days - 1)).isoformat()
    conn = get_conn()
    health = conn.execute(
        "SELECT log_date, sleep_hours FROM health_logs"
        " WHERE user_id=? AND log_date>=?", (user_id, cutoff)
    ).fetchall()
    tasks = conn.execute(
        "SELECT action_date, COUNT(*) c FROM task_history"
        " WHERE user_id=? AND action='success' AND action_date>=?"
        " GROUP BY action_date", (user_id, cutoff)
    ).fetchall()
    conn.close()

    sleep_map = {r["log_date"]: r["sleep_hours"] or 0 for r in health}
    task_map = {r["action_date"]: r["c"] for r in tasks}

    series = []
    for i in range(days):
        d = (date.today() - timedelta(days=days - 1 - i)).isoformat()
        series.append((d, sleep_map.get(d, 0.0), task_map.get(d, 0)))

    # Korelasi hanya di hari yang punya data tidur
    pairs = [(s, t) for _, s, t in series if s > 0]
    xs = [p[0] for p in pairs]
    ys = [p[1] for p in pairs]
    return {"series": series, "correlation": pearson(xs, ys),
            "days_with_sleep_data": len(pairs)}


# ── 🎁 YEAR WRAPPED ──────────────────────────────────────────────────────────

def get_year_wrapped(user_id, year=None):
    """Ringkasan setahun ala 'Wrapped'. Sumber: task_history, pomodoro, ekonomi."""
    year = year or date.today().year
    yprefix = f"{year:04d}-"
    conn = get_conn()

    # Penyelesaian per tipe
    done_by_type = {}
    for r in conn.execute(
        "SELECT task_type, COUNT(*) c FROM task_history"
        " WHERE user_id=? AND action='success' AND substr(action_date,1,5)=?"
        " GROUP BY task_type", (user_id, yprefix)
    ).fetchall():
        done_by_type[r["task_type"]] = r["c"]

    # Hari aktif + hari terbaik
    days = conn.execute(
        "SELECT action_date, COUNT(*) c FROM task_history"
        " WHERE user_id=? AND action='success' AND substr(action_date,1,5)=?"
        " GROUP BY action_date ORDER BY c DESC", (user_id, yprefix)
    ).fetchall()
    conn.close()

    # Top 3 habit paling sering
    conn = get_conn()
    top = conn.execute(
        "SELECT th.task_id, COUNT(*) c FROM task_history th"
        " WHERE th.user_id=? AND th.action='success' AND th.task_type='habit'"
        " AND substr(th.action_date,1,5)=? GROUP BY th.task_id ORDER BY c DESC LIMIT 3",
        (user_id, yprefix)
    ).fetchall()
    top_habits = []
    for t in top:
        name_row = conn.execute("SELECT name, icon FROM habits WHERE id=?", (t["task_id"],)).fetchone()
        if name_row:
            top_habits.append({"name": name_row["name"], "icon": name_row["icon"], "count": t["c"]})
    pomo = conn.execute(
        "SELECT COUNT(*) c, COALESCE(SUM(duration_minutes),0) m FROM pomodoro_sessions"
        " WHERE user_id=? AND substr(completed_at,1,5)=?", (user_id, yprefix)
    ).fetchone()
    eco = conn.execute(
        "SELECT type, COALESCE(SUM(amount),0) s FROM economy_items"
        " WHERE user_id=? AND substr(date,1,5)=? GROUP BY type", (user_id, yprefix)
    ).fetchall()
    conn.close()

    eco_map = {r["type"]: r["s"] for r in eco}
    u = get_user(user_id)
    total_done = sum(done_by_type.values())
    return {
        "year": year,
        "total_done": total_done,
        "by_type": done_by_type,
        "active_days": len(days),
        "best_day": days[0]["action_date"] if days else None,
        "best_day_count": days[0]["c"] if days else 0,
        "top_habits": top_habits,
        "focus_sessions": pomo["c"],
        "focus_minutes": pomo["m"],
        "income": eco_map.get("income", 0),
        "expense": eco_map.get("expense", 0),
        "level": u.get("level", 1),
        "longest_streak": u.get("longest_streak", 0) or 0,
    }


# ── 🌳 TALENT TREE PER CLASS ─────────────────────────────────────────────────

TALENTS = {
    # Warrior ⚔️
    "warrior_t1_skin":    {"class": "warrior", "tier": 1, "icon": "🛡️",
                           "name": ("Iron Skin", "Iron Skin"), "buff": {"hp_reduc": 4},
                           "desc": ("-4 HP dari tiap kerusakan", "-4 HP from every hit")},
    "warrior_t1_rage":    {"class": "warrior", "tier": 1, "icon": "💢",
                           "name": ("Berserker", "Berserker"), "buff": {"boss_dmg": 4},
                           "desc": ("+4 Boss DMG", "+4 Boss DMG")},
    "warrior_t2_chest":   {"class": "warrior", "tier": 2, "icon": "💰",
                           "name": ("War Chest", "War Chest"), "buff": {"gold_pct": 8},
                           "desc": ("+8% Gold", "+8% Gold")},
    "warrior_t2_veteran": {"class": "warrior", "tier": 2, "icon": "🎖️",
                           "name": ("Veteran", "Veteran"), "buff": {"xp_pct": 8},
                           "desc": ("+8% XP", "+8% XP")},
    "warrior_t3_jugg":    {"class": "warrior", "tier": 3, "icon": "🏰",
                           "name": ("Juggernaut", "Juggernaut"), "buff": {"hp_reduc": 10},
                           "desc": ("-10 HP dari tiap kerusakan", "-10 HP from every hit")},
    "warrior_t3_slayer":  {"class": "warrior", "tier": 3, "icon": "⚔️",
                           "name": ("Slayer", "Slayer"), "buff": {"boss_dmg": 10},
                           "desc": ("+10 Boss DMG", "+10 Boss DMG")},
    # Mage 🧙
    "mage_t1_mana":       {"class": "mage", "tier": 1, "icon": "🔮",
                           "name": ("Mana Well", "Mana Well"), "buff": {"mp_bonus": 10},
                           "desc": ("+10 Max MP", "+10 Max MP")},
    "mage_t1_mind":       {"class": "mage", "tier": 1, "icon": "📖",
                           "name": ("Arcane Mind", "Arcane Mind"), "buff": {"xp_pct": 6},
                           "desc": ("+6% XP", "+6% XP")},
    "mage_t2_alch":       {"class": "mage", "tier": 2, "icon": "⚗️",
                           "name": ("Alchemist", "Alchemist"), "buff": {"gold_pct": 8},
                           "desc": ("+8% Gold", "+8% Gold")},
    "mage_t2_flow":       {"class": "mage", "tier": 2, "icon": "🌊",
                           "name": ("Focus Flow", "Focus Flow"), "buff": {"mp_bonus": 15},
                           "desc": ("+15 Max MP", "+15 Max MP")},
    "mage_t3_arch":       {"class": "mage", "tier": 3, "icon": "🌟",
                           "name": ("Archmage", "Archmage"), "buff": {"xp_pct": 15},
                           "desc": ("+15% XP", "+15% XP")},
    "mage_t3_hex":        {"class": "mage", "tier": 3, "icon": "🪄",
                           "name": ("Hexer", "Hexer"), "buff": {"boss_dmg": 6},
                           "desc": ("+6 Boss DMG", "+6 Boss DMG")},
    # Archer 🏹
    "archer_t1_keen":     {"class": "archer", "tier": 1, "icon": "🎯",
                           "name": ("Keen Eye", "Keen Eye"), "buff": {"xp_pct": 6},
                           "desc": ("+6% XP", "+6% XP")},
    "archer_t1_scout":    {"class": "archer", "tier": 1, "icon": "🥾",
                           "name": ("Scout", "Scout"), "buff": {"gold_pct": 6},
                           "desc": ("+6% Gold", "+6% Gold")},
    "archer_t2_tracker":  {"class": "archer", "tier": 2, "icon": "🐾",
                           "name": ("Tracker", "Tracker"), "buff": {"boss_dmg": 5},
                           "desc": ("+5 Boss DMG", "+5 Boss DMG")},
    "archer_t2_quiver":   {"class": "archer", "tier": 2, "icon": "🏹",
                           "name": ("Endless Quiver", "Endless Quiver"), "buff": {"gold_pct": 8},
                           "desc": ("+8% Gold", "+8% Gold")},
    "archer_t3_deadeye":  {"class": "archer", "tier": 3, "icon": "☄️",
                           "name": ("Deadeye", "Deadeye"), "buff": {"boss_dmg": 12},
                           "desc": ("+12 Boss DMG", "+12 Boss DMG")},
    "archer_t3_ranger":   {"class": "archer", "tier": 3, "icon": "🌲",
                           "name": ("Ranger Lord", "Ranger Lord"), "buff": {"xp_pct": 15},
                           "desc": ("+15% XP", "+15% XP")},
    # Healer ❤️ — defensive/support
    "healer_t1_medic":    {"class": "healer", "tier": 1, "icon": "💊",
                           "name": ("Medic", "Medic"), "buff": {"hp_reduc": 4},
                           "desc": ("-4 HP dari tiap kerusakan", "-4 HP from every hit")},
    "healer_t1_bless":    {"class": "healer", "tier": 1, "icon": "🙏",
                           "name": ("Blessing", "Blessing"), "buff": {"xp_pct": 8},
                           "desc": ("+8% XP", "+8% XP")},
    "healer_t2_charity":  {"class": "healer", "tier": 2, "icon": "🎁",
                           "name": ("Charity", "Charity"), "buff": {"gold_pct": 8},
                           "desc": ("+8% Gold", "+8% Gold")},
    "healer_t2_prayer":   {"class": "healer", "tier": 2, "icon": "🕯️",
                           "name": ("Prayer", "Prayer"), "buff": {"mp_bonus": 15},
                           "desc": ("+15 Max MP", "+15 Max MP")},
    "healer_t3_guard":    {"class": "healer", "tier": 3, "icon": "😇",
                           "name": ("Guardian", "Guardian"), "buff": {"hp_reduc": 10},
                           "desc": ("-10 HP dari tiap kerusakan", "-10 HP from every hit")},
    "healer_t3_holy":     {"class": "healer", "tier": 3, "icon": "✨",
                           "name": ("Holy Strike", "Holy Strike"), "buff": {"boss_dmg": 6},
                           "desc": ("+6 Boss DMG", "+6 Boss DMG")},
    # Rogue 🗡
    "rogue_t1_pick":      {"class": "rogue", "tier": 1, "icon": "🪙",
                           "name": ("Pickpocket", "Pickpocket"), "buff": {"gold_pct": 6},
                           "desc": ("+6% Gold", "+6% Gold")},
    "rogue_t1_nimble":    {"class": "rogue", "tier": 1, "icon": "💨",
                           "name": ("Nimble", "Nimble"), "buff": {"xp_pct": 6},
                           "desc": ("+6% XP", "+6% XP")},
    "rogue_t2_shadow":    {"class": "rogue", "tier": 2, "icon": "🌫️",
                           "name": ("Shadow Mend", "Shadow Mend"), "buff": {"hp_reduc": 4},
                           "desc": ("-4 HP dari tiap kerusakan", "-4 HP from every hit")},
    "rogue_t2_loot":      {"class": "rogue", "tier": 2, "icon": "💎",
                           "name": ("Lootmaster", "Lootmaster"), "buff": {"gold_pct": 10},
                           "desc": ("+10% Gold", "+10% Gold")},
    "rogue_t3_phantom":   {"class": "rogue", "tier": 3, "icon": "👤",
                           "name": ("Phantom", "Phantom"), "buff": {"gold_pct": 15},
                           "desc": ("+15% Gold", "+15% Gold")},
    "rogue_t3_assassin":  {"class": "rogue", "tier": 3, "icon": "🗡️",
                           "name": ("Assassin", "Assassin"), "buff": {"boss_dmg": 8},
                           "desc": ("+8 Boss DMG", "+8 Boss DMG")},
}

TALENT_POINT_PER_LEVELS = 5   # 1 poin tiap 5 level


def get_talent_state(user_id):
    """{points, class, unlocked:set, tiers:{1:[node...],...}} untuk UI."""
    u = get_user(user_id)
    cls = u.get("avatar_class", "warrior")
    lang = u.get("language", "id")
    level = u.get("level", 1) or 1
    earned = level // TALENT_POINT_PER_LEVELS
    spent = u.get("talent_points_spent", 0) or 0
    conn = get_conn()
    rows = conn.execute("SELECT talent_key FROM user_talents WHERE user_id=?",
                        (user_id,)).fetchall()
    conn.close()
    unlocked = {r["talent_key"] for r in rows}

    tiers = {1: [], 2: [], 3: []}
    for key, t in TALENTS.items():
        if t["class"] != cls:
            continue
        idx = 0 if lang == "id" else 1
        tiers[t["tier"]].append({
            "key": key, "icon": t["icon"], "name": t["name"][idx],
            "desc": t["desc"][idx], "tier": t["tier"],
            "unlocked": key in unlocked,
            "tier_req_level": t["tier"] * TALENT_POINT_PER_LEVELS,
        })
    return {"points": max(0, earned - spent), "earned": earned,
            "class": cls, "level": level, "unlocked": unlocked, "tiers": tiers}


def can_unlock_talent(user_id, talent_key):
    """Aturan: class cocok, belum di-unlock, punya poin, level ≥ tier×5,
    dan tier N>1 butuh minimal 1 talent di tier N-1."""
    st = get_talent_state(user_id)
    t = TALENTS.get(talent_key)
    if not t or t["class"] != st["class"]:
        return {"ok": False, "reason": "class"}
    if talent_key in st["unlocked"]:
        return {"ok": False, "reason": "owned"}
    if st["points"] <= 0:
        return {"ok": False, "reason": "points"}
    tier = t["tier"]
    if st["level"] < tier * TALENT_POINT_PER_LEVELS:
        return {"ok": False, "reason": "level",
                "need_level": tier * TALENT_POINT_PER_LEVELS}
    if tier > 1:
        prev = [k for k, x in TALENTS.items()
                if x["class"] == st["class"] and x["tier"] == tier - 1]
        if not any(k in st["unlocked"] for k in prev):
            return {"ok": False, "reason": "prereq", "tier": tier}
    return {"ok": True}


@retry_on_lock
def unlock_talent(user_id, talent_key):
    """Beli 1 node talent dengan 1 poin, lalu hitung ulang buff."""
    chk = can_unlock_talent(user_id, talent_key)
    if not chk["ok"]:
        return {"ok": False, "msg": tr_db(user_id=user_id,
                                          key="db_talent_cant", reason=chk["reason"])}
    conn = get_conn()
    conn.execute("INSERT INTO user_talents(user_id, talent_key) VALUES(?,?)",
                 (user_id, talent_key))
    conn.execute("UPDATE users SET talent_points_spent=talent_points_spent+1 WHERE id=?",
                 (user_id,))
    conn.commit()
    conn.close()
    recalculate_all_buffs(user_id)
    t = TALENTS[talent_key]
    u = get_user(user_id)
    name = t["name"][0] if u.get("language", "id") == "id" else t["name"][1]
    log.info(f"Talent unlocked user_id={user_id}: {talent_key}")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_talent_unlocked",
                                     icon=t["icon"], name=name)}


# ── 👾 CUSTOM BOSS (buatan user/guild) ────────────────────────────────────────

CUSTOM_BOSS_ICONS = ["👾", "👹", "🤖", "🦖", "🐙", "🦂", "🧌", "🐉", "🦇", "🍄"]


def _boss_from_custom(row):
    """Ubah baris custom_bosses → dict ala BOSSES."""
    return {"name": row["name"], "icon": row["icon"], "tier": "custom",
            "hp": row["hp"], "atk": row["atk"],
            "xp": max(50, row["hp"] // 2), "gold": max(20, row["hp"] // 5),
            "min_level": row["min_level"], "custom": True}


def get_custom_boss(boss_key):
    """boss_key 'custom_<id>' → dict boss atau None."""
    if not str(boss_key).startswith("custom_"):
        return None
    try:
        cid = int(str(boss_key)[7:])
    except ValueError:
        return None
    conn = get_conn()
    row = conn.execute("SELECT * FROM custom_bosses WHERE id=?", (cid,)).fetchone()
    conn.close()
    return _boss_from_custom(row) if row else None


def get_effective_boss(boss_key):
    """BOSSES biasa dulu, fallback ke custom_bosses."""
    b = BOSSES.get(boss_key)
    if b:
        return b
    return get_custom_boss(boss_key)


def get_all_bosses_for_guild(guild_id):
    """BOSSES bawaan + custom milik guild ini: {boss_key: boss_dict}."""
    all_b = dict(BOSSES)
    conn = get_conn()
    if guild_id:
        rows = conn.execute(
            "SELECT * FROM custom_bosses WHERE guild_id=? OR guild_id IS NULL",
            (guild_id,)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM custom_bosses WHERE guild_id IS NULL",
                            ()).fetchall()
    conn.close()
    for r in rows:
        all_b[f"custom_{r['id']}"] = _boss_from_custom(r)
    return all_b


def create_custom_boss(user_id, guild_id, name, icon, hp, atk, min_level=1):
    """Validasi + simpan boss buatan user."""
    name = (name or "").strip()
    if not (3 <= len(name) <= 30):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_cboss_name_bad")}
    if icon not in CUSTOM_BOSS_ICONS:
        icon = CUSTOM_BOSS_ICONS[0]
    try:
        hp = int(hp); atk = int(atk); min_level = int(min_level)
    except (TypeError, ValueError):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_cboss_bad_stats")}
    if not (100 <= hp <= 10000) or not (1 <= atk <= 150) or not (1 <= min_level <= 99):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_cboss_bad_stats")}
    conn = get_conn()
    conn.execute(
        "INSERT INTO custom_bosses(guild_id, creator_id, name, icon, hp, atk, min_level)"
        " VALUES(?,?,?,?,?,?,?)",
        (guild_id, user_id, name, icon, hp, atk, min_level))
    conn.commit()
    conn.close()
    log.info(f"Custom boss dibuat user_id={user_id}: '{name}' (hp={hp}, atk={atk})")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_cboss_created", name=name)}


def delete_custom_boss(user_id, boss_db_id):
    conn = get_conn()
    conn.execute("DELETE FROM custom_bosses WHERE id=? AND creator_id=?",
                 (boss_db_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True}


# ── ⚔️ PvP STREAK BATTLE (tantangan produktivitas antar-teman) ───────────────

PVP_DURATION_DAYS = 7
PVP_WINNER_XP = 100
PVP_WINNER_GOLD = 50


def _are_friends(a, b):
    conn = get_conn()
    r = conn.execute(
        "SELECT 1 FROM friends WHERE user_id=? AND friend_id=? AND status='accepted'",
        (a, b)).fetchone()
    conn.close()
    return bool(r)


def _pvp_score(user_id, start, end):
    conn = get_conn()
    r = conn.execute(
        "SELECT COUNT(*) c FROM task_history"
        " WHERE user_id=? AND action='success' AND action_date>=? AND action_date<=?",
        (user_id, start, end)).fetchone()
    conn.close()
    return r["c"]


def send_pvp_challenge(user_id, friend_id):
    """Kirim tantangan PvP 7 hari ke teman (status accepted)."""
    if not _are_friends(user_id, friend_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pvp_not_friends")}
    conn = get_conn()
    dup = conn.execute(
        "SELECT 1 FROM pvp_challenges WHERE status IN ('pending','active')"
        " AND ((challenger_id=? AND opponent_id=?) OR (challenger_id=? AND opponent_id=?))",
        (user_id, friend_id, friend_id, user_id)).fetchone()
    if dup:
        conn.close()
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_pvp_dup")}
    conn.execute(
        "INSERT INTO pvp_challenges(challenger_id, opponent_id) VALUES(?,?)",
        (user_id, friend_id))
    conn.commit()
    conn.close()
    log.info(f"PvP challenge: {user_id} → {friend_id}")
    return {"ok": True, "msg": tr_db(user_id=user_id, key="db_pvp_sent")}


def respond_pvp_challenge(challenge_id, user_id, accept=True):
    """Terima/tolak tantangan (hanya oleh opponent)."""
    conn = get_conn()
    ch = conn.execute("SELECT * FROM pvp_challenges WHERE id=?", (challenge_id,)).fetchone()
    if not ch or ch["opponent_id"] != user_id or ch["status"] != "pending":
        conn.close()
        return {"ok": False}
    if accept:
        today = date.today()
        end = today + timedelta(days=PVP_DURATION_DAYS - 1)
        conn.execute(
            "UPDATE pvp_challenges SET status='active', start_date=?, end_date=? WHERE id=?",
            (today.isoformat(), end.isoformat(), challenge_id))
    else:
        conn.execute("UPDATE pvp_challenges SET status='declined' WHERE id=?",
                     (challenge_id,))
    conn.commit()
    conn.close()
    return {"ok": True}


def _finalize_pvp_if_due(ch):
    """Jika lewat end_date: tentukan pemenang + beri hadiah. Return row terbaru."""
    if ch["status"] != "active" or not ch["end_date"]:
        return ch
    if date.today() <= date.fromisoformat(ch["end_date"]):
        return ch
    ca = _pvp_score(ch["challenger_id"], ch["start_date"], ch["end_date"])
    cb = _pvp_score(ch["opponent_id"], ch["start_date"], ch["end_date"])
    winner = ch["challenger_id"] if ca > cb else (ch["opponent_id"] if cb > ca else None)
    conn = get_conn()
    conn.execute("UPDATE pvp_challenges SET status='finished', winner_id=? WHERE id=?",
                 (winner, ch["id"]))
    conn.commit()
    conn.close()
    if winner:
        gain_xp_gold(winner, PVP_WINNER_XP, PVP_WINNER_GOLD)
        add_notification(winner, tr_db(user_id=winner, key="db_pvp_won",
                                       xp=PVP_WINNER_XP, gold=PVP_WINNER_GOLD), "success")
    log.info(f"PvP #{ch['id']} selesai: {ca}-{cb}, winner={winner}")
    conn = get_conn()
    row = conn.execute("SELECT * FROM pvp_challenges WHERE id=?", (ch["id"],)).fetchone()
    conn.close()
    return row


def get_pvp_challenges(user_id):
    """Semua tantangan user (pending/active/finished) dengan skor live."""
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM pvp_challenges WHERE challenger_id=? OR opponent_id=?"
        " ORDER BY id DESC LIMIT 20", (user_id, user_id)).fetchall()
    conn.close()

    out = []
    for ch in rows:
        ch = _finalize_pvp_if_due(dict(ch))
        opp_id = ch["opponent_id"] if ch["challenger_id"] == user_id else ch["challenger_id"]
        ou = get_user(opp_id)
        item = {
            "id": ch["id"], "status": ch["status"],
            "is_challenger": ch["challenger_id"] == user_id,
            "opponent_name": ou.get("display_name") or ou.get("username", "?"),
            "opponent_id": opp_id,
            "start": ch["start_date"], "end": ch["end_date"],
            "winner_id": ch["winner_id"],
        }
        if ch["status"] in ("active", "finished") and ch["start_date"]:
            start, end = ch["start_date"], ch["end_date"]
            item["my_score"] = _pvp_score(user_id, start, end)
            item["opponent_score"] = _pvp_score(opp_id, start, end)
            try:
                item["days_left"] = max(0, (date.fromisoformat(end)
                                            - date.today()).days + 1)
            except (ValueError, TypeError):
                item["days_left"] = 0
        out.append(item)
    return out


# ── ♿ AKSESIBILITAS & ONBOARDING ─────────────────────────────────────────────

def set_font_scale(user_id, scale):
    """Skala font UI (80–140%)."""
    try:
        scale = int(scale)
    except (TypeError, ValueError):
        return False
    scale = max(80, min(140, scale))
    update_user(user_id, font_scale=scale)
    return True


def set_high_contrast(user_id, enabled):
    update_user(user_id, high_contrast=1 if enabled else 0)
    return True


def mark_onboarding_done(user_id):
    update_user(user_id, onboarding_done=1)


def set_dashboard_widgets(user_id, widgets):
    """Simpan konfigurasi widget dashboard (list of dict) sebagai JSON."""
    update_user(user_id, dashboard_widgets=json.dumps(widgets, ensure_ascii=False))


def get_dashboard_widgets(user_id):
    """Konfigurasi widget dashboard; default jika belum diatur."""
    u = get_user(user_id)
    raw = u.get("dashboard_widgets") or ""
    try:
        data = json.loads(raw)
        if isinstance(data, list) and data:
            return data
    except (ValueError, TypeError):
        pass
    return [
        {"key": "heatmap", "visible": True, "compact": False},
        {"key": "insights", "visible": True, "compact": False},
        {"key": "health_chart", "visible": True, "compact": False},
    ]


# ========== LEARNING PAGE (NotebookLM) ==========
def set_gemini_api_key(user_id, api_key):
    """Simpan API key Gemini untuk user."""
    conn = get_conn()
    conn.execute("UPDATE users SET gemini_api_key=? WHERE id=?", (api_key, user_id))
    conn.commit()
    conn.close()
    return {"ok": True}

def get_gemini_api_key(user_id):
    u = get_user(user_id)
    return (u.get("gemini_api_key") or "").strip() if u else ""

# ── Notebooks ──
def create_learning_notebook(user_id, title):
    conn = get_conn()
    cur = conn.execute("INSERT INTO learning_notebooks(user_id, title) VALUES(?,?)", (user_id, title))
    nid = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "notebook_id": nid}

def get_learning_notebooks(user_id):
    conn = get_conn()
    rows = conn.execute("SELECT * FROM learning_notebooks WHERE user_id=? ORDER BY created_at DESC", (user_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_learning_notebook(notebook_id, user_id):
    conn = get_conn()
    row = conn.execute("SELECT * FROM learning_notebooks WHERE id=? AND user_id=?", (notebook_id, user_id)).fetchone()
    conn.close()
    return dict(row) if row else None

def delete_learning_notebook(notebook_id, user_id):
    conn = get_conn()
    conn.execute("DELETE FROM learning_notebooks WHERE id=? AND user_id=?", (notebook_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True}

def update_learning_notebook(notebook_id, user_id, title):
    conn = get_conn()
    conn.execute("UPDATE learning_notebooks SET title=? WHERE id=? AND user_id=?", (title, notebook_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True}

# ── Sources ──
def add_learning_source(notebook_id, user_id, type_, title, path, content):
    conn = get_conn()
    cur = conn.execute("INSERT INTO learning_sources(notebook_id, user_id, type, title, path, content) VALUES(?,?,?,?,?,?)",
                       (notebook_id, user_id, type_, title, path, content))
    sid = cur.lastrowid
    conn.commit()
    conn.close()
    # Chunking
    try:
        import learning_helper as lh
        chunks = lh.chunk_text(content)
    except:
        # Fallback simple chunk
        chunks = [content[i:i+800] for i in range(0, len(content), 800)]
    conn = get_conn()
    for idx, ch in enumerate(chunks):
        conn.execute("INSERT INTO learning_chunks(source_id, chunk_text, chunk_index) VALUES(?,?,?)", (sid, ch, idx))
    conn.commit()
    conn.close()
    return {"ok": True, "source_id": sid, "chunks": len(chunks)}

def get_learning_sources(notebook_id, user_id=None):
    conn = get_conn()
    if user_id:
        rows = conn.execute("SELECT * FROM learning_sources WHERE notebook_id=? AND user_id=? ORDER BY created_at DESC", (notebook_id, user_id)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM learning_sources WHERE notebook_id=? ORDER BY created_at DESC", (notebook_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def delete_learning_source(source_id, user_id):
    conn = get_conn()
    conn.execute("DELETE FROM learning_sources WHERE id=? AND user_id=?", (source_id, user_id))
    conn.commit()
    conn.close()
    return {"ok": True}

def get_learning_chunks(notebook_id, user_id=None):
    conn = get_conn()
    # Ambil semua chunks untuk notebook ini
    rows = conn.execute("""
        SELECT c.chunk_text, c.source_id, s.title as source_title
        FROM learning_chunks c
        JOIN learning_sources s ON s.id = c.source_id
        WHERE s.notebook_id=? 
        ORDER BY c.chunk_index
    """, (notebook_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_relevant_chunks(notebook_id, query, top_k=3):
    """Cari chunks relevan via keyword overlap (simple RAG)."""
    all_chunks = get_learning_chunks(notebook_id)
    if not all_chunks:
        return []
    try:
        import learning_helper as lh
        texts = [c["chunk_text"] for c in all_chunks]
        relevant = lh.find_relevant_chunks(texts, query, top_k)
        # Map back to dict with source
        result = []
        for r in relevant:
            for ch in all_chunks:
                if ch["chunk_text"] == r:
                    result.append(ch)
                    break
        return result[:top_k]
    except:
        # Fallback: just return first few
        return all_chunks[:top_k]

# ── Chats ──
def add_learning_chat(notebook_id, role, content, citations=None):
    conn = get_conn()
    conn.execute("INSERT INTO learning_chats(notebook_id, role, content, citations) VALUES(?,?,?,?)",
                 (notebook_id, role, content, json.dumps(citations) if citations else None))
    conn.commit()
    conn.close()
    return {"ok": True}

def get_learning_chats(notebook_id):
    conn = get_conn()
    rows = conn.execute("SELECT * FROM learning_chats WHERE notebook_id=? ORDER BY created_at ASC", (notebook_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def clear_learning_chats(notebook_id):
    conn = get_conn()
    conn.execute("DELETE FROM learning_chats WHERE notebook_id=?", (notebook_id,))
    conn.commit()
    conn.close()
    return {"ok": True}

# ── Generations (Studio) ──
def add_learning_generation(notebook_id, type_, title, content):
    conn = get_conn()
    cur = conn.execute("INSERT INTO learning_generations(notebook_id, type, title, content) VALUES(?,?,?,?)",
                       (notebook_id, type_, title, content))
    gid = cur.lastrowid
    conn.commit()
    conn.close()
    return {"ok": True, "generation_id": gid}

def get_learning_generations(notebook_id, type_=None):
    conn = get_conn()
    if type_:
        rows = conn.execute("SELECT * FROM learning_generations WHERE notebook_id=? AND type=? ORDER BY created_at DESC", (notebook_id, type_)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM learning_generations WHERE notebook_id=? ORDER BY created_at DESC", (notebook_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def delete_learning_generation(gen_id, notebook_id):
    conn = get_conn()
    conn.execute("DELETE FROM learning_generations WHERE id=? AND notebook_id=?", (gen_id, notebook_id))
    conn.commit()
    conn.close()
    return {"ok": True}

# ========== LOVE PAGE ==========
def get_relationship_profile(user_id):
    owner_id = _love_primary_user_id(user_id)
    conn = get_conn()
    row = conn.execute("SELECT * FROM relationship_profiles WHERE user_id=?", (owner_id,)).fetchone()
    conn.close()
    return dict(row) if row else {}


def save_relationship_profile(user_id, partner_name, partner_gender, partner_age,
                              relationship_type, start_date):
    if is_account_locked(user_id):
        return {"ok": False, "msg": tr_db(user_id=user_id, key="db_account_locked_msg")}
    owner_id = _love_primary_user_id(user_id)
    conn = get_conn()
    conn.execute("""
        INSERT INTO relationship_profiles(user_id, partner_name, partner_gender, partner_age,
                                          relationship_type, start_date, updated_at)
        VALUES(?,?,?,?,?,?,datetime('now'))
        ON CONFLICT(user_id) DO UPDATE SET
            partner_name=excluded.partner_name, partner_gender=excluded.partner_gender,
            partner_age=excluded.partner_age, relationship_type=excluded.relationship_type,
            start_date=excluded.start_date, updated_at=datetime('now')
    """, (owner_id, partner_name.strip(), partner_gender, int(partner_age),
          relationship_type, start_date))
    conn.commit(); conn.close()
    return {"ok": True}


def add_relationship_event(user_id, title, event_date, category="date", notes=""):
    conn = get_conn()
    cur = conn.execute("INSERT INTO relationship_events(user_id,title,event_date,category,notes) VALUES(?,?,?,?,?)",
                       (user_id, title.strip(), event_date, category, notes.strip()))
    conn.commit(); conn.close()
    return {"ok": True, "event_id": cur.lastrowid}


def get_relationship_events(user_id, upcoming_only=False, limit=100):
    scope = _love_scope_user_ids(user_id); marks = ",".join("?" for _ in scope)
    conn = get_conn()
    if upcoming_only:
        rows = conn.execute(f"SELECT * FROM relationship_events WHERE user_id IN ({marks}) AND event_date>=date('now') ORDER BY event_date LIMIT ?",
                            (*scope, limit)).fetchall()
    else:
        rows = conn.execute(f"SELECT * FROM relationship_events WHERE user_id IN ({marks}) ORDER BY event_date DESC LIMIT ?",
                            (*scope, limit)).fetchall()
    conn.close(); return [dict(row) for row in rows]


def delete_relationship_event(user_id, event_id):
    conn = get_conn(); conn.execute("DELETE FROM relationship_events WHERE id=? AND user_id=?", (event_id, user_id))
    conn.commit(); conn.close(); return {"ok": True}


def add_relationship_memory(user_id, title, memory_date, notes=""):
    conn = get_conn()
    cur = conn.execute("INSERT INTO relationship_memories(user_id,title,memory_date,notes) VALUES(?,?,?,?)",
                       (user_id, title.strip(), memory_date, notes.strip()))
    conn.commit(); conn.close(); return {"ok": True, "memory_id": cur.lastrowid}


def get_relationship_memories(user_id, limit=100):
    scope = _love_scope_user_ids(user_id); marks = ",".join("?" for _ in scope)
    conn = get_conn(); rows = conn.execute(
        f"SELECT * FROM relationship_memories WHERE user_id IN ({marks}) ORDER BY memory_date DESC LIMIT ?",
        (*scope, limit)).fetchall()
    conn.close(); return [dict(row) for row in rows]


def delete_relationship_memory(user_id, memory_id):
    conn = get_conn(); conn.execute("DELETE FROM relationship_memories WHERE id=? AND user_id=?", (memory_id, user_id))
    conn.commit(); conn.close(); return {"ok": True}


def save_relationship_checkin(user_id, checkin_date, my_mood, partner_mood, connection_score, note=""):
    conn = get_conn()
    conn.execute("""
        INSERT INTO relationship_checkins(user_id,checkin_date,my_mood,partner_mood,connection_score,note)
        VALUES(?,?,?,?,?,?)
        ON CONFLICT(user_id,checkin_date) DO UPDATE SET
            my_mood=excluded.my_mood, partner_mood=excluded.partner_mood,
            connection_score=excluded.connection_score, note=excluded.note
    """, (user_id, checkin_date, my_mood, partner_mood, connection_score, note.strip()))
    conn.commit(); conn.close(); return {"ok": True}


def get_relationship_checkins(user_id, limit=14):
    scope = _love_scope_user_ids(user_id); marks = ",".join("?" for _ in scope)
    conn = get_conn(); rows = conn.execute(
        f"SELECT * FROM relationship_checkins WHERE user_id IN ({marks}) ORDER BY checkin_date DESC LIMIT ?",
        (*scope, limit)).fetchall()
    conn.close(); return [dict(row) for row in rows]


def add_relationship_prompt_response(user_id, prompt_key, category, prompt_text,
                                     my_answer, partner_answer, response_date=None):
    response_date = response_date or date.today().isoformat()
    conn = get_conn(); cur = conn.execute("""
        INSERT INTO relationship_prompt_responses
        (user_id,prompt_key,category,prompt_text,my_answer,partner_answer,response_date)
        VALUES(?,?,?,?,?,?,?)
    """, (user_id, prompt_key, category, prompt_text,
          my_answer.strip(), partner_answer.strip(), response_date))
    conn.commit(); conn.close()
    return {"ok": True, "response_id": cur.lastrowid}


def get_relationship_prompt_responses(user_id, limit=30):
    scope = _love_scope_user_ids(user_id); marks = ",".join("?" for _ in scope)
    conn = get_conn(); rows = conn.execute(f"""
        SELECT * FROM relationship_prompt_responses
        WHERE user_id IN ({marks}) ORDER BY response_date DESC, id DESC LIMIT ?
    """, (*scope, limit)).fetchall()
    conn.close(); return [dict(row) for row in rows]


def delete_relationship_prompt_response(user_id, response_id):
    conn = get_conn(); conn.execute(
        "DELETE FROM relationship_prompt_responses WHERE id=? AND user_id=?",
        (response_id, user_id))
    conn.commit(); conn.close(); return {"ok": True}


def get_relationship_prompt_favorites(user_id):
    conn = get_conn(); rows = conn.execute(
        "SELECT prompt_key FROM relationship_prompt_favorites WHERE user_id=?",
        (user_id,)).fetchall()
    conn.close(); return {row["prompt_key"] for row in rows}


def toggle_relationship_prompt_favorite(user_id, prompt_key):
    conn = get_conn(); row = conn.execute(
        "SELECT 1 FROM relationship_prompt_favorites WHERE user_id=? AND prompt_key=?",
        (user_id, prompt_key)).fetchone()
    if row:
        conn.execute("DELETE FROM relationship_prompt_favorites WHERE user_id=? AND prompt_key=?",
                     (user_id, prompt_key)); favorite = False
    else:
        conn.execute("INSERT INTO relationship_prompt_favorites(user_id,prompt_key) VALUES(?,?)",
                     (user_id, prompt_key)); favorite = True
    conn.commit(); conn.close(); return {"ok": True, "favorite": favorite}


def save_relationship_weekly_review(user_id, week_start, appreciation, wins,
                                    support_needed, shared_intention):
    conn = get_conn(); conn.execute("""
        INSERT INTO relationship_weekly_reviews
        (user_id,week_start,appreciation,wins,support_needed,shared_intention)
        VALUES(?,?,?,?,?,?)
        ON CONFLICT(user_id,week_start) DO UPDATE SET
            appreciation=excluded.appreciation, wins=excluded.wins,
            support_needed=excluded.support_needed,
            shared_intention=excluded.shared_intention, updated_at=datetime('now')
    """, (user_id, week_start, appreciation.strip(), wins.strip(),
          support_needed.strip(), shared_intention.strip()))
    conn.commit(); conn.close(); return {"ok": True}


def get_relationship_weekly_reviews(user_id, limit=12):
    scope = _love_scope_user_ids(user_id); marks = ",".join("?" for _ in scope)
    conn = get_conn(); rows = conn.execute(f"""
        SELECT * FROM relationship_weekly_reviews
        WHERE user_id IN ({marks}) ORDER BY week_start DESC LIMIT ?
    """, (*scope, limit)).fetchall()
    conn.close(); return [dict(row) for row in rows]


def delete_relationship_weekly_review(user_id, review_id):
    conn = get_conn(); conn.execute(
        "DELETE FROM relationship_weekly_reviews WHERE id=? AND user_id=?",
        (review_id, user_id))
    conn.commit(); conn.close(); return {"ok": True}


def add_relationship_bucket_item(user_id, title, category="dream", target_date=None):
    conn = get_conn(); cur = conn.execute(
        "INSERT INTO relationship_bucket_items(user_id,title,category,target_date) VALUES(?,?,?,?)",
        (user_id, title.strip(), category, target_date))
    conn.commit(); conn.close(); return {"ok": True, "item_id": cur.lastrowid}


def get_relationship_bucket_items(user_id):
    scope = _love_scope_user_ids(user_id); marks = ",".join("?" for _ in scope)
    conn = get_conn(); rows = conn.execute(
        f"SELECT * FROM relationship_bucket_items WHERE user_id IN ({marks}) ORDER BY is_done, COALESCE(target_date,'9999-12-31'), created_at DESC",
        tuple(scope)).fetchall()
    conn.close(); return [dict(row) for row in rows]


def toggle_relationship_bucket_item(user_id, item_id, done):
    conn = get_conn(); conn.execute(
        "UPDATE relationship_bucket_items SET is_done=?, completed_at=CASE WHEN ?=1 THEN datetime('now') ELSE NULL END WHERE id=? AND user_id=?",
        (1 if done else 0, 1 if done else 0, item_id, user_id))
    conn.commit(); conn.close(); return {"ok": True}


def delete_relationship_bucket_item(user_id, item_id):
    conn = get_conn(); conn.execute("DELETE FROM relationship_bucket_items WHERE id=? AND user_id=?", (item_id, user_id))
    conn.commit(); conn.close(); return {"ok": True}


def get_menstrual_settings(user_id):
    owner_id = _love_primary_user_id(user_id)
    conn = get_conn(); row = conn.execute("SELECT * FROM menstrual_settings WHERE user_id=?", (owner_id,)).fetchone()
    conn.close()
    return dict(row) if row else {"tracked_person": "partner", "last_period_start": "", "cycle_length": 28, "period_length": 5}


def save_menstrual_settings(user_id, tracked_person, last_period_start, cycle_length, period_length):
    owner_id = _love_primary_user_id(user_id)
    conn = get_conn(); conn.execute("""
        INSERT INTO menstrual_settings(user_id,tracked_person,last_period_start,cycle_length,period_length,updated_at)
        VALUES(?,?,?,?,?,datetime('now'))
        ON CONFLICT(user_id) DO UPDATE SET tracked_person=excluded.tracked_person,
            last_period_start=excluded.last_period_start, cycle_length=excluded.cycle_length,
            period_length=excluded.period_length, updated_at=datetime('now')
    """, (owner_id, tracked_person, last_period_start, int(cycle_length), int(period_length)))
    conn.commit(); conn.close(); return {"ok": True}


def add_menstrual_cycle(user_id, start_date, end_date=None, notes=""):
    owner_id = _love_primary_user_id(user_id)
    conn = get_conn(); cur = conn.execute(
        "INSERT INTO menstrual_cycles(user_id,start_date,end_date,notes) VALUES(?,?,?,?)",
        (user_id, start_date, end_date, notes.strip()))
    conn.execute("""
        INSERT INTO menstrual_settings(user_id,last_period_start) VALUES(?,?)
        ON CONFLICT(user_id) DO UPDATE SET last_period_start=excluded.last_period_start, updated_at=datetime('now')
    """, (owner_id, start_date))
    conn.commit(); conn.close(); return {"ok": True, "cycle_id": cur.lastrowid}


def get_menstrual_cycles(user_id, limit=12):
    scope = _love_scope_user_ids(user_id); marks = ",".join("?" for _ in scope)
    conn = get_conn(); rows = conn.execute(
        f"SELECT * FROM menstrual_cycles WHERE user_id IN ({marks}) ORDER BY start_date DESC LIMIT ?",
        (*scope, limit)).fetchall()
    conn.close(); return [dict(row) for row in rows]


def delete_menstrual_cycle(user_id, cycle_id):
    conn = get_conn(); conn.execute("DELETE FROM menstrual_cycles WHERE id=? AND user_id=?", (cycle_id, user_id))
    conn.commit(); conn.close(); return {"ok": True}


def get_menstrual_prediction(user_id, on_date=None):
    settings = get_menstrual_settings(user_id)
    cycles = get_menstrual_cycles(user_id, 1)
    raw_start = cycles[0]["start_date"] if cycles else settings.get("last_period_start")
    if not raw_start:
        return {}
    try:
        start = date.fromisoformat(raw_start)
        today = date.fromisoformat(on_date) if on_date else date.today()
        cycle_length = max(20, min(45, int(settings.get("cycle_length") or 28)))
        period_length = max(2, min(10, int(settings.get("period_length") or 5)))
        predicted = start
        while predicted + timedelta(days=period_length - 1) < today:
            predicted += timedelta(days=cycle_length)
        return {
            "tracked_person": settings.get("tracked_person", "partner"),
            "predicted_start": predicted.isoformat(),
            "predicted_end": (predicted + timedelta(days=period_length - 1)).isoformat(),
            "days_until": (predicted - today).days,
            "cycle_length": cycle_length,
            "period_length": period_length,
            "is_estimate": True,
        }
    except (ValueError, TypeError):
        return {}


# ========== OPTIONAL SUPABASE OFFLINE-FIRST METADATA ==========
def save_cloud_user_link(local_user_id, cloud_user_id, email, status="linked"):
    conn=get_conn(); conn.execute("""INSERT INTO cloud_user_links
        (local_user_id,cloud_user_id,email,status,linked_at)
        VALUES(?,?,?,?,datetime('now'))
        ON CONFLICT(local_user_id) DO UPDATE SET cloud_user_id=excluded.cloud_user_id,
            email=excluded.email,status=excluded.status""",
        (local_user_id,cloud_user_id,email,status))
    conn.execute("UPDATE users SET cloud_user_id=?,is_cloud_shadow=0 WHERE id=?", (cloud_user_id,local_user_id))
    conn.commit();conn.close();return {"ok":True}


def get_cloud_user_link(local_user_id):
    conn=get_conn();row=conn.execute("SELECT * FROM cloud_user_links WHERE local_user_id=?",(local_user_id,)).fetchone();conn.close()
    return dict(row) if row else {}


def get_local_user_id_for_cloud(cloud_user_id):
    conn=get_conn();row=conn.execute("SELECT id FROM users WHERE cloud_user_id=?",(cloud_user_id,)).fetchone();conn.close()
    return row["id"] if row else None


def mark_cloud_sync_complete(local_user_id):
    conn=get_conn();conn.execute("UPDATE cloud_user_links SET last_sync_at=datetime('now'),status='linked' WHERE local_user_id=?",(local_user_id,));conn.commit();conn.close()


def enqueue_sync(local_user_id,entity_type,entity_local_id,operation,payload=None):
    payload_text=json.dumps(payload or {},ensure_ascii=False)
    conn=get_conn();existing=conn.execute("""SELECT id FROM sync_queue WHERE local_user_id=? AND entity_type=?
        AND COALESCE(entity_local_id,'')=COALESCE(?,'') AND operation=? AND status IN ('pending','retry')
        ORDER BY id DESC LIMIT 1""",(local_user_id,entity_type,str(entity_local_id) if entity_local_id is not None else None,operation)).fetchone()
    if existing:
        conn.execute("UPDATE sync_queue SET payload=?,status='pending',next_retry_at=NULL,last_error=NULL,updated_at=datetime('now') WHERE id=?",(payload_text,existing["id"]));job_id=existing["id"]
    else:
        job_id=conn.execute("INSERT INTO sync_queue(local_user_id,entity_type,entity_local_id,operation,payload) VALUES(?,?,?,?,?)",
            (local_user_id,entity_type,str(entity_local_id) if entity_local_id is not None else None,operation,payload_text)).lastrowid
    conn.commit();conn.close();return {"ok":True,"job_id":job_id}


def get_pending_sync_jobs(local_user_id,limit=25):
    conn=get_conn();rows=conn.execute("""SELECT * FROM sync_queue WHERE local_user_id=?
        AND status IN ('pending','retry') AND (next_retry_at IS NULL OR next_retry_at<=datetime('now'))
        ORDER BY created_at,id LIMIT ?""",(local_user_id,limit)).fetchall();conn.close();return [dict(r) for r in rows]


def mark_sync_job_done(job_id):
    conn=get_conn();conn.execute("UPDATE sync_queue SET status='done',last_error=NULL,updated_at=datetime('now') WHERE id=?",(job_id,));conn.commit();conn.close()


def mark_sync_job_failed(job_id,error,retry_count):
    delay=min(3600,30*(2**min(int(retry_count),6)))
    conn=get_conn();conn.execute("""UPDATE sync_queue SET status='retry',retry_count=?,last_error=?,
        next_retry_at=datetime('now', ?),updated_at=datetime('now') WHERE id=?""",
        (retry_count,str(error)[:1000],f'+{delay} seconds',job_id));conn.commit();conn.close()


def sync_queue_summary(local_user_id):
    conn=get_conn();rows=conn.execute("SELECT status,COUNT(*) c FROM sync_queue WHERE local_user_id=? GROUP BY status",(local_user_id,)).fetchall();conn.close()
    return {r["status"]:r["c"] for r in rows}


# ---------- Phase 3 personal realtime snapshot state ----------
def get_or_create_cloud_device(local_user_id, app_version="1.0-phase3"):
    """Return a stable, non-secret UUID for this local account/device pair."""
    conn=get_conn();row=conn.execute("SELECT * FROM cloud_device_state WHERE local_user_id=?",(local_user_id,)).fetchone()
    if row:
        conn.close();return dict(row)
    name=(os.environ.get("COMPUTERNAME") or os.environ.get("HOSTNAME") or "CraftLife Desktop")[:120]
    platform_name=(sys.platform or "desktop")[:80]
    device_id=str(uuid.uuid4())
    conn.execute("""INSERT INTO cloud_device_state(local_user_id,device_id,device_name,platform,app_version)
        VALUES(?,?,?,?,?)""",(local_user_id,device_id,name,platform_name,str(app_version)[:40]))
    conn.commit();conn.close()
    return {"local_user_id":local_user_id,"device_id":device_id,"device_name":name,
            "platform":platform_name,"app_version":str(app_version)[:40]}


def mark_cloud_device_registered(local_user_id, registered_at=None, last_seen_at=None):
    conn=get_conn();conn.execute("""UPDATE cloud_device_state SET registered_at=COALESCE(?,registered_at,datetime('now')),
        last_seen_at=COALESCE(?,datetime('now')) WHERE local_user_id=?""",
        (registered_at,last_seen_at,local_user_id));conn.commit();conn.close()


def rename_local_cloud_device(local_user_id,name):
    conn=get_conn();conn.execute("UPDATE cloud_device_state SET device_name=? WHERE local_user_id=?",(str(name)[:120],local_user_id));conn.commit();conn.close()


def build_cloud_personal_snapshot(local_user_id):
    """Build the private tracker_v1 document; BLOB/media and credentials are excluded."""
    snapshot=export_tracker_data(local_user_id)
    # Absolute custom-sound paths are device-local and can reveal Windows folder
    # names. Sync reminder metadata without uploading that unusable path.
    for reminder in snapshot.get("tables",{}).get("reminders",[]):
        reminder["sound_file"]=None
    snapshot["cloud_schema"]=1
    snapshot["document_key"]="tracker_v1"
    return snapshot


def cloud_personal_snapshot_hash(snapshot):
    # SQL row order is not semantic. Sort every table by canonical row JSON so
    # query-plan/order differences do not create phantom cloud revisions.
    normalized=json.loads(json.dumps(snapshot,ensure_ascii=False))
    for rows in (normalized.get("tables") or {}).values():
        if isinstance(rows,list):
            rows.sort(key=lambda row:json.dumps(row,ensure_ascii=False,sort_keys=True,separators=(",",":")))
    canonical=json.dumps(normalized,ensure_ascii=False,sort_keys=True,separators=(",",":"))
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def cloud_personal_snapshot_is_empty(snapshot):
    tables=(snapshot or {}).get("tables") or {}
    defaults={
        "user_nutrition_goals":{"daily_calories":2000,"daily_protein":50,"daily_carbs":250,"daily_fat":70},
        "user_water_goals":{"daily_ml":2000},
        "user_health_goals":{"daily_steps":10000,"daily_sleep_hours":7.0,"height_cm":170,
                             "weight_kg":70,"age":25,"gender":"Laki-laki","activity_factor":1.55},
    }
    for table,rows in tables.items():
        if not rows:continue
        if table not in defaults:return False
        expected=defaults[table]
        if any(any(row.get(key)!=value for key,value in expected.items()) for row in rows):return False
    return True


def _sanitize_cloud_personal_snapshot(snapshot):
    """Whitelist tables/columns before remote JSON reaches dynamic SQLite INSERTs."""
    if not isinstance(snapshot,dict) or not isinstance(snapshot.get("tables"),dict):
        raise ValueError("invalid cloud tracker snapshot")
    tables=snapshot["tables"]
    missing=[table for table in EXPORT_TABLES if table not in tables]
    if missing:
        raise ValueError("incomplete cloud tracker snapshot: "+", ".join(missing[:5]))
    if any(table not in EXPORT_TABLES for table in tables):
        raise ValueError("cloud tracker snapshot contains an unsupported table")
    raw_size=len(json.dumps(snapshot,ensure_ascii=False).encode("utf-8"))
    if raw_size>8*1024*1024:
        raise ValueError("cloud tracker snapshot exceeds 8 MB")
    conn=get_conn();allowed={}
    try:
        for table in EXPORT_TABLES:
            allowed[table]={str(row["name"]) for row in conn.execute(f"PRAGMA table_info({table})").fetchall()}
    finally:
        conn.close()
    cleaned={"version":int(snapshot.get("version") or 1),"cloud_schema":1,
             "document_key":"tracker_v1","tables":{}}
    total_rows=0
    for table in EXPORT_TABLES:
        rows=tables.get(table)
        if not isinstance(rows,list):raise ValueError(f"invalid rows for {table}")
        total_rows+=len(rows)
        if total_rows>200000:raise ValueError("cloud tracker snapshot has too many rows")
        extras={"food_name"} if table in ("food_logs","recipe_items") else set()
        clean_rows=[]
        for row in rows:
            if not isinstance(row,dict):raise ValueError(f"invalid row for {table}")
            clean_rows.append({key:value for key,value in row.items() if key in allowed[table] or key in extras})
        cleaned["tables"][table]=clean_rows
    return cleaned


def apply_cloud_personal_snapshot(local_user_id, snapshot):
    """Atomically replace cloud-covered tracker tables while preserving progress."""
    clean=_sanitize_cloud_personal_snapshot(snapshot)
    local_backup=build_cloud_personal_snapshot(local_user_id)
    import_tracker_data(local_user_id,clean,preserve_progress=True)
    # SQLite remaps relational integer IDs on import. Baseline the actual local
    # representation so an unchanged device does not upload a duplicate revision.
    imported_local=build_cloud_personal_snapshot(local_user_id)
    return {"ok":True,"local_backup":local_backup,
            "local_hash":cloud_personal_snapshot_hash(imported_local)}


def get_cloud_personal_sync_state(local_user_id, document_key="tracker_v1"):
    conn=get_conn();row=conn.execute("""SELECT * FROM cloud_personal_sync_state
        WHERE local_user_id=? AND document_key=?""",(local_user_id,document_key)).fetchone();conn.close()
    if not row:return {}
    result=dict(row)
    for key in ("remote_payload","local_backup_payload"):
        try:result[key]=json.loads(result[key]) if result.get(key) else None
        except (TypeError,ValueError):result[key]=None
    return result


def record_cloud_personal_synced(local_user_id, snapshot_row, local_hash, direction="push", local_backup=None):
    revision=int((snapshot_row or {}).get("revision") or 0);remote_hash=(snapshot_row or {}).get("content_hash") or local_hash
    updated=(snapshot_row or {}).get("server_updated_at")
    backup_text=json.dumps(local_backup,ensure_ascii=False) if local_backup is not None else None
    pushed=local_now().isoformat() if direction=="push" else None
    pulled=local_now().isoformat() if direction=="pull" else None
    conn=get_conn();conn.execute("""INSERT INTO cloud_personal_sync_state(
        local_user_id,document_key,last_local_hash,remote_revision,remote_hash,remote_updated_at,
        conflict_status,remote_payload,local_backup_payload,last_pushed_at,last_pulled_at)
        VALUES(?,'tracker_v1',?,?,?,?,'none',NULL,?,?,?)
        ON CONFLICT(local_user_id,document_key) DO UPDATE SET
          last_local_hash=excluded.last_local_hash,remote_revision=excluded.remote_revision,
          remote_hash=excluded.remote_hash,remote_updated_at=excluded.remote_updated_at,
          conflict_status='none',remote_payload=NULL,
          local_backup_payload=COALESCE(excluded.local_backup_payload,cloud_personal_sync_state.local_backup_payload),
          last_pushed_at=COALESCE(excluded.last_pushed_at,cloud_personal_sync_state.last_pushed_at),
          last_pulled_at=COALESCE(excluded.last_pulled_at,cloud_personal_sync_state.last_pulled_at)""",
        (local_user_id,local_hash,revision,remote_hash,updated,backup_text,pushed,pulled))
    conn.commit();conn.close()


def record_cloud_personal_conflict(local_user_id, snapshot_row, local_snapshot):
    remote_payload=(snapshot_row or {}).get("payload") or {};revision=int((snapshot_row or {}).get("revision") or 0)
    conn=get_conn();conn.execute("""INSERT INTO cloud_personal_sync_state(
        local_user_id,document_key,last_local_hash,remote_revision,remote_hash,remote_updated_at,
        conflict_status,remote_payload,local_backup_payload,last_pulled_at)
        VALUES(?,'tracker_v1',?,?,?,?,'needs_resolution',?,?,?)
        ON CONFLICT(local_user_id,document_key) DO UPDATE SET
          remote_revision=excluded.remote_revision,remote_hash=excluded.remote_hash,
          remote_updated_at=excluded.remote_updated_at,conflict_status='needs_resolution',
          remote_payload=excluded.remote_payload,local_backup_payload=excluded.local_backup_payload,
          last_pulled_at=excluded.last_pulled_at""",
        (local_user_id,cloud_personal_snapshot_hash(local_snapshot),revision,
         (snapshot_row or {}).get("content_hash"),(snapshot_row or {}).get("server_updated_at"),
         json.dumps(remote_payload,ensure_ascii=False),json.dumps(local_snapshot,ensure_ascii=False),local_now().isoformat()))
    conn.commit();conn.close()


def mark_cloud_personal_resolution_pending(local_user_id):
    conn=get_conn();conn.execute("""UPDATE cloud_personal_sync_state SET conflict_status='keep_local_pending'
        WHERE local_user_id=? AND document_key='tracker_v1'""",(local_user_id,));conn.commit();conn.close()


# ---------- Phase 4A cloud-native shared Love Space cache ----------
_LOVE_RECORD_TABLES={
    "event":"relationship_events","memory":"relationship_memories",
    "checkin":"relationship_checkins","prompt_response":"relationship_prompt_responses",
    "weekly_review":"relationship_weekly_reviews","bucket_item":"relationship_bucket_items",
    "cycle":"menstrual_cycles",
}


def get_cloud_love_space_id(local_user_id):
    context=get_couple_context(local_user_id);local_space=context.get("love_space_id")
    if not local_space:return None
    conn=get_conn();row=conn.execute("SELECT cloud_id FROM love_spaces WHERE id=?",(local_space,)).fetchone();conn.close()
    return row["cloud_id"] if row and row["cloud_id"] else get_cloud_id("love_space",local_space)


def get_love_record_cloud_id(record_type,local_id):
    table=_LOVE_RECORD_TABLES.get(record_type)
    if not table or local_id is None:return None
    conn=get_conn();row=conn.execute(f"SELECT cloud_id FROM {table} WHERE id=?",(local_id,)).fetchone();conn.close()
    return row["cloud_id"] if row else None


def ensure_love_record_cloud_id(record_type,local_id):
    table=_LOVE_RECORD_TABLES.get(record_type)
    if not table:raise ValueError("unsupported Love record type")
    conn=get_conn();row=conn.execute(f"SELECT cloud_id FROM {table} WHERE id=?",(local_id,)).fetchone()
    if not row:conn.close();return None
    cloud_id=row["cloud_id"] or str(uuid.uuid4())
    if not row["cloud_id"]:
        conn.execute(f"UPDATE {table} SET cloud_id=? WHERE id=?",(cloud_id,local_id));conn.commit()
    conn.close();return cloud_id


def get_local_love_record(record_type,local_id):
    table=_LOVE_RECORD_TABLES.get(record_type)
    if not table:return {}
    conn=get_conn();row=conn.execute(f"SELECT * FROM {table} WHERE id=?",(local_id,)).fetchone();conn.close()
    return dict(row) if row else {}


def delete_cached_cloud_love_record(record_type,cloud_id):
    table=_LOVE_RECORD_TABLES.get(record_type)
    if not table or not cloud_id:return
    conn=get_conn();conn.execute(f"DELETE FROM {table} WHERE cloud_id=?",(str(cloud_id),));conn.commit();conn.close()


def mirror_cloud_love_profile(local_user_id,row):
    if not row:return None
    owner=_love_primary_user_id(local_user_id);space=str(row["love_space_id"]);conn=get_conn()
    conn.execute("""INSERT INTO relationship_profiles(
        user_id,partner_name,partner_gender,partner_age,relationship_type,start_date,updated_at,cloud_id,cloud_updated_at
      ) VALUES(?,?,?,?,?,?,?,?,?)
      ON CONFLICT(user_id) DO UPDATE SET partner_name=excluded.partner_name,
        partner_gender=excluded.partner_gender,partner_age=excluded.partner_age,
        relationship_type=excluded.relationship_type,start_date=excluded.start_date,
        updated_at=excluded.updated_at,cloud_id=excluded.cloud_id,cloud_updated_at=excluded.cloud_updated_at""",
      (owner,row.get("partner_name") or "",row.get("partner_gender") or "female",int(row.get("partner_age") or 25),
       row.get("relationship_type") or "dating",row.get("start_date"),row.get("updated_at") or local_now().isoformat(),
       space,row.get("updated_at")))
    conn.commit();conn.close();return owner


def mirror_cloud_love_cycle_settings(local_user_id,row):
    if not row:return None
    owner=_love_primary_user_id(local_user_id);space=str(row["love_space_id"]);conn=get_conn()
    conn.execute("""INSERT INTO menstrual_settings(
        user_id,tracked_person,last_period_start,cycle_length,period_length,updated_at,cloud_id,cloud_updated_at
      ) VALUES(?,?,?,?,?,?,?,?)
      ON CONFLICT(user_id) DO UPDATE SET tracked_person=excluded.tracked_person,
        last_period_start=excluded.last_period_start,cycle_length=excluded.cycle_length,
        period_length=excluded.period_length,updated_at=excluded.updated_at,
        cloud_id=excluded.cloud_id,cloud_updated_at=excluded.cloud_updated_at""",
      (owner,row.get("tracked_person") or "partner",row.get("last_period_start"),int(row.get("cycle_length") or 28),
       int(row.get("period_length") or 5),row.get("updated_at") or local_now().isoformat(),space,row.get("updated_at")))
    conn.commit();conn.close();return owner


def _love_creator_local(local_user_id,row):
    return get_local_user_id_for_cloud(row.get("created_by")) or local_user_id


def mirror_cloud_love_record(local_user_id,record_type,row):
    if not row:return None
    uid=_love_creator_local(local_user_id,row);cloud_id=str(row["id"]);updated=row.get("updated_at") or row.get("created_at")
    conn=get_conn();existing=conn.execute(f"SELECT id FROM {_LOVE_RECORD_TABLES[record_type]} WHERE cloud_id=?",(cloud_id,)).fetchone()
    if record_type=="event":
        values=(uid,row.get("title") or "",row.get("event_date"),row.get("category") or "date",row.get("notes") or "",row.get("created_at"),cloud_id,updated)
        if existing:conn.execute("""UPDATE relationship_events SET user_id=?,title=?,event_date=?,category=?,notes=?,created_at=?,cloud_id=?,cloud_updated_at=? WHERE id=?""",(*values,existing["id"]));local_id=existing["id"]
        else:local_id=conn.execute("""INSERT INTO relationship_events(user_id,title,event_date,category,notes,created_at,cloud_id,cloud_updated_at) VALUES(?,?,?,?,?,?,?,?)""",values).lastrowid
    elif record_type=="memory":
        values=(uid,row.get("title") or "",row.get("memory_date"),row.get("notes") or "",row.get("created_at"),cloud_id,updated)
        if existing:conn.execute("""UPDATE relationship_memories SET user_id=?,title=?,memory_date=?,notes=?,created_at=?,cloud_id=?,cloud_updated_at=? WHERE id=?""",(*values,existing["id"]));local_id=existing["id"]
        else:local_id=conn.execute("""INSERT INTO relationship_memories(user_id,title,memory_date,notes,created_at,cloud_id,cloud_updated_at) VALUES(?,?,?,?,?,?,?)""",values).lastrowid
    elif record_type=="checkin":
        values=(uid,row.get("checkin_date"),int(row.get("my_mood") or 3),int(row.get("partner_mood") or 3),int(row.get("connection_score") or 3),row.get("note") or "",row.get("created_at"),cloud_id,updated)
        same=existing or conn.execute("SELECT id FROM relationship_checkins WHERE user_id=? AND checkin_date=?",values[:2]).fetchone()
        if same:conn.execute("""UPDATE relationship_checkins SET user_id=?,checkin_date=?,my_mood=?,partner_mood=?,connection_score=?,note=?,created_at=?,cloud_id=?,cloud_updated_at=? WHERE id=?""",(*values,same["id"]));local_id=same["id"]
        else:local_id=conn.execute("""INSERT INTO relationship_checkins(user_id,checkin_date,my_mood,partner_mood,connection_score,note,created_at,cloud_id,cloud_updated_at) VALUES(?,?,?,?,?,?,?,?,?)""",values).lastrowid
    elif record_type=="prompt_response":
        values=(uid,row.get("prompt_key") or "",row.get("category") or "connection",row.get("prompt_text") or "",row.get("my_answer") or "",row.get("partner_answer") or "",row.get("response_date"),row.get("created_at"),cloud_id,updated)
        if existing:conn.execute("""UPDATE relationship_prompt_responses SET user_id=?,prompt_key=?,category=?,prompt_text=?,my_answer=?,partner_answer=?,response_date=?,created_at=?,cloud_id=?,cloud_updated_at=? WHERE id=?""",(*values,existing["id"]));local_id=existing["id"]
        else:local_id=conn.execute("""INSERT INTO relationship_prompt_responses(user_id,prompt_key,category,prompt_text,my_answer,partner_answer,response_date,created_at,cloud_id,cloud_updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)""",values).lastrowid
    elif record_type=="weekly_review":
        values=(uid,row.get("week_start"),row.get("appreciation") or "",row.get("wins") or "",row.get("support_needed") or "",row.get("shared_intention") or "",row.get("created_at"),updated or row.get("created_at"),cloud_id,updated)
        same=existing or conn.execute("SELECT id FROM relationship_weekly_reviews WHERE user_id=? AND week_start=?",values[:2]).fetchone()
        if same:conn.execute("""UPDATE relationship_weekly_reviews SET user_id=?,week_start=?,appreciation=?,wins=?,support_needed=?,shared_intention=?,created_at=?,updated_at=?,cloud_id=?,cloud_updated_at=? WHERE id=?""",(*values,same["id"]));local_id=same["id"]
        else:local_id=conn.execute("""INSERT INTO relationship_weekly_reviews(user_id,week_start,appreciation,wins,support_needed,shared_intention,created_at,updated_at,cloud_id,cloud_updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)""",values).lastrowid
    elif record_type=="bucket_item":
        values=(uid,row.get("title") or "",row.get("category") or "dream",row.get("target_date"),1 if row.get("is_done") else 0,row.get("created_at"),row.get("completed_at"),cloud_id,updated)
        if existing:conn.execute("""UPDATE relationship_bucket_items SET user_id=?,title=?,category=?,target_date=?,is_done=?,created_at=?,completed_at=?,cloud_id=?,cloud_updated_at=? WHERE id=?""",(*values,existing["id"]));local_id=existing["id"]
        else:local_id=conn.execute("""INSERT INTO relationship_bucket_items(user_id,title,category,target_date,is_done,created_at,completed_at,cloud_id,cloud_updated_at) VALUES(?,?,?,?,?,?,?,?,?)""",values).lastrowid
    elif record_type=="cycle":
        values=(uid,row.get("start_date"),row.get("end_date"),row.get("notes") or "",row.get("created_at"),cloud_id,updated)
        if existing:conn.execute("""UPDATE menstrual_cycles SET user_id=?,start_date=?,end_date=?,notes=?,created_at=?,cloud_id=?,cloud_updated_at=? WHERE id=?""",(*values,existing["id"]));local_id=existing["id"]
        else:local_id=conn.execute("""INSERT INTO menstrual_cycles(user_id,start_date,end_date,notes,created_at,cloud_id,cloud_updated_at) VALUES(?,?,?,?,?,?,?)""",values).lastrowid
    else:conn.close();raise ValueError("unsupported Love record type")
    conn.commit();conn.close();return local_id


def mirror_cloud_love_favorite(local_user_id,row,favorite=True,prompt_key=None):
    uid=_love_creator_local(local_user_id,row or {});key=(row or {}).get("prompt_key") or prompt_key
    if not key:return
    conn=get_conn()
    if favorite:
        conn.execute("""INSERT INTO relationship_prompt_favorites(user_id,prompt_key,created_at,cloud_id,cloud_updated_at)
          VALUES(?,?,?,?,?) ON CONFLICT(user_id,prompt_key) DO UPDATE SET
          cloud_id=excluded.cloud_id,cloud_updated_at=excluded.cloud_updated_at""",
          (uid,key,(row or {}).get("created_at") or local_now().isoformat(),str((row or {}).get("id") or "") or None,(row or {}).get("created_at")))
    else:conn.execute("DELETE FROM relationship_prompt_favorites WHERE user_id=? AND prompt_key=?",(uid,key))
    conn.commit();conn.close()


def mirror_cloud_love_bundle(local_user_id,bundle):
    profile=(bundle.get("profile") or []);settings=(bundle.get("cycle_settings") or [])
    if profile:mirror_cloud_love_profile(local_user_id,profile[0])
    if settings:mirror_cloud_love_cycle_settings(local_user_id,settings[0])
    mapping={"events":"event","memories":"memory","checkins":"checkin",
             "prompt_responses":"prompt_response","weekly_reviews":"weekly_review",
             "bucket_items":"bucket_item","cycles":"cycle"}
    scope=_love_scope_user_ids(local_user_id);marks=",".join("?" for _ in scope);conn=get_conn()
    for bundle_key,record_type in mapping.items():
        rows=bundle.get(bundle_key) or []
        for row in rows:mirror_cloud_love_record(local_user_id,record_type,row)
        ids=[str(row["id"]) for row in rows];table=_LOVE_RECORD_TABLES[record_type]
        if ids:
            cloud_marks=",".join("?" for _ in ids);conn.execute(f"DELETE FROM {table} WHERE cloud_id IS NOT NULL AND user_id IN ({marks}) AND cloud_id NOT IN ({cloud_marks})",(*scope,*ids))
        else:conn.execute(f"DELETE FROM {table} WHERE cloud_id IS NOT NULL AND user_id IN ({marks})",tuple(scope))
    favorite_rows=bundle.get("prompt_favorites") or []
    for row in favorite_rows:mirror_cloud_love_favorite(local_user_id,row,True)
    fav_ids=[str(row["id"]) for row in favorite_rows]
    if fav_ids:
        fm=",".join("?" for _ in fav_ids);conn.execute(f"DELETE FROM relationship_prompt_favorites WHERE cloud_id IS NOT NULL AND user_id IN ({marks}) AND cloud_id NOT IN ({fm})",(*scope,*fav_ids))
    else:conn.execute(f"DELETE FROM relationship_prompt_favorites WHERE cloud_id IS NOT NULL AND user_id IN ({marks})",tuple(scope))
    conn.commit();conn.close()
    return {"records":sum(len(bundle.get(key) or []) for key in mapping)+len(favorite_rows),"profile":bool(profile),"cycle_settings":bool(settings)}


def build_local_love_migration_bundle(local_user_id):
    return {
      "profile":get_relationship_profile(local_user_id),"events":get_relationship_events(local_user_id),
      "memories":get_relationship_memories(local_user_id),"checkins":get_relationship_checkins(local_user_id,1000),
      "prompt_responses":get_relationship_prompt_responses(local_user_id,1000),
      "prompt_favorites":sorted(get_relationship_prompt_favorites(local_user_id)),
      "weekly_reviews":get_relationship_weekly_reviews(local_user_id,1000),
      "bucket_items":get_relationship_bucket_items(local_user_id),"cycle_settings":get_menstrual_settings(local_user_id),
      "cycles":get_menstrual_cycles(local_user_id,1000),
    }


def save_cloud_entity_map(local_user_id,entity_type,local_id,cloud_id):
    conn=get_conn();conn.execute("""INSERT INTO cloud_entity_map(local_user_id,entity_type,local_id,cloud_id,updated_at)
        VALUES(?,?,?,?,datetime('now')) ON CONFLICT(local_user_id,entity_type,local_id)
        DO UPDATE SET cloud_id=excluded.cloud_id,updated_at=datetime('now')""",
        (local_user_id,entity_type,str(local_id),str(cloud_id)));conn.commit();conn.close()


def get_cloud_id(entity_type,local_id):
    if local_id is None:return None
    conn=get_conn();row=conn.execute("SELECT cloud_id FROM cloud_entity_map WHERE entity_type=? AND local_id=?",
                                     (entity_type,str(local_id))).fetchone();conn.close();return row["cloud_id"] if row else None


def save_cloud_conversation(local_user_id,other_local_user_id,cloud_id,updated_at=None):
    conn=get_conn();conn.execute("""INSERT INTO cloud_conversations(cloud_id,local_user_id,other_local_user_id,updated_at)
        VALUES(?,?,?,?) ON CONFLICT(cloud_id) DO UPDATE SET local_user_id=excluded.local_user_id,
        other_local_user_id=excluded.other_local_user_id,updated_at=COALESCE(excluded.updated_at,cloud_conversations.updated_at)""",
        (str(cloud_id),local_user_id,other_local_user_id,updated_at));conn.commit();conn.close()


def get_cloud_conversation(local_user_id,other_local_user_id):
    conn=get_conn();row=conn.execute("SELECT * FROM cloud_conversations WHERE local_user_id=? AND other_local_user_id=?",
        (local_user_id,other_local_user_id)).fetchone();conn.close();return dict(row) if row else {}


def cache_cloud_messages(messages):
    conn=get_conn()
    for message in messages:
        if message.get("client_message_id"):
            conn.execute("DELETE FROM cloud_messages WHERE sender_cloud_id=? AND client_message_id=? AND sync_status='pending'",
                         (str(message["sender_id"]),str(message["client_message_id"])))
        conn.execute("""INSERT INTO cloud_messages(cloud_id,conversation_cloud_id,sender_cloud_id,client_message_id,
            reply_to_cloud_id,body,created_at,edited_at,deleted_at,sync_status)
            VALUES(?,?,?,?,?,?,?,?,?,'synced') ON CONFLICT(cloud_id) DO UPDATE SET
            reply_to_cloud_id=excluded.reply_to_cloud_id,body=excluded.body,created_at=excluded.created_at,
            edited_at=excluded.edited_at,deleted_at=excluded.deleted_at,sync_status='synced'""",
            (str(message["id"]),str(message["conversation_id"]),str(message["sender_id"]),
             str(message.get("client_message_id") or ""),message.get("reply_to_id"),message.get("body") or "",
             message.get("created_at") or datetime.now().isoformat(),message.get("edited_at"),message.get("deleted_at")))
        if message.get("deleted_at"):
            conn.execute("UPDATE chat_attachments_cache SET deleted_at=? WHERE message_cloud_id=?",(message.get("deleted_at"),str(message["id"])))
    conn.commit();conn.close()


def cache_cloud_message_reactions(rows,message_ids=None):
    conn=get_conn();ids=[str(value) for value in (message_ids or [])]
    if ids:
        marks=','.join('?' for _ in ids);conn.execute(f"DELETE FROM cloud_message_reactions WHERE message_cloud_id IN ({marks})",tuple(ids))
    for row in rows:
        conn.execute("""INSERT INTO cloud_message_reactions(message_cloud_id,user_cloud_id,reaction,updated_at)
          VALUES(?,?,?,?) ON CONFLICT(message_cloud_id,user_cloud_id) DO UPDATE SET
          reaction=excluded.reaction,updated_at=excluded.updated_at""",
          (str(row["message_id"]),str(row["user_id"]),row.get("reaction") or "",row.get("updated_at")))
    conn.commit();conn.close()


def get_cached_cloud_messages(conversation_cloud_id,limit=200):
    conn=get_conn();rows=conn.execute("SELECT * FROM cloud_messages WHERE conversation_cloud_id=? ORDER BY created_at DESC LIMIT ?",
        (str(conversation_cloud_id),limit)).fetchall();result=[dict(r) for r in reversed(rows)]
    for message in result:
        reactions=conn.execute("SELECT user_cloud_id,reaction FROM cloud_message_reactions WHERE message_cloud_id=?",(message["cloud_id"],)).fetchall()
        message["reactions"]={row["user_cloud_id"]:row["reaction"] for row in reactions}
        attachments=conn.execute("SELECT * FROM chat_attachments_cache WHERE message_cloud_id=? AND deleted_at IS NULL ORDER BY id",(message["cloud_id"],)).fetchall()
        message["attachments"]=[dict(row) for row in attachments]
    conn.close();return result


def cache_pending_cloud_message(conversation_cloud_id,sender_cloud_id,client_message_id,body,reply_to_cloud_id=None):
    conn=get_conn();conn.execute("""INSERT OR REPLACE INTO cloud_messages
        (cloud_id,conversation_cloud_id,sender_cloud_id,client_message_id,reply_to_cloud_id,body,created_at,sync_status)
        VALUES(?,?,?,?,?,?,?,'pending')""",
        (f"pending:{client_message_id}",conversation_cloud_id,sender_cloud_id,client_message_id,
         reply_to_cloud_id,body,local_now().isoformat()))
    conn.commit();conn.close()


def create_pending_chat_attachment(local_user_id,cloud_id,filename,mime_type,file_data,sha256,width=None,height=None,thumbnail_data=None):
    blob=bytes(file_data)
    if not blob or len(blob)>10*1024*1024:raise ValueError("attachment size invalid")
    conn=get_conn();cur=conn.execute("""INSERT INTO chat_attachments_cache(
      cloud_id,local_user_id,original_filename,mime_type,size_bytes,width,height,sha256,file_data,thumbnail_data,sync_status)
      VALUES(?,?,?,?,?,?,?,?,?,?,'pending')""",
      (str(cloud_id),local_user_id,str(filename)[:160],mime_type,len(blob),width,height,sha256,blob,bytes(thumbnail_data) if thumbnail_data else None))
    conn.commit();local_id=cur.lastrowid;conn.close();return local_id


def get_chat_attachment(local_id):
    conn=get_conn();row=conn.execute("SELECT * FROM chat_attachments_cache WHERE id=?",(local_id,)).fetchone();conn.close()
    if not row:return {}
    result=dict(row)
    if result.get("file_data") is not None:result["file_data"]=bytes(result["file_data"])
    return result


def get_chat_attachments_for_message(message_id,cloud=True):
    column="message_cloud_id" if cloud else "local_message_id"
    conn=get_conn();rows=conn.execute(f"SELECT * FROM chat_attachments_cache WHERE {column}=? AND deleted_at IS NULL ORDER BY id",(str(message_id) if cloud else message_id,)).fetchall();conn.close()
    return [dict(row) for row in rows]


def link_local_chat_attachments(local_user_id,local_message_id,attachment_ids):
    if not attachment_ids:return
    conn=get_conn();marks=','.join('?' for _ in attachment_ids)
    conn.execute(f"UPDATE chat_attachments_cache SET local_message_id=?,sync_status='local' WHERE local_user_id=? AND id IN ({marks})",
                 (local_message_id,local_user_id,*attachment_ids));conn.commit();conn.close()


def mark_chat_attachments_synced(local_user_id,message_cloud_id,conversation_cloud_id,rows):
    conn=get_conn()
    for row in rows:
        conn.execute("""UPDATE chat_attachments_cache SET message_cloud_id=?,conversation_cloud_id=?,
          uploader_cloud_id=?,storage_path=?,original_filename=?,mime_type=?,size_bytes=?,width=?,height=?,sha256=?,
          sync_status='synced',deleted_at=? WHERE local_user_id=? AND cloud_id=?""",
          (str(message_cloud_id),str(conversation_cloud_id),row.get("uploader_id"),row.get("storage_path"),
           row.get("original_filename") or "attachment",row.get("mime_type"),int(row.get("size_bytes") or 0),
           row.get("width"),row.get("height"),row.get("sha256"),row.get("deleted_at"),local_user_id,str(row["id"])))
    conn.commit();conn.close()


def cache_cloud_chat_attachments(local_user_id,rows):
    conn=get_conn()
    for row in rows:
        existing=conn.execute("SELECT id FROM chat_attachments_cache WHERE cloud_id=?",(str(row["id"]),)).fetchone()
        values=(local_user_id,str(row.get("conversation_id") or ""),str(row.get("message_id") or ""),row.get("uploader_id"),
                row.get("storage_path"),row.get("original_filename") or "attachment",row.get("mime_type"),
                int(row.get("size_bytes") or 0),row.get("width"),row.get("height"),row.get("sha256") or "",row.get("created_at"),row.get("deleted_at"))
        if existing:conn.execute("""UPDATE chat_attachments_cache SET local_user_id=?,conversation_cloud_id=?,message_cloud_id=?,
          uploader_cloud_id=?,storage_path=?,original_filename=?,mime_type=?,size_bytes=?,width=?,height=?,sha256=?,
          created_at=?,deleted_at=?,sync_status='synced' WHERE id=?""",(*values,existing["id"]))
        else:conn.execute("""INSERT INTO chat_attachments_cache(local_user_id,conversation_cloud_id,message_cloud_id,
          uploader_cloud_id,storage_path,original_filename,mime_type,size_bytes,width,height,sha256,created_at,deleted_at,cloud_id,sync_status)
          VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,'synced')""",(*values,str(row["id"])))
    conn.commit();conn.close()


def cache_chat_attachment_data(cloud_id,file_data,thumbnail_data=None):
    blob=bytes(file_data);thumb=bytes(thumbnail_data) if thumbnail_data else None;conn=get_conn();conn.execute("UPDATE chat_attachments_cache SET file_data=?,thumbnail_data=COALESCE(?,thumbnail_data) WHERE cloud_id=?",(blob,thumb,str(cloud_id)));conn.commit();conn.close()


def delete_pending_chat_attachments(local_user_id,attachment_ids):
    if not attachment_ids:return
    conn=get_conn();marks=','.join('?' for _ in attachment_ids);conn.execute(f"DELETE FROM chat_attachments_cache WHERE local_user_id=? AND sync_status='pending' AND id IN ({marks})",(local_user_id,*attachment_ids));conn.commit();conn.close()


def cache_cloud_guild_messages(rows):
    conn=get_conn()
    for row in rows:
        conn.execute("""INSERT INTO cloud_guild_messages_cache(cloud_id,guild_cloud_id,sender_cloud_id,client_message_id,
          reply_to_cloud_id,body,created_at,edited_at,deleted_at) VALUES(?,?,?,?,?,?,?,?,?)
          ON CONFLICT(cloud_id) DO UPDATE SET reply_to_cloud_id=excluded.reply_to_cloud_id,body=excluded.body,
          edited_at=excluded.edited_at,deleted_at=excluded.deleted_at""",
          (str(row["id"]),str(row["guild_id"]),str(row["sender_id"]),str(row.get("client_message_id") or ""),
           row.get("reply_to_id"),row.get("body") or "",row.get("created_at") or local_now().isoformat(),row.get("edited_at"),row.get("deleted_at")))
    conn.commit();conn.close()


def cache_cloud_guild_reactions(rows,message_ids=None):
    conn=get_conn();ids=[str(value) for value in (message_ids or [])]
    if ids:
        marks=','.join('?' for _ in ids);conn.execute(f"DELETE FROM cloud_guild_reactions_cache WHERE message_cloud_id IN ({marks})",tuple(ids))
    for row in rows:
        conn.execute("""INSERT INTO cloud_guild_reactions_cache(message_cloud_id,user_cloud_id,reaction,updated_at)
          VALUES(?,?,?,?) ON CONFLICT(message_cloud_id,user_cloud_id) DO UPDATE SET reaction=excluded.reaction,updated_at=excluded.updated_at""",
          (str(row["message_id"]),str(row["user_id"]),row.get("reaction") or "",row.get("updated_at")))
    conn.commit();conn.close()


def get_cached_cloud_guild_messages(guild_id,limit=200):
    conn=get_conn();rows=conn.execute("SELECT * FROM cloud_guild_messages_cache WHERE guild_cloud_id=? ORDER BY created_at DESC LIMIT ?",(str(guild_id),limit)).fetchall();result=[dict(row) for row in reversed(rows)]
    for message in result:
        reactions=conn.execute("SELECT user_cloud_id,reaction FROM cloud_guild_reactions_cache WHERE message_cloud_id=?",(message["cloud_id"],)).fetchall()
        message["reactions"]={row["user_cloud_id"]:row["reaction"] for row in reactions}
    conn.close();return result


def cache_cloud_conversation_summaries(local_user_id,rows):
    conn=get_conn()
    for row in rows:
        other_local=get_local_user_id_for_cloud(row.get("other_user_id"))
        if not other_local:continue
        conn.execute("""INSERT INTO cloud_conversations(cloud_id,local_user_id,other_local_user_id,updated_at,
          unread_count,last_message_at,last_message_preview) VALUES(?,?,?,?,?,?,?)
          ON CONFLICT(cloud_id) DO UPDATE SET other_local_user_id=excluded.other_local_user_id,
          updated_at=excluded.updated_at,unread_count=excluded.unread_count,
          last_message_at=excluded.last_message_at,last_message_preview=excluded.last_message_preview""",
          (str(row["conversation_id"]),local_user_id,other_local,row.get("conversation_updated_at"),
           int(row.get("unread_count") or 0),row.get("last_message_at"),row.get("last_message_preview") or ""))
    conn.commit();conn.close()


def get_cloud_unread_count(local_user_id,other_local_user_id):
    row=get_cloud_conversation(local_user_id,other_local_user_id);return int(row.get("unread_count") or 0)


def mark_cloud_conversation_read_local(conversation_cloud_id):
    conn=get_conn();conn.execute("UPDATE cloud_conversations SET unread_count=0 WHERE cloud_id=?",(str(conversation_cloud_id),));conn.commit();conn.close()


def cache_cloud_presence(rows):
    conn=get_conn()
    for row in rows:
        conn.execute("""INSERT INTO cloud_presence_cache(cloud_user_id,status,device_name,last_seen_at,updated_at)
            VALUES(?,?,?,?,?) ON CONFLICT(cloud_user_id) DO UPDATE SET status=excluded.status,
            device_name=excluded.device_name,last_seen_at=excluded.last_seen_at,updated_at=excluded.updated_at""",
            (str(row["user_id"]),row.get("status","offline"),row.get("device_name","") or "",row.get("last_seen_at"),row.get("updated_at")))
    conn.commit();conn.close()


def get_cached_presence(cloud_user_id):
    conn=get_conn();row=conn.execute("SELECT * FROM cloud_presence_cache WHERE cloud_user_id=?",(str(cloud_user_id),)).fetchone();conn.close();return dict(row) if row else {}


def cache_cloud_notifications(local_user_id,rows):
    conn=get_conn()
    for row in rows:
        payload=row.get("payload") or {};cloud_id=str(row["id"]);is_read=1 if row.get("is_read") else 0
        conn.execute("""INSERT INTO cloud_notifications_cache(cloud_id,local_user_id,notification_type,actor_cloud_id,
            entity_type,entity_id,payload,is_read,created_at) VALUES(?,?,?,?,?,?,?,?,?)
            ON CONFLICT(cloud_id) DO UPDATE SET is_read=excluded.is_read,payload=excluded.payload""",
            (cloud_id,local_user_id,row.get("notification_type","info"),row.get("actor_id"),
             row.get("entity_type"),str(row.get("entity_id") or ""),json.dumps(payload),is_read,row.get("created_at")))
        preview=payload.get("preview") or payload.get("message") or row.get("notification_type","Cloud update").replace("_"," ").title()
        existing=conn.execute("SELECT id FROM notifications WHERE cloud_id=?",(cloud_id,)).fetchone()
        if existing:
            conn.execute("UPDATE notifications SET message=?,is_read=?,created_at=? WHERE id=?",
                         (preview,is_read,row.get("created_at") or local_now().isoformat(),existing["id"]))
        else:
            conn.execute("INSERT INTO notifications(user_id,message,type,is_read,created_at,cloud_id) VALUES(?,?,?,?,?,?)",
                         (local_user_id,preview,"info",is_read,row.get("created_at") or local_now().isoformat(),cloud_id))
    conn.commit();conn.close()


def enqueue_productivity_event(user_id,event_type,source_local_id,idempotency_key,completed_at,payload=None):
    if not get_cloud_user_link(user_id):return {"ok":False,"code":"not_linked"}
    return enqueue_sync(user_id,"productivity_event",source_local_id,"record",{
        "event_type":event_type,"source_local_id":str(source_local_id),"idempotency_key":idempotency_key,
        "completed_at":completed_at,"device_id":"CraftLife Desktop","payload":payload or {}})


def upsert_cloud_profile(cloud_user_id,profile):
    existing=get_local_user_id_for_cloud(cloud_user_id)
    if existing:
        update_user(existing,display_name=profile.get("display_name") or profile.get("username") or "Cloud User",
                    bio=profile.get("bio") or "",avatar_class=profile.get("avatar_class") or "warrior",
                    avatar_color=profile.get("avatar_color") or "#5a8a2e",avatar_emoji=profile.get("avatar_emoji") or "⚔️")
        return existing
    base=(profile.get("username") or f"cloud_{str(cloud_user_id)[:8]}").lower().strip()
    base="".join(ch for ch in base if ch.isalnum() or ch in "_-") or f"cloud_{str(cloud_user_id)[:8]}"
    conn=get_conn();username=base;n=2
    while conn.execute("SELECT 1 FROM users WHERE username=?",(username,)).fetchone():
        username=f"{base}_{n}";n+=1
    cur=conn.execute("""INSERT INTO users(username,password_hash,display_name,bio,avatar_class,avatar_color,avatar_emoji,
        cloud_user_id,is_cloud_shadow) VALUES(?,?,?,?,?,?,?,?,1)""",
        (username,f"!cloud-only!{cloud_user_id}",profile.get("display_name") or username,profile.get("bio") or "",
         profile.get("avatar_class") or "warrior",profile.get("avatar_color") or "#5a8a2e",profile.get("avatar_emoji") or "⚔️",cloud_user_id))
    conn.commit();local_id=cur.lastrowid;conn.close();return local_id


def prune_cloud_friendships(local_user_id,active_cloud_ids):
    conn=get_conn()
    if active_cloud_ids:
        marks=','.join('?' for _ in active_cloud_ids)
        conn.execute(f"DELETE FROM friends WHERE cloud_id IS NOT NULL AND (user_id=? OR friend_id=?) AND cloud_id NOT IN ({marks})",
                     (local_user_id,local_user_id,*active_cloud_ids))
    else:conn.execute("DELETE FROM friends WHERE cloud_id IS NOT NULL AND (user_id=? OR friend_id=?)",(local_user_id,local_user_id))
    conn.commit();conn.close()


def mirror_cloud_friendship(cloud_row,profile_by_cloud):
    requester=upsert_cloud_profile(cloud_row["requester_id"],profile_by_cloud[cloud_row["requester_id"]])
    addressee=upsert_cloud_profile(cloud_row["addressee_id"],profile_by_cloud[cloud_row["addressee_id"]])
    conn=get_conn();row=conn.execute("SELECT id FROM friends WHERE cloud_id=?",(cloud_row["id"],)).fetchone()
    status=cloud_row.get("status","pending")
    if row:
        conn.execute("UPDATE friends SET user_id=?,friend_id=?,status=?,action_user_id=? WHERE id=?",
                     (requester,addressee,status,requester,row["id"]))
    else:
        # Remove an equivalent non-cloud mirror before insertion to satisfy directional uniqueness.
        eq=conn.execute("SELECT id FROM friends WHERE user_id=? AND friend_id=?",(requester,addressee)).fetchone()
        if eq: conn.execute("UPDATE friends SET status=?,cloud_id=? WHERE id=?",(status,cloud_row["id"],eq["id"]))
        else: conn.execute("INSERT INTO friends(user_id,friend_id,status,action_user_id,cloud_id) VALUES(?,?,?,?,?)",
                           (requester,addressee,status,requester,cloud_row["id"]))
    conn.commit();conn.close()


def mirror_cloud_couple(cloud_row,profile_by_cloud):
    a=upsert_cloud_profile(cloud_row["user_a_id"],profile_by_cloud[cloud_row["user_a_id"]]);b=upsert_cloud_profile(cloud_row["user_b_id"],profile_by_cloud[cloud_row["user_b_id"]])
    a,b=sorted((a,b));requested=get_local_user_id_for_cloud(cloud_row["requested_by"])
    status="cancelled" if cloud_row.get("status")=="ended" else cloud_row.get("status","pending")
    conn=get_conn();row=conn.execute("SELECT id FROM couple_relationships WHERE cloud_id=?",(cloud_row["id"],)).fetchone()
    cloud_status=cloud_row.get("status");ended_at=cloud_row.get("ended_at");grace_ends=cloud_row.get("grace_ends_at")
    if row: conn.execute("UPDATE couple_relationships SET user_a_id=?,user_b_id=?,requested_by=?,status=?,responded_at=?,cloud_status=?,cloud_ended_at=?,cloud_grace_ends_at=? WHERE id=?",
                         (a,b,requested or a,status,cloud_row.get("responded_at"),cloud_status,ended_at,grace_ends,row["id"]));local_rel=row["id"]
    else:
        pair=conn.execute("SELECT id FROM couple_relationships WHERE user_a_id=? AND user_b_id=?",(a,b)).fetchone()
        if pair: conn.execute("UPDATE couple_relationships SET cloud_id=?,requested_by=?,status=?,responded_at=?,cloud_status=?,cloud_ended_at=?,cloud_grace_ends_at=? WHERE id=?",
                              (cloud_row["id"],requested or a,status,cloud_row.get("responded_at"),cloud_status,ended_at,grace_ends,pair["id"]));local_rel=pair["id"]
        else: local_rel=conn.execute("INSERT INTO couple_relationships(user_a_id,user_b_id,requested_by,status,responded_at,cloud_id,cloud_status,cloud_ended_at,cloud_grace_ends_at) VALUES(?,?,?,?,?,?,?,?,?)",
                                     (a,b,requested or a,status,cloud_row.get("responded_at"),cloud_row["id"],cloud_status,ended_at,grace_ends)).lastrowid
    conn.commit();conn.close();return local_rel


def mirror_cloud_love_space(cloud_space,local_relationship_id,member_local_ids):
    conn=get_conn();row=conn.execute("SELECT id FROM love_spaces WHERE cloud_id=?",(cloud_space["id"],)).fetchone()
    if row: local_space=row["id"]
    else:
        existing=conn.execute("SELECT id FROM love_spaces WHERE couple_relationship_id=?",(local_relationship_id,)).fetchone()
        if existing: local_space=existing["id"];conn.execute("UPDATE love_spaces SET cloud_id=? WHERE id=?",(cloud_space["id"],local_space))
        else: local_space=conn.execute("INSERT INTO love_spaces(couple_relationship_id,cloud_id) VALUES(?,?)",(local_relationship_id,cloud_space["id"])).lastrowid
    conn.executemany("INSERT OR IGNORE INTO love_space_members(love_space_id,user_id) VALUES(?,?)",[(local_space,uid) for uid in member_local_ids])
    conn.commit();conn.close();return local_space


def get_love_photo_owner(local_photo_id):
    conn=get_conn();row=conn.execute("SELECT owner_user_id FROM love_space_photos WHERE id=?",(local_photo_id,)).fetchone();conn.close();return row["owner_user_id"] if row else None


def mark_love_photo_synced(local_photo_id,cloud_id,storage_path):
    conn=get_conn();conn.execute("UPDATE love_space_photos SET cloud_id=?,cloud_storage_path=?,sync_status='synced' WHERE id=?",
                                 (cloud_id,storage_path,local_photo_id));conn.commit();conn.close()


def cache_cloud_profile_photo(local_user_id,image_data,mime_type,width,height):
    ok,code=_validate_stored_image(image_data,mime_type,width,height,PROFILE_PHOTO_MAX_BYTES)
    if not ok:return {"ok":False,"code":code}
    blob=bytes(image_data);conn=get_conn();conn.execute("""INSERT INTO user_profile_photos
        (user_id,image_data,mime_type,width,height,size_bytes,updated_at) VALUES(?,?,?,?,?,?,datetime('now'))
        ON CONFLICT(user_id) DO UPDATE SET image_data=excluded.image_data,mime_type=excluded.mime_type,
        width=excluded.width,height=excluded.height,size_bytes=excluded.size_bytes,updated_at=datetime('now')""",
        (local_user_id,blob,mime_type,width,height,len(blob)));conn.commit();conn.close();return {"ok":True}


def cache_cloud_gallery_photo(cloud_row,image_data,mime_type,width,height):
    uploader=get_local_user_id_for_cloud(cloud_row.get("uploader_id"));
    if not uploader:return {"ok":False,"code":"unknown_uploader"}
    ok,code=_validate_stored_image(image_data,mime_type,width,height,LOVE_PHOTO_MAX_BYTES)
    if not ok:return {"ok":False,"code":code}
    conn=get_conn();space=None
    if cloud_row.get("love_space_id"):
        row=conn.execute("SELECT id FROM love_spaces WHERE cloud_id=?",(cloud_row["love_space_id"],)).fetchone();space=row["id"] if row else None
    existing=conn.execute("SELECT id FROM love_space_photos WHERE cloud_id=?",(cloud_row["id"],)).fetchone();blob=bytes(image_data)
    values=(space,uploader,cloud_row.get("visibility","shared"),blob,mime_type,width,height,len(blob),
            cloud_row.get("caption") or "",cloud_row.get("photo_date"),cloud_row.get("id"),cloud_row.get("storage_path"))
    if existing:
        conn.execute("""UPDATE love_space_photos SET love_space_id=?,owner_user_id=?,visibility=?,image_data=?,mime_type=?,
            width=?,height=?,size_bytes=?,caption=?,photo_date=?,cloud_id=?,cloud_storage_path=?,sync_status='synced' WHERE id=?""",
            (*values,existing["id"]))
    else:
        conn.execute("""INSERT INTO love_space_photos(love_space_id,owner_user_id,visibility,image_data,mime_type,width,height,
            size_bytes,caption,photo_date,cloud_id,cloud_storage_path,sync_status) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,'synced')""",values)
    conn.commit();conn.close();return {"ok":True}


# ===========================================================================================================================#
if __name__ == "__main__":
    init_db()
    print("Database OK!")
