"""Life OS API helpers (sport, food, economy, notes, calendar, health). Wraps database.py."""
from __future__ import annotations

from datetime import date, datetime, timedelta

import database as db


def _today() -> str:
    return date.today().isoformat()


def _parse_day(qs, default_today=True) -> str:
    """Parse ?date=YYYY-MM-DD; clamp ke hari ini (parity HealthFoodPage:
    tombol ▶ dinonaktifkan untuk tanggal setelah hari ini)."""
    raw = None
    try:
        raw = (qs.get("date") or [None])[0]
    except Exception:
        raw = None
    try:
        d = datetime.strptime(str(raw), "%Y-%m-%d").date() if raw else date.today()
    except ValueError:
        d = date.today()
    if d > date.today():
        d = date.today()
    return d.isoformat()


def _monthday() -> str:
    """Format tanggal dipakai label UI (fallback bila library locale tak tersedia)."""
    return date.today().isoformat()


def _to_idr(uid: int, value) -> float:
    """Konversi nominal input Web UI (selalu dalam mata uang pilihan user) ke IDR.

    Parity dengan dialog PyQt (mis. AddEconomyDialog._save) yang memanggil
    db.convert_to_idr(...) sebelum menyimpan. Database SELALU menyimpan IDR;
    konversi dari/ke mata uang user hanya di batas input/output.
    """
    try:
        amt = float(value or 0)
    except (TypeError, ValueError):
        amt = 0.0
    try:
        curr = db.get_user_currency(uid)
        return float(db.convert_to_idr(amt, curr))
    except Exception:
        return amt


def map_sport(row: dict) -> dict:
    diff = (row.get("difficulty") or "medium").lower()
    if diff not in ("trivial", "easy", "medium", "hard", "epic"):
        diff = "medium"
    return {
        "id": str(row.get("id")),
        "sportType": row.get("sport_type") or "running",
        "sportName": row.get("name") or "",
        "icon": row.get("icon") or "🏃",
        "durationMinutes": int(row.get("duration_minutes") or 0),
        "caloriesBurned": int(row.get("calories_burned") or 0),
        "difficulty": diff,
        # tetap sediakan `intensity` untuk kompatibilitas UI lama
        "intensity": {"easy": "light", "medium": "moderate", "hard": "vigorous", "epic": "vigorous"}.get(diff, "moderate"),
        "xpReward": int(row.get("xp_reward") or 0),
        "goldReward": int(row.get("gold_reward") or 0),
        "sportPointsReward": int(row.get("sport_points_reward") or 0),
        "streak": int(row.get("streak") or 0),
        "notes": row.get("notes") or "",
        "date": row.get("last_done") or row.get("created_at") or _today(),
        "sportXpEarned": int(row.get("xp_reward") or 0),
        "done": bool(row.get("done_today")),
        "totalReps": int(getattr(db, "get_sport_rep_total")(row.get("user_id"), row.get("id")) or 0)
        if getattr(db, "get_sport_rep_total", None) else 0,
    }


def map_meal(row: dict) -> dict:
    mt = (row.get("meal_type") or "snack").lower()
    if mt not in ("breakfast", "lunch", "dinner", "snack"):
        mt = "snack"
    return {
        "id": str(row.get("id")),
        "mealType": mt,
        "foodName": row.get("name") or "",
        "icon": row.get("icon") or "🍽️",
        "portion": float(row.get("serving") or 1),
        "calories": int(row.get("calories") or 0),
        "protein": float(row.get("protein") or 0),
        "carbs": float(row.get("carbs") or 0),
        "fat": float(row.get("fat") or 0),
        "date": row.get("log_date") or _today(),
    }


def map_tx(row: dict) -> dict:
    t = row.get("type") or "expense"
    if t not in ("income", "expense"):
        t = "expense" if t in ("pengeluaran", "out") else "income"
    return {
        "id": str(row.get("id")),
        "type": t,
        "category": row.get("category") or "",
        "amount": float(row.get("amount") or 0),
        "date": row.get("date") or row.get("date_str") or _today(),
        "notes": row.get("notes") or "",
        "name": row.get("name") or "",
        "icon": row.get("icon") or "💰",
        "folderId": str(row.get("folder_id")) if row.get("folder_id") is not None else None,
    }


def map_debt(row: dict) -> dict:
    total = float(row.get("amount") or row.get("total_amount") or 0)
    remaining = float(row.get("remaining") or row.get("remaining_amount") or total)
    return {
        "id": str(row.get("id")),
        "title": row.get("name") or "",
        "type": "payable",
        "totalAmount": total,
        "remainingAmount": remaining,
        "dueDate": row.get("due_date") or "",
        "notes": row.get("notes") or "",
        "isPaid": bool(row.get("paid") or row.get("is_paid")),
    }


def map_note(row: dict) -> dict:
    return {
        "id": str(row.get("id")),
        "folderId": str(row["folder_id"]) if row.get("folder_id") else None,
        "title": row.get("title") or "",
        "content": row.get("content") or "",
        "isPinned": bool(row.get("pinned")),
        "isArchived": bool(row.get("archived") or row.get("is_archived")),
        "zoomLevel": int(row.get("zoom_level") or 100),
        "updatedAt": row.get("updated_at") or "",
    }


def map_folder(row: dict) -> dict:
    return {
        "id": str(row.get("id")),
        "name": row.get("name") or "",
        "icon": row.get("icon") or "📁",
        "parentId": str(row["parent_id"]) if row.get("parent_id") else None,
    }


def map_reminder(row: dict) -> dict:
    # Parity RemindersPage.load: status 🔔/🔕 dari is_active; ✅ dari triggered;
    # tampil reminder_datetime[:16]. Sound & repeat fields seluruhnya terbuka.
    dt = str(row.get("reminder_datetime") or "")
    time = dt[11:16] if len(dt) >= 16 else (row.get("time") or "08:00")
    rep = (row.get("repeat_type") or "none").lower()
    if rep not in ("none", "daily", "weekly", "custom"):
        rep = "none"
    return {
        "id": str(row.get("id")),
        "title": row.get("title") or "",
        "description": row.get("description") or "",
        "datetime": dt,
        "time": time,
        "repeat": rep,
        "repeatDays": row.get("repeat_days") or "",
        "isActive": bool(row.get("is_active", 1)),
        "triggered": bool(row.get("triggered")),
        "sound": row.get("sound_type") or "default",
        "soundFile": row.get("sound_file") or "",
    }


def map_health(row: dict) -> dict:
    mood = (row.get("mood") or "good").lower()
    if mood not in ("great", "good", "neutral", "tired", "stressed"):
        mood = "good"
    return {
        "id": str(row.get("id")),
        "date": row.get("log_date") or _today(),
        "steps": int(row.get("steps") or 0),
        "sleepHours": float(row.get("sleep_hours") or 0),
        "weightKg": row.get("weight_kg"),
        "heartRate": int(row.get("resting_hr") or 0) or None,
        "mood": mood,
        "notes": row.get("notes") or "",
    }


def _pomodoro_stats_safe(uid: int) -> dict:
    """Statistik PomodoroPage (today/total sessions+minutes) — aman kalau gagal."""
    try:
        s = db.get_pomodoro_stats(uid) or {}
        return {
            "todaySessions": int(s.get("today_sessions") or 0),
            "todayMinutes": int(s.get("today_minutes") or 0),
            "totalSessions": int(s.get("total_sessions") or 0),
            "totalMinutes": int(s.get("total_minutes") or 0),
        }
    except Exception:
        return {"todaySessions": 0, "todayMinutes": 0, "totalSessions": 0, "totalMinutes": 0}


def map_pomo(row: dict) -> dict:
    mins = int(row.get("duration_minutes") or row.get("minutes") or 0)
    return {
        "id": str(row.get("id")),
        "durationMinutes": mins,
        "completedAt": row.get("completed_at") or row.get("created_at") or "",
        "xpEarned": int(row.get("xp_earned") or row.get("xp") or mins * 2),
        "goldEarned": int(row.get("gold_earned") or row.get("gold") or 0),
        "label": row.get("task_name") or row.get("label") or "",
    }


def map_supply(row: dict) -> dict:
    return {
        "id": str(row.get("id")),
        "name": row.get("name") or "",
        "category": row.get("category") or "",
        "unit": row.get("unit") or "pcs",
        "stock": float(row.get("stock") or 0),
        "minStock": float(row.get("min_stock") or 0),
        "price": float(row.get("price") or 0),
        "location": row.get("location") or "",
        "notes": row.get("notes") or "",
    }


def map_task_folder(row: dict) -> dict:
    return {
        "id": str(row.get("id")),
        "name": row.get("name") or "",
        "icon": row.get("icon") or "📁",
        "mode": row.get("mode") or "habit",
        "color": "#10b981",
    }


def snapshot(uid: int) -> dict:
    try:
        sports = [map_sport(r) for r in db.get_sport_activities(uid)]
    except Exception:
        sports = []
    try:
        meals = [map_meal(r) for r in db.get_food_logs(uid)]
    except Exception:
        meals = []
    try:
        goal = int(db.get_water_goal(uid) or 2500)
        total = int(db.get_water_total(uid) or 0)
        water = {"date": _today(), "amountMl": total, "targetMl": goal}
    except Exception:
        water = {"date": _today(), "amountMl": 0, "targetMl": 2500}
    try:
        txs = [map_tx(r) for r in db.get_economy_items(uid)]
    except Exception:
        txs = []
    try:
        debts = [map_debt(r) for r in db.get_debts(uid, include_paid=True)]
    except Exception:
        debts = []
    try:
        notes = [map_note(r) for r in db.get_notes(uid, include_archived=True)]
    except Exception:
        notes = []
    try:
        note_folders = [map_folder(r) for r in db.get_note_folders(uid)]
    except Exception:
        note_folders = []
    try:
        rems = [map_reminder(r) for r in db.get_reminders(uid)]
    except Exception:
        rems = []
    try:
        health = [map_health(r) for r in db.get_health_logs(uid, days=30)]
    except Exception:
        health = []
    try:
        pomos = [map_pomo(r) for r in db.get_recent_pomodoros(uid, limit=20)]
    except Exception:
        pomos = []
    try:
        cal = db.get_calendar_notes(uid) or {}
        # get_calendar_notes returns a dict {note_date: note}; some callers pass
        # a rows-list. Handle both so calendarNotes is never silently empty.
        if isinstance(cal, dict):
            calendarNotes = [{"date": d, "note": v} for d, v in cal.items()]
        else:
            calendarNotes = [
                {"date": r.get("note_date") or r.get("date"), "note": r.get("note") or ""}
                for r in cal
            ]
    except Exception:
        calendarNotes = []
    try:
        supplies = [map_supply(r) for r in db.get_supplies_items(uid)]
    except Exception:
        supplies = []
    try:
        task_folders = []
        for mode in ("habit", "daily", "todo", "sport", "economy"):
            task_folders.extend(map_task_folder(r) for r in (db.get_task_folders(uid, mode) or []))
    except Exception:
        task_folders = []
    try:
        savings = [
            {
                "id": str(r.get("id")),
                "name": r.get("name") or "",
                "icon": r.get("icon") or "🏦",
                "targetAmount": float(r.get("target_amount") or 0),
                "currentAmount": float(r.get("current_amount") or 0),
                "targetDate": r.get("target_date") or "",
                "notes": r.get("notes") or "",
            }
            for r in (db.get_savings(uid) or [])
        ]
    except Exception:
        savings = []
    try:
        investments = [
            {
                "id": str(r.get("id")),
                "name": r.get("name") or "",
                "icon": r.get("icon") or "📈",
                "amount": float(r.get("amount") or 0),
                "notes": r.get("notes") or "",
            }
            for r in (db.get_investments(uid) or [])
        ]
    except Exception:
        investments = []
    try:
        subscriptions = [
            {
                "id": str(r.get("id")),
                "name": r.get("name") or "",
                "icon": r.get("icon") or "📅",
                "amount": float(r.get("amount") or 0),
                "dueDate": r.get("due_date") or "",
                "period": r.get("period") or "monthly",
                "notes": r.get("notes") or "",
            }
            for r in (db.get_subscriptions(uid) or [])
        ]
    except Exception:
        subscriptions = []
    try:
        debt_notes = [
            {
                "id": str(r.get("id")),
                "personName": r.get("person_name") or "",
                "amount": float(r.get("amount") or 0),
                "date": r.get("date") or "",
                "notes": r.get("notes") or "",
                "status": r.get("status") or "unpaid",
            }
            for r in (db.get_debt_notes(uid) or [])
        ]
    except Exception:
        debt_notes = []
    return {
        "sportLogs": sports,
        "mealLogs": meals,
        "waterLog": water,
        "transactions": txs,
        "debts": debts,
        "notes": notes,
        "noteFolders": note_folders,
        "reminders": rems,
        "healthLogs": health,
        "pomodoroSessions": pomos,
        "pomodoroStats": _pomodoro_stats_safe(uid),
        "calendarNotes": calendarNotes,
        # Heatmap data: jumlah task sukses per hari (28 hari terakhir).
        "dailyTaskCounts": db.get_daily_task_counts(uid, 28),
        "supplies": supplies,
        "taskFolders": task_folders,
        "savings": savings,
        "investments": investments,
        "subscriptions": subscriptions,
        "debtNotes": debt_notes,
    }


# ══════════════════════════════════════════════════════════════════════════
# P4 — HealthFoodPage parity: paket data per-tanggal + riwayat + ekspor
# ══════════════════════════════════════════════════════════════════════════

def _health_log_for_date(uid: int, log_date: str):
    try:
        conn = db.get_conn()
        row = conn.execute(
            "SELECT * FROM health_logs WHERE user_id=? AND log_date=?",
            (uid, log_date)).fetchone()
        conn.close()
        return dict(row) if row else None
    except Exception:
        return None


def _healthfood_day(uid: int, log_date: str) -> dict:
    """Paket lengkap HealthFoodPage untuk satu tanggal (parity load())."""
    try:
        summary = db.get_nutrition_summary(uid, log_date) or {}
        ngoals_raw = db.get_nutrition_goals(uid) or {}
        ngoals = {
            "calories": int(ngoals_raw.get("daily_calories") or 2000),
            "protein": float(ngoals_raw.get("daily_protein") or 50),
            "carbs": float(ngoals_raw.get("daily_carbs") or 250),
            "fat": float(ngoals_raw.get("daily_fat") or 70),
        }
    except Exception:
        summary, ngoals = {"calories": 0, "protein": 0, "carbs": 0, "fat": 0}, {"calories": 2000, "protein": 50, "carbs": 250, "fat": 70}
    try:
        water_total = int(db.get_water_total(uid, log_date) or 0)
        water_goal = int(db.get_water_goal(uid) or 2500)
        wlogs = []
        for w in (db.get_water_logs(uid, log_date) or []):
            created = str(w.get("created_at") or "")
            wlogs.append({
                "id": str(w.get("id")),
                "amountMl": int(w.get("amount_ml") or 0),
                "time": created[11:16] if len(created) >= 16 else "",
            })
    except Exception:
        water_total, water_goal, wlogs = 0, 2500, []
    try:
        flags = db.get_user(uid) or {}
        lang = flags.get("language") or "id"
        flogs = []
        for l in (db.get_food_logs(uid, log_date) or []):
            name = l.get("name") or ""
            try:
                name = db.get_food_name(l.get("name") or "", lang) or name
            except Exception:
                pass
            flogs.append({
                "id": str(l.get("id")),
                "name": name,
                "icon": l.get("icon") or "🍽️",
                "serving": float(l.get("serving") or 1),
                "calories": int(l.get("calories") or 0),
                "protein": float(l.get("protein") or 0),
                "carbs": float(l.get("carbs") or 0),
                "fat": float(l.get("fat") or 0),
                "mealType": (l.get("meal_type") or "snack"),
                "notes": l.get("notes") or "",
            })
    except Exception:
        flogs = []
    try:
        hgoals = db.get_health_goals(uid) or {}
    except Exception:
        hgoals = {}
    hlog = _health_log_for_date(uid, log_date)
    try:
        bmi = db.get_user_bmi_settings(uid) or {}
    except Exception:
        bmi = {}
    try:
        burned = int(db.get_total_calories_burned_today(uid, log_date) or 0)
    except Exception:
        burned = 0
    return {
        "ok": True,
        "date": log_date,
        "isToday": log_date == _today(),
        "nutrition": {
            "calories": float(summary.get("calories") or 0),
            "protein": float(summary.get("protein") or 0),
            "carbs": float(summary.get("carbs") or 0),
            "fat": float(summary.get("fat") or 0),
        },
        "goals": ngoals,
        "water": {"totalMl": water_total, "goalMl": water_goal, "logs": wlogs},
        "foodLogs": flogs,
        "healthLog": hlog,
        "healthGoals": {
            "dailySteps": int(hgoals.get("daily_steps") or 10000),
            "dailySleepHours": float(hgoals.get("daily_sleep_hours") or 7.0),
            "heightCm": float(hgoals.get("height_cm") or 170),
            "weightKg": float(hgoals.get("weight_kg") or 70),
        },
        "bmi": {
            "heightCm": float(bmi.get("height_cm") or 170),
            "weightKg": float(bmi.get("weight_kg") or 70),
            "age": int(bmi.get("age") or 25),
            "gender": str(bmi.get("gender") or "male").lower(),
            "activityFactor": float(bmi.get("activity_factor") or 1.375),
        },
        "caloriesBurned": burned,
        "netCalories": int(summary.get("calories") or 0) - burned,
    }


def _healthfood_history(uid: int) -> dict:
    """Riwayat 7 hari (parity _refresh_history_charts + _update_tips)."""
    try:
        logs = db.get_health_logs(uid, days=7) or []
    except Exception:
        logs = []
    n = max(1, len(logs))
    avg = {
        "steps": sum(int(l.get("steps") or 0) for l in logs) // n if logs else 0,
        "sleep": round(sum(float(l.get("sleep_hours") or 0) for l in logs) / n, 1) if logs else 0.0,
        "water": sum(int(l.get("water_ml") or 0) for l in logs) // n if logs else 0,
        "hr": sum(int(l.get("resting_hr") or 0) for l in logs) // n if logs else 0,
    }
    by_date = {str(l.get("log_date")): l for l in logs}
    weight_series, height_series = [], []
    try:
        hgoals = db.get_health_goals(uid) or {}
        height_cm = float(hgoals.get("height_cm") or 170)
    except Exception:
        height_cm = 170.0
    for i in range(7):
        d = (date.today() - timedelta(days=6 - i)).isoformat()
        l = by_date.get(d)
        w = float(l.get("weight_kg") or 0) if l and l.get("weight_kg") is not None else 0.0
        weight_series.append({"label": d[5:], "value": w})
        height_series.append({"label": d[5:], "value": height_cm})
    # Tips (parity _update_tips): dinamis berdasar net kalori hari ini + 1 statis acak
    try:
        today = _today()
        nutrition = db.get_nutrition_summary(uid, today) or {}
        burned = int(db.get_total_calories_burned_today(uid, today) or 0)
        net_cal = int(nutrition.get("calories") or 0) - burned
        goal_cal = int((db.get_nutrition_goals(uid) or {}).get("daily_calories") or 2000)
    except Exception:
        net_cal, goal_cal = 0, 2000
    if net_cal < 0:
        dynamic = "health_tip_calorie_deficit"
    elif goal_cal > 0 and net_cal > goal_cal * 1.1:
        dynamic = "health_tip_calorie_surplus"
    else:
        dynamic = "health_tip_calorie_normal"
    import random
    static_idx = random.randint(1, 7)
    return {
        "ok": True,
        "avg7": avg,
        "weightSeries": weight_series,
        "heightSeries": height_series,
        "tips": {"dynamic": dynamic, "static": f"health_tip_static_{static_idx}", "netCal": net_cal},
    }


def _nutrition_export(uid: int, fmt: str, days: int = 30):
    """Ekspor data nutrisi (parity _export_nutrition: CSV/XLSX/DOCX via 0 extra
    dep → openpyxl/python-docx bila tersedia). Mengembalikan marker bytes."""
    try:
        import io
        data = db.get_food_export_data(uid, days=days) or []
    except Exception:
        data = []
    u = db.get_user(uid) or {}
    lang = u.get("language") or "id"
    try:
        from translations import TRANSLATIONS as _TR
        def T(k):
            pair = _TR.get(k) or (k, k)
            return pair[0] if lang == "id" else pair[1]
    except Exception:
        T = lambda k: k  # noqa: E731
    headers = [T("export_date"), T("export_calories"), T("export_protein"), T("export_carbs"),
               T("export_fat"), T("export_water_ml"), T("export_calories_burned"), T("export_net_calories")]
    rows = [[fd.get("date"), fd.get("calories"), fd.get("protein"), fd.get("carbs"),
             fd.get("fat"), fd.get("water_ml"), fd.get("calories_burned"), fd.get("net_calories")]
            for fd in data]
    payload = None
    mime = "text/csv"
    name = "craftlife_nutrition.csv"
    if fmt == "xlsx":
        try:
            import openpyxl
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Nutrition"
            ws.append(headers)
            for r in rows:
                ws.append(r)
            # Kolom lebar dasar
            for col, w in zip("ABCDEFGH", (12, 12, 10, 10, 10, 12, 16, 14)):
                ws.column_dimensions[col].width = w
            bio = io.BytesIO()
            wb.save(bio)
            payload = bio.getvalue()
            mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            name = "craftlife_nutrition.xlsx"
        except Exception:
            payload = None
    elif fmt == "docx":
        try:
            import docx
            document = docx.Document()
            document.add_heading("CraftLife — Nutrition Export", 0)
            table = document.add_table(rows=1, cols=8)
            table.style = "Table Grid"
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = str(h)
            for r in rows:
                cells = table.add_row().cells
                for i, v in enumerate(r):
                    cells[i].text = str(v)
            bio = io.BytesIO()
            document.save(bio)
            payload = bio.getvalue()
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            name = "craftlife_nutrition.docx"
        except Exception:
            payload = None
    elif fmt == "pdf":
        # Parity HealthFoodPage._export_food_pdf: laporan satu halaman berisi
        # ringkasan 30 hari + tabel per-tanggal (reportlab, guarded import).
        try:
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

            bio = io.BytesIO()
            doc = SimpleDocTemplate(bio, pagesize=landscape(A4))
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(name="Title", parent=styles["Title"], alignment=1, fontSize=16)
            cap_style = ParagraphStyle(name="Cap", parent=styles["Normal"], fontSize=11)
            story = [Paragraph("CraftLife — Nutrition & Water Export", title_style), Spacer(1, 0.15 * inch)]
            total_cal = sum(float(r[1] or 0) for r in rows)
            avg_cal = total_cal / len(rows) if rows else 0
            total_water = sum(float(r[5] or 0) for r in rows)
            avg_water = total_water / len(rows) if rows else 0
            summary = Table(
                [[T("export_summary_total_cal"), f"{total_cal:.0f} kcal"],
                 [T("export_summary_avg_cal"), f"{avg_cal:.0f} kcal"],
                 [T("export_summary_total_water"), f"{total_water:.0f} ml"],
                 [T("export_summary_avg_water"), f"{avg_water:.0f} ml"]],
                colWidths=[2.8 * inch, 2.0 * inch],
            )
            summary.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 11),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (0, -1), colors.whitesmoke),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(Paragraph(T("export_summary_label"), cap_style))
            story.append(summary)
            story.append(Spacer(1, 0.2 * inch))
            if rows:
                table = Table([[Paragraph(f"<b>{h}</b>", cap_style) for h in headers]] +
                              [[str(c) for c in r] for r in rows], repeatRows=1)
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                story.append(table)
            else:
                story.append(Paragraph(T("export_no_data"), cap_style))
            doc.build(story)
            payload = bio.getvalue()
            mime = "application/pdf"
            name = "craftlife_nutrition.pdf"
        except Exception:
            payload = None
    if payload is None:
        # CSV fallback (selalu tersedia, parity _export_food_csv)
        import csv
        bio = io.StringIO(newline="")
        writer = csv.writer(bio)
        writer.writerow(headers)
        writer.writerows(rows)
        payload = bio.getvalue().encode("utf-8")
        mime = "text/csv"
        name = "craftlife_nutrition.csv"
    return {"__file_bytes__": payload, "mime": mime, "name": name}


def handle_get(path: str, uid: int, qs=None):
    qs = qs or {}
    # ── P7: Reminder parity ──
    if path == "/api/reminders/due":
        # Parity MainWindow._check_reminders (QTimer 5 detik): pending =
        # waktu <= sekarang, aktif, belum terpicu. Front-end poling ini,
        # bunyikan alarm loop, lalu POST /api/reminders/{id}/trigger.
        return {"ok": True,
                "due": [map_reminder(r) for r in db.get_pending_reminders(uid)]}
    # ── P4: HealthFoodPage parity (per-tanggal) ──
    if path == "/api/healthfood/day":
        day = _parse_day(qs)
        return _healthfood_day(uid, day)
    if path == "/api/healthfood/history":
        return _healthfood_history(uid)
    if path == "/api/food/logs/date":
        # Parity HealthFoodPage: food logs per-tanggal (bukan snapshot hari ini).
        day = _parse_day(qs)
        try:
            u = db.get_user(uid) or {}
            lang = u.get("language") or "id"
            logs = []
            for l in (db.get_food_logs(uid, day) or []):
                name = l.get("name") or ""
                try:
                    name = db.get_food_name(name, lang) or name
                except Exception:
                    pass
                logs.append({
                    "id": str(l.get("id")), "name": name, "icon": l.get("icon") or "🍽️",
                    "serving": float(l.get("serving") or 1),
                    "calories": int(l.get("calories") or 0), "protein": float(l.get("protein") or 0),
                    "carbs": float(l.get("carbs") or 0), "fat": float(l.get("fat") or 0),
                    "mealType": l.get("meal_type") or "snack",
                    "notes": l.get("notes") or "",
                })
        except Exception as e:
            return {"ok": False, "error": str(e)}
        return {"ok": True, "date": day, "logs": logs}
    if path == "/api/water/logs":
        day = _parse_day(qs)
        try:
            logs = [
                {
                    "id": str(w.get("id")),
                    "amountMl": int(w.get("amount_ml") or 0),
                    "time": str(w.get("created_at") or "")[11:16],
                }
                for w in (db.get_water_logs(uid, day) or [])
            ]
        except Exception:
            logs = []
        return {
            "ok": True,
            "date": day,
            "totalMl": int(db.get_water_total(uid, day) or 0),
            "goalMl": int(db.get_water_goal(uid) or 2500),
            "logs": logs,
        }
    if path == "/api/nutrition/export":
        fmt = ((qs.get("format") or ["csv"])[0] or "csv").lower()
        try:
            days = int((qs.get("days") or [30])[0] or 30)
        except (TypeError, ValueError):
            days = 30
        return _nutrition_export(uid, fmt, days)
    if path == "/api/sport":
        return {"ok": True, "sportLogs": snapshot(uid)["sportLogs"]}
    if path == "/api/sport/reps":
        try:
            acts = db.get_sport_activities(uid) or []
        except Exception:
            acts = []
        activities = []
        for a in acts:
            aid = a.get("id")
            total = 0
            try:
                total = int(db.get_sport_rep_total(uid, aid))
            except Exception:
                pass
            rank = {"key": "unranked", "icon": "⭐", "index": 0, "name": ""}
            try:
                rank = db.get_rep_rank(total)
            except Exception:
                rank = rank
            rep_name = a.get("name") or ""
            activities.append({
                "id": str(aid),
                "name": rep_name,
                "icon": a.get("icon") or "💪",
                "sportType": a.get("sport_type") or "other",
                "totalReps": total,
                "rank": rank,
            })
        return {"ok": True, "activities": activities}
    if path.startswith("/api/sport/") and path.endswith("/reps"):
        try:
            sid = int(path.split("/")[3])
            total = int(db.get_sport_rep_total(uid, sid))
            rank = db.get_rep_rank(total)
            series = db.get_sport_rep_series(uid, days=7, activity_id=sid)
            history = db.get_sport_rep_history(uid, sid, 5)
        except Exception as e:
            return {"ok": False, "error": str(e)}
        return {"ok": True, "total": total, "rank": rank, "series": series, "history": history}
    if path == "/api/food/logs":
        return {"ok": True, "mealLogs": snapshot(uid)["mealLogs"]}
    if path == "/api/food/items":
        try:
            items = db.get_food_items(uid, include_default=True) or []
        except Exception:
            items = []
        return {
            "ok": True,
            "items": [
                {
                    "id": str(r.get("id")),
                    "name": r.get("name") or "",
                    "icon": r.get("icon") or "🍽️",
                    "calories": int(r.get("calories") or 0),
                    "protein": float(r.get("protein") or 0),
                    "carbs": float(r.get("carbs") or 0),
                    "fat": float(r.get("fat") or 0),
                    "isCustom": bool(r.get("is_custom")),
                }
                for r in items
            ],
        }
    if path == "/api/nutrition/goals":
        try:
            g = db.get_nutrition_goals(uid) or {}
        except Exception:
            g = {}
        return {
            "ok": True,
            "goals": {
                "calories": int(g.get("daily_calories") or 2000),
                "protein": float(g.get("daily_protein") or 50),
                "carbs": float(g.get("daily_carbs") or 250),
                "fat": float(g.get("daily_fat") or 70),
            },
        }
    if path == "/api/water":
        s = snapshot(uid)
        return {"ok": True, "waterLog": s["waterLog"]}
    if path == "/api/economy":
        return {"ok": True, "transactions": snapshot(uid)["transactions"]}
    if path == "/api/debts":
        return {"ok": True, "debts": snapshot(uid)["debts"]}
    if path == "/api/notes":
        s = snapshot(uid)
        return {"ok": True, "notes": s["notes"], "noteFolders": s["noteFolders"]}
    if path == "/api/reminders":
        return {"ok": True, "reminders": snapshot(uid)["reminders"]}
    if path == "/api/health":
        return {"ok": True, "healthLogs": snapshot(uid)["healthLogs"]}
    if path == "/api/health/bmi":
        try:
            bmi = db.get_user_bmi_settings(uid) or {}
        except Exception:
            bmi = {}
        return {"ok": True, "bmi": bmi}
    if path == "/api/pomodoro":
        return {"ok": True, "pomodoroSessions": snapshot(uid)["pomodoroSessions"]}
    if path == "/api/calendar":
        return {"ok": True, "calendarNotes": snapshot(uid)["calendarNotes"]}
    if path == "/api/supplies":
        # Parity SuppliesPage: filter pencarian & filter kategori.
        search = ((qs.get("search") or [""])[0] or "").strip() or None
        category = ((qs.get("category") or [""])[0] or "").strip() or None
        try:
            rows = db.get_supplies_items(uid, category=category, search=search) or []
        except TypeError:
            rows = db.get_supplies_items(uid) or []
        items_out = [map_supply(r) for r in rows]
        try:
            cats = []
            for c in (db.get_supplies_categories(uid) or []):
                try:
                    subset = db.get_supplies_items(uid, category=c) or []
                except Exception:
                    subset = []
                val = sum((float(r.get("stock") or 0) * float(r.get("price") or 0)) for r in subset)
                cats.append({"name": c, "count": len(subset), "value": val})
        except Exception:
            cats = []
        return {"ok": True, "items": items_out, "categories": cats}
    if path == "/api/supplies/categories":
        try:
            cats = db.get_supplies_categories(uid) or []
        except Exception:
            cats = []
        out = []
        for c in cats:
            try:
                rows = db.get_supplies_items(uid, category=c) or []
            except Exception:
                rows = []
            val = sum((float(r.get("stock") or 0) * float(r.get("price") or 0)) for r in rows)
            out.append({"name": c, "count": len(rows), "value": val})
        return {"ok": True, "categories": out}
    if path == "/api/supplies/history":
        try:
            iid = (qs.get("itemId") or [None])[0]
            iid = int(iid) if iid else None
        except (TypeError, ValueError):
            iid = None
        try:
            rows = db.get_supplies_tx(uid, iid, limit=100) or []
        except Exception:
            rows = []
        try:
            items_map = {str(it.get("id")): it.get("name") or "" for it in (db.get_supplies_items(uid) or [])}
        except Exception:
            items_map = {}
        return {"ok": True, "history": [
            {
                "id": str(r.get("id")), "itemId": str(r.get("item_id")),
                "itemName": r.get("item_name") or r.get("name") or items_map.get(str(r.get("item_id")), ""),
                "kind": r.get("kind") or "in", "qty": float(r.get("qty") or 0),
                "note": r.get("note") or "",
                "createdAt": str(r.get("created_at") or ""),
            } for r in rows
        ]}
    if path == "/api/recipes":
        try:
            recipes = []
            for r in (db.get_recipes(uid) or []):
                rid = r.get("id")
                detail = db.get_recipe_details(rid)
                items = detail["items"] if detail else []
                recipes.append({
                    "id": str(rid),
                    "name": r.get("name") or "",
                    "icon": r.get("icon") or "🍲",
                    "servingSize": int(r.get("serving_size") or 1),
                    "notes": r.get("notes") or "",
                    "items": [
                        {"foodId": str(i.get("food_id")), "name": i.get("name") or "", "quantity": float(i.get("quantity") or 0)}
                        for i in items
                    ],
                })
        except Exception:
            recipes = []
        return {"ok": True, "recipes": recipes}
    if path.startswith("/api/templates/"):
        mode = path.split("/api/templates/", 1)[-1] or "habit"
        if mode in ("apply",):
            return None
        # Parity HabitTemplateDialog: get_templates_by_mode(mode, lang) → nama/desc
        # mengikuti bahasa user, bukan selalu "id".
        try:
            lang = (db.get_user(uid) or {}).get("language", "id")
        except Exception:
            lang = "id"
        try:
            items = db.get_templates_by_mode(mode, lang) or []
        except Exception:
            items = []
        return {"ok": True, "templates": items}
    if path == "/api/task-folders":
        return {"ok": True, "taskFolders": snapshot(uid)["taskFolders"]}
    if path == "/api/holidays":
        # Parity CalendarPage._fetch_holidays: 3 tahun (y-1, y, y+1) sekaligus.
        from datetime import date as _d
        try:
            year = int(((qs.get("year") or [""])[0] or "").strip())
        except Exception:
            year = 0
        if not year:
            year = _d.today().year
        items = []
        try:
            import holidays as hol
            for y in (year - 1, year, year + 1):
                data = hol.get_holidays_for_year(y) or {}
                for ds, names in data.items():
                    if isinstance(names, (list, tuple)) and len(names) >= 2:
                        name_id, name_en = names[0], names[1]
                    else:
                        name_id = name_en = str(names)
                    items.append({"date": ds, "nameId": name_id, "nameEn": name_en, "type": "national"})
        except Exception:
            pass
        return {"ok": True, "year": year, "holidays": items}
    return None


def _last_sport_id(uid):
    rows = db.get_sport_activities(uid) or []
    if not rows:
        return None
    return rows[-1].get("id") if isinstance(rows[-1], dict) else rows[0].get("id")


def handle_post(path: str, uid: int, body: dict, parts: list):
    if path == "/api/sport":
        raw_diff = (body.get("difficulty") or "").lower()
        if raw_diff not in ("easy", "medium", "hard", "epic"):
            intensity = (body.get("intensity") or "moderate").lower()
            raw_diff = {"light": "easy", "moderate": "medium", "vigorous": "hard"}.get(intensity, "medium")
        diff = raw_diff
        db.add_sport_activity(
            uid,
            body.get("sportName") or body.get("name") or "Workout",
            sport_type=body.get("sportType") or "running",
            icon=body.get("icon") or "🏃",
            difficulty=diff,
            notes=body.get("notes") or "",
            calories_burned=int(body.get("caloriesBurned") or 0),
            duration_minutes=int(body.get("durationMinutes") or 30),
        )
        result = {"ok": True}
        if body.get("complete", False):
            sid = _last_sport_id(uid)
            if sid is not None:
                try:
                    result = db.complete_sport_activity(uid, sid) or result
                except Exception as e:
                    result = {"ok": False, "msg": str(e)}
        return {"result": result}

    if len(parts) >= 4 and parts[1] == "sport":
        sid = int(parts[2])
        if parts[3] == "complete":
            return {"result": db.complete_sport_activity(uid, sid)}
        if parts[3] == "delete":
            db.delete_sport_activity(uid, sid)
            return {"result": {"ok": True}}
        if parts[3] == "duplicate":
            return {"result": db.duplicate_sport_activity(uid, sid)}
        if parts[3] == "update":
            kw = {}
            if body.get("name"):
                kw["name"] = str(body.get("name")).strip()
            if body.get("sportType"):
                sport_type = str(body.get("sportType"))
                kw["sport_type"] = sport_type if sport_type in db.SPORT_TYPES else "other"
                kw["icon"] = db.SPORT_TYPES.get(sport_type if sport_type in db.SPORT_TYPES else "other", {}).get("icon", "🏅")
            if body.get("difficulty"):
                kw["difficulty"] = str(body.get("difficulty"))
            if body.get("notes") is not None:
                kw["notes"] = str(body.get("notes") or "")
            for fk, dk in (("caloriesBurned", "calories_burned"), ("durationMinutes", "duration_minutes")):
                if body.get(fk) is not None:
                    try:
                        kw[dk] = int(body.get(fk) or 0)
                    except (TypeError, ValueError):
                        pass
            if "folderId" in body:
                fid = body.get("folderId")
                kw["folder_id"] = int(fid) if fid not in (None, "", "null") else None
            res = db.update_sport_activity(sid, uid, **kw)
            if isinstance(res, dict) and res.get("ok") is False:
                return {"result": res}
            return {"result": {"ok": True}}
        if parts[3] == "reps":
            return {"result": db.add_sport_rep_log(
                uid, sid, int(body.get("reps") or 0), int(body.get("sets") or 1),
                note=body.get("note") or "")}

    if path == "/api/sport/template":
        key = body.get("key") or "running_starter_s"
        n = db.apply_template_by_mode(uid, "sport", key)
        return {"result": {"ok": True, "n": n}}

    if path == "/api/food/custom":
        result = db.add_custom_food(
            uid,
            (body.get("foodName") or body.get("name") or "Food").strip(),
            body.get("icon") or "🍽️",
            int(body.get("calories") or 0),
            float(body.get("protein") or 0),
            float(body.get("carbs") or 0),
            float(body.get("fat") or 0),
        )
        return {"result": result}

    if path == "/api/food/log":
        fid = body.get("foodId") or body.get("food_id")
        if fid not in (None, "", "null"):
            try:
                fid = int(fid)
            except Exception:
                fid = None
        else:
            fid = None
        name = (body.get("foodName") or body.get("name") or "Food").strip()
        if not fid:
            try:
                fid = db.get_default_food_id_by_name(name)
            except Exception:
                fid = None
        if not fid:
            created = db.add_custom_food(
                uid, name, body.get("icon") or "🍽️",
                int(body.get("calories") or 0),
                float(body.get("protein") or 0),
                float(body.get("carbs") or 0),
                float(body.get("fat") or 0),
            )
            fid = created.get("id") if isinstance(created, dict) else created
        result = db.log_food(
            uid, fid,
            float(body.get("portion") or 1),
            body.get("mealType") or "snack",
            body.get("date") or _today(),
            body.get("notes") or "",
        )
        return {"result": result}

    if len(parts) >= 5 and parts[1] == "food" and parts[2] == "log" and parts[4] == "delete":
        db.delete_food_log(uid, int(parts[3]))
        return {"result": {"ok": True}}

    if path == "/api/nutrition/goals":
        result = db.update_nutrition_goals(
            uid,
            float(body.get("calories") or 2000),
            float(body.get("protein") or 50),
            float(body.get("carbs") or 250),
            float(body.get("fat") or 70),
        )
        return {"result": result}

    if path == "/api/health/goals":
        result = db.update_health_goals(
            uid,
            float(body.get("dailySteps") or body.get("steps") or 10000),
            float(body.get("dailySleepHours") or body.get("sleepHours") or 7.0),
            height_cm=body.get("heightCm"),
            weight_kg=body.get("weightKg"),
        )
        return {"result": result}

    if path == "/api/water":
        amount = int(body.get("amountMl") or 0)
        log_date = str(body.get("logDate") or body.get("date") or _today())
        goal = int(db.get_water_goal(uid) or 2500)
        before = int(db.get_water_total(uid, log_date) or 0)
        result = db.add_water_log(uid, amount, log_date)
        # Parity HealthFoodPage._add_water: saat total air melintasi target pada
        # entri ini → bonus 10 XP + 2 gold + notifikasi.
        try:
            after = int(db.get_water_total(uid, log_date) or 0)
            if after >= goal > before:
                db.gain_xp_gold(uid, 10, 2)
                try:
                    import database as _d
                    u = _d.get_user(uid) or {}
                    lang = u.get("language") or "id"
                    from translations import TRANSLATIONS as _TR
                    msg = (_TR.get("water_goal_reached") or ("🎉 Target air harian tercapai! +10 XP +2 Gold", "🎉 Daily water goal reached! +10 XP +2 Gold"))[0 if lang == "id" else 1]
                    db.add_notification(uid, msg, "success")
                except Exception:
                    pass
        except Exception:
            pass
        return {"result": result}
    if path == "/api/water/log/delete":
        result = db.delete_water_log(uid, int(body.get("id") or 0))
        return {"result": result if isinstance(result, dict) else {"ok": True}}
    if path == "/api/food/log/move":
        result = db.move_item_to_folder(uid, "food", int(body.get("id") or 0),
                                        body.get("mealType") or "snack")
        return {"result": result if isinstance(result, dict) else {"ok": True}}
    if path == "/api/nutrition/bonus":
        result = db.check_daily_nutrition_bonus(uid, str(body.get("logDate") or body.get("date") or _today()))
        return {"result": result if isinstance(result, dict) else {"ok": False}}
    if path == "/api/nutrition/goals/auto":
        # Parity _set_auto_goals: Mifflin-St Jeor + faktor aktivitas, lalu
        # set goal nutrisi + log berat badan.
        w = float(body.get("weightKg") or 70)
        h = float(body.get("heightCm") or 170)
        age = int(body.get("age") or 25)
        gender = str(body.get("gender") or "male").lower()
        factors = [1.2, 1.375, 1.55, 1.725, 1.9]
        try:
            factor = float(body.get("activityFactor"))
            if factor not in factors:
                factor = factors[min(4, max(0, int(body.get("activityIndex") or 1)))]
        except (TypeError, ValueError):
            factor = factors[min(4, max(0, int(body.get("activityIndex") or 1)))]
        bmr = 10 * w + 6.25 * h - 5 * age + (5 if gender == "male" else -161)
        cal = int(bmr * factor)
        protein = int(w * 1.6)
        carbs = int(cal * 0.5 / 4)
        fat = int(cal * 0.3 / 9)
        res1 = db.update_nutrition_goals(uid, cal, protein, carbs, fat)
        db.update_user_bmi_settings(uid, h, w, age, gender, factor)
        db.log_user_weight(uid, w)
        return {"result": {"ok": True, "calories": cal, "protein": protein,
                           "carbs": carbs, "fat": fat, "goals": res1}}

    if path == "/api/water/reset":
        result = db.reset_water_today(uid)
        return {"result": result or {"ok": True}}

    if path == "/api/water/goal":
        db.set_water_goal(uid, int(body.get("targetMl") or body.get("daily_ml") or 2000))
        return {"result": {"ok": True}}

    if path == "/api/task-folders":
        mode = (body.get("mode") or "habit").strip() or "habit"
        result = db.add_task_folder(uid, mode, body.get("name") or "Folder", body.get("icon") or "📁")
        return {"result": result}

    if len(parts) >= 4 and parts[1] == "task-folders" and parts[3] == "delete":
        mode = (body.get("mode") or "habit").strip() or "habit"
        db.delete_task_folder(uid, int(parts[2]), mode)
        return {"result": {"ok": True}}

    if len(parts) >= 4 and parts[1] == "task-folders" and parts[3] == "update":
        mode = (body.get("mode") or "habit").strip() or "habit"
        kw = {}
        if body.get("name") not in (None, ""):
            kw["name"] = body.get("name")
        if body.get("icon") not in (None, ""):
            kw["icon"] = body.get("icon")
        db.update_task_folder(int(parts[2]), uid, **kw) if kw else None
        return {"result": {"ok": True}}

    if len(parts) >= 4 and parts[1] == "task-folders" and parts[3] == "duplicate":
        mode = (body.get("mode") or "habit").strip() or "habit"
        return {"result": db.duplicate_task_folder(uid, int(parts[2]), mode)}

    if path == "/api/templates/apply":
        n = db.apply_template_by_mode(uid, body.get("mode") or "habit", body.get("key") or "")
        return {"result": {"ok": True, "count": n}}

    if path == "/api/savings":
        result = db.add_saving(
            uid,
            body.get("name") or "Saving",
            body.get("icon") or "🏦",
            _to_idr(uid, body.get("targetAmount")),
            _to_idr(uid, body.get("currentAmount")),
            body.get("targetDate"),
            body.get("notes") or "",
        )
        return {"result": result}
    if len(parts) >= 4 and parts[1] == "savings":
        sid = int(parts[2])
        if parts[3] == "add":
            return {"result": db.add_to_saving(sid, uid, _to_idr(uid, body.get("amount")))}
        if parts[3] == "withdraw":
            return {"result": db.withdraw_from_saving(sid, uid, _to_idr(uid, body.get("amount")))}
        if parts[3] == "delete":
            db.delete_saving(sid, uid)
            return {"result": {"ok": True}}

    if path == "/api/investments":
        result = db.add_investment(
            uid,
            body.get("name") or "Investment",
            body.get("icon") or "📈",
            _to_idr(uid, body.get("amount")),
            body.get("notes") or "",
        )
        return {"result": result}
    if len(parts) >= 4 and parts[1] == "investments":
        iid = int(parts[2])
        if parts[3] == "return":
            return {"result": db.collect_investment_return(iid, uid, float(body.get("percent") or 5))}
        if parts[3] == "withdraw":
            return {"result": db.withdraw_investment(iid, uid)}

    if path == "/api/subscriptions":
        result = db.add_subscription(
            uid,
            body.get("name") or "Subscription",
            body.get("icon") or "📅",
            _to_idr(uid, body.get("amount")),
            body.get("dueDate") or _today(),
            body.get("period") or "monthly",
            bool(body.get("isRecurring", True)),
            body.get("notes") or "",
        )
        return {"result": result}
    if len(parts) >= 4 and parts[1] == "subscriptions":
        sid = int(parts[2])
        if parts[3] == "delete":
            db.delete_subscription(sid, uid)
            return {"result": {"ok": True}}
        if parts[3] == "renew":
            return {"result": db.renew_subscription(sid, uid, bool(body.get("autoPay", True)))}

    if path == "/api/debt-notes":
        result = db.add_debt_note(
            uid,
            body.get("personName") or body.get("name") or "",
            _to_idr(uid, body.get("amount")),
            body.get("date") or _today(),
            body.get("notes") or "",
        )
        return {"result": result}
    if len(parts) >= 4 and parts[1] == "debt-notes":
        nid = int(parts[2])
        if parts[3] == "settle":
            return {"result": db.settle_debt_note(uid, nid)}
        if parts[3] == "delete":
            return {"result": db.delete_debt_note(uid, nid)}

    if path == "/api/economy":
        typ = body.get("type") or "expense"
        fid = body.get("folderId")
        result = db.add_economy_item(
            uid,
            body.get("name") or body.get("category") or typ,
            "💰",
            typ,
            _to_idr(uid, body.get("amount")),
            body.get("category") or "",
            body.get("date") or _today(),
            notes=body.get("notes") or "",
            folder_id=int(fid) if fid not in (None, "", "null") else None,
        )
        return {"result": result}

    if len(parts) >= 4 and parts[1] == "economy" and parts[3] == "update":
        kw = {}
        if body.get("name") is not None:
            kw["name"] = str(body.get("name")).strip() or "Transaksi"
        if body.get("category") is not None:
            kw["category"] = str(body.get("category")).strip() or "other"
        if body.get("type") in ("income", "expense"):
            kw["type"] = str(body.get("type"))
        if body.get("amount") is not None:
            kw["amount"] = _to_idr(uid, body.get("amount"))
        if body.get("date"):
            kw["date"] = str(body.get("date"))[:10]
        if body.get("notes") is not None:
            kw["notes"] = str(body.get("notes"))
        if "folderId" in body:
            fid = body.get("folderId")
            kw["folder_id"] = int(fid) if fid not in (None, "", "null") else None
        icon = body.get("icon")
        if icon:
            kw["icon"] = str(icon)[:16]
        db.update_economy_item(int(parts[2]), uid, **kw)
        return {"result": {"ok": True}}

    if len(parts) >= 4 and parts[1] == "economy" and parts[3] == "move":
        fid = body.get("folderId")
        db.move_item_to_folder(uid, "economy", int(parts[2]), int(fid) if fid not in (None, "", "null") else None)
        return {"result": {"ok": True}}

    if len(parts) >= 4 and parts[1] == "economy" and parts[3] == "delete":
        db.delete_economy_item(uid, int(parts[2]))
        return {"result": {"ok": True}}

    if path == "/api/debts":
        result = db.add_debt(
            uid,
            body.get("title") or body.get("name") or "Debt",
            _to_idr(uid, body.get("totalAmount") or body.get("amount")),
            due_date=body.get("dueDate"),
            notes=body.get("notes") or "",
        )
        return {"result": result}

    if len(parts) >= 4 and parts[1] == "debts":
        did = int(parts[2])
        if parts[3] == "pay":
            return {"result": db.pay_debt_installment(did, uid, _to_idr(uid, body.get("amount")))}
        if parts[3] == "update":
            # Parity EditDebtDialog-ish: update nama/jumlah/jatuh tempo/catatan.
            kw = {}
            if body.get("title") is not None:
                kw["name"] = str(body.get("title"))[:120]
            if body.get("totalAmount") is not None:
                kw["amount"] = float(_to_idr(uid, body.get("totalAmount")))
            if "dueDate" in body:
                kw["due_date"] = str(body.get("dueDate") or "")[:10] or None
            if "notes" in body:
                kw["notes"] = str(body.get("notes") or "")
            if kw:
                db.update_debt(did, uid, **kw)
            return {"result": {"ok": True}}
        if parts[3] == "delete":
            db.delete_debt(did, uid)
            return {"result": {"ok": True}}

    if path == "/api/notes":
        folder = body.get("folderId")
        folder_id = int(folder) if folder not in (None, "", "null") else None
        result = db.add_note(uid, folder_id, body.get("title") or "Note", body.get("content") or "")
        return {"result": result}

    if path == "/api/notes/preview-math":
        try:
            from mathtools import latex_to_unicode
            raw = body.get("content") or ""
            return {"result": {"ok": True, "preview": latex_to_unicode(raw)}, "skip_snap": True}
        except Exception as e:
            return {"result": {"ok": False, "msg": str(e)}, "skip_snap": True}

    if path == "/api/notes/math-chunks":
        # Parity NotesPage._latex_preview: daftar chunk LaTeX + hasil konversi
        # unicode masing-masing untuk dialog MathPreview.
        try:
            import mathtools
            raw = body.get("content") or ""
            chunks = mathtools.find_math_chunks(raw)
            out = []
            for c in chunks:
                seg = c.get("text", c) if isinstance(c, dict) else c
                out.append({"raw": seg, "converted": mathtools.latex_to_unicode(seg)})
            return {"result": {"ok": True, "chunks": out}, "skip_snap": True}
        except Exception as e:
            return {"result": {"ok": False, "msg": str(e)}, "skip_snap": True}

    if path == "/api/notes/reorder":
        items = body.get("items") or []
        result = db.reorder_notes(uid, items)
        return {"result": result}

    if path == "/api/note-folders":
        result = db.add_note_folder(uid, body.get("name") or "Folder", body.get("icon") or "📁", body.get("parentId"))
        return {"result": result}
    if len(parts) >= 4 and parts[1] == "note-folders" and parts[3] == "delete":
        db.delete_note_folder(int(parts[2]), uid)
        return {"result": {"ok": True}}
    if len(parts) >= 4 and parts[1] == "note-folders" and parts[3] == "update":
        # Parity NotesPage._edit_folder_name / _edit_folder_icon.
        fid = int(parts[2])
        if body.get("icon") is not None and body.get("name") is None:
            db.update_note_folder_icon(fid, uid, body.get("icon") or "📁")
        else:
            db.update_note_folder(fid, uid, name=body.get("name"), icon=body.get("icon"))
        return {"result": {"ok": True}}
    if len(parts) >= 4 and parts[1] == "note-folders" and parts[3] == "duplicate":
        # Parity NotesPage._duplicate_folder (deep-copy folder + isinya).
        return {"result": db.duplicate_note_folder(uid, int(parts[2]))}

    if len(parts) >= 4 and parts[1] == "notes":
        nid = int(parts[2])
        if parts[3] == "update":
            result = db.update_note(
                nid, uid,
                title=body.get("title"),
                content=body.get("content"),
                folder_id=int(body["folderId"]) if body.get("folderId") else None,
                zoom_level=body.get("zoomLevel") or body.get("zoom_level"),
            )
            return {"result": result}
        if parts[3] == "delete":
            db.delete_note(nid, uid)
            return {"result": {"ok": True}}
        if parts[3] == "archive":
            return {"result": db.archive_note(nid, uid, bool(body.get("archived", True)))}
        if parts[3] == "duplicate":
            dest = body.get("folderId")
            dest_id = int(dest) if dest not in (None, "", "null") else None
            return {"result": db.duplicate_note(uid, nid, dest_id)}

    if path == "/api/reminders":
        # Parity ReminderDialog._save: {title, description, reminder_datetime
        # yyyy-MM-dd HH:mm:ss, sound_type, sound_file, repeat_type, repeat_days}.
        # Kompat: {time:"HH:mm"} lama → hari ini jam tsb.
        dt = (body.get("reminderDatetime") or body.get("datetime") or "").strip()
        if not dt:
            t = body.get("time") or "08:00"
            dt = f"{_today()} {t}:00"
        if len(dt) == 16:
            dt = dt + ":00"
        st = body.get("soundType") or body.get("sound") or "default"
        result = db.add_reminder(
            uid, body.get("title") or "Reminder", body.get("description") or "",
            dt, sound_type=st,
            sound_file=(body.get("soundFile") or None) if st == "custom" else None,
            repeat_type=body.get("repeat") or "none",
            repeat_days=body.get("repeatDays") or "",
        )
        return {"result": result}

    if len(parts) >= 4 and parts[1] == "reminders":
        rid = int(parts[2])
        if parts[3] == "delete":
            db.delete_reminder(rid, uid)
            return {"result": {"ok": True}}
        if parts[3] == "trigger":
            # Parity MainWindow._check_reminders: setelah alarm dibunyikan —
            # non-repeat → mark triggered; repeat → jadwalkan ke next occurrence
            # (get_next_reminder_datetime); jika tak ada next → nonaktif+triggered.
            rem = db.get_reminder(rid, uid)
            if rem:
                rep = rem.get("repeat_type", "none")
                if rep and rep != "none":
                    next_dt = db.get_next_reminder_datetime(
                        rem["reminder_datetime"], rep, rem.get("repeat_days", "")
                    )
                    if next_dt:
                        db.update_reminder(rid, uid, reminder_datetime=next_dt,
                                           triggered=0, is_active=1)
                    else:
                        db.update_reminder(rid, uid, is_active=0, triggered=1)
                else:
                    db.mark_reminder_triggered(rid, uid)
            return {"result": {"ok": True}}
        if parts[3] == "update":
            # Parity ReminderDialog._save mode edit (triggered selalu reset 0).
            dt = (body.get("reminderDatetime") or "").strip()
            if len(dt) == 16:
                dt = dt + ":00"
            st = body.get("soundType") or "default"
            kwargs = {
                "title": body.get("title"), "description": body.get("description"),
                "sound_type": st,
                "sound_file": (body.get("soundFile") or None) if st == "custom" else None,
                "triggered": 0,
                "repeat_type": body.get("repeat") or "none",
                "repeat_days": body.get("repeatDays") or "",
            }
            if dt:
                kwargs["reminder_datetime"] = dt
            db.update_reminder(rid, uid, **{k: v for k, v in kwargs.items() if v is not None})
            return {"result": {"ok": True}}
        if parts[3] == "toggle":
            # Parity _toggle_selected: flip is_active; saat aktifkan kembali →
            # reset triggered agar berbunyi lagi.
            rem = db.get_reminder(rid, uid)
            if rem:
                new_state = 0 if rem.get("is_active") else 1
                db.update_reminder(rid, uid, is_active=new_state)
                if new_state:
                    db.reset_reminder_triggered(rid, uid)
            return {"result": {"ok": True}}

    if path == "/api/calendar/note":
        # Parity CalendarPage._save_note: catatan kosong = hapus.
        note = (body.get("note") or "")
        if note.strip():
            result = db.save_calendar_note(uid, body.get("date") or _today(), note.strip())
        else:
            db.delete_calendar_note(uid, body.get("date") or _today())
            result = {"ok": True}
        return {"result": result}

    if path == "/api/calendar/note/delete":
        # Parity CalendarPage._delete_note (tombol Hapus di dialog).
        db.delete_calendar_note(uid, body.get("date") or _today())
        return {"result": {"ok": True}}

    if path == "/api/health/bmi":
        result = db.update_user_bmi_settings(
            uid,
            float(body.get("height_cm") or body.get("heightCm") or 170),
            float(body.get("weight_kg") or body.get("weightKg") or 70),
            int(body.get("age") or 25),
            body.get("gender") or "male",
            float(body.get("activity_factor") or body.get("activityFactor") or 1.55),
        )
        return {"result": result}
    if path == "/api/health":
        log_date = str(body.get("logDate") or body.get("date") or _today())
        steps = int(body.get("steps") or 0)
        sleep = float(body.get("sleepHours") or 0)
        resting_hr = int(body.get("heartRate") or 0)
        weight = body.get("weightKg")
        weight = float(weight) if weight not in (None, "", 0) else None
        height = body.get("heightCm")
        height = float(height) if height not in (None, "", 0) else None
        mood = str(body.get("mood") or "normal")
        stress = str(body.get("stress") or "normal")
        notes = str(body.get("notes") or "")
        # Parity HealthFoodPage._save_health: total air & net kalori dihitung server.
        try:
            water_total = int(db.get_water_total(uid, log_date) or 0)
        except Exception:
            water_total = 0
        try:
            nutrition = db.get_nutrition_summary(uid, log_date) or {}
            burned = int(db.get_total_calories_burned_today(uid, log_date) or 0)
            net_cal = int(nutrition.get("calories") or 0) - burned
        except Exception:
            net_cal = None
        result = db.add_health_log(
            uid, log_date,
            steps=steps, sleep_hours=sleep, water_ml=water_total,
            weight_kg=weight, resting_hr=resting_hr,
            stress_level=stress, mood=mood, notes=notes, net_calories=net_cal,
        )
        # Parity: sinkronkan height/weight ke user_health_goals.
        try:
            goals = db.get_health_goals(uid) or {}
            if height is not None or weight is not None:
                db.update_health_goals(
                    uid,
                    float(goals.get("daily_steps") or 10000),
                    float(goals.get("daily_sleep_hours") or 7.0),
                    height_cm=height,
                    weight_kg=weight,
                )
        except Exception:
            pass
        return {"result": result}

    if path == "/api/pomodoro/complete":
        result = db.complete_pomodoro(uid, int(body.get("durationMinutes") or 25), body.get("label") or "")
        return {"result": result}

    if path == "/api/recipes":
        items = body.get("items") or []
        food_items = []
        for it in items:
            fid = it.get("foodId") or it.get("food_id")
            try:
                food_items.append((int(fid), float(it.get("quantity") or 0)))
            except (TypeError, ValueError):
                continue
        result = db.add_recipe(
            uid,
            (body.get("name") or "Recipe").strip(),
            body.get("icon") or "🍲",
            int(body.get("servingSize") or body.get("serving_size") or 1),
            body.get("notes") or "",
            food_items,
        )
        return {"result": result}
    if len(parts) >= 4 and parts[1] == "recipes":
        rid = int(parts[2])
        if parts[3] == "delete":
            db.delete_recipe(uid, rid)
            return {"result": {"ok": True}}
        if parts[3] == "log":
            return {"result": db.log_recipe(
                uid, rid,
                float(body.get("servingMultiplier") or body.get("serving") or 1),
                body.get("mealType") or "lunch",
                body.get("date") or _today(),
                body.get("notes") or "")}

    if path == "/api/supplies":
        result = db.add_supply_item(
            uid,
            body.get("name") or "Item",
            category=body.get("category") or "",
            unit=body.get("unit") or "pcs",
            stock=float(body.get("stock") or 0),
            min_stock=float(body.get("minStock") or body.get("min_stock") or 0),
            price=_to_idr(uid, body.get("price")),
            location=body.get("location") or "",
            notes=body.get("notes") or "",
            economy_category=body.get("economy_category") or body.get("economyCategory") or "",
        )
        return {"result": result}
    if len(parts) >= 4 and parts[1] == "supplies":
        sid = int(parts[2])
        action = parts[3]
        if action == "delete":
            db.delete_supply_item(uid, sid)
            return {"result": {"ok": True}}
        if action == "tx":
            qty = float(body.get("qty") or 0)
            note = body.get("note") or ""
            if body.get("logEconomy"):
                return {"result": db.record_supply_tx_with_economy(
                    uid, sid, body.get("kind") or "in", qty, note,
                    log_economy=True,
                    economy_amount=_to_idr(uid, body.get("economyAmount")),
                    economy_category=body.get("economyCategory") or "Supplies")}
            return {"result": db.record_supply_tx(uid, sid, body.get("kind") or "in", qty, note)}
        if action == "update":
            kw = {}
            for src, dst in (("name", "name"), ("location", "location"), ("category", "category"), ("unit", "unit"), ("notes", "notes")):
                if src in body:
                    kw[dst] = body.get(src)
            if "minStock" in body or "min_stock" in body:
                kw["min_stock"] = float(body.get("minStock") if "minStock" in body else body.get("min_stock") or 0)
            if "price" in body:
                kw["price"] = _to_idr(uid, body.get("price"))
            return {"result": db.update_supply_item(uid, sid, **kw)}

    return None
